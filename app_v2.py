

import os
import json
import re
from pathlib import Path
from datetime import datetime, date
from html import escape
from flask import (
    Flask,
    render_template,
    request,
    jsonify,
    redirect,
    url_for,
    flash,
    send_file,
    session,
    send_from_directory,
    abort,
)
from werkzeug.utils import secure_filename
import smtplib
import ssl
import requests
from typing import Dict, List

from dotenv import load_dotenv
app = Flask(__name__)

# ---- Optional RAG/Embeddings imports (best-effort, gracefully degrade) ----
try:
    import faiss  # type: ignore
    _FAISS_AVAILABLE = True
except Exception:
    faiss = None
    _FAISS_AVAILABLE = False
try:
    from sentence_transformers import SentenceTransformer  # type: ignore
    _ST_AVAILABLE = True
except Exception:
    SentenceTransformer = None  # type: ignore
    _ST_AVAILABLE = False
import hashlib
import json as _json_mod
from pathlib import Path
from threading import Lock
_CACHE_DIR = Path(os.getenv('CACHE_DIR') or (Path(__file__).parent / 'cache'))
_CACHE_DIR.mkdir(exist_ok=True)
_LLM_CACHE_DIR = _CACHE_DIR / 'llm'
_LLM_CACHE_DIR.mkdir(parents=True, exist_ok=True)
_FAISS_DIR = _CACHE_DIR / 'faiss'
_FAISS_DIR.mkdir(parents=True, exist_ok=True)
_cache_lock = Lock()
import time as _time
import torch
import numpy as np

# --- Simple circuit breaker for local Ollama to avoid repeated long timeouts ---
_OLLAMA_CB = {'opened_until': 0.0, 'failures': 0}

def _ollama_circuit_opened() -> bool:
    try:
        return _time.time() < float(_OLLAMA_CB.get('opened_until', 0.0) or 0.0)
    except Exception:
        return False

def _ollama_record_failure():
    try:
        failures = int(_OLLAMA_CB.get('failures', 0) or 0) + 1
        _OLLAMA_CB['failures'] = failures
        # After 2 consecutive failures, open the circuit for 5 minutes
        if failures >= 2:
            _OLLAMA_CB['opened_until'] = _time.time() + 300.0
            _OLLAMA_CB['failures'] = 0
    except Exception:
        pass

def _ollama_record_success():
    try:
        _OLLAMA_CB['failures'] = 0
        _OLLAMA_CB['opened_until'] = 0.0
    except Exception:
        pass

# --- Persistent HTTP session for Ollama with connection pooling ---
_OLLAMA_SESSION = None

def _get_ollama_session():
    global _OLLAMA_SESSION
    if _OLLAMA_SESSION is None:
        from requests.adapters import HTTPAdapter
        from urllib3.util.retry import Retry
        _OLLAMA_SESSION = requests.Session()
        # Configure connection pooling
        adapter = HTTPAdapter(
            pool_connections=10,
            pool_maxsize=20,
            max_retries=Retry(total=0),  # We handle retries ourselves
            pool_block=False
        )
        _OLLAMA_SESSION.mount('http://', adapter)
        _OLLAMA_SESSION.mount('https://', adapter)
    return _OLLAMA_SESSION

def _cache_key_from_dict(payload: dict) -> str:
    try:
        blob = _json_mod.dumps(payload, sort_keys=True, ensure_ascii=False).encode('utf-8')
    except Exception:
        blob = str(payload).encode('utf-8', 'ignore')
    return hashlib.sha256(blob).hexdigest()

def _llm_cache_get(key: str):
    try:
        p = _LLM_CACHE_DIR / f"{key}.json"
        if p.is_file():
            with p.open('r', encoding='utf-8') as f:
                return _json_mod.load(f)
    except Exception:
        return None
    return None

def _llm_cache_set(key: str, value: dict):
    try:
        p = _LLM_CACHE_DIR / f"{key}.json"
        with _cache_lock:
            with p.open('w', encoding='utf-8') as f:
                _json_mod.dump(value, f, ensure_ascii=False)
    except Exception:
        pass

def _text_chunks(text: str, chunk_size: int = 800, overlap: int = 120) -> list[str]:
    text = (text or '').strip()
    if not text:
        return []
    chunks = []
    start = 0
    n = len(text)
    while start < n:
        end = min(n, start + chunk_size)
        chunk = text[start:end]
        chunks.append(chunk)
        if end == n:
            break
        start = max(end - overlap, start + 1)
    return chunks

def _ollama_health_check():
    """Check Ollama health and resource usage for load balancing hints."""
    try:
        ollama_base = (app.config.get('OLLAMA_BASE_URL') or os.getenv('OLLAMA_BASE_URL') or 'http://127.0.0.1:11434').strip()

        # Check running models
        running_resp = _get_ollama_session().get(f"{ollama_base.rstrip('/')}/api/ps", timeout=5)
        if running_resp.status_code == 200:
            running_data = running_resp.json()
            running_models = running_data.get('models', [])
            if len(running_models) > 2:
                app.logger.warning(f"Ollama has {len(running_models)} models running - consider reducing concurrent load")

        # Check available models
        tags_resp = _get_ollama_session().get(f"{ollama_base.rstrip('/')}/api/tags", timeout=5)
        if tags_resp.status_code == 200:
            tags_data = tags_resp.json()
            models = tags_data.get('models', [])
            # Prefer smaller models if many are loaded
            if len(models) > 5:
                app.logger.info(f"Ollama has {len(models)} models available - consider using smaller models for better performance")

        return True
    except Exception as e:
        app.logger.warning(f"Ollama health check failed: {e}")
        return False

_emb_model_singleton = None
def _get_embedding_model():
    global _emb_model_singleton
    if not _ST_AVAILABLE:
        return None
    if _emb_model_singleton is None:
        try:
            # Small, fast default
            _emb_model_singleton = SentenceTransformer(os.getenv('EMBEDDING_MODEL') or 'sentence-transformers/all-MiniLM-L6-v2')
        except Exception:
            _emb_model_singleton = None
    return _emb_model_singleton

def _build_or_load_faiss(company: str, bid_id, docs: list[str]):
    if not (_FAISS_AVAILABLE and _ST_AVAILABLE):
        return None, docs
    try:
        key = f"{(company or '').strip().lower()}__{str(bid_id or '')}"
        index_path = _FAISS_DIR / f"{hashlib.sha256(key.encode()).hexdigest()}.index"
        meta_path = _FAISS_DIR / f"{hashlib.sha256((key+'__meta').encode()).hexdigest()}.json"
        model = _get_embedding_model()
        if model is None:
            return None, docs
        if index_path.is_file() and meta_path.is_file():
            index = faiss.read_index(str(index_path))  # type: ignore
            with meta_path.open('r', encoding='utf-8') as f:
                meta = _json_mod.load(f)
            stored_docs = meta.get('docs') or []
            if isinstance(stored_docs, list) and len(stored_docs) > 0:
                return index, stored_docs
        # Build fresh
        vectors = model.encode(docs, normalize_embeddings=True)
        index = faiss.IndexFlatIP(vectors.shape[1])  # type: ignore
        index.add(vectors)  # type: ignore
        faiss.write_index(index, str(index_path))  # type: ignore
        with meta_path.open('w', encoding='utf-8') as f:
            _json_mod.dump({'docs': docs}, f, ensure_ascii=False)
        return index, docs
    except Exception:
        return None, docs

def _retrieve_top_k(query: str, company: str, bid_id, docs: list[str], index=None, k: int = 8) -> list[str]:
    docs = docs or []
    if not docs:
        return []
    if not query:
        return docs[:k]
    if index is not None and _ST_AVAILABLE:
        try:
            model = _get_embedding_model()
            if model is not None:
                qv = model.encode([query], normalize_embeddings=True)
                D, I = index.search(qv, min(k, len(docs)))  # type: ignore
                top = []
                seen = set()
                for idx in (I[0] if len(I) else []):
                    if int(idx) not in seen and 0 <= int(idx) < len(docs):
                        top.append(docs[int(idx)])
                        seen.add(int(idx))
                return top
        except Exception:
            pass
    # Fallback: simple keyword scoring
    q = query.lower()
    scored = []
    for d in docs:
        score = 0
        dl = d.lower()
        for token in set(q.split()):
            if token and token in dl:
                score += 1
        scored.append((score, d))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [d for _, d in scored[:k]]

def _ollama_chat(messages: list[dict], temperature: float = 0.2, timeout_sec: int = 60):
    max_retries = 3
    for attempt in range(max_retries):
        try:
            ollama_base = (app.config.get('OLLAMA_BASE_URL') or os.getenv('OLLAMA_BASE_URL') or 'http://127.0.0.1:11434').strip()
            ollama_model = (app.config.get('OLLAMA_MODEL') or os.getenv('OLLAMA_MODEL') or 'llama3').strip()
            endpoint_ollama = ollama_base.rstrip('/') + '/api/chat'
            headers_ollama = {'Content-Type': 'application/json'}
            body_ollama = {'model': ollama_model, 'messages': messages, 'stream': False, 'options': {'temperature': temperature}}
            resp = _get_ollama_session().post(endpoint_ollama, headers=headers_ollama, json=body_ollama, timeout=timeout_sec)
            data_ollama = resp.json()
            if resp.status_code >= 400:
                if attempt < max_retries - 1:
                    # Exponential backoff: wait 1s, 2s, 4s...
                    _time.sleep(2 ** attempt)
                    continue
                return ''
            if isinstance(data_ollama, dict):
                if isinstance(data_ollama.get('message'), dict):
                    return data_ollama['message'].get('content', '') or ''
                if isinstance(data_ollama.get('messages'), list) and data_ollama['messages']:
                    return (data_ollama['messages'][-1] or {}).get('content', '') or ''
                if 'response' in data_ollama:
                    return data_ollama.get('response', '') or ''
            return ''
        except (requests.Timeout, requests.ConnectionError, requests.RequestException) as e:
            if attempt < max_retries - 1:
                # Exponential backoff for network errors
                _time.sleep(2 ** attempt)
                continue
            # Log the final failure
            app.logger.warning(f"Ollama request failed after {max_retries} attempts: {e}")
            return ''
        except Exception as e:
            # For non-network errors, don't retry
            app.logger.error(f"Unexpected error in Ollama request: {e}")
            return ''


def _load_dotenv_hierarchy():
    """Load environment variables from common .env locations.

    We look in the following order (later files override earlier values):
      1. Project root two levels up (handles running via different entrypoints)
      2. Directory that contains this file
      3. An `env/.env` helper directory beside the project (if present)
    """
    current_file = Path(__file__).resolve()
    candidates = [
        current_file.parent.parent / '.env',
        current_file.parent / '.env',
        current_file.parent / 'env' / '.env',
        Path.cwd() / '.env',
    ]

    loaded_any = False
    for candidate in candidates:
        if candidate.is_file():
            load_dotenv(candidate, override=True)
            loaded_any = True
    if not loaded_any:
        load_dotenv()


_load_dotenv_hierarchy()

def _apply_provider_defaults():
    provider = (os.getenv('LLM_PROVIDER') or '').strip().lower()
    if provider == 'openrouter':
        base = os.getenv('OPENAI_BASE_URL') or os.getenv('OPENROUTER_BASE_URL') or 'https://openrouter.ai/api/v1'
        model = os.getenv('OPENAI_MODEL') or os.getenv('OPENROUTER_MODEL') or 'meta-llama/llama-4-scout-17b-16e-instruct'
        key = os.getenv('OPENAI_API_KEY') or os.getenv('OPENROUTER_API_KEY') or ''
    elif provider == 'ollama':
        base = os.getenv('OPENAI_BASE_URL') or os.getenv('OLLAMA_BASE_URL') or 'http://127.0.0.1:11434/v1'
        model = os.getenv('OPENAI_MODEL') or os.getenv('OLLAMA_MODEL') or 'llama3.2:3b-instruct'
        key = os.getenv('OPENAI_API_KEY') or os.getenv('OLLAMA_API_KEY') or 'ollama'
    else:
        # Default to OpenRouter-compatible OpenAI schema to support multiple vendors
        base = os.getenv('OPENAI_BASE_URL') or os.getenv('OPENROUTER_BASE_URL') or 'https://openrouter.ai/api/v1'
        model = os.getenv('OPENAI_MODEL') or os.getenv('OPENROUTER_MODEL') or 'meta-llama/llama-4-scout-17b-16e-instruct'
        key = os.getenv('OPENAI_API_KEY') or os.getenv('OPENROUTER_API_KEY') or ''
    # Apply into app.config so downstream code reads a single place
    app.config['OPENAI_BASE_URL'] = base
    app.config['OPENAI_MODEL'] = model
    app.config['OPENAI_API_KEY'] = key

_apply_provider_defaults()

def _strip_html_tags(value):
    if not value:
        return ''
    text = re.sub(r'(<\s*br\s*/?>)', '\n', value, flags=re.IGNORECASE)
    text = re.sub(r'</p\s*>', '\n\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    # Remove stray Markdown/bullet markers that may remain in content or table cells
    # 1) Strip common bullet prefixes at the start of lines (e.g., "+ ", "- ", "* ", "• ")
    text = re.sub(r'(?m)^\s*[\+\-\*•·▪●◦‣–—]\s+', '', text)
    # 2) Remove leftover Markdown emphasis markers like '#' and '*'
    text = re.sub(r'[#*]+', '', text)
    return re.sub(r'\n{3,}', '\n\n', text).strip()

try:
    import fitz as _pymupdf  # PyMuPDF
    _FITZ_AVAILABLE = hasattr(_pymupdf, "open")
except ImportError:
    _pymupdf = None
    _FITZ_AVAILABLE = False

try:
    from PyPDF2 import PdfReader as _PdfReader
except ImportError:
    _PdfReader = None
from flask_mysqldb import MySQL
from flask_socketio import SocketIO
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from MySQLdb.cursors import DictCursor
from rfp_analyzer_routes import rfp_bp 
# --- App Initialization ---
app = Flask(__name__)
app.config['SECRET_KEY'] = 'a_very_secret_key_for_sessions'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file upload

# Ensure commonly used AI configuration keys are visible via app.config
app.config.setdefault('OPENAI_API_KEY', os.getenv('OPENAI_API_KEY', app.config.get('OPENAI_API_KEY', '')))
app.config.setdefault('OPENAI_BASE_URL', os.getenv('OPENAI_BASE_URL', app.config.get('OPENAI_BASE_URL', 'https://openrouter.ai/api/v1')))
app.config.setdefault('OPENAI_MODEL', os.getenv('OPENAI_MODEL', app.config.get('OPENAI_MODEL', 'meta-llama/llama-4-scout-17b-16e-instruct')))

# Register RFP Analyzer Blueprint
app.register_blueprint(rfp_bp)

# --- Error handling: ensure API endpoints return JSON on errors ---
@app.errorhandler(500)
def _handle_internal_server_error(e):
    try:
        path = request.path or ''
    except Exception:
        path = ''
    if isinstance(path, str) and path.startswith('/api/'):
        return jsonify({'error': 'internal_server_error', 'message': 'An unexpected error occurred.'}), 500
    # Fallback for non-API routes
    return "Internal Server Error", 500

@app.errorhandler(TypeError)
def _handle_type_error(e):
    # Catch cases like "view function did not return a valid response" and return JSON for APIs
    try:
        path = request.path or ''
    except Exception:
        path = ''
    if isinstance(path, str) and path.startswith('/api/'):
        return jsonify({'error': 'type_error', 'message': str(e)}), 500
    return str(e), 500

# --- Stage Constants ---
PIPELINE = ['analyzer', 'business', 'design', 'operations', 'engineer', 'handover']
LABELS = {
    'analyzer': 'BID Analyzer',
    'business': 'Business Development', 
    'design': 'Design Team',
    'operations': 'Operations Team',
    'engineer': 'Site Engineer',
    'handover': 'Handov'
}

def pct_for(stage: str) -> int:
    """Calculate progress percentage based on stage"""
    s = (stage or 'analyzer').lower()
    i = PIPELINE.index(s) if s in PIPELINE else 0
    return int(round(i * (100 / (len(PIPELINE) - 1))))  # 0,20,40,60,80,100

def status_texts(stage: str) -> tuple[str, str]:
    """Generate project status and work status texts based on stage"""
    s = (stage or 'analyzer').lower()
    proj = 'completed' if s == 'handover' else 'ongoing'
    if s == 'analyzer':
        work = 'Initiated by BID Analyzer'
    else:
        i = PIPELINE.index(s) if s in PIPELINE else 0
        prev = PIPELINE[i-1] if i > 0 else ''
        from_txt = LABELS.get(prev, '').replace(' Team', '')
        to_txt = LABELS.get(s, '')
        work = f'Updated by {from_txt} to {to_txt}'
    return proj, work

def generate_team_checklist(cur, g_id, team):
    """Generate team-specific default checklist tasks"""
    team_tasks = {
        'business': [
            {'name': 'Market Analysis', 'description': 'Analyze market potential and competition', 'priority': 'high'},
            {'name': 'Client Communication', 'description': 'Establish communication with client', 'priority': 'high'},
            {'name': 'Proposal Preparation', 'description': 'Prepare business proposal', 'priority': 'medium'},
            {'name': 'Budget Estimation', 'description': 'Estimate project budget and timeline', 'priority': 'high'}
        ],
        'design': [
            {'name': 'Initial Design Concept', 'description': 'Create initial design concepts', 'priority': 'high'},
            {'name': 'Technical Drawings', 'description': 'Prepare detailed technical drawings', 'priority': 'high'},
            {'name': 'Material Selection', 'description': 'Select appropriate materials and specifications', 'priority': 'medium'},
            {'name': 'Design Review', 'description': 'Review and finalize design with team', 'priority': 'medium'},
            {'name': 'Client Approval', 'description': 'Get client approval on design', 'priority': 'high'}
        ],
        'operations': [
            {'name': 'Project Planning', 'description': 'Create detailed project execution plan', 'priority': 'high'},
            {'name': 'Resource Allocation', 'description': 'Allocate resources and personnel', 'priority': 'high'},
            {'name': 'Timeline Management', 'description': 'Set up project timeline and milestones', 'priority': 'medium'},
            {'name': 'Quality Control Setup', 'description': 'Establish quality control procedures', 'priority': 'medium'},
            {'name': 'Risk Assessment', 'description': 'Identify and assess project risks', 'priority': 'high'}
        ],
        'engineer': [
            {'name': 'Site Survey', 'description': 'Conduct detailed site survey', 'priority': 'high'},
            {'name': 'Technical Specifications', 'description': 'Prepare technical specifications', 'priority': 'high'},
            {'name': 'Safety Planning', 'description': 'Develop safety protocols and procedures', 'priority': 'high'},
            {'name': 'Equipment Planning', 'description': 'Plan equipment and tool requirements', 'priority': 'medium'},
            {'name': 'Implementation Plan', 'description': 'Create detailed implementation plan', 'priority': 'high'}
        ]
    }
    
    if team in team_tasks:
        for task in team_tasks[team]:
            # Persist a stage on each seeded task so aggregation can be stage-aware
            stage_name = team.strip().lower()
            cur.execute("""
                INSERT INTO bid_checklists (g_id, task_name, description, priority, status, progress_pct, stage, created_by)
                VALUES (%s, %s, %s, %s, 'pending', %s, %s, %s)
            """, (g_id, task['name'], task['description'], task['priority'], 0, stage_name, current_user.id))

def log_write(action: str, details: str = ''):
    """Write to logs table and emit via Socket.IO"""
    try:
        cur = mysql.connection.cursor()
        user_id = getattr(current_user, 'id', None) if hasattr(current_user, 'id') else None
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)",
                    (f"{action} | {details}", user_id))
        mysql.connection.commit()
        cur.close()
        
        # Emit to master dashboard
        socketio.emit('master_update', {'log': {
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'action': f"{action} | {details}",
            'user_email': getattr(current_user, 'email', 'System') if hasattr(current_user, 'email') else 'System',
            'user_role': getattr(current_user, 'role', '') if hasattr(current_user, 'role') else ''
        }})
    except Exception as e:
        print(f"Log write error: {e}")

# Legacy function for backward compatibility
def stage_progress_pct(stage: str) -> int:
    return pct_for(stage)

# MySQL Config
app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = ''
app.config['MYSQL_DB'] = 'esco'
app.config['MYSQL_CURSORCLASS'] = 'DictCursor'

mysql = MySQL(app)
# Store mysql on app for blueprint access
app.mysql = mysql  # type: ignore[attr-defined]

socketio = SocketIO(app)

_db_init_done = False

@app.before_request
def _init_db_once():
    global _db_init_done
    if not _db_init_done:
        try:
            _ensure_tables_exist()
            # Sanity check for bid_checklists and recover from common errors (1932 orphan tablespace, 1146 missing table)
            def _ensure_bid_checklists_exists():
                create_sql = (
                    """
                    CREATE TABLE IF NOT EXISTS bid_checklists (
                        id INT AUTO_INCREMENT PRIMARY KEY,
                        g_id INT NOT NULL,
                        task_name VARCHAR(200) NOT NULL,
                        description TEXT,
                        assigned_to INT,
                        status VARCHAR(50) DEFAULT 'pending',
                        progress_pct INT DEFAULT NULL,
                        stage VARCHAR(50),
                        priority VARCHAR(20) DEFAULT 'medium',
                        due_date DATETIME,
                        attachment_path VARCHAR(255),
                        created_by INT,
                        team_archive VARCHAR(50),
                        created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                        updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                        FOREIGN KEY (g_id) REFERENCES go_bids(g_id),
                        FOREIGN KEY (assigned_to) REFERENCES employees(id),
                        FOREIGN KEY (created_by) REFERENCES users(id)
                    ) ENGINE=InnoDB
                    """
                )
                try:
                    c = mysql.connection.cursor()
                    c.execute(create_sql)
                    mysql.connection.commit()
                    c.close()
                except Exception as ce:
                    try:
                        err_no2 = ce.args[0] if hasattr(ce, 'args') and ce.args else None
                    except Exception:
                        err_no2 = None
                    if err_no2 == 1813:  # Orphan tablespace
                        try:
                            c2 = mysql.connection.cursor()
                            c2.execute("SHOW VARIABLES LIKE 'datadir'")
                            row = c2.fetchone()
                            data_dir = None
                            if isinstance(row, (list, tuple)) and len(row) >= 2:
                                data_dir = row[1]
                            elif isinstance(row, dict):
                                data_dir = row.get('Value') or row.get('value')
                            c2.close()
                            if data_dir:
                                db_name = app.config.get('MYSQL_DB', 'esco')
                                ibd_path = os.path.join(data_dir, db_name, 'bid_checklists.ibd')
                                cfg_path = os.path.join(data_dir, db_name, 'bid_checklists.cfg')
                                for p in (ibd_path, cfg_path):
                                    try:
                                        if os.path.exists(p):
                                            os.remove(p)
                                    except Exception:
                                        pass
                        except Exception:
                            pass
                        try:
                            c3 = mysql.connection.cursor()
                            c3.execute("DROP TABLE IF EXISTS bid_checklists")
                            mysql.connection.commit()
                            c3.close()
                        except Exception:
                            pass
                        c4 = mysql.connection.cursor()
                        c4.execute(create_sql)
                        mysql.connection.commit()
                        c4.close()
                    else:
                        raise
            # Verify existence
            success = False
            for _ in range(2):
                try:
                    cur = mysql.connection.cursor()
                    cur.execute("SELECT 1 FROM bid_checklists LIMIT 1")
                    cur.close()
                    success = True
                    break
                except Exception as e:
                    try:
                        err_no = e.args[0] if hasattr(e, 'args') and e.args else None
                    except Exception:
                        err_no = None
                    if err_no in (1146, 1932):
                        _ensure_bid_checklists_exists()
                    else:
                        break
            if success:
                _db_init_done = True
        except Exception as e:
            print(f"Database init error: {e}")

# Removed eager import-time init to avoid referencing functions before definition

@app.route('/profiling', methods=['GET', 'POST'], endpoint='profiling')
@login_required
def profiling_page():
    if not current_user.is_admin:
        return "Access Denied", 403
    # Knowledge Hub base dir
    base_dir = os.path.join(os.getcwd(), 'uploads', 'knowledge')
    os.makedirs(base_dir, exist_ok=True)
    message = ''
    preferences = None
    capabilities = []
    performances = []
    details = None
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Ensure mapping table exists for multi-assign (first thing)
        try:
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS bid_assignment_members (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    g_id INT NOT NULL,
                    employee_id INT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE KEY uniq_g_emp (g_id, employee_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
                """
            )
        except Exception:
            pass
        if request.method == 'POST':
            action = request.form.get('action')
            if action == 'save_preferences':
                states = (request.form.get('registered_states') or '').strip()
                cur.execute("DELETE FROM company_preferences")
                cur.execute("INSERT INTO company_preferences (registered_states) VALUES (%s)", (states,))
                mysql.connection.commit()
                message = 'Preferences saved.'
            elif action == 'add_capability':
                desc = request.form.get('cap_description') or ''
                file = request.files.get('cap_file')
                file_path = None
                if file and file.filename:
                    uploads_dir = os.path.join(os.getcwd(), 'uploads', 'capabilities')
                    os.makedirs(uploads_dir, exist_ok=True)
                    fname = secure_filename(file.filename)
                    save_path = os.path.join(uploads_dir, fname)
                    file.save(save_path)
                    file_path = f"capabilities/{fname}"
                cur.execute("INSERT INTO company_capabilities (description, file_path) VALUES (%s, %s)", (desc, file_path))
                mysql.connection.commit()
                message = 'Capability added.'
            elif action == 'add_performance':
                project = (request.form.get('project_name') or '').strip()
                year_raw = (request.form.get('project_year') or '').strip()
                year = None
                if year_raw:
                    try:
                        year = int(year_raw[:4])
                    except Exception:
                        year = None
                cur.execute("INSERT INTO company_performance (project_name, year) VALUES (%s, %s)", (project, year))
                mysql.connection.commit()
                message = 'Past performance saved.'
            elif action == 'save_details':
                name = (request.form.get('name') or '').strip()
                website = (request.form.get('website') or '').strip()
                email = (request.form.get('email') or '').strip()
                phone = (request.form.get('phone') or '').strip()
                about = (request.form.get('about') or '').strip()
                # Store a single row snapshot (replace existing)
                cur.execute("DELETE FROM company_details")
                cur.execute(
                    "INSERT INTO company_details (name, website, email, phone, about) VALUES (%s,%s,%s,%s,%s)",
                    (name, website, email, phone, about)
                )
                mysql.connection.commit()
                message = 'Company details saved.'
            elif action == 'create_folder':
                folder_name = (request.form.get('folder_name') or '').strip()
                if folder_name:
                    target = os.path.join(base_dir, folder_name)
                    os.makedirs(target, exist_ok=True)
                    message = f'Folder "{folder_name}" created.'
            elif action == 'upload_docs':
                target_folder = (request.form.get('target_folder') or '').strip()
                target_path = os.path.join(base_dir, target_folder) if target_folder else base_dir
                os.makedirs(target_path, exist_ok=True)
                files = request.files.getlist('documents')
                for f in files:
                    if not f.filename:
                        continue
                    safe_name = f.filename.replace('..','_')
                    f.save(os.path.join(target_path, safe_name))
                message = 'Documents uploaded.'
        else:
            # Fallback: accept GET query submission for company details if present (prevents user confusion)
            q = request.args
            if any(k in q for k in ('name','website','email','phone','about')):
                name = (q.get('name') or '').strip()
                website = (q.get('website') or '').strip()
                email = (q.get('email') or '').strip()
                phone = (q.get('phone') or '').strip()
                about = (q.get('about') or '').strip()
                if any([name, website, email, phone, about]):
                    cur.execute("DELETE FROM company_details")
                    cur.execute(
                        "INSERT INTO company_details (name, website, email, phone, about) VALUES (%s,%s,%s,%s,%s)",
                        (name, website, email, phone, about)
                    )
                    mysql.connection.commit()
                    message = 'Company details saved.'

        # Load current data for all tabs
        cur.execute("SELECT * FROM company_preferences ORDER BY id DESC LIMIT 1")
        preferences = cur.fetchone()
        # Some databases may not have the uploaded_at column from older migrations; fall back to id
        try:
            cur.execute("SELECT * FROM company_capabilities ORDER BY uploaded_at DESC")
        except Exception:
            cur.execute("SELECT * FROM company_capabilities ORDER BY id DESC")
        capabilities = cur.fetchall()
        cur.execute("SELECT * FROM company_performance ORDER BY year DESC, id DESC")
        performances = cur.fetchall()
        cur.execute("SELECT * FROM company_details ORDER BY id DESC LIMIT 1")
        details = cur.fetchone()
        # Load optional: past_performance (support common misspelling too)
        past_performances = []
        try:
            cur.execute("SELECT * FROM past_performance ORDER BY year DESC, id DESC")
            past_performances = cur.fetchall()
        except Exception:
            try:
                cur.execute("SELECT * FROM past_performace ORDER BY year DESC, id DESC")
                past_performances = cur.fetchall()
            except Exception:
                past_performances = []
        # Load optional: personnel (fallback to employees)
        personnel = []
        try:
            cur.execute("SELECT * FROM personnel ORDER BY id DESC")
            personnel = cur.fetchall()
        except Exception:
            try:
                cur.execute("""
                    SELECT name, department AS role, email 
                    FROM employees 
                    WHERE is_active = TRUE 
                    ORDER BY name ASC
                """)
                personnel = cur.fetchall()
            except Exception:
                personnel = []
        # Load optional: projects_bids
        projects_bids = []
        try:
            cur.execute("""
                SELECT * FROM projects_bids
                ORDER BY COALESCE(due_date, created_at, updated_at) DESC, id DESC
                LIMIT 20
            """)
            projects_bids = cur.fetchall()
        except Exception:
            projects_bids = []
    finally:
        cur.close()

    # Build knowledge hub lists
    folders = []
    files = []
    try:
        for entry in sorted(os.listdir(base_dir)):
            p = os.path.join(base_dir, entry)
            if os.path.isdir(p):
                folders.append(entry)
            else:
                files.append(entry)
    except Exception:
        pass

    return render_template('profiling.html',
                           preferences=preferences,
                           capabilities=capabilities,
                           performances=performances,
                           past_performances=past_performances,
                           personnel=personnel,
                           projects_bids=projects_bids,
                           details=details,
                           folders=folders,
                           files=files,
                           message=message)
# --- Ensure seed tasks exist for every team (parallel kickoff) ---
def ensure_tasks_for_team(team: str):
    cur = mysql.connection.cursor(DictCursor)
    try:
        cur.execute(
            """
            SELECT gb.g_id
            FROM go_bids gb
            LEFT JOIN bid_checklists bc
              ON bc.g_id = gb.g_id AND LOWER(COALESCE(bc.stage,'')) = %s
            WHERE bc.id IS NULL
            """, (team,)
        )
        missing = [row['g_id'] for row in cur.fetchall()]
        if missing:
            for g_id in missing:
                generate_team_checklist(cur, g_id, team)
            mysql.connection.commit()
    except Exception:
        mysql.connection.rollback()
    finally:
        cur.close()
@app.route('/knowledge', methods=['GET', 'POST'])
@login_required
def knowledge():
    if not current_user.is_admin:
        return "Access Denied", 403
    base_dir = os.path.join(os.getcwd(), 'uploads', 'knowledge')
    os.makedirs(base_dir, exist_ok=True)

    message = ''
    if request.method == 'POST':
        action = request.form.get('action')
        if action == 'create_folder':
            folder_name = (request.form.get('folder_name') or '').strip()
            if folder_name:
                target = os.path.join(base_dir, folder_name)
                os.makedirs(target, exist_ok=True)
                message = f'Folder "{folder_name}" created.'
        elif action == 'upload_docs':
            target_folder = (request.form.get('target_folder') or '').strip()
            target_path = os.path.join(base_dir, target_folder) if target_folder else base_dir
            os.makedirs(target_path, exist_ok=True)
            files = request.files.getlist('documents')
            for f in files:
                if not f.filename:
                    continue
                safe_name = f.filename.replace('..','_')
                f.save(os.path.join(target_path, safe_name))
            message = 'Documents uploaded.'

    # List folders and files
    folders = []
    files = []
    for entry in sorted(os.listdir(base_dir)):
        p = os.path.join(base_dir, entry)
        if os.path.isdir(p):
            folders.append(entry)
        else:
            files.append(entry)

    return render_template('profiling.html', folders=folders, files=files, message=message)
# --- Helper deprecated: revenue-based auto assignment removed ---
def assign_bids_by_revenue():
    return None

# assign_go_bids deprecated
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

@login_manager.unauthorized_handler
def _unauthorized_handler():
    # Return JSON errors for API endpoints instead of HTML redirects
    try:
        path = request.path or ''
    except Exception:
        path = ''
    if isinstance(path, str) and path.startswith('/api/'):
        return jsonify({'error': 'unauthorized', 'message': 'Login required. Please sign in and try again.'}), 401
    return redirect(url_for('login'))

# Custom User class
class User(UserMixin):
    def __init__(self, id, email, password, is_admin, role):
        self.id = id
        self.email = email
        self.password = password
        self.is_admin = bool(is_admin)
        self.role = role
    
    @property
    def is_supervisor(self):
        return self.role.lower() == 'supervisor'
    
    @property
    def can_assign_stages(self):
        return self.is_admin or self.is_supervisor
    
    @property
    def can_alter_timeline(self):
        return self.is_admin or self.is_supervisor
# --- User Loader for Flask-Login ---
@login_manager.user_loader
def load_user(user_id):
    cur = mysql.connection.cursor(DictCursor)
    cur.execute("SELECT * FROM users WHERE id=%s", (user_id,))
    data = cur.fetchone()
    cur.close()
    if data:
        return User(data['id'], data['email'], data['password'], data['is_admin'], data['role'])
    return None

# --- Utility: ensure GO bids from bid_incoming are mirrored in go_bids ---
def sync_go_bids() -> int:
    """Synchronize GO decisions from bid_incoming into go_bids.

    - Inserts any missing GO bids.
    - Updates existing go_bids rows with the latest information.
    - Removes entries that are no longer marked as GO.

    Returns number of newly inserted rows.
    """
    cur = mysql.connection.cursor(DictCursor)
    inserted = 0
    try:
        # Insert new GO bids that do not yet exist in go_bids
        cur.execute(
            """
            INSERT INTO go_bids (id, b_name, due_date, state, scope, type, scoring, company, decision, summary)
            SELECT bi.id,
                   bi.b_name,
                   bi.due_date,
                   COALESCE(NULLIF(TRIM(bi.state), ''), 'business'),
                   bi.scope,
                   bi.type,
                   bi.scoring,
                   bi.comp_name,
                   bi.decision,
                   bi.summary
            FROM bid_incoming bi
            LEFT JOIN go_bids gb ON gb.id = bi.id
            WHERE gb.id IS NULL
              AND UPPER(TRIM(COALESCE(bi.decision, ''))) = 'GO'
            """
        )
        inserted = cur.rowcount or 0

        # Update existing go_bids rows to reflect latest bid_incoming data
        cur.execute(
            """
            UPDATE go_bids gb
            JOIN bid_incoming bi ON bi.id = gb.id
            SET gb.b_name   = bi.b_name,
                gb.due_date = COALESCE(bi.due_date, gb.due_date),
                gb.state    = COALESCE(NULLIF(TRIM(bi.state), ''), gb.state),
                gb.scope    = bi.scope,
                gb.type     = bi.type,
                gb.scoring  = bi.scoring,
                gb.company  = COALESCE(bi.comp_name, gb.company),
                gb.decision = bi.decision,
                gb.summary  = bi.summary
            WHERE UPPER(TRIM(COALESCE(bi.decision, ''))) = 'GO'
            """
        )

        # Remove go_bids entries that are no longer GO decisions or missing source
        # First, get the g_id values that need to be deleted
        cur.execute(
            """
            SELECT gb.g_id
            FROM go_bids gb
            LEFT JOIN bid_incoming bi ON bi.id = gb.id
            WHERE gb.id IS NOT NULL
              AND (
                    bi.id IS NULL
                 OR UPPER(TRIM(COALESCE(bi.decision, ''))) <> 'GO'
              )
            """
        )
        g_ids_to_delete = [row['g_id'] for row in cur.fetchall()]
        
        if g_ids_to_delete:
            # Delete from all tables that have foreign key constraints
            placeholders = ','.join(['%s'] * len(g_ids_to_delete))
            
            # Delete from bid_checklists
            try:
                cur.execute(f"DELETE FROM bid_checklists WHERE g_id IN ({placeholders})", g_ids_to_delete)
            except Exception:
                pass
            
            # Delete from bid_stage_exclusions
            try:
                cur.execute(f"DELETE FROM bid_stage_exclusions WHERE g_id IN ({placeholders})", g_ids_to_delete)
            except Exception:
                pass
            
            # Delete from bid_custom_stages
            try:
                cur.execute(f"DELETE FROM bid_custom_stages WHERE g_id IN ({placeholders})", g_ids_to_delete)
            except Exception:
                pass
            
            # Delete from bid_assignment_members
            try:
                cur.execute(f"DELETE FROM bid_assignment_members WHERE g_id IN ({placeholders})", g_ids_to_delete)
            except Exception:
                pass
            
            # Get a_id values from bid_assign before deleting
            cur.execute(f"SELECT a_id FROM bid_assign WHERE g_id IN ({placeholders})", g_ids_to_delete)
            a_ids = [row['a_id'] for row in cur.fetchall() if row.get('a_id')]
            
            # Delete from bid_assign
            try:
                cur.execute(f"DELETE FROM bid_assign WHERE g_id IN ({placeholders})", g_ids_to_delete)
            except Exception:
                pass
            
            # Delete related win_lost_results, won_bids_result, and work_progress_status
            if a_ids:
                a_placeholders = ','.join(['%s'] * len(a_ids))
                try:
                    cur.execute(f"SELECT w_id FROM win_lost_results WHERE a_id IN ({a_placeholders})", a_ids)
                    w_ids = [row['w_id'] for row in cur.fetchall() if row.get('w_id')]
                    
                    if w_ids:
                        w_placeholders = ','.join(['%s'] * len(w_ids))
                        try:
                            cur.execute(f"SELECT won_id FROM won_bids_result WHERE w_id IN ({w_placeholders})", w_ids)
                            won_ids = [row['won_id'] for row in cur.fetchall() if row.get('won_id')]
                            
                            if won_ids:
                                won_placeholders = ','.join(['%s'] * len(won_ids))
                                try:
                                    cur.execute(f"DELETE FROM work_progress_status WHERE won_id IN ({won_placeholders})", won_ids)
                                except Exception:
                                    pass
                            
                            try:
                                cur.execute(f"DELETE FROM won_bids_result WHERE w_id IN ({w_placeholders})", w_ids)
                            except Exception:
                                pass
                        except Exception:
                            pass
                    
                    try:
                        cur.execute(f"DELETE FROM win_lost_results WHERE a_id IN ({a_placeholders})", a_ids)
                    except Exception:
                        pass
                except Exception:
                    pass
            
            # Finally, delete from go_bids
            cur.execute(
                """
                DELETE gb FROM go_bids gb
                LEFT JOIN bid_incoming bi ON bi.id = gb.id
                WHERE gb.id IS NOT NULL
                  AND (
                        bi.id IS NULL
                     OR UPPER(TRIM(COALESCE(bi.decision, ''))) <> 'GO'
                  )
                """
            )

        mysql.connection.commit()
        return inserted
    except Exception as sync_err:
        mysql.connection.rollback()
        print(f"Error synchronizing GO bids: {sync_err}")
        return 0
    finally:
        cur.close()

# --- Authentication Routes ---
@app.route('/register', methods=['GET', 'POST'])
def register():
    cur = mysql.connection.cursor(DictCursor)
    
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '').strip()
        role = request.form.get('role', 'member').strip() or 'member'
        company_id = request.form.get('company_id', '').strip()
        
        if not email or not password:
            cur.execute("SELECT id, name FROM companies ORDER BY name")
            companies = cur.fetchall()
            cur.close()
            return render_template('register.html', companies=companies, error='Email and password are required')
        
        # Check if email already exists
        cur.execute("SELECT id FROM users WHERE email=%s", (email,))
        if cur.fetchone():
            cur.execute("SELECT id, name FROM companies ORDER BY name")
            companies = cur.fetchall()
            cur.close()
            return render_template('register.html', companies=companies, error='Email already exists. Please use a different email or login.')
        
        is_admin = 1 if role.lower() == "admin" else 0
        cur.execute("INSERT INTO users (email, password, role, is_admin) VALUES (%s, %s, %s, %s)",
                    (email, password, role, is_admin))
        mysql.connection.commit()
        
        # Fetch companies for the template (even after success)
        cur.execute("SELECT id, name FROM companies ORDER BY name")
        companies = cur.fetchall()
        for company in companies:
            company_name_lower = company['name'].lower()
            if 'ikio' in company_name_lower:
                company['logo_path'] = 'ikio.png'
            elif 'metco' in company_name_lower:
                company['logo_path'] = 'metco.png'
            elif 'sunsprint' in company_name_lower:
                company['logo_path'] = 'sunsprint.png'
            else:
                company['logo_path'] = None
        
        cur.close()
        
        return render_template('register.html', 
                             companies=companies, 
                             success='Account created successfully! You can now login.',
                             show_login=True)
    
    # GET request - show registration form
    cur.execute("SELECT id, name FROM companies ORDER BY name")
    companies = cur.fetchall()
    
    # Add logo paths based on company name (fallback if logo_path column doesn't exist)
    for company in companies:
        company_name_lower = company['name'].lower()
        if 'ikio' in company_name_lower:
            company['logo_path'] = 'ikio.png'
        elif 'metco' in company_name_lower:
            company['logo_path'] = 'metco.png'
        elif 'sunsprint' in company_name_lower:
            company['logo_path'] = 'sunsprint.png'
        else:
            company['logo_path'] = None
    
    cur.close()
    return render_template('register.html', companies=companies)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        cur = mysql.connection.cursor(DictCursor)
        cur.execute("SELECT * FROM users WHERE email=%s AND password=%s", (email, password))
        user = cur.fetchone()
        if user:
            cur.close()
            user_obj = User(user['id'], user['email'], user['password'], user['is_admin'], user['role'])
            login_user(user_obj)
            log_write('login', f"role={user['role']}")
            if user['is_admin']:
                return redirect(url_for('master_dashboard'))
            
            # Role-based redirects
            role = user['role'].lower()
            if role == 'supervisor':
                return redirect(url_for('supervisor_dashboard'))
            elif role == 'business dev':
                return redirect(url_for('team_dashboard', team='business'))
            elif role == 'design':
                return redirect(url_for('team_dashboard', team='design'))
            elif role == 'operations':
                return redirect(url_for('team_dashboard', team='operations'))
            elif role == 'site manager':
                return redirect(url_for('team_dashboard', team='engineer'))
            else:
                return redirect(url_for('dashboard'))
        # Fallback: try authenticating as employee
        try:
            # Ensure password column exists before querying
            cur.execute("SHOW COLUMNS FROM employees LIKE 'password'")
            if cur.fetchone() is None:
                cur.close()
                return 'Employee login unavailable: employees.password missing', 500
            cur.execute(
                """
                SELECT * FROM employees
                WHERE email=%s AND password=%s AND is_active=1
                """,
                (email, password)
            )
            employee = cur.fetchone()
        finally:
            cur.close()

        if employee:
            session['employee_id'] = employee['id']
            session['employee_name'] = employee['name']
            session['employee_email'] = employee['email']
            session['employee_department'] = employee['department']
            log_write('employee_login', f"Employee {employee['name']} logged in via /login")
            return redirect(url_for('employee_dashboard', employee_id=employee['id']))

        return 'Invalid credentials', 401
    return render_template('login.html')

@app.route('/admin/users', methods=['GET', 'POST'])
@login_required
def users_admin():
    if not current_user.is_admin:
        return "Access Denied", 403
    cur = mysql.connection.cursor(DictCursor)
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '').strip()
        role = request.form.get('role', 'member').strip() or 'member'
        if not email or not password:
            return 'Email and password required', 400
        # Check if email already exists
        cur.execute("SELECT id FROM users WHERE email=%s", (email,))
        if cur.fetchone():
            cur.close()
            return 'Email already exists', 400
        is_admin = 1 if role.lower() == "admin" else 0
        cur.execute("INSERT INTO users (email, password, role, is_admin) VALUES (%s, %s, %s, %s)",
                    (email, password, role, is_admin))
        mysql.connection.commit()
        cur.close()
        return redirect(url_for('users_admin'))
    cur.execute("SELECT * FROM users ORDER BY id ASC")
    users = cur.fetchall()
    cur.close()
    return render_template('users.html', users=users)

@app.route('/logout')
def logout():
    log_write('logout')
    logout_user()
    return redirect(url_for('login'))

@app.route('/employee/login', methods=['POST'])
def employee_login():
    """Employee login route"""
    try:
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '').strip()
        # employee_id may be provided from the UI but we don't require it for auth
        if not email or not password:
            return jsonify({'success': False, 'error': 'Missing credentials'}), 400

        cur = mysql.connection.cursor(DictCursor)
        # Ensure password column exists; if not, fail with a clear message
        try:
            cur.execute("SHOW COLUMNS FROM employees LIKE 'password'")
            has_password_col = cur.fetchone() is not None
        except Exception:
            has_password_col = False
        if not has_password_col:
            cur.close()
            return jsonify({'success': False, 'error': "employees.password column missing. Run add_password_column.sql or migration."}), 500
        cur = mysql.connection.cursor(DictCursor)
        cur.execute(
            """
            SELECT * FROM employees 
            WHERE email = %s AND password = %s AND is_active = 1
            """,
            (email, password),
        )
        employee = cur.fetchone()
        cur.close()

        if employee:
            # Create a simple session for employee (not using Flask-Login for employees)
            session['employee_id'] = employee['id']
            session['employee_name'] = employee['name']
            session['employee_email'] = employee['email']
            session['employee_department'] = employee['department']

            log_write('employee_login', f"Employee {employee['name']} logged in")
            return jsonify({'success': True, 'employee_id': employee['id']})
        else:
            return jsonify({'success': False, 'error': 'Invalid credentials'}), 401

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/employee/logout')
def employee_logout():
    """Employee logout route"""
    if 'employee_id' in session:
        log_write('employee_logout', f"Employee {session.get('employee_name', 'Unknown')} logged out")
        session.pop('employee_id', None)
        session.pop('employee_name', None)
        session.pop('employee_email', None)
        session.pop('employee_department', None)
    return redirect(url_for('login'))

# --- Dashboard Routes ---
@app.route('/')
def index():
    if current_user.is_authenticated:
        if current_user.is_admin:
            return redirect(url_for('master_dashboard'))
        elif current_user.is_supervisor:
            return redirect(url_for('supervisor_dashboard'))
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/dashboard')
@login_required
def dashboard():
    # Redirect non-admins to their role-specific dashboards
    if current_user.is_admin:
        return redirect(url_for('master_dashboard'))
    elif current_user.is_supervisor:
        return redirect(url_for('supervisor_dashboard'))
    return redirect(url_for('role_dashboard'))

@app.route('/dashboard/role')
@login_required
def role_dashboard():
    if current_user.is_admin:
        return redirect(url_for('master_dashboard'))
    role = (current_user.role or 'member').strip()
    role_key = role.lower()

    # Map business flow stages to roles
    role_to_stage = {
        'business dev': 'business',
        'business': 'business',  # backward compatibility
        'bdm': 'business',
        'design': 'design',
        'operations': 'operations',
        'site manager': 'site_manager'
    }
    next_stage_map = {
        'analyzer': 'business',
        'business': 'design',
        'design': 'operations',
        'operations': 'site_manager',
        'site_manager': 'handover'
    }

    current_stage_for_role = role_to_stage.get(role_key)
    cur = mysql.connection.cursor(DictCursor)
    bids = []
    # Query by go_bids.state (not bid_assign.depart) and join assignee info
    role_to_stage = {
        'business dev': 'business',
        'business': 'business', 
        'bdm': 'business',
        'design': 'design',
        'operations': 'operations',
        'site manager': 'engineer',
        'site_manager': 'engineer'
    }
    current_stage_for_role = role_to_stage.get((current_user.role or 'member').lower())
    
    sql = """
        SELECT gb.g_id AS id,
               gb.b_name AS name,
               gb.company, 
               gb.due_date,
               COALESCE(gb.scoring, 0) AS progress,
               LOWER(COALESCE(gb.state, 'analyzer')) AS current_stage,
               ba.person_name,
               ba.assignee_email AS person_email,
               gb.summary
        FROM go_bids gb
        LEFT JOIN bid_assign ba ON ba.g_id = gb.g_id
        {where}
        ORDER BY gb.due_date ASC
    """
    where = "WHERE LOWER(COALESCE(gb.state, 'analyzer')) = %s" if current_stage_for_role else ""
    try:
        cur.execute(sql.format(where=where), (current_stage_for_role,) if current_stage_for_role else ())
        bids = cur.fetchall()
    except Exception as e:
        # Fallback for databases that don't yet have go_bids.state
        if "Unknown column 'state'" in str(e):
            sql_fallback = """
                SELECT gb.g_id AS id,
                       gb.b_name AS name,
                       gb.company, 
                       gb.due_date,
                       COALESCE(gb.scoring, 0) AS progress,
                       LOWER(COALESCE(ba.depart, 'analyzer')) AS current_stage,
                       ba.person_name,
                       ba.assignee_email AS person_email,
                       gb.summary
                FROM go_bids gb
                LEFT JOIN bid_assign ba ON ba.g_id = gb.g_id
                ORDER BY gb.due_date ASC
            """
            cur.execute(sql_fallback)
            rows = cur.fetchall()
            bids = [r for r in rows if (not current_stage_for_role or r.get('current_stage') == current_stage_for_role)]
        else:
            raise
    # Keep next_stage mapping in Python and pass it to the template
    next_stage_map = {
        'analyzer': 'business',
        'business': 'design', 
        'design': 'operations',
        'operations': 'engineer',
        'engineer': 'handover'
    }
    for bid in bids:
        bid['user'] = {'email': bid.get('person_email'), 'role': current_stage_for_role}
        bid['next_stage'] = next_stage_map.get(bid['current_stage'])
        # Add dynamic progress and status texts
        stage_key = (bid.get('current_stage') or 'analyzer').lower()
        bid['work_progress_pct'] = stage_progress_pct(stage_key)
        bid['project_status'], bid['work_status'] = status_texts(stage_key)
    # Build dynamic stage lists per bid similar to supervisor view
    # Ensure tables exist
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS bid_stage_exclusions (
            id INT AUTO_INCREMENT PRIMARY KEY,
            g_id INT NOT NULL,
            stage VARCHAR(50) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE KEY uniq_bid_stage (g_id, stage)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS bid_custom_stages (
            id INT AUTO_INCREMENT PRIMARY KEY,
            g_id INT NOT NULL,
            stage VARCHAR(50) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE KEY uniq_custom_stage (g_id, stage)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """
    )
    # Load exclusions and customs
    cur.execute("SELECT g_id, stage FROM bid_stage_exclusions")
    ex_rows = cur.fetchall()
    exclusions = {}
    for r in ex_rows:
        exclusions.setdefault(r['g_id'], set()).add((r['stage'] or '').strip().lower())
    cur.execute("SELECT g_id, stage FROM bid_custom_stages")
    cs_rows = cur.fetchall()
    customs = {}
    for r in cs_rows:
        customs.setdefault(r['g_id'], []).append((r['stage'] or '').strip().lower())
    default_stages = ['analyzer', 'business', 'design', 'operations', 'engineer']
    for bid in bids:
        bid_id = bid.get('id')
        excl = exclusions.get(bid_id, set())
        cust = [s for s in customs.get(bid_id, []) if s not in excl]
        dyn_stages = [s for s in default_stages if s not in excl] + [s for s in cust if s not in default_stages]
        bid['dyn_stages'] = dyn_stages

    cur.close()

    # Use a generic role dashboard that includes an advance-stage button
    normalized_role = ' '.join([w.capitalize() for w in role_key.split()])
    return render_template('dashboard_role.html', bids=bids, user=current_user, role=normalized_role,
                           current_stage=current_stage_for_role, next_stage_map=next_stage_map)

@app.route('/dashboard/<team>')
@login_required
def team_dashboard(team):
    """Team-specific dashboard for Business Dev, Design, Operations, Site Manager"""
    if current_user.is_admin:
        return redirect(url_for('master_dashboard'))
    
    # Keep go_bids in sync with GO decisions so teams see new bids immediately
    try:
        sync_go_bids()
    except Exception:
        pass

    # Map team names to stages
    team_to_stage = {
        'business': 'business',
        'design': 'design', 
        'operations': 'operations',
        'engineer': 'engineer'
    }
    
    if team not in team_to_stage:
        return "Invalid team", 404
    
    current_stage = team_to_stage[team]
    # Ensure seed tasks exist so all bids appear on this team dashboard
    ensure_tasks_for_team(current_stage)
    cur = mysql.connection.cursor(DictCursor)
    
    # Company access filter based on user email domain
    email = (current_user.email or '').lower()
    company_filter_name = None
    if email.endswith('@ikioledlighting.com'):
        company_filter_name = 'Ikio'
    elif email.endswith('@metcoengineering.com'):
        company_filter_name = 'Metco'
    elif email.endswith('@sunsprintengineering.com'):
        company_filter_name = 'Sunsprint'

    company_sql_clause = ""
    company_params = tuple()
    if company_filter_name:
        # Accept aliases like 'METCO Engineering' by matching on keyword
        pattern = None
        name_lower = company_filter_name.lower()
        if name_lower == 'metco':
            pattern = '%metco%'
        elif name_lower == 'ikio':
            pattern = '%ikio%'
        elif name_lower == 'sunsprint':
            pattern = '%sunsprint%'
        if pattern:
            company_sql_clause = " AND LOWER(COALESCE(gb.company,'')) LIKE %s"
            company_params = (pattern,)

    # Fetch all bids for this company scope; filtering by team is done in Python
    try:
        cur.execute("""
        SELECT gb.g_id AS id,
               gb.b_name AS name,
               gb.company, 
               gb.due_date,
               COALESCE(gb.scoring, 0) AS progress,
               LOWER(COALESCE(gb.state, 'analyzer')) AS current_stage,
               ba.person_name,
               ba.assignee_email AS person_email,
               ba.depart,
               wps.pr_completion_status AS work_status,
               wbr.closure_status AS project_status,
               wbr.work_progress_status AS work_progress_status,
               wlr.result AS wl_result,
               gb.summary
        FROM go_bids gb
        LEFT JOIN bid_assign ba ON ba.g_id = gb.g_id
        LEFT JOIN win_lost_results wlr ON wlr.a_id = ba.a_id
        LEFT JOIN won_bids_result wbr ON wbr.w_id = wlr.w_id
        LEFT JOIN work_progress_status wps ON wps.won_id = wbr.won_id
        WHERE 1=1
        """ + company_sql_clause + """
        ORDER BY gb.due_date ASC
    """, company_params)
        rows = cur.fetchall()
    except Exception as e:
        # Fallback when go_bids.state column doesn't exist yet
        if "Unknown column 'state'" in str(e):
            cur.execute("""
                SELECT gb.g_id AS id,
                       gb.b_name AS name,
                       gb.company, 
                       gb.due_date,
                       COALESCE(gb.scoring, 0) AS progress,
                       LOWER(COALESCE(ba.depart, 'analyzer')) AS current_stage,
                       ba.person_name,
                       ba.assignee_email AS person_email,
                       ba.depart,
                       wps.pr_completion_status AS work_status,
                       wbr.closure_status AS project_status,
                       wbr.work_progress_status AS work_progress_status,
                       wlr.result AS wl_result,
                       gb.summary
                FROM go_bids gb
                LEFT JOIN bid_assign ba ON ba.g_id = gb.g_id
                LEFT JOIN win_lost_results wlr ON wlr.a_id = ba.a_id
                LEFT JOIN won_bids_result wbr ON wbr.w_id = wlr.w_id
                LEFT JOIN work_progress_status wps ON wps.won_id = wbr.won_id
                WHERE 1=1
                """ + company_sql_clause + """
                ORDER BY gb.due_date ASC
            """, company_params)
            rows = cur.fetchall()
        else:
            raise
    # Remove duplicate bids by same name and due date (front-end filtering only)
    _seen_pairs = set()
    _unique_rows = []
    for _r in rows:
        _key = (((_r.get('name') or '').strip().lower()), str(_r.get('due_date') or ''))
        if _key in _seen_pairs:
            continue
        _seen_pairs.add(_key)
        _unique_rows.append(_r)
    rows = _unique_rows

    # Ensure dynamic stage tables exist and load per-bid configuration
    try:
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_stage_exclusions (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_bid_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_custom_stages (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_custom_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
    except Exception:
        # If creation fails (e.g., no rights), continue – downstream selects may still work
        pass

    # Safely load bid stage exclusions; if table is missing or broken, fall back to empty
    try:
        cur.execute("SELECT g_id, stage FROM bid_stage_exclusions")
        ex_rows = cur.fetchall()
    except Exception:
        ex_rows = []
    exclusions = {}
    for r in ex_rows:
        exclusions.setdefault(r['g_id'], set()).add((r['stage'] or '').strip().lower())

    # Safely load custom stages; if table is missing or broken, fall back to empty
    try:
        cur.execute("SELECT g_id, stage FROM bid_custom_stages")
        cs_rows = cur.fetchall()
    except Exception:
        cs_rows = []
    customs = {}
    for r in cs_rows:
        customs.setdefault(r['g_id'], []).append((r['stage'] or '').strip().lower())

    # Preload set of bids that already have tasks for this team's stage
    try:
        cur.execute("SELECT DISTINCT g_id FROM bid_checklists WHERE LOWER(COALESCE(stage,''))=%s", (current_stage,))
        task_ids = {row['g_id'] for row in cur.fetchall()}
    except Exception:
        task_ids = set()

    default_stages = ['analyzer', 'business', 'design', 'operations', 'engineer']
    bids = []
    for row in rows:
        bid_id = row.get('id')
        excl = exclusions.get(bid_id, set())
        cust = [s for s in customs.get(bid_id, []) if s not in excl]
        dyn_stages = [s for s in default_stages if s not in excl] + [s for s in cust if s not in default_stages]
        # If no dynamic config exists, keep default stages
        if not dyn_stages:
            dyn_stages = default_stages.copy()

        stage_now = (row.get('current_stage') or 'analyzer').lower()
        should_include = (
            stage_now == current_stage
            or bid_id in task_ids
            or current_stage in dyn_stages
        )
        if not should_include:
            continue

        # Shape bid dict like before and attach dynamic stages
        row['user'] = {'email': row.get('person_email'), 'role': current_stage}
        row['next_stage'] = None  # resolved in template via get_next_stage
        stage_key = stage_now
        row['work_progress_pct'] = stage_progress_pct(stage_key)
        row['project_status'], row['work_status'] = status_texts(stage_key)
        row['dyn_stages'] = dyn_stages
        bids.append(row)

    # Compute stats from filtered bids
    total_bids = len(bids)
    completed_bids = sum(1 for b in bids if (b.get('wl_result') == 'WON'))

    # Load assigned members for each bid from bid_assignment_members table
    bid_ids = [b['id'] for b in bids]
    assigned_map = {}
    try:
        if bid_ids:
            # Ensure mapping table exists
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS bid_assignment_members (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    g_id INT NOT NULL,
                    employee_id INT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE KEY uniq_g_emp (g_id, employee_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
                """
            )
            placeholders = ','.join(['%s'] * len(bid_ids))
            cur.execute(
                f"""
                SELECT bam.g_id, e.id AS employee_id, e.name, e.email
                FROM bid_assignment_members bam
                JOIN employees e ON e.id = bam.employee_id
                WHERE bam.g_id IN ({placeholders}) AND e.department = %s AND e.is_active = TRUE
                ORDER BY e.name
                """,
                (*bid_ids, team)
            )
            rows = cur.fetchall()
            for r in rows:
                assigned_map.setdefault(r['g_id'], []).append({'id': r['employee_id'], 'name': r['name'], 'email': r['email']})
        # Attach to bid rows
        for b in bids:
            b['assigned_members'] = assigned_map.get(b['id'], [])
    except Exception:
        # If there's an error, just set empty list for each bid
        for b in bids:
            b['assigned_members'] = []
    
    # Load RFP files for each bid from uploaded_rfp_files table
    rfp_files_map = {}
    try:
        if bid_ids:
            # Ensure uploaded_rfp_files table exists
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS uploaded_rfp_files (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    bid_id INT,
                    g_id INT,
                    filename VARCHAR(500) NOT NULL,
                    file_path VARCHAR(1000) NOT NULL,
                    file_size BIGINT,
                    uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    uploaded_by INT,
                    INDEX idx_bid_id (bid_id),
                    INDEX idx_g_id (g_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
                """
            )
            # Check if g_id column exists, if not add it
            try:
                cur.execute("SHOW COLUMNS FROM uploaded_rfp_files LIKE 'g_id'")
                if not cur.fetchone():
                    cur.execute("ALTER TABLE uploaded_rfp_files ADD COLUMN g_id INT, ADD INDEX idx_g_id (g_id)")
                    mysql.connection.commit()
            except Exception:
                pass
            
            placeholders = ','.join(['%s'] * len(bid_ids))
            # Try to query with g_id first, fallback to bid_id if g_id doesn't exist in results
            try:
                cur.execute(
                    f"""
                    SELECT id, g_id, bid_id, filename, file_path, file_size, uploaded_at
                    FROM uploaded_rfp_files
                    WHERE g_id IN ({placeholders}) OR bid_id IN ({placeholders})
                    ORDER BY uploaded_at DESC
                    """,
                    (*bid_ids, *bid_ids)
                )
            except Exception:
                # Fallback if g_id column doesn't exist in query results
                cur.execute(
                    f"""
                    SELECT id, bid_id, filename, file_path, file_size, uploaded_at
                    FROM uploaded_rfp_files
                    WHERE bid_id IN ({placeholders})
                    ORDER BY uploaded_at DESC
                    """,
                    bid_ids
                )
            rows = cur.fetchall()
            for r in rows:
                # bid_ids contains g_id values from go_bids (aliased as 'id')
                g_id = r.get('g_id') or r.get('bid_id')
                if g_id and g_id in bid_ids:
                    if g_id not in rfp_files_map:
                        rfp_files_map[g_id] = []
                    rfp_files_map[g_id].append({
                        'id': r['id'],
                        'filename': r['filename'],
                        'file_path': r['file_path'],
                        'file_size': r.get('file_size', 0),
                        'uploaded_at': r.get('uploaded_at')
                    })
        # Attach RFP files to bid rows
        for b in bids:
            b['rfp_files'] = rfp_files_map.get(b['id'], [])
    except Exception as e:
        # If there's an error, just set empty list for each bid
        print(f"Error loading RFP files: {str(e)}")
        for b in bids:
            b['rfp_files'] = []

    cur.close()
    
    # Map stage names for display
    stage_display_names = {
        'business': 'Business Development',
        'design': 'Design Team', 
        'operations': 'Operations Team',
        'engineer': 'Site Engineer'
    }
    
    team_display_name = stage_display_names.get(team, team.title())
    
    # Add dynamic progress and status texts to each bid
    for bid in bids:
        stage_key = (bid.get('current_stage') or 'analyzer').lower()
        bid['work_progress_pct'] = stage_progress_pct(stage_key)
        bid['project_status'], bid['work_status'] = status_texts(stage_key)
    
    # Define next stage mapping for template
    def get_next_stage(current_stage):
        stage_flow = {
            'analyzer': 'business',
            'business': 'design',
            'design': 'operations',
            'operations': 'engineer',
            'engineer': 'handover'
        }
        return stage_flow.get(current_stage, None)
    
    return render_template('team_dashboard.html', 
                         bids=bids, 
                         team=team,
                         team_display_name=team_display_name,
                         current_stage=current_stage,
                         total_bids=total_bids,
                         completed_bids=completed_bids,
                         user=current_user,
                         get_next_stage=get_next_stage)

main_stats = {
    'total_bids': 350, 'live_bids': 75, 'bids_won': 120, 'projects_completed': 95,
}

# Assuming you have a Log model, e.g., from SQLAlchemy


# ... other imports and app setup

@app.route('/logs')
@login_required
def logs_page():
    # Query all logs, ordering by the most recent first
    cur = mysql.connection.cursor(DictCursor)
    cur.execute("SELECT * FROM logs ORDER BY timestamp DESC")
    all_logs = cur.fetchall()
    cur.close()
    return render_template('logs.html', logs=all_logs)

@app.route('/supervisor-dashboard')
@login_required
def supervisor_dashboard():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    
    cur = mysql.connection.cursor(DictCursor)
    # Ensure exclusions table exists
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS bid_stage_exclusions (
            id INT AUTO_INCREMENT PRIMARY KEY,
            g_id INT NOT NULL,
            stage VARCHAR(50) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE KEY uniq_bid_stage (g_id, stage)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """
    )

    # Ensure custom stages table exists
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS bid_custom_stages (
            id INT AUTO_INCREMENT PRIMARY KEY,
            g_id INT NOT NULL,
            stage VARCHAR(50) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE KEY uniq_custom_stage (g_id, stage)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """
    )

    # Auto-sync GO bids from bid_incoming to go_bids table
    try:
        cur.execute("""
            INSERT INTO go_bids (id, b_name, due_date, state, scope, type, scoring, company, decision, summary)
            SELECT bi.id, bi.b_name, bi.due_date, bi.state, bi.scope, bi.type, bi.scoring, bi.comp_name, bi.decision, bi.summary
            FROM bid_incoming bi
            LEFT JOIN go_bids gb ON gb.id = bi.id
            WHERE UPPER(bi.decision) = 'GO' AND gb.id IS NULL
        """)
        mysql.connection.commit()
    except Exception as e:
        mysql.connection.rollback()
        # Log error but continue - don't break the dashboard
        pass

    # Get all GO bids for stage assignment (including summary and scope)
    cur.execute("""
        SELECT gb.g_id, gb.b_name, gb.state, gb.company, gb.type, gb.summary, gb.scope,
               ba.person_name, ba.depart
        FROM go_bids gb
        LEFT JOIN bid_assign ba ON ba.g_id = gb.g_id
        ORDER BY gb.g_id DESC
    """)
    go_bids = cur.fetchall()
    # Load companies for dropdowns
    cur.execute("SELECT id, name FROM companies ORDER BY name ASC")
    companies = cur.fetchall()
    # Fetch stage exclusions per bid
    cur.execute("SELECT g_id, stage FROM bid_stage_exclusions")
    rows = cur.fetchall()
    stage_exclusions = {}
    for r in rows:
        stage_exclusions.setdefault(r['g_id'], []).append(r['stage'])

    # Fetch custom stages per bid
    cur.execute("SELECT g_id, stage FROM bid_custom_stages")
    rows = cur.fetchall()
    custom_stages = {}
    for r in rows:
        custom_stages.setdefault(r['g_id'], []).append(r['stage'])

    # Get stage options
    stages = ['analyzer', 'business', 'design', 'operations', 'engineer']

    # Attach dynamic stages and progress to each bid row for table rendering
    default_stages = stages.copy()
    for bid in go_bids:
        bid_id = bid['g_id']
        excluded = set(stage_exclusions.get(bid_id, []))
        customs = [s for s in custom_stages.get(bid_id, []) if s not in excluded]
        dyn_stages = [s for s in default_stages if s not in excluded] + [s for s in customs if s not in default_stages]
        current_stage = (bid.get('state') or 'analyzer').strip().lower()
        if dyn_stages and current_stage in dyn_stages:
            idx = dyn_stages.index(current_stage)
            pct = int(round((idx / (max(1, len(dyn_stages) - 1))) * 100))
        else:
            pct = 0
        bid['dyn_stages'] = dyn_stages
        bid['dyn_progress_pct'] = pct
    
    cur.close()
    
    return render_template('supervisor_dashboard.html', 
                         go_bids=go_bids, 
                         stages=stages,
                         stage_exclusions=stage_exclusions,
                         custom_stages=custom_stages,
                         companies=companies,
                         user=current_user)
@app.route('/supervisor-projects')
@login_required
def supervisor_projects():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    cur = mysql.connection.cursor(DictCursor)
    cur.execute(
        """
        SELECT wlr.w_id, wlr.a_id, wlr.b_name, wlr.due_date,
               wlr.state, wlr.scope, wlr.value, wlr.company, wlr.department,
               wlr.person_name, wlr.status, wlr.result
        FROM win_lost_results wlr
        ORDER BY wlr.w_id DESC
        """
    )
    projects = cur.fetchall()
    cur.close()
    return render_template('supervisor_projects.html', projects=projects, user=current_user)
@app.route('/supervisor/assign-stage', methods=['POST'])
@login_required
def supervisor_assign_stage():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    
    g_id = request.form.get('g_id')
    new_stage = request.form.get('stage')
    person_name = request.form.get('person_name', '').strip()
    department = request.form.get('department', '').strip()
    
    if not g_id or not new_stage:
        flash('Missing required fields', 'error')
        return redirect(url_for('supervisor_dashboard'))
    
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Update go_bids state
        cur.execute("UPDATE go_bids SET state=%s WHERE g_id=%s", (new_stage, g_id))
        
        # Update or create bid_assign entry
        cur.execute("SELECT a_id FROM bid_assign WHERE g_id=%s", (g_id,))
        row = cur.fetchone()
        
        if row:
            # Update existing assignment
            cur.execute("""
                UPDATE bid_assign 
                SET state=%s, depart=%s, person_name=%s, status='assigned' 
                WHERE g_id=%s
            """, (new_stage, department, person_name, g_id))
        else:
            # Create new assignment
            cur.execute("""
                INSERT INTO bid_assign (g_id, b_name, due_date, state, scope, type, company, depart, person_name, status)
                SELECT g_id, b_name, due_date, %s, scope, type, company, %s, %s, 'assigned'
                FROM go_bids WHERE g_id=%s
            """, (new_stage, department, person_name, g_id))
        
        # Log the action
        cur.execute("SELECT b_name FROM go_bids WHERE g_id=%s", (g_id,))
        bid_name = cur.fetchone()['b_name']
        log_action = f"Supervisor '{current_user.email}' assigned bid '{bid_name}' (ID: {g_id}) to {new_stage} stage"
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)", (log_action, current_user.id))
        
        mysql.connection.commit()
        flash(f'Bid "{bid_name}" assigned to {new_stage} stage successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Assignment failed: {str(e)}', 'error')
    finally:
        cur.close()
    
    return redirect(url_for('supervisor_dashboard'))

@app.route('/supervisor/update-company', methods=['POST'])
@login_required
def supervisor_update_company():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    g_id = request.form.get('g_id')
    company_id = request.form.get('company_id')
    if not g_id:
        flash('Missing bid id', 'error')
        return redirect(url_for('supervisor_dashboard'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        company_name = None
        if company_id:
            cur.execute("SELECT name FROM companies WHERE id=%s", (company_id,))
            row = cur.fetchone()
            if row:
                company_name = row['name']
        if not company_name:
            # allow free text fallback
            company_name = (request.form.get('company_name') or '').strip()
        cur.execute("UPDATE go_bids SET company=%s WHERE g_id=%s", (company_name, g_id))
        mysql.connection.commit()
        flash('Company updated.', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Failed to update company: {e}', 'error')
    finally:
        cur.close()
    return redirect(url_for('supervisor_dashboard'))

@app.route('/uploads/<path:subpath>')
@login_required
def serve_uploads(subpath):
    base_dir = os.path.join(os.getcwd(), 'uploads')
    return send_from_directory(base_dir, subpath)

# [REMOVED duplicate /profiling handler consolidated above]

@app.route('/team/stage/add', methods=['POST'])
@login_required
def team_add_stage():
    """Allow team members to add a custom stage for a bid without supervisor role.
    The stage is stored in bid_custom_stages and un-excluded if previously excluded.
    """
    g_id = request.form.get('g_id')
    stage = (request.form.get('new_stage') or '').strip().lower()
    if not g_id or not stage:
        return jsonify({'ok': False, 'error': 'Missing parameters'}), 400
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Ensure required tables exist
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_stage_exclusions (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_bid_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_custom_stages (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_custom_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
        cur.execute("INSERT IGNORE INTO bid_custom_stages (g_id, stage) VALUES (%s, %s)", (g_id, stage))
        cur.execute("DELETE FROM bid_stage_exclusions WHERE g_id=%s AND stage=%s", (g_id, stage))
        mysql.connection.commit()
        return jsonify({'ok': True})
    except Exception as e:
        mysql.connection.rollback()
        return jsonify({'ok': False, 'error': str(e)}), 500
    finally:
        cur.close()

@app.route('/team/stage/delete', methods=['POST'])
@login_required
def team_delete_stage():
    """Allow team members to delete a stage that comes after the current stage."""
    g_id = request.form.get('g_id')
    stage = (request.form.get('stage') or '').strip().lower()
    if not g_id or not stage:
        return jsonify({'ok': False, 'error': 'Missing parameters'}), 400
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Ensure required tables exist
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_stage_exclusions (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_bid_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_custom_stages (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_custom_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
        # Determine current stage and dynamic stage ordering
        cur.execute("SELECT state FROM go_bids WHERE g_id=%s", (g_id,))
        row = cur.fetchone()
        current_stage = (row.get('state') or 'analyzer').strip().lower() if row else 'analyzer'

        cur.execute("SELECT stage FROM bid_stage_exclusions WHERE g_id=%s", (g_id,))
        excluded = { (r['stage'] or '').strip().lower() for r in cur.fetchall() }
        cur.execute("SELECT stage FROM bid_custom_stages WHERE g_id=%s", (g_id,))
        customs = [ (r['stage'] or '').strip().lower() for r in cur.fetchall() ]
        default_stages = ['analyzer', 'business', 'design', 'operations', 'engineer']
        stages = [s for s in default_stages if s not in excluded] + [s for s in customs if s not in excluded and s not in default_stages]

        if stage not in stages:
            return jsonify({'ok': False, 'error': 'Stage not found for this bid.'}), 400
        try:
            curr_idx = stages.index(current_stage)
        except ValueError:
            curr_idx = 0
        target_idx = stages.index(stage)
        if target_idx <= curr_idx:
            return jsonify({'ok': False, 'error': 'Cannot delete current or previous stages.'}), 400

        cur.execute("INSERT IGNORE INTO bid_stage_exclusions (g_id, stage) VALUES (%s, %s)", (g_id, stage))
        cur.execute("DELETE FROM bid_custom_stages WHERE g_id=%s AND stage=%s", (g_id, stage))
        mysql.connection.commit()
        return jsonify({'ok': True})
    except Exception as e:
        mysql.connection.rollback()
        return jsonify({'ok': False, 'error': str(e)}), 500
    finally:
        cur.close()

@app.route('/supervisor/delete-stage', methods=['POST'])
@login_required
def supervisor_delete_stage():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    g_id = request.form.get('g_id')
    deleted_stage = request.form.get('stage')
    if not g_id:
        flash('Missing bid id', 'error')
        return redirect(url_for('supervisor_dashboard'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Remove explicit assignment for this bid; keep current go_bids.state unchanged
        cur.execute("DELETE FROM bid_assign WHERE g_id=%s", (g_id,))
        # Persist exclusion so UI can hide this stage next time
        if deleted_stage:
            cur.execute(
                "INSERT IGNORE INTO bid_stage_exclusions (g_id, stage) VALUES (%s, %s)",
                (g_id, deleted_stage)
            )
        # Log action
        cur.execute("SELECT b_name FROM go_bids WHERE g_id=%s", (g_id,))
        row = cur.fetchone()
        bid_name = row['b_name'] if row else g_id
        stage_info = f" (stage: {deleted_stage})" if deleted_stage else ""
        log_action = f"Supervisor '{current_user.email}' deleted stage assignment for bid '{bid_name}' (ID: {g_id}){stage_info}"
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)", (log_action, current_user.id))
        mysql.connection.commit()
        flash('Stage assignment deleted.', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Could not delete stage: {e}', 'error')
    finally:
        cur.close()
    return redirect(url_for('supervisor_dashboard'))
@app.route('/supervisor/delete-bid', methods=['POST'])
@login_required
def supervisor_delete_bid():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    g_id = request.form.get('g_id')
    if not g_id:
        flash('Missing bid id', 'error')
        return redirect(url_for('supervisor_dashboard'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Fetch name for logs
        cur.execute("SELECT b_name FROM go_bids WHERE g_id=%s", (g_id,))
        row = cur.fetchone()
        bid_name = (row or {}).get('b_name', g_id)

        # Delete related records to avoid foreign key constraint errors
        cur.execute("DELETE FROM bid_checklists WHERE g_id=%s", (g_id,))
        cur.execute("DELETE FROM bid_stage_exclusions WHERE g_id=%s", (g_id,))
        cur.execute("DELETE FROM bid_custom_stages WHERE g_id=%s", (g_id,))
        
        # Delete from bid_assign and cascading tables
        cur.execute("SELECT a_id FROM bid_assign WHERE g_id=%s", (g_id,))
        assign_ids = [r['a_id'] for r in cur.fetchall()]
        
        for a_id in assign_ids:
            cur.execute("SELECT w_id FROM win_lost_results WHERE a_id=%s", (a_id,))
            w_ids = [r['w_id'] for r in cur.fetchall()]
            
            for w_id in w_ids:
                cur.execute("SELECT won_id FROM won_bids_result WHERE w_id=%s", (w_id,))
                won_ids = [r['won_id'] for r in cur.fetchall()]
                
                for won_id in won_ids:
                    cur.execute("DELETE FROM work_progress_status WHERE won_id=%s", (won_id,))
                
                cur.execute("DELETE FROM won_bids_result WHERE w_id=%s", (w_id,))
            
            cur.execute("DELETE FROM win_lost_results WHERE a_id=%s", (a_id,))
        
        cur.execute("DELETE FROM bid_assign WHERE g_id=%s", (g_id,))
        cur.execute("DELETE FROM go_bids WHERE g_id=%s", (g_id,))

        # Log
        log_action = f"Supervisor '{current_user.email}' deleted bid '{bid_name}' (ID: {g_id})"
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)", (log_action, current_user.id))
        mysql.connection.commit()
        flash('Bid deleted successfully.', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Failed to delete bid: {e}', 'error')
    finally:
        cur.close()
    return redirect(url_for('supervisor_dashboard'))

@app.route('/supervisor/set-stage', methods=['POST'])
@login_required
def supervisor_set_stage():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    g_id = request.form.get('g_id')
    new_stage = (request.form.get('new_stage') or '').strip().lower()
    if not g_id or not new_stage:
        flash('Missing bid id or stage.', 'error')
        return redirect(url_for('supervisor_dashboard'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Update go_bids current stage only
        cur.execute("UPDATE go_bids SET state=%s WHERE g_id=%s", (new_stage, g_id))
        # Log
        cur.execute("SELECT b_name FROM go_bids WHERE g_id=%s", (g_id,))
        bid_name = (cur.fetchone() or {}).get('b_name', g_id)
        log_action = f"Supervisor '{current_user.email}' set stage to '{new_stage}' for bid '{bid_name}' (ID: {g_id})"
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)", (log_action, current_user.id))
        mysql.connection.commit()
        flash('Stage updated.', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Could not update stage: {e}', 'error')
    finally:
        cur.close()
    return redirect(url_for('supervisor_dashboard'))

@app.route('/supervisor/add-stage', methods=['POST'])
@login_required
def supervisor_add_stage():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    g_id = request.form.get('g_id')
    stage = (request.form.get('new_stage') or '').strip().lower()
    if not g_id or not stage:
        flash('Bid id and stage name are required.', 'error')
        return redirect(url_for('supervisor_dashboard'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        cur.execute("INSERT IGNORE INTO bid_custom_stages (g_id, stage) VALUES (%s, %s)", (g_id, stage))
        # If stage was previously excluded, un-exclude it
        cur.execute("DELETE FROM bid_stage_exclusions WHERE g_id=%s AND stage=%s", (g_id, stage))
        cur.execute("SELECT b_name FROM go_bids WHERE g_id=%s", (g_id,))
        bid_name = (cur.fetchone() or {}).get('b_name', g_id)
        log_action = f"Supervisor '{current_user.email}' added custom stage '{stage}' for bid '{bid_name}' (ID: {g_id})"
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)", (log_action, current_user.id))
        mysql.connection.commit()
        flash('Stage added.', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Could not add stage: {e}', 'error')
    finally:
        cur.close()
    return redirect(url_for('supervisor_dashboard'))

@app.route('/supervisor/remove-stage', methods=['POST'])
@login_required
def supervisor_remove_stage():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    g_id = request.form.get('g_id')
    stage = (request.form.get('stage') or '').strip().lower()
    if not g_id or not stage:
        return {"ok": False, "error": "Missing parameters"}, 400
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Fetch current stage and dynamic stages for this bid
        cur.execute("SELECT state FROM go_bids WHERE g_id=%s", (g_id,))
        row = cur.fetchone()
        current_stage = (row.get('state') or 'analyzer').strip().lower() if row else 'analyzer'

        # Build dynamic stages order (defaults - exclusions + customs)
        cur.execute("SELECT stage FROM bid_stage_exclusions WHERE g_id=%s", (g_id,))
        excluded = { (r['stage'] or '').strip().lower() for r in cur.fetchall() }
        cur.execute("SELECT stage FROM bid_custom_stages WHERE g_id=%s", (g_id,))
        customs = [ (r['stage'] or '').strip().lower() for r in cur.fetchall() ]
        default_stages = ['analyzer', 'business', 'design', 'operations', 'engineer']
        stages = [s for s in default_stages if s not in excluded] + [s for s in customs if s not in excluded and s not in default_stages]

        if stage not in stages:
            return {"ok": False, "error": "Stage not found for this bid."}, 400
        try:
            curr_idx = stages.index(current_stage)
        except ValueError:
            curr_idx = 0
        target_idx = stages.index(stage)
        if target_idx <= curr_idx:
            return {"ok": False, "error": "Cannot delete current or previous stages."}, 400

        # Mark as excluded and remove from custom if present
        cur.execute("INSERT IGNORE INTO bid_stage_exclusions (g_id, stage) VALUES (%s, %s)", (g_id, stage))
        cur.execute("DELETE FROM bid_custom_stages WHERE g_id=%s AND stage=%s", (g_id, stage))
        mysql.connection.commit()
        return {"ok": True, "message": "Stage deleted."}
    except Exception as e:
        mysql.connection.rollback()
        return {"ok": False, "error": str(e)}, 500
    finally:
        cur.close()
@app.route('/master-dashboard')
@login_required
def master_dashboard():
    if not current_user.is_admin:
        return "Access Denied", 403
    
    # Auto-sync incoming GO bids so dashboards always reflect latest data
    try:
        sync_go_bids()
    except Exception:
        pass

    cur = mysql.connection.cursor(DictCursor)
    
    # Company access filter based on user email domain
    email = (current_user.email or '').lower()
    company_filter_name = None
    if email.endswith('@ikioledlighting.com'):
        company_filter_name = 'Ikio'
    elif email.endswith('@metcoengineering.com'):
        company_filter_name = 'Metco'
    elif email.endswith('@sunsprintengineering.com'):
        company_filter_name = 'Sunsprint'

    company_sql_clause = ""
    company_params = tuple()
    if company_filter_name:
        # Accept aliases like 'METCO Engineering' by matching on keyword
        pattern = None
        name_lower = company_filter_name.lower()
        if name_lower == 'metco':
            pattern = '%metco%'
        elif name_lower == 'ikio':
            pattern = '%ikio%'
        elif name_lower == 'sunsprint':
            pattern = '%sunsprint%'
        if pattern:
            company_sql_clause = " AND LOWER(COALESCE(gb.company,'')) LIKE %s"
            company_params = (pattern,)

    # Get all bids from go_bids
    cur.execute(
        """
        SELECT gb.g_id AS id,
               gb.b_name AS name,
               COALESCE(gb.state, 'analyzer') AS current_stage,
               '' AS user_email,
               '' AS user_role
        FROM go_bids gb
        """ + company_sql_clause, company_params
    )
    all_bids = cur.fetchall()
    # Attach nested user object for template compatibility (bid.user.email, bid.user.role)
    for bid in all_bids:
        bid['user'] = {
            'email': bid.get('user_email'),
            'role': bid.get('user_role')
        }
    
    # Get all companies
    cur.execute("SELECT * FROM companies")
    companies = cur.fetchall()

    # Get top 5 projects from all companies (most urgent due dates) with company name.
    # If the projects table is missing/broken, gracefully fall back to an empty list.
    try:
        cur.execute("""
            SELECT p.*, c.name AS company_name
            FROM projects p
            JOIN companies c ON p.company_id = c.id
            ORDER BY p.due_date ASC
            LIMIT 5
        """)
        top_projects = cur.fetchall()
    except Exception:
        top_projects = []
    # Attach nested company object for template compatibility (project.company.name)
    for project in top_projects:
        project['company'] = {'name': project.get('company_name')}
    has_projects = len(top_projects) > 0

    # Get all tasks. If the tasks table is missing/broken, fall back to an empty list.
    try:
        cur.execute("SELECT * FROM tasks")
        all_tasks = cur.fetchall()
    except Exception:
        all_tasks = []

    # Compute real-time stats for cards across the three companies only using go_bids
    cur.execute("SELECT name FROM companies WHERE name IN ('Ikio','Metco','Sunsprint')")
    target_company_names = [row['name'] for row in cur.fetchall()]
    if not target_company_names:
        target_company_names = ['__none__']
    in_clause = ','.join(['%s'] * len(target_company_names))
    # Also fetch ids for the projects metrics below
    cur.execute("SELECT id FROM companies WHERE name IN ('Ikio','Metco','Sunsprint')")
    target_company_ids = [row['id'] for row in cur.fetchall()] or [-1]
    in_clause_ids = ','.join(['%s'] * len(target_company_ids))

    cur.execute(f"SELECT COUNT(*) AS total_bids FROM go_bids WHERE company IN ({in_clause})", target_company_names)
    total_bids = cur.fetchone()['total_bids']

    cur.execute(f"""
        SELECT COUNT(*) AS live_bids
        FROM go_bids
        WHERE COALESCE(state,'analyzer') IN ('business','design','operations','engineer')
          AND company IN ({in_clause})
    """, target_company_names)
    live_bids = cur.fetchone()['live_bids']

    cur.execute(f"SELECT COUNT(*) AS bids_won FROM go_bids WHERE decision='WON' AND company IN ({in_clause})", target_company_names)
    bids_won = cur.fetchone()['bids_won']

    # Build per-company stage counts from go_bids for the timeline trackers
    # Per-state counts
    cur.execute(
        """
        SELECT COALESCE(company, '') AS company,
               LOWER(COALESCE(state, '')) AS state,
               COUNT(*) AS total
        FROM go_bids
        GROUP BY COALESCE(company, ''), LOWER(COALESCE(state, ''))
        """
    )
    rows = cur.fetchall()
    go_counts = {}
    for row in rows:
        company_key = row['company'] or ''
        if company_key not in go_counts:
            go_counts[company_key] = {'analyzer': 0, 'business': 0, 'design': 0, 'operations': 0, 'engineer': 0, 'handover': 0}
        # Map known downstream stages directly; others will be reflected in analyzer via totals below
        if row['state'] in ('business', 'design', 'operations', 'engineer', 'handover'):
            go_counts[company_key][row['state']] = row['total']

    # Total per company should power the Analyzer start point
    cur.execute(
        """
        SELECT COALESCE(company, '') AS company,
               COUNT(*) AS total
        FROM go_bids
        GROUP BY COALESCE(company, '')
        """
    )
    totals = cur.fetchall()
    for t in totals:
        company_key = t['company'] or ''
        if company_key not in go_counts:
            go_counts[company_key] = {'analyzer': 0, 'business': 0, 'design': 0, 'operations': 0, 'engineer': 0, 'handover': 0}
        go_counts[company_key]['analyzer'] = t['total']
    company_names = [c['name'] for c in companies]

    # Build go_bids by company for dropdown summaries, enriched with status from downstream tables
    cur.execute(
        """
        SELECT gb.g_id,
               gb.b_name,
        
               gb.due_date,
               gb.state,
               gb.type,
               gb.company,
               gb.decision,
               gb.summary,
               gb.scoring AS progress,
               wps.pr_completion_status AS work_status,
               wps.dept_bde,
               wps.dept_m_d,
               wps.dept_op,
               wps.dept_site,
               wbr.closure_status AS project_status,
               wbr.work_progress_status AS work_progress_status,
               wlr.result AS wl_result
        FROM go_bids gb
        LEFT JOIN bid_assign ba ON ba.g_id = gb.g_id
        LEFT JOIN win_lost_results wlr ON wlr.a_id = ba.a_id
        LEFT JOIN won_bids_result wbr ON wbr.w_id = wlr.w_id
        LEFT JOIN work_progress_status wps ON wps.won_id = wbr.won_id
        ORDER BY gb.company ASC, gb.due_date ASC
        """
    )
    go_rows = cur.fetchall()
    go_bids_by_company = {}
    for r in go_rows:
        cname = (r.get('company') or '').strip()
        go_bids_by_company.setdefault(cname, []).append(r)
    go_company_names = [name for name, rows in go_bids_by_company.items() if name and len(rows) > 0]

    # assigned_bids handles assignments now; no assign_go_bids refresh

    # Build a flat list of go_bids projects with normalized tracker progress
    def _normalize_stage(raw_state: str) -> str:
        s = (raw_state or '').strip().lower()
        mapping = {
            'analyzer': 'analyzer',
            'business': 'business',
            'business dev': 'business',
            'bdm': 'business',
            'design': 'design',
            'operations': 'operations',
            'operation': 'operations',
            'engineer': 'engineer',
            'site_manager': 'engineer',
            'site manager': 'engineer',
            'handover': 'handover',
            'won': 'handover'
        }
        # Treat unknown values like submitted/under_review/pending as analyzer
        return mapping.get(s, 'analyzer')

    stage_to_percent = {
        'analyzer': 0,
        'business': 20,
        'design': 40,
        'operations': 60,
        'engineer': 80,
        'handover': 100,
    }

    # Dynamic stages (respect supervisor-managed deletions and additions)
    cur.execute("SELECT g_id, stage FROM bid_stage_exclusions")
    ex_rows = cur.fetchall()
    exclusions_by_bid = {}
    for row in ex_rows:
        exclusions_by_bid.setdefault(row['g_id'], set()).add((row['stage'] or '').strip().lower())
    cur.execute("SELECT g_id, stage FROM bid_custom_stages")
    cs_rows = cur.fetchall()
    customs_by_bid = {}
    for row in cs_rows:
        customs_by_bid.setdefault(row['g_id'], []).append((row['stage'] or '').strip().lower())

    default_stages = ['analyzer', 'business', 'design', 'operations', 'engineer', 'handover']

    def _parse_percent(value: str) -> int:
        try:
            if value is None:
                return 0
            s = str(value)
            # extract digits first
            import re
            m = re.search(r"(\d{1,3})", s)
            if m:
                pct = max(0, min(100, int(m.group(1))))
                return pct
            s = s.strip().lower()
            if s in ('done','completed','closed','handover','100%'):
                return 100
            if s in ('in_progress','ongoing'):
                return 50
            if s in ('pending','todo','not_started'):
                return 0
            return 0
        except Exception:
            return 0

    # Helper: compute stage progress map from bid_checklists per bid
    def _compute_stage_map_for_bid(g_id: int) -> dict:
        try:
            cur2 = mysql.connection.cursor(DictCursor)
            cur2.execute("""
                SELECT bc.progress_pct, bc.status, COALESCE(bc.stage, u.role) AS stage_source
                FROM bid_checklists bc
                LEFT JOIN users u ON bc.created_by = u.id
                WHERE bc.g_id = %s
            """, (g_id,))
            rows = cur2.fetchall()
            cur2.close()
            role_to_stage = {
                'business dev': 'business',
                'design': 'design',
                'operations': 'operations',
                'site manager': 'engineer',
                'engineer': 'engineer'
            }
            buckets: dict[str, list] = {'analyzer': [], 'business': [], 'design': [], 'operations': [], 'engineer': [], 'handover': []}
            for r in rows:
                source = (r.get('stage_source') or '').strip().lower()
                stage = role_to_stage.get(source, source if source in buckets else None)
                if not stage:
                    continue
                pct = r.get('progress_pct')
                if pct is None:
                    s = (r.get('status') or '').strip().lower()
                    pct = 100 if s == 'completed' else 50 if s == 'in_progress' else 0
                try:
                    pct = max(0, min(100, int(pct)))
                except Exception:
                    pct = 0
                buckets[stage].append(pct)
            def avg(lst):
                if not lst:
                    return 0
                return int(round(sum(lst) / len(lst)))
            return {
                'analyzer': avg(buckets['analyzer']),
                'business': avg(buckets['business']),
                'design': avg(buckets['design']),
                'operations': avg(buckets['operations']),
                'engineer': avg(buckets['engineer']),
                'handover': avg(buckets['handover']),
            }
        except Exception:
            return {'analyzer': 0, 'business': 0, 'design': 0, 'operations': 0, 'engineer': 0, 'handover': 0}

    go_projects = []
    for r in go_rows:
        stage_key = _normalize_stage(r.get('state'))
        excluded = exclusions_by_bid.get(r.get('g_id'), set())
        custom = [s for s in customs_by_bid.get(r.get('g_id'), []) if s not in excluded]
        stages = [s for s in default_stages if s not in excluded] + [s for s in custom if s not in default_stages]
        if stages and stage_key in stages:
            idx = stages.index(stage_key)
            stage_progress = int(round((idx / (max(1, len(stages) - 1))) * 100))
        else:
            stage_progress = stage_to_percent.get(stage_key, 0)
        progress_pct = r.get('progress') if r.get('progress') is not None else r.get('scoring')
        try:
            progress_pct = max(0, min(100, int(progress_pct))) if progress_pct is not None else None
        except Exception:
            progress_pct = None
        
        # Calculate dynamic progress and status texts
        item_pct = stage_progress_pct(stage_key)
        proj_status, work_status = status_texts(stage_key)
        
        # Prefer task-based progress; fallback to work_progress_status when tasks are absent
        task_stage_map = _compute_stage_map_for_bid(r.get('g_id'))
        if any(v > 0 for v in task_stage_map.values()):
            stage_progress_map = {
                'analyzer': 0,
                'business': task_stage_map.get('business', 0),
                'design': task_stage_map.get('design', 0),
                'operations': task_stage_map.get('operations', 0),
                'engineer': task_stage_map.get('engineer', 0),
                'handover': 0,
            }
        else:
            stage_progress_map = {
                'analyzer': 0,
                'business': _parse_percent(r.get('dept_bde')),
                'design': _parse_percent(r.get('dept_m_d')),
                'operations': _parse_percent(r.get('dept_op')),
                'engineer': _parse_percent(r.get('dept_site')),
                'handover': 100 if (r.get('project_status') or '').strip().lower() in ('closed','completed','handover','done') else 0,
            }

        # If still no per-stage progress was found (all zeros),
        # estimate based on current stage so the timeline tracker reflects movement.
        if not any(v > 0 for v in stage_progress_map.values()):
            try:
                ordered = ['analyzer','business','design','operations','engineer','handover']
                cur_idx = ordered.index(stage_key) if stage_key in ordered else 0
                estimated = {}
                for i, s in enumerate(ordered):
                    if i < cur_idx:
                        estimated[s] = 100
                    elif i == cur_idx:
                        # show partial progress for the current stage
                        estimated[s] = 40
                    else:
                        estimated[s] = 0
                # keep analyzer and handover conventions
                estimated['analyzer'] = estimated.get('analyzer', 0)
                estimated['handover'] = estimated.get('handover', 0)
                stage_progress_map.update(estimated)
            except Exception:
                pass

        go_projects.append({
            'g_id': r.get('g_id'),
            'b_name': r.get('b_name'),
            'company': r.get('company'),
            'state': r.get('state'),
            'stage_key': stage_key,
            'stage_progress': stage_progress,
            'stages': stages,
            'current_stage': stage_key,
         
            'due_date': r.get('due_date'),
            'type': r.get('type'),
            'decision': r.get('decision') or r.get('wl_result'),
            'project_status': proj_status,  # New dynamic project status
            'work_status': work_status,     # New dynamic work status
            'summary': r.get('summary'),
            'progress_pct': progress_pct,
            'work_progress_pct': item_pct,  # New dynamic progress
            'stage_progress_map': stage_progress_map,
        })

    # Deduplicate entries for the same bid under Metco aliases (e.g., 'Metco' vs 'Metco Engineering')
    # Rule: if name and due_date match and company contains 'metco', keep only one record.
    try:
        def _normalize_company_alias(name: str) -> str:
            s = (name or '').strip().lower()
            if 'metco' in s:
                return 'metco'
            return s
        def _normalize_name(n: str) -> str:
            return ' '.join((n or '').strip().lower().split())
        def _normalize_date(dv) -> str:
            # Return YYYY-MM-DD string or empty if unknown
            from datetime import datetime, date
            if dv is None or dv == '':
                return ''
            if isinstance(dv, datetime):
                return dv.date().isoformat()
            if isinstance(dv, date):
                return dv.isoformat()
            s = str(dv).strip()
            # Try common formats
            for fmt in ('%Y-%m-%d', '%Y-%m-%d %H:%M:%S', '%d-%m-%Y', '%m/%d/%Y', '%Y/%m/%d'):
                try:
                    return datetime.strptime(s[:19], fmt).date().isoformat()
                except Exception:
                    continue
            # Fallback ISO parse
            try:
                return datetime.fromisoformat(s[:19]).date().isoformat()
            except Exception:
                return s  # leave raw for stable keying
        seen_keys = set()
        deduped = []
        for it in go_projects:
            alias = _normalize_company_alias(it.get('company'))
            if alias == 'metco':
                key = (alias, _normalize_name(it.get('b_name')), _normalize_date(it.get('due_date')))
                if key in seen_keys:
                    continue
                seen_keys.add(key)
            deduped.append(it)
        go_projects = deduped
    except Exception:
        # On any failure, fall back to original list
        pass

    # Total projects across target companies (fallback to all projects if none linked)
    cur.execute(f"SELECT COUNT(*) AS projects_linked FROM projects WHERE company_id IN ({in_clause_ids})", target_company_ids)
    projects_linked = cur.fetchone()['projects_linked']
    if projects_linked > 0:
        cur.execute(f"SELECT COUNT(*) AS projects_total FROM projects WHERE company_id IN ({in_clause_ids})", target_company_ids)
        projects_total = cur.fetchone()['projects_total']
    else:
        cur.execute("SELECT COUNT(*) AS projects_total FROM projects")
        projects_total = cur.fetchone()['projects_total']

    # Project-based dashboard cards (map to existing template keys)
    # total_bids -> Total Projects
    # live_bids -> Projects Completed
    # bids_won -> Projects (in progress)
    # projects_completed -> Projects (on hold)
    if projects_linked > 0:
        cur.execute(f"SELECT COUNT(*) AS projects_completed FROM projects WHERE status='completed' AND company_id IN ({in_clause_ids})", target_company_ids)
        p_completed = cur.fetchone()['projects_completed']
        cur.execute(f"SELECT COUNT(*) AS projects_in_progress FROM projects WHERE status IN ('active','in_progress') AND company_id IN ({in_clause_ids})", target_company_ids)
        p_in_progress = cur.fetchone()['projects_in_progress']
        cur.execute(f"SELECT COUNT(*) AS projects_on_hold FROM projects WHERE status='on_hold' AND company_id IN ({in_clause_ids})", target_company_ids)
        p_on_hold = cur.fetchone()['projects_on_hold']
    else:
        cur.execute("SELECT COUNT(*) AS projects_completed FROM projects WHERE status='completed'")
        p_completed = cur.fetchone()['projects_completed']
        cur.execute("SELECT COUNT(*) AS projects_in_progress FROM projects WHERE status IN ('active','in_progress')")
        p_in_progress = cur.fetchone()['projects_in_progress']
        cur.execute("SELECT COUNT(*) AS projects_on_hold FROM projects WHERE status='on_hold'")
        p_on_hold = cur.fetchone()['projects_on_hold']

    # Override dashboard cards using win_lost_results as requested
    try:
        cur.execute("SELECT COUNT(*) AS total_projects_wlr FROM win_lost_results")
        total_projects_wlr = cur.fetchone()['total_projects_wlr']
    except Exception:
        total_projects_wlr = projects_total

    try:
        cur.execute("""
            SELECT COUNT(*) AS in_progress_wlr
            FROM win_lost_results
            WHERE COALESCE(UPPER(result), '') NOT IN ('WON','LOST')
        """)
        in_progress_wlr = cur.fetchone()['in_progress_wlr']
    except Exception:
        in_progress_wlr = p_in_progress

    data = {
        'total_bids': total_projects_wlr,      # Total Projects
        'live_bids': p_completed,               # Projects Completed (unchanged)
        'bids_won': in_progress_wlr,            # Projects (in progress)
        'projects_completed': p_on_hold         # Projects (on hold)
    }

    # Compute bid analyzer stats from bid_incoming table
    cur.execute("SELECT COUNT(*) AS total_bids FROM bid_incoming")
    total_bids_analyzer = cur.fetchone()['total_bids']
    
    cur.execute("SELECT COUNT(*) AS bids_go FROM bid_incoming WHERE decision = 'GO'")
    bids_go_analyzer = cur.fetchone()['bids_go']
    
    cur.execute("SELECT COUNT(*) AS bids_no_go FROM bid_incoming WHERE decision = 'NO-GO'")
    bids_no_go_analyzer = cur.fetchone()['bids_no_go']
    
    cur.execute("SELECT COUNT(*) AS bids_submitted FROM bid_incoming WHERE state IN ('submitted', 'under_review')")
    bids_submitted_analyzer = cur.fetchone()['bids_submitted']
    
    cur.execute("SELECT COUNT(*) AS bids_won FROM bid_incoming WHERE decision = 'WON'")
    bids_won_analyzer = cur.fetchone()['bids_won']
    
    cur.execute("SELECT COUNT(*) AS bids_lost FROM bid_incoming WHERE decision = 'LOST'")
    bids_lost_analyzer = cur.fetchone()['bids_lost']

    bid_stats = {
        'total_bids': total_bids_analyzer,
        'bids_go': bids_go_analyzer,
        'bids_no_go': bids_no_go_analyzer,
        'bids_submitted': bids_submitted_analyzer,
        'bids_won': bids_won_analyzer,
        'bids_lost': bids_lost_analyzer
    }
    
    # Ensure won_bids_result contains all WON decisions from bid_incoming
    # Check the actual column name in bid_incoming table
    try:
        cur.execute("DESCRIBE bid_incoming")
        describe_results = cur.fetchall()
        # Handle both dict and tuple formats
        if describe_results and isinstance(describe_results[0], dict):
            bid_incoming_columns = [row['Field'] for row in describe_results]
        else:
            bid_incoming_columns = [row[0] for row in describe_results]  # First column is Field name
        
        bid_incoming_id_col = 'id' if 'id' in bid_incoming_columns else ('id' if 'id' in bid_incoming_columns else None)
        
        if bid_incoming_id_col:
            cur.execute(
                f"""
                INSERT INTO won_bids_result (w_id)
                SELECT DISTINCT wlr.w_id
                FROM bid_incoming bi
                LEFT JOIN go_bids gb ON gb.id = bi.{bid_incoming_id_col}
                LEFT JOIN bid_assign ba ON ba.g_id = gb.g_id
                LEFT JOIN win_lost_results wlr ON wlr.a_id = ba.a_id
                LEFT JOIN won_bids_result wbr ON wbr.w_id = wlr.w_id
                WHERE UPPER(bi.decision) = 'WON'
                  AND wlr.w_id IS NOT NULL
                  AND wbr.w_id IS NULL
                """
            )
            mysql.connection.commit()
    except Exception as e:
        # If the join fails (e.g., no relationship between tables), skip this step
        # won_bids_result will be populated from other sources
        mysql.connection.rollback()
        pass

    # Build Won Projects (latest 5) timeline data
    # Ensure won_bids_result has rows for any WIN in win_lost_results
    cur.execute(
        """
        INSERT INTO won_bids_result (w_id)
        SELECT wlr.w_id
        FROM win_lost_results wlr
        LEFT JOIN won_bids_result wbr ON wbr.w_id = wlr.w_id
        WHERE UPPER(COALESCE(wlr.result,'')) = 'WON'
          AND wbr.w_id IS NULL
        """
    )
    mysql.connection.commit()

    cur.execute(
        """
        SELECT wbr.*, gb.g_id, gb.b_name, gb.company,
               COALESCE(wps.pr_completion_status, wbr.work_progress_status) AS work_progress_status
        FROM won_bids_result wbr
        LEFT JOIN win_lost_results wlr ON wlr.w_id = wbr.w_id
        LEFT JOIN bid_assign ba ON ba.a_id = wlr.a_id
        LEFT JOIN go_bids gb ON gb.g_id = ba.g_id
        LEFT JOIN work_progress_status wps ON wps.won_id = wbr.won_id
        ORDER BY wbr.w_id DESC
        LIMIT 5
        """
    )
    won_rows = cur.fetchall()

    def _won_stage_key(row: dict) -> str:
        closure = (row.get('closure_status') or '').strip().lower()
        work_prog = (row.get('work_progress_status') or '').strip().lower()
        if closure in ('closed', 'completed', 'handover', 'done'):
            return 'closure'
        if work_prog in ('operations', 'operation', 'ops', 'in_progress'):
            return 'operations'
        if work_prog == 'design':
            return 'design'
        return 'business'

    won_stages_default = ['won', 'business', 'design', 'operations', 'closure']
    won_projects = []
    for r in won_rows:
        stage_key = _won_stage_key(r)
        idx = won_stages_default.index(stage_key) if stage_key in won_stages_default else 0
        stage_progress = int(round((idx / (len(won_stages_default) - 1)) * 100)) if len(won_stages_default) > 1 else 0
        won_projects.append({
            'g_id': r.get('g_id'),
            'b_name': r.get('b_name') or f"Won #{r.get('won_id')}",
            'company': r.get('company'),
            'current_stage': stage_key,
            'stage_progress': stage_progress,
            'stages': won_stages_default,
        })

    # Fallback: include directly WON bids from bid_incoming not present above
    cur.execute("""
        SELECT bi.id, bi.b_name, COALESCE(gb.company,'') AS company
        FROM bid_incoming bi
        LEFT JOIN go_bids gb ON gb.id = bi.id
        WHERE UPPER(bi.decision) = 'WON'
        ORDER BY bi.id DESC
        LIMIT 20
    """)
    bi_won = cur.fetchall() or []
    existing_names = { (p['b_name'] or '').strip() for p in won_projects }
    for r in bi_won:
        name = r.get('b_name')
        if name and name.strip() in existing_names:
            continue
        won_projects.append({
            'g_id': r.get('id'),  # Changed from id to id to match the SELECT query
            'b_name': name,
            'company': r.get('company'),
            'current_stage': 'business',
            'stage_progress': int(round((won_stages_default.index('business')/(len(won_stages_default)-1))*100)),
            'stages': won_stages_default,
        })

    # Split go_projects into top 5 by due_date and the rest
    def _date_key(item):
        try:
            from datetime import datetime, date
            v = item.get('due_date')
            if v is None or v == '':
                return datetime.max
            # If already a datetime/date, convert to datetime
            if isinstance(v, datetime):
                return v
            if isinstance(v, date):
                return datetime.combine(v, datetime.min.time())
            # Try common string formats
            s = str(v).strip()
            for fmt in ('%Y-%m-%d', '%Y-%m-%d %H:%M:%S', '%d-%m-%Y', '%m/%d/%Y', '%Y/%m/%d'):
                try:
                    return datetime.strptime(s[:19], fmt)
                except Exception:
                    pass
            # Fallback: try fromisoformat
            try:
                return datetime.fromisoformat(s[:19])
            except Exception:
                return datetime.max
        except Exception:
            from datetime import datetime
            return datetime.max
    go_projects_sorted = sorted(go_projects, key=_date_key)
    go_projects_top5 = go_projects_sorted[:5]
    go_projects_more = go_projects_sorted[5:]

    # Parallel kickoff: ensure every bid appears in each team dashboard
    for team_key in ['business','design','operations','engineer']:
        ensure_tasks_for_team(team_key)

    cur.close()

    # ---- Company/user analytics from database (employees + users + bid_checklists) ----
    def _company_from_email(email: str) -> str:
        e = (email or '').lower()
        if 'metco' in e:
            return 'METCO Engineering'
        if 'ikio' in e:
            return 'IKIO'
        if 'sunsprint' in e:
            return 'Sunsprint Engineering'
        return 'Other'
    users_company_stats = {}
    members_status_data = {}
    try:
        curu = mysql.connection.cursor(DictCursor)
        # Determine viewer's company from email; if recognized, restrict analytics to it
        viewer_company = _company_from_email(getattr(current_user, 'email', ''))
        restrict_company = viewer_company if viewer_company and viewer_company != 'Other' else None
        # Users by company (from users)
        curu.execute("SELECT id, email FROM users")
        user_rows = curu.fetchall() or []
        company_to_user_ids: dict[str, set] = {}
        for r in user_rows:
            comp = _company_from_email(r.get('email') or '')
            # Respect restriction if set
            if restrict_company and comp != restrict_company:
                continue
            company_to_user_ids.setdefault(comp, set()).add(r.get('id'))
        # Employees by company (augment using employees table when department suggests company)
        try:
            curu.execute("SELECT id, email, department, is_active FROM employees")
            emp_rows = curu.fetchall() or []
            for r in emp_rows:
                comp = (r.get('department') or '').strip()
                if not comp:
                    comp = _company_from_email(r.get('email') or '')
                if comp and (not restrict_company or comp == restrict_company):
                    company_to_user_ids.setdefault(comp, set()).add(r.get('id') or -100000 - (r.get('id') or 0))
        except Exception:
            pass
        # Average completion by company (from bid_checklists created_by)
        curu.execute("SELECT created_by, COALESCE(progress_pct,0) AS pct FROM bid_checklists")
        checklist_rows = curu.fetchall() or []
        company_sum: dict[str, float] = {}
        company_cnt: dict[str, int] = {}
        # Map user id -> email to infer company
        id_to_email = {r.get('id'): r.get('email') for r in user_rows if r.get('id') is not None}
        for row in checklist_rows:
            uid = row.get('created_by')
            comp = _company_from_email(id_to_email.get(uid, ''))
            if restrict_company and comp != restrict_company:
                continue
            company_sum[comp] = company_sum.get(comp, 0.0) + float(row.get('pct') or 0.0)
            company_cnt[comp] = company_cnt.get(comp, 0) + 1
        labels = list(company_to_user_ids.keys())
        labels.sort()
        users_counts = [len(company_to_user_ids.get(c, set())) for c in labels]
        avg_completion = []
        for c in labels:
            if company_cnt.get(c, 0) > 0:
                avg_completion.append(round(company_sum[c] / company_cnt[c], 1))
            else:
                avg_completion.append(0)
        # If restricted, collapse to a single company label to keep charts scoped
        if restrict_company:
            labels = [restrict_company]
            users_counts = [len(company_to_user_ids.get(restrict_company, set()))]
            if company_cnt.get(restrict_company, 0) > 0:
                avg_val = round(company_sum[restrict_company] / company_cnt[restrict_company], 1)
            else:
                avg_val = 0
            avg_completion = [avg_val]
        users_company_stats = {
            'labels': labels,
            'users': users_counts,
            'avg_completion': avg_completion
        }
        # Members status data (bucket by user via bid_checklists)
        curu.execute("""
            SELECT u.name, u.email, bc.progress_pct
            FROM bid_checklists bc
            LEFT JOIN users u ON u.id = bc.created_by
        """)
        rows = curu.fetchall() or []
        per_user = {}
        for r in rows:
            # Apply company restriction using user email
            if restrict_company and _company_from_email(r.get('email') or '') != restrict_company:
                continue
            name = (r.get('name') or '') or (r.get('email') or 'Member')
            pct = r.get('progress_pct')
            try:
                pct = int(pct) if pct is not None else 0
            except Exception:
                pct = 0
            bucket = 'in_progress'
            if pct >= 95:
                bucket = 'completed'
            elif pct <= 0:
                bucket = 'not_started'
            # late cannot be inferred from checklist alone; leave as 0
            if name not in per_user:
                per_user[name] = {'not_started': 0, 'in_progress': 0, 'late': 0, 'completed': 0}
            per_user[name][bucket] += 1
        labels_users = list(per_user.keys())
        labels_users.sort(key=lambda u: sum(per_user[u].values()), reverse=True)
        # limit for card view computed on client; send all
        members_status_data = {
            'labels': labels_users,
            'not_started': [per_user[u]['not_started'] for u in labels_users],
            'in_progress': [per_user[u]['in_progress'] for u in labels_users],
            'late': [per_user[u]['late'] for u in labels_users],
            'completed': [per_user[u]['completed'] for u in labels_users],
        }
        curu.close()
    except Exception:
        users_company_stats = {}
        members_status_data = {}

    # Stage display name mapping
    def get_stage_display_name(stage_key):
        mapping = {
            'analyzer': 'BID Analyzer',
            'business': 'Business Development',
            'design': 'Design & Marketing',
            'operations': 'Operation Team',
            'engineer': 'Engineering Team',
            'handover': 'Submitted'
        }
        return mapping.get(stage_key.lower(), stage_key.title())

    return render_template(
        'master_dashboard.html', 
        page='dashboard',
        title='Master Dashboard',
        data=data,
        bids=all_bids,
        companies=companies,
        projects=top_projects,
        tasks=all_tasks,
        bid_stats=bid_stats,
        go_counts=go_counts,
        has_projects=has_projects,
        company_names=company_names,
        go_company_names=go_company_names,
        go_bids_by_company=go_bids_by_company,
        go_projects_top5=go_projects_top5,
        go_projects_more=go_projects_more,
        won_projects=won_projects,
        users_company_stats=users_company_stats,
        members_status_data=members_status_data,
        get_stage_display_name=get_stage_display_name
    )
@app.route('/company/<company_name>')
@login_required
def company_dashboard(company_name):
    if not current_user.is_admin:
        return "Access Denied", 403
    
    # Map URL names to database names
    company_mapping = {
        'ikio': ['Ikio', 'IKIO'],
        'metco': ['Metco', 'Metco Engineering', 'METCO'],
        'sunsprint': ['Sunsprint', 'Sunsprint Engineering', 'SUNSPRINT']
    }
    
    aliases = company_mapping.get(company_name.lower())
    if not aliases:
        return "Company not found", 404
    if isinstance(aliases, str):
        aliases = [aliases]
        
    cur = mysql.connection.cursor(DictCursor)
    
    # Get company (try aliases, allow case-insensitive match)
    company = None
    for alias in aliases:
        cur.execute("SELECT * FROM companies WHERE LOWER(name) = %s", (alias.lower(),))
        company = cur.fetchone()
        if company:
            break
    # Fallback to LIKE search (handles names such as 'IKIO LED Lighting')
    if not company:
        for alias in aliases:
            pattern = f"%{alias.lower()}%"
            cur.execute("SELECT * FROM companies WHERE LOWER(name) LIKE %s", (pattern,))
            company = cur.fetchone()
            if company:
                break
    # If still not found, continue with a virtual company so the page can render using go_bids
    if not company:
        company = {'id': None, 'name': aliases[0]}
    
    # Get top 5 projects for this company with nested company info
    projects = []
    if company.get('id') is not None:
        cur.execute("""
            SELECT p.*, c.name AS company_name
            FROM projects p
            JOIN companies c ON p.company_id = c.id
            WHERE p.company_id=%s
            ORDER BY p.due_date ASC
            LIMIT 5
        """, (company['id'],))
        projects = cur.fetchall()
        for project in projects:
            project['company'] = {'name': project.get('company_name')}
    
    # Get tasks for this company
    tasks = []
    if company.get('id') is not None:
        cur.execute("""
            SELECT t.* FROM tasks t 
            JOIN projects p ON t.project_id = p.id 
            WHERE p.company_id = %s
        """, (company['id'],))
        tasks = cur.fetchall()
    
    # Build set of aliases for filtering bids (case-insensitive)
    company_aliases = { (company.get('name') or '').lower() }
    company_aliases.update(a.lower() for a in aliases if a)
    # Use LIKE patterns to match variations (e.g., 'ikio led lighting')
    like_patterns = tuple(f"%{a.strip().lower()}%" for a in company_aliases if a)
    where_clause = ' OR '.join(['LOWER(gb.company) LIKE %s'] * len(like_patterns)) or '1=0'

    # Get bids for this company from go_bids by company name (enriched)
    cur.execute(
        f"""
        SELECT gb.*, 
               wps.pr_completion_status AS work_status,
               wbr.closure_status AS project_status,
               wbr.work_progress_status AS work_progress_status,
               wlr.result AS wl_result
        FROM go_bids gb
        LEFT JOIN bid_assign ba ON ba.g_id = gb.g_id
        LEFT JOIN win_lost_results wlr ON wlr.a_id = ba.a_id
        LEFT JOIN won_bids_result wbr ON wbr.w_id = wlr.w_id
        LEFT JOIN work_progress_status wps ON wps.won_id = wbr.won_id
        WHERE {where_clause}
        ORDER BY gb.due_date ASC
        """,
        like_patterns
    )
    bids = cur.fetchall()
    
    # Fetch stage exclusions and custom stages for each bid
    cur.execute("SELECT g_id, stage FROM bid_stage_exclusions")
    rows = cur.fetchall()
    stage_exclusions = {}
    for r in rows:
        stage_exclusions.setdefault(r['g_id'], []).append(r['stage'])

    cur.execute("SELECT g_id, stage FROM bid_custom_stages")
    rows = cur.fetchall()
    custom_stages = {}
    for r in rows:
        custom_stages.setdefault(r['g_id'], []).append(r['stage'])

    # Prepare go_projects for this company similar to master dashboard
    def _normalize_stage_company(raw_state: str) -> str:
        s = (raw_state or '').strip().lower()
        mapping = {
            'analyzer': 'analyzer',
            'business': 'business', 'business dev': 'business', 'bdm': 'business',
            'design': 'design',
            'operations': 'operations', 'operation': 'operations',
            'engineer': 'engineer', 'site_manager': 'engineer', 'site manager': 'engineer',
            'handover': 'handover', 'won': 'handover'
        }
        return mapping.get(s, 'analyzer')

    stage_to_percent_company = {
        'analyzer': 0,
        'business': 20,
        'design': 40,
        'operations': 60,
        'engineer': 80,
        'handover': 100,
    }

    # Default stages list
    default_stages = ['analyzer', 'business', 'design', 'operations', 'engineer']
    
    go_projects = []
    for r in bids:
        bid_id = r.get('g_id')
        
        # Get dynamic stages for this bid
        excluded = set(stage_exclusions.get(bid_id, []))
        customs = [s for s in custom_stages.get(bid_id, []) if s not in excluded]
        stages = [s for s in default_stages if s not in excluded] + [s for s in customs if s not in default_stages]
        
        # If no stages remain, fall back to default
        if not stages:
            stages = default_stages.copy()
        
        stage_key = _normalize_stage_company(r.get('state'))
        
        # Calculate stage_progress based on dynamic stages
        if stage_key in stages:
            idx = stages.index(stage_key)
            stage_progress = int(round((idx / max(1, len(stages) - 1)) * 100))
        else:
            stage_progress = 0
        
        progress_pct = r.get('scoring')
        try:
            progress_pct = max(0, min(100, int(progress_pct))) if progress_pct is not None else None
        except Exception:
            progress_pct = None
        
        # Calculate dynamic progress and status texts
        item_pct = stage_progress_pct(stage_key)
        proj_status, work_status = status_texts(stage_key)
        
        go_projects.append({
            'g_id': bid_id,
            'b_name': r.get('b_name'),
            'company': r.get('company'),
            'state': r.get('state'),
            'stage_key': stage_key,
            'stage_progress': stage_progress,
            'stages': stages,  # Add dynamic stages list
         
            'due_date': r.get('due_date'),
            'type': r.get('type'),
            'decision': r.get('decision') or r.get('wl_result'),
            'project_status': proj_status,  # New dynamic project status
            'work_status': work_status,     # New dynamic work status
            'summary': r.get('summary'),
            'progress_pct': progress_pct,
            'work_progress_pct': item_pct,  # New dynamic progress
        })
    
    # Compute company metrics
    cur.execute("""
        SELECT 
            COUNT(*) AS total_projects,
            SUM(CASE WHEN COALESCE(status,'') = 'completed' OR COALESCE(progress,0) >= 100 THEN 1 ELSE 0 END) AS completed_projects,
            AVG(COALESCE(progress,0)) AS avg_progress
        FROM projects
        WHERE company_id = %s
    """, (company['id'],))
    proj_stats = cur.fetchone() or {}
    total_projects = int(proj_stats.get('total_projects') or 0)
    completed_projects = int(proj_stats.get('completed_projects') or 0)
    active_projects = max(0, total_projects - completed_projects)
    avg_progress_pct = int(round(proj_stats.get('avg_progress') or 0))

    company_metrics = {
        'active_projects': active_projects,
        'total_tasks': len(tasks or []),
        'completed_projects': completed_projects,
        'avg_progress_pct': avg_progress_pct,
    }

    # Load multi-assignees per bid
    bid_ids = []
    for b in (bids or []):
        bid_key = b.get('id') or b.get('g_id')
        if bid_key is not None:
            bid_ids.append(bid_key)
    assigned_map = {}
    try:
        if bid_ids:
            # Ensure mapping table exists
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS bid_assignment_members (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    g_id INT NOT NULL,
                    employee_id INT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE KEY uniq_g_emp (g_id, employee_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
                """
            )
            placeholders = ','.join(['%s'] * len(bid_ids))
            cur.execute(
                f"""
                SELECT bam.g_id, e.id AS employee_id, e.name, e.email
                FROM bid_assignment_members bam
                JOIN employees e ON e.id = bam.employee_id
                WHERE bam.g_id IN ({placeholders}) AND e.is_active = TRUE
                ORDER BY e.name
                """,
                tuple(bid_ids)
            )
            rows = cur.fetchall()
            for r in rows:
                assigned_map.setdefault(r['g_id'], []).append({'id': r['employee_id'], 'name': r['name'], 'email': r['email']})
        # Attach to bid rows
        for b in bids:
            bid_key = b.get('id') or b.get('g_id')
            b['assigned_members'] = assigned_map.get(bid_key, [])
    except Exception:
        for b in bids:
            b['assigned_members'] = []

    cur.close()
    
    # Stage display name mapping
    def get_stage_display_name(stage_key):
        mapping = {
            'analyzer': 'BID Analyzer',
            'business': 'Business Development',
            'design': 'Design & Marketing',
            'operations': 'Operation Team',
            'engineer': 'Engineering Team',
            'handover': 'Submitted'
        }
        return mapping.get(stage_key.lower(), stage_key.title())
    
    # --- Bid Analyzer mini-metrics for this company (filtered by comp_name) ---
    try:
        cur2 = mysql.connection.cursor(DictCursor)
        # Build LIKE patterns to match variations of the company name
        like_patterns_bi = tuple(f"%{a.strip().lower()}%" for a in company_aliases if a)
        where_bi = ' OR '.join(['LOWER(comp_name) LIKE %s'] * len(like_patterns_bi)) or '1=0'
        # Totals
        cur2.execute(f"SELECT COUNT(*) AS total_bids FROM bid_incoming WHERE {where_bi}", like_patterns_bi)
        total_bids_ci = (cur2.fetchone() or {}).get('total_bids', 0)
        # GO / NO-GO / Submitted / WON / LOST
        cur2.execute(f"SELECT COUNT(*) AS c FROM bid_incoming WHERE decision='GO' AND ({where_bi})", like_patterns_bi)
        bids_go_ci = (cur2.fetchone() or {}).get('c', 0)
        cur2.execute(f"SELECT COUNT(*) AS c FROM bid_incoming WHERE decision='NO-GO' AND ({where_bi})", like_patterns_bi)
        bids_no_go_ci = (cur2.fetchone() or {}).get('c', 0)
        cur2.execute(f"""SELECT COUNT(*) AS c 
                        FROM bid_incoming 
                        WHERE LOWER(COALESCE(state,'')) IN ('submitted','under_review') AND ({where_bi})""", like_patterns_bi)
        bids_submitted_ci = (cur2.fetchone() or {}).get('c', 0)
        cur2.execute(f"SELECT COUNT(*) AS c FROM bid_incoming WHERE decision='WON' AND ({where_bi})", like_patterns_bi)
        bids_won_ci = (cur2.fetchone() or {}).get('c', 0)
        cur2.execute(f"SELECT COUNT(*) AS c FROM bid_incoming WHERE decision='LOST' AND ({where_bi})", like_patterns_bi)
        bids_lost_ci = (cur2.fetchone() or {}).get('c', 0)
        cur2.close()
        bid_stats_company = {
            'total_bids': int(total_bids_ci or 0),
            'bids_go': int(bids_go_ci or 0),
            'bids_no_go': int(bids_no_go_ci or 0),
            'bids_submitted': int(bids_submitted_ci or 0),
            'bids_won': int(bids_won_ci or 0),
            'bids_lost': int(bids_lost_ci or 0),
        }
    except Exception:
        bid_stats_company = {
            'total_bids': 0, 'bids_go': 0, 'bids_no_go': 0, 'bids_submitted': 0, 'bids_won': 0, 'bids_lost': 0
        }
    
    return render_template(
        'company_dashboard.html', 
        company=company,
        projects=projects,
        tasks=tasks,
        bids=bids,
        go_projects=go_projects,
        company_metrics=company_metrics,
        get_stage_display_name=get_stage_display_name,
        current_company=company_name.lower(),
        bid_stats=bid_stats_company
    )

@app.route('/bid-analyzer')
@login_required
def bid_analyzer():
    if not current_user.is_admin:
        return "Access Denied", 403
    
    # Ensure latest GO bids are mirrored into go_bids whenever analyzer loads
    try:
        sync_go_bids()
    except Exception:
        pass
    
    cur = mysql.connection.cursor(DictCursor)
    
    # Get bid_incoming data for the table
    cur.execute("SELECT * FROM bid_incoming ORDER BY id DESC")
    bid_incoming_data = cur.fetchall()
    
    # Calculate bid stats from bid_incoming table
    cur.execute("SELECT COUNT(*) AS total_bids FROM bid_incoming")
    total_bids = cur.fetchone()['total_bids']
    
    cur.execute("SELECT COUNT(*) AS bids_go FROM bid_incoming WHERE decision = 'GO'")
    bids_go = cur.fetchone()['bids_go']
    
    cur.execute("SELECT COUNT(*) AS bids_no_go FROM bid_incoming WHERE decision = 'NO-GO'")
    bids_no_go = cur.fetchone()['bids_no_go']
    
    cur.execute("SELECT COUNT(*) AS bids_submitted FROM bid_incoming WHERE state IN ('submitted', 'under_review')")
    bids_submitted = cur.fetchone()['bids_submitted']
    
    cur.execute("SELECT COUNT(*) AS bids_won FROM bid_incoming WHERE decision = 'WON'")
    bids_won = cur.fetchone()['bids_won']
    
    cur.execute("SELECT COUNT(*) AS bids_lost FROM bid_incoming WHERE decision = 'LOST'")
    bids_lost = cur.fetchone()['bids_lost']
    
    cur.close()

    return render_template('bid_analyzer_landing.html', bid_cards={
        'total_bids': total_bids,
        'bids_go': bids_go,
        'bids_no_go': bids_no_go,
        'bids_submitted': bids_submitted,
        'bids_won': bids_won,
        'bids_lost': bids_lost
    }, bid_incoming_data=bid_incoming_data)
@app.route('/databases')
@login_required
def databases():
    if not current_user.is_admin:
        return redirect(url_for('role_dashboard'))
    
    # Get all companies and their projects
    cur = mysql.connection.cursor(DictCursor)
    cur.execute("SELECT * FROM companies")
    companies = cur.fetchall()
    company_data = {}
    
    for company in companies:
        cur.execute("SELECT * FROM projects WHERE company_id=%s", (company['id'],))
        projects = cur.fetchall()
        company_data[company['name']] = {
            'company': company,
            'projects': projects
        }
    
    cur.close()
    return render_template('databases.html', company_data=company_data)
@app.route('/databases/create_project', methods=['POST'])
@login_required
def create_project():
    if not current_user.is_admin:
        return redirect(url_for('role_dashboard'))
    
    try:
        company_name = request.form.get('company_name')
        project_name = request.form.get('project_name')
        due_date = request.form.get('due_date')
        status = request.form.get('status', 'active')
        progress = int(request.form.get('progress', 0))
        
        cur = mysql.connection.cursor()
        
        # Find or create company
        cur.execute("SELECT id FROM companies WHERE name=%s", (company_name,))
        company = cur.fetchone()
        if not company:
            cur.execute("INSERT INTO companies (name, description) VALUES (%s, %s)", 
                       (company_name, f'{company_name} Projects'))
            company_id = cur.lastrowid
        else:
            company_id = company['id']
        
        # Create project
        cur.execute("""
            INSERT INTO projects (name, company_id, start_date, due_date, status, progress) 
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (project_name, company_id, datetime.utcnow(), 
              datetime.strptime(due_date, '%Y-%m-%d'), status, progress))
        
        mysql.connection.commit()
        cur.close()
        
        flash(f'Project "{project_name}" created successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error creating project: {str(e)}', 'error')
    
    return redirect(url_for('databases'))

@app.route('/databases/update_project/<int:project_id>', methods=['POST'])
@login_required
def update_project(project_id):
    if not current_user.is_admin:
        return redirect(url_for('role_dashboard'))
    
    try:
        cur = mysql.connection.cursor()
        
        # Check if project exists
        cur.execute("SELECT * FROM projects WHERE id=%s", (project_id,))
        project = cur.fetchone()
        if not project:
            cur.close()
            return "Project not found", 404
        
        # Update project
        project_name = request.form.get('project_name', project['name'])
        due_date = datetime.strptime(request.form.get('due_date'), '%Y-%m-%d')
        status = request.form.get('status', project['status'])
        progress = int(request.form.get('progress', project['progress']))
        
        cur.execute("""
            UPDATE projects 
            SET name=%s, due_date=%s, status=%s, progress=%s 
            WHERE id=%s
        """, (project_name, due_date, status, progress, project_id))
        
        mysql.connection.commit()
        cur.close()
        
        flash(f'Project "{project_name}" updated successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error updating project: {str(e)}', 'error')
    
    return redirect(url_for('databases'))

@app.route('/databases/delete_project/<int:project_id>')
@login_required
def delete_project(project_id):
    if not current_user.is_admin:
        return redirect(url_for('role_dashboard'))
    
    try:
        cur = mysql.connection.cursor()
        
        # Get project name before deletion
        cur.execute("SELECT name FROM projects WHERE id=%s", (project_id,))
        project = cur.fetchone()
        if not project:
            cur.close()
            return "Project not found", 404
        
        project_name = project['name']
        
        # Delete project
        cur.execute("DELETE FROM projects WHERE id=%s", (project_id,))
        mysql.connection.commit()
        cur.close()
        
        flash(f'Project "{project_name}" deleted successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error deleting project: {str(e)}', 'error')
    
    return redirect(url_for('databases'))

@app.route('/databases/import_excel', methods=['POST'])
@login_required
def import_excel():
    if not current_user.is_admin:
        return redirect(url_for('role_dashboard'))
    
    try:
        if 'excel_file' not in request.files:
            flash('No file selected', 'error')
            return redirect(url_for('databases'))
        
        file = request.files['excel_file']
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(url_for('databases'))
        
        if file and file.filename.endswith('.xlsx'):
            import pandas as pd
            import io
            
            # Read Excel file
            df = pd.read_excel(io.BytesIO(file.read()))
            
            # Expected columns: Company, Project Name, Due Date, Revenue, Status, Progress
            required_columns = ['Company', 'Project Name', 'Due Date', 'Status', 'Progress']
            
            if not all(col in df.columns for col in required_columns):
                flash(f'Excel file must contain columns: {", ".join(required_columns)}', 'error')
                return redirect(url_for('databases'))
            
            projects_created = 0
            cur = mysql.connection.cursor()
            
            for _, row in df.iterrows():
                try:
                    # Find or create company
                    company_name = str(row['Company']).strip()
                    cur.execute("SELECT id FROM companies WHERE name=%s", (company_name,))
                    company = cur.fetchone()
                    if not company:
                        cur.execute("INSERT INTO companies (name, description) VALUES (%s, %s)", 
                                   (company_name, f'{company_name} Projects'))
                        company_id = cur.lastrowid
                    else:
                        company_id = company['id']
                    
                    # Create project
                    cur.execute("""
                        INSERT INTO projects (name, company_id, start_date, due_date, status, progress) 
                        VALUES (%s, %s, %s, %s, %s, %s, %s)
                    """, (str(row['Project Name']).strip(), company_id, datetime.utcnow(),
                          pd.to_datetime(row['Due Date']), 
                          str(row['Status']).strip() if pd.notna(row['Status']) else 'active',
                          int(row['Progress']) if pd.notna(row['Progress']) else 0))
                    
                    projects_created += 1
                    
                except Exception as e:
                    print(f"Error processing row: {e}")
                    continue
            
            mysql.connection.commit()
            cur.close()
            flash(f'Successfully imported {projects_created} projects from Excel file!', 'success')
            
        else:
            flash('Please upload a valid Excel file (.xlsx)', 'error')
            
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error importing Excel file: {str(e)}', 'error')
    
    return redirect(url_for('databases'))

@app.route('/databases/drop_all')
@login_required
def drop_all_databases():
    if not current_user.is_admin:
        return redirect(url_for('role_dashboard'))
    
    try:
        cur = mysql.connection.cursor()
        
        # Delete all projects
        cur.execute("DELETE FROM projects")
        
        # Delete all companies
        cur.execute("DELETE FROM companies")
        
        # Delete all tasks
        cur.execute("DELETE FROM tasks")
        
        mysql.connection.commit()
        cur.close()
        flash('All databases have been cleared successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error clearing databases: {str(e)}', 'error')
    
    return redirect(url_for('databases'))

# --- Bid Incoming CRUD Routes ---
@app.route('/bid-analyzer/create', methods=['POST'])
@login_required
def create_bid_incoming():
    if not current_user.is_admin:
        return "Access Denied", 403
    
    try:
        cur = mysql.connection.cursor()
        
        b_name = request.form.get('b_name', '').strip()
      
        due_date = request.form.get('due_date', '').strip()
        state = request.form.get('state', '').strip()
        scope = request.form.get('scope', '').strip()
        type_val = request.form.get('type', '').strip()
        scoring = int(request.form.get('scoring', 0)) if request.form.get('scoring') else None
        comp_name = request.form.get('comp_name', '').strip()
        decision = request.form.get('decision', '').strip()
        summary = request.form.get('summary', '').strip()
        
        if not b_name or not due_date:
            return 'Bid name, incoming date and due date are required', 400
        
        cur.execute("""
            INSERT INTO bid_incoming (b_name, due_date, state, scope, type, scoring, comp_name, decision, summary)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (b_name, due_date, state, scope, type_val, scoring, comp_name, decision, summary))
        
        bid_id = cur.lastrowid
        mysql.connection.commit()
        
        # Automatically sync GO decisions from bid_incoming to go_bids
        if (decision or '').upper() == 'GO':
            cur2 = mysql.connection.cursor(DictCursor)
            cur2.execute("SELECT g_id FROM go_bids WHERE id=%s", (bid_id,))
            row = cur2.fetchone()
            args = (b_name, due_date, state if state else 'business', scope, type_val, scoring, comp_name, decision, summary)
            if row:
                cur2.execute("""UPDATE go_bids SET b_name=%s,due_date=%s,state=%s,scope=%s,
                                type=%s,scoring=%s,company=%s,decision=%s,summary=%s WHERE id=%s""", (*args, bid_id))
            else:
                cur2.execute("""INSERT INTO go_bids (id,b_name,due_date,state,scope,type,scoring,company,decision,summary)
                                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                             (bid_id, *args))
            mysql.connection.commit()
            cur2.close()
            log_write('sync', f"Auto-synced GO bid '{b_name}' (id={bid_id}) from bid_incoming to go_bids")
        
        cur.close()
        log_write('create', f"table=bid_incoming, id={bid_id}")
        flash(f'Bid "{b_name}" created successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error creating bid: {str(e)}', 'error')
    
    return redirect(url_for('bid_analyzer'))

@app.route('/bid-analyzer/update/<int:bid_id>', methods=['POST'])
@login_required
def update_bid_incoming(bid_id):
    if not current_user.is_admin:
        return "Access Denied", 403
    
    try:
        cur = mysql.connection.cursor()
        
        # Check if bid exists
        cur.execute("SELECT * FROM bid_incoming WHERE id=%s", (bid_id,))
        bid = cur.fetchone()
        if not bid:
            cur.close()
            return "Bid not found", 404
        
        b_name = request.form.get('b_name', bid['b_name']).strip()
        due_date = request.form.get('due_date', str(bid['due_date'])).strip()
        state = request.form.get('state', bid['state'] or '').strip()
        scope = request.form.get('scope', bid['scope'] or '').strip()
        type_val = request.form.get('type', bid['type'] or '').strip()
        scoring = int(request.form.get('scoring', 0)) if request.form.get('scoring') else bid['scoring']
        comp_name = request.form.get('comp_name', bid['comp_name'] or '').strip()
        decision = request.form.get('decision', bid['decision'] or '').strip()
        summary = request.form.get('summary', bid['summary'] or '').strip()
        
        cur.execute("""
            UPDATE bid_incoming 
            SET b_name=%s, due_date=%s, state=%s, scope=%s, type=%s, scoring=%s, comp_name=%s, decision=%s, summary=%s
            WHERE id=%s
        """, (b_name, due_date, state, scope, type_val, scoring, comp_name, decision, summary, bid_id))
        
        mysql.connection.commit()
        
        # Automatically sync GO decisions from bid_incoming to go_bids
        if (decision or '').upper() == 'GO':
            cur2 = mysql.connection.cursor(DictCursor)
            cur2.execute("SELECT g_id FROM go_bids WHERE id=%s", (bid_id,))
            row = cur2.fetchone()
            args = (b_name, due_date, state if state else 'business', scope, type_val, scoring, comp_name, decision, summary)
            if row:
                cur2.execute("""UPDATE go_bids SET b_name=%s,due_date=%s,state=%s,scope=%s,
                                type=%s,scoring=%s,company=%s,decision=%s,summary=%s WHERE id=%s""", (*args, bid_id))
            else:
                cur2.execute("""INSERT INTO go_bids (id,b_name,due_date,state,scope,type,scoring,company,decision,summary)
                                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                             (bid_id, *args))
            mysql.connection.commit()
            cur2.close()
            log_write('sync', f"Auto-synced GO bid '{b_name}' (id={bid_id}) from bid_incoming to go_bids")
        else:
            # If decision changed from GO to something else, remove from go_bids
            cur2 = mysql.connection.cursor(DictCursor)
            cur2.execute("DELETE FROM go_bids WHERE id=%s", (bid_id,))
            mysql.connection.commit()
            cur2.close()
            log_write('sync', f"Removed bid '{b_name}' (id={bid_id}) from go_bids (decision changed from GO)")
        
        cur.close()
        log_write('update', f"table=bid_incoming, id={bid_id}")
        flash(f'Bid "{b_name}" updated successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error updating bid: {str(e)}', 'error')
    
    return redirect(url_for('bid_analyzer'))

@app.route('/bid-analyzer/delete/<int:bid_id>')
@login_required
def delete_bid_incoming(bid_id):
    if not current_user.is_admin:
        return "Access Denied", 403
    
    try:
        cur = mysql.connection.cursor()
        
        # Get bid name before deletion
        cur.execute("SELECT b_name FROM bid_incoming WHERE id=%s", (bid_id,))
        bid = cur.fetchone()
        if not bid:
            cur.close()
            return "Bid not found", 404
        
        bid_name = bid['b_name']
        
        # Delete bid
        cur.execute("DELETE FROM bid_incoming WHERE id=%s", (bid_id,))
        mysql.connection.commit()
        cur.close()
        
        log_write('delete', f"table=bid_incoming, id={bid_id}")
        flash(f'Bid "{bid_name}" deleted successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error deleting bid: {str(e)}', 'error')
    
    return redirect(url_for('bid_analyzer'))

@app.route('/database-management')
@login_required
def database_management():
    if not current_user.is_admin:
        return "Access Denied", 403
    
    cur = mysql.connection.cursor(DictCursor)
    
    # Get all available tables
    cur.execute("SHOW TABLES")
    all_tables = [list(table.values())[0] for table in cur.fetchall()]
    
    # Filter to show only specific tables in proper sequence
    allowed_tables = ['bid_incoming', 'go_bids', 'bid_assign', 'win_lost_results', 'won_bids_result', 'work_progress_status']
    # Maintain the sequence by preserving the order in allowed_tables
    tables = [t for t in allowed_tables if t in all_tables]
    
    # Create mapping for display names (capital letters with spaces)
    table_display_names = {
        'bid_incoming': 'BID INCOMING',
        'go_bids': 'GO BIDS',
        'bid_assign': 'BID ASSIGN',
        'win_lost_results': 'WIN LOST RESULTS',
        'won_bids_result': 'WON BIDS RESULT',
        'work_progress_status': 'WORK PROGRESS STATUS'
    }
    
    # Function to format column names (e.g., 'id' -> 'BID ID', 'comp_name' -> 'COMP NAME')
    def format_column_name(col_name):
        # Replace underscores with spaces and convert to uppercase
        formatted = col_name.replace('_', ' ').upper()
        return formatted
    
    # Get selected table (default to first table)
    selected_table = request.args.get('table', tables[0] if tables else '')
    search_query = request.args.get('search', '')
    decision_filter = request.args.get('decision', '').strip()
    
    # Get table data
    table_data = []
    table_columns = []
    companies = []
    
    if selected_table:
        # Get table structure
        cur.execute(f"DESCRIBE `{selected_table}`")
        table_columns = cur.fetchall()
        
        # Auto-sync GO bids into go_bids table
        if selected_table == 'go_bids':
            cur.execute("""
                INSERT INTO go_bids (id, b_name, due_date, state, scope, type, scoring, company, decision, summary)
                SELECT bi.id, bi.b_name, bi.due_date, bi.state, bi.scope, bi.type, bi.scoring, bi.comp_name, bi.decision, bi.summary
                FROM bid_incoming bi
                LEFT JOIN go_bids gb ON gb.id = bi.id
                WHERE UPPER(bi.decision) = 'GO' AND gb.id IS NULL
            """)
            mysql.connection.commit()
            # After syncing GO bids, refresh revenue-based assignments into assigned_bids
            try:
                assign_bids_by_revenue()
            except Exception:
                pass

        # Note: bid_assign is now populated ONLY via explicit Assign action from go_bids

        # Auto-sync bid_assign into win_lost_results (one row per assignment)
        if selected_table == 'win_lost_results':
            cur.execute("""
                INSERT INTO win_lost_results (a_id, b_name, due_date, state, scope, value, company, department, person_name, status, result)
                SELECT ba.a_id, ba.b_name, ba.due_date, ba.state, ba.scope, ba.value, ba.company, ba.depart, ba.person_name, ba.status, NULL
                FROM bid_assign ba
                LEFT JOIN win_lost_results wlr ON wlr.a_id = ba.a_id
                WHERE wlr.a_id IS NULL
            """)
            mysql.connection.commit()

        # Auto-sync win_lost_results into won_bids_result (link by w_id)
        if selected_table == 'won_bids_result':
            cur.execute("""
                INSERT INTO won_bids_result (w_id)
                SELECT wlr.w_id
                FROM win_lost_results wlr
                LEFT JOIN won_bids_result wbr ON wbr.w_id = wlr.w_id
                WHERE wbr.w_id IS NULL
            """)
            mysql.connection.commit()

        # Auto-sync won_bids_result into work_progress_status (link by won_id)
        if selected_table == 'work_progress_status':
            cur.execute("""
                INSERT INTO work_progress_status (won_id, company, b_name, dept_bde, dept_m_d, dept_op, dept_site, pr_completion_status)
                SELECT wbr.won_id, COALESCE(gb.company, ''), COALESCE(gb.b_name, ''), '', '', '', '', NULL
                FROM won_bids_result wbr
                LEFT JOIN win_lost_results wlr ON wlr.w_id = wbr.w_id
                LEFT JOIN bid_assign ba ON ba.a_id = wlr.a_id
                LEFT JOIN go_bids gb ON gb.g_id = ba.g_id
                LEFT JOIN work_progress_status wps ON wps.won_id = wbr.won_id
                WHERE wps.won_id IS NULL
            """)
            mysql.connection.commit()

        # Build dynamic WHERE for search and decision filter
        where_parts = []
        params = []

        # Text search across varchar/text columns
        if search_query:
            search_conditions = []
            for col in table_columns:
                if col['Type'].startswith(('varchar', 'text', 'char')):
                    search_conditions.append(f"`{col['Field']}` LIKE %s")
            if search_conditions:
                search_pattern = f"%{search_query}%"
                params.extend([search_pattern] * len(search_conditions))
                where_parts.append("(" + " OR ".join(search_conditions) + ")")

        # Decision filter for bid_incoming or go_bids
        if selected_table in ('bid_incoming', 'go_bids') and decision_filter:
            where_parts.append("UPPER(`decision`) = %s")
            params.append(decision_filter.upper())

        where_sql = (" WHERE " + " AND ".join(where_parts)) if where_parts else ""
        cur.execute(f"SELECT * FROM `{selected_table}`{where_sql}", params)
        
        table_data = cur.fetchall()
        if selected_table == 'bid_incoming':
            cur.execute("SELECT id, name FROM companies ORDER BY name")
            companies = cur.fetchall()
    
    cur.close()
    
    return render_template('database_management.html', 
                         tables=tables, 
                         selected_table=selected_table,
                         table_columns=table_columns,
                         table_data=table_data,
                         search_query=search_query,
                         decision_filter=decision_filter,
                         companies=companies,
                         table_display_names=table_display_names,
                         format_column_name=format_column_name)

@app.route('/admin/refresh-assign-go')
@login_required
def admin_refresh_assign_go():
    if not current_user.is_admin:
        return "Access Denied", 403
    try:
        assign_bids_by_revenue()
        flash('assigned_bids refreshed from go_bids.', 'success')
    except Exception as e:
        flash(f'Failed to refresh: {e}', 'error')
    return redirect(url_for('database_management', table='assigned_bids'))
@app.route('/database-management/create', methods=['POST'])
@login_required
def dbm_create():
    if not current_user.is_admin:
        return "Access Denied", 403
    table = request.form.get('table')
    if not table:
        return redirect(url_for('database_management'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Get columns and build insert
        cur.execute(f"DESCRIBE `{table}`")
        cols = cur.fetchall()
        fields = []
        values = []
        params = []
        for c in cols:
            if c['Key'] == 'PRI':
                continue
            field = c['Field']
            fields.append(f"`{field}`")
            values.append('%s')
            params.append(request.form.get(field))
        sql = f"INSERT INTO `{table}` ({', '.join(fields)}) VALUES ({', '.join(values)})"
        cur.execute(sql, params)
        mysql.connection.commit()

        # If a bid was created in bid_incoming, offer auto-assign path via flash info
        if table == 'bid_incoming':
            flash('Bid created in Bid Incoming. Assign it to a company from the actions column.', 'success')
    finally:
        cur.close()
    return redirect(url_for('database_management', table=table))

@app.route('/database-management/update/<int:row_id>', methods=['POST'])
@login_required
def dbm_update(row_id):
    if not current_user.is_admin:
        return "Access Denied", 403
    table = request.form.get('table')
    if not table:
        return redirect(url_for('database_management'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        cur.execute(f"DESCRIBE `{table}`")
        cols = cur.fetchall()
        pk = cols[0]['Field']
        sets = []
        params = []
        for c in cols:
            if c['Key'] == 'PRI':
                pk = c['Field']
                continue
            field = c['Field']
            sets.append(f"`{field}`=%s")
            params.append(request.form.get(field))
        params.append(row_id)
        sql = f"UPDATE `{table}` SET {', '.join(sets)} WHERE `{pk}`=%s"
        cur.execute(sql, params)
        mysql.connection.commit()
    finally:
        cur.close()
    return redirect(url_for('database_management', table=table))

@app.route('/database-management/assign', methods=['POST'])
@login_required
def dbm_assign():
    if not current_user.is_admin:
        return "Access Denied", 403
    bid_incoming_id = request.form.get('bid_id')
    company_id = request.form.get('company_id')
    if not bid_incoming_id or not company_id:
        return redirect(url_for('database_management', table='bid_incoming'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Read bid_incoming
        cur.execute("SELECT * FROM bid_incoming WHERE id=%s", (bid_incoming_id,))
        inc = cur.fetchone()
        if not inc:
            cur.close()
            return redirect(url_for('database_management', table='bid_incoming'))
        # Create bid in main bids with analyzer -> business based on decision
        stage = 'analyzer'
        if (inc.get('decision') or '').upper() == 'GO':
            stage = 'business'
        cur.execute("INSERT INTO bids (name, current_stage, user_id, company_id) VALUES (%s,%s,%s,%s)",
                    (inc.get('b_name'), stage, current_user.id, company_id))
        new_bid_id = cur.lastrowid
        # Timeline
        cur.execute("INSERT INTO bid_timeline (bid_id, event, details) VALUES (%s,%s,%s)",
                    (new_bid_id, 'assigned', f"Assigned to company_id={company_id} from bid_incoming {bid_incoming_id}"))
        # Optional: keep in bid_incoming or delete; here we keep it
        mysql.connection.commit()
        # Emit via socket if needed later; for now just flash
        flash('Bid assigned to company successfully.', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Assignment failed: {str(e)}', 'error')
    finally:
        cur.close()
    return redirect(url_for('database_management', table='bid_incoming'))

@app.route('/database-management/assign-go', methods=['POST'])
@login_required
def dbm_assign_go():
    if not current_user.is_admin:
        return "Access Denied", 403
    g_id = request.form.get('g_id')
    depart = request.form.get('depart', '').strip()
    person_name = request.form.get('person_name', '').strip()
    email_to = request.form.get('email', '').strip()
    if not g_id or not depart or not person_name:
        flash('Please provide department and person name', 'error')
        return redirect(url_for('database_management', table='go_bids'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Ensure bid_assign exists for this go bid; insert if missing else update
        cur.execute("SELECT a_id FROM bid_assign WHERE g_id=%s", (g_id,))
        row = cur.fetchone()
        if not row:
            # Pull from go_bids to seed
            cur.execute("SELECT * FROM go_bids WHERE g_id=%s", (g_id,))
            gb = cur.fetchone()
            if not gb:
                flash('GO bid not found', 'error')
                cur.close()
                return redirect(url_for('database_management', table='go_bids'))
            
            # Map department to stage for go_bids.state update
            dept_to_stage = {
                'business dev': 'business',
                'business': 'business',
                'design': 'design',
                'operations': 'operations',
                'site manager': 'engineer',
                'engineer': 'engineer'
            }
            new_stage = dept_to_stage.get(depart.lower(), depart.lower())
            
            # Update go_bids.state to match assigned department
            cur.execute("UPDATE go_bids SET state=%s WHERE g_id=%s", (new_stage, g_id))
            
            # Check if assignee_email column exists before inserting
            try:
                cur.execute("""
                    INSERT INTO bid_assign (g_id, b_name, due_date, state, scope, type, company, depart, person_name, assignee_email, status, value)
                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s)
                """, (
                    gb['g_id'], gb['b_name'], gb['due_date'], new_stage, gb['scope'], gb['type'], gb['company'], depart, person_name, email_to,
                    gb.get('scoring', 0)
                ))
            except Exception as e:
                if "Unknown column 'assignee_email'" in str(e):
                    # Fallback insert without assignee_email
                    cur.execute("""
                        INSERT INTO bid_assign (g_id, b_name, due_date, state, scope, type, company, depart, person_name, status, value)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s)
                    """, (
                        gb['g_id'], gb['b_name'], gb['due_date'], new_stage, gb['scope'], gb['type'], gb['company'], depart, person_name,
                        gb.get('scoring', 0)
                    ))
                else:
                    raise e
        else:
            # Map department to stage for go_bids.state update
            dept_to_stage = {
                'business dev': 'business',
                'business': 'business',
                'design': 'design',
                'operations': 'operations',
                'site manager': 'engineer',
                'engineer': 'engineer'
            }
            new_stage = dept_to_stage.get(depart.lower(), depart.lower())
            
            # Update go_bids.state to match assigned department
            cur.execute("UPDATE go_bids SET state=%s WHERE g_id=%s", (new_stage, g_id))
            
            # Check if assignee_email column exists before updating
            try:
                cur.execute("UPDATE bid_assign SET depart=%s, person_name=%s, assignee_email=%s, state=%s, status='assigned' WHERE g_id=%s", 
                           (depart, person_name, email_to, new_stage, g_id))
            except Exception as e:
                if "Unknown column 'assignee_email'" in str(e):
                    # Fallback update without assignee_email
                    cur.execute("UPDATE bid_assign SET depart=%s, person_name=%s, state=%s, status='assigned' WHERE g_id=%s", 
                               (depart, person_name, new_stage, g_id))
                else:
                    raise e
        # Log the assignment action
        cur.execute("SELECT b_name FROM go_bids WHERE g_id=%s", (g_id,))
        bid_name = cur.fetchone()['b_name']
        log_action = f"Admin '{current_user.email}' assigned bid '{bid_name}' (ID: {g_id}) to {depart} department - {person_name} ({email_to})"
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)", (log_action, current_user.id))
        
        mysql.connection.commit()
        
        # Emit Socket.IO update for real-time master dashboard sync
        socketio.emit('master_update', {
            'bid': {
                'id': g_id,
                'name': bid_name,
                'current_stage': new_stage,
                'assigned_to': person_name,
                'department': depart
            },
            'log': {
                'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'action': log_action,
                'user_email': current_user.email,
                'user_role': 'admin'
            },
            'assignment': True
        })
        
        flash('Bid assigned to person successfully.', 'success')
        # Send email notification if provided
        if email_to:
            try:
                send_assignment_email(email_to, gb['b_name'] if 'gb' in locals() else None, depart, person_name, gb['company'] if 'gb' in locals() else None)
            except Exception as _e:
                # Non-fatal
                flash('Assignment email could not be sent.', 'error')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Assignment failed: {str(e)}', 'error')
    finally:
        cur.close()
    return redirect(url_for('database_management', table='go_bids'))

# --- Email helper (adapted from mail_send.py) ---
SMTP_PORT = 587
SMTP_SERVER = "smtp.gmail.com"
EMAIL_FROM = "manuj@metcoengineering.com"
EMAIL_PASSWORD = "ksec iqja bmdg hrcv"

def send_assignment_email(email_to: str, bid_name: str, depart: str, person_name: str, company: str):
    subject = f"Bid Assignment: {bid_name or ''}"
    body = f"You have been assigned to bid '{bid_name or ''}' for company '{company or ''}'.\nDepartment: {depart}\nAssignee: {person_name}\n\nPlease log in to the ESCO suite to proceed."
    message = f"Subject: {subject}\n\n{body}"
    context = ssl.create_default_context()
    server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
    try:
        server.starttls(context=context)
        server.login(EMAIL_FROM, EMAIL_PASSWORD)
        server.sendmail(EMAIL_FROM, email_to, message)
    finally:
        try:
            server.quit()
        except Exception:
            pass
# --- Employee Management Routes ---
@app.route('/team/<team>/employees')
@login_required
def team_employees(team):
    """Manage employees for a specific team"""
    if current_user.is_admin:
        return redirect(url_for('master_dashboard'))
    
    # Map team names to stages
    team_to_stage = {
        'business': 'business',
        'design': 'design', 
        'operations': 'operations',
        'engineer': 'engineer'
    }
    
    if team not in team_to_stage:
        return "Invalid team", 404
    
    cur = mysql.connection.cursor(DictCursor)
    
    # Get employees for this team
    cur.execute("""
        SELECT e.*, u.email as team_lead_email
        FROM employees e
        LEFT JOIN users u ON e.team_lead_id = u.id
        WHERE e.department = %s AND e.is_active = TRUE
        ORDER BY e.name
    """, (team,))
    employees = cur.fetchall()
    
    # Get team leads. Map team -> acceptable user roles, and for non-admins
    # restrict to the current manager so the dropdown shows their email only.
    roles_for_team = {
        'business': ('business dev', 'business', 'bdm'),
        'design': ('design',),
        'operations': ('operations',),
        'engineer': ('site manager', 'engineer'),
    }
    acceptable_roles = roles_for_team.get(team, (team,))
    if getattr(current_user, 'is_admin', False):
        placeholders = ','.join(['%s'] * len(acceptable_roles))
        cur.execute(f"SELECT * FROM users WHERE role IN ({placeholders})", acceptable_roles)
        team_leads = cur.fetchall()
    else:
        # Only the current manager on their own dashboard
        cur.execute("SELECT * FROM users WHERE id=%s", (current_user.id,))
        team_leads = cur.fetchall()
    
    cur.close()
    
    return render_template('team_employees.html', 
                         team=team, 
                         employees=employees, 
                         team_leads=team_leads,
                         user=current_user)

@app.route('/team/<team>/employees/create', methods=['POST'])
@login_required
def create_employee(team):
    """Create a new employee for the team"""
    try:
        name = request.form.get('name', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '').strip()
        team_lead_id = request.form.get('team_lead_id') or (current_user.id if not getattr(current_user, 'is_admin', False) else None)
        
        if not name or not email or not password:
            flash('Name, email, and password are required', 'error')
            return redirect(url_for('team_employees', team=team))
        
        cur = mysql.connection.cursor()
        cur.execute("""
            INSERT INTO employees (name, email, password, department, team_lead_id) 
            VALUES (%s, %s, %s, %s, %s)
        """, (name, email, password, team, team_lead_id if team_lead_id else None))
        
        mysql.connection.commit()
        cur.close()
        
        log_write('create_employee', f"Created employee {name} for {team} team")
        flash(f'Employee "{name}" created successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error creating employee: {str(e)}', 'error')
    
    return redirect(url_for('team_employees', team=team))

@app.route('/employee/<int:employee_id>/dashboard')
def employee_dashboard(employee_id):
    """Employee-specific dashboard showing assigned tasks"""
    # Check if employee is logged in and matches the requested employee
    if 'employee_id' not in session or session['employee_id'] != employee_id:
        return "Access denied. Please login first.", 403
    
    cur = mysql.connection.cursor(DictCursor)
    
    # Get employee info
    cur.execute("SELECT * FROM employees WHERE id = %s", (employee_id,))
    employee = cur.fetchone()
    
    if not employee:
        cur.close()
        return "Employee not found", 404
    
    # Get assigned tasks for this employee
    cur.execute("""
        SELECT bc.*, gb.b_name, gb.company, gb.due_date as bid_due_date
        FROM bid_checklists bc
        JOIN go_bids gb ON bc.g_id = gb.g_id
        WHERE bc.assigned_to = %s
        ORDER BY bc.due_date ASC, bc.priority DESC
    """, (employee_id,))
    tasks = cur.fetchall()
    
    # Calculate task statistics
    total_tasks = len(tasks)
    completed_tasks = len([t for t in tasks if t['status'] == 'completed'])
    pending_tasks = len([t for t in tasks if t['status'] == 'pending'])
    in_progress_tasks = len([t for t in tasks if t['status'] == 'in_progress'])
    
    cur.close()
    
    return render_template('employee_dashboard.html',
                         employee=employee,
                         tasks=tasks,
                         total_tasks=total_tasks,
                         completed_tasks=completed_tasks,
                         pending_tasks=pending_tasks,
                         in_progress_tasks=in_progress_tasks,
                         user=current_user)

@app.route('/task/<int:task_id>/update_status', methods=['POST'])
def update_task_status(task_id):
    """Update task status and progress"""
    try:
        new_status = request.form.get('status', '').strip()
        progress_notes = request.form.get('progress_notes', '').strip()
        
        if not new_status:
            return jsonify({'error': 'Status is required'}), 400
        
        cur = mysql.connection.cursor(DictCursor)
        
        # Get task info
        cur.execute("""
            SELECT bc.*, gb.b_name, gb.company, e.name as employee_name
            FROM bid_checklists bc
            JOIN go_bids gb ON bc.g_id = gb.g_id
            JOIN employees e ON bc.assigned_to = e.id
            WHERE bc.id = %s
        """, (task_id,))
        task = cur.fetchone()
        
        if not task:
            cur.close()
            return jsonify({'error': 'Task not found'}), 404
        
        # Authorization: allow admins/managers (flask-login) OR the assigned employee via session
        is_flask_user = hasattr(current_user, 'is_authenticated') and current_user.is_authenticated
        employee_id = session.get('employee_id')
        if not is_flask_user:
            # When not logged in via Flask-Login, require employee session and ownership of the task
            if not employee_id or task.get('assigned_to') != employee_id:
                cur.close()
                return jsonify({'error': 'Forbidden'}), 403

        # Update task status (and map to a percentage for persistence)
        status_lower = (new_status or '').strip().lower()
        if status_lower == 'completed':
            pct_val = 100
        elif status_lower == 'in_progress':
            pct_val = 50
        else:
            pct_val = 0
        cur.execute("""
            UPDATE bid_checklists 
            SET status = %s, progress_pct = %s, updated_at = CURRENT_TIMESTAMP
            WHERE id = %s
        """, (new_status, pct_val, task_id))
        
        # Log the update
        log_write('task_update', 
                 f"Task '{task['task_name']}' for bid '{task['b_name']}' updated to {new_status} by {task['employee_name']}")
        
        mysql.connection.commit()
        cur.close()
        
        # Emit real-time update to team dashboard
        socketio.emit('task_update', {
            'task_id': task_id,
            'status': new_status,
            'bid_name': task['b_name'],
            'employee_name': task['employee_name'],
            'company': task['company']
        })

        # Also emit master_update with per-stage progress for this bid
        try:
            bid_id = task['g_id']
            cur2 = mysql.connection.cursor(DictCursor)
            # Compute per-team completion rates by tasks
            def pct_for(role_expr):
                cur2.execute(f"SELECT status FROM bid_checklists bc JOIN users u ON bc.created_by=u.id WHERE bc.g_id=%s AND u.role {role_expr}", (bid_id,))
                rows = cur2.fetchall()
                if not rows:
                    return 0
                done = len([r for r in rows if (r.get('status') or '').lower()=='completed'])
                return int(round((done/max(1,len(rows)))*100))
            stage_progress_map = {
                'business': pct_for("='business dev'"),
                'design': pct_for("='design'"),
                'operations': pct_for("='operations'"),
                'engineer': pct_for("IN ('site manager','engineer')"),
            }
            cur2.close()
            socketio.emit('master_update', {
                'summary': {
                    'bid_id': bid_id,
                    'work_progress_pct': stage_progress_map.get('design',0),
                    'project_status': 'ongoing',
                    'work_status': f"Task '{task['task_name']}' -> {new_status}",
                    'stage_progress_map': stage_progress_map
                }
            })
        except Exception:
            pass
        
        return jsonify({'success': 'Task status updated successfully'})
        
    except Exception as e:
        mysql.connection.rollback()
        if 'cur' in locals():
            cur.close()
        return jsonify({'error': f'Error updating task: {str(e)}'}), 500

@app.route('/task/<int:task_id>/update', methods=['POST'])
def update_task(task_id):
    """Update task fields (name, description, status, due_date, priority, assigned_to)."""
    try:
        cur = mysql.connection.cursor(DictCursor)
        # Fetch existing task
        cur.execute("SELECT * FROM bid_checklists WHERE id=%s", (task_id,))
        task = cur.fetchone()
        if not task:
            cur.close()
            return jsonify({'error': 'Task not found'}), 404

        # Authorization: admin/manager or assigned employee via session
        is_flask_user = hasattr(current_user, 'is_authenticated') and current_user.is_authenticated
        employee_id = session.get('employee_id')
        if not is_flask_user and (not employee_id or task.get('assigned_to') != employee_id):
            cur.close()
            return jsonify({'error': 'Forbidden'}), 403

        # Collect fields (optional updates)
        fields = []
        values = []
        m = request.form
        if 'task_name' in m:
            fields.append('task_name=%s'); values.append(m.get('task_name').strip())
        if 'description' in m:
            fields.append('description=%s'); values.append(m.get('description').strip())
        if 'status' in m:
            fields.append('status=%s'); values.append(m.get('status').strip())
        if 'due_date' in m:
            fields.append('due_date=%s'); values.append(m.get('due_date'))
        if 'priority' in m:
            fields.append('priority=%s'); values.append(m.get('priority').strip())
        if 'assigned_to' in m:
            at = m.get('assigned_to') or None
            fields.append('assigned_to=%s'); values.append(at)
        if 'progress_pct' in m:
            try:
                pp = int(m.get('progress_pct'))
                pp = max(0, min(100, pp))
            except Exception:
                pp = 0
            fields.append('progress_pct=%s'); values.append(pp)
        if not fields:
            cur.close()
            return jsonify({'error': 'No fields to update'}), 400
        set_clause = ', '.join(fields) + ', updated_at = CURRENT_TIMESTAMP'
        values.append(task_id)
        cur.execute(f"UPDATE bid_checklists SET {set_clause} WHERE id=%s", tuple(values))
        mysql.connection.commit()
        cur.close()
        log_write('task_update_fields', f"task_id={task_id}")
        return jsonify({'success': 'Task updated'})
    except Exception as e:
        mysql.connection.rollback()
        if 'cur' in locals():
            cur.close()
        return jsonify({'error': f'Error updating task: {str(e)}'}), 500

@app.route('/task/<int:task_id>/delete', methods=['POST'])
def delete_task(task_id):
    """Delete a task if authorized."""
    try:
        cur = mysql.connection.cursor(DictCursor)
        cur.execute("SELECT assigned_to FROM bid_checklists WHERE id=%s", (task_id,))
        row = cur.fetchone()
        if not row:
            cur.close()
            return jsonify({'error': 'Task not found'}), 404
        is_flask_user = hasattr(current_user, 'is_authenticated') and current_user.is_authenticated
        employee_id = session.get('employee_id')
        if not is_flask_user and (not employee_id or row.get('assigned_to') != employee_id):
            cur.close()
            return jsonify({'error': 'Forbidden'}), 403
        cur.execute("DELETE FROM bid_checklists WHERE id=%s", (task_id,))
        mysql.connection.commit()
        cur.close()
        log_write('task_delete', f"task_id={task_id}")
        return jsonify({'success': 'Task deleted'})
    except Exception as e:
        mysql.connection.rollback()
        if 'cur' in locals():
            cur.close()
        return jsonify({'error': f'Error deleting task: {str(e)}'}), 500

@app.route('/team/<team>/bids/<int:g_id>/checklist')
@login_required
def bid_checklist(team, g_id):
    """Manage checklist/tasks for a specific bid"""
    cur = mysql.connection.cursor(DictCursor)
    
    # Get bid info
    cur.execute("SELECT * FROM go_bids WHERE g_id = %s", (g_id,))
    bid = cur.fetchone()
    
    if not bid:
        cur.close()
        return "Bid not found", 404
    
    # Get checklist items for this bid
    cur.execute("""
        SELECT bc.*, e.name as assigned_employee_name
        FROM bid_checklists bc
        LEFT JOIN employees e ON bc.assigned_to = e.id
        WHERE bc.g_id = %s
        ORDER BY bc.priority DESC, bc.created_at ASC
    """, (g_id,))
    checklist_items = cur.fetchall()
    
    # Get team employees for assignment
    cur.execute("""
        SELECT * FROM employees 
        WHERE department = %s AND is_active = TRUE
        ORDER BY name
    """, (team,))
    team_employees = cur.fetchall()
    
    cur.close()
    
    return render_template('bid_checklist.html',
                         team=team,
                         bid=bid,
                         checklist_items=checklist_items,
                         team_employees=team_employees,
                         user=current_user)

@app.route('/team/<team>/bids/<int:g_id>/checklist/create', methods=['POST'])
@login_required
def create_checklist_item(team, g_id):
    """Create a new checklist item for a bid"""
    try:
        task_name = request.form.get('task_name', '').strip()
        description = request.form.get('description', '').strip()
        assigned_to = request.form.get('assigned_to')
        priority = request.form.get('priority', 'medium')
        due_date = request.form.get('due_date')
        file_obj = request.files.get('attachment')
        saved_path = None
        
        if not task_name:
            flash('Task name is required', 'error')
            return redirect(url_for('bid_checklist', team=team, g_id=g_id))
        
        cur = mysql.connection.cursor()
        # Explicit stage name for parallel tracking
        stage_name = team.strip().lower()
        # Handle optional attachment upload
        if file_obj and getattr(file_obj, 'filename', None):
            try:
                os.makedirs(os.path.join(os.getcwd(), 'uploads', 'checklists', str(g_id)), exist_ok=True)
                safe_name = secure_filename(file_obj.filename)
                saved_path = os.path.join('uploads', 'checklists', str(g_id), safe_name)
                file_obj.save(os.path.join(os.getcwd(), saved_path))
            except Exception:
                saved_path = None
        cur.execute("""
            INSERT INTO bid_checklists (g_id, task_name, description, assigned_to, priority, due_date, progress_pct, stage, created_by, attachment_path)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (g_id, task_name, description, assigned_to if assigned_to else None, 
              priority, due_date if due_date else None, 0, stage_name, current_user.id, saved_path))
        
        mysql.connection.commit()
        cur.close()
        
        log_write('create_checklist_item', f"Created task '{task_name}' for bid {g_id}")
        flash(f'Task "{task_name}" created successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error creating task: {str(e)}', 'error')
    
    return redirect(url_for('bid_checklist', team=team, g_id=g_id))

@app.route('/task/<int:task_id>/attach', methods=['POST'])
@login_required
def attach_file_to_task(task_id):
    """Attach or replace a file for a checklist task."""
    try:
        file_obj = request.files.get('attachment')
        if not file_obj or not getattr(file_obj, 'filename', None):
            return jsonify({'error': 'No file provided'}), 400
        cur = mysql.connection.cursor(DictCursor)
        cur.execute("SELECT id, g_id, assigned_to FROM bid_checklists WHERE id=%s", (task_id,))
        task = cur.fetchone()
        if not task:
            cur.close()
            return jsonify({'error': 'Task not found'}), 404
        # Authorization: admin/manager or assigned employee via session
        is_flask_user = hasattr(current_user, 'is_authenticated') and current_user.is_authenticated
        employee_id = session.get('employee_id')
        if not is_flask_user and (not employee_id or task.get('assigned_to') != employee_id):
            cur.close()
            return jsonify({'error': 'Forbidden'}), 403
        g_id = task['g_id']
        os.makedirs(os.path.join(os.getcwd(), 'uploads', 'checklists', str(g_id)), exist_ok=True)
        safe_name = secure_filename(file_obj.filename)
        saved_rel = os.path.join('uploads', 'checklists', str(g_id), safe_name)
        file_obj.save(os.path.join(os.getcwd(), saved_rel))
        cur.execute("UPDATE bid_checklists SET attachment_path=%s, updated_at=CURRENT_TIMESTAMP WHERE id=%s", (saved_rel, task_id))
        mysql.connection.commit()
        cur.close()
        return jsonify({'success': True, 'attachment_path': saved_rel})
    except Exception as e:
        mysql.connection.rollback()
        if 'cur' in locals():
            cur.close()
        return jsonify({'error': str(e)}), 500

@app.route('/task/<int:task_id>/attachment')
@login_required
def get_task_attachment(task_id):
    """Serve the attachment file for a task if present."""
    try:
        cur = mysql.connection.cursor(DictCursor)
        cur.execute("SELECT attachment_path FROM bid_checklists WHERE id=%s", (task_id,))
        row = cur.fetchone()
        cur.close()
        path = (row or {}).get('attachment_path') if row else None
        if not path:
            return "No attachment", 404
        abs_path = os.path.join(os.getcwd(), path)
        directory = os.path.dirname(abs_path)
        filename = os.path.basename(abs_path)
        if not os.path.exists(abs_path):
            return "Attachment missing on server", 404
        return send_from_directory(directory, filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}", 500

@app.route('/team/<team>/transfer_project', methods=['POST'])
@login_required
def transfer_project(team):
    """Transfer a project to another team"""
    try:
        g_id = request.form.get('g_id')
        to_team = request.form.get('to_team')
        transfer_reason = request.form.get('transfer_reason', '').strip()
        
        if not g_id or not to_team:
            flash('Project and destination team are required', 'error')
            return redirect(url_for('team_dashboard', team=team))
        
        cur = mysql.connection.cursor(DictCursor)
        
        # Get bid info
        cur.execute("SELECT * FROM go_bids WHERE g_id = %s", (g_id,))
        bid = cur.fetchone()
        
        if not bid:
            cur.close()
            flash('Bid not found', 'error')
            return redirect(url_for('team_dashboard', team=team))
        
        # Create transfer record
        cur.execute("""
            INSERT INTO project_transfers (g_id, from_team, to_team, transferred_by, transfer_reason)
            VALUES (%s, %s, %s, %s, %s)
        """, (g_id, team, to_team, current_user.id, transfer_reason))
        
        # Update bid state to next team
        cur.execute("UPDATE go_bids SET state = %s WHERE g_id = %s", (to_team, g_id))

        # When transferring, preserve current team's tasks and generate new tasks for receiving team
        # First, mark current team's tasks as completed and archive them
        cur.execute("""
            UPDATE bid_checklists 
            SET status = 'completed', progress_pct = 100, team_archive = %s, updated_at = CURRENT_TIMESTAMP
            WHERE g_id = %s AND created_by IN (
                SELECT id FROM users WHERE role = %s
            )
        """, (team, g_id, team))
        
        # Generate team-specific default checklist for the receiving team
        generate_team_checklist(cur, g_id, to_team)
        
        mysql.connection.commit()
        cur.close()
        
        log_write('project_transfer', 
                 f"Transferred bid '{bid['b_name']}' from {team} to {to_team}")
        flash(f'Project "{bid["b_name"]}" transferred to {to_team} team successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error transferring project: {str(e)}', 'error')
    
    return redirect(url_for('team_dashboard', team=team))

# --- RFP File Routes ---

def _ensure_uploaded_rfp_table_exists(cur):
    """Ensure the uploaded_rfp_files table exists before querying."""
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS uploaded_rfp_files (
            id INT AUTO_INCREMENT PRIMARY KEY,
            bid_id INT,
            g_id INT,
            filename VARCHAR(500) NOT NULL,
            file_path VARCHAR(1000) NOT NULL,
            file_size BIGINT,
            uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            uploaded_by INT,
            INDEX idx_bid_id (bid_id),
            INDEX idx_g_id (g_id)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """
    )
    mysql.connection.commit()

    def _column_exists(column_name: str) -> bool:
        cur.execute("SHOW COLUMNS FROM uploaded_rfp_files LIKE %s", (column_name,))
        return cur.fetchone() is not None

    def _index_exists(index_name: str) -> bool:
        cur.execute("SHOW INDEX FROM uploaded_rfp_files WHERE Key_name = %s", (index_name,))
        return cur.fetchone() is not None

    def _ensure_column(column_name: str, ddl: str):
        if _column_exists(column_name):
            return
        try:
            cur.execute(f"ALTER TABLE uploaded_rfp_files ADD COLUMN {ddl}")
            mysql.connection.commit()
        except Exception:
            mysql.connection.rollback()

    def _ensure_index(index_name: str, ddl: str):
        if _index_exists(index_name):
            return
        try:
            cur.execute(f"ALTER TABLE uploaded_rfp_files ADD INDEX {ddl}")
            mysql.connection.commit()
        except Exception:
            mysql.connection.rollback()

    # Ensure legacy tables have required columns and indexes, ignoring duplicate errors safely
    _ensure_column('bid_id', "bid_id INT AFTER id")
    _ensure_index('idx_bid_id', "idx_bid_id (bid_id)")

    _ensure_column('g_id', "g_id INT AFTER bid_id")
    _ensure_index('idx_g_id', "idx_g_id (g_id)")

    _ensure_column('filename', "filename VARCHAR(500) DEFAULT NULL")
    _ensure_column('original_filename', "original_filename VARCHAR(500) DEFAULT NULL")
    _ensure_column('saved_filename', "saved_filename VARCHAR(500) DEFAULT NULL")
    _ensure_column('file_type', "file_type VARCHAR(50) DEFAULT 'pdf'")
    _ensure_column('file_hash', "file_hash VARCHAR(128) DEFAULT NULL")
    _ensure_column('file_size', "file_size BIGINT DEFAULT NULL")
    _ensure_column('uploaded_at', "uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP")
    _ensure_column('uploaded_by', "uploaded_by INT DEFAULT NULL")
    _ensure_column('file_path', "file_path VARCHAR(1000) DEFAULT NULL")
    _ensure_column('section_id', "section_id VARCHAR(100) DEFAULT NULL")


def _get_latest_rfp_file_for_bid(g_id):
    """Fetch the most recent RFP file record for a bid using g_id."""
    cur = mysql.connection.cursor(DictCursor)
    try:
        _ensure_uploaded_rfp_table_exists(cur)

        # Inspect available columns to support legacy schemas
        cur.execute("SHOW COLUMNS FROM uploaded_rfp_files")
        column_rows = cur.fetchall() or []
        columns = {row['Field'] for row in column_rows if 'Field' in row}

        has_g_id = 'g_id' in columns
        has_bid_id = 'bid_id' in columns
        has_filename = 'filename' in columns
        has_original_filename = 'original_filename' in columns
        has_saved_filename = 'saved_filename' in columns
        has_file_path = 'file_path' in columns
        has_file_size = 'file_size' in columns
        has_uploaded_at = 'uploaded_at' in columns

        select_fields = ['id']
        if has_g_id:
            select_fields.append('g_id')
        else:
            select_fields.append('NULL AS g_id')
        if has_bid_id:
            select_fields.append('bid_id')
        else:
            select_fields.append('NULL AS bid_id')

        if has_filename:
            select_fields.append('filename')
        else:
            # Fall back to original/saved filename columns if present
            coalesce_parts = []
            if has_original_filename:
                coalesce_parts.append('original_filename')
            if has_saved_filename:
                coalesce_parts.append('saved_filename')
            if coalesce_parts:
                select_fields.append(f"COALESCE({', '.join(coalesce_parts)}) AS filename")
            else:
                select_fields.append("NULL AS filename")

        if has_file_path:
            select_fields.append('file_path')
        else:
            select_fields.append('NULL AS file_path')

        if has_file_size:
            select_fields.append('file_size')
        else:
            select_fields.append('NULL AS file_size')

        if has_uploaded_at:
            select_fields.append('uploaded_at')
            order_clause = "uploaded_at DESC, id DESC"
        else:
            # Legacy tables might use created_at/updated_at; prefer whichever exists
            timestamp_col = None
            for legacy_col in ('created_at', 'updated_at', 'saved_at'):
                if legacy_col in columns:
                    timestamp_col = legacy_col
                    break
            if timestamp_col:
                select_fields.append(f"{timestamp_col} AS uploaded_at")
                order_clause = f"{timestamp_col} DESC, id DESC"
            else:
                select_fields.append('NULL AS uploaded_at')
                order_clause = "id DESC"

        where_clauses = []
        params = []
        if has_g_id:
            where_clauses.append("g_id = %s")
            params.append(g_id)
        if has_bid_id:
            where_clauses.append("bid_id = %s")
            params.append(g_id)

        if not where_clauses:
            # As a last resort, attempt to match against legacy columns such as bid or project id
            for legacy_col in ('b_id', 'project_id'):
                if legacy_col in columns:
                    where_clauses.append(f"{legacy_col} = %s")
                    params.append(g_id)
                    break

        if not where_clauses:
            return None

        select_clause = ", ".join(select_fields)
        base_where = " OR ".join(f"({clause})" for clause in where_clauses)
        # Only consider true RFP source PDFs (exclude per-section attachments)
        extra_filters = []
        # Exclude rows tied to a section if the column exists
        extra_filters.append("(section_id IS NULL OR section_id = '')")
        # Ensure the file is a PDF based on file_type or filename columns if present
        pdf_checks = ["LOWER(file_path) LIKE '%%.pdf'"]
        if 'file_type' in columns:
            pdf_checks.append("LOWER(file_type) = 'pdf'")
        if 'filename' in columns:
            pdf_checks.append("LOWER(filename) LIKE '%%.pdf'")
        if 'saved_filename' in columns:
            pdf_checks.append("LOWER(saved_filename) LIKE '%%.pdf'")
        extra_filters.append("(" + " OR ".join(pdf_checks) + ")")
        where_clause_sql = f"({base_where}) AND " + " AND ".join(extra_filters)

        try:
            cur.execute(
                f"""
                SELECT {select_clause}
                FROM uploaded_rfp_files
                WHERE {where_clause_sql}
                ORDER BY {order_clause}
                LIMIT 1
                """,
                tuple(params),
            )
        except Exception:
            # The table may not exist yet—create it and retry once.
            mysql.connection.rollback()
            _ensure_uploaded_rfp_table_exists(cur)
            cur.execute(
                f"""
                SELECT {select_clause}
                FROM uploaded_rfp_files
                WHERE {where_clause_sql}
                ORDER BY {order_clause}
                LIMIT 1
                """,
                tuple(params),
            )
        row = cur.fetchone()

        # Fallback: if nothing matched strict filters (e.g., legacy rows with section_id populated),
        # relax the section_id constraint and try again for a PDF.
        if not row:
            relaxed_filters = []
            relaxed_filters.append(base_where)
            relaxed_pdf_checks = []
            relaxed_pdf_checks.append("LOWER(file_path) LIKE '%%.pdf'")
            if 'file_type' in columns:
                relaxed_pdf_checks.append("LOWER(file_type) = 'pdf'")
            if 'filename' in columns:
                relaxed_pdf_checks.append("LOWER(filename) LIKE '%%.pdf'")
            if 'saved_filename' in columns:
                relaxed_pdf_checks.append("LOWER(saved_filename) LIKE '%%.pdf'")
            relaxed_where_sql = f"({base_where}) AND (" + " OR ".join(relaxed_pdf_checks) + ")"
            try:
                cur.execute(
                    f"""
                    SELECT {select_clause}
                    FROM uploaded_rfp_files
                    WHERE {relaxed_where_sql}
                    ORDER BY {order_clause}
                    LIMIT 1
                    """,
                    tuple(params),
                )
                row = cur.fetchone()
            except Exception:
                pass

        if row and 'id' in row:
            update_clauses = []
            update_params = []
            # Backfill missing identifiers for legacy rows
            if has_g_id and (not row.get('g_id')):
                update_clauses.append("g_id = %s")
                update_params.append(g_id)
                row['g_id'] = g_id
            if has_bid_id and (not row.get('bid_id')):
                update_clauses.append("bid_id = %s")
                update_params.append(g_id)
                row['bid_id'] = g_id

            if update_clauses:
                update_params.append(row['id'])
                try:
                    cur.execute(
                        f"UPDATE uploaded_rfp_files SET {', '.join(update_clauses)} WHERE id = %s",
                        tuple(update_params),
                    )
                    mysql.connection.commit()
                except Exception:
                    mysql.connection.rollback()

        return row
    finally:
        cur.close()


def _extract_pdf_pages(file_path, start_page=1, limit=None):
    """Extract text from a PDF file path using PyMuPDF or fall back to PyPDF2."""
    pages = []
    total_pages = 0
    start_index = max(start_page - 1, 0)
    use_limit = None if (limit is None or limit <= 0) else limit

    if _FITZ_AVAILABLE:
        with _pymupdf.open(file_path) as doc:
            total_pages = doc.page_count
            if start_index >= total_pages:
                return pages, total_pages
            end_index = total_pages if use_limit is None else min(total_pages, start_index + use_limit)
            for index in range(start_index, end_index):
                page = doc.load_page(index)
                text = page.get_text("text") or ""
                pages.append(
                    {
                        "number": index + 1,
                        "text": text,
                        "characters": len(text),
                    }
                )
        return pages, total_pages

    if _PdfReader is not None:
        with open(file_path, "rb") as fh:
            reader = _PdfReader(fh)
            total_pages = len(reader.pages)
            if start_index >= total_pages:
                return pages, total_pages
            end_index = total_pages if use_limit is None else min(total_pages, start_index + use_limit)
            for index in range(start_index, end_index):
                page = reader.pages[index]
                text = page.extract_text() or ""
                pages.append(
                    {
                        "number": index + 1,
                        "text": text,
                        "characters": len(text),
                    }
                )
        return pages, total_pages

    raise RuntimeError(
        "PDF text extraction requires PyMuPDF (`pip install PyMuPDF`) or PyPDF2 (`pip install PyPDF2`)."
    )


@app.route('/api/rfp-file/<int:g_id>/upload', methods=['POST'])
@login_required
def api_upload_rfp_file(g_id):
    """Upload a PDF for the specified g_id and store metadata in uploaded_rfp_files."""
    upload_file = request.files.get('file')
    if not upload_file or not upload_file.filename:
        return jsonify({'error': 'missing_file', 'message': 'Please choose a PDF to upload.'}), 400

    filename_lower = upload_file.filename.lower()
    if not filename_lower.endswith('.pdf'):
        return jsonify({'error': 'invalid_type', 'message': 'Only PDF files are supported.'}), 400

    from werkzeug.utils import secure_filename
    import uuid
    import hashlib

    safe_original = secure_filename(upload_file.filename) or f"rfp_{uuid.uuid4().hex}.pdf"
    unique_token = uuid.uuid4().hex
    saved_filename = f"{unique_token}_{safe_original}"

    uploads_dir = os.path.join(os.getcwd(), 'uploads', 'rfp')
    os.makedirs(uploads_dir, exist_ok=True)
    absolute_path = os.path.join(uploads_dir, saved_filename)

    try:
        upload_file.save(absolute_path)
    except Exception as err:
        return jsonify({'error': 'save_failed', 'message': f'Could not save file: {err}'}), 500

    file_size = 0
    file_hash = ''
    try:
        file_size = os.path.getsize(absolute_path)
        with open(absolute_path, 'rb') as handler:
            file_hash = hashlib.sha256(handler.read()).hexdigest()
    except Exception:
        pass

    # Lookup bid_id (legacy) from go_bids table if available
    bid_id = None
    cur_lookup = mysql.connection.cursor(DictCursor)
    try:
        cur_lookup.execute("SELECT id FROM go_bids WHERE g_id = %s", (g_id,))
        go_bid_row = cur_lookup.fetchone()
        if go_bid_row and go_bid_row.get('id'):
            bid_id = go_bid_row['id']
    finally:
        cur_lookup.close()

    cur = mysql.connection.cursor(DictCursor)
    try:
        _ensure_uploaded_rfp_table_exists(cur)
        cur.execute("SHOW COLUMNS FROM uploaded_rfp_files")
        column_rows = cur.fetchall() or []
        available_columns = {row['Field'] for row in column_rows if 'Field' in row}

        columns = []
        values = []

        def add_column(col_name, value):
            if col_name in available_columns:
                columns.append(col_name)
                values.append(value)

        add_column('bid_id', bid_id)
        add_column('g_id', g_id)
        add_column('filename', safe_original)
        add_column('original_filename', upload_file.filename)
        add_column('saved_filename', saved_filename)
        add_column('file_path', absolute_path)
        add_column('file_type', 'pdf')
        add_column('file_size', file_size)
        add_column('file_hash', file_hash)
        add_column('uploaded_by', getattr(current_user, 'id', None))

        if not columns:
            return jsonify({'error': 'schema_error', 'message': 'Unable to persist the uploaded file.'}), 500

        placeholders = ','.join(['%s'] * len(columns))
        column_sql = ','.join(columns)
        cur.execute(
            f"INSERT INTO uploaded_rfp_files ({column_sql}) VALUES ({placeholders})",
            tuple(values),
        )
        mysql.connection.commit()

        # Fetch parsed preview for immediate display
        parsed_pages = []
        total_pages = 0
        next_start = None
        try:
            parsed_pages, total_pages = _extract_pdf_pages(absolute_path, start_page=1, limit=3)
            if parsed_pages:
                last_number = parsed_pages[-1].get('number')
                if last_number and total_pages and last_number < total_pages:
                    next_start = last_number + 1
        except Exception:
            parsed_pages = []
            total_pages = 0
            next_start = None

        return jsonify(
            {
                'success': True,
                'file_id': cur.lastrowid,
                'filename': safe_original,
                'total_pages': total_pages,
                'next_start': next_start,
                'pages': parsed_pages,
            }
        )
    except Exception as err:
        mysql.connection.rollback()
        if os.path.exists(absolute_path):
            try:
                os.remove(absolute_path)
            except Exception:
                pass
        return jsonify({'error': 'upload_failed', 'message': str(err)}), 500
    finally:
        cur.close()


@app.route('/api/section-attachment/<int:g_id>/upload', methods=['POST'])
@login_required
def api_upload_section_attachment(g_id):
    """Upload attachments for a specific section (images or PDFs)."""
    section_id = request.args.get('section_id') or request.form.get('section_id') or ''
    section_id = section_id.strip()
    if not section_id:
        return jsonify({'error': 'missing_section', 'message': 'Section ID is required.'}), 400
    upload_file = request.files.get('file')
    if not upload_file or not upload_file.filename:
        return jsonify({'error': 'missing_file', 'message': 'Please choose a file to upload.'}), 400
    filename_lower = upload_file.filename.lower()
    # Allow images and PDF files
    allowed_ext = ('.png', '.jpg', '.jpeg', '.webp', '.pdf')
    if not filename_lower.endswith(allowed_ext):
        return jsonify({'error': 'invalid_type', 'message': 'Only images (.png, .jpg, .jpeg, .webp) or PDF files are allowed.'}), 400
    from werkzeug.utils import secure_filename
    import uuid
    import hashlib
    safe_original = secure_filename(upload_file.filename) or f"img_{uuid.uuid4().hex}.png"
    unique_token = uuid.uuid4().hex
    saved_filename = f"{unique_token}_{safe_original}"
    uploads_dir = os.path.join(os.getcwd(), 'uploads', 'sections', section_id)
    os.makedirs(uploads_dir, exist_ok=True)
    absolute_path = os.path.join(uploads_dir, saved_filename)
    try:
        upload_file.save(absolute_path)
    except Exception as err:
        return jsonify({'error': 'save_failed', 'message': f'Could not save file: {err}'}), 500
    file_size = 0
    file_hash = ''
    try:
        file_size = os.path.getsize(absolute_path)
        with open(absolute_path, 'rb') as handler:
            file_hash = hashlib.sha256(handler.read()).hexdigest()
    except Exception:
        pass
    bid_id = None
    cur_lookup = mysql.connection.cursor(DictCursor)
    try:
        cur_lookup.execute("SELECT id FROM go_bids WHERE g_id = %s", (g_id,))
        go_bid_row = cur_lookup.fetchone()
        if go_bid_row and go_bid_row.get('id'):
            bid_id = go_bid_row['id']
    finally:
        cur_lookup.close()
    cur = mysql.connection.cursor(DictCursor)
    try:
        _ensure_uploaded_rfp_table_exists(cur)
        cur.execute("SHOW COLUMNS FROM uploaded_rfp_files")
        column_rows = cur.fetchall() or []
        available_columns = {row['Field'] for row in column_rows if 'Field' in row}
        columns = []
        values = []
        def add_column(col_name, value):
            if col_name in available_columns:
                columns.append(col_name)
                values.append(value)
        add_column('bid_id', bid_id)
        add_column('g_id', g_id)
        add_column('filename', safe_original)
        add_column('original_filename', upload_file.filename)
        add_column('saved_filename', saved_filename)
        add_column('file_path', absolute_path)
        add_column('file_type', 'pdf' if filename_lower.endswith('.pdf') else 'image')
        add_column('file_size', file_size)
        add_column('file_hash', file_hash)
        add_column('uploaded_by', getattr(current_user, 'id', None))
        add_column('section_id', section_id)
        if not columns:
            return jsonify({'error': 'schema_error', 'message': 'Unable to persist the uploaded image.'}), 500
        placeholders = ','.join(['%s'] * len(columns))
        column_sql = ','.join(columns)
        cur.execute(
            f"INSERT INTO uploaded_rfp_files ({column_sql}) VALUES ({placeholders})",
            tuple(values),
        )
        mysql.connection.commit()
        return jsonify(
            {
                'success': True,
                'file_id': cur.lastrowid,
                'filename': safe_original,
                'section_id': section_id,
                'file_type': 'pdf' if filename_lower.endswith('.pdf') else 'image',
            }
        )
    except Exception as err:
        mysql.connection.rollback()
        if os.path.exists(absolute_path):
            try:
                os.remove(absolute_path)
            except Exception:
                pass
        return jsonify({'error': 'upload_failed', 'message': str(err)}), 500
    finally:
        cur.close()


@app.route('/rfp-file/view/by-g/<int:g_id>')
@login_required
def view_rfp_file_by_g(g_id):
    """View the latest RFP file for a bid using its g_id."""
    try:
        file_data = _get_latest_rfp_file_for_bid(g_id)
    except Exception as err:
        app.logger.exception("Error locating RFP file for g_id %s: %s", g_id, err)
        abort(500, description="Failed to locate RFP file.")

    if not file_data:
        abort(404, description="No RFP file associated with this project.")

    file_path = file_data.get('file_path')
    filename = file_data.get('filename') or os.path.basename(file_path or '')

    if not file_path or not os.path.exists(file_path):
        abort(404, description="RFP file not found on server.")

    if file_path.lower().endswith('.pdf'):
        return send_file(file_path, mimetype='application/pdf', as_attachment=False, download_name=filename)

    return send_file(file_path, as_attachment=True, download_name=filename)


@app.route('/api/rfp-file/<int:g_id>/parsed', methods=['GET'])
@login_required
def api_rfp_file_parsed(g_id):
    """Return parsed PDF text for the latest RFP file associated with g_id."""
    start = request.args.get('start', default=1, type=int)
    limit = request.args.get('limit', default=3, type=int)
    page_limit = None if limit is None or limit <= 0 else limit
    start = start if start and start > 0 else 1

    try:
        file_data = _get_latest_rfp_file_for_bid(g_id)
    except Exception as err:
        app.logger.exception("Error locating RFP file for g_id %s: %s", g_id, err)
        return jsonify({'error': 'internal_error'}), 500

    if not file_data:
        return jsonify({'error': 'not_found', 'message': 'No RFP file associated with this project.'}), 404

    file_path = file_data.get('file_path')
    if not file_path or not os.path.exists(file_path):
        return jsonify({'error': 'file_missing', 'message': 'RFP file not found on server.'}), 404

    if not file_path.lower().endswith('.pdf'):
        return jsonify({'error': 'unsupported_type', 'message': 'Only PDF files can be parsed.'}), 400

    try:
        pages, total_pages = _extract_pdf_pages(file_path, start_page=start, limit=page_limit)
    except Exception as err:
        app.logger.exception("Error parsing PDF for g_id %s: %s", g_id, err)
        return jsonify({'error': 'parse_failed', 'message': 'Unable to parse the PDF file.'}), 500

    next_start = None
    if pages:
        last_page_number = pages[-1]['number']
        if last_page_number < total_pages:
            next_start = last_page_number + 1

    return jsonify(
        {
            'g_id': g_id,
            'file_id': file_data.get('id'),
            'filename': file_data.get('filename'),
            'total_pages': total_pages,
            'start': start,
            'limit': page_limit,
            'next_start': next_start,
            'pages': pages,
        }
    )


@app.route('/rfp-file/view/<int:file_id>')
@login_required
def view_rfp_file(file_id):
    """View RFP file in browser"""
    try:
        cur = mysql.connection.cursor(DictCursor)
        cur.execute("""
            SELECT file_path, filename
            FROM uploaded_rfp_files
            WHERE id = %s
        """, (file_id,))
        file_data = cur.fetchone()
        cur.close()
        
        if not file_data:
            flash('RFP file not found', 'error')
            return redirect(request.referrer or url_for('master_dashboard'))
        
        file_path = file_data['file_path']
        if not os.path.exists(file_path):
            flash('RFP file not found on server', 'error')
            return redirect(request.referrer or url_for('master_dashboard'))
        
        # For PDF files, send as inline to view in browser
        if file_path.lower().endswith('.pdf'):
            return send_file(file_path, mimetype='application/pdf', as_attachment=False)
        else:
            # For other file types, download
            return send_file(file_path, as_attachment=True)
    except Exception as e:
        flash(f'Error viewing file: {str(e)}', 'error')
        return redirect(request.referrer or url_for('master_dashboard'))

@app.route('/rfp-file/download/<int:file_id>')
@login_required
def download_rfp_file(file_id):
    """Download RFP file"""
    try:
        cur = mysql.connection.cursor(DictCursor)
        cur.execute("""
            SELECT file_path, filename
            FROM uploaded_rfp_files
            WHERE id = %s
        """, (file_id,))
        file_data = cur.fetchone()
        cur.close()
        
        if not file_data:
            flash('RFP file not found', 'error')
            return redirect(request.referrer or url_for('master_dashboard'))
        
        file_path = file_data['file_path']
        filename = file_data['filename']
        
        if not os.path.exists(file_path):
            flash('RFP file not found on server', 'error')
            return redirect(request.referrer or url_for('master_dashboard'))
        
        return send_file(file_path, as_attachment=True, download_name=filename)
    except Exception as e:
        flash(f'Error downloading file: {str(e)}', 'error')
        return redirect(request.referrer or url_for('master_dashboard'))


@app.route('/api/rfp-file/decrypt', methods=['POST'])
@login_required
def decrypt_rfp_file():
    """Decrypt an encrypted PDF file and return the decrypted version."""
    try:
        pdf_url = request.args.get('url')
        if not pdf_url:
            return jsonify({'error': 'URL parameter is required'}), 400

        # Extract file path from URL or fetch the file
        from urllib.parse import urlparse, unquote
        import tempfile
        
        # If it's a local file URL, extract the path
        if pdf_url.startswith('/'):
            # It's a relative URL, construct full path
            parsed = urlparse(pdf_url)
            # Try to find the file in the uploads directory
            if 'rfp-file/view' in pdf_url:
                # Extract file_id or g_id from URL
                import re
                match = re.search(r'/rfp-file/view/(\d+)', pdf_url)
                if match:
                    file_id = int(match.group(1))
                    cur = mysql.connection.cursor(DictCursor)
                    cur.execute("SELECT file_path FROM uploaded_rfp_files WHERE id = %s", (file_id,))
                    file_data = cur.fetchone()
                    cur.close()
                    if file_data and os.path.exists(file_data['file_path']):
                        file_path = file_data['file_path']
                    else:
                        return jsonify({'error': 'File not found'}), 404
                else:
                    # Try g_id
                    match = re.search(r'/rfp-file/view/by-g/(\d+)', pdf_url)
                    if match:
                        g_id = int(match.group(1))
                        file_data = _get_latest_rfp_file_for_bid(g_id)
                        if file_data and os.path.exists(file_data.get('file_path')):
                            file_path = file_data['file_path']
                        else:
                            return jsonify({'error': 'File not found'}), 404
                    else:
                        return jsonify({'error': 'Invalid URL format'}), 400
            else:
                return jsonify({'error': 'Invalid URL format'}), 400
        else:
            return jsonify({'error': 'Only local file URLs are supported'}), 400

        if not _FITZ_AVAILABLE:
            return jsonify({'error': 'PyMuPDF is not available for PDF decryption'}), 500

        # Attempt to decrypt the PDF
        try:
            # Try opening with empty password first (some PDFs have empty password)
            doc = _pymupdf.open(file_path)
            
            # If PDF is encrypted, try common passwords or empty password
            if doc.is_encrypted:
                # Try empty password
                if not doc.authenticate(""):
                    # Try some common passwords
                    common_passwords = ["", "password", "admin", "1234", "12345"]
                    authenticated = False
                    for pwd in common_passwords:
                        if doc.authenticate(pwd):
                            authenticated = True
                            break
                    
                    if not authenticated:
                        # Try to decrypt without password (some PDFs can be decrypted this way)
                        try:
                            doc.close()
                            # Create a new document by copying pages (this sometimes works for encrypted PDFs)
                            doc = _pymupdf.open(file_path)
                            if doc.is_encrypted and not doc.authenticate(""):
                                return jsonify({'error': 'PDF is password protected and cannot be decrypted automatically'}), 400
                        except:
                            return jsonify({'error': 'PDF is password protected and cannot be decrypted automatically'}), 400

            # Create a new decrypted PDF in memory
            decrypted_doc = _pymupdf.open()  # Create new empty PDF
            decrypted_doc.insert_pdf(doc)  # Copy all pages from encrypted to decrypted
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            temp_path = temp_file.name
            temp_file.close()
            
            decrypted_doc.save(temp_path)
            decrypted_doc.close()
            doc.close()
            
            # Return the decrypted PDF
            return send_file(temp_path, mimetype='application/pdf', as_attachment=False, download_name='decrypted_rfp.pdf')
            
        except Exception as decrypt_error:
            app.logger.exception("Error decrypting PDF: %s", decrypt_error)
            return jsonify({'error': f'Failed to decrypt PDF: {str(decrypt_error)}'}), 500

    except Exception as e:
        app.logger.exception("Error in decrypt_rfp_file: %s", e)
        return jsonify({'error': str(e)}), 500


@app.route('/api/rfp-file/<int:g_id>/save-edited', methods=['POST'])
@login_required
def save_edited_rfp_file(g_id):
    """Save an edited PDF file."""
    try:
        upload_file = request.files.get('file')
        if not upload_file or not upload_file.filename:
            return jsonify({'error': 'missing_file', 'message': 'Please provide a PDF file to save.'}), 400

        filename_lower = upload_file.filename.lower()
        if not filename_lower.endswith('.pdf'):
            return jsonify({'error': 'invalid_type', 'message': 'Only PDF files are supported.'}), 400

        from werkzeug.utils import secure_filename
        import uuid
        import hashlib

        safe_original = secure_filename(upload_file.filename) or f"edited_rfp_{uuid.uuid4().hex}.pdf"
        unique_token = uuid.uuid4().hex
        saved_filename = f"{unique_token}_{safe_original}"

        uploads_dir = os.path.join(os.getcwd(), 'uploads', 'rfp')
        os.makedirs(uploads_dir, exist_ok=True)
        absolute_path = os.path.join(uploads_dir, saved_filename)

        try:
            upload_file.save(absolute_path)
        except Exception as err:
            return jsonify({'error': 'save_failed', 'message': f'Could not save file: {err}'}), 500

        file_size = 0
        file_hash = ''
        try:
            file_size = os.path.getsize(absolute_path)
            with open(absolute_path, 'rb') as handler:
                file_hash = hashlib.sha256(handler.read()).hexdigest()
        except Exception:
            pass

        # Lookup bid_id (legacy) from go_bids table if available
        bid_id = None
        cur_lookup = mysql.connection.cursor(DictCursor)
        try:
            cur_lookup.execute("SELECT id FROM go_bids WHERE g_id = %s", (g_id,))
            go_bid_row = cur_lookup.fetchone()
            if go_bid_row and go_bid_row.get('id'):
                bid_id = go_bid_row['id']
        finally:
            cur_lookup.close()

        cur = mysql.connection.cursor(DictCursor)
        try:
            _ensure_uploaded_rfp_table_exists(cur)
            cur.execute("SHOW COLUMNS FROM uploaded_rfp_files")
            column_rows = cur.fetchall() or []
            available_columns = {row['Field'] for row in column_rows if 'Field' in row}

            columns = []
            values = []

            def add_column(col_name, value):
                if col_name in available_columns:
                    columns.append(col_name)
                    values.append(value)

            add_column('bid_id', bid_id)
            add_column('g_id', g_id)
            add_column('filename', safe_original)
            add_column('original_filename', upload_file.filename)
            add_column('saved_filename', saved_filename)
            add_column('file_path', absolute_path)
            add_column('file_type', 'pdf')
            add_column('file_size', file_size)
            add_column('file_hash', file_hash)
            add_column('uploaded_by', getattr(current_user, 'id', None))

            if not columns:
                return jsonify({'error': 'schema_error', 'message': 'Unable to persist the uploaded file.'}), 500

            placeholders = ','.join(['%s'] * len(columns))
            column_sql = ','.join(columns)
            cur.execute(
                f"INSERT INTO uploaded_rfp_files ({column_sql}) VALUES ({placeholders})",
                tuple(values),
            )
            mysql.connection.commit()

            return jsonify({
                'success': True,
                'message': 'Edited PDF saved successfully',
                'file_id': cur.lastrowid,
                'filename': safe_original
            })

        except Exception as db_error:
            mysql.connection.rollback()
            app.logger.exception("Database error saving edited PDF: %s", db_error)
            return jsonify({'error': 'database_error', 'message': f'Failed to save file metadata: {db_error}'}), 500
        finally:
            cur.close()

    except Exception as e:
        app.logger.exception("Error in save_edited_rfp_file: %s", e)
        return jsonify({'error': 'server_error', 'message': str(e)}), 500

# --- API Endpoints for Team Dashboard ---
@app.route('/api/team/<team>/employees')
@login_required
def api_team_employees(team):
    """API endpoint to get team employees"""
    cur = mysql.connection.cursor(DictCursor)
    
    cur.execute("""
        SELECT e.*, u.email as team_lead_email
        FROM employees e
        LEFT JOIN users u ON e.team_lead_id = u.id
        WHERE e.department = %s AND e.is_active = TRUE
        ORDER BY e.name
    """, (team,))
    employees = cur.fetchall()
    
    cur.close()
    
    return jsonify({'employees': employees})

@app.route('/api/team/<team>/assign', methods=['POST'])
@login_required
def api_team_assign(team):
    """Assign a bid to an employee from the current team.
    Body: JSON { g_id: int, employee_id: int }
    Note: Includes safe fallbacks when bid_assign.state or assignee_email
    columns are missing in older schemas."""
    try:
        data = request.get_json(silent=True) or {}
        g_id = int(data.get('g_id')) if data.get('g_id') is not None else None
        # accept single or multiple
        employee_ids = data.get('employee_ids')
        if employee_ids is None and data.get('employee_id') is not None:
            employee_ids = [data.get('employee_id')]
        if isinstance(employee_ids, list):
            try:
                employee_ids = [int(x) for x in employee_ids if x is not None]
            except Exception:
                employee_ids = []
        else:
            employee_ids = []
    except Exception:
        return jsonify({'ok': False, 'error': 'Invalid payload'}), 400

    if not g_id:
        return jsonify({'ok': False, 'error': 'g_id is required'}), 400

    cur = mysql.connection.cursor(DictCursor)
    try:
        # Validate employee belongs to this team
        # Validate employees belong to this team
        valid_emps = []
        if employee_ids:
            placeholders = ','.join(['%s'] * len(employee_ids))
            cur.execute(
                f"SELECT id, name, email FROM employees WHERE id IN ({placeholders}) AND department=%s AND is_active=TRUE",
                (*employee_ids, team)
            )
            valid_emps = cur.fetchall()
        # For legacy single assignment
        first_emp = (valid_emps[0] if valid_emps else None)

        # Get bid basic details
        cur.execute("SELECT * FROM go_bids WHERE g_id=%s", (g_id,))
        gb = cur.fetchone()
        if not gb:
            cur.close()
            return jsonify({'ok': False, 'error': 'Bid not found'}), 404

        # Upsert into bid_assign
        cur.execute("SELECT a_id FROM bid_assign WHERE g_id=%s", (g_id,))
        row = cur.fetchone()
        new_stage = team  # team name aligns to stage keys in our system
        if not row:
            try:
                # Full insert (newer schema with in_date, due_date, assignee_email)
                cur.execute(
                    """
                    INSERT INTO bid_assign (g_id, b_name, in_date, due_date, state, scope, type, company, depart, person_name, assignee_email, status, value, revenue)
                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s,%s)
                    """,
                    (
                        gb['g_id'], gb.get('b_name'), gb.get('in_date'), gb.get('due_date'), new_stage,
                        gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'), (first_emp or {}).get('email'),
                        gb.get('scoring', 0), gb.get('revenue', gb.get('scoring', 0))
                    )
                )
            except Exception as e:
                msg = str(e)
                if "Unknown column 'revenue'" in msg:
                    # Retry without 'revenue' column; keep dates and assignee_email
                    try:
                        cur.execute(
                            """
                            INSERT INTO bid_assign (g_id, b_name, in_date, due_date, state, scope, type, company, depart, person_name, assignee_email, status, value)
                            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s)
                            """,
                            (
                                gb['g_id'], gb.get('b_name'), gb.get('in_date'), gb.get('due_date'), new_stage,
                                gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'), (first_emp or {}).get('email'),
                                gb.get('scoring', 0)
                            )
                        )
                    except Exception as e_revenue_dates:
                        # If dates/assignee_email also missing, progressively fallback
                        msg2 = str(e_revenue_dates)
                        if "Unknown column 'in_date'" in msg2 or "Unknown column 'due_date'" in msg2:
                            try:
                                cur.execute(
                                    """
                                    INSERT INTO bid_assign (g_id, b_name, state, scope, type, company, depart, person_name, assignee_email, status, value)
                                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s)
                                    """,
                                    (
                                        gb['g_id'], gb.get('b_name'), new_stage,
                                        gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'), (first_emp or {}).get('email'),
                                        gb.get('scoring', 0)
                                    )
                                )
                            except Exception as e_revenue_dates_email:
                                if "Unknown column 'assignee_email'" in str(e_revenue_dates_email):
                                    cur.execute(
                                        """
                                        INSERT INTO bid_assign (g_id, b_name, state, scope, type, company, depart, person_name, status, value)
                                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s)
                                        """,
                                        (
                                            gb['g_id'], gb.get('b_name'), new_stage,
                                            gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'),
                                            gb.get('scoring', 0)
                                        )
                                    )
                                else:
                                    raise
                        elif "Unknown column 'assignee_email'" in msg2:
                            cur.execute(
                                """
                                INSERT INTO bid_assign (g_id, b_name, in_date, due_date, state, scope, type, company, depart, person_name, status, value)
                                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s)
                                """,
                                (
                                    gb['g_id'], gb.get('b_name'), gb.get('in_date'), gb.get('due_date'), new_stage,
                                    gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'),
                                    gb.get('scoring', 0)
                                )
                            )
                        else:
                            raise
                elif "Unknown column 'in_date'" in msg or "Unknown column 'due_date'" in msg:
                    # Older schema without date columns
                    try:
                        cur.execute(
                            """
                            INSERT INTO bid_assign (g_id, b_name, state, scope, type, company, depart, person_name, assignee_email, status, value, revenue)
                            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s,%s)
                            """,
                            (
                                gb['g_id'], gb.get('b_name'), new_stage,
                                gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'), (first_emp or {}).get('email'),
                                gb.get('scoring', 0), gb.get('revenue', gb.get('scoring', 0))
                            )
                        )
                    except Exception as e2:
                        msg2 = str(e2)
                        if "Unknown column 'revenue'" in msg2:
                            try:
                                cur.execute(
                                    """
                                    INSERT INTO bid_assign (g_id, b_name, state, scope, type, company, depart, person_name, assignee_email, status, value)
                                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s)
                                    """,
                                    (
                                        gb['g_id'], gb.get('b_name'), new_stage,
                                        gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'), (first_emp or {}).get('email'),
                                        gb.get('scoring', 0)
                                    )
                                )
                            except Exception as e2b:
                                if "Unknown column 'assignee_email'" in str(e2b):
                                    cur.execute(
                                        """
                                        INSERT INTO bid_assign (g_id, b_name, state, scope, type, company, depart, person_name, status, value)
                                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s)
                                        """,
                                        (
                                            gb['g_id'], gb.get('b_name'), new_stage,
                                            gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'),
                                            gb.get('scoring', 0)
                                        )
                                    )
                                else:
                                    raise
                        elif "Unknown column 'assignee_email'" in msg2:
                            # Oldest schema: no assignee_email either
                            cur.execute(
                                """
                                INSERT INTO bid_assign (g_id, b_name, state, scope, type, company, depart, person_name, status, value, revenue)
                                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s,%s)
                                """,
                                (
                                    gb['g_id'], gb.get('b_name'), new_stage,
                                    gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'),
                                    gb.get('scoring', 0), gb.get('revenue', gb.get('scoring', 0))
                                )
                            )
                        else:
                            raise
                elif "Unknown column 'assignee_email'" in msg:
                    # No assignee_email but has in_date/due_date
                    try:
                        cur.execute(
                            """
                            INSERT INTO bid_assign (g_id, b_name, in_date, due_date, state, scope, type, company, depart, person_name, status, value, revenue)
                            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s,%s)
                            """,
                            (
                                gb['g_id'], gb.get('b_name'), gb.get('in_date'), gb.get('due_date'), new_stage,
                                gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'),
                                gb.get('scoring', 0), gb.get('revenue', gb.get('scoring', 0))
                            )
                        )
                    except Exception as e3:
                        # If that also fails due to dates missing, drop those too
                        msg3 = str(e3)
                        if "Unknown column 'revenue'" in msg3:
                            try:
                                cur.execute(
                                    """
                                    INSERT INTO bid_assign (g_id, b_name, in_date, due_date, state, scope, type, company, depart, person_name, status, value)
                                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s)
                                    """,
                                    (
                                        gb['g_id'], gb.get('b_name'), gb.get('in_date'), gb.get('due_date'), new_stage,
                                        gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'),
                                        gb.get('scoring', 0)
                                    )
                                )
                            except Exception as e3b:
                                if "Unknown column 'in_date'" in str(e3b) or "Unknown column 'due_date'" in str(e3b):
                                    cur.execute(
                                        """
                                        INSERT INTO bid_assign (g_id, b_name, state, scope, type, company, depart, person_name, status, value)
                                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s)
                                        """,
                                        (
                                            gb['g_id'], gb.get('b_name'), new_stage,
                                            gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'),
                                            gb.get('scoring', 0)
                                        )
                                    )
                                else:
                                    raise
                        elif "Unknown column 'in_date'" in msg3 or "Unknown column 'due_date'" in msg3:
                            cur.execute(
                                """
                                INSERT INTO bid_assign (g_id, b_name, state, scope, type, company, depart, person_name, status, value, revenue)
                                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s,%s)
                                """,
                                (
                                    gb['g_id'], gb.get('b_name'), new_stage,
                                    gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'),
                                    gb.get('scoring', 0), gb.get('revenue', gb.get('scoring', 0))
                                )
                            )
                        else:
                            raise
                else:
                    raise
        else:
            try:
                # Safe fallbacks for legacy schemas:
                # If 'state' and/or 'assignee_email' columns are missing in bid_assign,
                # we progressively retry the UPDATE without them to keep assignment working.
                cur.execute(
                    "UPDATE bid_assign SET depart=%s, person_name=%s, assignee_email=%s, state=%s, status='assigned' WHERE g_id=%s",
                    (team, (first_emp or {}).get('name'), (first_emp or {}).get('email'), new_stage, g_id)
                )
            except Exception as e:
                msg = str(e)
                if "Unknown column 'assignee_email'" in msg:
                    # Try without assignee_email; if state also missing, fall back again
                    try:
                        cur.execute(
                            "UPDATE bid_assign SET depart=%s, person_name=%s, state=%s, status='assigned' WHERE g_id=%s",
                            (team, (first_emp or {}).get('name'), new_stage, g_id)
                        )
                    except Exception as e2:
                        if "Unknown column 'state'" in str(e2):
                            cur.execute(
                                "UPDATE bid_assign SET depart=%s, person_name=%s, status='assigned' WHERE g_id=%s",
                                (team, (first_emp or {}).get('name'), g_id)
                            )
                        else:
                            raise
                elif "Unknown column 'state'" in msg:
                    # Try without state; if assignee_email also missing, fall back again
                    try:
                        cur.execute(
                            "UPDATE bid_assign SET depart=%s, person_name=%s, assignee_email=%s, status='assigned' WHERE g_id=%s",
                            (team, (first_emp or {}).get('name'), (first_emp or {}).get('email'), g_id)
                        )
                    except Exception as e3:
                        if "Unknown column 'assignee_email'" in str(e3):
                            cur.execute(
                                "UPDATE bid_assign SET depart=%s, person_name=%s, status='assigned' WHERE g_id=%s",
                                (team, (first_emp or {}).get('name'), g_id)
                            )
                        else:
                            raise
                else:
                    raise

        # Keep go_bids.state aligned with department
        try:
            cur.execute("UPDATE go_bids SET state=%s WHERE g_id=%s", (new_stage, g_id))
        except Exception as e:
            # Allow running against older schemas that don't yet have go_bids.state
            if "Unknown column 'state'" in str(e):
                pass
            else:
                raise

        # Update mapping table for multiple assignees (team-scoped replace)
        cur.execute(
            "DELETE FROM bid_assignment_members WHERE g_id=%s AND employee_id IN (SELECT id FROM employees WHERE department=%s)",
            (g_id, team)
        )
        if employee_ids:
            for emp_id in employee_ids:
                try:
                    cur.execute(
                        "INSERT IGNORE INTO bid_assignment_members (g_id, employee_id) VALUES (%s,%s)",
                        (g_id, emp_id)
                    )
                except Exception:
                    pass

        # Log + emit
        display_name = (first_emp or {}).get('name') or (f"{len(employee_ids)} member(s)" if employee_ids else 'Unassigned')
        log_action = f"Assigned bid '{gb.get('b_name')}' (ID: {g_id}) to {display_name} [{team}]"
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)", (log_action, current_user.id))
        mysql.connection.commit()

        try:
            socketio.emit('master_update', {
                'log': {'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'), 'action': log_action, 'user_email': current_user.email, 'user_role': getattr(current_user, 'role', '')},
                'assignment': True
            })
        except Exception:
            pass

        cur.close()
        return jsonify({'ok': True})
    except Exception as e:
        mysql.connection.rollback()
        try:
            cur.close()
        except Exception:
            pass
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/api/team/<team>/bids/<int:g_id>/tasks')
@login_required
def api_bid_tasks(team, g_id):
    """API endpoint to get tasks for a specific bid - only show tasks created by current team"""
    cur = mysql.connection.cursor(DictCursor)
    
    # Allow access if bid is in this team OR there are tasks for this team's stage
    try:
        cur.execute("SELECT state FROM go_bids WHERE g_id = %s", (g_id,))
        bid = cur.fetchone()
    except Exception as e:
        if "Unknown column 'state'" in str(e):
            bid = {'state': team}
        else:
            raise
    cur.execute("SELECT 1 FROM bid_checklists WHERE g_id=%s AND LOWER(COALESCE(stage,''))=%s LIMIT 1", (g_id, team))
    has_team_tasks = cur.fetchone() is not None
    if not bid and not has_team_tasks:
        cur.close();
        return jsonify({'tasks': []})
    if bid and bid.get('state') != team and not has_team_tasks:
        cur.close();
        return jsonify({'tasks': []})
    
    # Get checklist items for this bid - show current team's active tasks and their archived tasks
    roles_for_team = {
        'business': ('business dev', 'business', 'bdm'),
        'design': ('design',),
        'operations': ('operations',),
        'engineer': ('site manager', 'engineer'),
    }
    acceptable_roles = roles_for_team.get(team, (team,))
    placeholders = ','.join(['%s'] * len(acceptable_roles))
    cur.execute(f"""
        SELECT bc.*, e.name as assigned_employee_name, e.email as employee_email, e.department
        FROM bid_checklists bc
        LEFT JOIN employees e ON bc.assigned_to = e.id
        LEFT JOIN users u ON bc.created_by = u.id
        WHERE bc.g_id = %s 
        AND (
            (u.role IN ({placeholders}) AND bc.team_archive IS NULL) OR  -- Current team's active tasks
            (bc.team_archive = %s)  -- This team's archived tasks
        )
        ORDER BY bc.priority DESC, bc.created_at ASC
    """, (g_id, *acceptable_roles, team))
    tasks = cur.fetchall()
    # Add attachment_url for easier consumption in UI
    try:
        for t in tasks:
            ap = (t.get('attachment_path') or '').strip() if isinstance(t, dict) else None
            t['attachment_url'] = url_for('get_task_attachment', task_id=t.get('id')) if ap else None
    except Exception:
        pass

    # Emit per-stage map for this bid so master dashboard reflects progress
    try:
        def _pct(rows):
            if not rows:
                return 0
            t = len(rows)
            d = len([r for r in rows if (r.get('status') or '').lower() == 'completed'])
            return int(round((d / max(1, t)) * 100))

        def _q(role_filter):
            cur.execute(f"SELECT status FROM bid_checklists WHERE g_id=%s AND created_by IN (SELECT id FROM users WHERE role {role_filter})", (g_id,))
            return cur.fetchall()

        spm = {
            'business': _pct(_q("='business dev'")),
            'design': _pct(_q("='design'")),
            'operations': _pct(_q("='operations'")),
            'engineer': _pct(_q("IN ('site manager','engineer')")),
        }
        socketio.emit('master_update', {
            'summary': {
                'bid_id': g_id,
                'work_progress_pct': spm.get(team, 0),
                'project_status': 'ongoing',
                'work_status': f'{team.title()} tasks updated',
                'stage_progress_map': spm
            }
        })
    except Exception:
        pass

    cur.close()
    return jsonify({'tasks': tasks})

# --- Simple API to update a bid summary (notes) ---
@app.route('/api/bids/<int:g_id>/summary', methods=['POST'])
@login_required
def api_update_bid_summary(g_id):
    try:
        summary = (request.form.get('summary') or '').strip()
        cur = mysql.connection.cursor()
        cur.execute("UPDATE go_bids SET summary=%s WHERE g_id=%s", (summary, g_id))
        mysql.connection.commit()
        cur.close()
        return jsonify({'ok': True})
    except Exception as e:
        try:
            mysql.connection.rollback()
        except Exception:
            pass
        return jsonify({'ok': False, 'error': str(e)}), 500

# Update bid dates (start/due)
@app.route('/api/bids/<int:g_id>/dates', methods=['POST'])
@login_required
def api_update_bid_dates(g_id):
    try:
        
        due_date = request.form.get('due_date')
        cur = mysql.connection.cursor()
        # Build dynamic set
        fields = []
        values = []
        if due_date:
            fields.append('due_date=%s'); values.append(due_date)
        if not fields:
            return jsonify({'ok': False, 'error': 'No fields provided'}), 400
        values.append(g_id)
        cur.execute(f"UPDATE go_bids SET {', '.join(fields)} WHERE g_id=%s", tuple(values))
        mysql.connection.commit()
        cur.close()
        return jsonify({'ok': True})
    except Exception as e:
        try:
            mysql.connection.rollback()
        except Exception:
            pass
        return jsonify({'ok': False, 'error': str(e)}), 500

# Comments CRUD (list + add)
@app.route('/api/bids/<int:g_id>/comments', methods=['GET', 'POST'])
@login_required
def api_bid_comments(g_id):
    if request.method == 'GET':
        cur = mysql.connection.cursor(DictCursor)
        cur.execute(
            """
            SELECT c.id, c.g_id, c.comment_text, c.created_at, u.email AS user_email
            FROM bid_comments c
            LEFT JOIN users u ON u.id = c.user_id
            WHERE c.g_id=%s
            ORDER BY c.created_at DESC
            """,
            (g_id,)
        )
        rows = cur.fetchall()
        cur.close()
        return jsonify({'comments': rows})
    else:
        text = (request.form.get('comment') or '').strip()
        if not text:
            return jsonify({'ok': False, 'error': 'Empty comment'}), 400
        try:
            cur = mysql.connection.cursor()
            uid = getattr(current_user, 'id', None)
            cur.execute("INSERT INTO bid_comments (g_id, user_id, comment_text) VALUES (%s,%s,%s)", (g_id, uid, text))
            mysql.connection.commit()
            cur.close()
            return jsonify({'ok': True})
        except Exception as e:
            try:
                mysql.connection.rollback()
            except Exception:
                pass
            return jsonify({'ok': False, 'error': str(e)}), 500
@app.route('/database-management/delete/<int:row_id>')
@login_required
def dbm_delete(row_id):
    if not current_user.is_admin:
        return "Access Denied", 403
    table = request.args.get('table')
    if not table:
        return redirect(url_for('database_management'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        cur.execute(f"DESCRIBE `{table}`")
        cols = cur.fetchall()
        pk = next((c['Field'] for c in cols if c['Key'] == 'PRI'), cols[0]['Field'])
        
        # Handle foreign key constraints for specific tables
        if table == 'go_bids':
            # Delete related records first to avoid foreign key constraint errors
            cur.execute("DELETE FROM bid_checklists WHERE g_id=%s", (row_id,))
            cur.execute("DELETE FROM bid_stage_exclusions WHERE g_id=%s", (row_id,))
            cur.execute("DELETE FROM bid_custom_stages WHERE g_id=%s", (row_id,))
            
            # Delete from bid_assign (which may have win_lost_results, won_bids_result, work_progress_status cascading)
            cur.execute("SELECT a_id FROM bid_assign WHERE g_id=%s", (row_id,))
            assign_ids = [row['a_id'] for row in cur.fetchall()]
            
            for a_id in assign_ids:
                # Get w_id from win_lost_results
                cur.execute("SELECT w_id FROM win_lost_results WHERE a_id=%s", (a_id,))
                w_ids = [row['w_id'] for row in cur.fetchall()]
                
                for w_id in w_ids:
                    # Get won_id from won_bids_result
                    cur.execute("SELECT won_id FROM won_bids_result WHERE w_id=%s", (w_id,))
                    won_ids = [row['won_id'] for row in cur.fetchall()]
                    
                    # Delete work_progress_status
                    for won_id in won_ids:
                        cur.execute("DELETE FROM work_progress_status WHERE won_id=%s", (won_id,))
                    
                    # Delete won_bids_result
                    cur.execute("DELETE FROM won_bids_result WHERE w_id=%s", (w_id,))
                
                # Delete win_lost_results
                cur.execute("DELETE FROM win_lost_results WHERE a_id=%s", (a_id,))
            
            # Delete bid_assign
            cur.execute("DELETE FROM bid_assign WHERE g_id=%s", (row_id,))
        
        elif table == 'bid_incoming':
            # Find related go_bids and delete them (which will cascade)
            cur.execute("SELECT g_id FROM go_bids WHERE id=%s", (row_id,))
            g_ids = [row['g_id'] for row in cur.fetchall()]
            
            for g_id in g_ids:
                # Recursively delete go_bids using the same logic
                cur.execute("DELETE FROM bid_checklists WHERE g_id=%s", (g_id,))
                cur.execute("DELETE FROM bid_stage_exclusions WHERE g_id=%s", (g_id,))
                cur.execute("DELETE FROM bid_custom_stages WHERE g_id=%s", (g_id,))
                
                cur.execute("SELECT a_id FROM bid_assign WHERE g_id=%s", (g_id,))
                assign_ids = [row['a_id'] for row in cur.fetchall()]
                
                for a_id in assign_ids:
                    cur.execute("SELECT w_id FROM win_lost_results WHERE a_id=%s", (a_id,))
                    w_ids = [row['w_id'] for row in cur.fetchall()]
                    
                    for w_id in w_ids:
                        cur.execute("SELECT won_id FROM won_bids_result WHERE w_id=%s", (w_id,))
                        won_ids = [row['won_id'] for row in cur.fetchall()]
                        
                        for won_id in won_ids:
                            cur.execute("DELETE FROM work_progress_status WHERE won_id=%s", (won_id,))
                        
                        cur.execute("DELETE FROM won_bids_result WHERE w_id=%s", (w_id,))
                    
                    cur.execute("DELETE FROM win_lost_results WHERE a_id=%s", (a_id,))
                
                cur.execute("DELETE FROM bid_assign WHERE g_id=%s", (g_id,))
                cur.execute("DELETE FROM go_bids WHERE g_id=%s", (g_id,))
        
        elif table == 'bid_assign':
            # Get w_id from win_lost_results
            cur.execute("SELECT w_id FROM win_lost_results WHERE a_id=%s", (row_id,))
            w_ids = [row['w_id'] for row in cur.fetchall()]
            
            for w_id in w_ids:
                # Get won_id from won_bids_result
                cur.execute("SELECT won_id FROM won_bids_result WHERE w_id=%s", (w_id,))
                won_ids = [row['won_id'] for row in cur.fetchall()]
                
                # Delete work_progress_status
                for won_id in won_ids:
                    cur.execute("DELETE FROM work_progress_status WHERE won_id=%s", (won_id,))
                
                # Delete won_bids_result
                cur.execute("DELETE FROM won_bids_result WHERE w_id=%s", (w_id,))
            
            # Delete win_lost_results
            cur.execute("DELETE FROM win_lost_results WHERE a_id=%s", (row_id,))
        
        elif table == 'win_lost_results':
            # Get won_id from won_bids_result
            cur.execute("SELECT won_id FROM won_bids_result WHERE w_id=%s", (row_id,))
            won_ids = [row['won_id'] for row in cur.fetchall()]
            
            # Delete work_progress_status
            for won_id in won_ids:
                cur.execute("DELETE FROM work_progress_status WHERE won_id=%s", (won_id,))
            
            # Delete won_bids_result
            cur.execute("DELETE FROM won_bids_result WHERE w_id=%s", (row_id,))
        
        elif table == 'won_bids_result':
            # Delete work_progress_status
            cur.execute("DELETE FROM work_progress_status WHERE won_id=%s", (row_id,))
        
        # Now delete the main record
        cur.execute(f"DELETE FROM `{table}` WHERE `{pk}`=%s", (row_id,))
        mysql.connection.commit()
        flash(f'Record deleted successfully from {table}!', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Error deleting record: {str(e)}', 'error')
    finally:
        cur.close()
    return redirect(url_for('database_management', table=table))

@app.route('/database-management/drop')
@login_required
def dbm_drop():
    if not current_user.is_admin:
        return "Access Denied", 403
    table = request.args.get('table')
    if not table:
        return redirect(url_for('database_management'))
    cur = mysql.connection.cursor()
    try:
        cur.execute(f"DROP TABLE IF EXISTS `{table}`")
        mysql.connection.commit()
    finally:
        cur.close()
    return redirect(url_for('database_management'))

@app.route('/database-management/export')
@login_required
def dbm_export():
    if not current_user.is_admin:
        return "Access Denied", 403
    table = request.args.get('table')
    if not table:
        return redirect(url_for('database_management'))
    try:
        import pandas as pd
        import io
        cur = mysql.connection.cursor(DictCursor)
        cur.execute(f"SELECT * FROM `{table}`")
        rows = cur.fetchall()
        cur.close()
        df = pd.DataFrame(rows)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name=table[:31])
        output.seek(0)
        filename = f"{table}.xlsx"
        return send_file(output, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    except Exception as e:
        flash(f'Export failed: {str(e)}', 'error')
        return redirect(url_for('database_management', table=table))

@app.route('/database-management/import', methods=['POST'])
@login_required
def dbm_import():
    if not current_user.is_admin:
        return "Access Denied", 403
    table = request.form.get('table')
    file = request.files.get('excel_file')
    if not table or not file or file.filename == '':
        flash('Please choose a table and an Excel file.', 'error')
        return redirect(url_for('database_management', table=table or ''))
    try:
        import pandas as pd
        import io
        df = pd.read_excel(io.BytesIO(file.read()))
        cur = mysql.connection.cursor(DictCursor)
        # Fetch columns and ignore primary key
        cur.execute(f"DESCRIBE `{table}`")
        cols = cur.fetchall()
        non_pk_cols = [c['Field'] for c in cols if c['Key'] != 'PRI']
        # Filter dataframe to only known columns
        df = df[[c for c in df.columns if c in non_pk_cols]]
        if df.empty:
            flash('No matching columns found in Excel file for this table.', 'error')
            cur.close()
            return redirect(url_for('database_management', table=table))
        placeholders = ','.join(['%s'] * len(df.columns))
        fields_sql = ','.join([f"`{c}`" for c in df.columns])
        sql = f"INSERT INTO `{table}` ({fields_sql}) VALUES ({placeholders})"
        for _, row in df.iterrows():
            cur.execute(sql, [None if pd.isna(v) else v for v in row.tolist()])
        mysql.connection.commit()
        cur.close()
        flash(f'Successfully imported {len(df)} rows into `{table}`.', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Import failed: {str(e)}', 'error')
    return redirect(url_for('database_management', table=table))

@app.route('/database-management/update-decision', methods=['POST'])
@login_required
def dbm_update_decision():
    try:
        payload = request.get_json(force=True) or {}
        table = (payload.get('table') or '').strip().lower()
        row_id = int(payload.get('row_id') or 0)
        decision = (payload.get('decision') or '').strip().upper()

        if table not in ('bid_incoming', 'go_bids'):
            return jsonify({'success': False, 'error': 'Invalid table'}), 400
        if row_id <= 0:
            return jsonify({'success': False, 'error': 'Invalid row id'}), 400
        # normalize decision
        mapping = {
            'GO': 'GO', 'NO GO': 'NO-GO', 'NO-GO': 'NO-GO', 'NOGO': 'NO-GO',
            'WON': 'WON', 'LOST': 'LOST', 'PENDING': 'PENDING'
        }
        decision = mapping.get(decision, 'PENDING')

        cur = mysql.connection.cursor(DictCursor)
        if table == 'bid_incoming':
            cur.execute("UPDATE bid_incoming SET decision=%s WHERE id=%s", (decision, row_id))
        else:
            cur.execute("UPDATE go_bids SET decision=%s WHERE g_id=%s", (decision, row_id))
        mysql.connection.commit()
        cur.close()

        # Keep go_bids synchronized with bid_incoming GO decisions
        try:
            sync_go_bids()
        except Exception as sync_err:
            # Do not fail the request if sync encounters an issue; just log it.
            print(f"Warning: sync_go_bids failed after decision update: {sync_err}")

        log_write('dbm_update_decision', f"table={table} id={row_id} decision={decision}")
        return jsonify({'success': True, 'decision': decision})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

def generate_proposal_with_local_rag(rfp_text, company_name, project_title, deadline, sections_outline, user_email=None):
    """
    Generate proposal content using local RAG approach without API calls.
    Uses template-based generation with contextual retrieval from company knowledge bases.
    """
    try:
        import os
        import json
        import re
        from datetime import datetime
        from sentence_transformers import SentenceTransformer
        import numpy as np

        # Initialize embedding model (reuse from RFP analyzer)
        device = "cuda" if torch.cuda.is_available() else "cpu"
        try:
            embedder = SentenceTransformer("BAAI/bge-large-en-v1.5", device=device)
        except:
            embedder = SentenceTransformer("all-MiniLM-L6-v2", device=device)

        # Load company knowledge base if available
        company_texts = []
        company_db_path = os.path.join("company_db", f"{company_name}.json")
        if os.path.exists(company_db_path):
            try:
                with open(company_db_path, 'r', encoding='utf-8') as f:
                    company_data = json.load(f)
                    company_texts = company_data.get('texts', [])
            except Exception as e:
                print(f"Warning: Could not load company knowledge base: {e}")

        # Extract key RFP information using pattern matching
        rfp_info = extract_rfp_patterns(rfp_text)

        # Generate content for each section using templates and retrieval
        sections_content = {}

        for section in sections_outline:
            section_title = section.get('title', 'Section')
            guidance = section.get('guidance', '')
            sub_categories = section.get('sub_categories', [])

            # Retrieve relevant context from RFP and company knowledge
            relevant_context = retrieve_relevant_context(
                rfp_text, company_texts, f"{section_title}: {guidance}", embedder
            )

            # Generate section content using templates
            section_content = generate_section_content(
                section_title, guidance, relevant_context, rfp_info,
                company_name, project_title, deadline
            )

            # Handle sub-sections
            sub_sections = []
            for sub_cat in sub_categories:
                sub_title = sub_cat.get('title', '')
                if sub_title:
                    sub_guidance = sub_cat.get('guidance', f'Provide details for {sub_title}')
                    sub_context = retrieve_relevant_context(
                        rfp_text, company_texts, f"{sub_title}: {sub_guidance}", embedder
                    )
                    sub_content = generate_section_content(
                        sub_title, sub_guidance, sub_context, rfp_info,
                        company_name, project_title, deadline
                    )
                    sub_sections.append({
                        'title': sub_title,
                        'content': sub_content
                    })

            sections_content[section_title] = {
                'title': section_title,
                'content': section_content,
                'sub_sections': sub_sections
            }

        return {'sections': list(sections_content.values())}

    except Exception as e:
        print(f"Error in local RAG generation: {e}")
        import traceback
        traceback.print_exc()
        # Fallback to basic template
        return {'sections': generate_fallback_proposal_sections(sections_outline, company_name, project_title)}


def extract_rfp_patterns(rfp_text: str) -> Dict:
    """Extract key patterns and information from RFP text"""
    patterns = {
        'project_requirements': [],
        'technical_specs': [],
        'deadlines': [],
        'budget_info': [],
        'compliance_items': [],
        'evaluation_criteria': []
    }

    # Extract project requirements
    req_patterns = [
        r'(?:requirements?|must|shall|required).*?([^\n]{20,200})',
        r'(?:scope).*?([^\n]{20,200})',
        r'(?:deliverables).*?([^\n]{20,200})'
    ]

    for pattern in req_patterns:
        matches = re.findall(pattern, rfp_text, re.IGNORECASE | re.DOTALL)
        patterns['project_requirements'].extend(matches[:5])  # Limit to top 5

    # Extract technical specifications
    tech_patterns = [
        r'(?:technical|specifications|specs).*?([^\n]{20,150})',
        r'(?:system|software|hardware).*?([^\n]{20,150})'
    ]

    for pattern in tech_patterns:
        matches = re.findall(pattern, rfp_text, re.IGNORECASE | re.DOTALL)
        patterns['technical_specs'].extend(matches[:3])

    # Extract deadlines
    deadline_patterns = [
        r'(?:deadline|due).*?(\d{1,2}[-/]\d{1,2}[-/]\d{2,4}|\w+\s+\d{1,2},?\s+\d{4})',
        r'(?:submission).*?(\d{1,2}[-/]\d{1,2}[-/]\d{2,4}|\w+\s+\d{1,2},?\s+\d{4})'
    ]

    for pattern in deadline_patterns:
        matches = re.findall(pattern, rfp_text, re.IGNORECASE)
        patterns['deadlines'].extend(matches[:2])

    return patterns


def retrieve_relevant_context(rfp_text: str, company_texts: List[str], query: str, embedder) -> str:
    """Retrieve relevant context using semantic similarity"""
    try:
        # Prepare texts for retrieval
        all_texts = [rfp_text] + company_texts
        if not all_texts:
            return ""

        # Encode query and texts
        query_embedding = embedder.encode([query], convert_to_numpy=True)
        text_embeddings = embedder.encode(all_texts, convert_to_numpy=True)

        # Calculate similarities
        similarities = np.dot(text_embeddings, query_embedding.T).flatten()
        similarities = similarities / (np.linalg.norm(text_embeddings, axis=1) * np.linalg.norm(query_embedding))

        # Get top 3 most similar texts
        top_indices = np.argsort(similarities)[-3:][::-1]
        relevant_texts = [all_texts[i] for i in top_indices if similarities[i] > 0.3]

        # Limit context length
        context = "\n\n".join(relevant_texts)
        return context[:2000] if len(context) > 2000 else context

    except Exception as e:
        print(f"Error in context retrieval: {e}")
        return rfp_text[:1000]


def generate_section_content(section_title: str, guidance: str, context: str, rfp_info: Dict,
                           company_name: str, project_title: str, deadline: str) -> List[str]:
    """Generate content for a proposal section using templates and context"""

    # Template-based content generation
    templates = {
        'Executive Summary': [
            f"{company_name} is pleased to submit this comprehensive proposal for the {project_title} project.",
            f"Our team brings extensive experience and proven capabilities to deliver exceptional results that meet or exceed all requirements.",
            f"This proposal demonstrates our thorough understanding of the project scope, technical requirements, and compliance needs.",
            f"With our proven track record and commitment to excellence, we are positioned to successfully deliver this project on time and within budget.",
            f"We look forward to the opportunity to partner with you on this important initiative."
        ],

        'Company Profile': [
            f"{company_name} is a leading provider of professional services with extensive experience in delivering complex projects.",
            f"Our team consists of qualified professionals with specialized expertise in project management, technical implementation, and quality assurance.",
            f"We maintain all necessary certifications, licenses, and insurance coverage required for this type of work.",
            f"Our commitment to safety, quality, and customer satisfaction has earned us a reputation for excellence in the industry.",
            f"We have successfully completed numerous similar projects, demonstrating our capability to handle complex requirements and deliver results."
        ],

        'Technical Approach': [
            f"Our technical approach is designed to meet all specified requirements while optimizing performance, reliability, and maintainability.",
            f"We will utilize industry-standard methodologies and best practices throughout the project lifecycle.",
            f"Our solution incorporates robust quality control measures and comprehensive testing protocols.",
            f"We have extensive experience implementing similar technical solutions with a proven track record of success.",
            f"Our technical team includes certified professionals with specialized expertise in the required technologies and methodologies."
        ],

        'Project Management': [
            f"Our project management methodology ensures clear communication, rigorous quality control, and on-time delivery.",
            f"We will establish regular progress meetings, detailed reporting, and proactive issue resolution processes.",
            f"Our experienced project managers will oversee all aspects of implementation with dedicated resources.",
            f"We utilize industry-standard project management tools and techniques to ensure successful project execution.",
            f"Risk management and change control procedures will be implemented to address potential challenges proactively."
        ],

        'Price Proposal Form (Attachment A)': [
            f"Our pricing structure is designed to provide exceptional value while ensuring project profitability and sustainability.",
            f"We have carefully analyzed all requirements and developed a comprehensive cost estimate based on industry standards.",
            f"Our pricing includes all necessary labor, materials, equipment, and overhead costs required for successful project completion.",
            f"We offer competitive pricing with transparent cost breakdowns and detailed line item descriptions.",
            f"Our cost proposal reflects our commitment to delivering high-quality results within the established budget parameters."
        ],

        'Compliance Matrix & Deliverables': [
            f"{company_name} is fully committed to meeting all contractual requirements and compliance standards.",
            f"We have reviewed all specifications, terms, and conditions and confirm our ability to meet all mandatory requirements.",
            f"Our compliance approach includes comprehensive documentation, regular audits, and quality assurance processes.",
            f"We maintain all necessary certifications and can provide evidence of compliance with all applicable standards.",
            f"Our team is experienced in working with government agencies and understands the importance of regulatory compliance."
        ],

        'Project Team': [
            f"Our project team consists of experienced professionals with specialized skills and proven track records.",
            f"Key team members include certified project managers, technical specialists, and quality assurance experts.",
            f"All team members have undergone appropriate background checks and maintain required security clearances.",
            f"We provide ongoing training and professional development to ensure our team remains current with industry best practices.",
            f"Our team structure ensures appropriate staffing levels and expertise allocation throughout the project lifecycle."
        ]
    }

    # Get base template content
    base_content = templates.get(section_title, [
        f"This section provides detailed information about {section_title.lower()}.",
        f"Our approach to {section_title.lower()} is based on industry best practices and proven methodologies.",
        f"We have extensive experience in delivering {section_title.lower()} services that meet client expectations.",
        f"Our team is fully prepared to implement all aspects of {section_title.lower()} as specified in the requirements."
    ])

    # Enhance with RFP-specific information
    enhanced_content = []
    for paragraph in base_content:
        enhanced_paragraph = paragraph

        # Add project-specific details
        if project_title and project_title != "Project":
            enhanced_paragraph = enhanced_paragraph.replace("this project", f"the {project_title} project")

        # Add deadline information if relevant
        if deadline and "schedule" in section_title.lower():
            enhanced_paragraph += f" We are committed to meeting all project deadlines, including the submission deadline of {deadline}."

        # Add compliance information
        if "compliance" in section_title.lower() and rfp_info.get('compliance_items'):
            enhanced_paragraph += f" We will ensure compliance with all specified requirements including {', '.join(rfp_info['compliance_items'][:2])}."

        enhanced_content.append(enhanced_paragraph)

    return enhanced_content


def generate_fallback_proposal_sections(sections_outline, company_name: str, project_title: str) -> List[Dict]:
    """Generate basic fallback proposal sections when RAG fails"""
    fallback_sections = []

    for section in sections_outline:
        title = section.get('title', 'Section')
        content = [
            f"{company_name} is pleased to submit this proposal section for {title}.",
            f"Our team has reviewed all requirements and is prepared to deliver comprehensive solutions.",
            f"We bring extensive experience and proven capabilities to ensure project success.",
            f"Our approach incorporates industry best practices and quality assurance measures.",
            f"We are committed to meeting all project requirements and delivering exceptional results."
        ]

        fallback_sections.append({
            'title': title,
            'content': content,
            'sub_sections': []
        })

    return fallback_sections


def generate_complete_proposal(pdf_path, company, template, output_path, bid_name=None, sections_outline=None, user_email=None):
    """
    Generate a complete structured proposal document from RFP PDF using AI.
    Creates a comprehensive proposal with all required sections.
    """
    try:
        import os
        from datetime import datetime
        from rfp_analyzer_routes import extract_page_texts, ollama_json, trim_text_to_token_limit
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.enum.table import WD_TABLE_ALIGNMENT
        import re
        
        # Read PDF file
        with open(pdf_path, 'rb') as pdf_file:
            pages = extract_page_texts(pdf_file)
        
        if not pages or len(pages) == 0:
            raise Exception("Could not extract text from PDF")
        
        full_text = "\n\n".join(pages)
        
        # Extract key information from RFP
        rfp_text_limited = trim_text_to_token_limit(full_text, 8000)  # Limit for prompt
        
        # Extract project title
        project_title = bid_name or "Project"
        title_match = re.search(r'(?i)(?:project|title|proposal|rfp)[:\s]+([A-Z][^\n]{10,100})', full_text[:2000])
        if title_match:
            project_title = title_match.group(1).strip()
        
        # Extract deadline
        deadline = None
        deadline_patterns = [
            r'(?i)(?:deadline|due\s+date|submission\s+date)[:\s]+([A-Za-z]+\s+\d{1,2},?\s+\d{4})',
            r'(?i)(?:deadline|due\s+date)[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
            r'(?i)August\s+31,?\s+2026',
        ]
        for pattern in deadline_patterns:
            match = re.search(pattern, full_text[:5000])
            if match:
                deadline = match.group(1) if match.groups() else match.group(0)
                break
        
        # Build dynamic proposal prompt based on provided sections
        company_display = company or "Our Company"
        project_label = project_title or bid_name or "Project"

        if not sections_outline:
            sections_outline = [
                {'title': 'Executive Summary', 'guidance': 'Summarize the opportunity, objectives, and how the proposed solution delivers value.'},
                {'title': 'Company Profile', 'guidance': 'Describe the company background, capabilities, and relevant experience.'},
                {'title': 'Technical Approach & Scope', 'guidance': 'Detail the technical solution, methodology, schedule, and risk mitigation strategies.'},
                {'title': 'Project Management & Schedule', 'guidance': 'Outline management structure, milestones, resource plan, and communication strategy.'},
                {'title': 'Price Proposal Form (Attachment A)', 'guidance': 'Explain pricing assumptions, cost elements, and any exclusions.'},
                {'title': 'Compliance Matrix & Deliverables', 'guidance': 'Summarize compliance approach, required deliverables, and submission checklist.'},
                {'title': 'Project Team', 'guidance': '.'},
                {'title': 'Corporate Qualifications', 'guidance': ''}
            ]

        section_prompt_lines = []
        for sec in sections_outline:
            guidance = sec.get('guidance') or 'Provide a comprehensive, proposal-ready narrative aligned with this heading.'
            sub_titles = ', '.join(sub.get('title') for sub in sec.get('sub_categories') or [] if sub.get('title'))
            if sub_titles:
                guidance += f" Include coverage of: {sub_titles}."
            section_prompt_lines.append(f'- {sec.get("title", "Section")}: {guidance}')
        sections_prompt = "\n".join(section_prompt_lines)

        proposal_prompt = f"""You are a professional proposal generation assistant specialized in government and infrastructure RFPs. Analyze the uploaded RFP document in full detail and automatically create a draft proposal narrative that follows the section headers provided.

RFP Document Content:
{rfp_text_limited}

Company Name: {company_display}
Project Title: {project_label}
Deadline: {deadline or 'To be determined'}

Sections to include (in this exact order):
{sections_prompt}

Return JSON with the following structure:
{{
  "sections": [
    {{
      "title": "Exact Section Title",
      "content": [
        "Paragraph 1",
        "Paragraph 2",
        "Paragraph 3"
      ],
      "sub_sections": [
        {{
          "title": "Optional Subheading",
          "content": [
            "Paragraph 1",
            "Paragraph 2"
          ]
        }}
      ]
    }}
  ]
}}

Ensure each section contains multiple detailed paragraphs tailored to the RFP requirements, referencing specific compliance, schedule, and technical expectations whenever possible."""

        # Generate proposal content using Local RAG (no API calls)
        print("Generating proposal content with Local RAG...")
        proposal_result = generate_proposal_with_local_rag(
            rfp_text_limited, company_display, project_label, deadline, sections_outline, user_email
        )

        def normalize_title(value):
            return re.sub(r'\s+', ' ', str(value).strip().lower())

        sections_output = []
        if isinstance(proposal_result, dict):
            if isinstance(proposal_result.get('sections'), list):
                sections_output = proposal_result['sections']
            else:
                # Interpret as dictionary of section -> content
                sections_output = [
                    {'title': key, 'content': value}
                    for key, value in proposal_result.items()
                    if key not in ('raw', 'error')
                ]

        sections_map = {}
        for entry in sections_output:
            title = entry.get('title')
            if not title:
                continue
            sections_map[normalize_title(title)] = entry

        normalized_sections = []
        for sec in sections_outline:
            title = sec.get('title', 'Section').strip()
            entry = sections_map.get(normalize_title(title))

            content = []
            sub_sections = []

            if entry:
                entry_content = entry.get('content')
                if isinstance(entry_content, str):
                    content = [entry_content]
                elif isinstance(entry_content, list):
                    content = [item for item in entry_content if item]
                elif isinstance(entry_content, dict):
                    # Flatten dict content
                    content = [v for v in entry_content.values() if v]
                sub_sections = entry.get('sub_sections') or []

            if not content:
                fallback_content = sec.get('content') or []
                if isinstance(fallback_content, str):
                    fallback_content = [fallback_content]
                if not fallback_content and sec.get('guidance'):
                    fallback_content = [sec['guidance']]
                if not fallback_content:
                    fallback_content = ['Content will be added once capture information is finalized.']
                content = fallback_content

            normalized_subsections = []
            for sub in sub_sections:
                sub_title = sub.get('title')
                if not sub_title:
                    continue
                sub_content = sub.get('content')
                if isinstance(sub_content, str):
                    sub_content = [sub_content]
                elif isinstance(sub_content, list):
                    sub_content = [item for item in sub_content if item]
                else:
                    sub_content = []
                if not sub_content:
                    sub_content = [f'Additional guidance for {sub_title} will be incorporated.']
                normalized_subsections.append({
                    'title': sub_title,
                    'content': sub_content
                })

            normalized_sections.append({
                'title': title,
                'content': content,
                'sub_sections': normalized_subsections
            })

        # Create Word document using brand-aware template selection
        try:
            chosen_path = choose_company_template(company_display)
            if chosen_path:
                doc = Document(chosen_path)
            else:
                doc = Document()
        except Exception:
            doc = Document()

        # Title Page
        doc.add_heading(project_label or 'Proposal Draft', 0)
        doc.add_paragraph(f'Prepared for: {project_label}')
        doc.add_paragraph(f'Prepared by: {company_display}')
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        doc.add_paragraph()
        doc.add_page_break()

        # Table of Contents
        doc.add_heading('Table of Contents', 1)
        for idx, section in enumerate(normalized_sections, 1):
            doc.add_paragraph(f'{idx}. {section["title"]}')
        doc.add_page_break()

        # Sections
        for idx, section in enumerate(normalized_sections):
            doc.add_heading(section['title'], level=1)
            for paragraph in section.get('content', []):
                for part in str(paragraph).split('\n\n'):
                    if part.strip():
                        doc.add_paragraph(part.strip())

            for sub_section in section.get('sub_sections', []):
                sub_title = sub_section.get('title')
                if sub_title:
                    doc.add_heading(sub_title, level=2)
                for paragraph in sub_section.get('content', []):
                    for part in str(paragraph).split('\n\n'):
                        if part.strip():
                            doc.add_paragraph(part.strip())

            if idx != len(normalized_sections) - 1:
                doc.add_page_break()

        # IKIO-only appendices: include data from IKIO-specific tables when BD user from ikioledlighting
        try:
            viewer_email = (user_email or '').lower()
            is_ikio_user = viewer_email.endswith('@ikioledlighting.com')
            is_bd_user = False
            if is_ikio_user:
                try:
                    curx = mysql.connection.cursor(DictCursor)
                    # Try to determine department from employees/users
                    curx.execute("SELECT department FROM employees WHERE LOWER(email)=%s LIMIT 1", (viewer_email,))
                    row_dep = curx.fetchone()
                    dep_val = (row_dep or {}).get('department') or ''
                    if dep_val:
                        dep_val_l = str(dep_val).strip().lower()
                        is_bd_user = ('bd' in dep_val_l) or ('business' in dep_val_l and 'dev' in dep_val_l)
                    else:
                        # Fall back to users.role if department not found
                        curx.execute("SELECT role FROM users WHERE LOWER(email)=%s LIMIT 1", (viewer_email,))
                        row_role = curx.fetchone()
                        role_val = (row_role or {}).get('role') or ''
                        role_l = str(role_val).strip().lower()
                        is_bd_user = ('business' in role_l) or ('bd' in role_l) or ('bde' in role_l)
                except Exception:
                    is_bd_user = False
                if is_bd_user:
                    def fetch_rows(table_name: str, limit: int = 12) -> list[dict]:
                        try:
                            curx.execute(f"SELECT * FROM `{table_name}` LIMIT %s", (limit,))
                            return curx.fetchall() or []
                        except Exception:
                            return []
                    def choose_cols(rows: list[dict], max_cols: int = 6) -> list[str]:
                        if not rows:
                            return []
                        # Prefer commonly useful columns if present
                        preferred = ['project_name','bid_name','title','scope','client','owner','status','value','amount','date','due_date','role','name','designation','qualification','experience','email','phone']
                        keys = list(rows[0].keys())
                        cols = [c for c in preferred if c in keys]
                        # Append additional keys (skip internal ids/paths)
                        extra = [k for k in keys if k not in cols and k.lower() not in ('id','created_at','updated_at')]
                        cols = (cols + extra)[:max_cols]
                        return cols
                    def add_table_section(title: str, rows: list[dict]):
                        if not rows:
                            return
                        doc.add_page_break()
                        doc.add_heading(title, level=1)
                        columns = choose_cols(rows)
                        if not columns:
                            # Fallback to simple bullet list of first 10 rows
                            for r in rows:
                                line = "; ".join(f"{k}: {v}" for k, v in r.items() if k.lower() not in ('id','created_at','updated_at') and v not in (None,''))
                                if line:
                                    doc.add_paragraph(line)
                            return
                        tbl = doc.add_table(rows=1, cols=len(columns))
                        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                        hdr_cells = tbl.rows[0].cells
                        for i, col_name in enumerate(columns):
                            hdr_cells[i].text = col_name.replace('_',' ').title()
                        for r in rows:
                            row_cells = tbl.add_row().cells
                            for i, col_name in enumerate(columns):
                                val = r.get(col_name)
                                row_cells[i].text = '' if val is None else str(val)
                    # Fetch IKIO tables
                    rows_pb = fetch_rows('project_bids')
                    rows_pp = fetch_rows('past_performance')
                    rows_pe = fetch_rows('personnel')
                    if rows_pb or rows_pp or rows_pe:
                        doc.add_page_break()
                        doc.add_heading('IKIO Supplemental Information', level=1)
                        doc.add_paragraph('The following sections include IKIO-specific data curated for Business Development use.')
                        add_table_section('IKIO Project Bids', rows_pb)
                        add_table_section('IKIO Past Performance', rows_pp)
                        add_table_section('IKIO Key Personnel', rows_pe)
                    try:
                        curx.close()
                    except Exception:
                        pass
        except Exception:
            # Never fail proposal generation due to supplemental data
            pass

        # Save document
        try:
            # Link headers/footers to the first section so they are consistent
            for _s_idx, _section in enumerate(doc.sections):
                if _s_idx > 0:
                    try:
                        _section.header.is_linked_to_previous = True
                        _section.footer.is_linked_to_previous = True
                    except Exception:
                        pass
                # Disable special first/odd-even to make header/footer consistent on all pages
                try:
                    _section.different_first_page_header_footer = False
                except Exception:
                    pass
                try:
                    _section.odd_and_even_pages_header_footer = False
                except Exception:
                    pass
            # Ensure continuous page numbering across sections
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            for _s_idx, _section in enumerate(doc.sections):
                _sectPr = _section._sectPr
                _pgNumType = _sectPr.find(qn('w:pgNumType'))
                if _pgNumType is None:
                    _pgNumType = OxmlElement('w:pgNumType')
                    _sectPr.append(_pgNumType)
                if _s_idx == 0:
                    _pgNumType.set(qn('w:start'), "1")
                else:
                    try:
                        _pgNumType.attrib.pop(qn('w:start'), None)
                    except Exception:
                        pass
        except Exception:
            pass
        doc.save(output_path)
        print(f"Proposal saved to: {output_path}")
        return True
        
    except Exception as e:
        print(f"Error generating complete proposal: {str(e)}")
        import traceback
        traceback.print_exc()
        # Create a basic document as fallback
        try:
            from docx import Document
            from datetime import datetime
            doc = Document()
            doc.add_heading('Proposal Generation Error', 0)
            doc.add_paragraph(f'RFP Document: {os.path.basename(pdf_path)}')
            doc.add_paragraph(f'Company: {company}')
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph('\nNote: An error occurred during proposal generation.')
            doc.add_paragraph(f'Error: {str(e)}')
            doc.save(output_path)
            return True
        except:
            return False

def generate_executive_summary_with_toc(pdf_path, company, template, output_path):
    """
    Generate an executive summary document with table of contents from RFP PDF.
    Uses RFP analyzer functions to extract and summarize content.
    """
    try:
        import os
        from datetime import datetime
        from rfp_analyzer_routes import extract_page_texts, summarize_batch_with_llama, build_master_summary, ollama_json, trim_text_to_token_limit
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re
        
        # Read PDF file
        with open(pdf_path, 'rb') as pdf_file:
            pages = extract_page_texts(pdf_file)
        
        if not pages or len(pages) == 0:
            raise Exception("Could not extract text from PDF")
        
        full_text = "\n\n".join(pages)
        
        # Process in batches for summarization
        DEFAULT_BATCH_PAGES = 10
        batch_summaries = []
        total_pages = len(pages)
        
        for i in range(0, total_pages, DEFAULT_BATCH_PAGES):
            batch_pages = pages[i:min(i + DEFAULT_BATCH_PAGES, total_pages)]
            batch_text = "\n\n".join(batch_pages)
            batch_idx = (i // DEFAULT_BATCH_PAGES) + 1
            summary = summarize_batch_with_llama(batch_text, batch_idx)
            if summary and not summary.get('error'):
                batch_summaries.append(summary)
        
        # Build master summary
        if batch_summaries:
            master_summary = build_master_summary(batch_summaries)
        else:
            # Fallback: create basic summary from text
            master_summary = {
                "project_type": "General Project",
                "scope": full_text[:500] + "..." if len(full_text) > 500 else full_text,
                "summary": full_text[:1000] + "..." if len(full_text) > 1000 else full_text,
                "key_requirements": [],
                "technical_requirements": [],
                "bid_requirements": []
            }
        
        # Generate executive summary using LLM
        summary_text = trim_text_to_token_limit(master_summary.get('summary', full_text[:2000]), 4000)
        
        
        # Use OpenAI to generate executive summary
        api_key = (app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or os.getenv('OPENROUTER_API_KEY') or '').strip()
        if not api_key:
            raise Exception("OPENAI_API_KEY is not configured (check .env in project folder)")
        base_url = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or 'https://openrouter.ai/api/v1').strip()
        model = (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'meta-llama/llama-4-scout-17b-16e-instruct').strip()
        endpoint = base_url.rstrip('/')
        if not endpoint.endswith('/chat/completions'):
            endpoint = f"{endpoint}/chat/completions"
        headers = {
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json',
        }
        body = {
            'model': model,
            'messages': [
                {'role': 'system', 'content': 'You are a proposal generation assistant. Return only the executive summary text.'}
            ],
            'temperature': 0.2,
            'max_tokens': 2000
        }
        try:
            response = requests.post(endpoint, headers=headers, json=body, timeout=90)
        except requests.RequestException as net_err:
            raise Exception(f"OpenAI network error: {net_err}")
        try:
            data = response.json()
        except ValueError:
            raise Exception(f"OpenAI returned non-JSON response (status {response.status_code})")
        if response.status_code >= 400:
            err_message = data.get('error') if isinstance(data, dict) else None
            detail = ''
            if isinstance(err_message, dict):
                detail = err_message.get('message') or err_message.get('code') or ''
            elif isinstance(err_message, str):
                detail = err_message
            if not detail:
                detail = f"Upstream error {response.status_code}"
            raise Exception(f"OpenAI error: {detail}")
        choices = data.get('choices') or []
        summary_text_generated = (choices[0].get('message', {}) or {}).get('content', '') if choices else ''
        executive_summary = summary_text_generated.strip() or master_summary.get('summary', 'Executive summary could not be generated.')
        
        # Extract table of contents from PDF text
        toc_items = []
        toc_patterns = [
            r'(?i)^\s*(?:table\s+of\s+contents?|contents?|toc)\s*$',
            r'(?i)^\s*\d+\.?\s+[A-Z][^\n]{10,100}\s+\.\.\.?\s*\d+',
            r'(?i)^\s*[A-Z][A-Z\s]{5,50}\s+\.\.\.?\s*\d+',
        ]
        
        # Try to find existing TOC in PDF
        lines = full_text.split('\n')
        in_toc = False
        for line in lines[:100]:  # Check first 100 lines
            line_clean = line.strip()
            if re.search(toc_patterns[0], line_clean):
                in_toc = True
                continue
            if in_toc and line_clean:
                # Look for TOC entries
                match = re.search(r'^(.+?)\s+\.\.\.?\s*(\d+)$', line_clean)
                if match:
                    toc_items.append({
                        'title': match.group(1).strip(),
                        'page': match.group(2).strip()
                    })
                elif len(toc_items) > 0 and len(line_clean) > 5:
                    # Continue previous entry
                    toc_items[-1]['title'] += ' ' + line_clean
                if len(toc_items) > 20:  # Limit TOC items
                    break
        
        # If no TOC found, generate one from sections
        if not toc_items:
            section_patterns = [
                r'(?i)^\s*\d+\.?\s+([A-Z][^\n]{10,80})',
                r'(?i)^\s*([A-Z][A-Z\s]{5,50})\s*$',
                r'(?i)^\s*(?:section|chapter|part)\s+\d+[:\s]+([^\n]{10,80})',
            ]
            seen_sections = set()
            for line in lines[:200]:
                for pattern in section_patterns:
                    match = re.search(pattern, line)
                    if match:
                        section_title = match.group(1).strip()
                        if len(section_title) > 10 and section_title not in seen_sections:
                            toc_items.append({'title': section_title, 'page': ''})
                            seen_sections.add(section_title)
                            if len(toc_items) >= 15:
                                break
                if len(toc_items) >= 15:
                    break
        
        # Create Word document (prefer H & F template to inherit header/footer)
        try:
            base_dir = app.root_path if hasattr(app, 'root_path') else os.getcwd()
            hf_template_path = os.path.join(base_dir, 'H & F.docx')
            if os.path.exists(hf_template_path):
                doc = Document(hf_template_path)
            else:
                doc = Document()
        except Exception:
            doc = Document()
        
        # Title page
        title_para = doc.add_heading('Executive Summary', 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'RFP Document: {os.path.basename(pdf_path)}')
        doc.add_paragraph(f'Company: {company}')
        doc.add_paragraph(f'Template: {template}')
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_page_break()
        
        # Table of Contents
        doc.add_heading('Table of Contents', 1)
        toc_para = doc.add_paragraph()
        
        if toc_items:
            for item in toc_items:
                p = doc.add_paragraph()
                run1 = p.add_run(item['title'])
                run1.font.size = Pt(11)
                if item['page']:
                    run2 = p.add_run(f'\t{item["page"]}')
                    run2.font.size = Pt(11)
                    # Add tab stop for alignment
                    p.paragraph_format.tab_stops.add_tab_stop(Inches(5.5))
        else:
            doc.add_paragraph('1. Executive Summary')
            doc.add_paragraph('2. Project Overview')
            doc.add_paragraph('3. Scope of Work')
            doc.add_paragraph('4. Key Requirements')
            doc.add_paragraph('5. Technical Specifications')
            doc.add_paragraph('6. Timeline and Deadlines')
            doc.add_paragraph('7. Budget Information')
            doc.add_paragraph('8. Submission Requirements')
        
        doc.add_page_break()
        
        # Executive Summary Section
        doc.add_heading('Executive Summary', 1)
        exec_para = doc.add_paragraph(executive_summary)
        exec_para_format = exec_para.paragraph_format
        exec_para_format.space_after = Pt(12)
        
        # Project Overview
        if master_summary.get('project_type') or master_summary.get('scope'):
            doc.add_heading('Project Overview', 1)
            if master_summary.get('project_type'):
                doc.add_paragraph(f"Project Type: {master_summary['project_type']}", style='List Bullet')
            if master_summary.get('scope'):
                scope_text = master_summary['scope'][:2000] if len(master_summary.get('scope', '')) > 2000 else master_summary.get('scope', '')
                doc.add_paragraph(scope_text)
        
        # Key Requirements
        if master_summary.get('key_requirements') and len(master_summary['key_requirements']) > 0:
            doc.add_heading('Key Requirements', 1)
            for req in master_summary['key_requirements'][:10]:  # Limit to 10
                if isinstance(req, str) and len(req.strip()) > 0:
                    doc.add_paragraph(req, style='List Bullet')
        
        # Technical Requirements
        if master_summary.get('technical_requirements') and len(master_summary['technical_requirements']) > 0:
            doc.add_heading('Technical Requirements', 1)
            for req in master_summary['technical_requirements'][:10]:  # Limit to 10
                if isinstance(req, str) and len(req.strip()) > 0:
                    doc.add_paragraph(req, style='List Bullet')
        
        # Bid Requirements
        if master_summary.get('bid_requirements') and len(master_summary['bid_requirements']) > 0:
            doc.add_heading('Bid Submission Requirements', 1)
            for req in master_summary['bid_requirements'][:10]:  # Limit to 10
                if isinstance(req, str) and len(req.strip()) > 0:
                    doc.add_paragraph(req, style='List Bullet')
        
        # Additional Information
        doc.add_heading('Document Information', 1)
        doc.add_paragraph(f'Total Pages: {total_pages}', style='List Bullet')
        doc.add_paragraph(f'Total Characters: {len(full_text):,}', style='List Bullet')
        if master_summary.get('due_date'):
            doc.add_paragraph(f'Due Date: {master_summary["due_date"]}', style='List Bullet')
        if master_summary.get('total_cost'):
            doc.add_paragraph(f'Estimated Cost: {master_summary["total_cost"]}', style='List Bullet')
        
        # Save document
        try:
            # Link headers/footers to the first section so they are consistent
            for _s_idx, _section in enumerate(doc.sections):
                if _s_idx > 0:
                    try:
                        _section.header.is_linked_to_previous = True
                        _section.footer.is_linked_to_previous = True
                    except Exception:
                        pass
                # Disable special first/odd-even to make header/footer consistent on all pages
                try:
                    _section.different_first_page_header_footer = False
                except Exception:
                    pass
                try:
                    _section.odd_and_even_pages_header_footer = False
                except Exception:
                    pass
            # Ensure continuous page numbering across sections
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            for _s_idx, _section in enumerate(doc.sections):
                _sectPr = _section._sectPr
                _pgNumType = _sectPr.find(qn('w:pgNumType'))
                if _pgNumType is None:
                    _pgNumType = OxmlElement('w:pgNumType')
                    _sectPr.append(_pgNumType)
                if _s_idx == 0:
                    _pgNumType.set(qn('w:start'), "1")
                else:
                    try:
                        _pgNumType.attrib.pop(qn('w:start'), None)
                    except Exception:
                        pass
        except Exception:
            pass
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error generating executive summary: {str(e)}")
        import traceback
        traceback.print_exc()
        # Create a basic document as fallback
        try:
            from docx import Document
            from datetime import datetime
            doc = Document()
            doc.add_heading('Executive Summary', 0)
            doc.add_paragraph(f'RFP Document: {os.path.basename(pdf_path)}')
            doc.add_paragraph(f'Company: {company}')
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph('\nNote: An error occurred during executive summary generation.')
            doc.add_paragraph(f'Error: {str(e)}')
            doc.save(output_path)
            return True
        except:
            return False

@app.route('/proposals-making', methods=['GET', 'POST'])
@login_required
def proposals_making():
    # Allow all logged-in users to access proposals-making (removed admin-only restriction)
    import os
    import re
    from datetime import datetime, date

    contact_email = getattr(current_user, 'email', '')

    def safe_str(value):
        if value is None:
            return ''
        return str(value).strip()

    def normalize_company_choice(value: str) -> str:
        """
        Map various database/company aliases to canonical select values used by the UI:
          - any name containing 'ikio' -> 'IKIO'
          - any name containing 'metco' -> 'METCO'
          - any name containing 'sunsprint' -> 'SUNSPRINT'
        Returns empty string if the input is falsy.
        """
        s = (value or '').strip().lower()
        if not s:
            return ''
        if 'ikio' in s:
            return 'IKIO'
        if 'metco' in s:
            return 'METCO'
        if 'sunsprint' in s:
            return 'SUNSPRINT'
        return value or ''

    def format_date_for_display(value):
        if not value:
            return ''
        if isinstance(value, (datetime, date)):
            return value.strftime('%B %d, %Y')
        try:
            parsed = datetime.strptime(str(value), '%Y-%m-%d')
            return parsed.strftime('%B %d, %Y')
        except Exception:
            return str(value)

    def split_text_blocks(text, max_blocks=3):
        if not text:
            return []
        blocks = [
            safe_str(piece).strip('•- ') for piece in re.split(r'(?:\r?\n){2,}', str(text))
            if safe_str(piece)
        ]
        if not blocks and safe_str(text):
            blocks = [safe_str(text)]
        return blocks[:max_blocks]

    def compact_list(items):
        result = []
        for item in items:
            if item is None:
                continue
            text = safe_str(item)
            if text:
                result.append(text)
        return result

    def get_bid_metadata(g_id):
        if not g_id:
            return None
        meta_cursor = mysql.connection.cursor(DictCursor)
        try:
            meta_cursor.execute("SELECT * FROM go_bids WHERE g_id=%s", (g_id,))
            return meta_cursor.fetchone()
        except Exception as meta_err:
            print(f"Error loading bid context for proposal view: {meta_err}")
            return None
        finally:
            meta_cursor.close()

    def build_sections_outline(bid_info, fallback_title, company_name, primary_contact):
        bid = bid_info or {}
        project_name = safe_str(bid.get('b_name')) or safe_str(fallback_title) or 'Current Opportunity'
        due_text = format_date_for_display(bid.get('due_date'))
        client_name = safe_str(
            bid.get('agency') or bid.get('customer') or bid.get('client') or bid.get('owner') or bid.get('comp_name')
        )
        scope_text = safe_str(bid.get('scope'))
        summary_text = safe_str(bid.get('summary'))
        type_text = safe_str(bid.get('type'))
        location_text = safe_str(bid.get('location') or bid.get('city') or bid.get('state'))
        stage_text = safe_str(bid.get('state'))
        company_display = safe_str(company_name) or safe_str(bid.get('company')) or 'Offer or'
        revenue_value = bid.get('revenue') if bid.get('revenue') not in (None, '') else bid.get('value')
        scoring_value = bid.get('scoring')

        summary_blocks = split_text_blocks(summary_text, 4)
        scope_blocks = split_text_blocks(scope_text, 5)

        revenue_text = ''
        if revenue_value not in (None, ''):
            try:
                revenue_text = f"${float(revenue_value):,.0f}"
            except Exception:
                revenue_text = safe_str(revenue_value)

        scoring_text = ''
        if scoring_value not in (None, ''):
            try:
                scoring_text = f"{float(scoring_value):.0f}"
            except Exception:
                scoring_text = safe_str(scoring_value)

        sections = []

        cover_content = compact_list([
            f"{company_display} is pleased to submit this proposal for {project_name}.",
            f"Issuing agency / customer: {client_name}." if client_name else '',
            f"Solicitation reference: {project_name}.",
            f"Proposal due date: {due_text}." if due_text else '',
            f"Primary point of contact: {primary_contact}." if primary_contact else '',
            f"Contract type: {type_text}." if type_text else '',
            f"Performance location: {location_text}." if location_text else '',
            f"Opportunity priority score: {scoring_text}." if scoring_text else ''
        ])
        sections.append({
            'id': 'section-cover',
            'title': 'Letter of Transmittal',
            'status': 'Saved' if cover_content else 'Draft',
            'guidance': 'Confirm the solicitation title, client information, offer or contact details, and submission identifiers before releasing the cover page.',
            'content': cover_content or ['Add cover page details once the capture team confirms the solicitation data.'],
            'sub_categories': [
                {'id': 'cover-company-info', 'title': 'Company Information'},
                {'id': 'cover-contact-details', 'title': 'Contact Details'},
                {'id': 'cover-solicitation-ref', 'title': 'Solicitation Reference'},
                {'id': 'cover-submission-date', 'title': 'Submission Date & Deadline'}
            ]
        })

        exec_content = summary_blocks or (
            [f"{company_display} will deliver a compliant, value-focused response for {project_name}."]
            if project_name else []
        )
        if scope_text and not summary_blocks:
            exec_content.append(scope_text)
        sections.append({
            'id': 'section-executive-summary',
            'title': 'Executive Summary',
            'status': 'In Progress' if summary_blocks else 'Draft',
            'guidance': "Summarize win themes, differentiators, and alignment to the customer's objectives captured during bid qualification.",
            'content': exec_content or ['Executive summary content will be generated as soon as capture notes are available.'],
            'sub_categories': [
                {'id': 'exec-overview', 'title': 'Project Overview'},
                {'id': 'exec-win-themes', 'title': 'Win Themes & Differentiators'},
                {'id': 'exec-value-proposition', 'title': 'Value Proposition'},
                {'id': 'exec-key-highlights', 'title': 'Key Highlights'}
            ]
        })

        compliance_content = compact_list([
            "Track each requirement in the compliance matrix to ensure every section of the RFP is addressed.",
            scope_blocks[0] if scope_blocks else '',
            summary_blocks[0] if summary_blocks else ''
        ])
        sections.append({
            'id': 'section-compliance',
            'title': 'Requirements Compliance',
            'status': 'In Progress',
            'guidance': 'Reference the compliance matrix and list every deliverable, attachment, and form demanded by the solicitation.',
            'content': compliance_content or [
                'Populate this section with compliance findings once the RFP analysis is complete.'
            ],
            'sub_categories': [
                {'id': 'comp-matrix', 'title': 'Compliance Matrix'},
                {'id': 'comp-deliverables', 'title': 'Deliverables List'},
                {'id': 'comp-attachments', 'title': 'Required Attachments'},
                {'id': 'comp-forms', 'title': 'Required Forms'}
            ]
        })

        sections.append({
            'id': 'section-scope-objectives',
            'title': 'Understanding of Scope & Objectives',
            'status': 'In Progress',
            'guidance': 'Demonstrate a clear understanding of the project scope, objectives, and expected outcomes.',
            'content': scope_blocks or [
                'Detail your understanding of the project scope and objectives based on the solicitation requirements.'
            ],
            'sub_categories': [
                {'id': 'scope-project-scope', 'title': 'Project Scope'},
                {'id': 'scope-objectives', 'title': 'Project Objectives'},
                {'id': 'scope-expected-outcomes', 'title': 'Expected Outcomes'},
                {'id': 'scope-assumptions', 'title': 'Key Assumptions'}
            ]
        })

        sections.append({
            'id': 'section-deviations',
            'title': 'Deviations, Assumptions, and Dependencies (if any)',
            'status': 'Draft',
            'guidance': 'Document any deviations from the solicitation requirements, key assumptions, and dependencies that may impact project delivery.',
            'content': [
                'List any deviations from the solicitation requirements, if applicable.',
                'Document key assumptions made during proposal development.',
                'Identify dependencies that may affect project execution.'
            ],
            'sub_categories': [
                {'id': 'dev-deviations', 'title': 'Deviations'},
                {'id': 'dev-assumptions', 'title': 'Assumptions'},
                {'id': 'dev-dependencies', 'title': 'Dependencies'},
                {'id': 'dev-risks', 'title': 'Risk Considerations'}
            ]
        })

        tech_content = scope_blocks or [
            'Detail the technical solution, materials, and implementation methodology once the RFP requirements are parsed.'
        ]
        sections.append({
            'id': 'section-technical',
            'title': 'Proposed Technical Solution',
            'status': 'In Progress' if scope_blocks else 'Draft',
            'guidance': 'Translate the solicitation Statement of Work into your delivery plan, highlighting compliance, innovations, and risk mitigations.',
            'content': tech_content,
            'sub_categories': [
                {'id': 'tech-approach', 'title': 'Technical Approach'},
                {'id': 'tech-methodology', 'title': 'Methodology'},
                {'id': 'tech-innovations', 'title': 'Innovations & Best Practices'},
                {'id': 'tech-compliance', 'title': 'Compliance & Standards'},
                {'id': 'tech-risk-mitigation', 'title': 'Risk Mitigation'}
            ]
        })

        sections.append({
            'id': 'section-implementation',
            'title': 'Implementation & Work Plan',
            'status': 'In Progress',
            'guidance': 'Provide a detailed implementation plan and work breakdown structure for project execution.',
            'content': [
                'Outline the step-by-step implementation approach.',
                'Detail the work breakdown structure and key activities.',
                'Describe the workflow and processes for project delivery.'
            ],
            'sub_categories': [
                {'id': 'impl-approach', 'title': 'Implementation Approach'},
                {'id': 'impl-work-breakdown', 'title': 'Work Breakdown Structure'},
                {'id': 'impl-phases', 'title': 'Implementation Phases'},
                {'id': 'impl-processes', 'title': 'Key Processes'}
            ]
        })

        sections.append({
            'id': 'section-deliverables',
            'title': 'Solution Deliverables',
            'status': 'In Progress',
            'guidance': 'List all deliverables that will be provided as part of the solution, including documentation, reports, and physical items.',
            'content': [
                'Comprehensive list of all solution deliverables.',
                'Documentation and reports to be provided.',
                'Physical deliverables, if applicable.'
            ],
            'sub_categories': [
                {'id': 'del-documentation', 'title': 'Documentation Deliverables'},
                {'id': 'del-reports', 'title': 'Reports & Analysis'},
                {'id': 'del-physical', 'title': 'Physical Deliverables'},
                {'id': 'del-schedule', 'title': 'Delivery Schedule'}
            ]
        })

        management_content = compact_list([
            f"Target submission date: {due_text}." if due_text else 'Submission date will be updated when the procurement schedule is confirmed.',
            f"Contract type guidance: {type_text}." if type_text else '',
            "Outline major milestones, design reviews, and government touchpoints that support a compliant response."
        ])
        sections.append({
            'id': 'section-management',
            'title': 'Project Management Plan',
            'status': 'In Progress',
            'guidance': 'Map the internal review cycle, color team dates, and delivery milestones to ensure on-time submission.',
            'content': management_content or [
                'Add the project management narrative and schedule once the capture calendar is finalized.'
            ],
            'sub_categories': [
                {'id': 'mgmt-organization', 'title': 'Project Organization'},
                {'id': 'mgmt-resources', 'title': 'Resource Management'},
                {'id': 'mgmt-quality', 'title': 'Quality Assurance'},
                {'id': 'mgmt-communication', 'title': 'Communication Plan'}
            ]
        })

        sections.append({
            'id': 'section-schedule',
            'title': 'Project Schedule & Milestones',
            'status': 'In Progress',
            'guidance': 'Provide a detailed project schedule with key milestones, deliverables, and timeline.',
            'content': [
                'Detailed project timeline with start and end dates.',
                'Key milestones and their target dates.',
                'Critical path and dependencies.'
            ],
            'sub_categories': [
                {'id': 'schedule-timeline', 'title': 'Project Timeline'},
                {'id': 'schedule-milestones', 'title': 'Key Milestones'},
                {'id': 'schedule-critical-path', 'title': 'Critical Path'},
                {'id': 'schedule-dependencies', 'title': 'Dependencies'}
            ]
        })

        sections.append({
            'id': 'section-team',
            'title': 'Project Team',
            'status': 'In Progress',
            'guidance': 'Introduce the project team members, their roles, qualifications, and responsibilities.',
            'content': [
                'List of key team members and their roles.',
                'Qualifications and experience of team members.',
                'Organizational structure and reporting lines.'
            ],
            'sub_categories': [
                {'id': 'team-members', 'title': 'Team Members'},
                {'id': 'team-roles', 'title': 'Roles & Responsibilities'},
                {'id': 'team-qualifications', 'title': 'Qualifications'},
                {'id': 'team-organization', 'title': 'Organization Structure'}
            ]
        })

        sections.append({
            'id': 'section-corporate',
            'title': 'Corporate Qualifications',
            'status': 'Saved',
            'guidance': 'Highlight corporate experience, past performance, certifications, and qualifications relevant to this opportunity.',
            'content': [
                'Company background and history.',
                'Relevant past performance and case studies.',
                'Certifications, accreditations, and qualifications.'
            ],
            'sub_categories': [
                {'id': 'corp-background', 'title': 'Company Background'},
                {'id': 'corp-past-performance', 'title': 'Past Performance'},
                {'id': 'corp-certifications', 'title': 'Certifications & Accreditations'},
                {'id': 'corp-capabilities', 'title': 'Key Capabilities'}
            ]
        })

        price_content = compact_list([
            f"Estimated contract value: {revenue_text}." if revenue_text else '',
            "Attachment A pricing will mirror the government-provided workbook with CLIN-level details.",
            "Document all pricing assumptions, escalation factors, and exclusions for contracting officer review."
        ])
        sections.append({
            'id': 'section-pricing',
            'title': 'Pricing Proposals',
            'status': 'Saved' if revenue_text else 'Draft',
            'guidance': 'Ensure Attachment A reflects current quantities, rates, and assumptions. Cross-check CLIN totals before submission.',
            'content': price_content or [
                'Populate Attachment A once final quantities and labor categories are confirmed.'
            ],
            'sub_categories': [
                {'id': 'price-clin-breakdown', 'title': 'CLIN Breakdown'},
                {'id': 'price-labor-rates', 'title': 'Labor Rates & Categories'},
                {'id': 'price-materials', 'title': 'Materials & Equipment'},
                {'id': 'price-assumptions', 'title': 'Pricing Assumptions'},
                {'id': 'price-total-summary', 'title': 'Total Cost Summary'}
            ]
        })

        sections.append({
            'id': 'section-contractual',
            'title': 'Contractual Terms',
            'status': 'Saved',
            'guidance': 'Outline key contractual terms, conditions, and agreements proposed for this engagement.',
            'content': [
                'Key contractual terms and conditions.',
                'Payment terms and conditions.',
                'Terms of service and deliverables.'
            ],
            'sub_categories': [
                {'id': 'contract-terms', 'title': 'Terms & Conditions'},
                {'id': 'contract-payment', 'title': 'Payment Terms'},
                {'id': 'contract-service', 'title': 'Service Terms'},
                {'id': 'contract-legal', 'title': 'Legal Considerations'}
            ]
        })

        sections.append({
            'id': 'section-project-schedule',
            'title': 'Project Schedule & Milestones',
            'status': 'Pending',
            'guidance': 'Confirm the bond amount, surety approval, and Treasury Circular 570 listing for the guarantor.',
            'content': [
                f"{company_display} will coordinate with the approved surety to furnish the required Project Schedule & Milestones prior to submission.",
                "Include the original bid bond or irrevocable letter of credit documentation in the final package."
            ],
           
        })

        sections.append({
            'id': 'requirements-attachments',
            'title': 'Requirements and Attachments',
            'status': 'Special Section',
            'guidance': 'List mandatory submission requirements, certifications, and attachment references extracted from the RFP.',
            'content': [
                'Compile all mandatory attachments, compliance documents, and certifications referenced throughout the solicitation.',
                'Ensure each attachment is labeled clearly and referenced in the compliance matrix.'
            ],
            'sub_categories': [
                {'id': 'req-mandatory', 'title': 'Mandatory Documents'},
                {'id': 'req-forms', 'title': 'Forms & Certifications'},
                {'id': 'req-checklist', 'title': 'Submission Checklist'}
            ]
        })

        sections.append({
            'id': 'section-amendments',
            'title': 'Addenda',
            'status': 'Pending',
            'guidance': 'Log each amendment number and release date. If no amendments were issued, include a statement confirming that status.',
            'content': [
                f"The offer or acknowledges receipt of all amendments released for {project_name} as of {datetime.now().strftime('%B %d, %Y')}." if project_name else
                f"The offer or acknowledges receipt of all solicitation amendments issued as of {datetime.now().strftime('%B %d, %Y')}.",
                "Update this section with individual amendment numbers and dates when the contracting officer publishes them."
            ],
            'sub_categories': [
                {'id': 'amend-list', 'title': 'Amendment List'},
                {'id': 'amend-acknowledgment', 'title': 'Acknowledgment Statement'},
                {'id': 'amend-impact', 'title': 'Impact Assessment'}
            ]
        })

        sections.append({
            'id': 'section-appendix',
            'title': 'Appendix',
            'status': 'Draft',
            'guidance': 'Include all supporting documents, references, and additional materials referenced in the proposal.',
            'content': [
                'Supporting documents and references.',
                'Additional technical specifications.',
                'Supplementary materials and exhibits.'
            ],
            'sub_categories': [
                {'id': 'appendix-documents', 'title': 'Supporting Documents'},
                {'id': 'appendix-references', 'title': 'References'},
                {'id': 'appendix-exhibits', 'title': 'Exhibits'},
                {'id': 'appendix-other', 'title': 'Other Materials'}
            ]
        })

        sections.append({
            'id': 'section-safety',
            'title': 'Safety',
            'status': 'Draft',
            'guidance': 'Describe safety management systems, OSHA compliance, hazard mitigation, and jobsite safety protocols. Include training, PPE, incident reporting, and continuous improvement practices.',
            'content': [
                'Overview of Safety Management System (SMS) and accountability.',
                'OSHA compliance approach and hazard identification and mitigation processes.',
                'Jobsite safety plan, tailgate meetings, and Job Hazard Analysis (JHA) procedures.',
                'Training matrix, certifications, and PPE standards.',
                'Incident reporting, root cause analysis, and corrective actions.'
            ],
            'sub_categories': [
                {'id': 'safety-plan', 'title': 'Safety Plan'},
                {'id': 'safety-compliance', 'title': 'Regulatory Compliance'},
                {'id': 'safety-training', 'title': 'Training & Certifications'},
                {'id': 'safety-incident', 'title': 'Incident Response & Reporting'}
            ]
        })

        sections.append({
            'id': 'section-interconnection',
            'title': 'Interconnection',
            'status': 'Draft',
            'guidance': 'Detail utility interconnection strategy including application, studies, protection coordination, metering, and commissioning. Align with applicable tariffs, IEEE 1547, and utility standards.',
            'content': [
                'Interconnection application workflow and milestone tracking.',
                'Load flow, short circuit, and protection coordination study assumptions.',
                'Protection scheme, relays, and anti-islanding compliance.',
                'Metering, telemetry/SCADA, and communications architecture.',
                'Commissioning, witness testing, and as-built documentation.'
            ],
            'sub_categories': [
                {'id': 'ic-applications', 'title': 'Applications & Permits'},
                {'id': 'ic-studies', 'title': 'Engineering Studies'},
                {'id': 'ic-protection', 'title': 'Protection & Controls'},
                {'id': 'ic-commissioning', 'title': 'Commissioning & Testing'}
            ]
        })

        sections.append({
            'id': 'section-safety',
            'title': 'Safety',
            'status': 'Draft',
            'guidance': 'Describe safety management systems, OSHA compliance, hazard mitigation, and jobsite safety protocols. Include training, PPE, incident reporting, and continuous improvement practices.',
            'content': [
                'Overview of Safety Management System (SMS) and accountability.',
                'OSHA compliance approach and hazard identification and mitigation processes.',
                'Jobsite safety plan, tailgate meetings, and Job Hazard Analysis (JHA) procedures.',
                'Training matrix, certifications, and PPE standards.',
                'Incident reporting, root cause analysis, and corrective actions.'
            ],
            'sub_categories': [
                {'id': 'safety-plan', 'title': 'Safety Plan'},
                {'id': 'safety-compliance', 'title': 'Regulatory Compliance'},
                {'id': 'safety-training', 'title': 'Training & Certifications'},
                {'id': 'safety-incident', 'title': 'Incident Response & Reporting'}
            ]
        })

        sections.append({
            'id': 'section-interconnection',
            'title': 'Interconnection',
            'status': 'Draft',
            'guidance': 'Detail utility interconnection strategy including application, studies, protection coordination, metering, and commissioning. Align with applicable tariffs, IEEE 1547, and utility standards.',
            'content': [
                'Interconnection application workflow and milestone tracking.',
                'Load flow, short circuit, and protection coordination study assumptions.',
                'Protection scheme, relays, and anti-islanding compliance.',
                'Metering, telemetry/SCADA, and communications architecture.',
                'Commissioning, witness testing, and as-built documentation.'
            ],
            'sub_categories': [
                {'id': 'ic-applications', 'title': 'Applications & Permits'},
                {'id': 'ic-studies', 'title': 'Engineering Studies'},
                {'id': 'ic-protection', 'title': 'Protection & Controls'},
                {'id': 'ic-commissioning', 'title': 'Commissioning & Testing'}
            ]
        })

        return sections

    if request.method == 'POST':
        try:
            files = request.files.getlist('rfp_files')
            company = normalize_company_choice(request.form.get('company', '').strip())
            template = request.form.get('template', 'standard').strip()
            bid_name = request.args.get('bid_name', '') or request.form.get('bid_name', '')
            bid_id_value = (
                request.args.get('bid_id', type=int)
                or request.form.get('bid_id', type=int)
                or request.args.get('g_id', type=int)
                or request.form.get('g_id', type=int)
            )
            bid_meta_for_generation = get_bid_metadata(bid_id_value)
            sections_outline_for_generation = build_sections_outline(
                bid_meta_for_generation,
                bid_name,
                company,
                contact_email
            )
            
            if not files or len(files) == 0:
                # No new file uploaded; try using the latest RFP file already saved for this bid
                try:
                    g_id_for_lookup = (
                        request.args.get('bid_id', type=int)
                        or request.form.get('bid_id', type=int)
                        or request.args.get('g_id', type=int)
                        or request.form.get('g_id', type=int)
                        or bid_id_value
                    )
                    rfp_row = _get_latest_rfp_file_for_bid(g_id_for_lookup) if g_id_for_lookup else None
                    rfp_path = rfp_row.get('file_path') if rfp_row else None
                    rfp_name = rfp_row.get('filename') if rfp_row else None
                except Exception:
                    rfp_path = None
                    rfp_name = None
                if not rfp_path or not os.path.exists(rfp_path):
                    flash('Please attach an RFP PDF or ensure an RFP is uploaded for this bid.', 'error')
                    return redirect(url_for('proposals_making'))
                
                # Generate proposal from database RFP file
                import uuid
                proposals_dir = 'uploads/proposals'
                os.makedirs(proposals_dir, exist_ok=True)
                processed_files = []
                errors = []
                try:
                    unique_id = str(uuid.uuid4())[:8]
                    filename = rfp_name or os.path.basename(rfp_path)
                    proposal_filename = f"proposal_{unique_id}_{(filename or 'rfp').replace('.pdf', '.docx')}"
                    proposal_path = os.path.join(proposals_dir, proposal_filename)
                    
                    success = generate_complete_proposal(
                        rfp_path,
                        company,
                        template,
                        proposal_path,
                        bid_name=bid_name,
                        sections_outline=sections_outline_for_generation,
                        user_email=contact_email
                    )
                    if success:
                        cur = mysql.connection.cursor(DictCursor)
                        try:
                            # Ensure table exists (idempotent)
                            cur.execute("""
                                CREATE TABLE IF NOT EXISTS proposals (
                                    id INT AUTO_INCREMENT PRIMARY KEY,
                                    unique_id VARCHAR(50) NOT NULL UNIQUE,
                                    rfp_name VARCHAR(500) NOT NULL,
                                    rfp_filename VARCHAR(500),
                                    proposal_filename VARCHAR(500) NOT NULL,
                                    proposal_path VARCHAR(1000) NOT NULL,
                                    company VARCHAR(100) NOT NULL,
                                    template VARCHAR(100),
                                    bid_id INT NULL,
                                    bid_name VARCHAR(500) NULL,
                                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                                    INDEX idx_unique_id (unique_id),
                                    INDEX idx_rfp_name (rfp_name),
                                    INDEX idx_company (company),
                                    INDEX idx_bid_id (bid_id)
                                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
                            """)
                            bid_id = (
                                request.args.get('bid_id', type=int)
                                or request.form.get('bid_id', type=int)
                                or request.args.get('g_id', type=int)
                                or request.form.get('g_id', type=int)
                            )
                            cur.execute("""
                                INSERT INTO proposals 
                                (unique_id, rfp_name, rfp_filename, proposal_filename, proposal_path, company, template, bid_id, bid_name)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                            """, (
                                unique_id,
                                filename or '',
                                filename or '',
                                proposal_filename,
                                proposal_path,
                                company,
                                template,
                                bid_id if bid_id else None,
                                bid_name if bid_name else None
                            ))
                            mysql.connection.commit()
                            processed_files.append({
                                'original': filename,
                                'proposal': proposal_filename,
                                'path': proposal_path,
                                'unique_id': unique_id,
                                'rfp_name': filename
                            })
                            log_write('proposal_generated', f"file={filename} company={company} template={template} unique_id={unique_id} bid_id={bid_id}")
                        except Exception as db_error:
                            mysql.connection.rollback()
                            print(f"Database error: {str(db_error)}")
                            processed_files.append({
                                'original': filename,
                                'proposal': proposal_filename,
                                'path': proposal_path,
                                'unique_id': unique_id,
                                'rfp_name': filename
                            })
                        finally:
                            cur.close()
                    else:
                        errors.append(f"Failed to generate proposal for {rfp_name or os.path.basename(rfp_path)}")
                except Exception as e:
                    errors.append(f"Error processing database RFP: {str(e)}")
                
                if processed_files:
                    flash(f'Successfully generated {len(processed_files)} proposal(s).', 'success')
                if errors:
                    flash(f'Errors: {"; ".join(errors)}', 'error')
                return redirect(url_for('proposals_making'))
            
            if not company:
                flash('Please select a company.', 'error')
                return redirect(url_for('proposals_making'))
            
            import uuid
            from werkzeug.utils import secure_filename
            
            # Create proposals directory
            proposals_dir = 'uploads/proposals'
            os.makedirs(proposals_dir, exist_ok=True)
            
            processed_files = []
            errors = []
            
            for file in files:
                if file and file.filename:
                    if not file.filename.lower().endswith('.pdf'):
                        errors.append(f"{file.filename} is not a PDF file")
                        continue
                    
                    # Save uploaded file temporarily
                    filename = secure_filename(file.filename)
                    unique_id = str(uuid.uuid4())[:8]
                    temp_filename = f"{unique_id}_{filename}"
                    temp_path = os.path.join('uploads/rfp', temp_filename)
                    os.makedirs('uploads/rfp', exist_ok=True)
                    file.save(temp_path)
                    
                    try:
                        # Generate complete proposal with all sections
                        proposal_filename = f"proposal_{unique_id}_{filename.replace('.pdf', '.docx')}"
                        proposal_path = os.path.join(proposals_dir, proposal_filename)
                        
                        # Generate complete proposal using AI
                        success = generate_complete_proposal(
                            temp_path,
                            company,
                            template,
                            proposal_path,
                            bid_name=bid_name,
                            sections_outline=sections_outline_for_generation,
                            user_email=contact_email
                        )
                        
                        if success:
                            # Save proposal data to database
                            cur = mysql.connection.cursor(DictCursor)
                            try:
                                # Create proposals table if it doesn't exist
                                cur.execute("""
                                    CREATE TABLE IF NOT EXISTS proposals (
                                        id INT AUTO_INCREMENT PRIMARY KEY,
                                        unique_id VARCHAR(50) NOT NULL UNIQUE,
                                        rfp_name VARCHAR(500) NOT NULL,
                                        rfp_filename VARCHAR(500),
                                        proposal_filename VARCHAR(500) NOT NULL,
                                        proposal_path VARCHAR(1000) NOT NULL,
                                        company VARCHAR(100) NOT NULL,
                                        template VARCHAR(100),
                                        bid_id INT NULL,
                                        bid_name VARCHAR(500) NULL,
                                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                                        INDEX idx_unique_id (unique_id),
                                        INDEX idx_rfp_name (rfp_name),
                                        INDEX idx_company (company),
                                        INDEX idx_bid_id (bid_id)
                                    ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
                                """)
                                
                                # Get bid_id (aka g_id) and bid_name from query parameters if available
                                bid_id = (
                                    request.args.get('bid_id', type=int)
                                    or request.form.get('bid_id', type=int)
                                    or request.args.get('g_id', type=int)
                                    or request.form.get('g_id', type=int)
                                )
                                bid_name = request.args.get('bid_name', '') or request.form.get('bid_name', '')
                                
                                # Insert proposal record
                                cur.execute("""
                                    INSERT INTO proposals 
                                    (unique_id, rfp_name, rfp_filename, proposal_filename, proposal_path, company, template, bid_id, bid_name)
                                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """, (
                                    unique_id,
                                    filename,  # RFP name
                                    filename,  # RFP filename
                                    proposal_filename,
                                    proposal_path,
                                    company,
                                    template,
                                    bid_id if bid_id else None,
                                    bid_name if bid_name else None
                                ))
                                mysql.connection.commit()
                                
                                processed_files.append({
                                    'original': filename,
                                    'proposal': proposal_filename,
                                    'path': proposal_path,
                                    'unique_id': unique_id,
                                    'rfp_name': filename
                                })
                                log_write('proposal_generated', f"file={filename} company={company} template={template} unique_id={unique_id} bid_id={bid_id}")
                            except Exception as db_error:
                                mysql.connection.rollback()
                                print(f"Database error: {str(db_error)}")
                                # Still add to processed_files even if DB save fails
                                processed_files.append({
                                    'original': filename,
                                    'proposal': proposal_filename,
                                    'path': proposal_path,
                                    'unique_id': unique_id,
                                    'rfp_name': filename
                                })
                            finally:
                                cur.close()
                        else:
                            errors.append(f"Failed to generate executive summary for {filename}")
                    except Exception as e:
                        errors.append(f"Error processing {filename}: {str(e)}")
                        if os.path.exists(temp_path):
                            os.remove(temp_path)
            
            if processed_files:
                flash(f'Successfully generated {len(processed_files)} proposal(s).', 'success')
            if errors:
                flash(f'Errors: {"; ".join(errors)}', 'error')
            
            return redirect(url_for('proposals_making'))
        
        except Exception as e:
            flash(f'Error: {str(e)}', 'error')
            return redirect(url_for('proposals_making'))
    
    # GET request - show page
    import os
    from datetime import datetime
    
    # Get bid_id (a.k.a g_id) and bid_name from query parameters if available (from Generate Proposals button)
    bid_id = request.args.get('bid_id', type=int)
    if bid_id is None:
        bid_id = request.args.get('g_id', type=int)
    bid_name = request.args.get('bid_name', '')
    company = request.args.get('company', '')

    contact_email = getattr(current_user, 'email', '')

    def safe_str(value):
        if value is None:
            return ''
        return str(value).strip()

    def format_date_for_display(value):
        if not value:
            return ''
        if isinstance(value, (datetime, date)):
            return value.strftime('%B %d, %Y')
        try:
            parsed = datetime.strptime(str(value), '%Y-%m-%d')
            return parsed.strftime('%B %d, %Y')
        except Exception:
            return str(value)

    def split_text_blocks(text, max_blocks=3):
        if not text:
            return []
        blocks = [
            safe_str(piece).strip('•- ') for piece in re.split(r'(?:\r?\n){2,}', str(text))
            if safe_str(piece)
        ]
        if not blocks and safe_str(text):
            blocks = [safe_str(text)]
        return blocks[:max_blocks]

    def compact_list(items):
        result = []
        for item in items:
            if item is None:
                continue
            text = safe_str(item)
            if text:
                result.append(text)
        return result

    def build_sections_outline(bid_info, fallback_title, company_name, primary_contact):
        bid = bid_info or {}
        project_name = safe_str(bid.get('b_name')) or safe_str(fallback_title) or 'Current Opportunity'
        due_text = format_date_for_display(bid.get('due_date'))
        client_name = safe_str(
            bid.get('agency') or bid.get('customer') or bid.get('client') or bid.get('owner') or bid.get('comp_name')
        )
        scope_text = safe_str(bid.get('scope'))
        summary_text = safe_str(bid.get('summary'))
        type_text = safe_str(bid.get('type'))
        location_text = safe_str(bid.get('location') or bid.get('city') or bid.get('state'))
        stage_text = safe_str(bid.get('state'))
        company_display = safe_str(company_name) or safe_str(bid.get('company')) or 'Offer or'
        revenue_value = bid.get('revenue') if bid.get('revenue') not in (None, '') else bid.get('value')
        scoring_value = bid.get('scoring')

        summary_blocks = split_text_blocks(summary_text, 4)
        scope_blocks = split_text_blocks(scope_text, 5)

        revenue_text = ''
        if revenue_value not in (None, ''):
            try:
                revenue_text = f"${float(revenue_value):,.0f}"
            except Exception:
                revenue_text = safe_str(revenue_value)

        scoring_text = ''
        if scoring_value not in (None, ''):
            try:
                scoring_text = f"{float(scoring_value):.0f}"
            except Exception:
                scoring_text = safe_str(scoring_value)

        sections = []

        cover_content = compact_list([
            f"{company_display} is pleased to submit this proposal for {project_name}.",
            f"Issuing agency / customer: {client_name}." if client_name else '',
            f"Solicitation reference: {project_name}.",
            f"Proposal due date: {due_text}." if due_text else '',
            f"Primary point of contact: {primary_contact}." if primary_contact else '',
            f"Contract type: {type_text}." if type_text else '',
            f"Performance location: {location_text}." if location_text else '',
            f"Opportunity priority score: {scoring_text}." if scoring_text else ''
        ])
        sections.append({
            'id': 'section-cover',
            'title': 'Letter of Transmittal',
            'status': 'Saved' if cover_content else 'Draft',
            'guidance': 'Confirm the solicitation title, client information, offer or contact details, and submission identifiers before releasing the cover page.',
            'content': cover_content or ['Add cover page details once the capture team confirms the solicitation data.'],
            'sub_categories': [
                {'id': 'cover-company-info', 'title': 'Company Information'},
                {'id': 'cover-contact-details', 'title': 'Contact Details'},
                {'id': 'cover-solicitation-ref', 'title': 'Solicitation Reference'},
                {'id': 'cover-submission-date', 'title': 'Submission Date & Deadline'}
            ]
        })

        sections.append({
            'id': 'section-toc',
            'title': 'Table of Contents',
            'status': 'Saved',
            'guidance': 'Generate a comprehensive table of contents listing all proposal sections and their page numbers.',
            'content': [
                'The table of contents will be automatically generated once all sections are finalized.',
                'Ensure all major sections and subsections are included with accurate page references.'
            ],
            'sub_categories': [
                {'id': 'toc-main-sections', 'title': 'Main Sections'},
                {'id': 'toc-appendices', 'title': 'Appendices'},
                {'id': 'toc-figures-tables', 'title': 'Figures & Tables'}
            ]
        })

        exec_content = summary_blocks or (
            [f"{company_display} will deliver a compliant, value-focused response for {project_name}."]
            if project_name else []
        )
        if scope_text and not summary_blocks:
            exec_content.append(scope_text)
        sections.append({
            'id': 'section-executive-summary',
            'title': 'Executive Summary',
            'status': 'In Progress' if summary_blocks else 'Draft',
            'guidance': "Summarize win themes, differentiators, and alignment to the customer's objectives captured during bid qualification.",
            'content': exec_content or ['Executive summary content will be generated as soon as capture notes are available.'],
            'sub_categories': [
                {'id': 'exec-overview', 'title': 'Project Overview'},
                {'id': 'exec-win-themes', 'title': 'Win Themes & Differentiators'},
                {'id': 'exec-value-proposition', 'title': 'Value Proposition'},
                {'id': 'exec-key-highlights', 'title': 'Key Highlights'}
            ]
        })

        compliance_content = compact_list([
            "Track each requirement in the compliance matrix to ensure every section of the RFP is addressed.",
            scope_blocks[0] if scope_blocks else '',
            summary_blocks[0] if summary_blocks else ''
        ])
        sections.append({
            'id': 'section-compliance',
            'title': 'Requirements Compliance',
            'status': 'In Progress',
            'guidance': 'Reference the compliance matrix and list every deliverable, attachment, and form demanded by the solicitation.',
            'content': compliance_content or [
                'Populate this section with compliance findings once the RFP analysis is complete.'
            ],
            'sub_categories': [
                {'id': 'comp-matrix', 'title': 'Compliance Matrix'},
                {'id': 'comp-deliverables', 'title': 'Deliverables List'},
                {'id': 'comp-attachments', 'title': 'Required Attachments'},
                {'id': 'comp-forms', 'title': 'Required Forms'}
            ]
        })

        sections.append({
            'id': 'section-scope-objectives',
            'title': 'Understanding of Scope & Objectives',
            'status': 'In Progress',
            'guidance': 'Demonstrate a clear understanding of the project scope, objectives, and expected outcomes.',
            'content': scope_blocks or [
                'Detail your understanding of the project scope and objectives based on the solicitation requirements.'
            ],
            'sub_categories': [
                {'id': 'scope-project-scope', 'title': 'Project Scope'},
                {'id': 'scope-objectives', 'title': 'Project Objectives'},
                {'id': 'scope-expected-outcomes', 'title': 'Expected Outcomes'},
                {'id': 'scope-assumptions', 'title': 'Key Assumptions'}
            ]
        })

        sections.append({
            'id': 'section-deviations',
            'title': 'Deviations, Assumptions, and Dependencies (if any)',
            'status': 'Draft',
            'guidance': 'Document any deviations from the solicitation requirements, key assumptions, and dependencies that may impact project delivery.',
            'content': [
                'List any deviations from the solicitation requirements, if applicable.',
                'Document key assumptions made during proposal development.',
                'Identify dependencies that may affect project execution.'
            ],
            'sub_categories': [
                {'id': 'dev-deviations', 'title': 'Deviations'},
                {'id': 'dev-assumptions', 'title': 'Assumptions'},
                {'id': 'dev-dependencies', 'title': 'Dependencies'},
                {'id': 'dev-risks', 'title': 'Risk Considerations'}
            ]
        })

        tech_content = scope_blocks or [
            'Detail the technical solution, materials, and implementation methodology once the RFP requirements are parsed.'
        ]
        sections.append({
            'id': 'section-technical',
            'title': 'Proposed Technical Solution',
            'status': 'In Progress' if scope_blocks else 'Draft',
            'guidance': 'Translate the solicitation Statement of Work into your delivery plan, highlighting compliance, innovations, and risk mitigations.',
            'content': tech_content,
            'sub_categories': [
                {'id': 'tech-approach', 'title': 'Technical Approach'},
                {'id': 'tech-methodology', 'title': 'Methodology'},
                {'id': 'tech-innovations', 'title': 'Innovations & Best Practices'},
                {'id': 'tech-compliance', 'title': 'Compliance & Standards'},
                {'id': 'tech-risk-mitigation', 'title': 'Risk Mitigation'}
            ]
        })

        sections.append({
            'id': 'section-implementation',
            'title': 'Implementation & Work Plan',
            'status': 'In Progress',
            'guidance': 'Provide a detailed implementation plan and work breakdown structure for project execution.',
            'content': [
                'Outline the step-by-step implementation approach.',
                'Detail the work breakdown structure and key activities.',
                'Describe the workflow and processes for project delivery.'
            ],
            'sub_categories': [
                {'id': 'impl-approach', 'title': 'Implementation Approach'},
                {'id': 'impl-work-breakdown', 'title': 'Work Breakdown Structure'},
                {'id': 'impl-phases', 'title': 'Implementation Phases'},
                {'id': 'impl-processes', 'title': 'Key Processes'}
            ]
        })

        sections.append({
            'id': 'section-deliverables',
            'title': 'Solution Deliverables',
            'status': 'In Progress',
            'guidance': 'List all deliverables that will be provided as part of the solution, including documentation, reports, and physical items.',
            'content': [
                'Comprehensive list of all solution deliverables.',
                'Documentation and reports to be provided.',
                'Physical deliverables, if applicable.'
            ],
            'sub_categories': [
                {'id': 'del-documentation', 'title': 'Documentation Deliverables'},
                {'id': 'del-reports', 'title': 'Reports & Analysis'},
                {'id': 'del-physical', 'title': 'Physical Deliverables'},
                {'id': 'del-schedule', 'title': 'Delivery Schedule'}
            ]
        })

        management_content = compact_list([
            f"Target submission date: {due_text}." if due_text else 'Submission date will be updated when the procurement schedule is confirmed.',
          
            f"Contract type guidance: {type_text}." if type_text else '',
            "Outline major milestones, design reviews, and government touchpoints that support a compliant response."
        ])
        sections.append({
            'id': 'section-management',
            'title': 'Project Management Plan',
            'status': 'In Progress',
            'guidance': 'Map the internal review cycle, color team dates, and delivery milestones to ensure on-time submission.',
            'content': management_content or [
                'Add the project management narrative and schedule once the capture calendar is finalized.'
            ],
            'sub_categories': [
                {'id': 'mgmt-organization', 'title': 'Project Organization'},
                {'id': 'mgmt-schedule', 'title': 'Project Schedule & Milestones'},
                {'id': 'mgmt-resources', 'title': 'Resource Management'},
                {'id': 'mgmt-quality', 'title': 'Quality Assurance'},
                {'id': 'mgmt-communication', 'title': 'Communication Plan'}
            ]
        })

        price_content = compact_list([
            f"Estimated contract value: {revenue_text}." if revenue_text else '',
            "Attachment A pricing will mirror the government-provided workbook with CLIN-level details.",
            "Document all pricing assumptions, escalation factors, and exclusions for contracting officer review."
        ])
        sections.append({
            'id': 'section-pricing',
            'title': 'Pricing Proposals',
            'status': 'Saved' if revenue_text else 'Draft',
            'guidance': 'Ensure Attachment A reflects current quantities, rates, and assumptions. Cross-check CLIN totals before submission.',
            'content': price_content or [
                'Populate Attachment A once final quantities and labor categories are confirmed.'
            ],
            'sub_categories': [
                {'id': 'price-clin-breakdown', 'title': 'CLIN Breakdown'},
                {'id': 'price-labor-rates', 'title': 'Labor Rates & Categories'},
                {'id': 'price-materials', 'title': 'Materials & Equipment'},
                {'id': 'price-assumptions', 'title': 'Pricing Assumptions'},
                {'id': 'price-total-summary', 'title': 'Total Cost Summary'}
            ]
        })

        sections.append({
            'id': 'section-contractual',
            'title': 'Contractual Terms',
            'status': 'Saved',
            'guidance': 'Outline key contractual terms, conditions, and agreements proposed for this engagement.',
            'content': [
                'Key contractual terms and conditions.',
                'Payment terms and conditions.',
                'Terms of service and deliverables.'
            ],
            'sub_categories': [
                {'id': 'contract-terms', 'title': 'Terms & Conditions'},
                {'id': 'contract-payment', 'title': 'Payment Terms'},
                {'id': 'contract-service', 'title': 'Service Terms'},
                {'id': 'contract-legal', 'title': 'Legal Considerations'}
            ]
        })

        sections.append({
            'id': 'section-project-schedule',
            'title': 'Project Schedule & Milestones',
            'status': 'Pending',
            'guidance': 'Confirm the bond amount, surety approval, and Treasury Circular 570 listing for the guarantor.',
            'content': [
                f"{company_display} will coordinate with the approved surety to furnish the required Project Schedule & Milestones prior to submission.",
                "Include the original bid bond or irrevocable letter of credit documentation in the final package."
            ],
           
        })

        sections.append({
            'id': 'section-team',
            'title': 'Project Team',
            'status': 'In Progress',
            'guidance': 'Introduce the project team members, their roles, qualifications, and responsibilities.',
            'content': [
                'List of key team members and their roles.',
                'Qualifications and experience of team members.',
                'Organizational structure and reporting lines.'
            ],
            'sub_categories': [
                {'id': 'team-members', 'title': 'Team Members'},
                {'id': 'team-roles', 'title': 'Roles & Responsibilities'},
                {'id': 'team-qualifications', 'title': 'Qualifications'},
                {'id': 'team-organization', 'title': 'Organization Structure'}
            ]
        })

        sections.append({
            'id': 'section-corporate',
            'title': 'Corporate Qualifications',
            'status': 'Saved',
            'guidance': 'Highlight corporate experience, past performance, certifications, and qualifications relevant to this opportunity.',
            'content': [
                'Company background and history.',
                'Relevant past performance and case studies.',
                'Certifications, accreditations, and qualifications.'
            ],
            'sub_categories': [
                {'id': 'corp-background', 'title': 'Company Background'},
                {'id': 'corp-past-performance', 'title': 'Past Performance'},
                {'id': 'corp-certifications', 'title': 'Certifications & Accreditations'},
                {'id': 'corp-capabilities', 'title': 'Key Capabilities'}
            ]
        })

        sections.append({
            'id': 'section-reps-certs',
            'title': 'Representations and Certifications',
            'status': 'Saved',
            'guidance': 'Confirm FAR, DFARS, SAM.gov, and agency-specific representations are current and included.',
            'content': [
                'All FAR and DFARS representations will be validated against the active SAM registration prior to submission.',
                'Include any agency-specific attestations or forms requested in the solicitation instructions.'
            ],
            'sub_categories': [
                {'id': 'reps-far', 'title': 'FAR Representations'},
                {'id': 'reps-dfars', 'title': 'DFARS Representations'},
                {'id': 'reps-sam', 'title': 'SAM.gov Certifications'},
                {'id': 'reps-agency', 'title': 'Agency-Specific Certifications'}
            ]
        })

        sections.append({
            'id': 'section-amendments',
            'title': 'Addenda',
            'status': 'Pending',
            'guidance': 'Log each amendment number and release date. If no amendments were issued, include a statement confirming that status.',
            'content': [
                f"The offer or acknowledges receipt of all amendments released for {project_name} as of {datetime.now().strftime('%B %d, %Y')}." if project_name else
                f"The offer or acknowledges receipt of all solicitation amendments issued as of {datetime.now().strftime('%B %d, %Y')}.",
                "Update this section with individual amendment numbers and dates when the contracting officer publishes them."
            ],
            'sub_categories': [
                {'id': 'amend-list', 'title': 'Amendment List'},
                {'id': 'amend-acknowledgment', 'title': 'Acknowledgment Statement'},
                {'id': 'amend-impact', 'title': 'Impact Assessment'}
            ]
        })

        sections.append({
            'id': 'section-appendix',
            'title': 'Appendix',
            'status': 'Draft',
            'guidance': 'Include all supporting documents, references, and additional materials referenced in the proposal.',
            'content': [
                'Supporting documents and references.',
                'Additional technical specifications.',
                'Supplementary materials and exhibits.'
            ],
            'sub_categories': [
                {'id': 'appendix-documents', 'title': 'Supporting Documents'},
                {'id': 'appendix-references', 'title': 'References'},
                {'id': 'appendix-exhibits', 'title': 'Exhibits'},
                {'id': 'appendix-other', 'title': 'Other Materials'}
            ]
        })

        return sections

    def build_bid_context(bid_info, fallback_title, company_name):
        info = bid_info or {}
        project_name = safe_str(info.get('b_name')) or safe_str(fallback_title)
        due_text = format_date_for_display(info.get('due_date'))
        summary_text = safe_str(info.get('summary'))
        scope_text = safe_str(info.get('scope'))
        location_text = safe_str(info.get('location') or info.get('city') or info.get('state'))
        type_text = safe_str(info.get('type'))
        stage_text = safe_str(info.get('state'))
        revenue_value = info.get('revenue') if info.get('revenue') not in (None, '') else info.get('value')
        scoring_value = info.get('scoring')

        revenue_text = ''
        if revenue_value not in (None, ''):
            try:
                revenue_text = f"${float(revenue_value):,.0f}"
            except Exception:
                revenue_text = safe_str(revenue_value)

        scoring_text = ''
        if scoring_value not in (None, ''):
            try:
                scoring_text = f"{float(scoring_value):.0f}"
            except Exception:
                scoring_text = safe_str(scoring_value)

        summary_excerpt = ''
        source_text = summary_text or scope_text
        if source_text:
            summary_excerpt = source_text if len(source_text) <= 220 else source_text[:217].rstrip() + '…'

        info_chips = []
        if due_text:
            info_chips.append({'label': 'Due', 'value': due_text, 'icon': 'calendar-days'})
        if type_text:
            info_chips.append({'label': 'Type', 'value': type_text, 'icon': 'diagram-project'})
        if location_text:
            info_chips.append({'label': 'Location', 'value': location_text, 'icon': 'location-dot'})
        if revenue_text:
            info_chips.append({'label': 'Est. Value', 'value': revenue_text, 'icon': 'sack-dollar'})
        if stage_text:
            info_chips.append({'label': 'Stage', 'value': stage_text.title() if stage_text else stage_text, 'icon': 'gauge-high'})
        if scoring_text:
            info_chips.append({'label': 'Score', 'value': scoring_text, 'icon': 'star-half-stroke'})

        context = {
            'project_name': project_name or safe_str(fallback_title),
            'client': safe_str(
                info.get('agency') or info.get('customer') or info.get('client') or info.get('owner')
            ),
            'company': safe_str(company_name) or safe_str(info.get('company')),
            'due_date_display': due_text,
            'summary_excerpt': summary_excerpt,
            'scope': scope_text,
            'type': type_text,
            'location': location_text,
            'rfp_label': project_name or safe_str(fallback_title),
            'info_chips': info_chips,
            'bid_id': info.get('g_id') or info.get('id'),
            'g_id': info.get('g_id')
        }
        # Remove empty values for cleaner JSON payloads
        return {k: v for k, v in context.items() if v}

    bid_meta_raw = None
    if bid_id:
        meta_cursor = mysql.connection.cursor(DictCursor)
        try:
            meta_cursor.execute("SELECT * FROM go_bids WHERE g_id=%s", (bid_id,))
            bid_meta_raw = meta_cursor.fetchone()
        except Exception as meta_err:
            print(f"Error loading bid context for proposal view: {meta_err}")
        finally:
            meta_cursor.close()

    if not company and bid_meta_raw:
        company = safe_str(bid_meta_raw.get('company'))
    # Normalize company so aliases like 'METCO Engineering' or 'Metco' select the Metco dashboard/workflow
    company = normalize_company_choice(company)

    sections_outline = build_sections_outline(bid_meta_raw, bid_name, company or (bid_meta_raw or {}).get('company'), contact_email)
    bid_context = build_bid_context(bid_meta_raw, bid_name, company or (bid_meta_raw or {}).get('company'))

    # Load proposals from database
    proposals_list = []
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Create table if it doesn't exist
        cur.execute("""
            CREATE TABLE IF NOT EXISTS proposals (
                id INT AUTO_INCREMENT PRIMARY KEY,
                unique_id VARCHAR(50) NOT NULL UNIQUE,
                rfp_name VARCHAR(500) NOT NULL,
                rfp_filename VARCHAR(500),
                proposal_filename VARCHAR(500) NOT NULL,
                proposal_path VARCHAR(1000) NOT NULL,
                company VARCHAR(100) NOT NULL,
                template VARCHAR(100),
                bid_id INT NULL,
                bid_name VARCHAR(500) NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                INDEX idx_unique_id (unique_id),
                INDEX idx_rfp_name (rfp_name),
                INDEX idx_company (company),
                INDEX idx_bid_id (bid_id)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """)
        
        # Query proposals from database
        if bid_id:
            cur.execute("""
                SELECT * FROM proposals 
                WHERE bid_id = %s 
                ORDER BY created_at DESC
            """, (bid_id,))
        else:
            cur.execute("""
                SELECT * FROM proposals 
                ORDER BY created_at DESC
                LIMIT 100
            """)
        
        db_proposals = cur.fetchall()
        
        # Convert database results to list format
        for prop in db_proposals:
            proposals_list.append({
                'id': prop['id'],
                'unique_id': prop['unique_id'],
                'name': prop['proposal_filename'],
                'rfp_name': prop['rfp_name'],
                'company': prop['company'],
                'template': prop.get('template', ''),
                'bid_id': prop.get('bid_id'),
                'bid_name': prop.get('bid_name', ''),
                'created_at': prop['created_at'].strftime('%Y-%m-%d %H:%M:%S') if prop['created_at'] else '',
                'download_url': url_for('download_proposal', filename=prop['proposal_filename'])
            })
    except Exception as e:
        print(f"Error loading proposals from database: {str(e)}")
        # Fallback to file system if database fails
        proposals_dir = 'uploads/proposals'
        if os.path.exists(proposals_dir):
            for filename in os.listdir(proposals_dir):
                if filename.endswith('.docx'):
                    filepath = os.path.join(proposals_dir, filename)
                    stat = os.stat(filepath)
                    proposals_list.append({
                        'name': filename,
                        'rfp_name': filename.replace('executive_summary_', '').replace('.docx', '.pdf'),
                        'created_at': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
                        'download_url': url_for('download_proposal', filename=filename)
                    })
    finally:
        cur.close()
    
    # Load RFP files from database for this bid
    rfp_files_list = []
    if bid_id:
        rfp_cursor = mysql.connection.cursor(DictCursor)
        try:
            _ensure_uploaded_rfp_table_exists(rfp_cursor)
            # Fetch only true RFP source PDFs for this bid (exclude per-section attachments and non-PDFs)
            rfp_cursor.execute("""
                SELECT id, g_id, bid_id, filename, original_filename, saved_filename, file_path, file_size, uploaded_at, file_type, section_id
                FROM uploaded_rfp_files
                WHERE (g_id = %s OR bid_id = %s)
                  AND (section_id IS NULL OR section_id = '')
                  AND (
                        LOWER(COALESCE(file_type, '')) = 'pdf'
                     OR LOWER(filename) LIKE '%%.pdf'
                     OR LOWER(saved_filename) LIKE '%%.pdf'
                     OR LOWER(file_path) LIKE '%%.pdf'
                  )
                ORDER BY uploaded_at DESC
            """, (bid_id, bid_id))
            rfp_files = rfp_cursor.fetchall()
            
            for rfp_file in rfp_files:
                file_path = rfp_file.get('file_path')
                if file_path and os.path.exists(file_path):
                    rfp_files_list.append({
                        'id': rfp_file.get('id'),
                        'filename': rfp_file.get('filename') or rfp_file.get('original_filename', ''),
                        'file_path': file_path,
                        'file_size': rfp_file.get('file_size', 0),
                        'uploaded_at': rfp_file.get('uploaded_at'),
                        'view_url': url_for('view_rfp_file', file_id=rfp_file.get('id')),
                        'parsed_url': url_for('api_rfp_file_parsed', g_id=bid_id)
                    })
        except Exception as e:
            print(f"Error loading RFP files from database: {str(e)}")
        finally:
            rfp_cursor.close()
    
    return render_template(
        'proposals_making.html',
        proposals=proposals_list,
        bid_id=bid_id,
        bid_name=bid_name,
        company=company,
        sections_outline=sections_outline,
        bid_context=bid_context,
        contact_email=contact_email,
        rfp_files=rfp_files_list
    )

@app.route('/api/section-attachments/<int:g_id>', methods=['GET'])
@login_required
def api_list_section_attachments(g_id):
    """Return JSON list of attachments for a given section of a bid."""
    section_id = request.args.get('section_id', '').strip()
    if not section_id:
        return jsonify({'success': False, 'message': 'section_id is required'}), 400
    limit = request.args.get('limit', type=int) or 30
    cur = mysql.connection.cursor(DictCursor)
    try:
        _ensure_uploaded_rfp_table_exists(cur)
        cur.execute(
            """
            SELECT id, filename, file_path, file_type, uploaded_at
            FROM uploaded_rfp_files
            WHERE (g_id = %s OR bid_id = %s)
              AND section_id = %s
              AND (
                   file_type IN ('image','pdf')
                OR LOWER(filename) LIKE '%%.png'
                OR LOWER(filename) LIKE '%%.jpg'
                OR LOWER(filename) LIKE '%%.jpeg'
                OR LOWER(filename) LIKE '%%.webp'
                OR LOWER(filename) LIKE '%%.pdf'
              )
            ORDER BY uploaded_at DESC
            LIMIT %s
            """,
            (g_id, g_id, section_id, limit),
        )
        rows = cur.fetchall() or []
        files = []
        for row in rows:
            files.append({
                'id': row.get('id'),
                'filename': row.get('filename'),
                'file_type': row.get('file_type') or '',
                'view_url': url_for('view_rfp_file', file_id=row.get('id')),
                'is_pdf': bool((row.get('file_type') or '').lower() == 'pdf' or str(row.get('filename') or '').lower().endswith('.pdf')),
            })
        return jsonify({'success': True, 'files': files})
    except Exception as exc:
        app.logger.warning(f"Error listing section attachments for {g_id}/{section_id}: {exc}")
        return jsonify({'success': False, 'message': 'Failed to load attachments'}), 500
    finally:
        cur.close()

@app.route('/api/proposals-agent', methods=['POST'])
@login_required
def proposals_agent():
    """Lightweight assistant endpoint for the proposals workspace."""
    payload = request.get_json(silent=True) or {}
    message = (payload.get('message') or '').strip()
    if not message:
        return jsonify({'error': 'missing_message', 'message': 'Please provide a question for the agent.'}), 400

    history = payload.get('history') or []
    if not isinstance(history, list):
        history = []
    history = history[-10:]

    context = payload.get('context') or {}
    try:
        context_summary = {
            'workspace': 'Proposals Making',
            'user': getattr(current_user, 'email', ''),
            'company': context.get('company') or '',
            'contact_email': context.get('contactEmail') or '',
        }
        bid_info = context.get('bid') or {}
        g_id = None
        if isinstance(bid_info, dict) and bid_info:
            g_id = bid_info.get('g_id') or bid_info.get('bid_id')
            context_summary['bid'] = {
                'project_name': bid_info.get('project_name') or bid_info.get('b_name') or '',
                'bid_id': bid_info.get('bid_id') or bid_info.get('g_id') or '',
                'due_date': bid_info.get('due_date') or '',
            }
        sections = context.get('sections') or []
        if isinstance(sections, list) and sections:
            context_summary['sections'] = [
                {
                    'title': sec.get('title', ''),
                    'status': sec.get('status', ''),
                    'guidance': (sec.get('guidance') or '')[:280],
                }
                for sec in sections[:5]
                if isinstance(sec, dict)
            ]
    except Exception:
        context_summary = {}

    # Retrieve RFP document content for RAG
    rfp_content = ""
    if g_id:
        cur = None
        try:
            cur = mysql.connection.cursor(DictCursor)
            _ensure_uploaded_rfp_table_exists(cur)
            # Get the latest RFP file for this bid
            rfp_file = _get_latest_rfp_file_for_bid(g_id)
            if rfp_file and rfp_file.get('file_path'):
                file_path = rfp_file['file_path']
                if os.path.exists(file_path):
                    # Extract text from first 50 pages (to avoid token limits)
                    pages, total_pages = _extract_pdf_pages(file_path, start_page=1, limit=50)
                    if pages:
                        # Combine all page texts
                        rfp_texts = [page.get('text', '') for page in pages if page.get('text')]
                        if rfp_texts:
                            # Limit total content to ~8000 characters to stay within token limits
                            combined_text = "\n\n--- PAGE BREAK ---\n\n".join(rfp_texts)
                            if len(combined_text) > 8000:
                                combined_text = combined_text[:8000] + "... [Content truncated]"
                            rfp_content = f"\n\n--- RFP DOCUMENT CONTENT (Pages 1-{min(50, total_pages)} of {total_pages}) ---\n{combined_text}\n--- END RFP CONTENT ---"
        except Exception as e:
            app.logger.warning(f"Error retrieving RFP content for RAG: {e}")
            rfp_content = ""
        finally:
            if cur:
                cur.close()

    system_prompt = (
        "You are the ESCO proposal workspace assistant. "
        "Provide concise, actionable support for crafting winning proposals. "
        "Use the workspace context when relevant, and be explicit when information is unavailable or uncertain."
    )
    if context_summary:
        try:
            system_prompt += "\n\nWorkspace context:\n" + json.dumps(context_summary, ensure_ascii=False, indent=2)
        except Exception:
            system_prompt += "\n\nWorkspace context:\n" + str(context_summary)
    
    # Add concise RFP summary for RAG (summarized locally via Ollama to save tokens)
    if rfp_content:
        try:
            summary = _ollama_chat(
                [
                    {'role': 'system', 'content': 'You are a concise technical summarizer for RFPs.'},
                    {'role': 'user', 'content': f"Summarize the key requirements, scope, deadlines, and evaluation criteria in <= 160 words:\n{rfp_content[:12000]}"}
                ],
                temperature=0.1,
                timeout_sec=20
            )
        except Exception:
            summary = ''
        if summary:
            system_prompt += "\n\nRFP summary (locally generated):\n" + summary.strip()

    messages = [{'role': 'system', 'content': system_prompt}]
    for entry in history:
        if not isinstance(entry, dict):
            continue
        role = entry.get('role', 'user')
        if role not in {'user', 'assistant', 'system'}:
            role = 'user'
        content = entry.get('content')
        if not content:
            continue
        messages.append({'role': role, 'content': str(content)})
    messages.append({'role': 'user', 'content': message})

    # Helper: try local Ollama as fallback
    def _try_ollama_fallback(messages_local, temperature_local=0.2):
        try:
            ollama_base = (app.config.get('OLLAMA_BASE_URL') or os.getenv('OLLAMA_BASE_URL') or 'http://localhost:11434').strip()
            ollama_model = (app.config.get('OLLAMA_MODEL') or os.getenv('OLLAMA_MODEL') or 'llama3').strip()
            endpoint_ollama = ollama_base.rstrip('/') + '/api/chat'
            headers_ollama = {'Content-Type': 'application/json'}
            body_ollama = {'model': ollama_model, 'messages': messages_local, 'stream': False, 'options': {'temperature': temperature_local}}
            resp = _get_ollama_session().post(endpoint_ollama, headers=headers_ollama, json=body_ollama, timeout=45)
            data_ollama = resp.json()
            if resp.status_code >= 400:
                app.logger.error("Ollama error %s: %s", resp.status_code, str(data_ollama)[:400])
                return None, None, None
            content_val = ''
            if isinstance(data_ollama, dict):
                if isinstance(data_ollama.get('message'), dict):
                    content_val = data_ollama['message'].get('content', '') or ''
                elif isinstance(data_ollama.get('messages'), list) and data_ollama['messages']:
                    content_val = (data_ollama['messages'][-1] or {}).get('content', '') or ''
                elif 'response' in data_ollama:
                    content_val = data_ollama.get('response', '') or ''
            if not content_val:
                return None, None, None
            return content_val, data_ollama, ollama_model
        except Exception as ollama_err:
            app.logger.warning("Ollama fallback failed: %s", ollama_err)
            return None, None, None

    # Attempt primary provider (OpenAI-compatible: OpenRouter/Together/Fireworks/OpenAI), then Ollama fallback
    api_key = (os.getenv('OPENAI_API_KEY') or os.getenv('OPENROUTER_API_KEY') or os.getenv('TOGETHER_API_KEY') or os.getenv('FIREWORKS_API_KEY') or '').strip()
    base_url = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or os.getenv('OPENROUTER_BASE_URL') or 'https://openrouter.ai/api/v1').strip()
    model = (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'meta-llama/llama-4-scout-17b-16e-instruct').strip()

    data = None
    reply = ''
    if api_key:
        endpoint = base_url.rstrip('/')
        if not endpoint.endswith('/chat/completions'):
            endpoint = f"{endpoint}/chat/completions"
        headers = {'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'}
        body = {'model': model, 'messages': messages, 'temperature': 0.2}
        try:
            response = requests.post(endpoint, headers=headers, json=body, timeout=60)
            raw_text = response.text
            try:
                data = response.json()
            except ValueError:
                app.logger.error("Primary provider returned non-JSON (status %s): %s", response.status_code, raw_text[:400])
                data = None
            if data and response.status_code < 400:
                choices = data.get('choices') or []
                if choices:
                    reply = choices[0].get('message', {}).get('content', '') or ''
        except requests.RequestException as err:
            app.logger.exception("Primary provider network failure: %s", err)

    # If primary failed or empty, try Ollama
    if not reply:
        content_val, data_ollama, ollama_model = _try_ollama_fallback(messages, temperature_local=0.2)
        if content_val:
            return jsonify({'reply': content_val.strip(), 'model': ollama_model, 'usage': data_ollama.get('usage', {}) if isinstance(data_ollama, dict) else {}})

        # If Ollama also failed and we have an upstream error payload, surface minimal error
        if isinstance(data, dict):
            err_message = data.get('error')
            detail = ''
            if isinstance(err_message, dict):
                detail = err_message.get('message') or err_message.get('code') or ''
            elif isinstance(err_message, str):
                detail = err_message
            if not detail:
                detail = 'Upstream error'
            return jsonify({'error': 'agent_unavailable', 'message': f'AI provider error: {detail}'}), 502
        return jsonify({'error': 'agent_unavailable', 'message': 'The proposal agent could not generate a response at this time.'}), 502

    return jsonify({'reply': reply.strip(), 'model': data.get('model', model) if isinstance(data, dict) else model, 'usage': data.get('usage', {}) if isinstance(data, dict) else {}})


def _fetch_rfp_content_for_bid(bid_id, page_limit=50, char_limit=12000):
    """Load RFP text from the uploaded files for a given bid or g_id."""
    if not bid_id:
        return ''

    rfp_content = ''
    try:
        cur = mysql.connection.cursor(DictCursor)
        cur.execute(
            """
            SELECT file_path
            FROM uploaded_rfp_files
            WHERE g_id = %s OR bid_id = %s
            ORDER BY uploaded_at DESC
            LIMIT 1
            """,
            (bid_id, bid_id),
        )
        rfp_file = cur.fetchone()
        cur.close()
    except Exception as exc:
        app.logger.warning(f"Error retrieving RFP content for bid {bid_id}: {exc}")
        return ''

    file_path = rfp_file.get('file_path') if rfp_file else None
    if not file_path or not os.path.exists(file_path):
        return ''

    try:
        if _FITZ_AVAILABLE:
            doc = _pymupdf.open(file_path)
            try:
                pages = []
                for page_num in range(min(page_limit, len(doc))):
                    page = doc[page_num]
                    pages.append(page.get_text())
                rfp_content = "\n\n--- PAGE BREAK ---\n\n".join(pages)
            finally:
                doc.close()
        else:
            app.logger.info("PyMuPDF not available; skipping PDF text extraction.")
    except Exception as exc:
        app.logger.warning(f"Error reading RFP PDF for bid {bid_id}: {exc}")
        rfp_content = ''

    if rfp_content and len(rfp_content) > char_limit:
        rfp_content = rfp_content[:char_limit] + "... [Content truncated]"

    return rfp_content


def _get_section_image_files_for_bid(bid_id, section_id, limit=10):
    """Return latest image file records for a bid and section."""
    if not bid_id or not section_id:
        return []
    try:
        cur = mysql.connection.cursor(DictCursor)
        cur.execute(
            """
            SELECT file_path, filename
            FROM uploaded_rfp_files
            WHERE (g_id = %s OR bid_id = %s)
              AND section_id = %s
              AND (file_type = 'image'
                   OR LOWER(filename) LIKE '%%.png'
                   OR LOWER(filename) LIKE '%%.jpg'
                   OR LOWER(filename) LIKE '%%.jpeg'
                   OR LOWER(filename) LIKE '%%.webp')
            ORDER BY uploaded_at DESC
            LIMIT %s
            """,
            (bid_id, bid_id, section_id, limit),
        )
        rows = cur.fetchall() or []
        cur.close()
        return rows
    except Exception as exc:
        app.logger.warning(f"Error fetching section images for bid {bid_id}: {exc}")
        return []


def _extract_text_from_image(image_path: str) -> str:
    """Best-effort OCR for image files; returns empty string on failure.

    Strategy:
    1) Try pytesseract with auto-detected tesseract.exe (Windows-friendly) and light preprocessing.
    2) Fallback to EasyOCR if available (no system tesseract dependency).
    3) Gracefully return '' on any error.
    """
    try:
        import os as _os
        import shutil as _shutil
        import importlib
        from PIL import Image as _PILImage, ImageOps as _ImageOps, ImageFilter as _ImageFilter

        # Open image
        pil_img = _PILImage.open(image_path)
        # Normalize mode and improve contrast for OCR
        if 'A' in pil_img.getbands():
            pil_img = pil_img.convert('RGB')  # type: ignore[assignment]
        try:
            # Convert to grayscale and apply slight sharpening to help OCR
            pil_img = _ImageOps.grayscale(pil_img)
            pil_img = pil_img.filter(_ImageFilter.UnsharpMask(radius=1, percent=150, threshold=3))
        except Exception:
            # Continue with original if preprocessing fails
            pass

        # Attempt pytesseract first
        try:
            _pytesseract = importlib.import_module('pytesseract')
            # If tesseract is not on PATH on Windows, try common install paths or env var
            tcmd_env = _os.getenv('TESSERACT_CMD') or _os.getenv('TESSERACT_EXE')
            probable_paths = [
                tcmd_env,
                r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            ]
            found_cmd = None
            for p in probable_paths:
                if p and _os.path.exists(p):
                    found_cmd = p
                    break
            if not found_cmd:
                found_cmd = _shutil.which('tesseract')
            if found_cmd:
                try:
                    _pytesseract.pytesseract.tesseract_cmd = found_cmd  # type: ignore[attr-defined]
                except Exception:
                    pass

            text = _pytesseract.image_to_string(pil_img) or ''  # type: ignore[attr-defined]
            text = (text or '').strip()
            if text:
                return text
        except Exception as _pt_err:
            app.logger.info(f"pytesseract OCR not available or failed for {image_path}: {_pt_err}")

        # Fallback to EasyOCR if installed (may be heavy, so optional)
        try:
            _easyocr = importlib.import_module('easyocr')
            reader = _easyocr.Reader(['en'], gpu=False)  # type: ignore[call-arg]
            # EasyOCR works with file paths directly
            result = reader.readtext(image_path, detail=0) or []  # type: ignore[assignment]
            text = '\n'.join([seg for seg in result if isinstance(seg, str)]).strip()
            if text:
                return text
        except Exception as _eo_err:
            app.logger.info(f"EasyOCR fallback not available or failed for {image_path}: {_eo_err}")

    except Exception as exc:
        app.logger.info(f"OCR pipeline failed for {image_path}: {exc}")
    return ''


def _extract_text_from_pdf(pdf_path: str, page_limit: int = 5) -> str:
    """Extract text from a PDF using PyMuPDF if available; returns '' on failure."""
    try:
        # These are defined earlier in the file for PDF support; reuse if present
        _has = globals().get('_FITZ_AVAILABLE', False)
        _mod = globals().get('_pymupdf')
        if not _has or not _mod:
            return ''
        text_parts = []
        with _mod.open(pdf_path) as doc:
            total = min(page_limit, doc.page_count or 0)
            for i in range(total):
                page = doc.load_page(i)
                text_parts.append(page.get_text('text') or '')
        return '\n'.join(text_parts).strip()
    except Exception as exc:
        app.logger.info(f"PDF extraction failed for {pdf_path}: {exc}")
        return ''


def _get_section_files_for_bid(bid_id, section_id, limit=20):
    """Return recent uploaded files for a specific section (images and PDFs)."""
    cur = mysql.connection.cursor(DictCursor)
    try:
        _ensure_uploaded_rfp_table_exists(cur)
        cur.execute(
            """
            SELECT id, filename, file_path, file_type, uploaded_at
            FROM uploaded_rfp_files
            WHERE (g_id = %s OR bid_id = %s)
              AND section_id = %s
              AND (
                   file_type IN ('image','pdf')
                   OR LOWER(filename) LIKE '%%.png'
                   OR LOWER(filename) LIKE '%%.jpg'
                   OR LOWER(filename) LIKE '%%.jpeg'
                   OR LOWER(filename) LIKE '%%.webp'
                   OR LOWER(filename) LIKE '%%.pdf'
              )
            ORDER BY uploaded_at DESC
            LIMIT %s
            """,
            (bid_id, bid_id, section_id, limit),
        )
        rows = cur.fetchall() or []
        cur.close()
        return rows
    except Exception as exc:
        app.logger.warning(f"Error fetching section files for bid {bid_id}: {exc}")
        return []


def _fetch_section_attachments_text(bid_id, section_id, char_limit=8000):
    """Aggregate text from section attachments (images + PDFs)."""
    attachments = _get_section_files_for_bid(bid_id, section_id, limit=20)
    if not attachments:
        return ''
    parts = []
    for idx, item in enumerate(attachments, start=1):
        fp = item.get('file_path')
        if not fp or not os.path.exists(fp):
            continue
        label = item.get('filename') or os.path.basename(fp)
        ftype = (item.get('file_type') or '').lower()
        extracted = ''
        try:
            if ftype == 'pdf' or fp.lower().endswith('.pdf'):
                extracted = _extract_text_from_pdf(fp, page_limit=5)
            else:
                extracted = _extract_text_from_image(fp)
        except Exception:
            extracted = ''
        if extracted:
            parts.append(f"[{idx}] {label}:\n{extracted}")
        else:
            parts.append(f"[{idx}] {label}: [No text detected]")
        if sum(len(p) for p in parts) > char_limit:
            break
    combined = "\n\n--- ATTACHMENT BREAK ---\n\n".join(parts)
    if len(combined) > char_limit:
        combined = combined[:char_limit] + "... [Content truncated]"
    return combined


def _fetch_section_images_text(bid_id, section_id, char_limit=6000):
    """Aggregate OCR text from section image attachments."""
    attachments = _get_section_image_files_for_bid(bid_id, section_id, limit=12)
    if not attachments:
        return ''
    parts = []
    for idx, item in enumerate(attachments, start=1):
        fp = item.get('file_path')
        if not fp or not os.path.exists(fp):
            continue
        ocr_text = _extract_text_from_image(fp)
        label = item.get('filename') or os.path.basename(fp)
        if ocr_text:
            parts.append(f"[{idx}] {label}:\n{ocr_text}")
        else:
            parts.append(f"[{idx}] {label}: [No text detected]")
        if sum(len(p) for p in parts) > char_limit:
            break
    combined = "\n\n--- IMAGE BREAK ---\n\n".join(parts)
    if len(combined) > char_limit:
        combined = combined[:char_limit] + "... [Content truncated]"
    return combined


@app.route('/api/analyze-rfp-master', methods=['POST'])
@login_required
def analyze_rfp_master():
    """Run an end-to-end compliance and outline analysis for the active RFP."""
    payload = request.get_json(silent=True) or {}
    bid_id = payload.get('bid_id')

    if not bid_id:
        return jsonify({'error': 'missing_params', 'message': 'Bid ID is required.'}), 400

    rfp_content = _fetch_rfp_content_for_bid(bid_id, page_limit=120, char_limit=22000)
    if not rfp_content:
        return jsonify({'error': 'no_rfp', 'message': 'No RFP content found for this bid. Please upload an RFP file first.'}), 404

    truncated_rfp = rfp_content[:22000]

    analysis_instructions = """
You are an expert Proposal Analyst and Writer specialized in public procurement (RFP/RFQ/IFB) across construction, energy, IT, and professional services. Using the provided RFP corpus, produce a compliance-first analysis and proposal blueprint.

Requirements:
- Output valid HTML only. No Markdown, JSON, or plain text.
- Use <h2> headings in this exact order:
  1. Executive Brief
  2. Key Dates and Submission Rules
  3. Compliance Matrix
  4. Proposal Table of Contents
  5. Section-by-Section Guidance and Drafts
  6. Pricing Package Instructions
  7. SBE/DBE and Compliance Plan
  8. Bonds, Insurance, and Financials
  9. Risk Register and Clarifications
  10. Submission Checklist
- Provide concise, action-oriented guidance. Quote or paraphrase RFP sections with citations (e.g., Section 4, Instructions to Bidders).
- The Compliance Matrix must be a <table> with columns: Requirement, RFP Reference, Mandatory?, Proposal Location, Owner Role, Notes. Keep rows CSV-friendly (no extra commas unless inside quotes).
- In Section-by-Section Guidance, create a <section> per entry that includes:
    * <h3> heading for the section name.
    * <p> summarizing purpose and linkage.
    * <ul> of required content bullets.
    * <ul class="placeholders"> for data placeholders (e.g., [Vendor_Name]).
    * <pre class="boilerplate"> containing draft boilerplate text with placeholders.
    * <ul class="evidence"> listing evidence artifacts/forms to attach.
- Supply copy-paste templates wrapped in <pre class="template"> blocks for: Cover Letter, Executive Summary, Technical Approach, Safety Program, Small Business Plan, Pricing Cover, Bonds/Insurance Acknowledgment, References, Appendices List.
- Flag assumptions as “Assumption – verify”.
- If SBE/DBE goals exist, include a subcontracting plan template and tracking approach.
- Emphasize sealed bid compliance, bonding thresholds, signatures/notarization, pricing forms, safety stats, and portal rules when relevant to construction/MEP.
- Highlight phasing, technical standards, monitoring, and access coordination for solar/energy scopes when applicable.
- Never invent data; if missing, call it out explicitly.
- Finish with a concise checklist aligned to submission portal/file rules and include a Go/No-Go gate summary.
"""

    user_prompt = f"""
RFP / IFB Source Content:
{truncated_rfp}

Generate the HTML deliverable now, following every rule above.
"""

    # Use OpenAI for master analysis
    api_key = (app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or os.getenv('OPENROUTER_API_KEY') or os.getenv('TOGETHER_API_KEY') or os.getenv('FIREWORKS_API_KEY') or '').strip()
    if not api_key:
        return jsonify({'error': 'api_unavailable', 'message': 'OpenAI API key is not configured on the server.'}), 500
    base_url = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or os.getenv('OPENROUTER_BASE_URL') or 'https://openrouter.ai/api/v1').strip()
    model = (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'meta-llama/llama-4-scout-17b-16e-instruct').strip()
    endpoint = base_url.rstrip('/')
    if not endpoint.endswith('/chat/completions'):
        endpoint = f"{endpoint}/chat/completions"
    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json',
    }
    body = {
        'model': model,
        'messages': [
            {'role': 'system', 'content': 'You are an expert proposal analyst. Return ONLY valid HTML per instructions.'},
            {'role': 'user', 'content': analysis_instructions + "\n\n" + user_prompt}
        ],
        'temperature': 0.2,
        'max_tokens': 6000
    }
    try:
        response = requests.post(endpoint, headers=headers, json=body, timeout=120)
    except requests.RequestException as err:
        app.logger.exception("OpenAI master analysis network failure: %s", err)
        return jsonify({'error': 'api_unavailable', 'message': 'Could not reach the AI service. Please try again shortly.'}), 502
    try:
        data = response.json()
    except ValueError:
        app.logger.error("OpenAI master analysis returned non-JSON response (status %s): %s", response.status_code, response.text[:400])
        return jsonify({'error': 'api_unavailable', 'message': 'Received an invalid response from the AI service.'}), 502
    if response.status_code >= 400:
        err_message = data.get('error') if isinstance(data, dict) else None
        detail = ''
        if isinstance(err_message, dict):
            detail = err_message.get('message') or err_message.get('code') or ''
        elif isinstance(err_message, str):
            detail = err_message
        if not detail:
            detail = f"Upstream error {response.status_code}"
        app.logger.error("OpenAI master analysis error %s: %s", response.status_code, detail)
        return jsonify({'error': 'api_unavailable', 'message': f'AI service returned an error: {detail}'}), 502
    choices = data.get('choices') or []
    analysis_html = (choices[0].get('message', {}) or {}).get('content', '') if choices else ''

    if not analysis_html:
        analysis_html = "<p class='text-sm text-gray-500'>No analysis output was generated.</p>"

    return jsonify({'analysis_html': analysis_html})


@app.route('/api/analyze-subcategory', methods=['POST'])
@login_required
def analyze_subcategory():
    """Analyze a sub-category section using AI based on RFP content."""
    payload = request.get_json(silent=True) or {}
    section_id = payload.get('section_id', '').strip()
    subcategory_id = payload.get('subcategory_id', '').strip()
    subcategory_title = payload.get('subcategory_title', '').strip()
    bid_id = payload.get('bid_id')
    rfp_content = payload.get('rfp_content', '')
    
    if not section_id or not subcategory_id or not subcategory_title:
        return jsonify({'error': 'missing_params', 'message': 'Section ID, subcategory ID, and title are required.'}), 400
    
    # Get RFP content from database if bid_id is provided
    if bid_id and not rfp_content:
        rfp_content = _fetch_rfp_content_for_bid(bid_id, page_limit=50, char_limit=8000)
    
    # Build AI prompt for sub-category analysis
    system_prompt = (
        "You are an expert proposal writer helping to draft specific sections of a government proposal. "
        "Analyze the RFP content and generate comprehensive, compliant content for the requested sub-category. "
        "Your response should be professional, detailed, and aligned with the RFP requirements."
    )
    
    user_prompt = f"""Generate content for the sub-category: "{subcategory_title}" within the section "{section_id}".

Requirements:
1. Analyze the provided RFP content to understand requirements relevant to this sub-category
2. Generate professional, compliant content that addresses the sub-category topic
3. Ensure the content is specific, actionable, and aligned with government proposal standards
4. If RFP content is limited, provide a well-structured template that can be customized

RFP Content:
{rfp_content[:8000] if rfp_content else "No RFP content available. Generate a professional template for this sub-category."}

Please provide the content for "{subcategory_title}" in a clear, well-formatted manner."""
    
    api_key = (app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or os.getenv('OPENROUTER_API_KEY') or os.getenv('TOGETHER_API_KEY') or os.getenv('FIREWORKS_API_KEY') or '').strip()
    if not api_key:
        return jsonify({'error': 'api_unavailable', 'message': 'OpenAI API key is not configured on the server.'}), 500
    
    base_url = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or os.getenv('OPENROUTER_BASE_URL') or 'https://openrouter.ai/api/v1').strip()
    model = (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'meta-llama/llama-4-scout-17b-16e-instruct').strip()
    endpoint = base_url.rstrip('/')
    if not endpoint.endswith('/chat/completions'):
        endpoint = f"{endpoint}/chat/completions"
    
    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json',
    }
    body = {
        'model': model,
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt}
        ],
        'temperature': 0.3,
    }
    
    try:
        response = requests.post(endpoint, headers=headers, json=body, timeout=90)
    except requests.RequestException as err:
        app.logger.exception("OpenAI subcategory analysis network failure: %s", err)
        return jsonify({'error': 'api_unavailable', 'message': 'Could not reach the AI service. Please try again shortly.'}), 502
    
    try:
        data = response.json()
    except ValueError:
        app.logger.error("OpenAI subcategory analysis returned non-JSON response (status %s): %s", response.status_code, response.text[:400])
        return jsonify({'error': 'api_unavailable', 'message': 'Received an invalid response from the AI service.'}), 502
    
    if response.status_code >= 400:
        err_message = data.get('error') if isinstance(data, dict) else None
        detail = ''
        if isinstance(err_message, dict):
            detail = err_message.get('message') or err_message.get('code') or ''
        elif isinstance(err_message, str):
            detail = err_message
        if not detail:
            detail = f"Upstream error {response.status_code}"
        app.logger.error("OpenAI subcategory analysis error %s: %s", response.status_code, detail)
        return jsonify({'error': 'api_unavailable', 'message': f'AI service returned an error: {detail}'}), 502
    
    choices = data.get('choices') or []
    if not choices:
        return jsonify({'error': 'api_unavailable', 'message': 'No response returned by the AI service.'}), 502
    
    content = choices[0].get('message', {}).get('content', '')
    if not content:
        return jsonify({'error': 'api_unavailable', 'message': 'The AI service returned an empty response.'}), 502
    
    return jsonify({
        'content': content.strip(),
        'subcategory_id': subcategory_id,
        'subcategory_title': subcategory_title,
        'model': data.get('model', model),
        'usage': data.get('usage', {}),
    })


@app.route('/api/analyze-section-outline', methods=['POST'])
@login_required
def analyze_section_outline():
    """Extract key requirements and a summary for a proposal section directly from the RFP."""
    payload = request.get_json(silent=True) or {}
    section_id = (payload.get('section_id') or '').strip()
    section_title = (payload.get('section_title') or '').strip() or section_id or 'Proposal Section'
    bid_id = payload.get('bid_id')
    rfp_content = payload.get('rfp_content', '')

    if not section_id:
        return jsonify({'error': 'missing_params', 'message': 'Section ID is required.'}), 400

    if bid_id and not rfp_content:
        rfp_content = _fetch_rfp_content_for_bid(bid_id, page_limit=60, char_limit=12000)

    truncated_rfp = rfp_content[:12000] if rfp_content else ''

    system_prompt = (
        "You are an expert RFP and bid document analyst. "
        "Read RFP, IFB, and tender documents carefully and extract only the actionable requirements, eligibility criteria, "
        "and specifications needed for a compliant proposal submission. Present results in clear proposal-ready language."
    )

    user_prompt = f"""
Carefully read the following RFP content and produce an output that follows the exact template below. Use concise bullet points. Do not invent information that is not present.

Template to follow:

📘 **Project Overview**
- Project title, number, client/issuing agency, location
- Procurement officer/contact email
- Submission portal/platform
- Important bid dates and schedule (release, pre-bid, submission, opening, award)
- Estimated project value or budget (if mentioned)

📋 **Scope of Work / Technical Requirements**
- Summarize required deliverables, materials, installation, testing, safety, coordination, warranty, and quality standards.

🧾 **Eligibility & Qualification Requirements**
- Experience, licenses, certifications, references, safety, financial, insurance, staffing requirements, mandatory forms, bonds.

⚖️ **Bid Submission Requirements**
- Submission method, required documents, pricing forms, conflict forms, bonding instructions, clarification procedures.

📊 **Evaluation & Award Criteria**
- List all evaluation factors, scoring parameters, participation goals, best-value or low-bid rules.

💼 **Contract & Performance Requirements**
- Contract term, start/completion dates, liquidated damages, compliance obligations, minimum self-performed work.

🔒 **Bonds, Warranty & Insurance**
- Bid/performance/payment bonds with percentages, warranty coverage, insurance types and limits.

💡 **Added Value / Special Notes**
- Unique clauses, sustainability requirements, community programs, optional value-add proposals.

🗓️ **Key Dates Summary Table**
| Milestone | Date |
|------------|-------|
| IFB/RFP Release | |
| Pre-Bid Meeting | |
| Questions Deadline | |
| Bid Submission | |
| Bid Opening | |
| Board/Contract Award | |

✅ **Summary for Proposal Preparation**
- Summarize everything a bidder must ensure for compliance (eligibility, documents, bonds, pricing, schedule, goals).

Formatting rules:
- Use the headings exactly as shown.
- Use bullet points for lists; keep them short and factual.
- Skip any sub-bullet if the RFP does not mention it (do not fabricate content).
- End the response with the sentence: “These are the complete actionable requirements, eligibility criteria, and specifications extracted from this RFP for proposal preparation.”

RFP_CONTENT_START
{truncated_rfp if truncated_rfp else "No RFP content is available."}
RFP_CONTENT_END
"""

    api_key = (app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or os.getenv('OPENROUTER_API_KEY') or os.getenv('TOGETHER_API_KEY') or os.getenv('FIREWORKS_API_KEY') or '').strip()
    if not api_key:
        return jsonify({'error': 'api_unavailable', 'message': 'OpenAI API key is not configured on the server.'}), 500

    base_url = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or os.getenv('OPENROUTER_BASE_URL') or 'https://openrouter.ai/api/v1').strip()
    model = (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'meta-llama/llama-4-scout-17b-16e-instruct').strip()
    endpoint = base_url.rstrip('/')
    if not endpoint.endswith('/chat/completions'):
        endpoint = f"{endpoint}/chat/completions"

    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json',
    }
    body = {
        'model': model,
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt}
        ],
        'temperature': 0.2,
    }

    try:
        response = requests.post(endpoint, headers=headers, json=body, timeout=90)
    except requests.RequestException as err:
        app.logger.exception("OpenAI section outline network failure: %s", err)
        return jsonify({'error': 'api_unavailable', 'message': 'Could not reach the AI service. Please try again shortly.'}), 502

    try:
        data = response.json()
    except ValueError:
        app.logger.error("OpenAI section outline returned non-JSON response (status %s): %s", response.status_code, response.text[:400])
        return jsonify({'error': 'api_unavailable', 'message': 'Received an invalid response from the AI service.'}), 502

    if response.status_code >= 400:
        err_message = data.get('error') if isinstance(data, dict) else None
        detail = ''
        if isinstance(err_message, dict):
            detail = err_message.get('message') or err_message.get('code') or ''
        elif isinstance(err_message, str):
            detail = err_message
        if not detail:
            detail = f"Upstream error {response.status_code}"
        app.logger.error("OpenAI section outline error %s: %s", response.status_code, detail)
        return jsonify({'error': 'api_unavailable', 'message': f'AI service returned an error: {detail}'}), 502

    choices = data.get('choices') or []
    if not choices:
        return jsonify({'error': 'api_unavailable', 'message': 'No response returned by the AI service.'}), 502

    content = choices[0].get('message', {}).get('content', '')
    if not content:
        return jsonify({'error': 'api_unavailable', 'message': 'The AI service returned an empty response.'}), 502

    return jsonify({
        'output': content.strip(),
        'raw': content.strip(),
        'model': data.get('model', model),
        'usage': data.get('usage', {}),
    })

@app.route('/api/analyze-requirements-attachments', methods=['POST'])
@login_required
def analyze_requirements_attachments():
    """Extract requirements and required attachments from RFP document."""
    payload = request.get_json(silent=True) or {}
    bid_id = payload.get('bid_id')
    
    if not bid_id:
        return jsonify({'error': 'missing_params', 'message': 'Bid ID is required.'}), 400

    # Fetch RFP content
    rfp_content = _fetch_rfp_content_for_bid(bid_id, page_limit=80, char_limit=15000)
    
    if not rfp_content:
        return jsonify({'error': 'no_rfp', 'message': 'No RFP content found for this bid. Please upload an RFP file first.'}), 404

    truncated_rfp = rfp_content[:15000] if rfp_content else ''

    system_prompt = (
        "You are a professional RFP analysis assistant. "
        "Your only task is to analyze the provided RFP / IFB document and extract **only the mandatory attachments, forms, and documents** that bidders are required to submit with their bid or proposal. "
        "Do NOT include scope of work, eligibility, or background text. "
        "Do NOT infer or assume anything — extract only what is explicitly listed in the RFP as required attachments, forms, or bid documents."
    )

    user_prompt = f"""
Analyze the following RFP content and extract **only the mandatory attachments, forms, and documents** that bidders are required to submit.

### 🔹 OUTPUT FORMAT (strictly follow this structure)

📑 **Mandatory Attachments & Documents to Include with the Bid**

List each required document or form in clean bullet points.  
If attachment numbers or names are mentioned, list them exactly as written in the RFP.

Example output format:

- Attachment 1 – Contract Award Form  
- Attachment 2 – Acknowledgment Form  
- Attachment 3 – Bidder's Certifications  
- Attachment 4 – Conflict of Interest Questionnaire  
- Attachment 5 – Financial Interests and Potential Conflicts Form  
- Attachment 6 – References  
- Attachment 7 – Insurance Requirements  
- Attachment 8 – Small Business Development (SBD) Forms  
- Attachment 9 – Contractor Certification Sheet  
- Bid Bond (if required)  
- Payment Bond (if contract value > $25,000)  
- Performance Bond (if contract value > $100,000)  
- Company Profile / Cover Letter  
- Safety Record / EMR Report  
- Proof of Experience and References  
- Licenses / Certifications  
- Financial Stability Statement  
- Bid Pricing Sheet / Cost Proposal  
- Small Business Participation Documentation  
- Any other forms explicitly listed in the RFP

If no attachments are mentioned in the document, write:  
**"No specific attachments or mandatory bid documents listed in this RFP."**

### 🔹 RULES FOR EXTRACTION

- Include **only** attachment titles, form names, or document names mentioned in the RFP.  
- Remove all explanation text or surrounding sentences.  
- One line per attachment or required document.  
- Use plain, clean formatting — perfect for UI display next to an "Attach File" button.  
- Do not generate or assume missing data.

### 🔹 GOAL

Output must contain **only the explicit list of required attachments and documents** found in the RFP.  
No commentary. No assumptions. No descriptions.

End your response with:  
**"These are the mandatory attachments and documents explicitly required for bid submission in this RFP."**

RFP_CONTENT_START
{truncated_rfp if truncated_rfp else "No RFP content is available."}
RFP_CONTENT_END
"""

    api_key = (app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or os.getenv('OPENROUTER_API_KEY') or os.getenv('TOGETHER_API_KEY') or os.getenv('FIREWORKS_API_KEY') or '').strip()
    if not api_key:
        return jsonify({'error': 'api_unavailable', 'message': 'OpenAI API key is not configured on the server.'}), 500

    base_url = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or os.getenv('OPENROUTER_BASE_URL') or 'https://openrouter.ai/api/v1').strip()
    model = (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'meta-llama/llama-4-scout-17b-16e-instruct').strip()
    endpoint = base_url.rstrip('/')
    if not endpoint.endswith('/chat/completions'):
        endpoint = f"{endpoint}/chat/completions"

    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json',
    }
    body = {
        'model': model,
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt}
        ],
        'temperature': 0.2,
    }

    try:
        response = requests.post(endpoint, headers=headers, json=body, timeout=90)
    except requests.RequestException as err:
        app.logger.exception("OpenAI requirements analysis network failure: %s", err)
        return jsonify({'error': 'api_unavailable', 'message': 'Could not reach the AI service. Please try again shortly.'}), 502

    try:
        data = response.json()
    except ValueError:
        app.logger.error("OpenAI requirements analysis returned non-JSON response (status %s): %s", response.status_code, response.text[:400])
        return jsonify({'error': 'api_unavailable', 'message': 'Received an invalid response from the AI service.'}), 502

    if response.status_code >= 400:
        err_message = data.get('error') if isinstance(data, dict) else None
        detail = ''
        if isinstance(err_message, dict):
            detail = err_message.get('message') or err_message.get('code') or ''
        elif isinstance(err_message, str):
            detail = err_message
        if not detail:
            detail = f"Upstream error {response.status_code}"
        app.logger.error("OpenAI requirements analysis error %s: %s", response.status_code, detail)
        return jsonify({'error': 'api_unavailable', 'message': f'AI service returned an error: {detail}'}), 502

    choices = data.get('choices') or []
    if not choices:
        return jsonify({'error': 'api_unavailable', 'message': 'No response returned by the AI service.'}), 502

    content = choices[0].get('message', {}).get('content', '')
    if not content:
        return jsonify({'error': 'api_unavailable', 'message': 'The AI service returned an empty response.'}), 502

    return jsonify({
        'output': content.strip(),
        'requirements': content.strip(),
        'raw': content.strip(),
        'model': data.get('model', model),
        'usage': data.get('usage', {}),
    })

@app.route('/api/generate-section-content', methods=['POST'])
@login_required
def generate_section_content():
    """Generate content for a specific proposal section, especially Letter of Transmittal."""
    payload = request.get_json(silent=True) or {}
    section_id = (payload.get('section_id') or '').strip()
    section_title = (payload.get('section_title') or '').strip()
    bid_id = payload.get('bid_id')
    company = (payload.get('company') or '').strip()
    contact_email = (payload.get('contact_email') or '').strip()
    override_api_key = (payload.get('api_key') or '').strip()
    
    if not section_id:
        return jsonify({'error': 'missing_params', 'message': 'Section ID is required.'}), 400
    
    # Get RFP content from database if bid_id is provided
    rfp_content = ''
    if bid_id:
        # Use higher limits for sections that need thorough analysis
        rfp_content = _fetch_rfp_content_for_bid(bid_id, page_limit=100, char_limit=25000)
    
    # Get bid information for context
    bid_info = {}
    if bid_id:
        try:
            cur = mysql.connection.cursor(DictCursor)
            cur.execute("SELECT * FROM go_bids WHERE g_id=%s", (bid_id,))
            bid_info = cur.fetchone() or {}
            cur.close()
        except Exception as e:
            app.logger.warning(f"Error loading bid info: {e}")
    
    # Extract relevant bid information
    project_name = bid_info.get('b_name') or bid_info.get('project_name') or ''
    client_name = bid_info.get('agency') or bid_info.get('customer') or bid_info.get('client') or ''
    due_date = bid_info.get('due_date') or ''
    
    # Format due date
    if due_date:
        try:
            from datetime import datetime, date
            if isinstance(due_date, (datetime, date)):
                due_date_str = due_date.strftime('%B %d, %Y')
            else:
                parsed = datetime.strptime(str(due_date), '%Y-%m-%d')
                due_date_str = parsed.strftime('%B %d, %Y')
        except Exception:
            due_date_str = str(due_date)
    else:
        due_date_str = ''
    
    # Check if this is Letter of Transmittal section
    is_letter_of_transmittal = (
        section_id == 'section-cover' or 
        'letter of transmittal' in (section_title or '').lower() or
        'transmittal' in (section_title or '').lower()
    )
    
    # Check if this is Executive Summary section
    is_executive_summary = (
        section_id == 'section-executive-summary' or
        'executive summary' in (section_title or '').lower()
    )
    
    # Check if this is Requirements Compliance section
    is_requirements_compliance = (
        section_id == 'section-compliance' or
        'requirements compliance' in (section_title or '').lower() or
        'compliance' in (section_title or '').lower() and 'requirements' in (section_title or '').lower()
    )
    
    # Check if this is Understanding of Scope & Objectives section
    is_scope_objectives = (
        section_id == 'section-scope-objectives' or
        'understanding of scope' in (section_title or '').lower() or
        'scope & objectives' in (section_title or '').lower() or
        ('scope' in (section_title or '').lower() and 'objectives' in (section_title or '').lower())
    )
    
    # Check if this is Deviations, Assumptions, and Dependencies section
    is_deviations_assumptions = (
        section_id == 'section-deviations' or
        'deviations' in (section_title or '').lower() or
        ('deviations' in (section_title or '').lower() and 'assumptions' in (section_title or '').lower() and 'dependencies' in (section_title or '').lower())
    )
    
    # Check if this is Proposed Technical Solution section
    is_proposed_technical_solution = (
        section_id == 'section-technical' or
        'proposed technical solution' in (section_title or '').lower() or
        ('technical solution' in (section_title or '').lower() and 'proposed' in (section_title or '').lower())
    )
    
    # Check if this is Implementation & Work Plan section
    is_implementation_work_plan = (
        section_id == 'section-implementation' or
        'implementation & work plan' in (section_title or '').lower() or
        'implementation and work plan' in (section_title or '').lower() or
        ('implementation' in (section_title or '').lower() and 'work plan' in (section_title or '').lower())
    )
    
    # Check if this is Solution Deliverables section
    is_solution_deliverables = (
        section_id == 'section-deliverables' or
        'solution deliverables' in (section_title or '').lower() or
        ('solution' in (section_title or '').lower() and 'deliverables' in (section_title or '').lower())
    )
    
    # Check if this is Project Management Plan section
    is_project_management_plan = (
        section_id == 'section-management' or
        'project management plan' in (section_title or '').lower() or
        ('project management' in (section_title or '').lower() and 'plan' in (section_title or '').lower())
    )
    
    # Check if this is Pricing Proposals section
    is_pricing_proposals = (
        section_id == 'section-pricing' or
        'pricing proposals' in (section_title or '').lower() or
        ('pricing' in (section_title or '').lower() and 'proposals' in (section_title or '').lower())
    )
    
    # Check if this is Contractual Terms section
    is_contractual_terms = (
        section_id == 'section-contractual' or
        'contractual terms' in (section_title or '').lower() or
        ('contractual' in (section_title or '').lower() and 'terms' in (section_title or '').lower())
    )
    
    # Check if this is Corporate Qualifications section
    is_corporate_qualifications = (
        section_id == 'section-corporate-qualifications' or
        section_id == 'section-corporate' or
        'corporate qualifications' in (section_title or '').lower() or
        ('corporate' in (section_title or '').lower() and 'qualifications' in (section_title or '').lower())
    )
    
    # Check if this is Project Schedule & Milestones section
    is_project_schedule_milestones = (
        section_id == 'section-schedule-milestones' or
        section_id == 'section-schedule' or
        ('project schedule' in (section_title or '').lower() and 'milestones' in (section_title or '').lower()) or
        ('schedule' in (section_title or '').lower() and 'milestone' in (section_title or '').lower())
    )
    
    # Check if this is Project Team section
    is_project_team = (
        section_id == 'section-project-team' or
        section_id == 'section-team' or
        'project team' in (section_title or '').lower() or
        ('team' in (section_title or '').lower() and 'project' in (section_title or '').lower())
    )
    
    # Check if this is Addenda section
    is_addenda = (
        section_id == 'section-amendments' or
        'addenda' in (section_title or '').lower() or
        'amendments' in (section_title or '').lower()
    )
    
    # Check if this is Appendix section
    is_appendix = (
        section_id == 'section-appendix' or
        'appendix' in (section_title or '').lower()
    )
    
    if is_letter_of_transmittal:
        # Use the specific prompt for Letter of Transmittal
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to draft a formal, professional Letter of Transmittal that meets all standard business letter requirements. "
            "The letter must be written from our company to the client, be technical in nature, and include all necessary points."
        )
        
        user_prompt = f"""Draft a formal, one-page Letter of Transmittal to be signed by an authorized executive. The letter must be in a standard business-letter format with clear paragraphs. It must:

1. State our formal intent to bid for this specific RFP.

2. Acknowledge receipt and review of all addenda.

3. Provide a concise statement of compliance with all mandatory requirements.

4. Briefly highlight our single most compelling technical differentiator.

5. Conclude by affirming our enthusiasm for the partnership.

This must be from our company to the client. It should be technical and should include all necessary points as mentioned in the RFP/RFI/RFQ/Bid Document.

At the very top of the letter, include exactly one header line with the client name, city/state, and the date (e.g., "The City of Oxford, Ohio — {due_date_str or 'October 17, 2025'}"). Do NOT include any company information before this line.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (for context):
{rfp_content[:15000] if rfp_content else 'No RFP content available. Generate a professional Letter of Transmittal template.'}

Please generate a complete, professional Letter of Transmittal that addresses all six requirements listed above. The letter should be technical, comprehensive, and written from our company to the client. Format it as a standard business letter with proper paragraphs and not write sub heading . Return the content as a single formatted text that can be split into paragraphs."""
    elif is_executive_summary:
        # Use the specific prompt for Executive Summary
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to draft a persuasive Executive Summary of about 500 words that is standalone and written for a high-level, non-technical decision-maker. "
            "The summary must be concise, non-repetitive, and avoid restating the same idea in different words."
        )
        
        user_prompt = f"""Prompt:  With reference to the requirement mentioned in the RFP/RFI/RFQ/BID document, generate a persuasive Executive Summary of about 500 words (ideally between 450 and 550 words). This summary must be standalone and written for a high-level, non-technical/technical decision-maker. It should be from our company to the client. This should not have any subheadings and should be in paragraph only. Avoid any repetition, filler, or generic statements. Do not include prices, headers, or bullet lists. Vary the sentence openings and ensure each sentence adds new information and does not repeat prior sentences. Use the following structure (paragraph-only):

• In a single paragraph, summarize our understanding of the client's core problem, goals, and desired outcomes as stated in the bid.

• In two paragraphs, provide a high-level, benefit-focused overview of our technical and management solution. Avoid deep jargon.

• Conclude with one short paragraph that highlights 3–5 specific, unique advantages our solution offers, written in prose (not bullets).

Company Information:
- Company Name: {company or 'Our Company'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (for context):
{rfp_content[:15000] if rfp_content else 'No RFP content available. Generate a professional Executive Summary template.'}

Important: Write as continuous paragraphs only (no headings or bullet lists). Each paragraph must be concise and non-redundant. Return only the paragraphs as plain text, separated by blank lines."""
    elif is_requirements_compliance:
        # Use the specific prompt for Requirements Compliance
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to draft a formal Statement of Compliance that thoroughly addresses all requirements and key activities from the RFP document."
        )
        
        user_prompt = f"""As per the requirement mentioned in the document, draft a formal 'Statement of Compliance.' State that we have read, understood, and will comply with all mandatory requirements, terms, and conditions set forth in the bid document. Include all key points in detail and in a technical manner. Do not skip any important points to be discussed.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify ALL requirements, key activities, deliverables, attachments, forms, technical specifications, and compliance obligations):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Statement of Compliance template.'}

Please generate the Statement of Compliance as exactly TWO paragraphs only, no headings, lists, bullets, or numbering:
- Paragraph 1: Provide the formal compliance assertion that we have read, understood, and will comply with all mandatory requirements, terms, and conditions of the bid.
- Paragraph 2: Note that any exceptions will be listed in the 'Deviations' section and summarize, in technical language, our comprehensive adherence to the RFP’s requirements, deliverables, forms, technical specifications, certifications, and compliance obligations.

Strictly limit the output to two paragraphs with complete sentences and professional tone."""
    elif is_scope_objectives:
        # Use the specific prompt for Understanding of Scope & Objectives
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to draft a comprehensive 'Understanding of Scope & Objectives' section that demonstrates deep subject-matter expertise by adding value and insight, not just repeating the RFP's text."
        )
        
        user_prompt = f"""As per the requirement mentioned in the document, draft a section titled 'Understanding of Scope & Objectives.' Using narrative paragraphs, analyse and re-state the client's core challenges, operational problems, and strategic goals. This section must demonstrate our deep subject-matter expertise by adding value and insight, not just repeating the RFP's text. Show that we understand the 'why' behind the 'what.' Include all key points in technical manner and in detail. Do not skip any important points to be discussed. Use pointers and table if necessary. Make it completely informative and elaborative.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to understand scope, objectives, challenges, phases, key activities, technical requirements, and all important points):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Understanding of Scope & Objectives template.'}

Please generate a comprehensive 'Understanding of Scope & Objectives' section that:
- Uses narrative paragraphs to demonstrate deep understanding, supplemented with pointers and tables where necessary
- Analyzes and re-states the client's core challenges, operational problems, and strategic goals in technical detail
- Adds value and insight beyond just repeating the RFP text - demonstrates subject-matter expertise
- Shows understanding of the 'why' behind the 'what' - explains the reasoning, context, strategic importance, and technical implications
- Includes ALL key points in technical manner and in detail - do not skip any important points to be discussed
- Uses pointers (bullet points) and tables where appropriate to organize information clearly and make it completely informative
- Makes it completely informative and elaborative - comprehensive coverage of all aspects with technical precision

Return the content as well-structured narrative paragraphs with appropriate use of pointers and tables to organize key activities, phases, technical requirements, and all important points. The content must be completely informative and elaborative, technically detailed, and comprehensive without omitting any important information."""
    elif is_deviations_assumptions:
        # Use the specific prompt for Deviations, Assumptions, and Dependencies
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a section that manages expectations by clearly documenting deviations, assumptions, and dependencies."
        )
        
        user_prompt = f"""As per mentioned in the document (if any), generate a section to manage expectations (if any). Use the following

• Use a numbered list to detail every item marked as 'Partial' or 'Exception'. Provide a clear justification for each.

• Use a numbered list to state all assumptions our technical solution and pricing are based on (e.g., 'We assume client will provide network access...').

• Use a numbered list to explicitly state all items, data, access, or resources we require from the client to ensure project success."

It should be completely technical and should include all key points. Do not skip any key points mentioned in RFP.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify ALL deviations, assumptions, dependencies, technical requirements, and key points):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Deviations, Assumptions, and Dependencies template.'}

Please generate a comprehensive, completely technical section that:
- Uses numbered lists for: Deviations and Exceptions; Technical & Operational Assumptions; Client-Furnished Dependencies
- Details every item marked as 'Partial' or 'Exception' in the compliance matrix (if applicable) with technical justifications
- States ALL assumptions our technical solution and pricing are based on
- Explicitly states ALL client-furnished items, data, access, resources, and dependencies required
- Includes ALL key points mentioned in the RFP - do not skip any important technical details, requirements, or specifications
- Is completely technical, specific, and actionable

Return the content with clearly separated parts and numbered lists as specified. The content must be completely technical, comprehensive, and include all key points from the RFP without omitting any important information."""
    elif is_proposed_technical_solution:
        # Use the specific prompt for Proposed Technical Solution
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive 'Proposed Technical Solution' section that addresses all technical requirements and work-streams from the RFP document."
        )
        
        # Include OCR text from section-specific image attachments (if any)
        section_images_text = _fetch_section_images_text(bid_id, 'section-technical', char_limit=6000) if bid_id else ''
        
        user_prompt = f"""As per the requirement in the document, generate the core 'Proposed Technical Solution' section. Structure these using headings and subheadings that directly correspond to the technical work-streams defined in the bid document.

• Before generating, ensure the user has uploaded/attached the RFP/SoW/addenda via the Attach File button. Incorporate all details/information from the attached documents and any provided data inputs.

• For each subsection, use technical paragraphs, bulleted lists for specifications, and tables for technical data.

• This section must be highly detailed, addressing how our solution fulfils each technical requirement. Include specific product models, software versions, and methodologies where relevant.

• Include all key activities/technical requirements. Do not skip any important points to be discussed. Use pointers and table if necessary.

Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify ALL technical work-streams, requirements, specifications, deliverables, phases, key activities, and technical details):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Proposed Technical Solution template.'}

Section Image Attachments (OCR text extracted from images uploaded for this section):
{section_images_text if section_images_text else 'No image attachments provided for this section.'}

Please generate a comprehensive, highly detailed, and completely technical 'Proposed Technical Solution' section that:
- Uses headings and subheadings that directly correspond to technical work-streams from the RFP
- For each subsection, uses technical paragraphs, bulleted lists (pointers) for specifications, and tables for technical data
- Is highly detailed and addresses how our solution fulfils each technical requirement with technical precision
- Includes specific product models, software versions, hardware specifications, methodologies, standards, and protocols where relevant
- Includes ALL key activities/technical requirements, specifications, and important points - does not skip any important points
- Uses pointers (bullet points) and tables where required to organize and present technical information clearly
- Is completely technical and comprehensive
- Is aligned with the requirements in RFP - addresses every technical requirement, specification, and work-stream
- Makes it in very detail by including all necessary points
- Includes all necessary technical details: configurations, specifications, methodologies, protocols, standards, tools, technologies, and implementation approaches

Return the content as well-structured, highly detailed, and completely technical content with appropriate headings and subheadings, paragraphs, bulleted lists (pointers), and tables. The response must be aligned with all requirements in the RFP and include all necessary points without omitting any important information."""
    elif is_implementation_work_plan:
        # Use the specific prompt for Implementation & Work Plan
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive 'Implementation & Work Plan' section that details the methodology, project phases, and key activities for successful project delivery."
        )
        
        user_prompt = f"""Generate a comprehensive 'Implementation & Work Plan' section with two main subsections:

1. Methodology

As per the requirement in the document, draft methodology section based on the project scope. In narrative paragraphs, describe the specific project management and technical implementation methodology we will use. Include all key activities. Do not skip any important points to be discussed. Use pointers and table if necessary. Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP. Use headings, subheadings, pointers and tables where required.

2. Project Phases & Activities

As per the requirement in the document, "daft a 'Project Phases & Activities' subsection based on the project scope. Use subheadings for each distinct project phase. Use weeks for each phase in the form Week1-Week2 etc. These weeks should be practically feasible as per the requirement in the project document. Includer all key activities in each phase. Do not skip any important points to be discussed. Use pointers and tables where necessary. Make it in very detail by including all necessary points. Also, make sure this is completely technical and feasible. Include pointers and table where required. The response should be aligned with the requirement in the project.

Include all key activities in each phase. Do not skip any important points to be discussed. Use pointers and tables where necessary. Make it in very detail by including all necessary points. Also, make sure this is completely technical and feasible. Include pointers and table where required. The response should be aligned with the requirement in RFP.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify ALL project phases, activities, tasks, objectives, methodology requirements, technical implementation details, and all important points):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Implementation & Work Plan template.'}

Please generate a comprehensive 'Implementation & Work Plan' section that includes:

1. Methodology (use headings/subheadings as appropriate):
   - Narrative paragraphs describing the specific project management and technical implementation methodology in very detail
   - Justification for why this methodology is the best fit for this specific project and client
   - ALL key activities related to methodology included - no important points skipped
   - Use of pointers (bullet points) and tables where required to organize and present technical information clearly
   - Very detailed coverage including all necessary points
   - Completely technical content
   - Response aligned with the requirements in RFP

2. Project Phases & Activities (use subheadings):
   - Subheadings for each distinct project phase (e.g., 'Phase 1: Discovery & Design,' 'Phase 2: Build & Test,' 'Phase 3: Deployment & Go-Live')
   - For each phase, include a practical week range label (e.g., 'Week 1–Week 2', 'Week 3–Week 6') aligned with the bid timeline
   - Under each phase, bulleted lists detailing key activities, tasks, and objectives in very detail
   - ALL key activities in each phase included - no important points skipped
   - Use of pointers (bullet points) and tables where required to organize and present technical information clearly
   - Very detailed coverage including all necessary points
   - Completely technical content
   - Response aligned with the requirements in RFP

Return the content as well-structured, highly detailed, and completely technical content with appropriate headings and subheadings, narrative paragraphs, bulleted lists (pointers), and tables. Both the Methodology and Project Phases & Activities subsections must be in very detail, completely technical, include all key activities, use pointers and tables where required, and be aligned with all requirements in the RFP."""
    elif is_solution_deliverables:
        # Use the specific prompt for Solution Deliverables
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive 'Solution Deliverables' section that lists all tangible items the client will receive in a structured table format."
        )
        
        user_prompt = f"""As per the requirement in the project information document, generate a 'Solution Deliverables' section based on the scope of the work. Create a table that lists all tangible items the client will receive. The table columns must be:

1. Deliverable Title: must include what will be delivered.

2. Description: (A brief explanation).

3. Format: (e.g., 'PDF Document,' 'Online Dashboard').

4. Delivery Milestone: (When it will be delivered, e.g., 'End of Phase 1').

Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify ALL deliverables, their descriptions, formats, delivery milestones, technical specifications, and all important points):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Solution Deliverables template.'}

Please generate a comprehensive, highly detailed, and completely technical 'Solution Deliverables' section that:
- Creates a table with exactly 4 columns: Deliverable Title, Description, Format, Delivery Milestone
- Lists ALL tangible items the client will receive with technical precision
- Includes ALL deliverables from the RFP - does not skip any items or important points
- Provides very specific and detailed technical descriptions for each deliverable
- Matches delivery milestones to project phases from the RFP
- Uses appropriate technical formats for each deliverable
- Is well-structured, comprehensive, and includes all technical details
- Is aligned with the requirements in RFP - addresses every deliverable requirement, specification, and item
- Includes pointers (bullet points) and tables where required to organize and present technical information clearly
- Makes it in very detail by including all necessary points
- Is completely technical

Return the content as a well-structured, highly detailed, and completely technical table with all deliverables listed. The table format must be clear, comprehensive, include all required columns, and be aligned with all requirements in the RFP. Include pointers and tables where required to make it completely informative and technical."""
    elif is_project_management_plan:
        # Use the specific prompt for Project Management Plan
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive 'Project Management Plan' section that includes project governance, communication plan, risk management, and quality assurance."
        )
        
        user_prompt = f"""As per the requirement in the bid document, generate a comprehensive, highly detailed, and completely technical 'Project Management Plan' with four main subsections:

1. Project Governance

Prompt: As per the requirement in the project document, draft a 'Project Governance' subsection based on the scope of the project. In paragraphs, describe the management structure involved in the project. Include all key activities in each phase. Do not skip any important points to be discussed. Use pointers and table if necessary. Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

2. Communication Plan

Prompt: As per the requirement in the project document, generate a 'Communication Plan as per the scope of the work. Create a table that defines the communication protocols. Includer all key activities. Do not skip any important points to be discussed. Use pointers and table if necessary. Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

3. Risk Management Plan

Prompt: As per the requirement in the project document, generate a 'Risk Management Plan based on the scope of the work. Create a table. Do not skip any important points to be discussed. Use pointers and table if necessary. Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

4. Quality Assurance Plan

Prompt: A per the requirement in the bid document, draft a 'Quality Assurance Plan based on the scope of work. Includer all key activities. Do not skip any important points to be discussed. Use pointers and table if necessary. Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

Include all key activities. Do not skip any important points to be discussed. Use pointers and tables if necessary.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify all project management requirements, communication protocols, risks, quality standards, and key activities):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Project Management Plan template.'}

Please generate a comprehensive, highly detailed, and completely technical 'Project Management Plan' section. Use appropriate H3 subheadings for each subsection and ensure alignment with the RFP. Include all key activities and do not skip any important points. Use pointers and tables where required."""
    elif is_pricing_proposals:
        # Use the specific prompt for Pricing Proposals
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive, highly detailed, and completely technical 'Pricing Proposals' section aligned with RFP requirements."
        )
        # Include text extracted from Pricing section image attachments only
        attachments_text = _fetch_section_images_text(bid_id, 'section-pricing', char_limit=8000) if bid_id else ''
        
        user_prompt = f"""Prompt: Generate space for this and give placeholders for each of this sections.

o Pricing Summary

o Detailed Cost Breakdown

o Payment Schedule

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content:
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Pricing Proposals template.'}"""
    elif is_contractual_terms:
        # Use the specific prompt for Contractual Terms
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive, highly detailed, and completely technical 'Acceptance of Terms and Conditions' section aligned with RFP requirements."
        )
        
        user_prompt = f"""PROMPT FOR  Contractual Terms:

o Acceptance of Terms and Conditions

Prompt: "Analyse the contract and Terms & Conditions (T&Cs) from the bid document.

If we accept all terms: Draft a single paragraph stating that 'We have reviewed all terms and conditions outlined in the bid document and confirm our full acceptance and compliance without exception.' This should be technical in nature based on the scope of the work.

If we have exceptions: Draft a paragraph stating we accept with exceptions, and then create a table with columns: Clause Number, Clause Title, Proposed Redline / Exception, and Business Justification." Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content:
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Contractual Terms template.'}"""
    elif is_corporate_qualifications:
        # Use the specific prompt for Corporate Qualifications
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive 'Corporate Qualifications' section that includes Company Overview, Relevant Past Performance, and Client References aligned with the RFP."
        )
        
        user_prompt = f"""PROMPT FOR Corporate Qualifications

o Company Overview

Prompt: "Draft a 'Company Overview.' In 2-3 concise paragraphs, describe our company's history, size, locations, and core competencies that are directly relevant to this project." Use technical statements.

o Relevant Past Performance

Prompt: As per the requirement in the bid document, generate a 'Relevant Past Performance' section. Under each client, create a table with rows for:

• Project Title:
• Period of Performance:
• Project Value:
• Challenge: (A brief description of the problem).
• Solution & Outcome: (A description of our work and the positive, quantifiable result)."

o Client References

Prompt: "Generate a 'Client References' section. Create a table providing 3-5 client references, corresponding to the projects listed in 'Past Performance.' Columns must be: Client Name, Contact Person, Title, Phone, and Email."

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content:
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Corporate Qualifications template.'}"""
    elif is_project_schedule_milestones:
        # Use the specific prompt for Project Schedule & Milestones
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive, highly detailed, and completely technical 'Project Schedule & Milestones' section that is complete, technical, and feasible practically, aligned with RFP requirements."
        )
        
        user_prompt = f"""As per the requirement in the bid document, generate a 'Project Schedule & Milestones' section. This should be complete technical and feasible practically. Use tables and pointers where necessary.

1. First, insert a placeholder for a visual Gantt chart (e.g., [Insert Gantt Chart Visual Here]).

2. Following the placeholder, create a table that lists the 'Critical Project Milestones.' Columns should be: Milestone, Description, and Target Completion Date (or 'Weeks from Kick-off').

Make it in very detail by including all necessary points. Also, make sure this is completely technical and feasible. Include pointers and table where required. The response should be aligned with the requirement in RFP.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify all project phases, milestones, timelines, delivery dates, and schedule requirements):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Project Schedule & Milestones template.'}
"""
    elif is_project_team:
        # Use the specific prompt for Project Team
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive, highly detailed, and completely technical 'Project Team' section that includes organizational chart, roles & responsibilities, and key personnel information, aligned with RFP requirements."
        )
        
        user_prompt = f"""PROMPT FOR Project Team

o Project Team Organizational Chart

Prompt: "Insert a placeholder for a 'Project Team Organizational Chart.' This visual diagram must show the proposed team structure, reporting lines, and key interfaces with the client's team." 

o Roles & Responsibilities

Prompt: "Generate a 'Roles & Responsibilities' section. Create a table with columns: Project Role, Proposed Name, and Key Responsibilities." Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

o Key Personnel Resumes

Prompt: insert a placeholder for key resumes or small bios

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content:
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Project Team template.'}"""
    elif is_addenda:
        # Use the specific prompt for Addenda
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate an acknowledgment statement for the Addenda section that acknowledges all addenda received with their names and dates."
        )
        
        user_prompt = f"""Generate an acknowledgment statement for the Addenda section. Analyze the RFP thoroughly to identify all addenda that have been issued.

RFP Content (analyze thoroughly to identify all addenda, amendments, and modification notices):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Addenda acknowledgment statement.'}

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

Prompt: "Write an acknowledgement statement in 5-6 lines acknowledging all addenda received. It should have addenda name and date."

IMPORTANT REQUIREMENTS:
- Write an acknowledgment statement in 5-6 lines (paragraphs)
- Acknowledge all addenda that have been issued for this solicitation
- Include the addenda name/number for each addendum
- Include the release date for each addendum
- If no addenda were issued, include a statement confirming that no addenda were received
- Use professional and formal language appropriate for government proposals
- Ensure the statement is clear and comprehensive

Return the content as a well-structured acknowledgment statement in 5-6 lines, with each addendum clearly identified by name/number and date."""
    elif is_appendix:
        # Use the specific prompt for Appendix
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive 'List of Appendices' section that identifies all supporting documents included at the end of the proposal."
        )
        
        user_prompt = f"""Generate a 'List of Appendices' section. Analyze the RFP thoroughly to identify all supporting documents, references, and additional materials that should be included in the appendices.

RFP Content (analyze thoroughly to identify all supporting documents, attachments, references, and additional materials referenced in the proposal):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional List of Appendices template.'}

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

Prompt: "Generate a 'List of Appendices.' Use a lettered, nested list to identify all supporting documents included at the end of the proposal. This should include (but is not limited to):

- Appendix A: 
- Appendix B: 
- Appendix C: 
- Appendix D: 
- Appendix E:"

IMPORTANT REQUIREMENTS:
- Generate a 'List of Appendices' section
- Use a lettered, nested list format (Appendix A, Appendix B, Appendix C, etc.)
- Identify all supporting documents that should be included at the end of the proposal
- Include documents such as (but not limited to):
  - Technical specifications and drawings
  - Certifications and accreditations
  - Past performance documentation
  - Key personnel resumes
  - Organizational charts
  - Compliance matrices
  - Pricing details and breakdowns
  - References and testimonials
  - Any other supporting materials referenced in the proposal
- Analyze the RFP to identify which specific appendices are required or referenced
- Provide clear, descriptive titles for each appendix
- Ensure the list is comprehensive and includes all relevant supporting documents
- Use professional formatting with proper lettering (A, B, C, D, E, etc.)

Return the content as a well-structured 'List of Appendices' with a lettered, nested list format. Each appendix should be clearly identified with a letter (A, B, C, D, E, etc.) and a descriptive title indicating what document or material it contains."""
    else:
        # Generic section generation
        system_prompt = (
            "You are an expert proposal writer helping to draft specific sections of a government proposal. "
            "Analyze the RFP content and generate comprehensive, compliant content for the requested section."
        )
        
        user_prompt = f"""Generate content for the section: "{section_title or section_id}".

Requirements:
1. Analyze the provided RFP content to understand requirements relevant to this section
2. Generate professional, compliant content that addresses the section topic
3. Ensure the content is specific, actionable, and aligned with government proposal standards
4. If RFP content is limited, provide a well-structured template that can be customized

RFP Content:
{rfp_content[:15000] if rfp_content else "No RFP content available. Generate a professional template for this section."}

Company: {company or 'Our Company'}
Project: {project_name or 'This Project'}

Please provide the content for "{section_title or section_id}" in a clear, well-formatted manner."""
    
    # Enrich all prompts with company profile context from database
    try:
        cur = mysql.connection.cursor(DictCursor)
        # Helpers for flexible table/column handling
        def _table_exists(table_name: str) -> bool:
            try:
                cur.execute("SHOW TABLES LIKE %s", (table_name,))
                return cur.fetchone() is not None
            except Exception:
                return False
        def _describe_columns(table_name: str) -> list[str]:
            try:
                cur.execute(f"DESCRIBE `{table_name}`")
                rows = cur.fetchall() or []
                cols = []
                for r in rows:
                    # MySQL DictCursor returns {'Field': 'col', 'Type': '...'}
                    cols.append(r.get('Field') or list(r.values())[0])
                return [c for c in cols if isinstance(c, str)]
            except Exception:
                return []
        def _company_where_clause(existing_cols: list[str], company_value: str):
            candidate_cols = [
                'company_name', 'company', 'comp_name', 'organization',
                'org_name', 'vendor', 'contractor', 'client_company'
            ]
            cols = [c for c in candidate_cols if c in (existing_cols or [])]
            if company_value and cols:
                clause = ' OR '.join([f"LOWER(COALESCE({c},'')) = LOWER(%s)"] * len(cols))
                params = tuple([company_value] * len(cols))
                return f"WHERE {clause}", params
            return "", tuple()
        # Latest company profile details
        cur.execute("SELECT * FROM company_details ORDER BY id DESC LIMIT 1")
        _details = cur.fetchone() or {}
        # Company preferences (e.g., registered states)
        cur.execute("SELECT * FROM company_preferences ORDER BY id DESC LIMIT 1")
        _prefs = cur.fetchone() or {}
        # Capabilities (latest uploads first)
        try:
            cur.execute("SELECT * FROM company_capabilities ORDER BY uploaded_at DESC LIMIT 5")
        except Exception:
            cur.execute("SELECT * FROM company_capabilities ORDER BY id DESC LIMIT 5")
        _caps = cur.fetchall() or []
        # Primary context: use projects_bids in place of company_performance, filtered by company
        _perf = []
        selected_company = (company or (_details.get('name') if _details else '') or '').strip()
        try:
            if _table_exists('projects_bids'):
                cols_b = _describe_columns('projects_bids')
                where_sql_b, where_params_b = _company_where_clause(cols_b, selected_company)
                sql_b = f"""
                    SELECT * FROM projects_bids
                    {where_sql_b}
                    ORDER BY COALESCE(due_date, created_at, updated_at) DESC, id DESC
                    LIMIT 10
                """
                cur.execute(sql_b, where_params_b)
                _perf = cur.fetchall() or []
            if not _perf and _table_exists('go_bids') and selected_company:
                cur.execute("""
                    SELECT g_id, b_name, due_date, state, company, summary
                    FROM go_bids
                    WHERE LOWER(COALESCE(company,'')) = LOWER(%s)
                    ORDER BY due_date DESC
                    LIMIT 10
                """, (selected_company,))
                _perf = cur.fetchall() or []
        except Exception:
            _perf = []
        # Optional: table past_performance (if present)
        _past_perf_extra = []
        try:
            pp_table = 'past_performance'
            if not _table_exists(pp_table) and _table_exists('past_performace'):
                pp_table = 'past_performace'  # common misspelling
            if _table_exists(pp_table):
                cols = _describe_columns(pp_table)
                where_sql, where_params = _company_where_clause(cols, selected_company)
                sql = f"SELECT * FROM `{pp_table}` {where_sql} ORDER BY COALESCE(year, 0) DESC, id DESC LIMIT 10"
                cur.execute(sql, where_params)
                _past_perf_extra = cur.fetchall() or []
        except Exception:
            _past_perf_extra = []
        # Optional: include company_performance records (global past performance)
        _company_perf = []
        try:
            if _table_exists('company_performance'):
                cur.execute("""
                    SELECT * 
                    FROM company_performance 
                    ORDER BY COALESCE(year, 0) DESC, id DESC 
                    LIMIT 10
                """)
                _company_perf = cur.fetchall() or []
        except Exception:
            _company_perf = []
        # Optional: personnel roster (prefer 'personnel' table; fallback to employees)
        _personnel = []
        try:
            if _table_exists('personnel'):
                cols = _describe_columns('personnel')
                where_sql, where_params = _company_where_clause(cols, selected_company)
                sql = f"SELECT * FROM personnel {where_sql} ORDER BY id DESC LIMIT 15"
                cur.execute(sql, where_params)
                _personnel = cur.fetchall() or []
            else:
                # Fallback to generic employees table if available
                if _table_exists('employees'):
                    cur.execute("""
                        SELECT name, department AS role, email 
                        FROM employees 
                        WHERE is_active = TRUE 
                        ORDER BY name ASC
                        LIMIT 15
                    """)
                    _personnel = cur.fetchall() or []
        except Exception:
            _personnel = []
        # Optional: active bids or projects_bids
        _projects_bids = []
        try:
            # projects_bids
            if _table_exists('projects_bids'):
                cols = _describe_columns('projects_bids')
                where_sql, where_params = _company_where_clause(cols, selected_company)
                sql = f"""
                    SELECT * FROM projects_bids
                    {where_sql}
                    ORDER BY COALESCE(due_date, created_at, updated_at) DESC, id DESC
                    LIMIT 10
                """
                cur.execute(sql, where_params)
                _projects_bids = cur.fetchall() or []
            # projects_contracts
            _projects_contracts = []
            if _table_exists('projects_contracts'):
                cols_pc = _describe_columns('projects_contracts')
                where_sql_pc, where_params_pc = _company_where_clause(cols_pc, selected_company)
                sql_pc = f"""
                    SELECT * FROM projects_contracts
                    {where_sql_pc}
                    ORDER BY COALESCE(contract_date, award_date, created_at) DESC, id DESC
                    LIMIT 10
                """
                cur.execute(sql_pc, where_params_pc)
                _projects_contracts = cur.fetchall() or []
            # Fallbacks when neither table is present or results empty
            if not _projects_bids and selected_company:
                if _table_exists('go_bids'):
                    cur.execute("""
                        SELECT g_id, b_name, due_date, state, company, summary
                        FROM go_bids
                        WHERE LOWER(COALESCE(company,'')) = LOWER(%s)
                        ORDER BY due_date DESC
                        LIMIT 10
                    """, (selected_company,))
                    _projects_bids = cur.fetchall() or []
        except Exception:
            _projects_bids, _projects_contracts = [], []
        cur.close()
    except Exception:
        _details, _prefs, _caps, _perf, _past_perf_extra, _company_perf, _personnel, _projects_bids = {}, {}, [], [], [], [], [], []
    
    company_context_lines = []
    if _details:
        company_context_lines.append(f"- Legal Name: {(_details.get('name') or company or 'Our Company')}")
        if _details.get('website'):
            company_context_lines.append(f"- Website: {_details.get('website')}")
        if _details.get('email'):
            company_context_lines.append(f"- Email: {_details.get('email')}")
        if _details.get('phone'):
            company_context_lines.append(f"- Phone: {_details.get('phone')}")
        if _details.get('about'):
            company_context_lines.append(f"- About: {_details.get('about')[:400]}")
    if _prefs and (_prefs.get('registered_states') or '').strip():
        company_context_lines.append(f"- Registered/Eligible States: {_prefs.get('registered_states')}")
    if _caps:
        company_context_lines.append("- Capabilities:")
        for c in _caps[:5]:
            desc = (c.get('description') or '').strip()
            if desc:
                company_context_lines.append(f"  • {desc[:160]}")
    if _perf:
        company_context_lines.append("- Relevant Projects/Bids (from projects_bids):")
        for p in _perf[:7]:
            # Support both projects_bids and go_bids fields
            pname = (p.get('project_name') or p.get('b_name') or p.get('name') or '').strip() or 'Project/Bid'
            due = p.get('due_date') or p.get('created_at') or ''
            company_name = (p.get('company_name') or p.get('company') or '').strip()
            psum = (p.get('summary') or p.get('description') or '').strip()
            due_str = str(due)[:10] if due else ''
            snippet = psum[:220] + ('…' if psum and len(psum) > 220 else '')
            tail = []
            if due_str: tail.append(f"Due {due_str}")
            if company_name: tail.append(company_name)
            meta = f" – {' | '.join(tail)}" if tail else ''
            company_context_lines.append(f"  • {pname}{meta}{f' – {snippet}' if snippet else ''}")
    # Include extra past_performance table if available
    if _past_perf_extra:
        company_context_lines.append("- Additional Past Performance (from past_performance):")
        for p in _past_perf_extra[:5]:
            pname = (p.get('project_name') or p.get('title') or '').strip() or 'Project'
            client = (p.get('client') or p.get('owner') or '').strip()
            pyear = p.get('year')
            value = p.get('contract_value') or p.get('value')
            psum = (p.get('summary') or p.get('description') or '').strip()
            snippet = psum[:200] + ('…' if psum and len(psum) > 200 else '')
            line_bits = [pname]
            if pyear: line_bits.append(f"{pyear}")
            if client: line_bits.append(client)
            if value: line_bits.append(str(value))
            line = " – ".join([b for b in line_bits if b])
            company_context_lines.append(f"  • {line}{f' – {snippet}' if snippet else ''}")
    # Include company_performance table if available
    if '_company_perf' in locals() and _company_perf:
        company_context_lines.append("- Additional Past Performance (from company_performance):")
        for p in _company_perf[:5]:
            pname = (p.get('project_name') or '').strip() or 'Project'
            pyear = p.get('year')
            line = f"{pname}{f' – {pyear}' if pyear else ''}"
            company_context_lines.append(f"  • {line}")
    # Personnel roster
    if _personnel:
        company_context_lines.append("- Key Personnel:")
        for m in _personnel[:7]:
            name = (m.get('name') or m.get('full_name') or '').strip()
            role = (m.get('title') or m.get('role') or m.get('position') or m.get('department') or '').strip()
            certs = (m.get('certifications') or '').strip()
            tail = f" – {role}" if role else ""
            if certs:
                tail += f" ({certs[:80]}{'…' if len(certs) > 80 else ''})"
            if name:
                company_context_lines.append(f"  • {name}{tail}")
    # Projects / Bids snapshot (kept only when not already used as _perf)
    if _projects_bids and not _perf:
        company_context_lines.append("- Active/Recent Bids & Projects:")
        for b in _projects_bids[:7]:
            name = (b.get('b_name') or b.get('project_name') or b.get('name') or '').strip() or 'Bid/Project'
            due = b.get('due_date') or b.get('created_at')
            state = (b.get('state') or b.get('status') or '').strip()
            summary = (b.get('summary') or b.get('description') or '').strip()
            due_str = str(due)[:10] if due else ''
            snippet = summary[:120] + ('…' if summary and len(summary) > 120 else '')
            label = f"{name}{f' – Due {due_str}' if due_str else ''}{f' – {state}' if state else ''}"
            company_context_lines.append(f"  • {label}{f' – {snippet}' if snippet else ''}")
    if '_projects_contracts' in locals() and _projects_contracts:
        company_context_lines.append("- Recent Contracts:")
        for c in _projects_contracts[:7]:
            name = (c.get('contract_name') or c.get('project_name') or c.get('name') or '').strip() or 'Contract'
            date = c.get('contract_date') or c.get('award_date') or c.get('created_at')
            value = c.get('contract_value') or c.get('value') or c.get('amount')
            due_str = str(date)[:10] if date else ''
            val_str = f"${value}" if value not in (None, '') else ''
            company_context_lines.append(f"  • {name}{f' – {due_str}' if due_str else ''}{f' – {val_str}' if val_str else ''}")
    company_context_block = "\n".join(company_context_lines).strip()
    if company_context_block:
        user_prompt = f"{user_prompt}\n\nCompany Profile Context (from database):\n{company_context_block}"
    
    # ---- Hybrid RAG: Retrieve top-K chunks and summarize locally via Ollama; cache results ----
    try:
        # Build corpus from RFP + company profile
        rfp_chunks = _text_chunks(rfp_content or '', 900, 140)
        company_chunks = _text_chunks(company_context_block or '', 900, 140)
        all_docs = [c for c in (rfp_chunks + company_chunks) if c.strip()]
        index = None
        if all_docs:
            index, all_docs = _build_or_load_faiss(company, bid_id, all_docs)
        # Query based on section title + id as hint
        retrieval_query = f"{section_title or section_id or 'proposal section'} {company or ''}".strip()
        top_docs = _retrieve_top_k(retrieval_query, company, bid_id, all_docs, index=index, k=8) if all_docs else []
        retrieved_block = "\n\n---\n\n".join(top_docs) if top_docs else ''
        # Summarize locally with Ollama to keep cloud tokens low
        summary = ''
        if retrieved_block:
            messages_local = [
                {'role': 'system', 'content': 'You are a concise technical summarizer. Summarize only the salient facts for proposal writing.'},
                {'role': 'user', 'content': f"Summarize briefly (<= 250 words) the following context relevant to '{section_title or section_id}':\n\n{retrieved_block}"}
            ]
            summary = _ollama_chat(messages_local, temperature=0.1, timeout_sec=45) or ''
        if retrieved_block:
            user_prompt = f"{user_prompt}\n\nRetrieved Context (top-k):\n{retrieved_block[:6000]}"
        if summary:
            user_prompt = f"{user_prompt}\n\nContext Summary (local Ollama):\n{summary}"
        # Section-level response cache (avoid repeated LLM calls)
        _sec_cache_key = _cache_key_from_dict({
            'kind': 'section_gen',
            'section_id': section_id,
            'section_title': section_title,
            'bid_id': bid_id,
            'company': company,
            'model': (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or ''),
            'user_prompt': user_prompt[-4000:],  # suffix captures specifics
        })
        _cached = _llm_cache_get(_sec_cache_key)
        if isinstance(_cached, dict) and _cached.get('content') and _cached.get('raw'):
            return jsonify(_cached)
    except Exception:
        pass
    
    # If no AI providers are configured, return a reasonable offline fallback instead of erroring
    _openai_key_present = (override_api_key or app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or os.getenv('OPENROUTER_API_KEY') or os.getenv('TOGETHER_API_KEY') or os.getenv('FIREWORKS_API_KEY') or '').strip()
    _enable_ollama = str(app.config.get('ENABLE_OLLAMA') or os.getenv('ENABLE_OLLAMA') or '').lower() in ('1','true','yes','on')
    if not (_openai_key_present or _enable_ollama):
        # Minimal offline templates so the user can proceed without AI
        def _fallback_content() -> str:
            if is_requirements_compliance:
                return (
                    f"Our organization confirms that we have carefully reviewed the solicitation and all associated documents, "
                    f"including addenda, forms, certifications, and submission requirements. We acknowledge and will comply with "
                    f"all mandatory instructions, general and special conditions, technical specifications, and submittal formats. "
                    f"We will submit every required attachment and form in the prescribed manner and by the deadline "
                    f"{f'({due_date_str})' if due_date_str else ''}. "
                    f"Any exceptions, if required, will be listed in the Deviations section; otherwise, none are taken."
                ) + "\n\n" + (
                    f"We further certify that our proposed approach will address all deliverables and key activities mandated by the RFP. "
                    f"We will provide complete documentation, maintain the requested quality standards, and coordinate with the client to "
                    f"ensure smooth execution in accordance with the scope and evaluation criteria. This Statement of Compliance is submitted "
                    f"in good faith and may be customized once the full RFP text is analyzed by the AI service."
                )
            # Generic single-section fallback
            return (
                f"AI is temporarily unavailable. This section can be drafted using the attached RFP and company profile. "
                f"Summarize the client objective, list all required deliverables, and briefly describe our compliant approach. "
                f"Once AI service is restored, click Generate again to expand this into a complete narrative."
            )
        _fc = _fallback_content()
        _paras = [p.strip() for p in re.split(r'\n\s*\n', _fc.strip()) if p.strip()] or [_fc.strip()]
        # Simple heuristic flags to surface potential suggestions
        def _heuristic_flags(section_paragraphs):
            import re as _re
            trig = _re.compile(r'\b(we\s+recommend|we\s+propose|the\s+client\s+should|it\s+is\s+advisable|best\s+practice)\b', _re.I)
            results = []
            for i, para in enumerate(section_paragraphs):
                for s in [s.strip() for s in _re.split(r'(?<=[.!?])\s+', para) if s.strip()]:
                    if trig.search(s):
                        results.append({'paragraph_index': i, 'sentence': s, 'reason': 'suggestion_outside_rfp', 'rationale': 'Heuristic match'})
            return results
        _flags = _heuristic_flags(_paras)
        return jsonify({
            'content': _paras,
            'raw': "\n\n".join(_paras),
            'flags': {'outside_rfp_recommendations': _flags, 'flag_count': len(_flags)},
            'model': 'offline-fallback',
            'usage': {},
        })
    
    # Respect safe local-only mode (do not send content to cloud providers)
    _safe_local_only = str(app.config.get('SAFE_LOCAL_ONLY') or os.getenv('SAFE_LOCAL_ONLY') or '').lower() in ('1','true','yes','on')
    def _chat_with_fallback(messages, temperature=0.3, max_tokens=None):
        # Attempt 1: explicit override or server OpenAI-compatible (OpenRouter/Together/Fireworks/OpenAI)
        api_key_a = ''
        base_url_a = ''
        model_a = ''
        if not _safe_local_only:
            api_key_a = (
                override_api_key
                or app.config.get('OPENAI_API_KEY')
                or os.getenv('OPENAI_API_KEY')
                or os.getenv('OPENROUTER_API_KEY')
                or os.getenv('TOGETHER_API_KEY')
                or os.getenv('FIREWORKS_API_KEY')
                or ''
            ).strip()
            base_url_a = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or os.getenv('OPENROUTER_BASE_URL') or 'https://openrouter.ai/api/v1').strip()
            model_a = (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'meta-llama/llama-4-scout-17b-16e-instruct').strip()
        attempts = []
        if api_key_a and not _safe_local_only:
            attempts.append(('openai', api_key_a, base_url_a, model_a))
        # Attempt 3: Local Ollama (enabled by default; set ENABLE_OLLAMA=0 to disable)
        enable_ollama = str(app.config.get('ENABLE_OLLAMA') or os.getenv('ENABLE_OLLAMA') or '1').lower() in ('1','true','yes','on')
        if enable_ollama and not _ollama_circuit_opened():
            ollama_base = (app.config.get('OLLAMA_BASE_URL') or os.getenv('OLLAMA_BASE_URL') or 'http://127.0.0.1:11434').strip()
            ollama_model = (app.config.get('OLLAMA_MODEL') or os.getenv('OLLAMA_MODEL') or 'llama3').strip()
            # Quick preflight to avoid long timeouts if Ollama is not running
            try:
                _ping = _get_ollama_session().get(ollama_base.rstrip('/') + '/api/tags', timeout=3)
                if _ping.status_code < 500:
                    attempts.append(('ollama', '', ollama_base, ollama_model))
                elif _safe_local_only:
                    # In safe-local-only, still try even if ping is flaky
                    attempts.append(('ollama', '', ollama_base, ollama_model))
            except requests.RequestException:
                # Skip Ollama attempt if not reachable
                if _safe_local_only and not _ollama_circuit_opened():
                    # In safe mode, attempt anyway to give it a chance
                    attempts.append(('ollama', '', ollama_base, ollama_model))
        last_error = None
        for provider, key, base_url, model in attempts:
            try:
                if provider == 'ollama':
                    endpoint = base_url.rstrip('/') + '/api/chat'
                    headers = {'Content-Type': 'application/json'}
                    options = {
                        'temperature': temperature,
                    }
                    if isinstance(max_tokens, int) and max_tokens > 0:
                        options['num_predict'] = max_tokens
                    # Nudge toward concise, non-repetitive outputs
                    options.update({
                        'top_p': 0.9,
                        'repeat_penalty': 1.12,
                        'repeat_last_n': 256,
                        'mirostat': 2,
                        'mirostat_tau': 5.0,
                        'mirostat_eta': 0.1,
                    })
                    body = {'model': model, 'messages': messages, 'stream': False, 'options': options}
                    # Health check for resource monitoring
                    _ollama_health_check()
                    # Improved retry logic with exponential backoff
                    _last_local_err = None
                    max_ollama_retries = 3
                    for _attempt_idx in range(max_ollama_retries):
                        try:
                            resp = _get_ollama_session().post(endpoint, headers=headers, json=body, timeout=(10, 60))
                            data = resp.json()
                            if resp.status_code >= 400:
                                last_error = f"ollama error {resp.status_code}: {data}"
                                _ollama_record_failure()
                                if _attempt_idx < max_ollama_retries - 1:
                                    _time.sleep(2 ** _attempt_idx)  # Exponential backoff
                                    continue
                                break
                            content_val = ''
                            if isinstance(data, dict):
                                if isinstance(data.get('message'), dict):
                                    content_val = data['message'].get('content', '') or ''
                                elif isinstance(data.get('messages'), list) and data['messages']:
                                    content_val = (data['messages'][-1] or {}).get('content', '') or ''
                                elif 'response' in data:
                                    content_val = data.get('response', '') or ''
                            if not content_val:
                                last_error = 'ollama returned empty content'
                                _ollama_record_failure()
                                if _attempt_idx < max_ollama_retries - 1:
                                    _time.sleep(2 ** _attempt_idx)
                                    continue
                                break
                            _ollama_record_success()
                            return (provider, content_val, data, model)
                        except requests.Timeout as err:
                            _last_local_err = err
                            last_error = f"{provider} network timeout: {err}"
                            if _attempt_idx < max_ollama_retries - 1:
                                _time.sleep(2 ** _attempt_idx)
                                continue
                            _ollama_record_failure()
                        except requests.RequestException as err:
                            _last_local_err = err
                            last_error = f"{provider} network error: {err}"
                            if _attempt_idx < max_ollama_retries - 1:
                                _time.sleep(2 ** _attempt_idx)
                                continue
                            _ollama_record_failure()
                        except ValueError:
                            last_error = f"{provider} returned non-JSON"
                            _ollama_record_failure()
                            break
                    # After retries, move on to next provider
                    continue
                else:
                    endpoint = base_url.rstrip('/')
                    if not endpoint.endswith('/chat/completions'):
                        endpoint = f"{endpoint}/chat/completions"
                    headers = {'Authorization': f'Bearer {key}', 'Content-Type': 'application/json'}
                    body = {'model': model, 'messages': messages, 'temperature': temperature}
                    if isinstance(max_tokens, int) and max_tokens > 0:
                        body['max_tokens'] = max_tokens
                    resp = requests.post(endpoint, headers=headers, json=body, timeout=90)
                    data = resp.json()
                    if resp.status_code >= 400:
                        err_message = data.get('error') if isinstance(data, dict) else None
                        detail = ''
                        if isinstance(err_message, dict):
                            detail = err_message.get('message') or err_message.get('code') or ''
                        elif isinstance(err_message, str):
                            detail = err_message
                        last_error = f"{provider} error {resp.status_code}: {detail or 'unknown'}"
                        continue
                    choices_local = data.get('choices') or []
                    if not choices_local:
                        last_error = f"{provider} returned no choices"
                        continue
                    return (provider, (choices_local[0].get('message', {}) or {}).get('content', ''), data, model)
            except requests.RequestException as err:
                last_error = f"{provider} network error: {err}"
                continue
            except ValueError:
                last_error = f"{provider} returned non-JSON"
                continue
        raise Exception(last_error or 'No AI provider configured')
    
    # Dynamic token budget for conciseness (does not change prompts)
    def _section_token_budget():
        if is_requirements_compliance:
            return 420
        if is_letter_of_transmittal:
            return 450
        if is_executive_summary:
            return 520
        if is_scope_objectives:
            return 700
        if is_proposed_technical_solution:
            return 900
        if is_implementation_work_plan:
            return 900
        if is_solution_deliverables:
            return 650
        if is_project_management_plan:
            return 800
        if is_pricing_proposals:
            return 450
        if is_contractual_terms:
            return 550
        if is_corporate_qualifications:
            return 720
        if is_project_schedule_milestones:
            return 550
        if is_project_team:
            return 550
        if is_addenda:
            return 260
        if is_appendix:
            return 420
        return 560
    try:
        provider_used, content, data, model = _chat_with_fallback(
            [
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_prompt}
            ],
            temperature=0.3,
            max_tokens=_section_token_budget()
        )
    except Exception as e:
        # Graceful fallback if all providers failed (e.g., Ollama not running)
        app.logger.error("Section generation failed, using offline fallback: %s", e)
        # Offline fallback text when all AI providers fail (e.g., Ollama not running or network issues)
        if is_requirements_compliance:
            _fc = (
                "AI service is temporarily unavailable. The following draft can be customized:\n\n"
                "Statement of Compliance: We confirm full compliance with all mandatory requirements, "
                "forms, deliverables, and conditions of the solicitation. Any exceptions will be listed "
                "in the Deviations section; otherwise, none are taken."
            )
        elif is_executive_summary:
            # More helpful Executive Summary fallback (~500 words, non-repetitive)
            _fc = (
                "AI service is temporarily unavailable. The following Executive Summary draft can be customized:\n\n"
                "Our company is pleased to submit this proposal in response to the client's solicitation. "
                "We understand that the client is seeking a qualified partner to deliver a reliable, "
                "cost-effective, and compliant solution that addresses current operational challenges and "
                "supports long-term strategic objectives. Based on our review of the bid document, the "
                "client requires improved performance, clear accountability, and measurable results, all "
                "delivered within the schedule and budget constraints outlined in the RFP.\n\n"
                "Our proposed approach is designed to fully meet these requirements while minimizing risk "
                "for the client. We combine proven methodologies, experienced personnel, and carefully "
                "selected technologies to ensure a solution that is both technically sound and practical "
                "to implement. The work plan emphasizes thorough planning, proactive communication, and "
                "disciplined project controls so that milestones are achieved on time and any issues are "
                "identified and addressed early. Operational processes are structured to comply with all "
                "specifications in the solicitation while remaining flexible enough to adapt to the "
                "client's evolving needs. Throughout the engagement, we will coordinate closely with the "
                "client's designated representatives to maintain transparency and ensure that expectations "
                "are consistently met.\n\n"
                "By selecting our company, the client gains a partner with relevant experience, a strong "
                "track record of successfully delivering similar projects, and a commitment to quality, "
                "safety, and customer service. Our team offers deep subject-matter expertise, disciplined "
                "execution, and responsive support during and after implementation. We are committed to "
                "providing a solution that not only satisfies the immediate requirements of this "
                "solicitation but also creates long-term value for the client and its stakeholders."
            )
        else:
            _fc = (
                "AI service is temporarily unavailable. The following draft can be customized:\n\n"
                "Draft section placeholder. Summarize the RFP objective, our compliant approach, and required deliverables."
            )
        _paras = [p.strip() for p in re.split(r'\n\s*\n', _fc.strip()) if p.strip()] or [_fc.strip()]
        # Heuristic flags
        def _heuristic_flags(section_paragraphs):
            import re as _re
            trig = _re.compile(r'\b(we\s+recommend|we\s+propose|the\s+client\s+should|it\s+is\s+advisable|best\s+practice)\b', _re.I)
            results = []
            for i, para in enumerate(section_paragraphs):
                for s in [s.strip() for s in _re.split(r'(?<=[.!?])\s+', para) if s.strip()]:
                    if trig.search(s):
                        results.append({'paragraph_index': i, 'sentence': s, 'reason': 'suggestion_outside_rfp', 'rationale': 'Heuristic match'})
            return results
        _flags = _heuristic_flags(_paras)
        return jsonify({
            'content': _paras,
            'raw': "\n\n".join(_paras),
            'flags': {'outside_rfp_recommendations': _flags, 'flag_count': len(_flags)},
            'model': 'offline-fallback',
            'usage': {},
        })
    
    if not content:
        return jsonify({'error': 'api_unavailable', 'message': 'The AI service returned an empty response.'}), 502
    
    # Split content into paragraphs for frontend display
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', content.strip()) if p.strip()]
    if not paragraphs:
        paragraphs = [content.strip()]
    # Executive Summary: defensively remove any accidental header line like "Client, City, State Date"
    if is_executive_summary and paragraphs:
        try:
            import re as _re
            first = paragraphs[0].strip()
            header_like = False
            if len(first) <= 120:
                # Matches "October 17, 2025" etc.
                month_pat = r'(jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*'
                if _re.search(rf'\b{month_pat}\b\s+\d{{1,2}},\s*\d{{4}}', first, _re.I):
                    header_like = True
            if not header_like:
                # Heuristic: mentions civic entities plus a year
                civic_terms = ('city', 'county', 'state', 'department', 'school district', 'university', 'college', 'authority')
                if any(t in first.lower() for t in civic_terms) and _re.search(r'\b\d{4}\b', first):
                    header_like = True
            if header_like and len(paragraphs) > 1:
                paragraphs = paragraphs[1:]
        except Exception:
            pass

    # AI guardrail: ask the LLM to identify recommendation/suggestion sentences that are not supported by the RFP
    def _ai_flag_outside_recommendations(section_paragraphs: list[str], rfp_text: str) -> list[dict]:
        try:
            joined_section = "\n\n".join(section_paragraphs)
            # Try local-only first if safe mode; else OpenAI-compatible
            attempts = []
            if not _safe_local_only:
                attempts.append((
                    (app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or os.getenv('OPENROUTER_API_KEY') or os.getenv('TOGETHER_API_KEY') or os.getenv('FIREWORKS_API_KEY') or '').strip(),
                    (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or os.getenv('OPENROUTER_BASE_URL') or 'https://openrouter.ai/api/v1').strip(),
                    (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'meta-llama/llama-4-scout-17b-16e-instruct').strip()
                ))

            # Trim RFP to keep request small but useful
            rfp_trimmed = (rfp_text or '')[:20000]
            system_msg = (
                "You are a compliance auditor for proposal content. "
                "Given an RFP and a draft proposal section, identify only those sentences in the proposal text that are recommendations/suggestions/proposals/advice introduced by the bidder and are not explicitly supported, requested, or quoted by the RFP. "
                "Be conservative: flag only when clearly unsupported. "
                "Return ONLY compact JSON: {\"outside\":[{\"sentence\":\"...\",\"rationale\":\"...\"}]}. No commentary."
            )
            user_msg = (
                f"RFP TEXT:\n{rfp_trimmed}\n\n---\n\n"
                f"PROPOSAL SECTION TEXT:\n{joined_section}\n\n"
                "Instructions:\n"
                "- Evaluate proposal sentences in context of the RFP.\n"
                "- A sentence is 'outside' if it expresses a recommendation/suggestion/advice (e.g., 'we recommend', 'we propose', 'the client should', 'it is advisable') that the RFP does not require, request, or imply.\n"
                "- Do NOT list factual descriptions, mandatory compliance assertions, or paraphrases of RFP requirements.\n"
                "- Output JSON only."
            )
            reply = ''
            if _safe_local_only:
                # Use local Ollama if available; otherwise fall back to heuristics
                messages_local = [
                    {'role': 'system', 'content': system_msg},
                    {'role': 'user', 'content': user_msg}
                ]
                try:
                    reply = _ollama_chat(messages_local, temperature=0.1, timeout_sec=45) or ''
                except Exception:
                    reply = ''
            else:
                for key, base_url_local, model_local in attempts:
                    if not key:
                        continue
                    endpoint_local = base_url_local.rstrip('/')
                    if not endpoint_local.endswith('/chat/completions'):
                        endpoint_local = f"{endpoint_local}/chat/completions"
                    headers_local = {'Authorization': f'Bearer {key}', 'Content-Type': 'application/json'}
                    body_local = {
                        'model': model_local,
                        'messages': [
                            {'role': 'system', 'content': system_msg},
                            {'role': 'user', 'content': user_msg}
                        ],
                        'temperature': 0.1,
                    }
                    try:
                        resp = requests.post(endpoint_local, headers=headers_local, json=body_local, timeout=90)
                        data = resp.json()
                        if resp.status_code >= 400:
                            continue
                        choices = data.get('choices') or []
                        reply = (choices[0].get('message', {}) or {}).get('content', '') if choices else ''
                        if reply:
                            break
                    except Exception:
                        continue
            # Attempt to parse JSON directly; strip code fences if present
            text = reply.strip()
            if text.startswith('```'):
                text = text.lstrip('`')
                fence = text.find('```')
                if fence != -1:
                    text = text[:fence]
            import json as _json
            parsed = {}
            try:
                parsed = _json.loads(text)
            except Exception:
                # Try extracting JSON substring
                m = re.search(r'\{[\s\S]*\}', text)
                if m:
                    try:
                        parsed = _json.loads(m.group(0))
                    except Exception:
                        parsed = {}
            outside = parsed.get('outside') if isinstance(parsed, dict) else None
            if not isinstance(outside, list):
                outside = []

            # Map sentences to paragraph indexes
            lower_paras = [p.lower() for p in section_paragraphs]
            results = []
            for item in outside:
                sentence = (item.get('sentence') or '').strip()
                rationale = (item.get('rationale') or '').strip()
                if not sentence:
                    continue
                target = sentence.lower()
                p_idx = -1
                for i, p in enumerate(lower_paras):
                    if target and target in p:
                        p_idx = i
                        break
                results.append({
                    'paragraph_index': p_idx if p_idx >= 0 else 0,
                    'sentence': sentence,
                    'reason': 'suggestion_outside_rfp',
                    'rationale': rationale
                })
            return results
        except Exception:
            return []

    outside_flags = _ai_flag_outside_recommendations(paragraphs, rfp_content or '')
    # Heuristic fallback when AI flagger is unavailable or returns nothing
    if not outside_flags:
        def _heuristic_flag_outside(section_paragraphs):
            import re as _re
            trigger_phrases = [
                r'\bwe\s+recommend\b', r'\bour\s+recommendation\b', r'\bwe\s+propose\b', r'\bour\s+proposal\b',
                r'\bthe\s+client\s+should\b', r'\bshould\b', r'\bit\s+is\s+advisable\b', r'\brecommend\b',
                r'\bsuggestion\b', r'\badvis(e|ory)\b', r'\bbest\s+practice\b'
            ]
            trig_re = _re.compile('|'.join(trigger_phrases), _re.IGNORECASE)
            results = []
            for idx, para in enumerate(section_paragraphs):
                # Split naive sentences
                sentences = [s.strip() for s in _re.split(r'(?<=[.!?])\s+', para) if s.strip()]
                for s in sentences:
                    if trig_re.search(s):
                        results.append({
                            'paragraph_index': idx,
                            'sentence': s,
                            'reason': 'suggestion_outside_rfp',
                            'rationale': 'Heuristic match for recommendation/suggestion phrasing; verify against RFP.'
                        })
            return results
        outside_flags = _heuristic_flag_outside(paragraphs)

    # Write section result to cache for re-use
    try:
        if '_sec_cache_key' in locals() and _sec_cache_key:
            _llm_cache_set(_sec_cache_key, {
                'content': paragraphs,
                'raw': "\n\n".join(paragraphs).strip(),
                'flags': {
                    'outside_rfp_recommendations': outside_flags,
                    'flag_count': len(outside_flags)
                },
                'model': data.get('model', model),
                'usage': data.get('usage', {}),
            })
    except Exception:
        pass
    return jsonify({
        'content': paragraphs,
        'raw': "\n\n".join(paragraphs).strip(),
        'flags': {
            'outside_rfp_recommendations': outside_flags,
            'flag_count': len(outside_flags)
        },
        'model': data.get('model', model),
        'usage': data.get('usage', {}),
    })


@app.route('/api/refine-section-content', methods=['POST'])
@login_required
def refine_section_content():
    """Refine existing section content to strictly align with the RFP and remove unsupported suggestions."""
    payload = request.get_json(silent=True) or {}
    paragraphs = payload.get('paragraphs') or []
    if isinstance(paragraphs, str):
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', paragraphs) if p.strip()]
    section_id = (payload.get('section_id') or '').strip()
    section_title = (payload.get('section_title') or '').strip()
    bid_id = payload.get('bid_id')
    company = (payload.get('company') or '').strip()
    contact_email = (payload.get('contact_email') or '').strip()
    override_api_key = (payload.get('api_key') or '').strip()
    override_model = (payload.get('model') or '').strip()

    if not paragraphs:
        return jsonify({'error': 'missing_params', 'message': 'No content to refine.'}), 400

    rfp_content = ''
    if bid_id:
        rfp_content = _fetch_rfp_content_for_bid(bid_id, page_limit=100, char_limit=25000)

    # Prepare LLM request with fallback (OpenAI-compatible -> Ollama)
    def _chat_with_fallback(messages, temperature=0.2, max_tokens=None):
        _safe_local_only_ref = str(app.config.get('SAFE_LOCAL_ONLY') or os.getenv('SAFE_LOCAL_ONLY') or '').lower() in ('1','true','yes','on')
        api_key_a = ''
        base_url_a = ''
        model_a = ''
        if not _safe_local_only_ref:
            api_key_a = (
                override_api_key
                or app.config.get('OPENAI_API_KEY')
                or os.getenv('OPENAI_API_KEY')
                or os.getenv('OPENROUTER_API_KEY')
                or os.getenv('TOGETHER_API_KEY')
                or os.getenv('FIREWORKS_API_KEY')
                or ''
            ).strip()
            base_url_a = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or os.getenv('OPENROUTER_BASE_URL') or 'https://openrouter.ai/api/v1').strip()
            model_a = (override_model or app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'meta-llama/llama-4-scout-17b-16e-instruct').strip()
        attempts = []
        if api_key_a and not _safe_local_only_ref:
            attempts.append((api_key_a, base_url_a, model_a, 'openai'))
        # Optional Ollama third attempt (enabled by default; set ENABLE_OLLAMA=0 to disable)
        enable_ollama = str(app.config.get('ENABLE_OLLAMA') or os.getenv('ENABLE_OLLAMA') or '1').lower() in ('1','true','yes','on')
        if enable_ollama and not _ollama_circuit_opened():
            ollama_base = (app.config.get('OLLAMA_BASE_URL') or os.getenv('OLLAMA_BASE_URL') or 'http://127.0.0.1:11434').strip()
            ollama_model = (app.config.get('OLLAMA_MODEL') or os.getenv('OLLAMA_MODEL') or 'llama3').strip()
            # Quick preflight to avoid long timeouts if Ollama is not running
            try:
                _ping = _get_ollama_session().get(ollama_base.rstrip('/') + '/api/tags', timeout=3)
                if _ping.status_code < 500:
                    attempts.append(('', ollama_base, ollama_model, 'ollama'))
                elif _safe_local_only_ref:
                    attempts.append(('', ollama_base, ollama_model, 'ollama'))
            except requests.RequestException:
                # Skip Ollama attempt if not reachable
                if _safe_local_only_ref and not _ollama_circuit_opened():
                    attempts.append(('', ollama_base, ollama_model, 'ollama'))
        last_error = None
        for key, base_url, model, provider in attempts:
            endpoint_local = base_url.rstrip('/')
            if provider == 'ollama':
                try:
                    headers_local = {'Content-Type': 'application/json'}
                    options_local = {
                        'temperature': temperature,
                    }
                    if isinstance(max_tokens, int) and max_tokens > 0:
                        options_local['num_predict'] = max_tokens
                    options_local.update({
                        'top_p': 0.9,
                        'repeat_penalty': 1.12,
                        'repeat_last_n': 256,
                        'mirostat': 2,
                        'mirostat_tau': 5.0,
                        'mirostat_eta': 0.1,
                    })
                    body_local = {'model': model, 'messages': messages, 'stream': False, 'options': options_local}
                    # Retry once with shorter read timeout to avoid long hangs
                    _last_local_err = None
                    for _attempt_idx in range(2):
                        try:
                            resp = requests.post(endpoint_local + '/api/chat', headers=headers_local, json=body_local, timeout=(5, 20))
                            data_local = resp.json()
                            if resp.status_code >= 400:
                                last_error = f"ollama error {resp.status_code}: {data_local}"
                                _ollama_record_failure()
                                break
                            content_val = ''
                            if isinstance(data_local, dict):
                                if isinstance(data_local.get('message'), dict):
                                    content_val = data_local['message'].get('content', '') or ''
                                elif isinstance(data_local.get('messages'), list) and data_local['messages']:
                                    content_val = (data_local['messages'][-1] or {}).get('content', '') or ''
                                elif 'response' in data_local:
                                    content_val = data_local.get('response', '') or ''
                            if not content_val:
                                last_error = 'ollama returned empty content'
                                _ollama_record_failure()
                                break
                            _ollama_record_success()
                            return (provider, content_val, data_local, model, base_url, endpoint_local)
                        except requests.Timeout as err:
                            _last_local_err = err
                            last_error = f"{provider} network timeout: {err}"
                            if _attempt_idx == 0:
                                _time.sleep(0.3)
                                continue
                            _ollama_record_failure()
                        except requests.RequestException as err:
                            _last_local_err = err
                            last_error = f"{provider} network error: {err}"
                            _ollama_record_failure()
                            break
                        except ValueError:
                            last_error = f"{provider} returned non-JSON"
                            _ollama_record_failure()
                            break
                    continue
                except Exception as err:
                    last_error = f"{provider} unexpected error: {err}"
                    _ollama_record_failure()
                    continue
            if provider != 'ollama':
                if not endpoint_local.endswith('/chat/completions'):
                    endpoint_local = f"{endpoint_local}/chat/completions"
                headers_local = {'Authorization': f'Bearer {key}', 'Content-Type': 'application/json'}
                body_local = {'model': model, 'messages': messages, 'temperature': temperature}
                if isinstance(max_tokens, int) and max_tokens > 0:
                    body_local['max_tokens'] = max_tokens
                try:
                    resp = requests.post(endpoint_local, headers=headers_local, json=body_local, timeout=120)
                except requests.RequestException as err:
                    last_error = f"{provider} network error: {err}"
                    continue
                try:
                    data_local = resp.json()
                except ValueError:
                    last_error = f"{provider} returned non-JSON (status {resp.status_code})"
                    continue
                if resp.status_code >= 400:
                    err_message = data_local.get('error') if isinstance(data_local, dict) else None
                    detail = ''
                    if isinstance(err_message, dict):
                        detail = err_message.get('message') or err_message.get('code') or ''
                    elif isinstance(err_message, str):
                        detail = err_message
                    last_error = f"{provider} error {resp.status_code}: {detail or 'unknown'}"
                    continue
                choices_local = data_local.get('choices') or []
                if not choices_local:
                    last_error = f"{provider} returned no choices"
                    continue
                return (provider, (choices_local[0].get('message', {}) or {}).get('content', ''), data_local, model, base_url, endpoint_local)
        raise Exception(last_error or 'No AI provider configured')

    system_prompt = (
        "You are an expert proposal editor for government RFP responses. "
        "Given the RFP and a draft section, rewrite the section to be strictly compliant: "
        "remove or rephrase any unsupported recommendations/suggestions, keep required assertions, "
        "and strengthen alignment with the RFP requirements. Maintain professional tone. "
        "Return only the refined section text that can be split into paragraphs."
    )
    user_prompt = f"""RFP TEXT:
{(rfp_content or '')[:20000]}

---

CURRENT SECTION TEXT:
{'\n\n'.join(paragraphs)}

Rewrite the section so it strictly aligns with the RFP. Do not add any new commitments that the RFP does not imply.
Return only the refined text."""

    try:
        provider_used, refined_text, data, model, base_url, endpoint = _chat_with_fallback(
            [
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_prompt}
            ],
            temperature=0.2,
            max_tokens=560
        )
    except Exception as e:
        return jsonify({'error': 'api_unavailable', 'message': f'AI service returned an error: {e}'}), 502
    if not refined_text:
        return jsonify({'error': 'api_unavailable', 'message': 'The AI service returned an empty response.'}), 502
    refined_paragraphs = [p.strip() for p in re.split(r'\n\s*\n', refined_text.strip()) if p.strip()] or [refined_text.strip()]

    # Re-run AI guardrail to report any remaining unsupported suggestions
    def _ai_flag_outside_recommendations_ref(section_paragraphs, rfp_text):
        try:
            joined_section = "\n\n".join(section_paragraphs)
            # Try using same provider used for refinement; if unavailable, fall back to OpenAI-compatible default
            attempts = []
            _safe_local_only_ref = str(app.config.get('SAFE_LOCAL_ONLY') or os.getenv('SAFE_LOCAL_ONLY') or '').lower() in ('1','true','yes','on')
            if not _safe_local_only_ref:
                attempts.append((base_url, model))
                attempts.append((app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or os.getenv('OPENROUTER_BASE_URL') or 'https://openrouter.ai/api/v1', app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'meta-llama/llama-4-scout-17b-16e-instruct'))
            rfp_trimmed = (rfp_text or '')[:20000]
            system_msg = (
                "You are a compliance auditor for proposal content. "
                "Given an RFP and a draft proposal section, identify only those sentences in the proposal text that are recommendations/suggestions/proposals/advice introduced by the bidder and are not explicitly supported by the RFP. "
                "Return ONLY JSON: {\"outside\":[{\"sentence\":\"...\",\"rationale\":\"...\"}]}"
            )
            user_msg = f"RFP TEXT:\n{rfp_trimmed}\n\n---\n\nPROPOSAL SECTION TEXT:\n{joined_section}\n\nOutput JSON only."
            reply = ''
            _safe_local_only_ref = str(app.config.get('SAFE_LOCAL_ONLY') or os.getenv('SAFE_LOCAL_ONLY') or '').lower() in ('1','true','yes','on')
            if _safe_local_only_ref:
                try:
                    messages_local = [
                        {'role': 'system', 'content': system_msg},
                        {'role': 'user', 'content': user_msg}
                    ]
                    reply = _ollama_chat(messages_local, temperature=0.1, timeout_sec=45) or ''
                except Exception:
                    reply = ''
            else:
                for base_url_local, model_local in attempts:
                    # Choose key based on base_url
                    key = (app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or os.getenv('OPENROUTER_API_KEY') or os.getenv('TOGETHER_API_KEY') or os.getenv('FIREWORKS_API_KEY') or '').strip()
                    if not key:
                        continue
                    endpoint_local = base_url_local.rstrip('/')
                    if not endpoint_local.endswith('/chat/completions'):
                        endpoint_local = f"{endpoint_local}/chat/completions"
                    headers_local = {'Authorization': f'Bearer {key}', 'Content-Type': 'application/json'}
                    body_local = {
                        'model': model_local,
                        'messages': [
                            {'role': 'system', 'content': system_msg},
                            {'role': 'user', 'content': user_msg}
                        ],
                        'temperature': 0.1,
                    }
                    try:
                        resp = requests.post(endpoint_local, headers=headers_local, json=body_local, timeout=90)
                        data_local = resp.json()
                        if resp.status_code >= 400:
                            continue
                        choices_local = data_local.get('choices') or []
                        reply = (choices_local[0].get('message', {}) or {}).get('content', '') if choices_local else ''
                        if reply:
                            break
                    except Exception:
                        continue
            text = reply.strip()
            if text.startswith('```'):
                text = text.lstrip('`')
                fence = text.find('```')
                if fence != -1:
                    text = text[:fence]
            import json as _json
            parsed = {}
            try:
                parsed = _json.loads(text)
            except Exception:
                m = re.search(r'\{[\s\S]*\}', text)
                if m:
                    try:
                        parsed = _json.loads(m.group(0))
                    except Exception:
                        parsed = {}
            outside = parsed.get('outside') if isinstance(parsed, dict) else None
            if not isinstance(outside, list):
                outside = []
            lower_paras = [p.lower() for p in section_paragraphs]
            results = []
            for item in outside:
                sentence = (item.get('sentence') or '').strip()
                rationale = (item.get('rationale') or '').strip()
                if not sentence:
                    continue
                target = sentence.lower()
                p_idx = -1
                for i, p in enumerate(lower_paras):
                    if target and target in p:
                        p_idx = i
                        break
                results.append({
                    'paragraph_index': p_idx if p_idx >= 0 else 0,
                    'sentence': sentence,
                    'reason': 'suggestion_outside_rfp',
                    'rationale': rationale
                })
            return results
        except Exception:
            return []

    outside_flags = _ai_flag_outside_recommendations_ref(refined_paragraphs, rfp_content or '')
    # Heuristic fallback when AI flagger is unavailable or returns nothing
    if not outside_flags:
        def _heuristic_flag_outside_ref(section_paragraphs):
            import re as _re
            trigger_phrases = [
                r'\bwe\s+recommend\b', r'\bour\s+recommendation\b', r'\bwe\s+propose\b', r'\bour\s+proposal\b',
                r'\bthe\s+client\s+should\b', r'\bshould\b', r'\bit\s+is\s+advisable\b', r'\brecommend\b',
                r'\bsuggestion\b', r'\badvis(e|ory)\b', r'\bbest\s+practice\b'
            ]
            trig_re = _re.compile('|'.join(trigger_phrases), _re.IGNORECASE)
            results = []
            for idx, para in enumerate(section_paragraphs):
                sentences = [s.strip() for s in _re.split(r'(?<=[.!?])\s+', para) if s.strip()]
                for s in sentences:
                    if trig_re.search(s):
                        results.append({
                            'paragraph_index': idx,
                            'sentence': s,
                            'reason': 'suggestion_outside_rfp',
                            'rationale': 'Heuristic match for recommendation/suggestion phrasing; verify against RFP.'
                        })
            return results
        outside_flags = _heuristic_flag_outside_ref(refined_paragraphs)

    return jsonify({
        'content': refined_paragraphs,
        'raw': refined_text.strip(),
        'flags': {
            'outside_rfp_recommendations': outside_flags,
            'flag_count': len(outside_flags)
        },
        'model': model,
        'usage': data.get('usage', {}),
    })
@app.route('/api/compile-proposal-sections', methods=['POST'])
@login_required
def compile_proposal_sections():
    """Compile all generated section content into a single downloadable Word document."""
    payload = request.get_json(silent=True) or {}
    sections = payload.get('sections') or []
    # Server-side guard: only include sections explicitly marked as Saved and with content
    filtered_sections = []
    for s in sections:
        try:
            status_ok = str((s.get('status') or '')).strip().lower() == 'saved'
            has_content = bool(s.get('raw_content')) or (isinstance(s.get('content'), list) and len(s.get('content')) > 0)
            if status_ok and has_content:
                filtered_sections.append(s)
        except Exception:
            continue
    sections = filtered_sections
    if not isinstance(sections, list) or not sections:
        return jsonify({'error': 'missing_sections', 'message': 'No sections were provided to compile.'}), 400

    project_title = (payload.get('project_title') or '').strip()
    company_name = (payload.get('company') or '').strip()
    bid_id = payload.get('bid_id') or payload.get('g_id') or None
    html_fragments = []

    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return jsonify({'error': 'dependency_missing', 'message': 'python-docx library is required to compile proposals.'}), 500

    try:
        # Always prefer the H & F template for consistent formatting
        template_used = None
        document = None
        try:
            chosen_path = choose_company_template(company_name)
            if chosen_path:
                document = Document(chosen_path)
                _bn_lower = os.path.basename(chosen_path).lower()
                if 'sunsprint' in _bn_lower:
                    template_used = 'sunsprint'
                elif 'Proposal ikio' in _bn_lower or _bn_lower.endswith('Proposal ikio.docx'):
                    # IKIO-branded template
                    template_used = 'ikio'
                elif 'h & f' in _bn_lower or 'h&f' in _bn_lower or _bn_lower == 'h' or _bn_lower.endswith('h & f.docx'):
                    # Treat H & F.docx as Metco-branded template
                    template_used = 'metco'
                else:
                    template_used = 'hf'
            else:
                # Fallback to blank document if no template is present
                document = Document()
                template_used = 'blank'
        except Exception:
            # Defensive fallback in case template loading fails
            document = Document()
            template_used = 'blank'
        
        # Enforce universal fonts only when using non-branded templates
        PREFERRED_BODY_FONT_NAME = 'Aptos Display'
        PREFERRED_HEADING_FONT_NAME = 'Aptos Bold'
        if template_used not in ('sunsprint', 'metco', 'ikio'):
            try:
                normal_style = document.styles['Normal']
                normal_style.font.name = PREFERRED_BODY_FONT_NAME
                try:
                    from docx.shared import Pt as _Pt
                    normal_style.font.size = _Pt(11)
                except Exception:
                    pass
            except Exception:
                pass
            # Also try to set common heading/title styles to Aptos
            for _style_name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6', 'Heading 7', 'Heading 8', 'Heading 9']:
                try:
                    document.styles[_style_name].font.name = PREFERRED_HEADING_FONT_NAME
                    try:
                        from docx.shared import Pt as _Pt
                        # Use 16pt for main section headings (Heading 1), keep others at 11pt
                        if _style_name == 'Heading 1':
                            document.styles[_style_name].font.size = _Pt(16)
                        else:
                            document.styles[_style_name].font.size = _Pt(11)
                        document.styles[_style_name].font.bold = True
                    except Exception:
                        pass
                except Exception:
                    continue
        
        # Ensure automatic header/footer fields:
        # - Header shows the current main heading via STYLEREF "Heading 1"
        # - Footer shows right-aligned "Page X of Y"
        try:
            section0 = document.sections[0]
            # Only add header/footer if template is blank/hf or existing parts are empty (avoid overriding branded templates)
            def _part_has_text(part):
                try:
                    return any((p.text or '').strip() for p in getattr(part, 'paragraphs', []) or [])
                except Exception:
                    return False
            
            # Header suppressed per requirement: no dynamic heading in header
            try:
                pass
            except Exception:
                pass
            
            # Footer page numbering logic removed per requirement
        except Exception:
            pass
        
        # Helper: style existence
        def _has_style(doc, style_name):
            try:
                _ = doc.styles[style_name]
                return True
            except Exception:
                return False
        
        # Add title page with robust style fallback
        if project_title:
            try:
                if _has_style(document, 'Title'):
                    title_para = document.add_paragraph(project_title, style=document.styles['Title'])
                elif _has_style(document, 'METCO Title'):
                    title_para = document.add_paragraph(project_title, style=document.styles['METCO Title'])
                else:
                    # Fall back to first-level heading if Title style is unavailable
                    title_para = document.add_heading(project_title, level=1)
                try:
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
            except Exception:
                # Absolute fallback to avoid style-related failures
                p = document.add_paragraph(project_title)
                try:
                    for run in p.runs:
                        run.font.name = PREFERRED_HEADING_FONT_NAME
                        from docx.shared import Pt as _Pt
                        run.font.bold = True
                        run.font.size = _Pt(11)
                except Exception:
                    pass
            html_fragments.append(f'<h1 style=\"text-align: center;\">{escape(project_title)}</h1>')
            document.add_paragraph()  # Add spacing
        
        # Collect section titles for Table of Contents (before processing content)
        section_titles = []
        for section in sections:
            if not isinstance(section, dict):
                continue
            title = (section.get('title') or 'Untitled Section').strip() or 'Untitled Section'
            content_blocks = section.get('content') if isinstance(section.get('content'), list) else []
            raw_html = section.get('raw_content') or ''
            if content_blocks or raw_html:
                section_titles.append(title)
        
        # Add Table of Contents page if we have sections
        if section_titles:
            document.add_page_break()
            toc_heading = document.add_heading('Table of Contents', level=1)
            try:
                from docx.shared import Pt as _Pt
                for run in toc_heading.runs:
                    run.font.name = PREFERRED_HEADING_FONT_NAME
                    run.font.bold = True
                    run.font.size = _Pt(11)
            except Exception:
                pass
            document.add_paragraph()  # Add spacing
            # Insert a dynamic, hyperlink-enabled TOC field that auto-calculates page numbers
            try:
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                toc_para = document.add_paragraph()
                r = toc_para.add_run()
                fldChar = OxmlElement('w:fldChar')
                fldChar.set(qn('w:fldCharType'), 'begin')
                # Mark as dirty so Word refreshes the TOC on open
                try:
                    fldChar.set(qn('w:dirty'), 'true')
                except Exception:
                    pass
                r._r.append(fldChar)
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                # \o "1-1" includes only Heading 1; \h hyperlinks; \z hides page numbers in web view; \u uses outline levels
                instrText.text = 'TOC \\o "1-1" \\h \\z \\u'
                r._r.append(instrText)
                fldChar = OxmlElement('w:fldChar')
                fldChar.set(qn('w:fldCharType'), 'separate')
                r._r.append(fldChar)
                # Placeholder text that Word replaces after field update
                toc_para.add_run('Table of Contents will update on open.')
                fldChar = OxmlElement('w:fldChar')
                fldChar.set(qn('w:fldCharType'), 'end')
                r._r.append(fldChar)
            except Exception:
                # Fallback: simple static list (no links)
                for idx, toc_title in enumerate(section_titles, 1):
                    toc_para = document.add_paragraph(f'{idx}. {toc_title}')
                    toc_para.paragraph_format.space_after = Pt(6)
                    try:
                        for run in toc_para.runs:
                            run.font.name = 'Aptos Display'
                    except Exception:
                        pass
            
            document.add_page_break()
            html_fragments.append('<h2>Table of Contents</h2>')
            html_fragments.append('<ol>' + ''.join(f'<li>{escape(title)}</li>' for title in section_titles) + '</ol>')
        
        # Process each section with generated content
        first_section_added = False
        # Determine a table style available in the template, prefer METCO-specific if present
        table_style_name = None
        try:
            from docx.enum.style import WD_STYLE_TYPE as _WD_STYLE_TYPE
            for s in document.styles:
                try:
                    if getattr(s, "type", None) == _WD_STYLE_TYPE.TABLE:
                        if s.name in ("METCO Table", "Table Grid", "Light Grid Accent 1", "Grid Table 4 Accent 1"):
                            table_style_name = s.name
                            if table_style_name == "METCO Table":
                                break
                except Exception:
                    continue
            if not table_style_name:
                table_style_name = "Table Grid"
        except Exception:
            table_style_name = "Table Grid"

        # End of initial preparation block
    except Exception as e:
        app.logger.exception('Unexpected error preparing document: %s', e)
        return jsonify({'error': 'unexpected_error', 'message': f'Compilation failed: {str(e)}'}), 500

    # Continue with compilation
    def _render_markdown_to_docx_paragraph(doc, text: str):
        """
        Render a single text block into the Word document, interpreting a minimal Markdown subset:
        - Lines starting with '### ' become Heading 3.
        - Bold spans wrapped in **double asterisks** become bold runs.
        Any remaining stray '*' or '#' characters are removed.
        """
        if not text:
            return
        try:
            import re as __re_md
            heading_match = __re_md.match(r'^\s*#{3,}\s*(.+?)\s*$', text)
            if heading_match:
                heading_text = heading_match.group(1)
                # Add a heading paragraph and style per requirement (Aptos, bold, 16pt)
                para = doc.add_heading(heading_text, level=3)
                try:
                    from docx.shared import Pt as _Pt
                    for run in para.runs:
                        run.font.bold = True
                        run.font.name = PREFERRED_HEADING_FONT_NAME
                        run.font.size = _Pt(11)
                except Exception:
                    pass
                return
            # Detect multi-line bullet list and render each as a bullet paragraph
            lines = [ln for ln in text.splitlines() if ln.strip()]
            has_bullets = any(__re_md.match(r'^\s*[\-\*\+•·▪●◦‣]\s+', ln) for ln in lines)
            if has_bullets and len(lines) >= 1:
                for raw in lines:
                    m = __re_md.match(r'^\s*[\-\*\+•·▪●◦‣]\s+(.*)$', raw)
                    if m:
                        content = m.group(1).strip()
                        # Strip markdown artifacts like **bold** within bullet items
                        try:
                            content = __re_md.sub(r'\*\*(.+?)\*\*', r'\1', content)
                            content = __re_md.sub(r'[#*]+', '', content)
                        except Exception:
                            pass
                        p = doc.add_paragraph()
                        applied_style = False
                        try:
                            # Prefer Word's built-in bullet list (disc/filled circle)
                            p.style = doc.styles['List Bullet']
                            applied_style = True
                        except Exception:
                            applied_style = False
                        # Fallback: manually prepend a filled circle bullet and a tab
                        try:
                            from docx.shared import Pt as _Pt
                            if not applied_style:
                                rb = p.add_run('•\t')
                                rb.font.name = PREFERRED_BODY_FONT_NAME
                                rb.font.size = _Pt(11)
                            rt = p.add_run(content)
                            rt.font.name = PREFERRED_BODY_FONT_NAME
                            rt.font.size = _Pt(11)
                        except Exception:
                            # Absolute fallback: simple text
                            p.add_run(content)
                    else:
                        # non-bullet line in the block -> normal paragraph
                        p = doc.add_paragraph(__re_md.sub(r'[#*]+', '', raw))
                        try:
                            from docx.shared import Pt as _Pt
                            for run in p.runs:
                                run.font.name = PREFERRED_BODY_FONT_NAME
                                run.font.size = _Pt(11)
                        except Exception:
                            pass
                return
            # Regular paragraph with bold spans
            paragraph = doc.add_paragraph()
            parts = __re_md.split(r'(\*\*.+?\*\*)', text)
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**') and len(part) >= 4:
                    content = part[2:-2]
                    run = paragraph.add_run(content)
                    try:
                        run.bold = True
                        run.font.name = PREFERRED_BODY_FONT_NAME
                        from docx.shared import Pt as _Pt
                        run.font.size = _Pt(11)
                    except Exception:
                        pass
                else:
                    # Remove leftover single '*' or '#' characters
                    cleaned = __re_md.sub(r'[#*]+', '', part)
                    run = paragraph.add_run(cleaned)
                    try:
                        run.font.name = PREFERRED_BODY_FONT_NAME
                        from docx.shared import Pt as _Pt
                        run.font.size = _Pt(11)
                    except Exception:
                        pass
        except Exception:
            # Fallback to simple paragraph if anything fails
            doc.add_paragraph(text)

    def _plaintext_markdown_to_html(text: str) -> str:
        """
        Convert a minimal Markdown subset to clean HTML string for preview.
        - '### Title' -> <h3>Title</h3>
        - '**bold**'  -> <strong>bold</strong>
        - Remove stray '*' and '#' characters elsewhere.
        """
        if not text:
            return ''
        import re as __re_md
        # Heading
        heading_match = __re_md.match(r'^\s*#{3,}\s*(.+?)\s*$', text)
        if heading_match:
            return f'<h3>{escape(heading_match.group(1))}</h3>'
        # Bullet lists
        if '\n' in text and any(l.strip().startswith(('-', '*', '+', '•')) for l in text.splitlines()):
            items = []
            for ln in [l for l in text.splitlines() if l.strip()]:
                ln_stripped = ln.strip()
                if ln_stripped[:1] in ('-', '*', '+') or ln_stripped.startswith('•'):
                    inner = ln_stripped[1:].strip()
                    # Remove simple markdown bold markers within list items for clean HTML preview
                    try:
                        import re as __re_md2
                        inner = __re_md2.sub(r'\*\*(.+?)\*\*', r'\1', inner)
                        inner = __re_md2.sub(r'[#*]+', '', inner)
                    except Exception:
                        pass
                    items.append(f'<li>{escape(inner)}</li>')
                else:
                    items.append(f'<li>{escape(ln_stripped)}</li>')
            return '<ul>' + ''.join(items) + '</ul>'
        # Bold spans
        def _bold_repl(m):
            return f'<strong>{escape(m.group(1))}</strong>'
        html = __re_md.sub(r'\*\*(.+?)\*\*', lambda m: _bold_repl(m), text)
        # Escape remaining content safely, but preserve the <strong> tags we just added
        # Strategy: temporarily replace strong tags, escape, then restore
        placeholder_open = '___STRONG_OPEN___'
        placeholder_close = '___STRONG_CLOSE___'
        html = html.replace('<strong>', placeholder_open).replace('</strong>', placeholder_close)
        html = escape(html)
        html = html.replace(placeholder_open, '<strong>').replace(placeholder_close, '</strong>')
        # Remove stray '*' and '#' characters
        html = __re_md.sub(r'[#*]+', '', html)
        # Wrap as paragraph by default
        return f'<p>{html}</p>'
    def _insert_html_with_tables(doc, html_text):
        """
        Render simple HTML content into the Word document.
        - Preserves tables (<table><tr><td>) as real Word tables with a template style.
        - Non-table HTML is converted to paragraphs using _strip_html_tags.
        """
        if not html_text:
            return
        # Helpers to style header rows with blue/green background and white bold text
        def _shade_cell(cell, fill_hex):
            try:
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fill_hex)
                tcPr.append(shd)
            except Exception:
                pass
        def _style_table_header_row(tbl):
            try:
                from docx.shared import RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH
                from docx.shared import Pt as _Pt
            except Exception:
                return
            if not getattr(tbl, "rows", None) or len(tbl.rows) == 0:
                return
            header_row = tbl.rows[0]
            header_colors = ['2E74B5', '70AD47']  # blue, green
            for idx, cell in enumerate(header_row.cells):
                _shade_cell(cell, header_colors[idx % 2])
                for paragraph in cell.paragraphs:
                    paragraph.alignment = _WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        paragraph.paragraph_format.space_before = _Pt(0)
                        paragraph.paragraph_format.space_after = _Pt(2)
                    except Exception:
                        pass
                    for run in paragraph.runs:
                        run.font.bold = True
                        try:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.name = PREFERRED_BODY_FONT_NAME
                            from docx.shared import Pt as _Pt
                            run.font.size = _Pt(11)
                        except Exception:
                            pass
        def _style_table_body(tbl):
            # Ensure body cells have consistent spacing and vertical alignment
            try:
                from docx.shared import Pt as _Pt
                from docx.enum.table import WD_ALIGN_VERTICAL as _WD_ALIGN_VERTICAL
            except Exception:
                return
            if not getattr(tbl, "rows", None) or len(tbl.rows) <= 1:
                return
            for r_idx, row in enumerate(tbl.rows):
                if r_idx == 0:
                    continue  # header handled separately
                for cell in row.cells:
                    try:
                        cell.vertical_alignment = _WD_ALIGN_VERTICAL.TOP
                    except Exception:
                        pass
                    for paragraph in cell.paragraphs:
                        try:
                            paragraph.paragraph_format.space_before = _Pt(0)
                            paragraph.paragraph_format.space_after = _Pt(2)
                        except Exception:
                            pass
                        try:
                            from docx.shared import Pt as _Pt
                            for run in paragraph.runs:
                                run.font.name = PREFERRED_BODY_FONT_NAME
                                run.font.size = _Pt(11)
                        except Exception:
                            pass
        # Helper: ensure visible grid borders on all tables and cells
        def _apply_table_borders(tbl, color="000000", size=8):
            try:
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                tblPr = tbl._tbl.tblPr
                # Remove existing borders to avoid duplicates
                for child in list(tblPr):
                    if child.tag == qn('w:tblBorders'):
                        tblPr.remove(child)
                borders = OxmlElement('w:tblBorders')
                for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                    ele = OxmlElement(f'w:{border_name}')
                    ele.set(qn('w:val'), 'single')
                    ele.set(qn('w:sz'), str(size))
                    ele.set(qn('w:color'), color)
                    borders.append(ele)
                tblPr.append(borders)
            except Exception:
                pass
        # Helper: apply pagination preferences so tables stay on same page if they fit
        def _apply_table_pagination_preferences(tbl):
            # Prevent splitting a row across pages and try to keep table together
            try:
                # Avoid splitting table rows across pages
                for row in getattr(tbl, "rows", []) or []:
                    try:
                        from docx.oxml import OxmlElement as _OX
                        trPr = row._tr.get_or_add_trPr()
                        cantSplit = _OX('w:cantSplit')
                        trPr.append(cantSplit)
                    except Exception:
                        pass
                # Keep paragraphs within the table together; allow break after the last row
                total_rows = len(getattr(tbl, "rows", []) or [])
                for r_idx, row in enumerate(getattr(tbl, "rows", []) or []):
                    for cell in getattr(row, "cells", []) or []:
                        for p in getattr(cell, "paragraphs", []) or []:
                            try:
                                pf = p.paragraph_format
                                pf.keep_together = True
                                # keep with next for all rows except last
                                pf.keep_with_next = (r_idx < total_rows - 1)
                            except Exception:
                                pass
            except Exception:
                pass
        # Render general delimited blocks (tabs, 2+ spaces, or key: value pairs)
        def _try_render_delimited_table_block(doc, block_text):
            """
            Detects:
            - Tab-delimited rows across multiple lines
            - Multi-space (2+) delimited rows across multiple lines
            - Repeated 'Key: Value' lines -> 2-column table
            Returns True if a table was rendered.
            """
            if not block_text or '\n' not in block_text:
                return False
            lines = [ln for ln in block_text.splitlines() if ln.strip()]
            if len(lines) < 2:
                return False
            import re as __re
            # 1) Key: Value pairs
            colon_pairs = [__re.split(r'\s*:\s+', ln, maxsplit=1) for ln in lines if ':' in ln]
            if len(colon_pairs) == len(lines) and all(len(p) == 2 for p in colon_pairs):
                # Build 2-col table, first row as header inferred if all keys look like titles
                tbl = doc.add_table(rows=(1 + len(colon_pairs)), cols=2)
                try:
                    if table_style_name:
                        tbl.style = table_style_name
                except Exception:
                    pass
                try:
                    from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
                    tbl.alignment = _WD_TABLE_ALIGNMENT.CENTER
                except Exception:
                    pass
                _apply_table_pagination_preferences(tbl)
                # Header
                tbl.cell(0, 0).text = 'Key'
                tbl.cell(0, 1).text = 'Value'
                # Body
                for r_off, (k, v) in enumerate(colon_pairs, start=1):
                    tbl.cell(r_off, 0).text = _strip_html_tags(k.strip())
                    tbl.cell(r_off, 1).text = _strip_html_tags(v.strip())
                _style_table_header_row(tbl)
                _style_table_body(tbl)
                _apply_table_borders(tbl)
                doc.add_paragraph()
                return True
            # 2) Tab-delimited rows
            if any('\t' in ln for ln in lines):
                rows = [ln.split('\t') for ln in lines]
                num_cols = max(len(r) for r in rows)
                if num_cols >= 2:
                    tbl = doc.add_table(rows=len(rows), cols=num_cols)
                    try:
                        if table_style_name:
                            tbl.style = table_style_name
                    except Exception:
                        pass
                    try:
                        from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
                        tbl.alignment = _WD_TABLE_ALIGNMENT.CENTER
                    except Exception:
                        pass
                    _apply_table_pagination_preferences(tbl)
                    for r_idx, row_vals in enumerate(rows):
                        for c_idx in range(num_cols):
                            val = row_vals[c_idx] if c_idx < len(row_vals) else ''
                            tbl.cell(r_idx, c_idx).text = _strip_html_tags(val.strip())
                    _style_table_header_row(tbl)
                    _style_table_body(tbl)
                    _apply_table_borders(tbl)
                    doc.add_paragraph()
                    return True
            # 3) Multi-space-delimited rows
            space_split_rows = [__re.split(r'\s{2,}', ln.strip()) for ln in lines]
            num_cols = max(len(r) for r in space_split_rows)
            if num_cols >= 2 and all(len(r) >= 2 for r in space_split_rows):
                tbl = doc.add_table(rows=len(space_split_rows), cols=num_cols)
                try:
                    if table_style_name:
                        tbl.style = table_style_name
                except Exception:
                    pass
                try:
                    from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
                    tbl.alignment = _WD_TABLE_ALIGNMENT.CENTER
                except Exception:
                    pass
                _apply_table_pagination_preferences(tbl)
                for r_idx, row_vals in enumerate(space_split_rows):
                    for c_idx in range(num_cols):
                        val = row_vals[c_idx] if c_idx < len(row_vals) else ''
                        tbl.cell(r_idx, c_idx).text = _strip_html_tags(val.strip())
                _style_table_header_row(tbl)
                _style_table_body(tbl)
                _apply_table_borders(tbl)
                doc.add_paragraph()
                return True
            return False
        # Attempt to detect and render Markdown/pipe tables from plain text blocks
        def _try_render_markdown_table_block(doc, block_text):
            """
            Detects Markdown-style tables:
            | H1 | H2 |
            | --- | --- |
            | v1 | v2 |
            Or pipe-delimited rows without explicit alignment row.
            Returns True if a table was rendered, False otherwise.
            """
            lines = [ln for ln in (block_text or '').splitlines() if ln.strip()]
            if len(lines) < 2:
                return False
            import re as __re
            def split_cells(line):
                # Keep inner pipes by splitting and trimming; drop leading/trailing empty caused by outer pipes
                parts = [c.strip() for c in line.strip().split('|')]
                if parts and parts[0] == '':
                    parts = parts[1:]
                if parts and parts[-1] == '':
                    parts = parts[:-1]
                return parts
            def is_pipe_line(ln):
                return '|' in ln and len(split_cells(ln)) >= 2
            def is_align_row(ln, cols):
                # e.g., | --- | :---: | --- |
                cells = split_cells(ln)
                if len(cells) != cols:
                    return False
                for c in cells:
                    if not __re.fullmatch(r':?-{3,}:?', c.replace(' ', '')):
                        return False
                return True
            if not is_pipe_line(lines[0]) or not is_pipe_line(lines[1]):
                return False
            header_cells = split_cells(lines[0])
            cols = len(header_cells)
            has_align = is_align_row(lines[1], cols)
            remaining = lines[2:] if has_align else lines[1:]
            if not remaining:
                return False
            # Assemble rows allowing multi-line cell continuations (lines without pipes)
            assembled_rows = []
            current_row = None
            for raw_ln in remaining:
                ln = raw_ln.rstrip()
                if is_pipe_line(ln):
                    # flush any current row
                    if current_row is not None:
                        if len(current_row) < cols:
                            current_row += [''] * (cols - len(current_row))
                        assembled_rows.append(current_row[:cols])
                    current_row = split_cells(ln)
                else:
                    if current_row is None:
                        return False
                    if not current_row:
                        current_row.append(ln.strip())
                    else:
                        current_row[-1] = (current_row[-1] + '\n' + ln.strip()).strip()
            if current_row is not None:
                if len(current_row) < cols:
                    current_row += [''] * (cols - len(current_row))
                assembled_rows.append(current_row[:cols])
            if not assembled_rows:
                return False
            # Build table
            tbl = doc.add_table(rows=(1 + len(assembled_rows)), cols=cols)
            try:
                if table_style_name:
                    tbl.style = table_style_name
            except Exception:
                pass
            try:
                from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
                tbl.alignment = _WD_TABLE_ALIGNMENT.CENTER
            except Exception:
                pass
            _apply_table_pagination_preferences(tbl)
            # Header
            for c_idx, text in enumerate(header_cells):
                tbl.cell(0, c_idx).text = _strip_html_tags(text)
            # Body
            for r_off, row_cells in enumerate(assembled_rows, start=1):
                # ensure length to cols
                if len(row_cells) < cols:
                    row_cells += [''] * (cols - len(row_cells))
                for c_idx, text in enumerate(row_cells[:cols]):
                    tbl.cell(r_off, c_idx).text = _strip_html_tags(text)
            _style_table_header_row(tbl)
            _style_table_body(tbl)
            _apply_table_borders(tbl)
            doc.add_paragraph()
            return True
        import re as _re
        table_splitter = _re.compile(r'(<table[\s\S]*?</table>)', _re.IGNORECASE)
        table_rows_re = _re.compile(r'<tr[\s\S]*?</tr>', _re.IGNORECASE)
        cell_contents_re = _re.compile(r'<t[dh][^>]*?>([\s\S]*?)</t[dh]>', _re.IGNORECASE)

        # Remove images, svg/canvas, and figure blocks (treat as "charts") before parsing
        try:
            html_text = _re.sub(r'<img\b[^>]*?/?>', '', html_text, flags=_re.IGNORECASE)
            html_text = _re.sub(r'<(svg|canvas)[\s\S]*?</\1\s*>', '', html_text, flags=_re.IGNORECASE)
            html_text = _re.sub(r'<figure[\s\S]*?</figure\s*>', '', html_text, flags=_re.IGNORECASE)
            html_text = _re.sub(r'```(?:mermaid|chart)[\s\S]*?```', '', html_text, flags=_re.IGNORECASE)
        except Exception:
            pass

        segments = table_splitter.split(html_text)
        for segment in segments:
            if not segment:
                continue
            if segment.lstrip().lower().startswith('<table'):
                rows = table_rows_re.findall(segment) or []
                num_cols = 0
                for rhtml in rows:
                    num_cols = max(num_cols, len(cell_contents_re.findall(rhtml)))
                if num_cols <= 0:
                    # Fallback: treat as paragraph text
                    text_version = _strip_html_tags(segment)
                    for p in [p for p in text_version.split('\n\n') if p.strip()]:
                        para = doc.add_paragraph(p.strip())
                        try:
                            from docx.shared import Pt as _Pt
                            for run in para.runs:
                                run.font.name = PREFERRED_BODY_FONT_NAME
                                run.font.size = _Pt(11)
                        except Exception:
                            pass
                    continue
                # Insert table inline; Word will move it to next page when it doesn't fit
                tbl = doc.add_table(rows=len(rows), cols=num_cols)
                try:
                    if table_style_name:
                        tbl.style = table_style_name
                except Exception:
                    pass
                try:
                    from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
                    tbl.alignment = _WD_TABLE_ALIGNMENT.CENTER
                except Exception:
                    pass
                _apply_table_pagination_preferences(tbl)
                for r_idx, rhtml in enumerate(rows):
                    cells = cell_contents_re.findall(rhtml)
                    for c_idx, cell_html in enumerate(cells):
                        if c_idx >= num_cols:
                            break
                        cell_text = _strip_html_tags(cell_html)
                        tbl.cell(r_idx, c_idx).text = cell_text
                # Style the header row with alternating blue/green colors
                _style_table_header_row(tbl)
                # Apply body cell formatting
                _style_table_body(tbl)
                # Ensure visible grid borders around all cells
                _apply_table_borders(tbl)
                # Add a small spacing after the table for readability
                doc.add_paragraph()
            else:
                text_version = _strip_html_tags(segment)
                # Try to identify markdown/pipe tables by blocks first
                blocks = [b for b in text_version.split('\n\n')]
                for block in blocks:
                    blk = block.strip()
                    if not blk:
                        continue
                    if _try_render_markdown_table_block(doc, blk) or _try_render_delimited_table_block(doc, blk):
                        continue
                    # Fallback: treat as plain paragraphs (split single newlines too)
                    paras = [p.strip() for p in blk.split('\n') if p.strip()]
                    for para_text in paras:
                        para = doc.add_paragraph(para_text)
                        try:
                            for run in para.runs:
                                run.font.name = 'Aptos Display'
                        except Exception:
                            pass

    for section in sections:
        if not isinstance(section, dict):
            continue
        title = (section.get('title') or 'Untitled Section').strip() or 'Untitled Section'
        content_blocks = section.get('content') if isinstance(section.get('content'), list) else []
        raw_html = section.get('raw_content') or ''
        section_id = (section.get('id') or '').strip()
        
        # Fetch any attachments for this section up-front so sections with only images are still rendered
        attachments = []
        if bid_id and section_id:
            try:
                attachments = _get_section_files_for_bid(bid_id, section_id, limit=20) or []
            except Exception:
                attachments = []
        
        # Skip sections that truly have no data at all (no text/html and no attachments)
        if not content_blocks and not raw_html and not attachments:
            continue

        # Ensure every new section begins on a new page
        if first_section_added:
            document.add_page_break()
        first_section_added = True

        # Add section heading
        heading_para = document.add_heading(title, level=1)
        try:
            from docx.shared import Pt as _Pt
            is_transmittal = ('transmittal' in (title or '').lower())
            for run in heading_para.runs:
                run.font.name = PREFERRED_HEADING_FONT_NAME
                run.font.bold = True
                run.font.size = _Pt(16)
                # For Letter of Transmittal, hide the heading in the body but keep it as Heading 1
                # so it still participates in TOC and header STYLEREF.
                if is_transmittal:
                    try:
                        run.font.hidden = True
                    except Exception:
                        pass
            if is_transmittal:
                try:
                    heading_para.paragraph_format.space_after = _Pt(0)
                    heading_para.paragraph_format.space_before = _Pt(0)
                except Exception:
                    pass
        except Exception:
            pass
        html_fragments.append(f'<h2>{escape(title)}</h2>')

        # Process content - prioritize raw_html if available (for HTML tables, etc.)
        if raw_html:
            # Render HTML preserving tables; append raw_html to preview
            _insert_html_with_tables(document, raw_html)
            try:
                # Apply minimal markdown cleanup inside raw_html for preview only
                import re as __re_md
                rh = raw_html
                # Convert bold markers inside raw HTML text content crudely
                rh = __re_md.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', rh)
                # Convert '### ' headings at the start of lines to <h3>
                rh = __re_md.sub(r'(^|\n)\s*#{3,}\s*(.+?)\s*(?=\n|$)', lambda m: f"{m.group(1)}<h3>{escape(m.group(2))}</h3>", rh)
                # Remove remaining stray markers
                rh = __re_md.sub(r'[#*]+', '', rh)
                html_fragments.append(rh)
            except Exception:
                html_fragments.append(raw_html)
        else:
            # Process regular content blocks
            fragment_html = ''
            for paragraph in content_blocks:
                if paragraph and paragraph.strip():
                    para_text = paragraph.strip()
                    # If the paragraph looks like a table block (pipe, tabs, multi-space, or key:value across lines), render it as a real table
                    if '\n' in para_text:
                        try:
                            import re as __re2
                            looks_table = (
                                ('|' in para_text) or
                                ('\t' in para_text) or
                                (__re2.search(r'\s{2,}', para_text) is not None) or
                                (__re2.search(r'^[^:\n]+:\s*.+$', para_text, flags=__re2.M) is not None)
                            )
                        except Exception:
                            looks_table = ('|' in para_text) or ('\t' in para_text)
                    else:
                        looks_table = False
                    if looks_table:
                        _insert_html_with_tables(document, para_text)
                        # Keep preview minimal; embed as preformatted text block
                        fragment_html += f'<pre>{escape(para_text)}</pre>'
                    else:
                        # Render with Markdown interpretation
                        _render_markdown_to_docx_paragraph(document, para_text)
                        fragment_html += _plaintext_markdown_to_html(para_text)
            
            if fragment_html:
                html_fragments.append(fragment_html)

        # Append any section attachments (PDFs/images) directly into the document
        try:
            if attachments:
                # Add a small label in preview
                html_fragments.append('<div><em>Attachments included:</em></div>')
            for att in attachments:
                    fp = (att.get('file_path') or '').strip()
                    name = (att.get('filename') or os.path.basename(fp) or 'Attachment').strip()
                    if not fp or not os.path.exists(fp):
                        continue
                    try:
                        from docx.shared import Inches
                    except Exception:
                        Inches = None  # type: ignore[assignment]
                    # Compute usable width if possible
                    usable_width = None
                    try:
                        sec = document.sections[-1]
                        usable_width = sec.page_width - sec.left_margin - sec.right_margin
                    except Exception:
                        pass
                    try:
                        is_pdf = fp.lower().endswith('.pdf') or (att.get('file_type') or '').lower() == 'pdf'
                        if is_pdf and globals().get('_FITZ_AVAILABLE', False):
                            # Render each PDF page as an image and insert
                            import tempfile
                            import shutil as _shutil
                            doc_mod = globals().get('_pymupdf')
                            with doc_mod.open(fp) as _pdf:
                                for i in range(min(_pdf.page_count or 0, 20)):
                                    page = _pdf.load_page(i)
                                    zoom = 2.0
                                    mat = doc_mod.Matrix(zoom, zoom)
                                    pix = page.get_pixmap(matrix=mat)
                                    tmpdir = tempfile.gettempdir()
                                    tmp_png = os.path.join(tmpdir, f"att_{os.path.basename(fp)}_{i}.png")
                                    pix.save(tmp_png)
                                    if Inches:
                                        if usable_width:
                                            document.add_picture(tmp_png, width=usable_width)  # type: ignore[arg-type]
                                        else:
                                            document.add_picture(tmp_png, width=Inches(6.5))
                                    else:
                                        # Fallback: insert a paragraph with the file name if images unsupported
                                        document.add_paragraph(f"[Attachment page] {name} - page {i+1}")
                                    html_fragments.append(f'<div class="text-xs text-gray-500">Attached PDF page: {name} (page {i+1})</div>')
                                    try:
                                        _shutil.which('true')  # no-op to keep import
                                    except Exception:
                                        pass
                        else:
                            # Treat as image
                            if Inches:
                                if usable_width:
                                    document.add_picture(fp, width=usable_width)  # type: ignore[arg-type]
                                else:
                                    document.add_picture(fp, width=Inches(6.5))
                            else:
                                document.add_paragraph(f"[Attachment] {name}")
                            html_fragments.append(f'<div class="text-xs text-gray-500">Attached image: {name}</div>')
                    except Exception:
                        # Continue on errors for individual attachments
                        continue
            if attachments:
                # Add a page break after attachments to separate from next section
                try:
                    document.add_page_break()
                except Exception:
                    pass
        except Exception:
            # Ignore attachment embedding failures silently to not break main flow
            pass
        # No extra paragraph needed here; page breaks handle separation

    if not html_fragments:
        return jsonify({'error': 'empty_sections', 'message': 'No compiled content available. Generate or enter section content first.'}), 400

    # Generate filename with project title if available
    os.makedirs('uploads/proposals', exist_ok=True)
    timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
    if project_title:
        # Sanitize project title for filename
        safe_title = ''.join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in project_title[:50])
        safe_title = safe_title.replace(' ', '_')
        filename = f"Proposal_{safe_title}_{timestamp}.docx"
    else:
        filename = f"Proposal_{timestamp}.docx"

    filepath = os.path.join('uploads/proposals', filename)

    try:
        # Ensure headers/footers from template apply to all pages/sections
        try:
            # For IKIO template, many templates place logo only in First Page header/footer.
            # Copy first-page header/footer into primary so it repeats on all pages,
            # then link subsequent sections to previous.
            if template_used == 'ikio' and getattr(document, "sections", None):
                from copy import deepcopy
                first_sec = document.sections[0]
                # Helper to clone header/footer XML from src to dst within same section
                def _clone_part_xml(src_part, dst_part):
                    try:
                        dst_el = dst_part._element  # w:hdr or w:ftr
                        # remove existing children
                        for child in list(dst_el):
                            dst_el.remove(child)
                        # append deep-copied children from source
                        for child in list(src_part._element):
                            dst_el.append(deepcopy(child))
                    except Exception:
                        pass
                # If first-page header has content, clone to primary header
                try:
                    if getattr(first_sec, "first_page_header", None):
                        if len(list(first_sec.first_page_header._element)) > 0:
                            _clone_part_xml(first_sec.first_page_header, first_sec.header)
                except Exception:
                    pass
                # If first-page footer has content, clone to primary footer
                try:
                    if getattr(first_sec, "first_page_footer", None):
                        if len(list(first_sec.first_page_footer._element)) > 0:
                            _clone_part_xml(first_sec.first_page_footer, first_sec.footer)
                except Exception:
                    pass
                # Disable special header/footer modes across ALL sections to ensure consistency
                try:
                    for _sec in document.sections:
                        try:
                            _sec.different_first_page_header_footer = False
                        except Exception:
                            pass
                        try:
                            _sec.odd_and_even_pages_header_footer = False
                        except Exception:
                            pass
                except Exception:
                    pass
            # Link headers/footers after all content is added
            for _s_idx, _section in enumerate(document.sections):
                if _s_idx > 0:
                    try:
                        _section.header.is_linked_to_previous = True
                        _section.footer.is_linked_to_previous = True
                    except Exception:
                        pass
            # Continuous page numbering: start at 1 on first section, continue thereafter
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            for _s_idx, _section in enumerate(document.sections):
                _sectPr = _section._sectPr
                _pgNumType = _sectPr.find(qn('w:pgNumType'))
                if _pgNumType is None:
                    _pgNumType = OxmlElement('w:pgNumType')
                    _sectPr.append(_pgNumType)
                if _s_idx == 0:
                    _pgNumType.set(qn('w:start'), "1")
                else:
                    try:
                        _pgNumType.attrib.pop(qn('w:start'), None)
                    except Exception:
                        pass
        except Exception:
            # If linking or numbering XML manipulation fails, proceed without blocking save
            pass
        # Ask Word to update fields (including TOC) on document open
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            # Get or create settings part and set updateFields=true
            settings_part = getattr(document._part, "get_or_add_settings", None)
            settings_el = settings_part() if callable(settings_part) else None
            if settings_el is None:
                # Fallback: try to reach settings element via package (best-effort)
                settings_el = getattr(document._part, "_settings", None)
            if settings_el is not None:
                upd = OxmlElement('w:updateFields')
                upd.set(qn('w:val'), 'true')
                try:
                    # Remove existing to avoid duplicates
                    for child in list(settings_el):
                        if child.tag == qn('w:updateFields'):
                            settings_el.remove(child)
                except Exception:
                    pass
                settings_el.append(upd)
        except Exception:
            pass
        document.save(filepath)
        app.logger.info(f'Successfully compiled proposal with {len(sections)} sections to {filename}')
    except Exception as error:
        app.logger.exception('Failed to save compiled proposal: %s', error)
        return jsonify({'error': 'save_failed', 'message': 'Could not save the compiled proposal. Please try again.'}), 500

    preview_html = ''.join(html_fragments)
    return jsonify({
        'message': f'Proposal compiled successfully with {len(sections)} section(s). Review the preview before downloading.',
        'filename': filename,
        'download_url': url_for('download_proposal', filename=filename),
        'html_preview': preview_html,
        'section_count': len(sections)
    })

def extract_template_header_footer(template_path):
    """
    Safely read a DOCX template and extract visible header/footer text,
    including text that may be inside tables for first/even/primary parts.
    Returns a dict: { 'header_text': str, 'footer_text': str }.
    """
    try:
        from docx import Document
    except Exception:
        return {"header_text": "", "footer_text": ""}
    try:
        doc = Document(template_path)
    except Exception:
        return {"header_text": "", "footer_text": ""}

    try:
        section = doc.sections[0]
    except Exception:
        return {"header_text": "", "footer_text": ""}

    def _collect_text(part):
        lines = []
        if not part:
            return lines
        try:
            for p in getattr(part, 'paragraphs', []) or []:
                t = (p.text or '').strip()
                if t:
                    lines.append(t)
        except Exception:
            pass
        try:
            for tbl in getattr(part, 'tables', []) or []:
                for row in getattr(tbl, 'rows', []) or []:
                    for cell in getattr(row, 'cells', []) or []:
                        for p in getattr(cell, 'paragraphs', []) or []:
                            t = (p.text or '').strip()
                            if t:
                                lines.append(t)
        except Exception:
            pass
        return lines

    # Gather header parts: primary, first-page, even-page
    header_parts = []
    try:
        header_parts.append(getattr(section, 'header', None))
        header_parts.append(getattr(section, 'first_page_header', None))
        header_parts.append(getattr(section, 'even_page_header', None))
    except Exception:
        pass
    # Gather footer parts
    footer_parts = []
    try:
        footer_parts.append(getattr(section, 'footer', None))
        footer_parts.append(getattr(section, 'first_page_footer', None))
        footer_parts.append(getattr(section, 'even_page_footer', None))
    except Exception:
        pass

    header_lines = []
    footer_lines = []
    seen = set()
    for part in header_parts:
        for ln in _collect_text(part):
            if ln not in seen:
                seen.add(ln)
                header_lines.append(ln)
    seen.clear()
    for part in footer_parts:
        for ln in _collect_text(part):
            if ln not in seen:
                seen.add(ln)
                footer_lines.append(ln)

    header_text = '\n'.join(header_lines).strip()
    footer_text = '\n'.join(footer_lines).strip()
    return {"header_text": header_text, "footer_text": footer_text}

def choose_company_template(company_display):
    """
    Returns an absolute path to the preferred template for the given company.
    Search order is brand-aware with sensible fallbacks, and supports common
    filename variants and locations.
    """
    import os as _os
    base_dir = app.root_path if hasattr(app, "root_path") else _os.getcwd()

    # Filename variants per brand (to handle '&' and path differences)
    variants = {
        "ikio": [
            # Primary IKIO template(s)
            "Proposal ikio.docx",
            _os.path.join("formatt", "uploads", "templates", "Proposal ikio.docx"),
        
        ],
        "metco": [
            "H & F.docx",
            "H&F.docx",
            _os.path.join("formatt", "uploads", "templates", "H & F.docx"),
            _os.path.join("formatt", "uploads", "templates", "H&F.docx"),
        ],
        "sunsprint": [
            "Proposal For Sunsprint.docx",
            _os.path.join("formatt", "uploads", "templates", "Proposal For Sunsprint.docx"),
        ],
    }

    company_lower = (company_display or "").lower()
    if "ikio" in company_lower:
        priority = ["ikio", "sunsprint", "metco"]
    elif "metco" in company_lower:
        priority = ["metco", "sunsprint", "ikio"]
    elif "sunsprint" in company_lower:
        priority = ["sunsprint", "metco", "ikio"]
    else:
        priority = ["sunsprint", "metco", "ikio"]

    # Build ordered absolute candidate list
    candidates = []
    for brand in priority:
        for rel in variants[brand]:
            candidates.append(_os.path.join(base_dir, rel))

    # Return first existing path
    for cand in candidates:
        try:
            if _os.path.exists(cand):
                return cand
        except Exception:
            continue
    return None

@app.route('/api/validate-requirement-attachment', methods=['POST'])
@login_required
def validate_requirement_attachment():
    """Validate that an uploaded document matches the required form type."""
    if 'file' not in request.files:
        return jsonify({'error': 'missing_file', 'message': 'No file provided.'}), 400
    
    file = request.files['file']
    requirement_text = request.form.get('requirement_text', '').strip()
    
    if not file or not file.filename:
        return jsonify({'error': 'missing_file', 'message': 'No file provided.'}), 400
    
    if not requirement_text:
        return jsonify({'error': 'missing_requirement', 'message': 'Requirement text is required.'}), 400
    
    # Validate file type
    filename = file.filename.lower()
    if not (filename.endswith('.pdf') or filename.endswith('.jpg') or filename.endswith('.jpeg')):
        return jsonify({'error': 'invalid_format', 'message': 'Only PDF and JPG files are allowed.'}), 400
    
    try:
        # Save file temporarily
        import tempfile
        import os
        from werkzeug.utils import secure_filename
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp_file:
            file.save(tmp_file.name)
            tmp_path = tmp_file.name
        
        try:
            # Extract text from document
            extracted_text = ''
            
            if filename.endswith('.pdf'):
                # Extract text from PDF
                if _FITZ_AVAILABLE:
                    with _pymupdf.open(tmp_path) as doc:
                        for page_num in range(min(3, doc.page_count)):  # First 3 pages
                            page = doc.load_page(page_num)
                            extracted_text += page.get_text("text") + "\n\n"
                else:
                    # Fallback to PyPDF2
                    if _PdfReader is not None:
                        with open(tmp_path, "rb") as fh:
                            reader = _PdfReader(fh)
                            for page_num in range(min(3, len(reader.pages))):
                                extracted_text += reader.pages[page_num].extract_text() + "\n\n"
                    else:
                        return jsonify({'error': 'pdf_library_missing', 'message': 'PDF processing library not available.'}), 500
            
            elif filename.endswith('.jpg') or filename.endswith('.jpeg'):
                # For images, we'll use OpenAI Vision API or return limited validation
                # For now, we'll use AI to analyze the image directly
                extracted_text = '[IMAGE_FILE]'  # Placeholder for image files
            
            # Use AI to validate document content
            api_key = (app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or os.getenv('OPENROUTER_API_KEY') or os.getenv('TOGETHER_API_KEY') or os.getenv('FIREWORKS_API_KEY') or '').strip()
            if not api_key:
                return jsonify({'error': 'api_unavailable', 'message': 'OpenAI API key is not configured.'}), 500
            
            base_url = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or 'https://api.openai.com/v1').strip()
            model = (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'gpt-4o-mini').strip()
            
            # For images, use vision model
            if filename.endswith('.jpg') or filename.endswith('.jpeg'):
                # Read image as base64
                import base64
                with open(tmp_path, 'rb') as img_file:
                    image_data = base64.b64encode(img_file.read()).decode('utf-8')
                
                endpoint = base_url.rstrip('/')
                if not endpoint.endswith('/chat/completions'):
                    endpoint = f"{endpoint}/chat/completions"
                
                headers = {
                    'Authorization': f'Bearer {api_key}',
                    'Content-Type': 'application/json',
                }
                
                system_prompt = (
                    "You are a document validation expert. Your task is to analyze uploaded documents "
                    "and determine if they match the required form type specified in the requirement text. "
                    "Return only 'YES' if the document matches the requirement, or 'NO' with a brief reason if it doesn't."
                )
                
                user_prompt = f"""
Analyze the uploaded image and determine if it matches the required document type.

REQUIREMENT: {requirement_text}

The requirement may specify forms like:
- W-9 Form (Request for Taxpayer Identification Number and Certification)
- W-8 Form (Certificate of Foreign Status)
- Insurance Certificate
- Bid Bond
- Performance Bond
- License/Certification
- Company Profile
- References
- Any other specific form or document

Examine the image carefully and check:
1. Does the document title/header match the required form name?
2. Are there identifying markers (form numbers, official headers) that match?
3. Is this the correct document type?

Respond with:
- "YES" if the document matches the requirement
- "NO: [brief reason]" if it doesn't match

Your response should be only "YES" or "NO: [reason]".
"""
                
                body = {
                    'model': 'gpt-4o',  # Use vision-capable model
                    'messages': [
                        {'role': 'system', 'content': system_prompt},
                        {
                            'role': 'user',
                            'content': [
                                {'type': 'text', 'text': user_prompt},
                                {
                                    'type': 'image_url',
                                    'image_url': {
                                        'url': f'data:image/jpeg;base64,{image_data}'
                                    }
                                }
                            ]
                        }
                    ],
                    'temperature': 0.1,
                    'max_tokens': 150
                }
            else:
                # For PDFs, use text-based validation
                endpoint = base_url.rstrip('/')
                if not endpoint.endswith('/chat/completions'):
                    endpoint = f"{endpoint}/chat/completions"
                
                headers = {
                    'Authorization': f'Bearer {api_key}',
                    'Content-Type': 'application/json',
                }
                
                system_prompt = (
                    "You are a document validation expert. Your task is to analyze extracted text from documents "
                    "and determine if they match the required form type specified in the requirement text. "
                    "Return only 'YES' if the document matches the requirement, or 'NO' with a brief reason if it doesn't."
                )
                
                user_prompt = f"""
Analyze the extracted text from the uploaded document and determine if it matches the required document type.

REQUIREMENT: {requirement_text}

The requirement may specify forms like:
- W-9 Form (Request for Taxpayer Identification Number and Certification)
- W-8 Form (Certificate of Foreign Status)
- Insurance Certificate
- Bid Bond
- Performance Bond
- License/Certification
- Company Profile
- References
- Any other specific form or document

EXTRACTED TEXT FROM DOCUMENT:
{extracted_text[:3000] if extracted_text else '[No text extracted]'}

Examine the text carefully and check:
1. Does the document title/header match the required form name?
2. Are there identifying markers (form numbers, official headers, form names) that match?
3. Is this the correct document type?

Respond with:
- "YES" if the document matches the requirement
- "NO: [brief reason]" if it doesn't match

Your response should be only "YES" or "NO: [reason]".
"""
                
                body = {
                    'model': model,
                    'messages': [
                        {'role': 'system', 'content': system_prompt},
                        {'role': 'user', 'content': user_prompt}
                    ],
                    'temperature': 0.1,
                    'max_tokens': 150
                }
            
            # Call OpenAI API
            import requests
            response = requests.post(endpoint, headers=headers, json=body, timeout=30)
            
            if response.status_code >= 400:
                data = response.json() if response.content else {}
                error_msg = data.get('error', {}).get('message', 'Validation failed') if isinstance(data, dict) else 'Validation failed'
                return jsonify({'error': 'validation_failed', 'message': error_msg}), 502
            
            data = response.json()
            choices = data.get('choices', [])
            if not choices:
                return jsonify({'error': 'validation_failed', 'message': 'No response from validation service.'}), 502
            
            validation_result = choices[0].get('message', {}).get('content', '').strip().upper()
            is_valid = validation_result.startswith('YES')
            
            return jsonify({
                'is_valid': is_valid,
                'validation_message': validation_result,
                'requirement': requirement_text
            })
        
        finally:
            # Clean up temporary file
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
    
    except Exception as e:
        app.logger.exception("Error validating attachment: %s", e)
        return jsonify({'error': 'validation_error', 'message': f'Validation failed: {str(e)}'}), 500

@app.route('/proposals-making/download/<filename>')
@login_required
def download_proposal(filename):
    # Allow all logged-in users to download proposals (removed admin-only restriction)
    from werkzeug.utils import secure_filename
    from flask import send_from_directory
    safe_filename = secure_filename(filename)
    proposals_dir = 'uploads/proposals'
    return send_from_directory(proposals_dir, safe_filename, as_attachment=True)

# --- API & Real-time Logic ---
@app.route('/api/update_stage/<int:bid_id>', methods=['POST'])
@login_required
def update_stage(bid_id):
    try:
        data = request.get_json()
        new_stage = (data.get('stage') or '').lower()
        
        # Validate stage
        allowed = {'analyzer', 'business', 'design', 'operations', 'engineer', 'handover'}
        if new_stage not in allowed:
            return jsonify({'error': 'invalid stage'}), 400
        
        cur = mysql.connection.cursor(DictCursor)
        
        # Get bid information from go_bids
        cur.execute("SELECT * FROM go_bids WHERE g_id=%s", (bid_id,))
        bid = cur.fetchone()
        
        if not bid:
            cur.close()
            return jsonify({'error': 'Bid not found or access denied'}), 404

        # Get old stage for logging
        old_stage = (bid.get('state') or 'analyzer').lower()
        
        # Update go_bids state
        cur.execute("UPDATE go_bids SET state=%s WHERE g_id=%s", (new_stage, bid_id))
        
        # Upsert assignment so the next team still sees it
        cur.execute("SELECT a_id FROM bid_assign WHERE g_id=%s", (bid_id,))
        row = cur.fetchone()
        if row:
            cur.execute("UPDATE bid_assign SET depart=%s, state=%s, status='pending' WHERE g_id=%s",
                        (new_stage, new_stage, bid_id))
        else:
            try:
                cur.execute("""
                    INSERT INTO bid_assign (g_id, b_name, due_date, state, scope, type, company, depart,
                                           person_name, assignee_email, status, value, revenue)
                    SELECT g_id, b_name, due_date, state, scope, type, company, %s, '', '', 'pending',
                           COALESCE(scoring, 0), COALESCE(revenue, 0)
                    FROM go_bids WHERE g_id=%s
                """, (new_stage, bid_id))
            except Exception as e_ins:
                if "Unknown column 'revenue'" in str(e_ins):
                    # Retry without 'revenue' column for legacy bid_assign schema
                    cur.execute("""
                        INSERT INTO bid_assign (g_id, b_name, due_date, state, scope, type, company, depart,
                                               person_name, assignee_email, status, value)
                        SELECT g_id, b_name, due_date, state, scope, type, company, %s, '', '', 'pending',
                               COALESCE(scoring, 0)
                        FROM go_bids WHERE g_id=%s
                    """, (new_stage, bid_id))
                else:
                    raise
        
        # Derive dynamic summary line
        from_txt = LABELS.get(old_stage, '')
        to_txt = LABELS.get(new_stage, '')
        summary_line = f"Updated by {from_txt} to {to_txt}"
        
        # Commit the transaction
        mysql.connection.commit()
        
        # Log the stage change
        log_write('stage_change', f"{bid.get('b_name')} | {old_stage} → {new_stage}")
        
        # Calculate dynamic progress and status texts for new stage
        pct = pct_for(new_stage)
        proj_status = 'completed' if new_stage == 'handover' else 'ongoing'
        
        # Recompute cards and analyzer stats after update
        # Summary cards across the three companies only
        cur2 = mysql.connection.cursor(DictCursor)
        cur2.execute("SELECT id FROM companies WHERE name IN ('Ikio','Metco','Sunsprint')")
        target_company_ids = [row['id'] for row in cur2.fetchall()]
        if not target_company_ids:
            target_company_ids = [-1]
        in_clause = ','.join(['%s'] * len(target_company_ids))

        # Compute totals from go_bids instead of bids
        cur2.execute("SELECT name FROM companies WHERE name IN ('Ikio','Metco','Sunsprint')")
        target_company_names = [row['name'] for row in cur2.fetchall()] or ['__none__']
        in_clause_names = ','.join(['%s'] * len(target_company_names))
        cur2.execute(f"SELECT COUNT(*) AS total_bids FROM go_bids WHERE company IN ({in_clause_names})", target_company_names)
        total_bids = cur2.fetchone()['total_bids']
        cur2.execute(f"SELECT COUNT(*) AS live_bids FROM go_bids WHERE COALESCE(state,'analyzer') IN ('business','design','operations','engineer') AND company IN ({in_clause_names})", target_company_names)
        live_bids = cur2.fetchone()['live_bids']
        cur2.execute(f"SELECT COUNT(*) AS bids_won FROM go_bids WHERE decision='WON' AND company IN ({in_clause_names})", target_company_names)
        bids_won = cur2.fetchone()['bids_won']
        # Total projects across target companies (fallback to all projects if none linked)
        cur2.execute(f"SELECT COUNT(*) AS projects_linked FROM projects WHERE company_id IN ({in_clause})", target_company_ids)
        projects_linked = cur2.fetchone()['projects_linked']
        if projects_linked > 0:
            cur2.execute(f"SELECT COUNT(*) AS projects_total FROM projects WHERE company_id IN ({in_clause})", target_company_ids)
            projects_total = cur2.fetchone()['projects_total']
        else:
            cur2.execute("SELECT COUNT(*) AS projects_total FROM projects")
            projects_total = cur2.fetchone()['projects_total']

        # Analyzer stats from bid_incoming table
        cur2.execute("SELECT COUNT(*) AS total_bids FROM bid_incoming")
        total_bids_analyzer = cur2.fetchone()['total_bids']
        
        cur2.execute("SELECT COUNT(*) AS bids_go FROM bid_incoming WHERE decision = 'GO'")
        bids_go_analyzer = cur2.fetchone()['bids_go']
        
        cur2.execute("SELECT COUNT(*) AS bids_no_go FROM bid_incoming WHERE decision = 'NO-GO'")
        bids_no_go_analyzer = cur2.fetchone()['bids_no_go']
        
        cur2.execute("SELECT COUNT(*) AS bids_submitted FROM bid_incoming WHERE state IN ('submitted', 'under_review')")
        bids_submitted_analyzer = cur2.fetchone()['bids_submitted']
        
        cur2.execute("SELECT COUNT(*) AS bids_won FROM bid_incoming WHERE decision = 'WON'")
        bids_won_analyzer = cur2.fetchone()['bids_won']
        
        cur2.execute("SELECT COUNT(*) AS bids_lost FROM bid_incoming WHERE decision = 'LOST'")
        bids_lost_analyzer = cur2.fetchone()['bids_lost']

        bid_stats = {
            'total_bids': total_bids_analyzer,
            'bids_go': bids_go_analyzer,
            'bids_no_go': bids_no_go_analyzer,
            'bids_submitted': bids_submitted_analyzer,
            'bids_won': bids_won_analyzer,
            'bids_lost': bids_lost_analyzer
        }

        socketio.emit('master_update', {
            'bid': {
                'id': bid_id,
                'name': bid.get('b_name'),
                'current_stage': new_stage,
                'user_email': getattr(current_user, 'email', '')
            },
            'summary': {
                'work_progress_pct': pct,
                'project_status': proj_status,
                'work_status': summary_line
            },
            'cards': {
                'total_bids': projects_total,
                'live_bids': 0,
                'bids_won': 0,
                'projects_completed': 0
            },
            'bid_stats': bid_stats
        })
        
        cur.close()
        cur2.close()
        return jsonify({'success': f'Bid {bid_id} updated to {new_stage}'})
    
    except Exception as e:
        mysql.connection.rollback()
        if 'cur' in locals():
            cur.close()
        if 'cur2' in locals():
            cur2.close()
        return jsonify({'error': f'Error updating stage: {str(e)}'}), 500
# --- Main execution ---
def _ensure_tables_exist():
    """Create tables if they don't exist"""
    try:
        cur = mysql.connection.cursor()
        
        # Create users table (with lightweight self-healing for common MySQL errors)
        users_create_sql = """
            CREATE TABLE IF NOT EXISTS users (
                id INT AUTO_INCREMENT PRIMARY KEY,
                email VARCHAR(100) UNIQUE NOT NULL,
                password VARCHAR(100) NOT NULL,
                is_admin BOOLEAN DEFAULT FALSE,
                role VARCHAR(50) DEFAULT 'member'
            )
        """
        cur.execute(users_create_sql)
        # Verify users table is readable; repair if it's missing or has orphaned metadata (errors 1146, 1932)
        try:
            cur.execute("SELECT 1 FROM users LIMIT 1")
        except Exception as e:
            try:
                err_no = e.args[0] if hasattr(e, 'args') and e.args else None
            except Exception:
                err_no = None
            if err_no in (1146, 1932):
                try:
                    cur.execute("DROP TABLE IF EXISTS users")
                    cur.execute(users_create_sql)
                except Exception:
                    # Let the outer handler report any persistent issues
                    pass
            else:
                # Unexpected error; let the outer handler surface it
                raise
        
        # Create bids table (with lightweight self-healing for common MySQL errors)
        bids_create_sql = """
            CREATE TABLE IF NOT EXISTS bids (
                id INT AUTO_INCREMENT PRIMARY KEY,
                name VARCHAR(100) NOT NULL,
                current_stage VARCHAR(50) DEFAULT 'analyzer',
                user_id INT,
                company_id INT,
                FOREIGN KEY (user_id) REFERENCES users(id)
            )
        """
        cur.execute(bids_create_sql)
        # Verify bids table is readable; repair if it's missing or has orphaned metadata (errors 1146, 1932)
        try:
            cur.execute("SELECT 1 FROM bids LIMIT 1")
        except Exception as e:
            try:
                err_no = e.args[0] if hasattr(e, 'args') and e.args else None
            except Exception:
                err_no = None
            if err_no in (1146, 1932):
                try:
                    cur.execute("DROP TABLE IF EXISTS bids")
                    cur.execute(bids_create_sql)
                except Exception:
                    # Let the outer handler report any persistent issues
                    pass
            else:
                # Unexpected error; let the outer handler surface it
                raise
        
        # Create bid_incoming table (with lightweight self-healing)
        bid_incoming_create_sql = """
            CREATE TABLE IF NOT EXISTS bid_incoming (
                id INT AUTO_INCREMENT PRIMARY KEY,
                b_name VARCHAR(100),
               
                due_date DATE NOT NULL DEFAULT CURRENT_TIMESTAMP,
                state VARCHAR(100),
                scope TEXT,
                type VARCHAR(100),
                scoring INT,
                comp_name VARCHAR(100),
                decision VARCHAR(100),
                summary TEXT
            )
        """
        cur.execute(bid_incoming_create_sql)
        try:
            cur.execute("SELECT 1 FROM bid_incoming LIMIT 1")
        except Exception as e:
            try:
                err_no = e.args[0] if hasattr(e, 'args') and e.args else None
            except Exception:
                err_no = None
            if err_no in (1146, 1932):
                try:
                    cur.execute("DROP TABLE IF EXISTS bid_incoming")
                    cur.execute(bid_incoming_create_sql)
                except Exception:
                    pass
            else:
                raise
        
        # Create go_bids table (with lightweight self-healing)
        go_bids_create_sql = """
            CREATE TABLE IF NOT EXISTS go_bids (
                g_id INT AUTO_INCREMENT PRIMARY KEY,
                id INT,
                b_name VARCHAR(100),
               
                due_date DATE NOT NULL DEFAULT CURRENT_TIMESTAMP,
                state VARCHAR(100),
                scope TEXT,
                type VARCHAR(100),
                scoring INT,
                company TEXT,
                decision TEXT,
                summary TEXT,
                revenue DECIMAL(15,2) DEFAULT 0.00
            )
        """
        cur.execute(go_bids_create_sql)
        try:
            cur.execute("SELECT 1 FROM go_bids LIMIT 1")
        except Exception as e:
            try:
                err_no = e.args[0] if hasattr(e, 'args') and e.args else None
            except Exception:
                err_no = None
            if err_no in (1146, 1932):
                try:
                    cur.execute("DROP TABLE IF EXISTS go_bids")
                    cur.execute(go_bids_create_sql)
                except Exception:
                    pass
            else:
                raise
        
        # Ensure summary and scope columns exist and are TEXT type for detailed content
        # Check and alter go_bids table
        try:
            cur.execute("SHOW COLUMNS FROM go_bids LIKE 'summary'")
            summary_col = cur.fetchone()
            if not summary_col:
                cur.execute("ALTER TABLE go_bids ADD COLUMN summary TEXT")
            else:
                # Check if it's not TEXT type and alter it
                col_type = summary_col[1] if isinstance(summary_col, tuple) else summary_col.get('Type', '')
                if col_type and 'text' not in str(col_type).lower():
                    cur.execute("ALTER TABLE go_bids MODIFY COLUMN summary TEXT")
        except Exception:
            pass
        
        try:
            cur.execute("SHOW COLUMNS FROM go_bids LIKE 'scope'")
            scope_col = cur.fetchone()
            if not scope_col:
                cur.execute("ALTER TABLE go_bids ADD COLUMN scope TEXT")
            else:
                # Check if it's not TEXT type and alter it
                col_type = scope_col[1] if isinstance(scope_col, tuple) else scope_col.get('Type', '')
                if col_type and 'text' not in str(col_type).lower():
                    cur.execute("ALTER TABLE go_bids MODIFY COLUMN scope TEXT")
        except Exception:
            pass
        
        # Ensure summary and scope columns exist in bid_incoming table
        try:
            cur.execute("SHOW COLUMNS FROM bid_incoming LIKE 'summary'")
            summary_col = cur.fetchone()
            if not summary_col:
                cur.execute("ALTER TABLE bid_incoming ADD COLUMN summary TEXT")
            else:
                # Check if it's not TEXT type and alter it
                col_type = summary_col[1] if isinstance(summary_col, tuple) else summary_col.get('Type', '')
                if col_type and 'text' not in str(col_type).lower():
                    cur.execute("ALTER TABLE bid_incoming MODIFY COLUMN summary TEXT")
        except Exception:
            pass
        
        try:
            cur.execute("SHOW COLUMNS FROM bid_incoming LIKE 'scope'")
            scope_col = cur.fetchone()
            if not scope_col:
                cur.execute("ALTER TABLE bid_incoming ADD COLUMN scope TEXT")
            else:
                # Check if it's not TEXT type and alter it
                col_type = scope_col[1] if isinstance(scope_col, tuple) else scope_col.get('Type', '')
                if col_type and 'text' not in str(col_type).lower():
                    cur.execute("ALTER TABLE bid_incoming MODIFY COLUMN scope TEXT")
        except Exception:
            pass
        
        # Create bid_assign table (with lightweight self-healing)
        bid_assign_create_sql = """
            CREATE TABLE IF NOT EXISTS bid_assign (
                a_id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT,
                b_name VARCHAR(100),
               
                due_date DATE NOT NULL DEFAULT CURRENT_TIMESTAMP,
                state VARCHAR(100),
                scope VARCHAR(100),
                type VARCHAR(100),
                company TEXT,
                depart TEXT,
                person_name TEXT,
                assignee_email VARCHAR(100),
                status TEXT,
                value INT,
                revenue DECIMAL(15,2) DEFAULT 0.00
            )
        """
        cur.execute(bid_assign_create_sql)
        try:
            cur.execute("SELECT 1 FROM bid_assign LIMIT 1")
        except Exception as e:
            try:
                err_no = e.args[0] if hasattr(e, 'args') and e.args else None
            except Exception:
                err_no = None
            if err_no in (1146, 1932):
                try:
                    cur.execute("DROP TABLE IF EXISTS bid_assign")
                    cur.execute(bid_assign_create_sql)
                except Exception:
                    pass
            else:
                raise
        
        # Create win_lost_results table (with lightweight self-healing)
        win_lost_create_sql = """
            CREATE TABLE IF NOT EXISTS win_lost_results (
                w_id INT AUTO_INCREMENT PRIMARY KEY,
                a_id INT,
                b_name TEXT,
              
                due_date INT,
                state TEXT,
                scope TEXT,
                value INT,
                company TEXT,
                department TEXT,
                person_name TEXT,
                status TEXT,
                result TEXT
            )
        """
        cur.execute(win_lost_create_sql)
        try:
            cur.execute("SELECT 1 FROM win_lost_results LIMIT 1")
        except Exception as e:
            try:
                err_no = e.args[0] if hasattr(e, 'args') and e.args else None
            except Exception:
                err_no = None
            if err_no in (1146, 1932):
                try:
                    cur.execute("DROP TABLE IF EXISTS win_lost_results")
                    cur.execute(win_lost_create_sql)
                except Exception:
                    pass
            else:
                raise
        
        # Create won_bids_result table (with lightweight self-healing)
        won_bids_create_sql = """
            CREATE TABLE IF NOT EXISTS won_bids_result (
                won_id INT AUTO_INCREMENT PRIMARY KEY,
                w_id INT,
                closure_status TEXT,
                work_progress_status TEXT
            )
        """
        cur.execute(won_bids_create_sql)
        try:
            cur.execute("SELECT 1 FROM won_bids_result LIMIT 1")
        except Exception as e:
            try:
                err_no = e.args[0] if hasattr(e, 'args') and e.args else None
            except Exception:
                err_no = None
            if err_no in (1146, 1932):
                try:
                    cur.execute("DROP TABLE IF EXISTS won_bids_result")
                    cur.execute(won_bids_create_sql)
                except Exception:
                    pass
            else:
                raise
        
        # Create assigned_bids table using the exact schema requested
        cur.execute("""
            CREATE TABLE IF NOT EXISTS assigned_bids (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT,
                b_name VARCHAR(100),
                company VARCHAR(100),
                revenue DECIMAL(15,2) DEFAULT 0.00,
                assigned_to VARCHAR(100),
                assigned_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Create work_progress_status table (extended schema used across the app, with self-healing)
        work_progress_create_sql = """
            CREATE TABLE IF NOT EXISTS work_progress_status (
                p_id INT AUTO_INCREMENT PRIMARY KEY,
                won_id INT,
                company TEXT,
                b_name TEXT,
                dept_bde TEXT,
                dept_m_d TEXT,
                dept_op TEXT,
                dept_site TEXT,
                pr_completion_status TEXT
            )
        """
        cur.execute(work_progress_create_sql)
        try:
            cur.execute("SELECT 1 FROM work_progress_status LIMIT 1")
        except Exception as e:
            try:
                err_no = e.args[0] if hasattr(e, 'args') and e.args else None
            except Exception:
                err_no = None
            if err_no in (1146, 1932):
                try:
                    cur.execute("DROP TABLE IF EXISTS work_progress_status")
                    cur.execute(work_progress_create_sql)
                except Exception:
                    pass
            else:
                raise

        # Create logs table for tracking user actions (with lightweight self-healing)
        logs_create_sql = """
            CREATE TABLE IF NOT EXISTS logs (
                id INT AUTO_INCREMENT PRIMARY KEY,
                action VARCHAR(255) NOT NULL,
                user_id INT,
                timestamp DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users(id)
            )
        """
        cur.execute(logs_create_sql)
        # Verify logs table is readable; repair if it's missing or has orphaned metadata (errors 1146, 1932)
        try:
            cur.execute("SELECT 1 FROM logs LIMIT 1")
        except Exception as e:
            try:
                err_no = e.args[0] if hasattr(e, 'args') and e.args else None
            except Exception:
                err_no = None
            if err_no in (1146, 1932):
                try:
                    cur.execute("DROP TABLE IF EXISTS logs")
                    cur.execute(logs_create_sql)
                except Exception:
                    # Let the outer handler report any persistent issues
                    pass
            else:
                # Unexpected error; let the outer handler surface it
                raise
        
        # Create companies table (with lightweight self-healing for common MySQL errors)
        companies_create_sql = """
            CREATE TABLE IF NOT EXISTS companies (
                id INT AUTO_INCREMENT PRIMARY KEY,
                name VARCHAR(100) UNIQUE NOT NULL,
                description TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """
        cur.execute(companies_create_sql)
        # Verify companies table is readable; repair if it's missing or has orphaned metadata (errors 1146, 1932)
        try:
            cur.execute("SELECT 1 FROM companies LIMIT 1")
        except Exception as e:
            try:
                err_no = e.args[0] if hasattr(e, 'args') and e.args else None
            except Exception:
                err_no = None
            if err_no in (1146, 1932):
                try:
                    cur.execute("DROP TABLE IF EXISTS companies")
                    cur.execute(companies_create_sql)
                except Exception:
                    # Let the outer handler report any persistent issues
                    pass
            else:
                # Unexpected error; let the outer handler surface it
                raise
        
        # Create projects table (with lightweight self-healing for common MySQL errors)
        projects_create_sql = """
            CREATE TABLE IF NOT EXISTS projects (
                id INT AUTO_INCREMENT PRIMARY KEY,
                name VARCHAR(200) NOT NULL,
                company_id INT NOT NULL,
                start_date DATETIME,
                due_date DATETIME NOT NULL,
                revenue FLOAT DEFAULT 0.0,
                status VARCHAR(50) DEFAULT 'active',
                progress INT DEFAULT 0,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (company_id) REFERENCES companies(id)
            )
        """
        cur.execute(projects_create_sql)
        try:
            cur.execute("SELECT 1 FROM projects LIMIT 1")
        except Exception as e:
            try:
                err_no = e.args[0] if hasattr(e, 'args') and e.args else None
            except Exception:
                err_no = None
            if err_no in (1146, 1932):
                try:
                    cur.execute("DROP TABLE IF EXISTS projects")
                    cur.execute(projects_create_sql)
                except Exception:
                    # Let the outer handler report any persistent issues
                    pass
            else:
                # Unexpected error; let the outer handler surface it
                raise

        # Create tasks table (with lightweight self-healing for common MySQL errors)
        tasks_create_sql = """
            CREATE TABLE IF NOT EXISTS tasks (
                id INT AUTO_INCREMENT PRIMARY KEY,
                name VARCHAR(200) NOT NULL,
                project_id INT NOT NULL,
                assigned_user_id INT,
                due_date DATETIME NOT NULL,
                status VARCHAR(50) DEFAULT 'pending',
                priority VARCHAR(20) DEFAULT 'medium',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (project_id) REFERENCES projects(id),
                FOREIGN KEY (assigned_user_id) REFERENCES users(id)
            )
        """
        cur.execute(tasks_create_sql)
        try:
            cur.execute("SELECT 1 FROM tasks LIMIT 1")
        except Exception as e:
            try:
                err_no = e.args[0] if hasattr(e, 'args') and e.args else None
            except Exception:
                err_no = None
            if err_no in (1146, 1932):
                try:
                    cur.execute("DROP TABLE IF EXISTS tasks")
                    cur.execute(tasks_create_sql)
                except Exception:
                    # Let the outer handler report any persistent issues
                    pass
            else:
                # Unexpected error; let the outer handler surface it
                raise
        
        # Create bid_timeline table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS bid_timeline (
                id INT AUTO_INCREMENT PRIMARY KEY,
                bid_id INT,
                event VARCHAR(200) NOT NULL,
                details TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Create employees table for team-specific employee management
        cur.execute("""
            CREATE TABLE IF NOT EXISTS employees (
                id INT AUTO_INCREMENT PRIMARY KEY,
                name VARCHAR(100) NOT NULL,
                email VARCHAR(100) UNIQUE NOT NULL,
                password VARCHAR(255) NOT NULL,
                department VARCHAR(50) NOT NULL,
                team_lead_id INT,
                is_active BOOLEAN DEFAULT TRUE,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (team_lead_id) REFERENCES users(id)
            )
        """)
        
        # Create bid_checklists table for task management per bid (with InnoDB orphan tablespace repair)
        create_sql = (
            """
            CREATE TABLE IF NOT EXISTS bid_checklists (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                task_name VARCHAR(200) NOT NULL,
                description TEXT,
                assigned_to INT,
                status VARCHAR(50) DEFAULT 'pending',
                progress_pct INT DEFAULT NULL,
                stage VARCHAR(50),
                priority VARCHAR(20) DEFAULT 'medium',
                due_date DATETIME,
                attachment_path VARCHAR(255),
                created_by INT,
                team_archive VARCHAR(50),
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                FOREIGN KEY (g_id) REFERENCES go_bids(g_id),
                FOREIGN KEY (assigned_to) REFERENCES employees(id),
                FOREIGN KEY (created_by) REFERENCES users(id)
            ) ENGINE=InnoDB
            """
        )
        try:
            cur.execute(create_sql)
        except Exception as e:
            try:
                err_no = e.args[0] if hasattr(e, 'args') and e.args else None
            except Exception:
                err_no = None
            # 1813: orphan tablespace exists on disk; attempt OS-level cleanup then recreate
            if err_no == 1813:
                try:
                    cur2 = mysql.connection.cursor()
                    cur2.execute("SHOW VARIABLES LIKE 'datadir'")
                    row = cur2.fetchone()
                    data_dir = None
                    if isinstance(row, (list, tuple)) and len(row) >= 2:
                        data_dir = row[1]
                    elif isinstance(row, dict):
                        data_dir = row.get('Value') or row.get('value')
                    cur2.close()
                    if data_dir:
                        db_name = app.config.get('MYSQL_DB', 'esco')
                        ibd_path = os.path.join(data_dir, db_name, 'bid_checklists.ibd')
                        cfg_path = os.path.join(data_dir, db_name, 'bid_checklists.cfg')
                        for p in (ibd_path, cfg_path):
                            try:
                                if os.path.exists(p):
                                    os.remove(p)
                            except Exception:
                                pass
                except Exception:
                    pass
                try:
                    cur.execute("DROP TABLE IF EXISTS bid_checklists")
                    mysql.connection.commit()
                except Exception:
                    pass
                # Retry create
                cur.execute(create_sql)
            else:
                raise

        # Ensure progress_pct column exists even on older databases
        cur.execute("SELECT COUNT(*) AS cnt FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME='bid_checklists' AND COLUMN_NAME='progress_pct'")
        row = cur.fetchone()
        if not row or int(row.get('cnt', 0)) == 0:
            try:
                cur.execute("ALTER TABLE bid_checklists ADD COLUMN progress_pct INT NULL AFTER status")
            except Exception:
                pass

        # Ensure stage column exists
        cur.execute("SELECT COUNT(*) AS cnt FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME='bid_checklists' AND COLUMN_NAME='stage'")
        row = cur.fetchone()
        if not row or int(row.get('cnt', 0)) == 0:
            try:
                cur.execute("ALTER TABLE bid_checklists ADD COLUMN stage VARCHAR(50) NULL AFTER progress_pct")
            except Exception:
                pass

        # Ensure attachment_path column exists
        cur.execute("SELECT COUNT(*) AS cnt FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME='bid_checklists' AND COLUMN_NAME='attachment_path'")
        row = cur.fetchone()
        if not row or int(row.get('cnt', 0)) == 0:
            try:
                cur.execute("ALTER TABLE bid_checklists ADD COLUMN attachment_path VARCHAR(255) NULL AFTER due_date")
            except Exception:
                pass
        # Create dynamic stage tables used across dashboards
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_stage_exclusions (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_bid_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_custom_stages (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_custom_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )

        # Create dynamic stage tables used across dashboards
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_stage_exclusions (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_bid_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_custom_stages (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_custom_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )

        # Create project_transfers table for tracking project handoffs between teams
        cur.execute("""
            CREATE TABLE IF NOT EXISTS project_transfers (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                from_team VARCHAR(50) NOT NULL,
                to_team VARCHAR(50) NOT NULL,
                transferred_by INT,
                transfer_reason TEXT,
                status VARCHAR(50) DEFAULT 'pending',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (g_id) REFERENCES go_bids(g_id),
                FOREIGN KEY (transferred_by) REFERENCES users(id)
            )
        """)

        # Comments table for board modal
        cur.execute("""
            CREATE TABLE IF NOT EXISTS bid_comments (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                user_id INT,
                comment_text TEXT NOT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (g_id) REFERENCES go_bids(g_id),
                FOREIGN KEY (user_id) REFERENCES users(id)
            )
        """)

        # Comments table for board modal
        cur.execute("""
            CREATE TABLE IF NOT EXISTS bid_comments (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                user_id INT,
                comment_text TEXT NOT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (g_id) REFERENCES go_bids(g_id),
                FOREIGN KEY (user_id) REFERENCES users(id)
            )
        """)

        # --- Profiling tables ---
        cur.execute("""
            CREATE TABLE IF NOT EXISTS company_preferences (
                id INT AUTO_INCREMENT PRIMARY KEY,
                registered_states TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS company_capabilities (
                id INT AUTO_INCREMENT PRIMARY KEY,
                description TEXT,
                file_path VARCHAR(255),
                uploaded_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS company_performance (
                id INT AUTO_INCREMENT PRIMARY KEY,
                project_name VARCHAR(200) NOT NULL,
                year INT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        # Mapping table for multi-employee bid assignments
        cur.execute("""
            CREATE TABLE IF NOT EXISTS bid_assignment_members (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                employee_id INT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_g_emp (g_id, employee_id)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """)

        # Company details (profile) table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS company_details (
                id INT AUTO_INCREMENT PRIMARY KEY,
                name VARCHAR(200),
                website VARCHAR(255),
                email VARCHAR(255),
                phone VARCHAR(50),
                about TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)

        mysql.connection.commit()
        cur.close()
        print("Database tables created/verified successfully")
    except Exception as e:
        print(f"Error creating tables: {e}")
        try:
            mysql.connection.rollback()
        except Exception:
            pass
        if 'cur' in locals():
            try:
                cur.close()
            except Exception:
                pass
if __name__ == '__main__':
    from datetime import datetime, timedelta
    
    with app.app_context():
        _ensure_tables_exist()
        
        cur = mysql.connection.cursor(DictCursor)
        
        # Always ensure companies exist
        cur.execute("SELECT COUNT(*) as count FROM companies")
        if cur.fetchone()['count'] == 0:
            cur.execute("""
                INSERT INTO companies (name, description) VALUES 
                ('Ikio', 'Renewable Energy Solutions'),
                ('Metco', 'Industrial Energy Management'),
                ('Sunsprint', 'Solar Power Systems')
            """)
            mysql.connection.commit()
            print("Companies created successfully")
        
        # Check if users exist
        cur.execute("SELECT COUNT(*) as count FROM users")
        if cur.fetchone()['count'] == 0:
            # Create admin user
            cur.execute("""
                INSERT INTO users (email, password, is_admin, role) VALUES 
                ('admin@example.com', 'admin', 1, 'admin')
            """)
            admin_user_id = cur.lastrowid
            
            # Create other users
            cur.execute("""
                INSERT INTO users (email, password, is_admin, role) VALUES 
                ('bd@example.com', 'user', 0, 'business dev'),
                ('designer@example.com', 'designer', 0, 'design'),
                ('ops@example.com', 'ops', 0, 'operations'),
                ('sitemgr@example.com', 'site', 0, 'site manager')
            """)
            # Fetch exact user IDs by email to avoid relying on lastrowid math
            cur.execute("SELECT id FROM users WHERE email=%s", ('bd@example.com',))
            user_bdm_id = cur.fetchone()['id']
            cur.execute("SELECT id FROM users WHERE email=%s", ('designer@example.com',))
            user_design_id = cur.fetchone()['id']
            cur.execute("SELECT id FROM users WHERE email=%s", ('ops@example.com',))
            user_ops_id = cur.fetchone()['id']
            cur.execute("SELECT id FROM users WHERE email=%s", ('sitemgr@example.com',))
            user_site_id = cur.fetchone()['id']
            
            # Get company ids for linking bids
            cur.execute("SELECT id FROM companies WHERE name='Ikio'")
            ikio_id = cur.fetchone()['id']
            cur.execute("SELECT id FROM companies WHERE name='Metco'")
            metco_id = cur.fetchone()['id']
            cur.execute("SELECT id FROM companies WHERE name='Sunsprint'")
            sunsprint_id = cur.fetchone()['id']

            # Create sample bids linked to companies
            cur.execute("""
            INSERT INTO bids (name, current_stage, user_id, company_id) VALUES 
            ('Project Alpha', 'business', %s, %s),
            ('Project Beta', 'design', %s, %s),
            ('Project Gamma', 'operations', %s, %s)
        """, (user_bdm_id, ikio_id, user_design_id, metco_id, user_ops_id, sunsprint_id))
            # Create sample bid_incoming data
            cur.execute("""
                INSERT INTO bid_incoming (b_name, due_date, state, scope, type, scoring, comp_name, decision, summary) VALUES 
                ('Solar Energy Project', %s, 'submitted', 'Installation of 500kW solar panels for commercial building', 'Renewable Energy', 85, 'Ikio', 'GO', 'High potential project with good ROI'),
                ('Wind Farm Development', %s, 'under_review', 'Development of 2MW wind farm in rural area', 'Wind Energy', 72, 'Metco', 'NO-GO', 'Land acquisition issues identified'),
                ('Energy Efficiency Audit', %s, 'pending', 'Comprehensive energy audit for manufacturing facility', 'Energy Management', 90, 'Sunsprint', 'WON', 'Excellent technical proposal with competitive pricing'),
                ('Battery Storage System', %s, 'submitted', 'Installation of 1MWh battery storage system', 'Energy Storage', 78, 'Ikio', 'LOST', 'Lost to competitor with lower bid'),
                ('Smart Grid Implementation', %s, 'completed', 'Implementation of smart grid technology for city', 'Smart Grid', 95, 'Metco', 'WON', 'Successfully completed project ahead of schedule')
            """, (
                datetime.now() + timedelta(days=30),
                datetime.now() + timedelta(days=45),
                datetime.now() + timedelta(days=15),
                datetime.now() + timedelta(days=60),
                datetime.now() + timedelta(days=90)
            ))
            
            # One sample bid ready for site manager handover
            cur.execute("UPDATE bids SET current_stage='site_manager' WHERE name='Project Beta'")
            
            mysql.connection.commit()
            
            # Company ids already loaded above
            
            # Create sample projects
            cur.execute("""
                INSERT INTO projects (name, company_id, start_date, due_date, status, progress) VALUES 
                ('Solar Farm Installation', %s, %s, %s, 50000, 'active', 45),
                ('Wind Energy Project', %s, %s, %s, 75000, 'active', 70),
                ('Energy Efficiency Audit', %s, %s, %s, 25000, 'active', 30),
                ('Industrial Solar Setup', %s, %s, %s, 100000, 'active', 15),
                ('Residential Solar Panel', %s, %s, %s, 30000, 'active', 80),
                ('Commercial Solar System', %s, %s, %s, 60000, 'active', 25)
            """, (
                ikio_id, datetime.now(), datetime.now() + timedelta(days=90),
                ikio_id, datetime.now() - timedelta(days=30), datetime.now() + timedelta(days=60),
                metco_id, datetime.now() - timedelta(days=15), datetime.now() + timedelta(days=45),
                metco_id, datetime.now(), datetime.now() + timedelta(days=120),
                sunsprint_id, datetime.now() - timedelta(days=10), datetime.now() + timedelta(days=30),
                sunsprint_id, datetime.now(), datetime.now() + timedelta(days=75)
            ))
            
            # Get project IDs for tasks
            cur.execute("SELECT id FROM projects ORDER BY id")
            project_ids = [row['id'] for row in cur.fetchall()]
            
            # Create sample tasks
            cur.execute("""
                INSERT INTO tasks (name, project_id, assigned_user_id, due_date, status, priority) VALUES 
                ('Site Survey', %s, %s, %s, 'in_progress', 'high'),
                ('Equipment Procurement', %s, %s, %s, 'pending', 'medium'),
                ('Installation Planning', %s, %s, %s, 'completed', 'high'),
                ('Energy Assessment', %s, %s, %s, 'in_progress', 'urgent'),
                ('Client Consultation', %s, %s, %s, 'pending', 'high'),
                ('System Testing', %s, %s, %s, 'in_progress', 'medium'),
                ('Documentation', %s, %s, %s, 'pending', 'low')
            """, (
                project_ids[0], user_bdm_id, datetime.now() + timedelta(days=5),
                project_ids[0], user_design_id, datetime.now() + timedelta(days=15),
                project_ids[1], admin_user_id, datetime.now() + timedelta(days=10),
                project_ids[2], user_ops_id, datetime.now() + timedelta(days=7),
                project_ids[3], user_site_id, datetime.now() + timedelta(days=3),
                project_ids[4], admin_user_id, datetime.now() + timedelta(days=2),
                project_ids[5], user_bdm_id, datetime.now() + timedelta(days=20)
            ))
            
            mysql.connection.commit()
            
        # Seed at least one log if none exist
        cur.execute("SELECT COUNT(*) as count FROM logs")
        if cur.fetchone()['count'] == 0:
            cur.execute("INSERT INTO logs (action) VALUES ('System initialized and sample data seeded.')")
            mysql.connection.commit()
        
        cur.close()
    
    socketio.run(app, debug=True, port=5001)