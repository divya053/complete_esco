

import os
import json
import re
from pathlib import Path
from datetime import datetime, date
from html import escape
from flask import (
    Flask,
    render_template,
    request,
    jsonify,
    redirect,
    url_for,
    flash,
    send_file,
    session,
    send_from_directory,
    abort,
)
from werkzeug.utils import secure_filename
import smtplib
import ssl
import requests

from dotenv import load_dotenv


def _load_dotenv_hierarchy():
    """Load environment variables from common .env locations.

    We look in the following order (later files override earlier values):
      1. Project root two levels up (handles running via different entrypoints)
      2. Directory that contains this file
      3. An `env/.env` helper directory beside the project (if present)
    """
    current_file = Path(__file__).resolve()
    candidates = [
        current_file.parent.parent / '.env',
        current_file.parent / '.env',
        current_file.parent / 'env' / '.env',
        Path.cwd() / '.env',
    ]

    loaded_any = False
    for candidate in candidates:
        if candidate.is_file():
            load_dotenv(candidate, override=True)
            loaded_any = True
    if not loaded_any:
        load_dotenv()


_load_dotenv_hierarchy()

def _apply_provider_defaults():
    provider = (os.getenv('LLM_PROVIDER') or '').strip().lower()
    if provider == 'groq':
        base = os.getenv('OPENAI_BASE_URL') or os.getenv('GROQ_BASE_URL') or 'https://api.groq.com/openai/v1'
        model = os.getenv('OPENAI_MODEL') or os.getenv('GROQ_MODEL') or 'meta-llama/llama-4-scout-17b-16e-instruct'
        key = os.getenv('OPENAI_API_KEY') or os.getenv('GROQ_API_KEY') or ''
    elif provider == 'openrouter':
        base = os.getenv('OPENAI_BASE_URL') or os.getenv('OPENROUTER_BASE_URL') or 'https://openrouter.ai/api/v1'
        model = os.getenv('OPENAI_MODEL') or os.getenv('OPENROUTER_MODEL') or 'openrouter/auto'
        key = os.getenv('OPENAI_API_KEY') or os.getenv('OPENROUTER_API_KEY') or ''
    elif provider == 'ollama':
        base = os.getenv('OPENAI_BASE_URL') or os.getenv('OLLAMA_BASE_URL') or 'http://localhost:11434/v1'
        model = os.getenv('OPENAI_MODEL') or os.getenv('OLLAMA_MODEL') or 'llama3.2:3b-instruct'
        key = os.getenv('OPENAI_API_KEY') or os.getenv('OLLAMA_API_KEY') or 'ollama'
    else:
        base = os.getenv('OPENAI_BASE_URL') or 'https://api.openai.com/v1'
        model = os.getenv('OPENAI_MODEL') or 'gpt-4o-mini'
        key = os.getenv('OPENAI_API_KEY') or ''
    app.config['OPENAI_BASE_URL'] = base
    app.config['OPENAI_MODEL'] = model
    app.config['OPENAI_API_KEY'] = key

_apply_provider_defaults()

def _strip_html_tags(value):
    if not value:
        return ''
    text = re.sub(r'(<\s*br\s*/?>)', '\n', value, flags=re.IGNORECASE)
    text = re.sub(r'</p\s*>', '\n\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    return re.sub(r'\n{3,}', '\n\n', text).strip()

try:
    import fitz as _pymupdf  # PyMuPDF
    _FITZ_AVAILABLE = hasattr(_pymupdf, "open")
except ImportError:
    _pymupdf = None
    _FITZ_AVAILABLE = False

try:
    from PyPDF2 import PdfReader as _PdfReader
except ImportError:
    _PdfReader = None
from flask_mysqldb import MySQL
from flask_socketio import SocketIO
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from MySQLdb.cursors import DictCursor
from rfp_analyzer_routes import rfp_bp 
from rfp_analyzer_routes import scorer, trim_text_to_token_limit
# --- App Initialization ---
app = Flask(__name__)
app.config['SECRET_KEY'] = 'a_very_secret_key_for_sessions'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file upload

# Ensure commonly used AI configuration keys are visible via app.config
app.config.setdefault('OPENAI_API_KEY', os.getenv('OPENAI_API_KEY', app.config.get('OPENAI_API_KEY', '')))
app.config.setdefault('OPENAI_BASE_URL', os.getenv('OPENAI_BASE_URL', app.config.get('OPENAI_BASE_URL', 'https://api.groq.com/openai/v1')))
app.config.setdefault('OPENAI_MODEL', os.getenv('OPENAI_MODEL', app.config.get('OPENAI_MODEL', 'meta-llama/llama-4-scout-17b-16e-instruct')))

# Register RFP Analyzer Blueprint
app.register_blueprint(rfp_bp)

# --- Error handling: ensure API endpoints return JSON on errors ---
@app.errorhandler(500)
def _handle_internal_server_error(e):
    try:
        path = request.path or ''
    except Exception:
        path = ''
    if isinstance(path, str) and path.startswith('/api/'):
        return jsonify({'error': 'internal_server_error', 'message': 'An unexpected error occurred.'}), 500
    # Fallback for non-API routes
    return "Internal Server Error", 500

@app.errorhandler(TypeError)
def _handle_type_error(e):
    # Catch cases like "view function did not return a valid response" and return JSON for APIs
    try:
        path = request.path or ''
    except Exception:
        path = ''
    if isinstance(path, str) and path.startswith('/api/'):
        return jsonify({'error': 'type_error', 'message': str(e)}), 500
    return str(e), 500

# --- Stage Constants ---
PIPELINE = ['analyzer', 'business', 'design', 'operations', 'engineer', 'handover']
LABELS = {
    'analyzer': 'BID Analyzer',
    'business': 'Business Development', 
    'design': 'Design Team',
    'operations': 'Operations Team',
    'engineer': 'Site Engineer',
    'handover': 'Handov'
}

def pct_for(stage: str) -> int:
    """Calculate progress percentage based on stage"""
    s = (stage or 'analyzer').lower()
    i = PIPELINE.index(s) if s in PIPELINE else 0
    return int(round(i * (100 / (len(PIPELINE) - 1))))  # 0,20,40,60,80,100

def status_texts(stage: str) -> tuple[str, str]:
    """Generate project status and work status texts based on stage"""
    s = (stage or 'analyzer').lower()
    proj = 'completed' if s == 'handover' else 'ongoing'
    if s == 'analyzer':
        work = 'Initiated by BID Analyzer'
    else:
        i = PIPELINE.index(s) if s in PIPELINE else 0
        prev = PIPELINE[i-1] if i > 0 else ''
        from_txt = LABELS.get(prev, '').replace(' Team', '')
        to_txt = LABELS.get(s, '')
        work = f'Updated by {from_txt} to {to_txt}'
    return proj, work

def generate_team_checklist(cur, g_id, team):
    """Generate team-specific default checklist tasks"""
    team_tasks = {
        'business': [
            {'name': 'Market Analysis', 'description': 'Analyze market potential and competition', 'priority': 'high'},
            {'name': 'Client Communication', 'description': 'Establish communication with client', 'priority': 'high'},
            {'name': 'Proposal Preparation', 'description': 'Prepare business proposal', 'priority': 'medium'},
            {'name': 'Budget Estimation', 'description': 'Estimate project budget and timeline', 'priority': 'high'}
        ],
        'design': [
            {'name': 'Initial Design Concept', 'description': 'Create initial design concepts', 'priority': 'high'},
            {'name': 'Technical Drawings', 'description': 'Prepare detailed technical drawings', 'priority': 'high'},
            {'name': 'Material Selection', 'description': 'Select appropriate materials and specifications', 'priority': 'medium'},
            {'name': 'Design Review', 'description': 'Review and finalize design with team', 'priority': 'medium'},
            {'name': 'Client Approval', 'description': 'Get client approval on design', 'priority': 'high'}
        ],
        'operations': [
            {'name': 'Project Planning', 'description': 'Create detailed project execution plan', 'priority': 'high'},
            {'name': 'Resource Allocation', 'description': 'Allocate resources and personnel', 'priority': 'high'},
            {'name': 'Timeline Management', 'description': 'Set up project timeline and milestones', 'priority': 'medium'},
            {'name': 'Quality Control Setup', 'description': 'Establish quality control procedures', 'priority': 'medium'},
            {'name': 'Risk Assessment', 'description': 'Identify and assess project risks', 'priority': 'high'}
        ],
        'engineer': [
            {'name': 'Site Survey', 'description': 'Conduct detailed site survey', 'priority': 'high'},
            {'name': 'Technical Specifications', 'description': 'Prepare technical specifications', 'priority': 'high'},
            {'name': 'Safety Planning', 'description': 'Develop safety protocols and procedures', 'priority': 'high'},
            {'name': 'Equipment Planning', 'description': 'Plan equipment and tool requirements', 'priority': 'medium'},
            {'name': 'Implementation Plan', 'description': 'Create detailed implementation plan', 'priority': 'high'}
        ]
    }
    
    if team in team_tasks:
        for task in team_tasks[team]:
            # Persist a stage on each seeded task so aggregation can be stage-aware
            stage_name = team.strip().lower()
            cur.execute("""
                INSERT INTO bid_checklists (g_id, task_name, description, priority, status, progress_pct, stage, created_by)
                VALUES (%s, %s, %s, %s, 'pending', %s, %s, %s)
            """, (g_id, task['name'], task['description'], task['priority'], 0, stage_name, current_user.id))

def log_write(action: str, details: str = ''):
    """Write to logs table and emit via Socket.IO"""
    try:
        cur = mysql.connection.cursor()
        user_id = getattr(current_user, 'id', None) if hasattr(current_user, 'id') else None
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)",
                    (f"{action} | {details}", user_id))
        mysql.connection.commit()
        cur.close()
        
        # Emit to master dashboard
        socketio.emit('master_update', {'log': {
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'action': f"{action} | {details}",
            'user_email': getattr(current_user, 'email', 'System') if hasattr(current_user, 'email') else 'System',
            'user_role': getattr(current_user, 'role', '') if hasattr(current_user, 'role') else ''
        }})
    except Exception as e:
        print(f"Log write error: {e}")

# Legacy function for backward compatibility
def stage_progress_pct(stage: str) -> int:
    return pct_for(stage)

# MySQL Config
app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = ''
app.config['MYSQL_DB'] = 'esco'
app.config['MYSQL_CURSORCLASS'] = 'DictCursor'

mysql = MySQL(app)
# Store mysql on app for blueprint access
app.mysql = mysql  # type: ignore[attr-defined]

socketio = SocketIO(app)

_db_init_done = False

@app.before_request
def _init_db_once():
    global _db_init_done
    if not _db_init_done:
        try:
            _ensure_tables_exist()
            # Sanity check for bid_checklists and recover from common errors (1932 orphan tablespace, 1146 missing table)
            def _ensure_bid_checklists_exists():
                create_sql = (
                    """
                    CREATE TABLE IF NOT EXISTS bid_checklists (
                        id INT AUTO_INCREMENT PRIMARY KEY,
                        g_id INT NOT NULL,
                        task_name VARCHAR(200) NOT NULL,
                        description TEXT,
                        assigned_to INT,
                        status VARCHAR(50) DEFAULT 'pending',
                        progress_pct INT DEFAULT NULL,
                        stage VARCHAR(50),
                        priority VARCHAR(20) DEFAULT 'medium',
                        due_date DATETIME,
                        attachment_path VARCHAR(255),
                        created_by INT,
                        team_archive VARCHAR(50),
                        created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                        updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                        FOREIGN KEY (g_id) REFERENCES go_bids(g_id),
                        FOREIGN KEY (assigned_to) REFERENCES employees(id),
                        FOREIGN KEY (created_by) REFERENCES users(id)
                    ) ENGINE=InnoDB
                    """
                )
                try:
                    c = mysql.connection.cursor()
                    c.execute(create_sql)
                    mysql.connection.commit()
                    c.close()
                except Exception as ce:
                    try:
                        err_no2 = ce.args[0] if hasattr(ce, 'args') and ce.args else None
                    except Exception:
                        err_no2 = None
                    if err_no2 == 1813:  # Orphan tablespace
                        try:
                            c2 = mysql.connection.cursor()
                            c2.execute("SHOW VARIABLES LIKE 'datadir'")
                            row = c2.fetchone()
                            data_dir = None
                            if isinstance(row, (list, tuple)) and len(row) >= 2:
                                data_dir = row[1]
                            elif isinstance(row, dict):
                                data_dir = row.get('Value') or row.get('value')
                            c2.close()
                            if data_dir:
                                db_name = app.config.get('MYSQL_DB', 'esco')
                                ibd_path = os.path.join(data_dir, db_name, 'bid_checklists.ibd')
                                cfg_path = os.path.join(data_dir, db_name, 'bid_checklists.cfg')
                                for p in (ibd_path, cfg_path):
                                    try:
                                        if os.path.exists(p):
                                            os.remove(p)
                                    except Exception:
                                        pass
                        except Exception:
                            pass
                        try:
                            c3 = mysql.connection.cursor()
                            c3.execute("DROP TABLE IF EXISTS bid_checklists")
                            mysql.connection.commit()
                            c3.close()
                        except Exception:
                            pass
                        c4 = mysql.connection.cursor()
                        c4.execute(create_sql)
                        mysql.connection.commit()
                        c4.close()
                    else:
                        raise
            # Verify existence
            success = False
            for _ in range(2):
                try:
                    cur = mysql.connection.cursor()
                    cur.execute("SELECT 1 FROM bid_checklists LIMIT 1")
                    cur.close()
                    success = True
                    break
                except Exception as e:
                    try:
                        err_no = e.args[0] if hasattr(e, 'args') and e.args else None
                    except Exception:
                        err_no = None
                    if err_no in (1146, 1932):
                        _ensure_bid_checklists_exists()
                    else:
                        break
            if success:
                _db_init_done = True
        except Exception as e:
            print(f"Database init error: {e}")

# Removed eager import-time init to avoid referencing functions before definition

@app.route('/profiling', methods=['GET', 'POST'], endpoint='profiling')
@login_required
def profiling_page():
    if not current_user.is_admin:
        return "Access Denied", 403
    # Knowledge Hub base dir
    base_dir = os.path.join(os.getcwd(), 'uploads', 'knowledge')
    os.makedirs(base_dir, exist_ok=True)
    message = ''
    preferences = None
    capabilities = []
    performances = []
    details = None
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Ensure mapping table exists for multi-assign (first thing)
        try:
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS bid_assignment_members (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    g_id INT NOT NULL,
                    employee_id INT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE KEY uniq_g_emp (g_id, employee_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
                """
            )
        except Exception:
            pass
        if request.method == 'POST':
            action = request.form.get('action')
            if action == 'save_preferences':
                states = (request.form.get('registered_states') or '').strip()
                cur.execute("DELETE FROM company_preferences")
                cur.execute("INSERT INTO company_preferences (registered_states) VALUES (%s)", (states,))
                mysql.connection.commit()
                message = 'Preferences saved.'
            elif action == 'add_capability':
                desc = request.form.get('cap_description') or ''
                file = request.files.get('cap_file')
                file_path = None
                if file and file.filename:
                    uploads_dir = os.path.join(os.getcwd(), 'uploads', 'capabilities')
                    os.makedirs(uploads_dir, exist_ok=True)
                    fname = secure_filename(file.filename)
                    save_path = os.path.join(uploads_dir, fname)
                    file.save(save_path)
                    file_path = f"capabilities/{fname}"
                cur.execute("INSERT INTO company_capabilities (description, file_path) VALUES (%s, %s)", (desc, file_path))
                mysql.connection.commit()
                message = 'Capability added.'
            elif action == 'add_performance':
                project = (request.form.get('project_name') or '').strip()
                year_raw = (request.form.get('project_year') or '').strip()
                year = None
                if year_raw:
                    try:
                        year = int(year_raw[:4])
                    except Exception:
                        year = None
                cur.execute("INSERT INTO company_performance (project_name, year) VALUES (%s, %s)", (project, year))
                mysql.connection.commit()
                message = 'Past performance saved.'
            elif action == 'save_details':
                name = (request.form.get('name') or '').strip()
                website = (request.form.get('website') or '').strip()
                email = (request.form.get('email') or '').strip()
                phone = (request.form.get('phone') or '').strip()
                about = (request.form.get('about') or '').strip()
                # Store a single row snapshot (replace existing)
                cur.execute("DELETE FROM company_details")
                cur.execute(
                    "INSERT INTO company_details (name, website, email, phone, about) VALUES (%s,%s,%s,%s,%s)",
                    (name, website, email, phone, about)
                )
                mysql.connection.commit()
                message = 'Company details saved.'
            elif action == 'create_folder':
                folder_name = (request.form.get('folder_name') or '').strip()
                if folder_name:
                    target = os.path.join(base_dir, folder_name)
                    os.makedirs(target, exist_ok=True)
                    message = f'Folder "{folder_name}" created.'
            elif action == 'upload_docs':
                target_folder = (request.form.get('target_folder') or '').strip()
                target_path = os.path.join(base_dir, target_folder) if target_folder else base_dir
                os.makedirs(target_path, exist_ok=True)
                files = request.files.getlist('documents')
                for f in files:
                    if not f.filename:
                        continue
                    safe_name = f.filename.replace('..','_')
                    f.save(os.path.join(target_path, safe_name))
                message = 'Documents uploaded.'
        else:
            # Fallback: accept GET query submission for company details if present (prevents user confusion)
            q = request.args
            if any(k in q for k in ('name','website','email','phone','about')):
                name = (q.get('name') or '').strip()
                website = (q.get('website') or '').strip()
                email = (q.get('email') or '').strip()
                phone = (q.get('phone') or '').strip()
                about = (q.get('about') or '').strip()
                if any([name, website, email, phone, about]):
                    cur.execute("DELETE FROM company_details")
                    cur.execute(
                        "INSERT INTO company_details (name, website, email, phone, about) VALUES (%s,%s,%s,%s,%s)",
                        (name, website, email, phone, about)
                    )
                    mysql.connection.commit()
                    message = 'Company details saved.'

        # Load current data for all tabs
        cur.execute("SELECT * FROM company_preferences ORDER BY id DESC LIMIT 1")
        preferences = cur.fetchone()
        # Some databases may not have the uploaded_at column from older migrations; fall back to id
        try:
            cur.execute("SELECT * FROM company_capabilities ORDER BY uploaded_at DESC")
        except Exception:
            cur.execute("SELECT * FROM company_capabilities ORDER BY id DESC")
        capabilities = cur.fetchall()
        cur.execute("SELECT * FROM company_performance ORDER BY year DESC, id DESC")
        performances = cur.fetchall()
        cur.execute("SELECT * FROM company_details ORDER BY id DESC LIMIT 1")
        details = cur.fetchone()
    finally:
        cur.close()

    # Build knowledge hub lists
    folders = []
    files = []
    try:
        for entry in sorted(os.listdir(base_dir)):
            p = os.path.join(base_dir, entry)
            if os.path.isdir(p):
                folders.append(entry)
            else:
                files.append(entry)
    except Exception:
        pass

    return render_template('profiling.html',
                           preferences=preferences,
                           capabilities=capabilities,
                           performances=performances,
                           details=details,
                           folders=folders,
                           files=files,
                           message=message)
# --- Ensure seed tasks exist for every team (parallel kickoff) ---
def ensure_tasks_for_team(team: str):
    cur = mysql.connection.cursor(DictCursor)
    try:
        cur.execute(
            """
            SELECT gb.g_id
            FROM go_bids gb
            LEFT JOIN bid_checklists bc
              ON bc.g_id = gb.g_id AND LOWER(COALESCE(bc.stage,'')) = %s
            WHERE bc.id IS NULL
            """, (team,)
        )
        missing = [row['g_id'] for row in cur.fetchall()]
        if missing:
            for g_id in missing:
                generate_team_checklist(cur, g_id, team)
            mysql.connection.commit()
    except Exception:
        mysql.connection.rollback()
    finally:
        cur.close()
@app.route('/knowledge', methods=['GET', 'POST'])
@login_required
def knowledge():
    if not current_user.is_admin:
        return "Access Denied", 403
    base_dir = os.path.join(os.getcwd(), 'uploads', 'knowledge')
    os.makedirs(base_dir, exist_ok=True)

    message = ''
    if request.method == 'POST':
        action = request.form.get('action')
        if action == 'create_folder':
            folder_name = (request.form.get('folder_name') or '').strip()
            if folder_name:
                target = os.path.join(base_dir, folder_name)
                os.makedirs(target, exist_ok=True)
                message = f'Folder "{folder_name}" created.'
        elif action == 'upload_docs':
            target_folder = (request.form.get('target_folder') or '').strip()
            target_path = os.path.join(base_dir, target_folder) if target_folder else base_dir
            os.makedirs(target_path, exist_ok=True)
            files = request.files.getlist('documents')
            for f in files:
                if not f.filename:
                    continue
                safe_name = f.filename.replace('..','_')
                f.save(os.path.join(target_path, safe_name))
            message = 'Documents uploaded.'

    # List folders and files
    folders = []
    files = []
    for entry in sorted(os.listdir(base_dir)):
        p = os.path.join(base_dir, entry)
        if os.path.isdir(p):
            folders.append(entry)
        else:
            files.append(entry)

    return render_template('profiling.html', folders=folders, files=files, message=message)
# --- Helper deprecated: revenue-based auto assignment removed ---
def assign_bids_by_revenue():
    return None

# assign_go_bids deprecated
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

@login_manager.unauthorized_handler
def _unauthorized_handler():
    # Return JSON errors for API endpoints instead of HTML redirects
    try:
        path = request.path or ''
    except Exception:
        path = ''
    if isinstance(path, str) and path.startswith('/api/'):
        return jsonify({'error': 'unauthorized', 'message': 'Login required. Please sign in and try again.'}), 401
    return redirect(url_for('login'))

# Custom User class
class User(UserMixin):
    def __init__(self, id, email, password, is_admin, role):
        self.id = id
        self.email = email
        self.password = password
        self.is_admin = bool(is_admin)
        self.role = role
    
    @property
    def is_supervisor(self):
        return self.role.lower() == 'supervisor'
    
    @property
    def can_assign_stages(self):
        return self.is_admin or self.is_supervisor
    
    @property
    def can_alter_timeline(self):
        return self.is_admin or self.is_supervisor
# --- User Loader for Flask-Login ---
@login_manager.user_loader
def load_user(user_id):
    cur = mysql.connection.cursor(DictCursor)
    cur.execute("SELECT * FROM users WHERE id=%s", (user_id,))
    data = cur.fetchone()
    cur.close()
    if data:
        return User(data['id'], data['email'], data['password'], data['is_admin'], data['role'])
    return None

# --- Utility: ensure GO bids from bid_incoming are mirrored in go_bids ---
def sync_go_bids() -> int:
    """Synchronize GO decisions from bid_incoming into go_bids.

    - Inserts any missing GO bids.
    - Updates existing go_bids rows with the latest information.
    - Removes entries that are no longer marked as GO.

    Returns number of newly inserted rows.
    """
    cur = mysql.connection.cursor(DictCursor)
    inserted = 0
    try:
        # Insert new GO bids that do not yet exist in go_bids
        cur.execute(
            """
            INSERT INTO go_bids (id, b_name, due_date, state, scope, type, scoring, company, decision, summary)
            SELECT bi.id,
                   bi.b_name,
                   bi.due_date,
                   COALESCE(NULLIF(TRIM(bi.state), ''), 'business'),
                   bi.scope,
                   bi.type,
                   bi.scoring,
                   bi.comp_name,
                   bi.decision,
                   bi.summary
            FROM bid_incoming bi
            LEFT JOIN go_bids gb ON gb.id = bi.id
            WHERE gb.id IS NULL
              AND UPPER(TRIM(COALESCE(bi.decision, ''))) = 'GO'
            """
        )
        inserted = cur.rowcount or 0

        # Update existing go_bids rows to reflect latest bid_incoming data
        cur.execute(
            """
            UPDATE go_bids gb
            JOIN bid_incoming bi ON bi.id = gb.id
            SET gb.b_name   = bi.b_name,
                gb.due_date = COALESCE(bi.due_date, gb.due_date),
                gb.state    = COALESCE(NULLIF(TRIM(bi.state), ''), gb.state),
                gb.scope    = bi.scope,
                gb.type     = bi.type,
                gb.scoring  = bi.scoring,
                gb.company  = COALESCE(bi.comp_name, gb.company),
                gb.decision = bi.decision,
                gb.summary  = bi.summary
            WHERE UPPER(TRIM(COALESCE(bi.decision, ''))) = 'GO'
            """
        )

        # Remove go_bids entries that are no longer GO decisions or missing source
        # First, get the g_id values that need to be deleted
        cur.execute(
            """
            SELECT gb.g_id
            FROM go_bids gb
            LEFT JOIN bid_incoming bi ON bi.id = gb.id
            WHERE gb.id IS NOT NULL
              AND (
                    bi.id IS NULL
                 OR UPPER(TRIM(COALESCE(bi.decision, ''))) <> 'GO'
              )
            """
        )
        g_ids_to_delete = [row['g_id'] for row in cur.fetchall()]
        
        if g_ids_to_delete:
            # Delete from all tables that have foreign key constraints
            placeholders = ','.join(['%s'] * len(g_ids_to_delete))
            
            # Delete from bid_checklists
            try:
                cur.execute(f"DELETE FROM bid_checklists WHERE g_id IN ({placeholders})", g_ids_to_delete)
            except Exception:
                pass
            
            # Delete from bid_stage_exclusions
            try:
                cur.execute(f"DELETE FROM bid_stage_exclusions WHERE g_id IN ({placeholders})", g_ids_to_delete)
            except Exception:
                pass
            
            # Delete from bid_custom_stages
            try:
                cur.execute(f"DELETE FROM bid_custom_stages WHERE g_id IN ({placeholders})", g_ids_to_delete)
            except Exception:
                pass
            
            # Delete from bid_assignment_members
            try:
                cur.execute(f"DELETE FROM bid_assignment_members WHERE g_id IN ({placeholders})", g_ids_to_delete)
            except Exception:
                pass
            
            # Get a_id values from bid_assign before deleting
            cur.execute(f"SELECT a_id FROM bid_assign WHERE g_id IN ({placeholders})", g_ids_to_delete)
            a_ids = [row['a_id'] for row in cur.fetchall() if row.get('a_id')]
            
            # Delete from bid_assign
            try:
                cur.execute(f"DELETE FROM bid_assign WHERE g_id IN ({placeholders})", g_ids_to_delete)
            except Exception:
                pass
            
            # Delete related win_lost_results, won_bids_result, and work_progress_status
            if a_ids:
                a_placeholders = ','.join(['%s'] * len(a_ids))
                try:
                    cur.execute(f"SELECT w_id FROM win_lost_results WHERE a_id IN ({a_placeholders})", a_ids)
                    w_ids = [row['w_id'] for row in cur.fetchall() if row.get('w_id')]
                    
                    if w_ids:
                        w_placeholders = ','.join(['%s'] * len(w_ids))
                        try:
                            cur.execute(f"SELECT won_id FROM won_bids_result WHERE w_id IN ({w_placeholders})", w_ids)
                            won_ids = [row['won_id'] for row in cur.fetchall() if row.get('won_id')]
                            
                            if won_ids:
                                won_placeholders = ','.join(['%s'] * len(won_ids))
                                try:
                                    cur.execute(f"DELETE FROM work_progress_status WHERE won_id IN ({won_placeholders})", won_ids)
                                except Exception:
                                    pass
                            
                            try:
                                cur.execute(f"DELETE FROM won_bids_result WHERE w_id IN ({w_placeholders})", w_ids)
                            except Exception:
                                pass
                        except Exception:
                            pass
                    
                    try:
                        cur.execute(f"DELETE FROM win_lost_results WHERE a_id IN ({a_placeholders})", a_ids)
                    except Exception:
                        pass
                except Exception:
                    pass
            
            # Finally, delete from go_bids
            cur.execute(
                """
                DELETE gb FROM go_bids gb
                LEFT JOIN bid_incoming bi ON bi.id = gb.id
                WHERE gb.id IS NOT NULL
                  AND (
                        bi.id IS NULL
                     OR UPPER(TRIM(COALESCE(bi.decision, ''))) <> 'GO'
                  )
                """
            )

        mysql.connection.commit()
        return inserted
    except Exception as sync_err:
        mysql.connection.rollback()
        print(f"Error synchronizing GO bids: {sync_err}")
        return 0
    finally:
        cur.close()

# --- Authentication Routes ---
@app.route('/register', methods=['GET', 'POST'])
def register():
    cur = mysql.connection.cursor(DictCursor)
    
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '').strip()
        role = request.form.get('role', 'member').strip() or 'member'
        company_id = request.form.get('company_id', '').strip()
        
        if not email or not password:
            cur.execute("SELECT id, name FROM companies ORDER BY name")
            companies = cur.fetchall()
            cur.close()
            return render_template('register.html', companies=companies, error='Email and password are required')
        
        # Check if email already exists
        cur.execute("SELECT id FROM users WHERE email=%s", (email,))
        if cur.fetchone():
            cur.execute("SELECT id, name FROM companies ORDER BY name")
            companies = cur.fetchall()
            cur.close()
            return render_template('register.html', companies=companies, error='Email already exists. Please use a different email or login.')
        
        is_admin = 1 if role.lower() == "admin" else 0
        cur.execute("INSERT INTO users (email, password, role, is_admin) VALUES (%s, %s, %s, %s)",
                    (email, password, role, is_admin))
        mysql.connection.commit()
        
        # Fetch companies for the template (even after success)
        cur.execute("SELECT id, name FROM companies ORDER BY name")
        companies = cur.fetchall()
        for company in companies:
            company_name_lower = company['name'].lower()
            if 'ikio' in company_name_lower:
                company['logo_path'] = 'ikio.png'
            elif 'metco' in company_name_lower:
                company['logo_path'] = 'metco.png'
            elif 'sunsprint' in company_name_lower:
                company['logo_path'] = 'sunsprint.png'
            else:
                company['logo_path'] = None
        
        cur.close()
        
        return render_template('register.html', 
                             companies=companies, 
                             success='Account created successfully! You can now login.',
                             show_login=True)
    
    # GET request - show registration form
    cur.execute("SELECT id, name FROM companies ORDER BY name")
    companies = cur.fetchall()
    
    # Add logo paths based on company name (fallback if logo_path column doesn't exist)
    for company in companies:
        company_name_lower = company['name'].lower()
        if 'ikio' in company_name_lower:
            company['logo_path'] = 'ikio.png'
        elif 'metco' in company_name_lower:
            company['logo_path'] = 'metco.png'
        elif 'sunsprint' in company_name_lower:
            company['logo_path'] = 'sunsprint.png'
        else:
            company['logo_path'] = None
    
    cur.close()
    return render_template('register.html', companies=companies)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        cur = mysql.connection.cursor(DictCursor)
        cur.execute("SELECT * FROM users WHERE email=%s AND password=%s", (email, password))
        user = cur.fetchone()
        if user:
            cur.close()
            user_obj = User(user['id'], user['email'], user['password'], user['is_admin'], user['role'])
            login_user(user_obj)
            log_write('login', f"role={user['role']}")
            if user['is_admin']:
                return redirect(url_for('master_dashboard'))
            
            # Role-based redirects
            role = user['role'].lower()
            if role == 'supervisor':
                return redirect(url_for('supervisor_dashboard'))
            elif role == 'business dev':
                return redirect(url_for('team_dashboard', team='business'))
            elif role == 'design':
                return redirect(url_for('team_dashboard', team='design'))
            elif role == 'operations':
                return redirect(url_for('team_dashboard', team='operations'))
            elif role == 'site manager':
                return redirect(url_for('team_dashboard', team='engineer'))
            else:
                return redirect(url_for('dashboard'))
        # Fallback: try authenticating as employee
        try:
            # Ensure password column exists before querying
            cur.execute("SHOW COLUMNS FROM employees LIKE 'password'")
            if cur.fetchone() is None:
                cur.close()
                return 'Employee login unavailable: employees.password missing', 500
            cur.execute(
                """
                SELECT * FROM employees
                WHERE email=%s AND password=%s AND is_active=1
                """,
                (email, password)
            )
            employee = cur.fetchone()
        finally:
            cur.close()

        if employee:
            session['employee_id'] = employee['id']
            session['employee_name'] = employee['name']
            session['employee_email'] = employee['email']
            session['employee_department'] = employee['department']
            log_write('employee_login', f"Employee {employee['name']} logged in via /login")
            return redirect(url_for('employee_dashboard', employee_id=employee['id']))

        return 'Invalid credentials', 401
    return render_template('login.html')

@app.route('/admin/users', methods=['GET', 'POST'])
@login_required
def users_admin():
    if not current_user.is_admin:
        return "Access Denied", 403
    cur = mysql.connection.cursor(DictCursor)
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '').strip()
        role = request.form.get('role', 'member').strip() or 'member'
        if not email or not password:
            return 'Email and password required', 400
        # Check if email already exists
        cur.execute("SELECT id FROM users WHERE email=%s", (email,))
        if cur.fetchone():
            cur.close()
            return 'Email already exists', 400
        is_admin = 1 if role.lower() == "admin" else 0
        cur.execute("INSERT INTO users (email, password, role, is_admin) VALUES (%s, %s, %s, %s)",
                    (email, password, role, is_admin))
        mysql.connection.commit()
        cur.close()
        return redirect(url_for('users_admin'))
    cur.execute("SELECT * FROM users ORDER BY id ASC")
    users = cur.fetchall()
    cur.close()
    return render_template('users.html', users=users)

@app.route('/logout')
def logout():
    log_write('logout')
    logout_user()
    return redirect(url_for('login'))

@app.route('/employee/login', methods=['POST'])
def employee_login():
    """Employee login route"""
    try:
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '').strip()
        # employee_id may be provided from the UI but we don't require it for auth
        if not email or not password:
            return jsonify({'success': False, 'error': 'Missing credentials'}), 400

        cur = mysql.connection.cursor(DictCursor)
        # Ensure password column exists; if not, fail with a clear message
        try:
            cur.execute("SHOW COLUMNS FROM employees LIKE 'password'")
            has_password_col = cur.fetchone() is not None
        except Exception:
            has_password_col = False
        if not has_password_col:
            cur.close()
            return jsonify({'success': False, 'error': "employees.password column missing. Run add_password_column.sql or migration."}), 500
        cur = mysql.connection.cursor(DictCursor)
        cur.execute(
            """
            SELECT * FROM employees 
            WHERE email = %s AND password = %s AND is_active = 1
            """,
            (email, password),
        )
        employee = cur.fetchone()
        cur.close()

        if employee:
            # Create a simple session for employee (not using Flask-Login for employees)
            session['employee_id'] = employee['id']
            session['employee_name'] = employee['name']
            session['employee_email'] = employee['email']
            session['employee_department'] = employee['department']

            log_write('employee_login', f"Employee {employee['name']} logged in")
            return jsonify({'success': True, 'employee_id': employee['id']})
        else:
            return jsonify({'success': False, 'error': 'Invalid credentials'}), 401

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/employee/logout')
def employee_logout():
    """Employee logout route"""
    if 'employee_id' in session:
        log_write('employee_logout', f"Employee {session.get('employee_name', 'Unknown')} logged out")
        session.pop('employee_id', None)
        session.pop('employee_name', None)
        session.pop('employee_email', None)
        session.pop('employee_department', None)
    return redirect(url_for('login'))

# --- Dashboard Routes ---
@app.route('/')
def index():
    if current_user.is_authenticated:
        if current_user.is_admin:
            return redirect(url_for('master_dashboard'))
        elif current_user.is_supervisor:
            return redirect(url_for('supervisor_dashboard'))
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/dashboard')
@login_required
def dashboard():
    # Redirect non-admins to their role-specific dashboards
    if current_user.is_admin:
        return redirect(url_for('master_dashboard'))
    elif current_user.is_supervisor:
        return redirect(url_for('supervisor_dashboard'))
    return redirect(url_for('role_dashboard'))

@app.route('/dashboard/role')
@login_required
def role_dashboard():
    if current_user.is_admin:
        return redirect(url_for('master_dashboard'))
    role = (current_user.role or 'member').strip()
    role_key = role.lower()

    # Map business flow stages to roles
    role_to_stage = {
        'business dev': 'business',
        'business': 'business',  # backward compatibility
        'bdm': 'business',
        'design': 'design',
        'operations': 'operations',
        'site manager': 'site_manager'
    }
    next_stage_map = {
        'analyzer': 'business',
        'business': 'design',
        'design': 'operations',
        'operations': 'site_manager',
        'site_manager': 'handover'
    }

    current_stage_for_role = role_to_stage.get(role_key)
    cur = mysql.connection.cursor(DictCursor)
    bids = []
    # Query by go_bids.state (not bid_assign.depart) and join assignee info
    role_to_stage = {
        'business dev': 'business',
        'business': 'business', 
        'bdm': 'business',
        'design': 'design',
        'operations': 'operations',
        'site manager': 'engineer',
        'site_manager': 'engineer'
    }
    current_stage_for_role = role_to_stage.get((current_user.role or 'member').lower())
    
    sql = """
        SELECT gb.g_id AS id,
               gb.b_name AS name,
               gb.company, 
               gb.due_date,
               COALESCE(gb.scoring, 0) AS progress,
               LOWER(COALESCE(gb.state, 'analyzer')) AS current_stage,
               ba.person_name,
               ba.assignee_email AS person_email,
               gb.summary
        FROM go_bids gb
        LEFT JOIN bid_assign ba ON ba.g_id = gb.g_id
        {where}
        ORDER BY gb.due_date ASC
    """
    where = "WHERE LOWER(COALESCE(gb.state, 'analyzer')) = %s" if current_stage_for_role else ""
    cur.execute(sql.format(where=where), (current_stage_for_role,) if current_stage_for_role else ())
    bids = cur.fetchall()
    # Keep next_stage mapping in Python and pass it to the template
    next_stage_map = {
        'analyzer': 'business',
        'business': 'design', 
        'design': 'operations',
        'operations': 'engineer',
        'engineer': 'handover'
    }
    for bid in bids:
        bid['user'] = {'email': bid.get('person_email'), 'role': current_stage_for_role}
        bid['next_stage'] = next_stage_map.get(bid['current_stage'])
        # Add dynamic progress and status texts
        stage_key = (bid.get('current_stage') or 'analyzer').lower()
        bid['work_progress_pct'] = stage_progress_pct(stage_key)
        bid['project_status'], bid['work_status'] = status_texts(stage_key)
    # Build dynamic stage lists per bid similar to supervisor view
    # Ensure tables exist
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS bid_stage_exclusions (
            id INT AUTO_INCREMENT PRIMARY KEY,
            g_id INT NOT NULL,
            stage VARCHAR(50) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE KEY uniq_bid_stage (g_id, stage)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS bid_custom_stages (
            id INT AUTO_INCREMENT PRIMARY KEY,
            g_id INT NOT NULL,
            stage VARCHAR(50) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE KEY uniq_custom_stage (g_id, stage)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """
    )
    # Load exclusions and customs
    cur.execute("SELECT g_id, stage FROM bid_stage_exclusions")
    ex_rows = cur.fetchall()
    exclusions = {}
    for r in ex_rows:
        exclusions.setdefault(r['g_id'], set()).add((r['stage'] or '').strip().lower())
    cur.execute("SELECT g_id, stage FROM bid_custom_stages")
    cs_rows = cur.fetchall()
    customs = {}
    for r in cs_rows:
        customs.setdefault(r['g_id'], []).append((r['stage'] or '').strip().lower())
    default_stages = ['analyzer', 'business', 'design', 'operations', 'engineer']
    for bid in bids:
        bid_id = bid.get('id')
        excl = exclusions.get(bid_id, set())
        cust = [s for s in customs.get(bid_id, []) if s not in excl]
        dyn_stages = [s for s in default_stages if s not in excl] + [s for s in cust if s not in default_stages]
        bid['dyn_stages'] = dyn_stages

    cur.close()

    # Use a generic role dashboard that includes an advance-stage button
    normalized_role = ' '.join([w.capitalize() for w in role_key.split()])
    return render_template('dashboard_role.html', bids=bids, user=current_user, role=normalized_role,
                           current_stage=current_stage_for_role, next_stage_map=next_stage_map)

@app.route('/dashboard/<team>')
@login_required
def team_dashboard(team):
    """Team-specific dashboard for Business Dev, Design, Operations, Site Manager"""
    if current_user.is_admin:
        return redirect(url_for('master_dashboard'))
    
    # Keep go_bids in sync with GO decisions so teams see new bids immediately
    try:
        sync_go_bids()
    except Exception:
        pass

    # Map team names to stages
    team_to_stage = {
        'business': 'business',
        'design': 'design', 
        'operations': 'operations',
        'engineer': 'engineer'
    }
    
    if team not in team_to_stage:
        return "Invalid team", 404
    
    current_stage = team_to_stage[team]
    # Ensure seed tasks exist so all bids appear on this team dashboard
    ensure_tasks_for_team(current_stage)
    cur = mysql.connection.cursor(DictCursor)
    
    # Company access filter based on user email domain
    email = (current_user.email or '').lower()
    company_filter_name = None
    if email.endswith('@ikioledlighting.com'):
        company_filter_name = 'Ikio'
    elif email.endswith('@metcoengineering.com'):
        company_filter_name = 'Metco'
    elif email.endswith('@sunsprintengineering.com'):
        company_filter_name = 'Sunsprint'

    company_sql_clause = ""
    company_params = tuple()
    if company_filter_name:
        company_sql_clause = " AND LOWER(COALESCE(gb.company,'')) = %s"
        company_params = (company_filter_name.lower(),)

    # Fetch all bids for this company scope; filtering by team is done in Python
    cur.execute("""
        SELECT gb.g_id AS id,
               gb.b_name AS name,
               gb.company, 
               gb.due_date,
               COALESCE(gb.scoring, 0) AS progress,
               LOWER(COALESCE(gb.state, 'analyzer')) AS current_stage,
               ba.person_name,
               ba.assignee_email AS person_email,
               ba.depart,
               wps.pr_completion_status AS work_status,
               wbr.closure_status AS project_status,
               wbr.work_progress_status AS work_progress_status,
               wlr.result AS wl_result,
               gb.summary
        FROM go_bids gb
        LEFT JOIN bid_assign ba ON ba.g_id = gb.g_id
        LEFT JOIN win_lost_results wlr ON wlr.a_id = ba.a_id
        LEFT JOIN won_bids_result wbr ON wbr.w_id = wlr.w_id
        LEFT JOIN work_progress_status wps ON wps.won_id = wbr.won_id
        WHERE 1=1
        """ + company_sql_clause + """
        ORDER BY gb.due_date ASC
    """, company_params)

    rows = cur.fetchall()

    # Ensure dynamic stage tables exist and load per-bid configuration
    try:
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_stage_exclusions (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_bid_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_custom_stages (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_custom_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
    except Exception:
        # If creation fails (e.g., no rights), continue – downstream selects may still work
        pass

    cur.execute("SELECT g_id, stage FROM bid_stage_exclusions")
    ex_rows = cur.fetchall()
    exclusions = {}
    for r in ex_rows:
        exclusions.setdefault(r['g_id'], set()).add((r['stage'] or '').strip().lower())

    cur.execute("SELECT g_id, stage FROM bid_custom_stages")
    cs_rows = cur.fetchall()
    customs = {}
    for r in cs_rows:
        customs.setdefault(r['g_id'], []).append((r['stage'] or '').strip().lower())

    # Preload set of bids that already have tasks for this team's stage
    try:
        cur.execute("SELECT DISTINCT g_id FROM bid_checklists WHERE LOWER(COALESCE(stage,''))=%s", (current_stage,))
        task_ids = {row['g_id'] for row in cur.fetchall()}
    except Exception:
        task_ids = set()

    default_stages = ['analyzer', 'business', 'design', 'operations', 'engineer']
    bids = []
    for row in rows:
        bid_id = row.get('id')
        excl = exclusions.get(bid_id, set())
        cust = [s for s in customs.get(bid_id, []) if s not in excl]
        dyn_stages = [s for s in default_stages if s not in excl] + [s for s in cust if s not in default_stages]
        # If no dynamic config exists, keep default stages
        if not dyn_stages:
            dyn_stages = default_stages.copy()

        stage_now = (row.get('current_stage') or 'analyzer').lower()
        should_include = (
            stage_now == current_stage
            or bid_id in task_ids
            or current_stage in dyn_stages
        )
        if not should_include:
            continue

        # Shape bid dict like before and attach dynamic stages
        row['user'] = {'email': row.get('person_email'), 'role': current_stage}
        row['next_stage'] = None  # resolved in template via get_next_stage
        stage_key = stage_now
        row['work_progress_pct'] = stage_progress_pct(stage_key)
        row['project_status'], row['work_status'] = status_texts(stage_key)
        row['dyn_stages'] = dyn_stages
        bids.append(row)

    # Compute stats from filtered bids
    total_bids = len(bids)
    completed_bids = sum(1 for b in bids if (b.get('wl_result') == 'WON'))

    # Load assigned members for each bid from bid_assignment_members table
    bid_ids = [b['id'] for b in bids]
    assigned_map = {}
    try:
        if bid_ids:
            # Ensure mapping table exists
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS bid_assignment_members (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    g_id INT NOT NULL,
                    employee_id INT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE KEY uniq_g_emp (g_id, employee_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
                """
            )
            placeholders = ','.join(['%s'] * len(bid_ids))
            cur.execute(
                f"""
                SELECT bam.g_id, e.id AS employee_id, e.name, e.email
                FROM bid_assignment_members bam
                JOIN employees e ON e.id = bam.employee_id
                WHERE bam.g_id IN ({placeholders}) AND e.department = %s AND e.is_active = TRUE
                ORDER BY e.name
                """,
                (*bid_ids, team)
            )
            rows = cur.fetchall()
            for r in rows:
                assigned_map.setdefault(r['g_id'], []).append({'id': r['employee_id'], 'name': r['name'], 'email': r['email']})
        # Attach to bid rows
        for b in bids:
            b['assigned_members'] = assigned_map.get(b['id'], [])
    except Exception:
        # If there's an error, just set empty list for each bid
        for b in bids:
            b['assigned_members'] = []
    
    # Load RFP files for each bid from uploaded_rfp_files table
    rfp_files_map = {}
    try:
        if bid_ids:
            # Ensure uploaded_rfp_files table exists
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS uploaded_rfp_files (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    bid_id INT,
                    g_id INT,
                    filename VARCHAR(500) NOT NULL,
                    file_path VARCHAR(1000) NOT NULL,
                    file_size BIGINT,
                    uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    uploaded_by INT,
                    INDEX idx_bid_id (bid_id),
                    INDEX idx_g_id (g_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
                """
            )
            # Check if g_id column exists, if not add it
            try:
                cur.execute("SHOW COLUMNS FROM uploaded_rfp_files LIKE 'g_id'")
                if not cur.fetchone():
                    cur.execute("ALTER TABLE uploaded_rfp_files ADD COLUMN g_id INT, ADD INDEX idx_g_id (g_id)")
                    mysql.connection.commit()
            except Exception:
                pass
            
            placeholders = ','.join(['%s'] * len(bid_ids))
            # Try to query with g_id first, fallback to bid_id if g_id doesn't exist in results
            try:
                cur.execute(
                    f"""
                    SELECT id, g_id, bid_id, filename, file_path, file_size, uploaded_at
                    FROM uploaded_rfp_files
                    WHERE g_id IN ({placeholders}) OR bid_id IN ({placeholders})
                    ORDER BY uploaded_at DESC
                    """,
                    (*bid_ids, *bid_ids)
                )
            except Exception:
                # Fallback if g_id column doesn't exist in query results
                cur.execute(
                    f"""
                    SELECT id, bid_id, filename, file_path, file_size, uploaded_at
                    FROM uploaded_rfp_files
                    WHERE bid_id IN ({placeholders})
                    ORDER BY uploaded_at DESC
                    """,
                    bid_ids
                )
            rows = cur.fetchall()
            for r in rows:
                # bid_ids contains g_id values from go_bids (aliased as 'id')
                g_id = r.get('g_id') or r.get('bid_id')
                if g_id and g_id in bid_ids:
                    if g_id not in rfp_files_map:
                        rfp_files_map[g_id] = []
                    rfp_files_map[g_id].append({
                        'id': r['id'],
                        'filename': r['filename'],
                        'file_path': r['file_path'],
                        'file_size': r.get('file_size', 0),
                        'uploaded_at': r.get('uploaded_at')
                    })
        # Attach RFP files to bid rows
        for b in bids:
            b['rfp_files'] = rfp_files_map.get(b['id'], [])
    except Exception as e:
        # If there's an error, just set empty list for each bid
        print(f"Error loading RFP files: {str(e)}")
        for b in bids:
            b['rfp_files'] = []

    cur.close()
    
    # Map stage names for display
    stage_display_names = {
        'business': 'Business Development',
        'design': 'Design Team', 
        'operations': 'Operations Team',
        'engineer': 'Site Engineer'
    }
    
    team_display_name = stage_display_names.get(team, team.title())
    
    # Add dynamic progress and status texts to each bid
    for bid in bids:
        stage_key = (bid.get('current_stage') or 'analyzer').lower()
        bid['work_progress_pct'] = stage_progress_pct(stage_key)
        bid['project_status'], bid['work_status'] = status_texts(stage_key)
    
    # Define next stage mapping for template
    def get_next_stage(current_stage):
        stage_flow = {
            'analyzer': 'business',
            'business': 'design',
            'design': 'operations',
            'operations': 'engineer',
            'engineer': 'handover'
        }
        return stage_flow.get(current_stage, None)
    
    return render_template('team_dashboard.html', 
                         bids=bids, 
                         team=team,
                         team_display_name=team_display_name,
                         current_stage=current_stage,
                         total_bids=total_bids,
                         completed_bids=completed_bids,
                         user=current_user,
                         get_next_stage=get_next_stage)

main_stats = {
    'total_bids': 350, 'live_bids': 75, 'bids_won': 120, 'projects_completed': 95,
}

# Assuming you have a Log model, e.g., from SQLAlchemy


# ... other imports and app setup

@app.route('/logs')
@login_required
def logs_page():
    # Query all logs, ordering by the most recent first
    cur = mysql.connection.cursor(DictCursor)
    cur.execute("SELECT * FROM logs ORDER BY timestamp DESC")
    all_logs = cur.fetchall()
    cur.close()
    return render_template('logs.html', logs=all_logs)

@app.route('/supervisor-dashboard')
@login_required
def supervisor_dashboard():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    
    cur = mysql.connection.cursor(DictCursor)
    # Ensure exclusions table exists
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS bid_stage_exclusions (
            id INT AUTO_INCREMENT PRIMARY KEY,
            g_id INT NOT NULL,
            stage VARCHAR(50) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE KEY uniq_bid_stage (g_id, stage)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """
    )

    # Ensure custom stages table exists
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS bid_custom_stages (
            id INT AUTO_INCREMENT PRIMARY KEY,
            g_id INT NOT NULL,
            stage VARCHAR(50) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE KEY uniq_custom_stage (g_id, stage)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """
    )

    # Auto-sync GO bids from bid_incoming to go_bids table
    try:
        cur.execute("""
            INSERT INTO go_bids (id, b_name, due_date, state, scope, type, scoring, company, decision, summary)
            SELECT bi.id, bi.b_name, bi.due_date, bi.state, bi.scope, bi.type, bi.scoring, bi.comp_name, bi.decision, bi.summary
            FROM bid_incoming bi
            LEFT JOIN go_bids gb ON gb.id = bi.id
            WHERE UPPER(bi.decision) = 'GO' AND gb.id IS NULL
        """)
        mysql.connection.commit()
    except Exception as e:
        mysql.connection.rollback()
        # Log error but continue - don't break the dashboard
        pass

    # Get all GO bids for stage assignment (including summary and scope)
    cur.execute("""
        SELECT gb.g_id, gb.b_name, gb.state, gb.company, gb.type, gb.summary, gb.scope,
               ba.person_name, ba.depart
        FROM go_bids gb
        LEFT JOIN bid_assign ba ON ba.g_id = gb.g_id
        ORDER BY gb.g_id DESC
    """)
    go_bids = cur.fetchall()
    # Load companies for dropdowns
    cur.execute("SELECT id, name FROM companies ORDER BY name ASC")
    companies = cur.fetchall()
    # Fetch stage exclusions per bid
    cur.execute("SELECT g_id, stage FROM bid_stage_exclusions")
    rows = cur.fetchall()
    stage_exclusions = {}
    for r in rows:
        stage_exclusions.setdefault(r['g_id'], []).append(r['stage'])

    # Fetch custom stages per bid
    cur.execute("SELECT g_id, stage FROM bid_custom_stages")
    rows = cur.fetchall()
    custom_stages = {}
    for r in rows:
        custom_stages.setdefault(r['g_id'], []).append(r['stage'])

    # Get stage options
    stages = ['analyzer', 'business', 'design', 'operations', 'engineer']

    # Attach dynamic stages and progress to each bid row for table rendering
    default_stages = stages.copy()
    for bid in go_bids:
        bid_id = bid['g_id']
        excluded = set(stage_exclusions.get(bid_id, []))
        customs = [s for s in custom_stages.get(bid_id, []) if s not in excluded]
        dyn_stages = [s for s in default_stages if s not in excluded] + [s for s in customs if s not in default_stages]
        current_stage = (bid.get('state') or 'analyzer').strip().lower()
        if dyn_stages and current_stage in dyn_stages:
            idx = dyn_stages.index(current_stage)
            pct = int(round((idx / (max(1, len(dyn_stages) - 1))) * 100))
        else:
            pct = 0
        bid['dyn_stages'] = dyn_stages
        bid['dyn_progress_pct'] = pct
    
    cur.close()
    
    return render_template('supervisor_dashboard.html', 
                         go_bids=go_bids, 
                         stages=stages,
                         stage_exclusions=stage_exclusions,
                         custom_stages=custom_stages,
                         companies=companies,
                         user=current_user)
@app.route('/supervisor-projects')
@login_required
def supervisor_projects():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    cur = mysql.connection.cursor(DictCursor)
    cur.execute(
        """
        SELECT wlr.w_id, wlr.a_id, wlr.b_name, wlr.due_date,
               wlr.state, wlr.scope, wlr.value, wlr.company, wlr.department,
               wlr.person_name, wlr.status, wlr.result
        FROM win_lost_results wlr
        ORDER BY wlr.w_id DESC
        """
    )
    projects = cur.fetchall()
    cur.close()
    return render_template('supervisor_projects.html', projects=projects, user=current_user)
@app.route('/supervisor/assign-stage', methods=['POST'])
@login_required
def supervisor_assign_stage():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    
    g_id = request.form.get('g_id')
    new_stage = request.form.get('stage')
    person_name = request.form.get('person_name', '').strip()
    department = request.form.get('department', '').strip()
    
    if not g_id or not new_stage:
        flash('Missing required fields', 'error')
        return redirect(url_for('supervisor_dashboard'))
    
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Update go_bids state
        cur.execute("UPDATE go_bids SET state=%s WHERE g_id=%s", (new_stage, g_id))
        
        # Update or create bid_assign entry
        cur.execute("SELECT a_id FROM bid_assign WHERE g_id=%s", (g_id,))
        row = cur.fetchone()
        
        if row:
            # Update existing assignment
            cur.execute("""
                UPDATE bid_assign 
                SET state=%s, depart=%s, person_name=%s, status='assigned' 
                WHERE g_id=%s
            """, (new_stage, department, person_name, g_id))
        else:
            # Create new assignment
            cur.execute("""
                INSERT INTO bid_assign (g_id, b_name, due_date, state, scope, type, company, depart, person_name, status)
                SELECT g_id, b_name, due_date, %s, scope, type, company, %s, %s, 'assigned'
                FROM go_bids WHERE g_id=%s
            """, (new_stage, department, person_name, g_id))
        
        # Log the action
        cur.execute("SELECT b_name FROM go_bids WHERE g_id=%s", (g_id,))
        bid_name = cur.fetchone()['b_name']
        log_action = f"Supervisor '{current_user.email}' assigned bid '{bid_name}' (ID: {g_id}) to {new_stage} stage"
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)", (log_action, current_user.id))
        
        mysql.connection.commit()
        flash(f'Bid "{bid_name}" assigned to {new_stage} stage successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Assignment failed: {str(e)}', 'error')
    finally:
        cur.close()
    
    return redirect(url_for('supervisor_dashboard'))

@app.route('/supervisor/update-company', methods=['POST'])
@login_required
def supervisor_update_company():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    g_id = request.form.get('g_id')
    company_id = request.form.get('company_id')
    if not g_id:
        flash('Missing bid id', 'error')
        return redirect(url_for('supervisor_dashboard'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        company_name = None
        if company_id:
            cur.execute("SELECT name FROM companies WHERE id=%s", (company_id,))
            row = cur.fetchone()
            if row:
                company_name = row['name']
        if not company_name:
            # allow free text fallback
            company_name = (request.form.get('company_name') or '').strip()
        cur.execute("UPDATE go_bids SET company=%s WHERE g_id=%s", (company_name, g_id))
        mysql.connection.commit()
        flash('Company updated.', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Failed to update company: {e}', 'error')
    finally:
        cur.close()
    return redirect(url_for('supervisor_dashboard'))

@app.route('/uploads/<path:subpath>')
@login_required
def serve_uploads(subpath):
    base_dir = os.path.join(os.getcwd(), 'uploads')
    return send_from_directory(base_dir, subpath)

# [REMOVED duplicate /profiling handler consolidated above]

@app.route('/team/stage/add', methods=['POST'])
@login_required
def team_add_stage():
    """Allow team members to add a custom stage for a bid without supervisor role.
    The stage is stored in bid_custom_stages and un-excluded if previously excluded.
    """
    g_id = request.form.get('g_id')
    stage = (request.form.get('new_stage') or '').strip().lower()
    if not g_id or not stage:
        return jsonify({'ok': False, 'error': 'Missing parameters'}), 400
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Ensure required tables exist
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_stage_exclusions (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_bid_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_custom_stages (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_custom_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
        cur.execute("INSERT IGNORE INTO bid_custom_stages (g_id, stage) VALUES (%s, %s)", (g_id, stage))
        cur.execute("DELETE FROM bid_stage_exclusions WHERE g_id=%s AND stage=%s", (g_id, stage))
        mysql.connection.commit()
        return jsonify({'ok': True})
    except Exception as e:
        mysql.connection.rollback()
        return jsonify({'ok': False, 'error': str(e)}), 500
    finally:
        cur.close()

@app.route('/team/stage/delete', methods=['POST'])
@login_required
def team_delete_stage():
    """Allow team members to delete a stage that comes after the current stage."""
    g_id = request.form.get('g_id')
    stage = (request.form.get('stage') or '').strip().lower()
    if not g_id or not stage:
        return jsonify({'ok': False, 'error': 'Missing parameters'}), 400
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Ensure required tables exist
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_stage_exclusions (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_bid_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_custom_stages (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_custom_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
        # Determine current stage and dynamic stage ordering
        cur.execute("SELECT state FROM go_bids WHERE g_id=%s", (g_id,))
        row = cur.fetchone()
        current_stage = (row.get('state') or 'analyzer').strip().lower() if row else 'analyzer'

        cur.execute("SELECT stage FROM bid_stage_exclusions WHERE g_id=%s", (g_id,))
        excluded = { (r['stage'] or '').strip().lower() for r in cur.fetchall() }
        cur.execute("SELECT stage FROM bid_custom_stages WHERE g_id=%s", (g_id,))
        customs = [ (r['stage'] or '').strip().lower() for r in cur.fetchall() ]
        default_stages = ['analyzer', 'business', 'design', 'operations', 'engineer']
        stages = [s for s in default_stages if s not in excluded] + [s for s in customs if s not in excluded and s not in default_stages]

        if stage not in stages:
            return jsonify({'ok': False, 'error': 'Stage not found for this bid.'}), 400
        try:
            curr_idx = stages.index(current_stage)
        except ValueError:
            curr_idx = 0
        target_idx = stages.index(stage)
        if target_idx <= curr_idx:
            return jsonify({'ok': False, 'error': 'Cannot delete current or previous stages.'}), 400

        cur.execute("INSERT IGNORE INTO bid_stage_exclusions (g_id, stage) VALUES (%s, %s)", (g_id, stage))
        cur.execute("DELETE FROM bid_custom_stages WHERE g_id=%s AND stage=%s", (g_id, stage))
        mysql.connection.commit()
        return jsonify({'ok': True})
    except Exception as e:
        mysql.connection.rollback()
        return jsonify({'ok': False, 'error': str(e)}), 500
    finally:
        cur.close()

@app.route('/supervisor/delete-stage', methods=['POST'])
@login_required
def supervisor_delete_stage():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    g_id = request.form.get('g_id')
    deleted_stage = request.form.get('stage')
    if not g_id:
        flash('Missing bid id', 'error')
        return redirect(url_for('supervisor_dashboard'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Remove explicit assignment for this bid; keep current go_bids.state unchanged
        cur.execute("DELETE FROM bid_assign WHERE g_id=%s", (g_id,))
        # Persist exclusion so UI can hide this stage next time
        if deleted_stage:
            cur.execute(
                "INSERT IGNORE INTO bid_stage_exclusions (g_id, stage) VALUES (%s, %s)",
                (g_id, deleted_stage)
            )
        # Log action
        cur.execute("SELECT b_name FROM go_bids WHERE g_id=%s", (g_id,))
        row = cur.fetchone()
        bid_name = row['b_name'] if row else g_id
        stage_info = f" (stage: {deleted_stage})" if deleted_stage else ""
        log_action = f"Supervisor '{current_user.email}' deleted stage assignment for bid '{bid_name}' (ID: {g_id}){stage_info}"
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)", (log_action, current_user.id))
        mysql.connection.commit()
        flash('Stage assignment deleted.', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Could not delete stage: {e}', 'error')
    finally:
        cur.close()
    return redirect(url_for('supervisor_dashboard'))
@app.route('/supervisor/delete-bid', methods=['POST'])
@login_required
def supervisor_delete_bid():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    g_id = request.form.get('g_id')
    if not g_id:
        flash('Missing bid id', 'error')
        return redirect(url_for('supervisor_dashboard'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Fetch name for logs
        cur.execute("SELECT b_name FROM go_bids WHERE g_id=%s", (g_id,))
        row = cur.fetchone()
        bid_name = (row or {}).get('b_name', g_id)

        # Delete related records to avoid foreign key constraint errors
        cur.execute("DELETE FROM bid_checklists WHERE g_id=%s", (g_id,))
        cur.execute("DELETE FROM bid_stage_exclusions WHERE g_id=%s", (g_id,))
        cur.execute("DELETE FROM bid_custom_stages WHERE g_id=%s", (g_id,))
        
        # Delete from bid_assign and cascading tables
        cur.execute("SELECT a_id FROM bid_assign WHERE g_id=%s", (g_id,))
        assign_ids = [r['a_id'] for r in cur.fetchall()]
        
        for a_id in assign_ids:
            cur.execute("SELECT w_id FROM win_lost_results WHERE a_id=%s", (a_id,))
            w_ids = [r['w_id'] for r in cur.fetchall()]
            
            for w_id in w_ids:
                cur.execute("SELECT won_id FROM won_bids_result WHERE w_id=%s", (w_id,))
                won_ids = [r['won_id'] for r in cur.fetchall()]
                
                for won_id in won_ids:
                    cur.execute("DELETE FROM work_progress_status WHERE won_id=%s", (won_id,))
                
                cur.execute("DELETE FROM won_bids_result WHERE w_id=%s", (w_id,))
            
            cur.execute("DELETE FROM win_lost_results WHERE a_id=%s", (a_id,))
        
        cur.execute("DELETE FROM bid_assign WHERE g_id=%s", (g_id,))
        cur.execute("DELETE FROM go_bids WHERE g_id=%s", (g_id,))

        # Log
        log_action = f"Supervisor '{current_user.email}' deleted bid '{bid_name}' (ID: {g_id})"
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)", (log_action, current_user.id))
        mysql.connection.commit()
        flash('Bid deleted successfully.', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Failed to delete bid: {e}', 'error')
    finally:
        cur.close()
    return redirect(url_for('supervisor_dashboard'))

@app.route('/supervisor/set-stage', methods=['POST'])
@login_required
def supervisor_set_stage():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    g_id = request.form.get('g_id')
    new_stage = (request.form.get('new_stage') or '').strip().lower()
    if not g_id or not new_stage:
        flash('Missing bid id or stage.', 'error')
        return redirect(url_for('supervisor_dashboard'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Update go_bids current stage only
        cur.execute("UPDATE go_bids SET state=%s WHERE g_id=%s", (new_stage, g_id))
        # Log
        cur.execute("SELECT b_name FROM go_bids WHERE g_id=%s", (g_id,))
        bid_name = (cur.fetchone() or {}).get('b_name', g_id)
        log_action = f"Supervisor '{current_user.email}' set stage to '{new_stage}' for bid '{bid_name}' (ID: {g_id})"
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)", (log_action, current_user.id))
        mysql.connection.commit()
        flash('Stage updated.', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Could not update stage: {e}', 'error')
    finally:
        cur.close()
    return redirect(url_for('supervisor_dashboard'))

@app.route('/supervisor/add-stage', methods=['POST'])
@login_required
def supervisor_add_stage():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    g_id = request.form.get('g_id')
    stage = (request.form.get('new_stage') or '').strip().lower()
    if not g_id or not stage:
        flash('Bid id and stage name are required.', 'error')
        return redirect(url_for('supervisor_dashboard'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        cur.execute("INSERT IGNORE INTO bid_custom_stages (g_id, stage) VALUES (%s, %s)", (g_id, stage))
        # If stage was previously excluded, un-exclude it
        cur.execute("DELETE FROM bid_stage_exclusions WHERE g_id=%s AND stage=%s", (g_id, stage))
        cur.execute("SELECT b_name FROM go_bids WHERE g_id=%s", (g_id,))
        bid_name = (cur.fetchone() or {}).get('b_name', g_id)
        log_action = f"Supervisor '{current_user.email}' added custom stage '{stage}' for bid '{bid_name}' (ID: {g_id})"
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)", (log_action, current_user.id))
        mysql.connection.commit()
        flash('Stage added.', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Could not add stage: {e}', 'error')
    finally:
        cur.close()
    return redirect(url_for('supervisor_dashboard'))

@app.route('/supervisor/remove-stage', methods=['POST'])
@login_required
def supervisor_remove_stage():
    if not current_user.can_assign_stages:
        return "Access Denied", 403
    g_id = request.form.get('g_id')
    stage = (request.form.get('stage') or '').strip().lower()
    if not g_id or not stage:
        return {"ok": False, "error": "Missing parameters"}, 400
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Fetch current stage and dynamic stages for this bid
        cur.execute("SELECT state FROM go_bids WHERE g_id=%s", (g_id,))
        row = cur.fetchone()
        current_stage = (row.get('state') or 'analyzer').strip().lower() if row else 'analyzer'

        # Build dynamic stages order (defaults - exclusions + customs)
        cur.execute("SELECT stage FROM bid_stage_exclusions WHERE g_id=%s", (g_id,))
        excluded = { (r['stage'] or '').strip().lower() for r in cur.fetchall() }
        cur.execute("SELECT stage FROM bid_custom_stages WHERE g_id=%s", (g_id,))
        customs = [ (r['stage'] or '').strip().lower() for r in cur.fetchall() ]
        default_stages = ['analyzer', 'business', 'design', 'operations', 'engineer']
        stages = [s for s in default_stages if s not in excluded] + [s for s in customs if s not in excluded and s not in default_stages]

        if stage not in stages:
            return {"ok": False, "error": "Stage not found for this bid."}, 400
        try:
            curr_idx = stages.index(current_stage)
        except ValueError:
            curr_idx = 0
        target_idx = stages.index(stage)
        if target_idx <= curr_idx:
            return {"ok": False, "error": "Cannot delete current or previous stages."}, 400

        # Mark as excluded and remove from custom if present
        cur.execute("INSERT IGNORE INTO bid_stage_exclusions (g_id, stage) VALUES (%s, %s)", (g_id, stage))
        cur.execute("DELETE FROM bid_custom_stages WHERE g_id=%s AND stage=%s", (g_id, stage))
        mysql.connection.commit()
        return {"ok": True, "message": "Stage deleted."}
    except Exception as e:
        mysql.connection.rollback()
        return {"ok": False, "error": str(e)}, 500
    finally:
        cur.close()
@app.route('/master-dashboard')
@login_required
def master_dashboard():
    if not current_user.is_admin:
        return "Access Denied", 403
    
    # Auto-sync incoming GO bids so dashboards always reflect latest data
    try:
        sync_go_bids()
    except Exception:
        pass

    cur = mysql.connection.cursor(DictCursor)
    
    # Company access filter based on user email domain
    email = (current_user.email or '').lower()
    company_filter_name = None
    if email.endswith('@ikioledlighting.com'):
        company_filter_name = 'Ikio'
    elif email.endswith('@metcoengineering.com'):
        company_filter_name = 'Metco'
    elif email.endswith('@sunsprintengineering.com'):
        company_filter_name = 'Sunsprint'

    company_sql_clause = ""
    company_params = tuple()
    if company_filter_name:
        company_sql_clause = " AND LOWER(COALESCE(gb.company,'')) = %s"
        company_params = (company_filter_name.lower(),)

    # Get all bids from go_bids
    cur.execute(
        """
        SELECT gb.g_id AS id,
               gb.b_name AS name,
               COALESCE(gb.state, 'analyzer') AS current_stage,
               '' AS user_email,
               '' AS user_role
        FROM go_bids gb
        """ + company_sql_clause, company_params
    )
    all_bids = cur.fetchall()
    # Attach nested user object for template compatibility (bid.user.email, bid.user.role)
    for bid in all_bids:
        bid['user'] = {
            'email': bid.get('user_email'),
            'role': bid.get('user_role')
        }
    
    # Get all companies
    cur.execute("SELECT * FROM companies")
    companies = cur.fetchall()
    
    # Get top 5 projects from all companies (most urgent due dates) with company name
    cur.execute("""
        SELECT p.*, c.name AS company_name
        FROM projects p
        JOIN companies c ON p.company_id = c.id
        ORDER BY p.due_date ASC
        LIMIT 5
    """)
    top_projects = cur.fetchall()
    # Attach nested company object for template compatibility (project.company.name)
    for project in top_projects:
        project['company'] = {'name': project.get('company_name')}
    has_projects = len(top_projects) > 0
    
    # Get all tasks
    cur.execute("SELECT * FROM tasks")
    all_tasks = cur.fetchall()

    # Compute real-time stats for cards across the three companies only using go_bids
    cur.execute("SELECT name FROM companies WHERE name IN ('Ikio','Metco','Sunsprint')")
    target_company_names = [row['name'] for row in cur.fetchall()]
    if not target_company_names:
        target_company_names = ['__none__']
    in_clause = ','.join(['%s'] * len(target_company_names))
    # Also fetch ids for the projects metrics below
    cur.execute("SELECT id FROM companies WHERE name IN ('Ikio','Metco','Sunsprint')")
    target_company_ids = [row['id'] for row in cur.fetchall()] or [-1]
    in_clause_ids = ','.join(['%s'] * len(target_company_ids))

    cur.execute(f"SELECT COUNT(*) AS total_bids FROM go_bids WHERE company IN ({in_clause})", target_company_names)
    total_bids = cur.fetchone()['total_bids']

    cur.execute(f"""
        SELECT COUNT(*) AS live_bids
        FROM go_bids
        WHERE COALESCE(state,'analyzer') IN ('business','design','operations','engineer')
          AND company IN ({in_clause})
    """, target_company_names)
    live_bids = cur.fetchone()['live_bids']

    cur.execute(f"SELECT COUNT(*) AS bids_won FROM go_bids WHERE decision='WON' AND company IN ({in_clause})", target_company_names)
    bids_won = cur.fetchone()['bids_won']

    # Build per-company stage counts from go_bids for the timeline trackers
    # Per-state counts
    cur.execute(
        """
        SELECT COALESCE(company, '') AS company,
               LOWER(COALESCE(state, '')) AS state,
               COUNT(*) AS total
        FROM go_bids
        GROUP BY COALESCE(company, ''), LOWER(COALESCE(state, ''))
        """
    )
    rows = cur.fetchall()
    go_counts = {}
    for row in rows:
        company_key = row['company'] or ''
        if company_key not in go_counts:
            go_counts[company_key] = {'analyzer': 0, 'business': 0, 'design': 0, 'operations': 0, 'engineer': 0, 'handover': 0}
        # Map known downstream stages directly; others will be reflected in analyzer via totals below
        if row['state'] in ('business', 'design', 'operations', 'engineer', 'handover'):
            go_counts[company_key][row['state']] = row['total']

    # Total per company should power the Analyzer start point
    cur.execute(
        """
        SELECT COALESCE(company, '') AS company,
               COUNT(*) AS total
        FROM go_bids
        GROUP BY COALESCE(company, '')
        """
    )
    totals = cur.fetchall()
    for t in totals:
        company_key = t['company'] or ''
        if company_key not in go_counts:
            go_counts[company_key] = {'analyzer': 0, 'business': 0, 'design': 0, 'operations': 0, 'engineer': 0, 'handover': 0}
        go_counts[company_key]['analyzer'] = t['total']
    company_names = [c['name'] for c in companies]

    # Build go_bids by company for dropdown summaries, enriched with status from downstream tables
    cur.execute(
        """
        SELECT gb.g_id,
               gb.b_name,
        
               gb.due_date,
               gb.state,
               gb.type,
               gb.company,
               gb.decision,
               gb.summary,
               gb.scoring AS progress,
               wps.pr_completion_status AS work_status,
               wps.dept_bde,
               wps.dept_m_d,
               wps.dept_op,
               wps.dept_site,
               wbr.closure_status AS project_status,
               wbr.work_progress_status AS work_progress_status,
               wlr.result AS wl_result
        FROM go_bids gb
        LEFT JOIN bid_assign ba ON ba.g_id = gb.g_id
        LEFT JOIN win_lost_results wlr ON wlr.a_id = ba.a_id
        LEFT JOIN won_bids_result wbr ON wbr.w_id = wlr.w_id
        LEFT JOIN work_progress_status wps ON wps.won_id = wbr.won_id
        ORDER BY gb.company ASC, gb.due_date ASC
        """
    )
    go_rows = cur.fetchall()
    go_bids_by_company = {}
    for r in go_rows:
        cname = (r.get('company') or '').strip()
        go_bids_by_company.setdefault(cname, []).append(r)
    go_company_names = [name for name, rows in go_bids_by_company.items() if name and len(rows) > 0]

    # assigned_bids handles assignments now; no assign_go_bids refresh

    # Build a flat list of go_bids projects with normalized tracker progress
    def _normalize_stage(raw_state: str) -> str:
        s = (raw_state or '').strip().lower()
        mapping = {
            'analyzer': 'analyzer',
            'business': 'business',
            'business dev': 'business',
            'bdm': 'business',
            'design': 'design',
            'operations': 'operations',
            'operation': 'operations',
            'engineer': 'engineer',
            'site_manager': 'engineer',
            'site manager': 'engineer',
            'handover': 'handover',
            'won': 'handover'
        }
        # Treat unknown values like submitted/under_review/pending as analyzer
        return mapping.get(s, 'analyzer')

    stage_to_percent = {
        'analyzer': 0,
        'business': 20,
        'design': 40,
        'operations': 60,
        'engineer': 80,
        'handover': 100,
    }

    # Dynamic stages (respect supervisor-managed deletions and additions)
    cur.execute("SELECT g_id, stage FROM bid_stage_exclusions")
    ex_rows = cur.fetchall()
    exclusions_by_bid = {}
    for row in ex_rows:
        exclusions_by_bid.setdefault(row['g_id'], set()).add((row['stage'] or '').strip().lower())
    cur.execute("SELECT g_id, stage FROM bid_custom_stages")
    cs_rows = cur.fetchall()
    customs_by_bid = {}
    for row in cs_rows:
        customs_by_bid.setdefault(row['g_id'], []).append((row['stage'] or '').strip().lower())

    default_stages = ['analyzer', 'business', 'design', 'operations', 'engineer', 'handover']

    def _parse_percent(value: str) -> int:
        try:
            if value is None:
                return 0
            s = str(value)
            # extract digits first
            import re
            m = re.search(r"(\d{1,3})", s)
            if m:
                pct = max(0, min(100, int(m.group(1))))
                return pct
            s = s.strip().lower()
            if s in ('done','completed','closed','handover','100%'):
                return 100
            if s in ('in_progress','ongoing'):
                return 50
            if s in ('pending','todo','not_started'):
                return 0
            return 0
        except Exception:
            return 0

    # Helper: compute stage progress map from bid_checklists per bid
    def _compute_stage_map_for_bid(g_id: int) -> dict:
        try:
            cur2 = mysql.connection.cursor(DictCursor)
            cur2.execute("""
                SELECT bc.progress_pct, bc.status, COALESCE(bc.stage, u.role) AS stage_source
                FROM bid_checklists bc
                LEFT JOIN users u ON bc.created_by = u.id
                WHERE bc.g_id = %s
            """, (g_id,))
            rows = cur2.fetchall()
            cur2.close()
            role_to_stage = {
                'business dev': 'business',
                'design': 'design',
                'operations': 'operations',
                'site manager': 'engineer',
                'engineer': 'engineer'
            }
            buckets: dict[str, list] = {'analyzer': [], 'business': [], 'design': [], 'operations': [], 'engineer': [], 'handover': []}
            for r in rows:
                source = (r.get('stage_source') or '').strip().lower()
                stage = role_to_stage.get(source, source if source in buckets else None)
                if not stage:
                    continue
                pct = r.get('progress_pct')
                if pct is None:
                    s = (r.get('status') or '').strip().lower()
                    pct = 100 if s == 'completed' else 50 if s == 'in_progress' else 0
                try:
                    pct = max(0, min(100, int(pct)))
                except Exception:
                    pct = 0
                buckets[stage].append(pct)
            def avg(lst):
                if not lst:
                    return 0
                return int(round(sum(lst) / len(lst)))
            return {
                'analyzer': avg(buckets['analyzer']),
                'business': avg(buckets['business']),
                'design': avg(buckets['design']),
                'operations': avg(buckets['operations']),
                'engineer': avg(buckets['engineer']),
                'handover': avg(buckets['handover']),
            }
        except Exception:
            return {'analyzer': 0, 'business': 0, 'design': 0, 'operations': 0, 'engineer': 0, 'handover': 0}

    go_projects = []
    for r in go_rows:
        stage_key = _normalize_stage(r.get('state'))
        excluded = exclusions_by_bid.get(r.get('g_id'), set())
        custom = [s for s in customs_by_bid.get(r.get('g_id'), []) if s not in excluded]
        stages = [s for s in default_stages if s not in excluded] + [s for s in custom if s not in default_stages]
        if stages and stage_key in stages:
            idx = stages.index(stage_key)
            stage_progress = int(round((idx / (max(1, len(stages) - 1))) * 100))
        else:
            stage_progress = stage_to_percent.get(stage_key, 0)
        progress_pct = r.get('progress') if r.get('progress') is not None else r.get('scoring')
        try:
            progress_pct = max(0, min(100, int(progress_pct))) if progress_pct is not None else None
        except Exception:
            progress_pct = None
        
        # Calculate dynamic progress and status texts
        item_pct = stage_progress_pct(stage_key)
        proj_status, work_status = status_texts(stage_key)
        
        # Prefer task-based progress; fallback to work_progress_status when tasks are absent
        task_stage_map = _compute_stage_map_for_bid(r.get('g_id'))
        if any(v > 0 for v in task_stage_map.values()):
            stage_progress_map = {
                'analyzer': 0,
                'business': task_stage_map.get('business', 0),
                'design': task_stage_map.get('design', 0),
                'operations': task_stage_map.get('operations', 0),
                'engineer': task_stage_map.get('engineer', 0),
                'handover': 0,
            }
        else:
            stage_progress_map = {
                'analyzer': 0,
                'business': _parse_percent(r.get('dept_bde')),
                'design': _parse_percent(r.get('dept_m_d')),
                'operations': _parse_percent(r.get('dept_op')),
                'engineer': _parse_percent(r.get('dept_site')),
                'handover': 100 if (r.get('project_status') or '').strip().lower() in ('closed','completed','handover','done') else 0,
            }

        # If still no per-stage progress was found (all zeros),
        # estimate based on current stage so the timeline tracker reflects movement.
        if not any(v > 0 for v in stage_progress_map.values()):
            try:
                ordered = ['analyzer','business','design','operations','engineer','handover']
                cur_idx = ordered.index(stage_key) if stage_key in ordered else 0
                estimated = {}
                for i, s in enumerate(ordered):
                    if i < cur_idx:
                        estimated[s] = 100
                    elif i == cur_idx:
                        # show partial progress for the current stage
                        estimated[s] = 40
                    else:
                        estimated[s] = 0
                # keep analyzer and handover conventions
                estimated['analyzer'] = estimated.get('analyzer', 0)
                estimated['handover'] = estimated.get('handover', 0)
                stage_progress_map.update(estimated)
            except Exception:
                pass

        go_projects.append({
            'g_id': r.get('g_id'),
            'b_name': r.get('b_name'),
            'company': r.get('company'),
            'state': r.get('state'),
            'stage_key': stage_key,
            'stage_progress': stage_progress,
            'stages': stages,
            'current_stage': stage_key,
         
            'due_date': r.get('due_date'),
            'type': r.get('type'),
            'decision': r.get('decision') or r.get('wl_result'),
            'project_status': proj_status,  # New dynamic project status
            'work_status': work_status,     # New dynamic work status
            'summary': r.get('summary'),
            'progress_pct': progress_pct,
            'work_progress_pct': item_pct,  # New dynamic progress
            'stage_progress_map': stage_progress_map,
        })

    # Total projects across target companies (fallback to all projects if none linked)
    cur.execute(f"SELECT COUNT(*) AS projects_linked FROM projects WHERE company_id IN ({in_clause_ids})", target_company_ids)
    projects_linked = cur.fetchone()['projects_linked']
    if projects_linked > 0:
        cur.execute(f"SELECT COUNT(*) AS projects_total FROM projects WHERE company_id IN ({in_clause_ids})", target_company_ids)
        projects_total = cur.fetchone()['projects_total']
    else:
        cur.execute("SELECT COUNT(*) AS projects_total FROM projects")
        projects_total = cur.fetchone()['projects_total']

    # Project-based dashboard cards (map to existing template keys)
    # total_bids -> Total Projects
    # live_bids -> Projects Completed
    # bids_won -> Projects (in progress)
    # projects_completed -> Projects (on hold)
    if projects_linked > 0:
        cur.execute(f"SELECT COUNT(*) AS projects_completed FROM projects WHERE status='completed' AND company_id IN ({in_clause_ids})", target_company_ids)
        p_completed = cur.fetchone()['projects_completed']
        cur.execute(f"SELECT COUNT(*) AS projects_in_progress FROM projects WHERE status IN ('active','in_progress') AND company_id IN ({in_clause_ids})", target_company_ids)
        p_in_progress = cur.fetchone()['projects_in_progress']
        cur.execute(f"SELECT COUNT(*) AS projects_on_hold FROM projects WHERE status='on_hold' AND company_id IN ({in_clause_ids})", target_company_ids)
        p_on_hold = cur.fetchone()['projects_on_hold']
    else:
        cur.execute("SELECT COUNT(*) AS projects_completed FROM projects WHERE status='completed'")
        p_completed = cur.fetchone()['projects_completed']
        cur.execute("SELECT COUNT(*) AS projects_in_progress FROM projects WHERE status IN ('active','in_progress')")
        p_in_progress = cur.fetchone()['projects_in_progress']
        cur.execute("SELECT COUNT(*) AS projects_on_hold FROM projects WHERE status='on_hold'")
        p_on_hold = cur.fetchone()['projects_on_hold']

    # Override dashboard cards using win_lost_results as requested
    try:
        cur.execute("SELECT COUNT(*) AS total_projects_wlr FROM win_lost_results")
        total_projects_wlr = cur.fetchone()['total_projects_wlr']
    except Exception:
        total_projects_wlr = projects_total

    try:
        cur.execute("""
            SELECT COUNT(*) AS in_progress_wlr
            FROM win_lost_results
            WHERE COALESCE(UPPER(result), '') NOT IN ('WON','LOST')
        """)
        in_progress_wlr = cur.fetchone()['in_progress_wlr']
    except Exception:
        in_progress_wlr = p_in_progress

    data = {
        'total_bids': total_projects_wlr,      # Total Projects
        'live_bids': p_completed,               # Projects Completed (unchanged)
        'bids_won': in_progress_wlr,            # Projects (in progress)
        'projects_completed': p_on_hold         # Projects (on hold)
    }

    # Compute bid analyzer stats from bid_incoming table
    cur.execute("SELECT COUNT(*) AS total_bids FROM bid_incoming")
    total_bids_analyzer = cur.fetchone()['total_bids']
    
    cur.execute("SELECT COUNT(*) AS bids_go FROM bid_incoming WHERE decision = 'GO'")
    bids_go_analyzer = cur.fetchone()['bids_go']
    
    cur.execute("SELECT COUNT(*) AS bids_no_go FROM bid_incoming WHERE decision = 'NO-GO'")
    bids_no_go_analyzer = cur.fetchone()['bids_no_go']
    
    cur.execute("SELECT COUNT(*) AS bids_submitted FROM bid_incoming WHERE state IN ('submitted', 'under_review')")
    bids_submitted_analyzer = cur.fetchone()['bids_submitted']
    
    cur.execute("SELECT COUNT(*) AS bids_won FROM bid_incoming WHERE decision = 'WON'")
    bids_won_analyzer = cur.fetchone()['bids_won']
    
    cur.execute("SELECT COUNT(*) AS bids_lost FROM bid_incoming WHERE decision = 'LOST'")
    bids_lost_analyzer = cur.fetchone()['bids_lost']

    bid_stats = {
        'total_bids': total_bids_analyzer,
        'bids_go': bids_go_analyzer,
        'bids_no_go': bids_no_go_analyzer,
        'bids_submitted': bids_submitted_analyzer,
        'bids_won': bids_won_analyzer,
        'bids_lost': bids_lost_analyzer
    }
    
    # Ensure won_bids_result contains all WON decisions from bid_incoming
    # Check the actual column name in bid_incoming table
    try:
        cur.execute("DESCRIBE bid_incoming")
        describe_results = cur.fetchall()
        # Handle both dict and tuple formats
        if describe_results and isinstance(describe_results[0], dict):
            bid_incoming_columns = [row['Field'] for row in describe_results]
        else:
            bid_incoming_columns = [row[0] for row in describe_results]  # First column is Field name
        
        bid_incoming_id_col = 'id' if 'id' in bid_incoming_columns else ('id' if 'id' in bid_incoming_columns else None)
        
        if bid_incoming_id_col:
            cur.execute(
                f"""
                INSERT INTO won_bids_result (w_id)
                SELECT DISTINCT wlr.w_id
                FROM bid_incoming bi
                LEFT JOIN go_bids gb ON gb.id = bi.{bid_incoming_id_col}
                LEFT JOIN bid_assign ba ON ba.g_id = gb.g_id
                LEFT JOIN win_lost_results wlr ON wlr.a_id = ba.a_id
                LEFT JOIN won_bids_result wbr ON wbr.w_id = wlr.w_id
                WHERE UPPER(bi.decision) = 'WON'
                  AND wlr.w_id IS NOT NULL
                  AND wbr.w_id IS NULL
                """
            )
            mysql.connection.commit()
    except Exception as e:
        # If the join fails (e.g., no relationship between tables), skip this step
        # won_bids_result will be populated from other sources
        mysql.connection.rollback()
        pass

    # Build Won Projects (latest 5) timeline data
    # Ensure won_bids_result has rows for any WIN in win_lost_results
    cur.execute(
        """
        INSERT INTO won_bids_result (w_id)
        SELECT wlr.w_id
        FROM win_lost_results wlr
        LEFT JOIN won_bids_result wbr ON wbr.w_id = wlr.w_id
        WHERE UPPER(COALESCE(wlr.result,'')) = 'WON'
          AND wbr.w_id IS NULL
        """
    )
    mysql.connection.commit()

    cur.execute(
        """
        SELECT wbr.*, gb.g_id, gb.b_name, gb.company,
               COALESCE(wps.pr_completion_status, wbr.work_progress_status) AS work_progress_status
        FROM won_bids_result wbr
        LEFT JOIN win_lost_results wlr ON wlr.w_id = wbr.w_id
        LEFT JOIN bid_assign ba ON ba.a_id = wlr.a_id
        LEFT JOIN go_bids gb ON gb.g_id = ba.g_id
        LEFT JOIN work_progress_status wps ON wps.won_id = wbr.won_id
        ORDER BY wbr.w_id DESC
        LIMIT 5
        """
    )
    won_rows = cur.fetchall()

    def _won_stage_key(row: dict) -> str:
        closure = (row.get('closure_status') or '').strip().lower()
        work_prog = (row.get('work_progress_status') or '').strip().lower()
        if closure in ('closed', 'completed', 'handover', 'done'):
            return 'closure'
        if work_prog in ('operations', 'operation', 'ops', 'in_progress'):
            return 'operations'
        if work_prog == 'design':
            return 'design'
        return 'business'

    won_stages_default = ['won', 'business', 'design', 'operations', 'closure']
    won_projects = []
    for r in won_rows:
        stage_key = _won_stage_key(r)
        idx = won_stages_default.index(stage_key) if stage_key in won_stages_default else 0
        stage_progress = int(round((idx / (len(won_stages_default) - 1)) * 100)) if len(won_stages_default) > 1 else 0
        won_projects.append({
            'g_id': r.get('g_id'),
            'b_name': r.get('b_name') or f"Won #{r.get('won_id')}",
            'company': r.get('company'),
            'current_stage': stage_key,
            'stage_progress': stage_progress,
            'stages': won_stages_default,
        })

    # Fallback: include directly WON bids from bid_incoming not present above
    cur.execute("""
        SELECT bi.id, bi.b_name, COALESCE(gb.company,'') AS company
        FROM bid_incoming bi
        LEFT JOIN go_bids gb ON gb.id = bi.id
        WHERE UPPER(bi.decision) = 'WON'
        ORDER BY bi.id DESC
        LIMIT 20
    """)
    bi_won = cur.fetchall() or []
    existing_names = { (p['b_name'] or '').strip() for p in won_projects }
    for r in bi_won:
        name = r.get('b_name')
        if name and name.strip() in existing_names:
            continue
        won_projects.append({
            'g_id': r.get('id'),  # Changed from id to id to match the SELECT query
            'b_name': name,
            'company': r.get('company'),
            'current_stage': 'business',
            'stage_progress': int(round((won_stages_default.index('business')/(len(won_stages_default)-1))*100)),
            'stages': won_stages_default,
        })

    # Split go_projects into top 5 by due_date and the rest
    def _date_key(item):
        try:
            from datetime import datetime, date
            v = item.get('due_date')
            if v is None or v == '':
                return datetime.max
            # If already a datetime/date, convert to datetime
            if isinstance(v, datetime):
                return v
            if isinstance(v, date):
                return datetime.combine(v, datetime.min.time())
            # Try common string formats
            s = str(v).strip()
            for fmt in ('%Y-%m-%d', '%Y-%m-%d %H:%M:%S', '%d-%m-%Y', '%m/%d/%Y', '%Y/%m/%d'):
                try:
                    return datetime.strptime(s[:19], fmt)
                except Exception:
                    pass
            # Fallback: try fromisoformat
            try:
                return datetime.fromisoformat(s[:19])
            except Exception:
                return datetime.max
        except Exception:
            from datetime import datetime
            return datetime.max
    go_projects_sorted = sorted(go_projects, key=_date_key)
    go_projects_top5 = go_projects_sorted[:5]
    go_projects_more = go_projects_sorted[5:]

    # Parallel kickoff: ensure every bid appears in each team dashboard
    for team_key in ['business','design','operations','engineer']:
        ensure_tasks_for_team(team_key)

    cur.close()

    # Stage display name mapping
    def get_stage_display_name(stage_key):
        mapping = {
            'analyzer': 'BID Analyzer',
            'business': 'Business Development',
            'design': 'Design & Marketing',
            'operations': 'Operation Team',
            'engineer': 'Engineering Team',
            'handover': 'Submitted'
        }
        return mapping.get(stage_key.lower(), stage_key.title())

    return render_template(
        'master_dashboard.html', 
        page='dashboard',
        title='Master Dashboard',
        data=data,
        bids=all_bids,
        companies=companies,
        projects=top_projects,
        tasks=all_tasks,
        bid_stats=bid_stats,
        go_counts=go_counts,
        has_projects=has_projects,
        company_names=company_names,
        go_company_names=go_company_names,
        go_bids_by_company=go_bids_by_company,
        go_projects_top5=go_projects_top5,
        go_projects_more=go_projects_more,
        won_projects=won_projects,
        get_stage_display_name=get_stage_display_name
    )
@app.route('/company/<company_name>')
@login_required
def company_dashboard(company_name):
    if not current_user.is_admin:
        return "Access Denied", 403
    
    # Map URL names to database names
    company_mapping = {
        'ikio': ['Ikio', 'IKIO'],
        'metco': ['Metco', 'Metco Engineering', 'METCO'],
        'sunsprint': ['Sunsprint', 'Sunsprint Engineering', 'SUNSPRINT']
    }
    
    aliases = company_mapping.get(company_name.lower())
    if not aliases:
        return "Company not found", 404
    if isinstance(aliases, str):
        aliases = [aliases]
        
    cur = mysql.connection.cursor(DictCursor)
    
    # Get company (try aliases, allow case-insensitive match)
    company = None
    for alias in aliases:
        cur.execute("SELECT * FROM companies WHERE LOWER(name) = %s", (alias.lower(),))
        company = cur.fetchone()
        if company:
            break
    if not company:
        cur.close()
        return "Company not found", 404
    
    # Get top 5 projects for this company with nested company info
    cur.execute("""
        SELECT p.*, c.name AS company_name
        FROM projects p
        JOIN companies c ON p.company_id = c.id
        WHERE p.company_id=%s
        ORDER BY p.due_date ASC
        LIMIT 5
    """, (company['id'],))
    projects = cur.fetchall()
    for project in projects:
        project['company'] = {'name': project.get('company_name')}
    
    # Get tasks for this company
    cur.execute("""
        SELECT t.* FROM tasks t 
        JOIN projects p ON t.project_id = p.id 
        WHERE p.company_id = %s
    """, (company['id'],))
    tasks = cur.fetchall()
    
    # Build set of aliases for filtering bids (case-insensitive)
    company_aliases = { (company.get('name') or '').lower() }
    company_aliases.update(a.lower() for a in aliases if a)
    placeholders = ','.join(['%s'] * len(company_aliases)) or '%s'

    # Get bids for this company from go_bids by company name (enriched)
    cur.execute(
        f"""
        SELECT gb.*, 
               wps.pr_completion_status AS work_status,
               wbr.closure_status AS project_status,
               wbr.work_progress_status AS work_progress_status,
               wlr.result AS wl_result
        FROM go_bids gb
        LEFT JOIN bid_assign ba ON ba.g_id = gb.g_id
        LEFT JOIN win_lost_results wlr ON wlr.a_id = ba.a_id
        LEFT JOIN won_bids_result wbr ON wbr.w_id = wlr.w_id
        LEFT JOIN work_progress_status wps ON wps.won_id = wbr.won_id
        WHERE LOWER(gb.company) IN ({placeholders})
        ORDER BY gb.due_date ASC
        """,
        tuple(company_aliases)
    )
    bids = cur.fetchall()
    
    # Fetch stage exclusions and custom stages for each bid
    cur.execute("SELECT g_id, stage FROM bid_stage_exclusions")
    rows = cur.fetchall()
    stage_exclusions = {}
    for r in rows:
        stage_exclusions.setdefault(r['g_id'], []).append(r['stage'])

    cur.execute("SELECT g_id, stage FROM bid_custom_stages")
    rows = cur.fetchall()
    custom_stages = {}
    for r in rows:
        custom_stages.setdefault(r['g_id'], []).append(r['stage'])

    # Prepare go_projects for this company similar to master dashboard
    def _normalize_stage_company(raw_state: str) -> str:
        s = (raw_state or '').strip().lower()
        mapping = {
            'analyzer': 'analyzer',
            'business': 'business', 'business dev': 'business', 'bdm': 'business',
            'design': 'design',
            'operations': 'operations', 'operation': 'operations',
            'engineer': 'engineer', 'site_manager': 'engineer', 'site manager': 'engineer',
            'handover': 'handover', 'won': 'handover'
        }
        return mapping.get(s, 'analyzer')

    stage_to_percent_company = {
        'analyzer': 0,
        'business': 20,
        'design': 40,
        'operations': 60,
        'engineer': 80,
        'handover': 100,
    }

    # Default stages list
    default_stages = ['analyzer', 'business', 'design', 'operations', 'engineer']
    
    go_projects = []
    for r in bids:
        bid_id = r.get('g_id')
        
        # Get dynamic stages for this bid
        excluded = set(stage_exclusions.get(bid_id, []))
        customs = [s for s in custom_stages.get(bid_id, []) if s not in excluded]
        stages = [s for s in default_stages if s not in excluded] + [s for s in customs if s not in default_stages]
        
        # If no stages remain, fall back to default
        if not stages:
            stages = default_stages.copy()
        
        stage_key = _normalize_stage_company(r.get('state'))
        
        # Calculate stage_progress based on dynamic stages
        if stage_key in stages:
            idx = stages.index(stage_key)
            stage_progress = int(round((idx / max(1, len(stages) - 1)) * 100))
        else:
            stage_progress = 0
        
        progress_pct = r.get('scoring')
        try:
            progress_pct = max(0, min(100, int(progress_pct))) if progress_pct is not None else None
        except Exception:
            progress_pct = None
        
        # Calculate dynamic progress and status texts
        item_pct = stage_progress_pct(stage_key)
        proj_status, work_status = status_texts(stage_key)
        
        go_projects.append({
            'g_id': bid_id,
            'b_name': r.get('b_name'),
            'company': r.get('company'),
            'state': r.get('state'),
            'stage_key': stage_key,
            'stage_progress': stage_progress,
            'stages': stages,  # Add dynamic stages list
         
            'due_date': r.get('due_date'),
            'type': r.get('type'),
            'decision': r.get('decision') or r.get('wl_result'),
            'project_status': proj_status,  # New dynamic project status
            'work_status': work_status,     # New dynamic work status
            'summary': r.get('summary'),
            'progress_pct': progress_pct,
            'work_progress_pct': item_pct,  # New dynamic progress
        })
    
    # Compute company metrics
    cur.execute("""
        SELECT 
            COUNT(*) AS total_projects,
            SUM(CASE WHEN COALESCE(status,'') = 'completed' OR COALESCE(progress,0) >= 100 THEN 1 ELSE 0 END) AS completed_projects,
            AVG(COALESCE(progress,0)) AS avg_progress
        FROM projects
        WHERE company_id = %s
    """, (company['id'],))
    proj_stats = cur.fetchone() or {}
    total_projects = int(proj_stats.get('total_projects') or 0)
    completed_projects = int(proj_stats.get('completed_projects') or 0)
    active_projects = max(0, total_projects - completed_projects)
    avg_progress_pct = int(round(proj_stats.get('avg_progress') or 0))

    company_metrics = {
        'active_projects': active_projects,
        'total_tasks': len(tasks or []),
        'completed_projects': completed_projects,
        'avg_progress_pct': avg_progress_pct,
    }

    # Load multi-assignees per bid
    bid_ids = []
    for b in (bids or []):
        bid_key = b.get('id') or b.get('g_id')
        if bid_key is not None:
            bid_ids.append(bid_key)
    assigned_map = {}
    try:
        if bid_ids:
            # Ensure mapping table exists
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS bid_assignment_members (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    g_id INT NOT NULL,
                    employee_id INT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE KEY uniq_g_emp (g_id, employee_id)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
                """
            )
            placeholders = ','.join(['%s'] * len(bid_ids))
            cur.execute(
                f"""
                SELECT bam.g_id, e.id AS employee_id, e.name, e.email
                FROM bid_assignment_members bam
                JOIN employees e ON e.id = bam.employee_id
                WHERE bam.g_id IN ({placeholders}) AND e.is_active = TRUE
                ORDER BY e.name
                """,
                tuple(bid_ids)
            )
            rows = cur.fetchall()
            for r in rows:
                assigned_map.setdefault(r['g_id'], []).append({'id': r['employee_id'], 'name': r['name'], 'email': r['email']})
        # Attach to bid rows
        for b in bids:
            bid_key = b.get('id') or b.get('g_id')
            b['assigned_members'] = assigned_map.get(bid_key, [])
    except Exception:
        for b in bids:
            b['assigned_members'] = []

    cur.close()
    
    # Stage display name mapping
    def get_stage_display_name(stage_key):
        mapping = {
            'analyzer': 'BID Analyzer',
            'business': 'Business Development',
            'design': 'Design & Marketing',
            'operations': 'Operation Team',
            'engineer': 'Engineering Team',
            'handover': 'Submitted'
        }
        return mapping.get(stage_key.lower(), stage_key.title())
    
    return render_template(
        'company_dashboard.html', 
        company=company,
        projects=projects,
        tasks=tasks,
        bids=bids,
        go_projects=go_projects,
        company_metrics=company_metrics,
        get_stage_display_name=get_stage_display_name,
        current_company=company_name.lower()
    )

@app.route('/bid-analyzer')
@login_required
def bid_analyzer():
    if not current_user.is_admin:
        return "Access Denied", 403
    
    # Ensure latest GO bids are mirrored into go_bids whenever analyzer loads
    try:
        sync_go_bids()
    except Exception:
        pass
    
    cur = mysql.connection.cursor(DictCursor)
    
    # Get bid_incoming data for the table
    cur.execute("SELECT * FROM bid_incoming ORDER BY id DESC")
    bid_incoming_data = cur.fetchall()
    
    # Calculate bid stats from bid_incoming table
    cur.execute("SELECT COUNT(*) AS total_bids FROM bid_incoming")
    total_bids = cur.fetchone()['total_bids']
    
    cur.execute("SELECT COUNT(*) AS bids_go FROM bid_incoming WHERE decision = 'GO'")
    bids_go = cur.fetchone()['bids_go']
    
    cur.execute("SELECT COUNT(*) AS bids_no_go FROM bid_incoming WHERE decision = 'NO-GO'")
    bids_no_go = cur.fetchone()['bids_no_go']
    
    cur.execute("SELECT COUNT(*) AS bids_submitted FROM bid_incoming WHERE state IN ('submitted', 'under_review')")
    bids_submitted = cur.fetchone()['bids_submitted']
    
    cur.execute("SELECT COUNT(*) AS bids_won FROM bid_incoming WHERE decision = 'WON'")
    bids_won = cur.fetchone()['bids_won']
    
    cur.execute("SELECT COUNT(*) AS bids_lost FROM bid_incoming WHERE decision = 'LOST'")
    bids_lost = cur.fetchone()['bids_lost']
    
    cur.close()

    return render_template('bid_analyzer_landing.html', bid_cards={
        'total_bids': total_bids,
        'bids_go': bids_go,
        'bids_no_go': bids_no_go,
        'bids_submitted': bids_submitted,
        'bids_won': bids_won,
        'bids_lost': bids_lost
    }, bid_incoming_data=bid_incoming_data)
@app.route('/databases')
@login_required
def databases():
    if not current_user.is_admin:
        return redirect(url_for('role_dashboard'))
    
    # Get all companies and their projects
    cur = mysql.connection.cursor(DictCursor)
    cur.execute("SELECT * FROM companies")
    companies = cur.fetchall()
    company_data = {}
    
    for company in companies:
        cur.execute("SELECT * FROM projects WHERE company_id=%s", (company['id'],))
        projects = cur.fetchall()
        company_data[company['name']] = {
            'company': company,
            'projects': projects
        }
    
    cur.close()
    return render_template('databases.html', company_data=company_data)
@app.route('/databases/create_project', methods=['POST'])
@login_required
def create_project():
    if not current_user.is_admin:
        return redirect(url_for('role_dashboard'))
    
    try:
        company_name = request.form.get('company_name')
        project_name = request.form.get('project_name')
        due_date = request.form.get('due_date')
        status = request.form.get('status', 'active')
        progress = int(request.form.get('progress', 0))
        
        cur = mysql.connection.cursor()
        
        # Find or create company
        cur.execute("SELECT id FROM companies WHERE name=%s", (company_name,))
        company = cur.fetchone()
        if not company:
            cur.execute("INSERT INTO companies (name, description) VALUES (%s, %s)", 
                       (company_name, f'{company_name} Projects'))
            company_id = cur.lastrowid
        else:
            company_id = company['id']
        
        # Create project
        cur.execute("""
            INSERT INTO projects (name, company_id, start_date, due_date, status, progress) 
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (project_name, company_id, datetime.utcnow(), 
              datetime.strptime(due_date, '%Y-%m-%d'), status, progress))
        
        mysql.connection.commit()
        cur.close()
        
        flash(f'Project "{project_name}" created successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error creating project: {str(e)}', 'error')
    
    return redirect(url_for('databases'))

@app.route('/databases/update_project/<int:project_id>', methods=['POST'])
@login_required
def update_project(project_id):
    if not current_user.is_admin:
        return redirect(url_for('role_dashboard'))
    
    try:
        cur = mysql.connection.cursor()
        
        # Check if project exists
        cur.execute("SELECT * FROM projects WHERE id=%s", (project_id,))
        project = cur.fetchone()
        if not project:
            cur.close()
            return "Project not found", 404
        
        # Update project
        project_name = request.form.get('project_name', project['name'])
        due_date = datetime.strptime(request.form.get('due_date'), '%Y-%m-%d')
        status = request.form.get('status', project['status'])
        progress = int(request.form.get('progress', project['progress']))
        
        cur.execute("""
            UPDATE projects 
            SET name=%s, due_date=%s, status=%s, progress=%s 
            WHERE id=%s
        """, (project_name, due_date, status, progress, project_id))
        
        mysql.connection.commit()
        cur.close()
        
        flash(f'Project "{project_name}" updated successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error updating project: {str(e)}', 'error')
    
    return redirect(url_for('databases'))

@app.route('/databases/delete_project/<int:project_id>')
@login_required
def delete_project(project_id):
    if not current_user.is_admin:
        return redirect(url_for('role_dashboard'))
    
    try:
        cur = mysql.connection.cursor()
        
        # Get project name before deletion
        cur.execute("SELECT name FROM projects WHERE id=%s", (project_id,))
        project = cur.fetchone()
        if not project:
            cur.close()
            return "Project not found", 404
        
        project_name = project['name']
        
        # Delete project
        cur.execute("DELETE FROM projects WHERE id=%s", (project_id,))
        mysql.connection.commit()
        cur.close()
        
        flash(f'Project "{project_name}" deleted successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error deleting project: {str(e)}', 'error')
    
    return redirect(url_for('databases'))

@app.route('/databases/import_excel', methods=['POST'])
@login_required
def import_excel():
    if not current_user.is_admin:
        return redirect(url_for('role_dashboard'))
    
    try:
        if 'excel_file' not in request.files:
            flash('No file selected', 'error')
            return redirect(url_for('databases'))
        
        file = request.files['excel_file']
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(url_for('databases'))
        
        if file and file.filename.endswith('.xlsx'):
            import pandas as pd
            import io
            
            # Read Excel file
            df = pd.read_excel(io.BytesIO(file.read()))
            
            # Expected columns: Company, Project Name, Due Date, Revenue, Status, Progress
            required_columns = ['Company', 'Project Name', 'Due Date', 'Status', 'Progress']
            
            if not all(col in df.columns for col in required_columns):
                flash(f'Excel file must contain columns: {", ".join(required_columns)}', 'error')
                return redirect(url_for('databases'))
            
            projects_created = 0
            cur = mysql.connection.cursor()
            
            for _, row in df.iterrows():
                try:
                    # Find or create company
                    company_name = str(row['Company']).strip()
                    cur.execute("SELECT id FROM companies WHERE name=%s", (company_name,))
                    company = cur.fetchone()
                    if not company:
                        cur.execute("INSERT INTO companies (name, description) VALUES (%s, %s)", 
                                   (company_name, f'{company_name} Projects'))
                        company_id = cur.lastrowid
                    else:
                        company_id = company['id']
                    
                    # Create project
                    cur.execute("""
                        INSERT INTO projects (name, company_id, start_date, due_date, status, progress) 
                        VALUES (%s, %s, %s, %s, %s, %s, %s)
                    """, (str(row['Project Name']).strip(), company_id, datetime.utcnow(),
                          pd.to_datetime(row['Due Date']), 
                          str(row['Status']).strip() if pd.notna(row['Status']) else 'active',
                          int(row['Progress']) if pd.notna(row['Progress']) else 0))
                    
                    projects_created += 1
                    
                except Exception as e:
                    print(f"Error processing row: {e}")
                    continue
            
            mysql.connection.commit()
            cur.close()
            flash(f'Successfully imported {projects_created} projects from Excel file!', 'success')
            
        else:
            flash('Please upload a valid Excel file (.xlsx)', 'error')
            
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error importing Excel file: {str(e)}', 'error')
    
    return redirect(url_for('databases'))

@app.route('/databases/drop_all')
@login_required
def drop_all_databases():
    if not current_user.is_admin:
        return redirect(url_for('role_dashboard'))
    
    try:
        cur = mysql.connection.cursor()
        
        # Delete all projects
        cur.execute("DELETE FROM projects")
        
        # Delete all companies
        cur.execute("DELETE FROM companies")
        
        # Delete all tasks
        cur.execute("DELETE FROM tasks")
        
        mysql.connection.commit()
        cur.close()
        flash('All databases have been cleared successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error clearing databases: {str(e)}', 'error')
    
    return redirect(url_for('databases'))

# --- Bid Incoming CRUD Routes ---
@app.route('/bid-analyzer/create', methods=['POST'])
@login_required
def create_bid_incoming():
    if not current_user.is_admin:
        return "Access Denied", 403
    
    try:
        cur = mysql.connection.cursor()
        
        b_name = request.form.get('b_name', '').strip()
      
        due_date = request.form.get('due_date', '').strip()
        state = request.form.get('state', '').strip()
        scope = request.form.get('scope', '').strip()
        type_val = request.form.get('type', '').strip()
        scoring = int(request.form.get('scoring', 0)) if request.form.get('scoring') else None
        comp_name = request.form.get('comp_name', '').strip()
        decision = request.form.get('decision', '').strip()
        summary = request.form.get('summary', '').strip()
        
        if not b_name or not due_date:
            return 'Bid name, incoming date and due date are required', 400
        
        cur.execute("""
            INSERT INTO bid_incoming (b_name, due_date, state, scope, type, scoring, comp_name, decision, summary)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (b_name, due_date, state, scope, type_val, scoring, comp_name, decision, summary))
        
        bid_id = cur.lastrowid
        mysql.connection.commit()
        
        # Automatically sync GO decisions from bid_incoming to go_bids
        if (decision or '').upper() == 'GO':
            cur2 = mysql.connection.cursor(DictCursor)
            cur2.execute("SELECT g_id FROM go_bids WHERE id=%s", (bid_id,))
            row = cur2.fetchone()
            args = (b_name, due_date, state if state else 'business', scope, type_val, scoring, comp_name, decision, summary)
            if row:
                cur2.execute("""UPDATE go_bids SET b_name=%s,due_date=%s,state=%s,scope=%s,
                                type=%s,scoring=%s,company=%s,decision=%s,summary=%s WHERE id=%s""", (*args, bid_id))
            else:
                cur2.execute("""INSERT INTO go_bids (id,b_name,due_date,state,scope,type,scoring,company,decision,summary)
                                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                             (bid_id, *args))
            mysql.connection.commit()
            cur2.close()
            log_write('sync', f"Auto-synced GO bid '{b_name}' (id={bid_id}) from bid_incoming to go_bids")
        
        cur.close()
        log_write('create', f"table=bid_incoming, id={bid_id}")
        flash(f'Bid "{b_name}" created successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error creating bid: {str(e)}', 'error')
    
    return redirect(url_for('bid_analyzer'))

@app.route('/bid-analyzer/update/<int:bid_id>', methods=['POST'])
@login_required
def update_bid_incoming(bid_id):
    if not current_user.is_admin:
        return "Access Denied", 403
    
    try:
        cur = mysql.connection.cursor()
        
        # Check if bid exists
        cur.execute("SELECT * FROM bid_incoming WHERE id=%s", (bid_id,))
        bid = cur.fetchone()
        if not bid:
            cur.close()
            return "Bid not found", 404
        
        b_name = request.form.get('b_name', bid['b_name']).strip()
        due_date = request.form.get('due_date', str(bid['due_date'])).strip()
        state = request.form.get('state', bid['state'] or '').strip()
        scope = request.form.get('scope', bid['scope'] or '').strip()
        type_val = request.form.get('type', bid['type'] or '').strip()
        scoring = int(request.form.get('scoring', 0)) if request.form.get('scoring') else bid['scoring']
        comp_name = request.form.get('comp_name', bid['comp_name'] or '').strip()
        decision = request.form.get('decision', bid['decision'] or '').strip()
        summary = request.form.get('summary', bid['summary'] or '').strip()
        
        cur.execute("""
            UPDATE bid_incoming 
            SET b_name=%s, due_date=%s, state=%s, scope=%s, type=%s, scoring=%s, comp_name=%s, decision=%s, summary=%s
            WHERE id=%s
        """, (b_name, due_date, state, scope, type_val, scoring, comp_name, decision, summary, bid_id))
        
        mysql.connection.commit()
        
        # Automatically sync GO decisions from bid_incoming to go_bids
        if (decision or '').upper() == 'GO':
            cur2 = mysql.connection.cursor(DictCursor)
            cur2.execute("SELECT g_id FROM go_bids WHERE id=%s", (bid_id,))
            row = cur2.fetchone()
            args = (b_name, due_date, state if state else 'business', scope, type_val, scoring, comp_name, decision, summary)
            if row:
                cur2.execute("""UPDATE go_bids SET b_name=%s,due_date=%s,state=%s,scope=%s,
                                type=%s,scoring=%s,company=%s,decision=%s,summary=%s WHERE id=%s""", (*args, bid_id))
            else:
                cur2.execute("""INSERT INTO go_bids (id,b_name,due_date,state,scope,type,scoring,company,decision,summary)
                                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                             (bid_id, *args))
            mysql.connection.commit()
            cur2.close()
            log_write('sync', f"Auto-synced GO bid '{b_name}' (id={bid_id}) from bid_incoming to go_bids")
        else:
            # If decision changed from GO to something else, remove from go_bids
            cur2 = mysql.connection.cursor(DictCursor)
            cur2.execute("DELETE FROM go_bids WHERE id=%s", (bid_id,))
            mysql.connection.commit()
            cur2.close()
            log_write('sync', f"Removed bid '{b_name}' (id={bid_id}) from go_bids (decision changed from GO)")
        
        cur.close()
        log_write('update', f"table=bid_incoming, id={bid_id}")
        flash(f'Bid "{b_name}" updated successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error updating bid: {str(e)}', 'error')
    
    return redirect(url_for('bid_analyzer'))

@app.route('/bid-analyzer/delete/<int:bid_id>')
@login_required
def delete_bid_incoming(bid_id):
    if not current_user.is_admin:
        return "Access Denied", 403
    
    try:
        cur = mysql.connection.cursor()
        
        # Get bid name before deletion
        cur.execute("SELECT b_name FROM bid_incoming WHERE id=%s", (bid_id,))
        bid = cur.fetchone()
        if not bid:
            cur.close()
            return "Bid not found", 404
        
        bid_name = bid['b_name']
        
        # Delete bid
        cur.execute("DELETE FROM bid_incoming WHERE id=%s", (bid_id,))
        mysql.connection.commit()
        cur.close()
        
        log_write('delete', f"table=bid_incoming, id={bid_id}")
        flash(f'Bid "{bid_name}" deleted successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error deleting bid: {str(e)}', 'error')
    
    return redirect(url_for('bid_analyzer'))

@app.route('/database-management')
@login_required
def database_management():
    if not current_user.is_admin:
        return "Access Denied", 403
    
    cur = mysql.connection.cursor(DictCursor)
    
    # Get all available tables
    cur.execute("SHOW TABLES")
    all_tables = [list(table.values())[0] for table in cur.fetchall()]
    
    # Filter to show only specific tables in proper sequence
    allowed_tables = ['bid_incoming', 'go_bids', 'bid_assign', 'win_lost_results', 'won_bids_result', 'work_progress_status']
    # Maintain the sequence by preserving the order in allowed_tables
    tables = [t for t in allowed_tables if t in all_tables]
    
    # Create mapping for display names (capital letters with spaces)
    table_display_names = {
        'bid_incoming': 'BID INCOMING',
        'go_bids': 'GO BIDS',
        'bid_assign': 'BID ASSIGN',
        'win_lost_results': 'WIN LOST RESULTS',
        'won_bids_result': 'WON BIDS RESULT',
        'work_progress_status': 'WORK PROGRESS STATUS'
    }
    
    # Function to format column names (e.g., 'id' -> 'BID ID', 'comp_name' -> 'COMP NAME')
    def format_column_name(col_name):
        # Replace underscores with spaces and convert to uppercase
        formatted = col_name.replace('_', ' ').upper()
        return formatted
    
    # Get selected table (default to first table)
    selected_table = request.args.get('table', tables[0] if tables else '')
    search_query = request.args.get('search', '')
    decision_filter = request.args.get('decision', '').strip()
    
    # Get table data
    table_data = []
    table_columns = []
    companies = []
    
    if selected_table:
        # Get table structure
        cur.execute(f"DESCRIBE `{selected_table}`")
        table_columns = cur.fetchall()
        
        # Auto-sync GO bids into go_bids table
        if selected_table == 'go_bids':
            cur.execute("""
                INSERT INTO go_bids (id, b_name, due_date, state, scope, type, scoring, company, decision, summary)
                SELECT bi.id, bi.b_name, bi.due_date, bi.state, bi.scope, bi.type, bi.scoring, bi.comp_name, bi.decision, bi.summary
                FROM bid_incoming bi
                LEFT JOIN go_bids gb ON gb.id = bi.id
                WHERE UPPER(bi.decision) = 'GO' AND gb.id IS NULL
            """)
            mysql.connection.commit()
            # After syncing GO bids, refresh revenue-based assignments into assigned_bids
            try:
                assign_bids_by_revenue()
            except Exception:
                pass

        # Note: bid_assign is now populated ONLY via explicit Assign action from go_bids

        # Auto-sync bid_assign into win_lost_results (one row per assignment)
        if selected_table == 'win_lost_results':
            cur.execute("""
                INSERT INTO win_lost_results (a_id, b_name, due_date, state, scope, value, company, department, person_name, status, result)
                SELECT ba.a_id, ba.b_name, ba.due_date, ba.state, ba.scope, ba.value, ba.company, ba.depart, ba.person_name, ba.status, NULL
                FROM bid_assign ba
                LEFT JOIN win_lost_results wlr ON wlr.a_id = ba.a_id
                WHERE wlr.a_id IS NULL
            """)
            mysql.connection.commit()

        # Auto-sync win_lost_results into won_bids_result (link by w_id)
        if selected_table == 'won_bids_result':
            cur.execute("""
                INSERT INTO won_bids_result (w_id)
                SELECT wlr.w_id
                FROM win_lost_results wlr
                LEFT JOIN won_bids_result wbr ON wbr.w_id = wlr.w_id
                WHERE wbr.w_id IS NULL
            """)
            mysql.connection.commit()

        # Auto-sync won_bids_result into work_progress_status (link by won_id)
        if selected_table == 'work_progress_status':
            cur.execute("""
                INSERT INTO work_progress_status (won_id, company, b_name, dept_bde, dept_m_d, dept_op, dept_site, pr_completion_status)
                SELECT wbr.won_id, COALESCE(gb.company, ''), COALESCE(gb.b_name, ''), '', '', '', '', NULL
                FROM won_bids_result wbr
                LEFT JOIN win_lost_results wlr ON wlr.w_id = wbr.w_id
                LEFT JOIN bid_assign ba ON ba.a_id = wlr.a_id
                LEFT JOIN go_bids gb ON gb.g_id = ba.g_id
                LEFT JOIN work_progress_status wps ON wps.won_id = wbr.won_id
                WHERE wps.won_id IS NULL
            """)
            mysql.connection.commit()

        # Build dynamic WHERE for search and decision filter
        where_parts = []
        params = []

        # Text search across varchar/text columns
        if search_query:
            search_conditions = []
            for col in table_columns:
                if col['Type'].startswith(('varchar', 'text', 'char')):
                    search_conditions.append(f"`{col['Field']}` LIKE %s")
            if search_conditions:
                search_pattern = f"%{search_query}%"
                params.extend([search_pattern] * len(search_conditions))
                where_parts.append("(" + " OR ".join(search_conditions) + ")")

        # Decision filter for bid_incoming or go_bids
        if selected_table in ('bid_incoming', 'go_bids') and decision_filter:
            where_parts.append("UPPER(`decision`) = %s")
            params.append(decision_filter.upper())

        where_sql = (" WHERE " + " AND ".join(where_parts)) if where_parts else ""
        cur.execute(f"SELECT * FROM `{selected_table}`{where_sql}", params)
        
        table_data = cur.fetchall()
        if selected_table == 'bid_incoming':
            cur.execute("SELECT id, name FROM companies ORDER BY name")
            companies = cur.fetchall()
    
    cur.close()
    
    return render_template('database_management.html', 
                         tables=tables, 
                         selected_table=selected_table,
                         table_columns=table_columns,
                         table_data=table_data,
                         search_query=search_query,
                         decision_filter=decision_filter,
                         companies=companies,
                         table_display_names=table_display_names,
                         format_column_name=format_column_name)

@app.route('/admin/refresh-assign-go')
@login_required
def admin_refresh_assign_go():
    if not current_user.is_admin:
        return "Access Denied", 403
    try:
        assign_bids_by_revenue()
        flash('assigned_bids refreshed from go_bids.', 'success')
    except Exception as e:
        flash(f'Failed to refresh: {e}', 'error')
    return redirect(url_for('database_management', table='assigned_bids'))
@app.route('/database-management/create', methods=['POST'])
@login_required
def dbm_create():
    if not current_user.is_admin:
        return "Access Denied", 403
    table = request.form.get('table')
    if not table:
        return redirect(url_for('database_management'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Get columns and build insert
        cur.execute(f"DESCRIBE `{table}`")
        cols = cur.fetchall()
        fields = []
        values = []
        params = []
        for c in cols:
            if c['Key'] == 'PRI':
                continue
            field = c['Field']
            fields.append(f"`{field}`")
            values.append('%s')
            params.append(request.form.get(field))
        sql = f"INSERT INTO `{table}` ({', '.join(fields)}) VALUES ({', '.join(values)})"
        cur.execute(sql, params)
        mysql.connection.commit()

        # If a bid was created in bid_incoming, offer auto-assign path via flash info
        if table == 'bid_incoming':
            flash('Bid created in Bid Incoming. Assign it to a company from the actions column.', 'success')
    finally:
        cur.close()
    return redirect(url_for('database_management', table=table))

@app.route('/database-management/update/<int:row_id>', methods=['POST'])
@login_required
def dbm_update(row_id):
    if not current_user.is_admin:
        return "Access Denied", 403
    table = request.form.get('table')
    if not table:
        return redirect(url_for('database_management'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        cur.execute(f"DESCRIBE `{table}`")
        cols = cur.fetchall()
        pk = cols[0]['Field']
        sets = []
        params = []
        for c in cols:
            if c['Key'] == 'PRI':
                pk = c['Field']
                continue
            field = c['Field']
            sets.append(f"`{field}`=%s")
            params.append(request.form.get(field))
        params.append(row_id)
        sql = f"UPDATE `{table}` SET {', '.join(sets)} WHERE `{pk}`=%s"
        cur.execute(sql, params)
        mysql.connection.commit()
    finally:
        cur.close()
    return redirect(url_for('database_management', table=table))

@app.route('/database-management/assign', methods=['POST'])
@login_required
def dbm_assign():
    if not current_user.is_admin:
        return "Access Denied", 403
    bid_incoming_id = request.form.get('bid_id')
    company_id = request.form.get('company_id')
    if not bid_incoming_id or not company_id:
        return redirect(url_for('database_management', table='bid_incoming'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Read bid_incoming
        cur.execute("SELECT * FROM bid_incoming WHERE id=%s", (bid_incoming_id,))
        inc = cur.fetchone()
        if not inc:
            cur.close()
            return redirect(url_for('database_management', table='bid_incoming'))
        # Create bid in main bids with analyzer -> business based on decision
        stage = 'analyzer'
        if (inc.get('decision') or '').upper() == 'GO':
            stage = 'business'
        cur.execute("INSERT INTO bids (name, current_stage, user_id, company_id) VALUES (%s,%s,%s,%s)",
                    (inc.get('b_name'), stage, current_user.id, company_id))
        new_bid_id = cur.lastrowid
        # Timeline
        cur.execute("INSERT INTO bid_timeline (bid_id, event, details) VALUES (%s,%s,%s)",
                    (new_bid_id, 'assigned', f"Assigned to company_id={company_id} from bid_incoming {bid_incoming_id}"))
        # Optional: keep in bid_incoming or delete; here we keep it
        mysql.connection.commit()
        # Emit via socket if needed later; for now just flash
        flash('Bid assigned to company successfully.', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Assignment failed: {str(e)}', 'error')
    finally:
        cur.close()
    return redirect(url_for('database_management', table='bid_incoming'))

@app.route('/database-management/assign-go', methods=['POST'])
@login_required
def dbm_assign_go():
    if not current_user.is_admin:
        return "Access Denied", 403
    g_id = request.form.get('g_id')
    depart = request.form.get('depart', '').strip()
    person_name = request.form.get('person_name', '').strip()
    email_to = request.form.get('email', '').strip()
    if not g_id or not depart or not person_name:
        flash('Please provide department and person name', 'error')
        return redirect(url_for('database_management', table='go_bids'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Ensure bid_assign exists for this go bid; insert if missing else update
        cur.execute("SELECT a_id FROM bid_assign WHERE g_id=%s", (g_id,))
        row = cur.fetchone()
        if not row:
            # Pull from go_bids to seed
            cur.execute("SELECT * FROM go_bids WHERE g_id=%s", (g_id,))
            gb = cur.fetchone()
            if not gb:
                flash('GO bid not found', 'error')
                cur.close()
                return redirect(url_for('database_management', table='go_bids'))
            
            # Map department to stage for go_bids.state update
            dept_to_stage = {
                'business dev': 'business',
                'business': 'business',
                'design': 'design',
                'operations': 'operations',
                'site manager': 'engineer',
                'engineer': 'engineer'
            }
            new_stage = dept_to_stage.get(depart.lower(), depart.lower())
            
            # Update go_bids.state to match assigned department
            cur.execute("UPDATE go_bids SET state=%s WHERE g_id=%s", (new_stage, g_id))
            
            # Check if assignee_email column exists before inserting
            try:
                cur.execute("""
                    INSERT INTO bid_assign (g_id, b_name, due_date, state, scope, type, company, depart, person_name, assignee_email, status, value)
                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s)
                """, (
                    gb['g_id'], gb['b_name'], gb['due_date'], new_stage, gb['scope'], gb['type'], gb['company'], depart, person_name, email_to,
                    gb.get('scoring', 0)
                ))
            except Exception as e:
                if "Unknown column 'assignee_email'" in str(e):
                    # Fallback insert without assignee_email
                    cur.execute("""
                        INSERT INTO bid_assign (g_id, b_name, due_date, state, scope, type, company, depart, person_name, status, value)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s)
                    """, (
                        gb['g_id'], gb['b_name'], gb['due_date'], new_stage, gb['scope'], gb['type'], gb['company'], depart, person_name,
                        gb.get('scoring', 0)
                    ))
                else:
                    raise e
        else:
            # Map department to stage for go_bids.state update
            dept_to_stage = {
                'business dev': 'business',
                'business': 'business',
                'design': 'design',
                'operations': 'operations',
                'site manager': 'engineer',
                'engineer': 'engineer'
            }
            new_stage = dept_to_stage.get(depart.lower(), depart.lower())
            
            # Update go_bids.state to match assigned department
            cur.execute("UPDATE go_bids SET state=%s WHERE g_id=%s", (new_stage, g_id))
            
            # Check if assignee_email column exists before updating
            try:
                cur.execute("UPDATE bid_assign SET depart=%s, person_name=%s, assignee_email=%s, state=%s, status='assigned' WHERE g_id=%s", 
                           (depart, person_name, email_to, new_stage, g_id))
            except Exception as e:
                if "Unknown column 'assignee_email'" in str(e):
                    # Fallback update without assignee_email
                    cur.execute("UPDATE bid_assign SET depart=%s, person_name=%s, state=%s, status='assigned' WHERE g_id=%s", 
                               (depart, person_name, new_stage, g_id))
                else:
                    raise e
        # Log the assignment action
        cur.execute("SELECT b_name FROM go_bids WHERE g_id=%s", (g_id,))
        bid_name = cur.fetchone()['b_name']
        log_action = f"Admin '{current_user.email}' assigned bid '{bid_name}' (ID: {g_id}) to {depart} department - {person_name} ({email_to})"
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)", (log_action, current_user.id))
        
        mysql.connection.commit()
        
        # Emit Socket.IO update for real-time master dashboard sync
        socketio.emit('master_update', {
            'bid': {
                'id': g_id,
                'name': bid_name,
                'current_stage': new_stage,
                'assigned_to': person_name,
                'department': depart
            },
            'log': {
                'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'action': log_action,
                'user_email': current_user.email,
                'user_role': 'admin'
            },
            'assignment': True
        })
        
        flash('Bid assigned to person successfully.', 'success')
        # Send email notification if provided
        if email_to:
            try:
                send_assignment_email(email_to, gb['b_name'] if 'gb' in locals() else None, depart, person_name, gb['company'] if 'gb' in locals() else None)
            except Exception as _e:
                # Non-fatal
                flash('Assignment email could not be sent.', 'error')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Assignment failed: {str(e)}', 'error')
    finally:
        cur.close()
    return redirect(url_for('database_management', table='go_bids'))

# --- Email helper (adapted from mail_send.py) ---
SMTP_PORT = 587
SMTP_SERVER = "smtp.gmail.com"
EMAIL_FROM = "manuj@metcoengineering.com"
EMAIL_PASSWORD = "ksec iqja bmdg hrcv"

def send_assignment_email(email_to: str, bid_name: str, depart: str, person_name: str, company: str):
    subject = f"Bid Assignment: {bid_name or ''}"
    body = f"You have been assigned to bid '{bid_name or ''}' for company '{company or ''}'.\nDepartment: {depart}\nAssignee: {person_name}\n\nPlease log in to the ESCO suite to proceed."
    message = f"Subject: {subject}\n\n{body}"
    context = ssl.create_default_context()
    server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
    try:
        server.starttls(context=context)
        server.login(EMAIL_FROM, EMAIL_PASSWORD)
        server.sendmail(EMAIL_FROM, email_to, message)
    finally:
        try:
            server.quit()
        except Exception:
            pass
# --- Employee Management Routes ---
@app.route('/team/<team>/employees')
@login_required
def team_employees(team):
    """Manage employees for a specific team"""
    if current_user.is_admin:
        return redirect(url_for('master_dashboard'))
    
    # Map team names to stages
    team_to_stage = {
        'business': 'business',
        'design': 'design', 
        'operations': 'operations',
        'engineer': 'engineer'
    }
    
    if team not in team_to_stage:
        return "Invalid team", 404
    
    cur = mysql.connection.cursor(DictCursor)
    
    # Get employees for this team
    cur.execute("""
        SELECT e.*, u.email as team_lead_email
        FROM employees e
        LEFT JOIN users u ON e.team_lead_id = u.id
        WHERE e.department = %s AND e.is_active = TRUE
        ORDER BY e.name
    """, (team,))
    employees = cur.fetchall()
    
    # Get team leads. Map team -> acceptable user roles, and for non-admins
    # restrict to the current manager so the dropdown shows their email only.
    roles_for_team = {
        'business': ('business dev', 'business', 'bdm'),
        'design': ('design',),
        'operations': ('operations',),
        'engineer': ('site manager', 'engineer'),
    }
    acceptable_roles = roles_for_team.get(team, (team,))
    if getattr(current_user, 'is_admin', False):
        placeholders = ','.join(['%s'] * len(acceptable_roles))
        cur.execute(f"SELECT * FROM users WHERE role IN ({placeholders})", acceptable_roles)
        team_leads = cur.fetchall()
    else:
        # Only the current manager on their own dashboard
        cur.execute("SELECT * FROM users WHERE id=%s", (current_user.id,))
        team_leads = cur.fetchall()
    
    cur.close()
    
    return render_template('team_employees.html', 
                         team=team, 
                         employees=employees, 
                         team_leads=team_leads,
                         user=current_user)

@app.route('/team/<team>/employees/create', methods=['POST'])
@login_required
def create_employee(team):
    """Create a new employee for the team"""
    try:
        name = request.form.get('name', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '').strip()
        team_lead_id = request.form.get('team_lead_id') or (current_user.id if not getattr(current_user, 'is_admin', False) else None)
        
        if not name or not email or not password:
            flash('Name, email, and password are required', 'error')
            return redirect(url_for('team_employees', team=team))
        
        cur = mysql.connection.cursor()
        cur.execute("""
            INSERT INTO employees (name, email, password, department, team_lead_id) 
            VALUES (%s, %s, %s, %s, %s)
        """, (name, email, password, team, team_lead_id if team_lead_id else None))
        
        mysql.connection.commit()
        cur.close()
        
        log_write('create_employee', f"Created employee {name} for {team} team")
        flash(f'Employee "{name}" created successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error creating employee: {str(e)}', 'error')
    
    return redirect(url_for('team_employees', team=team))

@app.route('/employee/<int:employee_id>/dashboard')
def employee_dashboard(employee_id):
    """Employee-specific dashboard showing assigned tasks"""
    # Check if employee is logged in and matches the requested employee
    if 'employee_id' not in session or session['employee_id'] != employee_id:
        return "Access denied. Please login first.", 403
    
    cur = mysql.connection.cursor(DictCursor)
    
    # Get employee info
    cur.execute("SELECT * FROM employees WHERE id = %s", (employee_id,))
    employee = cur.fetchone()
    
    if not employee:
        cur.close()
        return "Employee not found", 404
    
    # Get assigned tasks for this employee
    cur.execute("""
        SELECT bc.*, gb.b_name, gb.company, gb.due_date as bid_due_date
        FROM bid_checklists bc
        JOIN go_bids gb ON bc.g_id = gb.g_id
        WHERE bc.assigned_to = %s
        ORDER BY bc.due_date ASC, bc.priority DESC
    """, (employee_id,))
    tasks = cur.fetchall()
    
    # Calculate task statistics
    total_tasks = len(tasks)
    completed_tasks = len([t for t in tasks if t['status'] == 'completed'])
    pending_tasks = len([t for t in tasks if t['status'] == 'pending'])
    in_progress_tasks = len([t for t in tasks if t['status'] == 'in_progress'])
    
    cur.close()
    
    return render_template('employee_dashboard.html',
                         employee=employee,
                         tasks=tasks,
                         total_tasks=total_tasks,
                         completed_tasks=completed_tasks,
                         pending_tasks=pending_tasks,
                         in_progress_tasks=in_progress_tasks,
                         user=current_user)

@app.route('/task/<int:task_id>/update_status', methods=['POST'])
def update_task_status(task_id):
    """Update task status and progress"""
    try:
        new_status = request.form.get('status', '').strip()
        progress_notes = request.form.get('progress_notes', '').strip()
        
        if not new_status:
            return jsonify({'error': 'Status is required'}), 400
        
        cur = mysql.connection.cursor(DictCursor)
        
        # Get task info
        cur.execute("""
            SELECT bc.*, gb.b_name, gb.company, e.name as employee_name
            FROM bid_checklists bc
            JOIN go_bids gb ON bc.g_id = gb.g_id
            JOIN employees e ON bc.assigned_to = e.id
            WHERE bc.id = %s
        """, (task_id,))
        task = cur.fetchone()
        
        if not task:
            cur.close()
            return jsonify({'error': 'Task not found'}), 404
        
        # Authorization: allow admins/managers (flask-login) OR the assigned employee via session
        is_flask_user = hasattr(current_user, 'is_authenticated') and current_user.is_authenticated
        employee_id = session.get('employee_id')
        if not is_flask_user:
            # When not logged in via Flask-Login, require employee session and ownership of the task
            if not employee_id or task.get('assigned_to') != employee_id:
                cur.close()
                return jsonify({'error': 'Forbidden'}), 403

        # Update task status (and map to a percentage for persistence)
        status_lower = (new_status or '').strip().lower()
        if status_lower == 'completed':
            pct_val = 100
        elif status_lower == 'in_progress':
            pct_val = 50
        else:
            pct_val = 0
        cur.execute("""
            UPDATE bid_checklists 
            SET status = %s, progress_pct = %s, updated_at = CURRENT_TIMESTAMP
            WHERE id = %s
        """, (new_status, pct_val, task_id))
        
        # Log the update
        log_write('task_update', 
                 f"Task '{task['task_name']}' for bid '{task['b_name']}' updated to {new_status} by {task['employee_name']}")
        
        mysql.connection.commit()
        cur.close()
        
        # Emit real-time update to team dashboard
        socketio.emit('task_update', {
            'task_id': task_id,
            'status': new_status,
            'bid_name': task['b_name'],
            'employee_name': task['employee_name'],
            'company': task['company']
        })

        # Also emit master_update with per-stage progress for this bid
        try:
            bid_id = task['g_id']
            cur2 = mysql.connection.cursor(DictCursor)
            # Compute per-team completion rates by tasks
            def pct_for(role_expr):
                cur2.execute(f"SELECT status FROM bid_checklists bc JOIN users u ON bc.created_by=u.id WHERE bc.g_id=%s AND u.role {role_expr}", (bid_id,))
                rows = cur2.fetchall()
                if not rows:
                    return 0
                done = len([r for r in rows if (r.get('status') or '').lower()=='completed'])
                return int(round((done/max(1,len(rows)))*100))
            stage_progress_map = {
                'business': pct_for("='business dev'"),
                'design': pct_for("='design'"),
                'operations': pct_for("='operations'"),
                'engineer': pct_for("IN ('site manager','engineer')"),
            }
            cur2.close()
            socketio.emit('master_update', {
                'summary': {
                    'bid_id': bid_id,
                    'work_progress_pct': stage_progress_map.get('design',0),
                    'project_status': 'ongoing',
                    'work_status': f"Task '{task['task_name']}' -> {new_status}",
                    'stage_progress_map': stage_progress_map
                }
            })
        except Exception:
            pass
        
        return jsonify({'success': 'Task status updated successfully'})
        
    except Exception as e:
        mysql.connection.rollback()
        if 'cur' in locals():
            cur.close()
        return jsonify({'error': f'Error updating task: {str(e)}'}), 500

@app.route('/task/<int:task_id>/update', methods=['POST'])
def update_task(task_id):
    """Update task fields (name, description, status, due_date, priority, assigned_to)."""
    try:
        cur = mysql.connection.cursor(DictCursor)
        # Fetch existing task
        cur.execute("SELECT * FROM bid_checklists WHERE id=%s", (task_id,))
        task = cur.fetchone()
        if not task:
            cur.close()
            return jsonify({'error': 'Task not found'}), 404

        # Authorization: admin/manager or assigned employee via session
        is_flask_user = hasattr(current_user, 'is_authenticated') and current_user.is_authenticated
        employee_id = session.get('employee_id')
        if not is_flask_user and (not employee_id or task.get('assigned_to') != employee_id):
            cur.close()
            return jsonify({'error': 'Forbidden'}), 403

        # Collect fields (optional updates)
        fields = []
        values = []
        m = request.form
        if 'task_name' in m:
            fields.append('task_name=%s'); values.append(m.get('task_name').strip())
        if 'description' in m:
            fields.append('description=%s'); values.append(m.get('description').strip())
        if 'status' in m:
            fields.append('status=%s'); values.append(m.get('status').strip())
        if 'due_date' in m:
            fields.append('due_date=%s'); values.append(m.get('due_date'))
        if 'priority' in m:
            fields.append('priority=%s'); values.append(m.get('priority').strip())
        if 'assigned_to' in m:
            at = m.get('assigned_to') or None
            fields.append('assigned_to=%s'); values.append(at)
        if 'progress_pct' in m:
            try:
                pp = int(m.get('progress_pct'))
                pp = max(0, min(100, pp))
            except Exception:
                pp = 0
            fields.append('progress_pct=%s'); values.append(pp)
        if not fields:
            cur.close()
            return jsonify({'error': 'No fields to update'}), 400
        set_clause = ', '.join(fields) + ', updated_at = CURRENT_TIMESTAMP'
        values.append(task_id)
        cur.execute(f"UPDATE bid_checklists SET {set_clause} WHERE id=%s", tuple(values))
        mysql.connection.commit()
        cur.close()
        log_write('task_update_fields', f"task_id={task_id}")
        return jsonify({'success': 'Task updated'})
    except Exception as e:
        mysql.connection.rollback()
        if 'cur' in locals():
            cur.close()
        return jsonify({'error': f'Error updating task: {str(e)}'}), 500

@app.route('/task/<int:task_id>/delete', methods=['POST'])
def delete_task(task_id):
    """Delete a task if authorized."""
    try:
        cur = mysql.connection.cursor(DictCursor)
        cur.execute("SELECT assigned_to FROM bid_checklists WHERE id=%s", (task_id,))
        row = cur.fetchone()
        if not row:
            cur.close()
            return jsonify({'error': 'Task not found'}), 404
        is_flask_user = hasattr(current_user, 'is_authenticated') and current_user.is_authenticated
        employee_id = session.get('employee_id')
        if not is_flask_user and (not employee_id or row.get('assigned_to') != employee_id):
            cur.close()
            return jsonify({'error': 'Forbidden'}), 403
        cur.execute("DELETE FROM bid_checklists WHERE id=%s", (task_id,))
        mysql.connection.commit()
        cur.close()
        log_write('task_delete', f"task_id={task_id}")
        return jsonify({'success': 'Task deleted'})
    except Exception as e:
        mysql.connection.rollback()
        if 'cur' in locals():
            cur.close()
        return jsonify({'error': f'Error deleting task: {str(e)}'}), 500

@app.route('/team/<team>/bids/<int:g_id>/checklist')
@login_required
def bid_checklist(team, g_id):
    """Manage checklist/tasks for a specific bid"""
    cur = mysql.connection.cursor(DictCursor)
    
    # Get bid info
    cur.execute("SELECT * FROM go_bids WHERE g_id = %s", (g_id,))
    bid = cur.fetchone()
    
    if not bid:
        cur.close()
        return "Bid not found", 404
    
    # Get checklist items for this bid
    cur.execute("""
        SELECT bc.*, e.name as assigned_employee_name
        FROM bid_checklists bc
        LEFT JOIN employees e ON bc.assigned_to = e.id
        WHERE bc.g_id = %s
        ORDER BY bc.priority DESC, bc.created_at ASC
    """, (g_id,))
    checklist_items = cur.fetchall()
    
    # Get team employees for assignment
    cur.execute("""
        SELECT * FROM employees 
        WHERE department = %s AND is_active = TRUE
        ORDER BY name
    """, (team,))
    team_employees = cur.fetchall()
    
    cur.close()
    
    return render_template('bid_checklist.html',
                         team=team,
                         bid=bid,
                         checklist_items=checklist_items,
                         team_employees=team_employees,
                         user=current_user)

@app.route('/team/<team>/bids/<int:g_id>/checklist/create', methods=['POST'])
@login_required
def create_checklist_item(team, g_id):
    """Create a new checklist item for a bid"""
    try:
        task_name = request.form.get('task_name', '').strip()
        description = request.form.get('description', '').strip()
        assigned_to = request.form.get('assigned_to')
        priority = request.form.get('priority', 'medium')
        due_date = request.form.get('due_date')
        file_obj = request.files.get('attachment')
        saved_path = None
        
        if not task_name:
            flash('Task name is required', 'error')
            return redirect(url_for('bid_checklist', team=team, g_id=g_id))
        
        cur = mysql.connection.cursor()
        # Explicit stage name for parallel tracking
        stage_name = team.strip().lower()
        # Handle optional attachment upload
        if file_obj and getattr(file_obj, 'filename', None):
            try:
                os.makedirs(os.path.join(os.getcwd(), 'uploads', 'checklists', str(g_id)), exist_ok=True)
                safe_name = secure_filename(file_obj.filename)
                saved_path = os.path.join('uploads', 'checklists', str(g_id), safe_name)
                file_obj.save(os.path.join(os.getcwd(), saved_path))
            except Exception:
                saved_path = None
        cur.execute("""
            INSERT INTO bid_checklists (g_id, task_name, description, assigned_to, priority, due_date, progress_pct, stage, created_by, attachment_path)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (g_id, task_name, description, assigned_to if assigned_to else None, 
              priority, due_date if due_date else None, 0, stage_name, current_user.id, saved_path))
        
        mysql.connection.commit()
        cur.close()
        
        log_write('create_checklist_item', f"Created task '{task_name}' for bid {g_id}")
        flash(f'Task "{task_name}" created successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error creating task: {str(e)}', 'error')
    
    return redirect(url_for('bid_checklist', team=team, g_id=g_id))

@app.route('/task/<int:task_id>/attach', methods=['POST'])
@login_required
def attach_file_to_task(task_id):
    """Attach or replace a file for a checklist task."""
    try:
        file_obj = request.files.get('attachment')
        if not file_obj or not getattr(file_obj, 'filename', None):
            return jsonify({'error': 'No file provided'}), 400
        cur = mysql.connection.cursor(DictCursor)
        cur.execute("SELECT id, g_id, assigned_to FROM bid_checklists WHERE id=%s", (task_id,))
        task = cur.fetchone()
        if not task:
            cur.close()
            return jsonify({'error': 'Task not found'}), 404
        # Authorization: admin/manager or assigned employee via session
        is_flask_user = hasattr(current_user, 'is_authenticated') and current_user.is_authenticated
        employee_id = session.get('employee_id')
        if not is_flask_user and (not employee_id or task.get('assigned_to') != employee_id):
            cur.close()
            return jsonify({'error': 'Forbidden'}), 403
        g_id = task['g_id']
        os.makedirs(os.path.join(os.getcwd(), 'uploads', 'checklists', str(g_id)), exist_ok=True)
        safe_name = secure_filename(file_obj.filename)
        saved_rel = os.path.join('uploads', 'checklists', str(g_id), safe_name)
        file_obj.save(os.path.join(os.getcwd(), saved_rel))
        cur.execute("UPDATE bid_checklists SET attachment_path=%s, updated_at=CURRENT_TIMESTAMP WHERE id=%s", (saved_rel, task_id))
        mysql.connection.commit()
        cur.close()
        return jsonify({'success': True, 'attachment_path': saved_rel})
    except Exception as e:
        mysql.connection.rollback()
        if 'cur' in locals():
            cur.close()
        return jsonify({'error': str(e)}), 500

@app.route('/task/<int:task_id>/attachment')
@login_required
def get_task_attachment(task_id):
    """Serve the attachment file for a task if present."""
    try:
        cur = mysql.connection.cursor(DictCursor)
        cur.execute("SELECT attachment_path FROM bid_checklists WHERE id=%s", (task_id,))
        row = cur.fetchone()
        cur.close()
        path = (row or {}).get('attachment_path') if row else None
        if not path:
            return "No attachment", 404
        abs_path = os.path.join(os.getcwd(), path)
        directory = os.path.dirname(abs_path)
        filename = os.path.basename(abs_path)
        if not os.path.exists(abs_path):
            return "Attachment missing on server", 404
        return send_from_directory(directory, filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}", 500

@app.route('/team/<team>/transfer_project', methods=['POST'])
@login_required
def transfer_project(team):
    """Transfer a project to another team"""
    try:
        g_id = request.form.get('g_id')
        to_team = request.form.get('to_team')
        transfer_reason = request.form.get('transfer_reason', '').strip()
        
        if not g_id or not to_team:
            flash('Project and destination team are required', 'error')
            return redirect(url_for('team_dashboard', team=team))
        
        cur = mysql.connection.cursor(DictCursor)
        
        # Get bid info
        cur.execute("SELECT * FROM go_bids WHERE g_id = %s", (g_id,))
        bid = cur.fetchone()
        
        if not bid:
            cur.close()
            flash('Bid not found', 'error')
            return redirect(url_for('team_dashboard', team=team))
        
        # Create transfer record
        cur.execute("""
            INSERT INTO project_transfers (g_id, from_team, to_team, transferred_by, transfer_reason)
            VALUES (%s, %s, %s, %s, %s)
        """, (g_id, team, to_team, current_user.id, transfer_reason))
        
        # Update bid state to next team
        cur.execute("UPDATE go_bids SET state = %s WHERE g_id = %s", (to_team, g_id))

        # When transferring, preserve current team's tasks and generate new tasks for receiving team
        # First, mark current team's tasks as completed and archive them
        cur.execute("""
            UPDATE bid_checklists 
            SET status = 'completed', progress_pct = 100, team_archive = %s, updated_at = CURRENT_TIMESTAMP
            WHERE g_id = %s AND created_by IN (
                SELECT id FROM users WHERE role = %s
            )
        """, (team, g_id, team))
        
        # Generate team-specific default checklist for the receiving team
        generate_team_checklist(cur, g_id, to_team)
        
        mysql.connection.commit()
        cur.close()
        
        log_write('project_transfer', 
                 f"Transferred bid '{bid['b_name']}' from {team} to {to_team}")
        flash(f'Project "{bid["b_name"]}" transferred to {to_team} team successfully!', 'success')
        
    except Exception as e:
        mysql.connection.rollback()
        cur.close()
        flash(f'Error transferring project: {str(e)}', 'error')
    
    return redirect(url_for('team_dashboard', team=team))

# --- RFP File Routes ---

def _ensure_uploaded_rfp_table_exists(cur):
    """Ensure the uploaded_rfp_files table exists before querying."""
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS uploaded_rfp_files (
            id INT AUTO_INCREMENT PRIMARY KEY,
            bid_id INT,
            g_id INT,
            filename VARCHAR(500) NOT NULL,
            file_path VARCHAR(1000) NOT NULL,
            file_size BIGINT,
            uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            uploaded_by INT,
            INDEX idx_bid_id (bid_id),
            INDEX idx_g_id (g_id)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """
    )
    mysql.connection.commit()

    def _column_exists(column_name: str) -> bool:
        cur.execute("SHOW COLUMNS FROM uploaded_rfp_files LIKE %s", (column_name,))
        return cur.fetchone() is not None

    def _index_exists(index_name: str) -> bool:
        cur.execute("SHOW INDEX FROM uploaded_rfp_files WHERE Key_name = %s", (index_name,))
        return cur.fetchone() is not None

    def _ensure_column(column_name: str, ddl: str):
        if _column_exists(column_name):
            return
        try:
            cur.execute(f"ALTER TABLE uploaded_rfp_files ADD COLUMN {ddl}")
            mysql.connection.commit()
        except Exception:
            mysql.connection.rollback()

    def _ensure_index(index_name: str, ddl: str):
        if _index_exists(index_name):
            return
        try:
            cur.execute(f"ALTER TABLE uploaded_rfp_files ADD INDEX {ddl}")
            mysql.connection.commit()
        except Exception:
            mysql.connection.rollback()

    # Ensure legacy tables have required columns and indexes, ignoring duplicate errors safely
    _ensure_column('bid_id', "bid_id INT AFTER id")
    _ensure_index('idx_bid_id', "idx_bid_id (bid_id)")

    _ensure_column('g_id', "g_id INT AFTER bid_id")
    _ensure_index('idx_g_id', "idx_g_id (g_id)")

    _ensure_column('filename', "filename VARCHAR(500) DEFAULT NULL")
    _ensure_column('original_filename', "original_filename VARCHAR(500) DEFAULT NULL")
    _ensure_column('saved_filename', "saved_filename VARCHAR(500) DEFAULT NULL")
    _ensure_column('file_type', "file_type VARCHAR(50) DEFAULT 'pdf'")
    _ensure_column('file_hash', "file_hash VARCHAR(128) DEFAULT NULL")
    _ensure_column('file_size', "file_size BIGINT DEFAULT NULL")
    _ensure_column('uploaded_at', "uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP")
    _ensure_column('uploaded_by', "uploaded_by INT DEFAULT NULL")
    _ensure_column('file_path', "file_path VARCHAR(1000) DEFAULT NULL")


def _get_latest_rfp_file_for_bid(g_id):
    """Fetch the most recent RFP file record for a bid using g_id."""
    cur = mysql.connection.cursor(DictCursor)
    try:
        _ensure_uploaded_rfp_table_exists(cur)

        # Inspect available columns to support legacy schemas
        cur.execute("SHOW COLUMNS FROM uploaded_rfp_files")
        column_rows = cur.fetchall() or []
        columns = {row['Field'] for row in column_rows if 'Field' in row}

        has_g_id = 'g_id' in columns
        has_bid_id = 'bid_id' in columns
        has_filename = 'filename' in columns
        has_original_filename = 'original_filename' in columns
        has_saved_filename = 'saved_filename' in columns
        has_file_path = 'file_path' in columns
        has_file_size = 'file_size' in columns
        has_uploaded_at = 'uploaded_at' in columns

        select_fields = ['id']
        if has_g_id:
            select_fields.append('g_id')
        else:
            select_fields.append('NULL AS g_id')
        if has_bid_id:
            select_fields.append('bid_id')
        else:
            select_fields.append('NULL AS bid_id')

        if has_filename:
            select_fields.append('filename')
        else:
            # Fall back to original/saved filename columns if present
            coalesce_parts = []
            if has_original_filename:
                coalesce_parts.append('original_filename')
            if has_saved_filename:
                coalesce_parts.append('saved_filename')
            if coalesce_parts:
                select_fields.append(f"COALESCE({', '.join(coalesce_parts)}) AS filename")
            else:
                select_fields.append("NULL AS filename")

        if has_file_path:
            select_fields.append('file_path')
        else:
            select_fields.append('NULL AS file_path')

        if has_file_size:
            select_fields.append('file_size')
        else:
            select_fields.append('NULL AS file_size')

        if has_uploaded_at:
            select_fields.append('uploaded_at')
            order_clause = "uploaded_at DESC, id DESC"
        else:
            # Legacy tables might use created_at/updated_at; prefer whichever exists
            timestamp_col = None
            for legacy_col in ('created_at', 'updated_at', 'saved_at'):
                if legacy_col in columns:
                    timestamp_col = legacy_col
                    break
            if timestamp_col:
                select_fields.append(f"{timestamp_col} AS uploaded_at")
                order_clause = f"{timestamp_col} DESC, id DESC"
            else:
                select_fields.append('NULL AS uploaded_at')
                order_clause = "id DESC"

        where_clauses = []
        params = []
        if has_g_id:
            where_clauses.append("g_id = %s")
            params.append(g_id)
        if has_bid_id:
            where_clauses.append("bid_id = %s")
            params.append(g_id)

        if not where_clauses:
            # As a last resort, attempt to match against legacy columns such as bid or project id
            for legacy_col in ('b_id', 'project_id'):
                if legacy_col in columns:
                    where_clauses.append(f"{legacy_col} = %s")
                    params.append(g_id)
                    break

        if not where_clauses:
            return None

        select_clause = ", ".join(select_fields)
        where_clause_sql = " OR ".join(f"({clause})" for clause in where_clauses)

        try:
            cur.execute(
                f"""
                SELECT {select_clause}
                FROM uploaded_rfp_files
                WHERE {where_clause_sql}
                ORDER BY {order_clause}
                LIMIT 1
                """,
                tuple(params),
            )
        except Exception:
            # The table may not exist yet—create it and retry once.
            mysql.connection.rollback()
            _ensure_uploaded_rfp_table_exists(cur)
            cur.execute(
                f"""
                SELECT {select_clause}
                FROM uploaded_rfp_files
                WHERE {where_clause_sql}
                ORDER BY {order_clause}
                LIMIT 1
                """,
                tuple(params),
            )
        row = cur.fetchone()

        if row and 'id' in row:
            update_clauses = []
            update_params = []
            # Backfill missing identifiers for legacy rows
            if has_g_id and (not row.get('g_id')):
                update_clauses.append("g_id = %s")
                update_params.append(g_id)
                row['g_id'] = g_id
            if has_bid_id and (not row.get('bid_id')):
                update_clauses.append("bid_id = %s")
                update_params.append(g_id)
                row['bid_id'] = g_id

            if update_clauses:
                update_params.append(row['id'])
                try:
                    cur.execute(
                        f"UPDATE uploaded_rfp_files SET {', '.join(update_clauses)} WHERE id = %s",
                        tuple(update_params),
                    )
                    mysql.connection.commit()
                except Exception:
                    mysql.connection.rollback()

        return row
    finally:
        cur.close()


def _extract_pdf_pages(file_path, start_page=1, limit=None):
    """Extract text from a PDF file path using PyMuPDF or fall back to PyPDF2."""
    pages = []
    total_pages = 0
    start_index = max(start_page - 1, 0)
    use_limit = None if (limit is None or limit <= 0) else limit

    if _FITZ_AVAILABLE:
        with _pymupdf.open(file_path) as doc:
            total_pages = doc.page_count
            if start_index >= total_pages:
                return pages, total_pages
            end_index = total_pages if use_limit is None else min(total_pages, start_index + use_limit)
            for index in range(start_index, end_index):
                page = doc.load_page(index)
                text = page.get_text("text") or ""
                pages.append(
                    {
                        "number": index + 1,
                        "text": text,
                        "characters": len(text),
                    }
                )
        return pages, total_pages

    if _PdfReader is not None:
        with open(file_path, "rb") as fh:
            reader = _PdfReader(fh)
            total_pages = len(reader.pages)
            if start_index >= total_pages:
                return pages, total_pages
            end_index = total_pages if use_limit is None else min(total_pages, start_index + use_limit)
            for index in range(start_index, end_index):
                page = reader.pages[index]
                text = page.extract_text() or ""
                pages.append(
                    {
                        "number": index + 1,
                        "text": text,
                        "characters": len(text),
                    }
                )
        return pages, total_pages

    raise RuntimeError(
        "PDF text extraction requires PyMuPDF (`pip install PyMuPDF`) or PyPDF2 (`pip install PyPDF2`)."
    )


@app.route('/api/rfp-file/<int:g_id>/upload', methods=['POST'])
@login_required
def api_upload_rfp_file(g_id):
    """Upload a PDF for the specified g_id and store metadata in uploaded_rfp_files."""
    upload_file = request.files.get('file')
    if not upload_file or not upload_file.filename:
        return jsonify({'error': 'missing_file', 'message': 'Please choose a PDF to upload.'}), 400

    filename_lower = upload_file.filename.lower()
    if not filename_lower.endswith('.pdf'):
        return jsonify({'error': 'invalid_type', 'message': 'Only PDF files are supported.'}), 400

    from werkzeug.utils import secure_filename
    import uuid
    import hashlib

    safe_original = secure_filename(upload_file.filename) or f"rfp_{uuid.uuid4().hex}.pdf"
    unique_token = uuid.uuid4().hex
    saved_filename = f"{unique_token}_{safe_original}"

    uploads_dir = os.path.join(os.getcwd(), 'uploads', 'rfp')
    os.makedirs(uploads_dir, exist_ok=True)
    absolute_path = os.path.join(uploads_dir, saved_filename)

    try:
        upload_file.save(absolute_path)
    except Exception as err:
        return jsonify({'error': 'save_failed', 'message': f'Could not save file: {err}'}), 500

    file_size = 0
    file_hash = ''
    try:
        file_size = os.path.getsize(absolute_path)
        with open(absolute_path, 'rb') as handler:
            file_hash = hashlib.sha256(handler.read()).hexdigest()
    except Exception:
        pass

    # Lookup bid_id (legacy) from go_bids table if available
    bid_id = None
    cur_lookup = mysql.connection.cursor(DictCursor)
    try:
        cur_lookup.execute("SELECT id FROM go_bids WHERE g_id = %s", (g_id,))
        go_bid_row = cur_lookup.fetchone()
        if go_bid_row and go_bid_row.get('id'):
            bid_id = go_bid_row['id']
    finally:
        cur_lookup.close()

    cur = mysql.connection.cursor(DictCursor)
    try:
        _ensure_uploaded_rfp_table_exists(cur)
        cur.execute("SHOW COLUMNS FROM uploaded_rfp_files")
        column_rows = cur.fetchall() or []
        available_columns = {row['Field'] for row in column_rows if 'Field' in row}

        columns = []
        values = []

        def add_column(col_name, value):
            if col_name in available_columns:
                columns.append(col_name)
                values.append(value)

        add_column('bid_id', bid_id)
        add_column('g_id', g_id)
        add_column('filename', safe_original)
        add_column('original_filename', upload_file.filename)
        add_column('saved_filename', saved_filename)
        add_column('file_path', absolute_path)
        add_column('file_type', 'pdf')
        add_column('file_size', file_size)
        add_column('file_hash', file_hash)
        add_column('uploaded_by', getattr(current_user, 'id', None))

        if not columns:
            return jsonify({'error': 'schema_error', 'message': 'Unable to persist the uploaded file.'}), 500

        placeholders = ','.join(['%s'] * len(columns))
        column_sql = ','.join(columns)
        cur.execute(
            f"INSERT INTO uploaded_rfp_files ({column_sql}) VALUES ({placeholders})",
            tuple(values),
        )
        mysql.connection.commit()

        # Fetch parsed preview for immediate display
        parsed_pages = []
        total_pages = 0
        next_start = None
        try:
            parsed_pages, total_pages = _extract_pdf_pages(absolute_path, start_page=1, limit=3)
            if parsed_pages:
                last_number = parsed_pages[-1].get('number')
                if last_number and total_pages and last_number < total_pages:
                    next_start = last_number + 1
        except Exception:
            parsed_pages = []
            total_pages = 0
            next_start = None

        return jsonify(
            {
                'success': True,
                'file_id': cur.lastrowid,
                'filename': safe_original,
                'total_pages': total_pages,
                'next_start': next_start,
                'pages': parsed_pages,
            }
        )
    except Exception as err:
        mysql.connection.rollback()
        if os.path.exists(absolute_path):
            try:
                os.remove(absolute_path)
            except Exception:
                pass
        return jsonify({'error': 'upload_failed', 'message': str(err)}), 500
    finally:
        cur.close()


@app.route('/rfp-file/view/by-g/<int:g_id>')
@login_required
def view_rfp_file_by_g(g_id):
    """View the latest RFP file for a bid using its g_id."""
    try:
        file_data = _get_latest_rfp_file_for_bid(g_id)
    except Exception as err:
        app.logger.exception("Error locating RFP file for g_id %s: %s", g_id, err)
        abort(500, description="Failed to locate RFP file.")

    if not file_data:
        abort(404, description="No RFP file associated with this project.")

    file_path = file_data.get('file_path')
    filename = file_data.get('filename') or os.path.basename(file_path or '')

    if not file_path or not os.path.exists(file_path):
        abort(404, description="RFP file not found on server.")

    if file_path.lower().endswith('.pdf'):
        return send_file(file_path, mimetype='application/pdf', as_attachment=False, download_name=filename)

    return send_file(file_path, as_attachment=True, download_name=filename)


@app.route('/api/rfp-file/<int:g_id>/parsed', methods=['GET'])
@login_required
def api_rfp_file_parsed(g_id):
    """Return parsed PDF text for the latest RFP file associated with g_id."""
    start = request.args.get('start', default=1, type=int)
    limit = request.args.get('limit', default=3, type=int)
    page_limit = None if limit is None or limit <= 0 else limit
    start = start if start and start > 0 else 1

    try:
        file_data = _get_latest_rfp_file_for_bid(g_id)
    except Exception as err:
        app.logger.exception("Error locating RFP file for g_id %s: %s", g_id, err)
        return jsonify({'error': 'internal_error'}), 500

    if not file_data:
        return jsonify({'error': 'not_found', 'message': 'No RFP file associated with this project.'}), 404

    file_path = file_data.get('file_path')
    if not file_path or not os.path.exists(file_path):
        return jsonify({'error': 'file_missing', 'message': 'RFP file not found on server.'}), 404

    if not file_path.lower().endswith('.pdf'):
        return jsonify({'error': 'unsupported_type', 'message': 'Only PDF files can be parsed.'}), 400

    try:
        pages, total_pages = _extract_pdf_pages(file_path, start_page=start, limit=page_limit)
    except Exception as err:
        app.logger.exception("Error parsing PDF for g_id %s: %s", g_id, err)
        return jsonify({'error': 'parse_failed', 'message': 'Unable to parse the PDF file.'}), 500

    next_start = None
    if pages:
        last_page_number = pages[-1]['number']
        if last_page_number < total_pages:
            next_start = last_page_number + 1

    return jsonify(
        {
            'g_id': g_id,
            'file_id': file_data.get('id'),
            'filename': file_data.get('filename'),
            'total_pages': total_pages,
            'start': start,
            'limit': page_limit,
            'next_start': next_start,
            'pages': pages,
        }
    )


@app.route('/rfp-file/view/<int:file_id>')
@login_required
def view_rfp_file(file_id):
    """View RFP file in browser"""
    try:
        cur = mysql.connection.cursor(DictCursor)
        cur.execute("""
            SELECT file_path, filename
            FROM uploaded_rfp_files
            WHERE id = %s
        """, (file_id,))
        file_data = cur.fetchone()
        cur.close()
        
        if not file_data:
            flash('RFP file not found', 'error')
            return redirect(request.referrer or url_for('master_dashboard'))
        
        file_path = file_data['file_path']
        if not os.path.exists(file_path):
            flash('RFP file not found on server', 'error')
            return redirect(request.referrer or url_for('master_dashboard'))
        
        # For PDF files, send as inline to view in browser
        if file_path.lower().endswith('.pdf'):
            return send_file(file_path, mimetype='application/pdf', as_attachment=False)
        else:
            # For other file types, download
            return send_file(file_path, as_attachment=True)
    except Exception as e:
        flash(f'Error viewing file: {str(e)}', 'error')
        return redirect(request.referrer or url_for('master_dashboard'))

@app.route('/rfp-file/download/<int:file_id>')
@login_required
def download_rfp_file(file_id):
    """Download RFP file"""
    try:
        cur = mysql.connection.cursor(DictCursor)
        cur.execute("""
            SELECT file_path, filename
            FROM uploaded_rfp_files
            WHERE id = %s
        """, (file_id,))
        file_data = cur.fetchone()
        cur.close()
        
        if not file_data:
            flash('RFP file not found', 'error')
            return redirect(request.referrer or url_for('master_dashboard'))
        
        file_path = file_data['file_path']
        filename = file_data['filename']
        
        if not os.path.exists(file_path):
            flash('RFP file not found on server', 'error')
            return redirect(request.referrer or url_for('master_dashboard'))
        
        return send_file(file_path, as_attachment=True, download_name=filename)
    except Exception as e:
        flash(f'Error downloading file: {str(e)}', 'error')
        return redirect(request.referrer or url_for('master_dashboard'))


@app.route('/api/rfp-file/decrypt', methods=['POST'])
@login_required
def decrypt_rfp_file():
    """Decrypt an encrypted PDF file and return the decrypted version."""
    try:
        pdf_url = request.args.get('url')
        if not pdf_url:
            return jsonify({'error': 'URL parameter is required'}), 400

        # Extract file path from URL or fetch the file
        from urllib.parse import urlparse, unquote
        import tempfile
        
        # If it's a local file URL, extract the path
        if pdf_url.startswith('/'):
            # It's a relative URL, construct full path
            parsed = urlparse(pdf_url)
            # Try to find the file in the uploads directory
            if 'rfp-file/view' in pdf_url:
                # Extract file_id or g_id from URL
                import re
                match = re.search(r'/rfp-file/view/(\d+)', pdf_url)
                if match:
                    file_id = int(match.group(1))
                    cur = mysql.connection.cursor(DictCursor)
                    cur.execute("SELECT file_path FROM uploaded_rfp_files WHERE id = %s", (file_id,))
                    file_data = cur.fetchone()
                    cur.close()
                    if file_data and os.path.exists(file_data['file_path']):
                        file_path = file_data['file_path']
                    else:
                        return jsonify({'error': 'File not found'}), 404
                else:
                    # Try g_id
                    match = re.search(r'/rfp-file/view/by-g/(\d+)', pdf_url)
                    if match:
                        g_id = int(match.group(1))
                        file_data = _get_latest_rfp_file_for_bid(g_id)
                        if file_data and os.path.exists(file_data.get('file_path')):
                            file_path = file_data['file_path']
                        else:
                            return jsonify({'error': 'File not found'}), 404
                    else:
                        return jsonify({'error': 'Invalid URL format'}), 400
            else:
                return jsonify({'error': 'Invalid URL format'}), 400
        else:
            return jsonify({'error': 'Only local file URLs are supported'}), 400

        if not _FITZ_AVAILABLE:
            return jsonify({'error': 'PyMuPDF is not available for PDF decryption'}), 500

        # Attempt to decrypt the PDF
        try:
            # Try opening with empty password first (some PDFs have empty password)
            doc = _pymupdf.open(file_path)
            
            # If PDF is encrypted, try common passwords or empty password
            if doc.is_encrypted:
                # Try empty password
                if not doc.authenticate(""):
                    # Try some common passwords
                    common_passwords = ["", "password", "admin", "1234", "12345"]
                    authenticated = False
                    for pwd in common_passwords:
                        if doc.authenticate(pwd):
                            authenticated = True
                            break
                    
                    if not authenticated:
                        # Try to decrypt without password (some PDFs can be decrypted this way)
                        try:
                            doc.close()
                            # Create a new document by copying pages (this sometimes works for encrypted PDFs)
                            doc = _pymupdf.open(file_path)
                            if doc.is_encrypted and not doc.authenticate(""):
                                return jsonify({'error': 'PDF is password protected and cannot be decrypted automatically'}), 400
                        except:
                            return jsonify({'error': 'PDF is password protected and cannot be decrypted automatically'}), 400

            # Create a new decrypted PDF in memory
            decrypted_doc = _pymupdf.open()  # Create new empty PDF
            decrypted_doc.insert_pdf(doc)  # Copy all pages from encrypted to decrypted
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            temp_path = temp_file.name
            temp_file.close()
            
            decrypted_doc.save(temp_path)
            decrypted_doc.close()
            doc.close()
            
            # Return the decrypted PDF
            return send_file(temp_path, mimetype='application/pdf', as_attachment=False, download_name='decrypted_rfp.pdf')
            
        except Exception as decrypt_error:
            app.logger.exception("Error decrypting PDF: %s", decrypt_error)
            return jsonify({'error': f'Failed to decrypt PDF: {str(decrypt_error)}'}), 500

    except Exception as e:
        app.logger.exception("Error in decrypt_rfp_file: %s", e)
        return jsonify({'error': str(e)}), 500


@app.route('/api/rfp-file/<int:g_id>/save-edited', methods=['POST'])
@login_required
def save_edited_rfp_file(g_id):
    """Save an edited PDF file."""
    try:
        upload_file = request.files.get('file')
        if not upload_file or not upload_file.filename:
            return jsonify({'error': 'missing_file', 'message': 'Please provide a PDF file to save.'}), 400

        filename_lower = upload_file.filename.lower()
        if not filename_lower.endswith('.pdf'):
            return jsonify({'error': 'invalid_type', 'message': 'Only PDF files are supported.'}), 400

        from werkzeug.utils import secure_filename
        import uuid
        import hashlib

        safe_original = secure_filename(upload_file.filename) or f"edited_rfp_{uuid.uuid4().hex}.pdf"
        unique_token = uuid.uuid4().hex
        saved_filename = f"{unique_token}_{safe_original}"

        uploads_dir = os.path.join(os.getcwd(), 'uploads', 'rfp')
        os.makedirs(uploads_dir, exist_ok=True)
        absolute_path = os.path.join(uploads_dir, saved_filename)

        try:
            upload_file.save(absolute_path)
        except Exception as err:
            return jsonify({'error': 'save_failed', 'message': f'Could not save file: {err}'}), 500

        file_size = 0
        file_hash = ''
        try:
            file_size = os.path.getsize(absolute_path)
            with open(absolute_path, 'rb') as handler:
                file_hash = hashlib.sha256(handler.read()).hexdigest()
        except Exception:
            pass

        # Lookup bid_id (legacy) from go_bids table if available
        bid_id = None
        cur_lookup = mysql.connection.cursor(DictCursor)
        try:
            cur_lookup.execute("SELECT id FROM go_bids WHERE g_id = %s", (g_id,))
            go_bid_row = cur_lookup.fetchone()
            if go_bid_row and go_bid_row.get('id'):
                bid_id = go_bid_row['id']
        finally:
            cur_lookup.close()

        cur = mysql.connection.cursor(DictCursor)
        try:
            _ensure_uploaded_rfp_table_exists(cur)
            cur.execute("SHOW COLUMNS FROM uploaded_rfp_files")
            column_rows = cur.fetchall() or []
            available_columns = {row['Field'] for row in column_rows if 'Field' in row}

            columns = []
            values = []

            def add_column(col_name, value):
                if col_name in available_columns:
                    columns.append(col_name)
                    values.append(value)

            add_column('bid_id', bid_id)
            add_column('g_id', g_id)
            add_column('filename', safe_original)
            add_column('original_filename', upload_file.filename)
            add_column('saved_filename', saved_filename)
            add_column('file_path', absolute_path)
            add_column('file_type', 'pdf')
            add_column('file_size', file_size)
            add_column('file_hash', file_hash)
            add_column('uploaded_by', getattr(current_user, 'id', None))

            if not columns:
                return jsonify({'error': 'schema_error', 'message': 'Unable to persist the uploaded file.'}), 500

            placeholders = ','.join(['%s'] * len(columns))
            column_sql = ','.join(columns)
            cur.execute(
                f"INSERT INTO uploaded_rfp_files ({column_sql}) VALUES ({placeholders})",
                tuple(values),
            )
            mysql.connection.commit()

            return jsonify({
                'success': True,
                'message': 'Edited PDF saved successfully',
                'file_id': cur.lastrowid,
                'filename': safe_original
            })

        except Exception as db_error:
            mysql.connection.rollback()
            app.logger.exception("Database error saving edited PDF: %s", db_error)
            return jsonify({'error': 'database_error', 'message': f'Failed to save file metadata: {db_error}'}), 500
        finally:
            cur.close()

    except Exception as e:
        app.logger.exception("Error in save_edited_rfp_file: %s", e)
        return jsonify({'error': 'server_error', 'message': str(e)}), 500

# --- API Endpoints for Team Dashboard ---
@app.route('/api/team/<team>/employees')
@login_required
def api_team_employees(team):
    """API endpoint to get team employees"""
    cur = mysql.connection.cursor(DictCursor)
    
    cur.execute("""
        SELECT e.*, u.email as team_lead_email
        FROM employees e
        LEFT JOIN users u ON e.team_lead_id = u.id
        WHERE e.department = %s AND e.is_active = TRUE
        ORDER BY e.name
    """, (team,))
    employees = cur.fetchall()
    
    cur.close()
    
    return jsonify({'employees': employees})

@app.route('/api/team/<team>/assign', methods=['POST'])
@login_required
def api_team_assign(team):
    """Assign a bid to an employee from the current team.
    Body: JSON { g_id: int, employee_id: int }
    """
    try:
        data = request.get_json(silent=True) or {}
        g_id = int(data.get('g_id')) if data.get('g_id') is not None else None
        # accept single or multiple
        employee_ids = data.get('employee_ids')
        if employee_ids is None and data.get('employee_id') is not None:
            employee_ids = [data.get('employee_id')]
        if isinstance(employee_ids, list):
            try:
                employee_ids = [int(x) for x in employee_ids if x is not None]
            except Exception:
                employee_ids = []
        else:
            employee_ids = []
    except Exception:
        return jsonify({'ok': False, 'error': 'Invalid payload'}), 400

    if not g_id:
        return jsonify({'ok': False, 'error': 'g_id is required'}), 400

    cur = mysql.connection.cursor(DictCursor)
    try:
        # Validate employee belongs to this team
        # Validate employees belong to this team
        valid_emps = []
        if employee_ids:
            placeholders = ','.join(['%s'] * len(employee_ids))
            cur.execute(
                f"SELECT id, name, email FROM employees WHERE id IN ({placeholders}) AND department=%s AND is_active=TRUE",
                (*employee_ids, team)
            )
            valid_emps = cur.fetchall()
        # For legacy single assignment
        first_emp = (valid_emps[0] if valid_emps else None)

        # Get bid basic details
        cur.execute("SELECT * FROM go_bids WHERE g_id=%s", (g_id,))
        gb = cur.fetchone()
        if not gb:
            cur.close()
            return jsonify({'ok': False, 'error': 'Bid not found'}), 404

        # Upsert into bid_assign
        cur.execute("SELECT a_id FROM bid_assign WHERE g_id=%s", (g_id,))
        row = cur.fetchone()
        new_stage = team  # team name aligns to stage keys in our system
        if not row:
            try:
                # Full insert (newer schema with in_date, due_date, assignee_email)
                cur.execute(
                    """
                    INSERT INTO bid_assign (g_id, b_name, in_date, due_date, state, scope, type, company, depart, person_name, assignee_email, status, value, revenue)
                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s,%s)
                    """,
                    (
                        gb['g_id'], gb.get('b_name'), gb.get('in_date'), gb.get('due_date'), new_stage,
                        gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'), (first_emp or {}).get('email'),
                        gb.get('scoring', 0), gb.get('revenue', gb.get('scoring', 0))
                    )
                )
            except Exception as e:
                msg = str(e)
                if "Unknown column 'in_date'" in msg or "Unknown column 'due_date'" in msg:
                    # Older schema without date columns
                    try:
                        cur.execute(
                            """
                            INSERT INTO bid_assign (g_id, b_name, state, scope, type, company, depart, person_name, assignee_email, status, value, revenue)
                            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s,%s)
                            """,
                            (
                                gb['g_id'], gb.get('b_name'), new_stage,
                                gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'), (first_emp or {}).get('email'),
                                gb.get('scoring', 0), gb.get('revenue', gb.get('scoring', 0))
                            )
                        )
                    except Exception as e2:
                        if "Unknown column 'assignee_email'" in str(e2):
                            # Oldest schema: no assignee_email either
                            cur.execute(
                                """
                                INSERT INTO bid_assign (g_id, b_name, state, scope, type, company, depart, person_name, status, value, revenue)
                                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s,%s)
                                """,
                                (
                                    gb['g_id'], gb.get('b_name'), new_stage,
                                    gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'),
                                    gb.get('scoring', 0), gb.get('revenue', gb.get('scoring', 0))
                                )
                            )
                        else:
                            raise
                elif "Unknown column 'assignee_email'" in msg:
                    # No assignee_email but has in_date/due_date
                    try:
                        cur.execute(
                            """
                            INSERT INTO bid_assign (g_id, b_name, in_date, due_date, state, scope, type, company, depart, person_name, status, value, revenue)
                            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s,%s)
                            """,
                            (
                                gb['g_id'], gb.get('b_name'), gb.get('in_date'), gb.get('due_date'), new_stage,
                                gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'),
                                gb.get('scoring', 0), gb.get('revenue', gb.get('scoring', 0))
                            )
                        )
                    except Exception as e3:
                        # If that also fails due to dates missing, drop those too
                        if "Unknown column 'in_date'" in str(e3) or "Unknown column 'due_date'" in str(e3):
                            cur.execute(
                                """
                                INSERT INTO bid_assign (g_id, b_name, state, scope, type, company, depart, person_name, status, value, revenue)
                                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,'assigned',%s,%s)
                                """,
                                (
                                    gb['g_id'], gb.get('b_name'), new_stage,
                                gb.get('scope'), gb.get('type'), gb.get('company'), team, (first_emp or {}).get('name'),
                                    gb.get('scoring', 0), gb.get('revenue', gb.get('scoring', 0))
                                )
                            )
                        else:
                            raise
                else:
                    raise
        else:
            try:
                cur.execute(
                    "UPDATE bid_assign SET depart=%s, person_name=%s, assignee_email=%s, state=%s, status='assigned' WHERE g_id=%s",
                    (team, (first_emp or {}).get('name'), (first_emp or {}).get('email'), new_stage, g_id)
                )
            except Exception as e:
                if "Unknown column 'assignee_email'" in str(e):
                    cur.execute(
                        "UPDATE bid_assign SET depart=%s, person_name=%s, state=%s, status='assigned' WHERE g_id=%s",
                        (team, (first_emp or {}).get('name'), new_stage, g_id)
                    )
                else:
                    raise

        # Keep go_bids.state aligned with department
        cur.execute("UPDATE go_bids SET state=%s WHERE g_id=%s", (new_stage, g_id))

        # Update mapping table for multiple assignees (team-scoped replace)
        cur.execute(
            "DELETE FROM bid_assignment_members WHERE g_id=%s AND employee_id IN (SELECT id FROM employees WHERE department=%s)",
            (g_id, team)
        )
        if employee_ids:
            for emp_id in employee_ids:
                try:
                    cur.execute(
                        "INSERT IGNORE INTO bid_assignment_members (g_id, employee_id) VALUES (%s,%s)",
                        (g_id, emp_id)
                    )
                except Exception:
                    pass

        # Log + emit
        display_name = (first_emp or {}).get('name') or (f"{len(employee_ids)} member(s)" if employee_ids else 'Unassigned')
        log_action = f"Assigned bid '{gb.get('b_name')}' (ID: {g_id}) to {display_name} [{team}]"
        cur.execute("INSERT INTO logs (action, user_id) VALUES (%s, %s)", (log_action, current_user.id))
        mysql.connection.commit()

        try:
            socketio.emit('master_update', {
                'log': {'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'), 'action': log_action, 'user_email': current_user.email, 'user_role': getattr(current_user, 'role', '')},
                'assignment': True
            })
        except Exception:
            pass

        cur.close()
        return jsonify({'ok': True})
    except Exception as e:
        mysql.connection.rollback()
        try:
            cur.close()
        except Exception:
            pass
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/api/team/<team>/bids/<int:g_id>/tasks')
@login_required
def api_bid_tasks(team, g_id):
    """API endpoint to get tasks for a specific bid - only show tasks created by current team"""
    cur = mysql.connection.cursor(DictCursor)
    
    # Allow access if bid is in this team OR there are tasks for this team's stage
    cur.execute("SELECT state FROM go_bids WHERE g_id = %s", (g_id,))
    bid = cur.fetchone()
    cur.execute("SELECT 1 FROM bid_checklists WHERE g_id=%s AND LOWER(COALESCE(stage,''))=%s LIMIT 1", (g_id, team))
    has_team_tasks = cur.fetchone() is not None
    if not bid and not has_team_tasks:
        cur.close();
        return jsonify({'tasks': []})
    if bid and bid.get('state') != team and not has_team_tasks:
        cur.close();
        return jsonify({'tasks': []})
    
    # Get checklist items for this bid - show current team's active tasks and their archived tasks
    roles_for_team = {
        'business': ('business dev', 'business', 'bdm'),
        'design': ('design',),
        'operations': ('operations',),
        'engineer': ('site manager', 'engineer'),
    }
    acceptable_roles = roles_for_team.get(team, (team,))
    placeholders = ','.join(['%s'] * len(acceptable_roles))
    cur.execute(f"""
        SELECT bc.*, e.name as assigned_employee_name, e.email as employee_email, e.department
        FROM bid_checklists bc
        LEFT JOIN employees e ON bc.assigned_to = e.id
        LEFT JOIN users u ON bc.created_by = u.id
        WHERE bc.g_id = %s 
        AND (
            (u.role IN ({placeholders}) AND bc.team_archive IS NULL) OR  -- Current team's active tasks
            (bc.team_archive = %s)  -- This team's archived tasks
        )
        ORDER BY bc.priority DESC, bc.created_at ASC
    """, (g_id, *acceptable_roles, team))
    tasks = cur.fetchall()
    # Add attachment_url for easier consumption in UI
    try:
        for t in tasks:
            ap = (t.get('attachment_path') or '').strip() if isinstance(t, dict) else None
            t['attachment_url'] = url_for('get_task_attachment', task_id=t.get('id')) if ap else None
    except Exception:
        pass

    # Emit per-stage map for this bid so master dashboard reflects progress
    try:
        def _pct(rows):
            if not rows:
                return 0
            t = len(rows)
            d = len([r for r in rows if (r.get('status') or '').lower() == 'completed'])
            return int(round((d / max(1, t)) * 100))

        def _q(role_filter):
            cur.execute(f"SELECT status FROM bid_checklists WHERE g_id=%s AND created_by IN (SELECT id FROM users WHERE role {role_filter})", (g_id,))
            return cur.fetchall()

        spm = {
            'business': _pct(_q("='business dev'")),
            'design': _pct(_q("='design'")),
            'operations': _pct(_q("='operations'")),
            'engineer': _pct(_q("IN ('site manager','engineer')")),
        }
        socketio.emit('master_update', {
            'summary': {
                'bid_id': g_id,
                'work_progress_pct': spm.get(team, 0),
                'project_status': 'ongoing',
                'work_status': f'{team.title()} tasks updated',
                'stage_progress_map': spm
            }
        })
    except Exception:
        pass

    cur.close()
    return jsonify({'tasks': tasks})

# --- Simple API to update a bid summary (notes) ---
@app.route('/api/bids/<int:g_id>/summary', methods=['POST'])
@login_required
def api_update_bid_summary(g_id):
    try:
        summary = (request.form.get('summary') or '').strip()
        cur = mysql.connection.cursor()
        cur.execute("UPDATE go_bids SET summary=%s WHERE g_id=%s", (summary, g_id))
        mysql.connection.commit()
        cur.close()
        return jsonify({'ok': True})
    except Exception as e:
        try:
            mysql.connection.rollback()
        except Exception:
            pass
        return jsonify({'ok': False, 'error': str(e)}), 500

# Update bid dates (start/due)
@app.route('/api/bids/<int:g_id>/dates', methods=['POST'])
@login_required
def api_update_bid_dates(g_id):
    try:
        
        due_date = request.form.get('due_date')
        cur = mysql.connection.cursor()
        # Build dynamic set
        fields = []
        values = []
        if due_date:
            fields.append('due_date=%s'); values.append(due_date)
        if not fields:
            return jsonify({'ok': False, 'error': 'No fields provided'}), 400
        values.append(g_id)
        cur.execute(f"UPDATE go_bids SET {', '.join(fields)} WHERE g_id=%s", tuple(values))
        mysql.connection.commit()
        cur.close()
        return jsonify({'ok': True})
    except Exception as e:
        try:
            mysql.connection.rollback()
        except Exception:
            pass
        return jsonify({'ok': False, 'error': str(e)}), 500

# Comments CRUD (list + add)
@app.route('/api/bids/<int:g_id>/comments', methods=['GET', 'POST'])
@login_required
def api_bid_comments(g_id):
    if request.method == 'GET':
        cur = mysql.connection.cursor(DictCursor)
        cur.execute(
            """
            SELECT c.id, c.g_id, c.comment_text, c.created_at, u.email AS user_email
            FROM bid_comments c
            LEFT JOIN users u ON u.id = c.user_id
            WHERE c.g_id=%s
            ORDER BY c.created_at DESC
            """,
            (g_id,)
        )
        rows = cur.fetchall()
        cur.close()
        return jsonify({'comments': rows})
    else:
        text = (request.form.get('comment') or '').strip()
        if not text:
            return jsonify({'ok': False, 'error': 'Empty comment'}), 400
        try:
            cur = mysql.connection.cursor()
            uid = getattr(current_user, 'id', None)
            cur.execute("INSERT INTO bid_comments (g_id, user_id, comment_text) VALUES (%s,%s,%s)", (g_id, uid, text))
            mysql.connection.commit()
            cur.close()
            return jsonify({'ok': True})
        except Exception as e:
            try:
                mysql.connection.rollback()
            except Exception:
                pass
            return jsonify({'ok': False, 'error': str(e)}), 500
@app.route('/database-management/delete/<int:row_id>')
@login_required
def dbm_delete(row_id):
    if not current_user.is_admin:
        return "Access Denied", 403
    table = request.args.get('table')
    if not table:
        return redirect(url_for('database_management'))
    cur = mysql.connection.cursor(DictCursor)
    try:
        cur.execute(f"DESCRIBE `{table}`")
        cols = cur.fetchall()
        pk = next((c['Field'] for c in cols if c['Key'] == 'PRI'), cols[0]['Field'])
        
        # Handle foreign key constraints for specific tables
        if table == 'go_bids':
            # Delete related records first to avoid foreign key constraint errors
            cur.execute("DELETE FROM bid_checklists WHERE g_id=%s", (row_id,))
            cur.execute("DELETE FROM bid_stage_exclusions WHERE g_id=%s", (row_id,))
            cur.execute("DELETE FROM bid_custom_stages WHERE g_id=%s", (row_id,))
            
            # Delete from bid_assign (which may have win_lost_results, won_bids_result, work_progress_status cascading)
            cur.execute("SELECT a_id FROM bid_assign WHERE g_id=%s", (row_id,))
            assign_ids = [row['a_id'] for row in cur.fetchall()]
            
            for a_id in assign_ids:
                # Get w_id from win_lost_results
                cur.execute("SELECT w_id FROM win_lost_results WHERE a_id=%s", (a_id,))
                w_ids = [row['w_id'] for row in cur.fetchall()]
                
                for w_id in w_ids:
                    # Get won_id from won_bids_result
                    cur.execute("SELECT won_id FROM won_bids_result WHERE w_id=%s", (w_id,))
                    won_ids = [row['won_id'] for row in cur.fetchall()]
                    
                    # Delete work_progress_status
                    for won_id in won_ids:
                        cur.execute("DELETE FROM work_progress_status WHERE won_id=%s", (won_id,))
                    
                    # Delete won_bids_result
                    cur.execute("DELETE FROM won_bids_result WHERE w_id=%s", (w_id,))
                
                # Delete win_lost_results
                cur.execute("DELETE FROM win_lost_results WHERE a_id=%s", (a_id,))
            
            # Delete bid_assign
            cur.execute("DELETE FROM bid_assign WHERE g_id=%s", (row_id,))
        
        elif table == 'bid_incoming':
            # Find related go_bids and delete them (which will cascade)
            cur.execute("SELECT g_id FROM go_bids WHERE id=%s", (row_id,))
            g_ids = [row['g_id'] for row in cur.fetchall()]
            
            for g_id in g_ids:
                # Recursively delete go_bids using the same logic
                cur.execute("DELETE FROM bid_checklists WHERE g_id=%s", (g_id,))
                cur.execute("DELETE FROM bid_stage_exclusions WHERE g_id=%s", (g_id,))
                cur.execute("DELETE FROM bid_custom_stages WHERE g_id=%s", (g_id,))
                
                cur.execute("SELECT a_id FROM bid_assign WHERE g_id=%s", (g_id,))
                assign_ids = [row['a_id'] for row in cur.fetchall()]
                
                for a_id in assign_ids:
                    cur.execute("SELECT w_id FROM win_lost_results WHERE a_id=%s", (a_id,))
                    w_ids = [row['w_id'] for row in cur.fetchall()]
                    
                    for w_id in w_ids:
                        cur.execute("SELECT won_id FROM won_bids_result WHERE w_id=%s", (w_id,))
                        won_ids = [row['won_id'] for row in cur.fetchall()]
                        
                        for won_id in won_ids:
                            cur.execute("DELETE FROM work_progress_status WHERE won_id=%s", (won_id,))
                        
                        cur.execute("DELETE FROM won_bids_result WHERE w_id=%s", (w_id,))
                    
                    cur.execute("DELETE FROM win_lost_results WHERE a_id=%s", (a_id,))
                
                cur.execute("DELETE FROM bid_assign WHERE g_id=%s", (g_id,))
                cur.execute("DELETE FROM go_bids WHERE g_id=%s", (g_id,))
        
        elif table == 'bid_assign':
            # Get w_id from win_lost_results
            cur.execute("SELECT w_id FROM win_lost_results WHERE a_id=%s", (row_id,))
            w_ids = [row['w_id'] for row in cur.fetchall()]
            
            for w_id in w_ids:
                # Get won_id from won_bids_result
                cur.execute("SELECT won_id FROM won_bids_result WHERE w_id=%s", (w_id,))
                won_ids = [row['won_id'] for row in cur.fetchall()]
                
                # Delete work_progress_status
                for won_id in won_ids:
                    cur.execute("DELETE FROM work_progress_status WHERE won_id=%s", (won_id,))
                
                # Delete won_bids_result
                cur.execute("DELETE FROM won_bids_result WHERE w_id=%s", (w_id,))
            
            # Delete win_lost_results
            cur.execute("DELETE FROM win_lost_results WHERE a_id=%s", (row_id,))
        
        elif table == 'win_lost_results':
            # Get won_id from won_bids_result
            cur.execute("SELECT won_id FROM won_bids_result WHERE w_id=%s", (row_id,))
            won_ids = [row['won_id'] for row in cur.fetchall()]
            
            # Delete work_progress_status
            for won_id in won_ids:
                cur.execute("DELETE FROM work_progress_status WHERE won_id=%s", (won_id,))
            
            # Delete won_bids_result
            cur.execute("DELETE FROM won_bids_result WHERE w_id=%s", (row_id,))
        
        elif table == 'won_bids_result':
            # Delete work_progress_status
            cur.execute("DELETE FROM work_progress_status WHERE won_id=%s", (row_id,))
        
        # Now delete the main record
        cur.execute(f"DELETE FROM `{table}` WHERE `{pk}`=%s", (row_id,))
        mysql.connection.commit()
        flash(f'Record deleted successfully from {table}!', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Error deleting record: {str(e)}', 'error')
    finally:
        cur.close()
    return redirect(url_for('database_management', table=table))

@app.route('/database-management/drop')
@login_required
def dbm_drop():
    if not current_user.is_admin:
        return "Access Denied", 403
    table = request.args.get('table')
    if not table:
        return redirect(url_for('database_management'))
    cur = mysql.connection.cursor()
    try:
        cur.execute(f"DROP TABLE IF EXISTS `{table}`")
        mysql.connection.commit()
    finally:
        cur.close()
    return redirect(url_for('database_management'))

@app.route('/database-management/export')
@login_required
def dbm_export():
    if not current_user.is_admin:
        return "Access Denied", 403
    table = request.args.get('table')
    if not table:
        return redirect(url_for('database_management'))
    try:
        import pandas as pd
        import io
        cur = mysql.connection.cursor(DictCursor)
        cur.execute(f"SELECT * FROM `{table}`")
        rows = cur.fetchall()
        cur.close()
        df = pd.DataFrame(rows)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name=table[:31])
        output.seek(0)
        filename = f"{table}.xlsx"
        return send_file(output, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    except Exception as e:
        flash(f'Export failed: {str(e)}', 'error')
        return redirect(url_for('database_management', table=table))

@app.route('/database-management/import', methods=['POST'])
@login_required
def dbm_import():
    if not current_user.is_admin:
        return "Access Denied", 403
    table = request.form.get('table')
    file = request.files.get('excel_file')
    if not table or not file or file.filename == '':
        flash('Please choose a table and an Excel file.', 'error')
        return redirect(url_for('database_management', table=table or ''))
    try:
        import pandas as pd
        import io
        df = pd.read_excel(io.BytesIO(file.read()))
        cur = mysql.connection.cursor(DictCursor)
        # Fetch columns and ignore primary key
        cur.execute(f"DESCRIBE `{table}`")
        cols = cur.fetchall()
        non_pk_cols = [c['Field'] for c in cols if c['Key'] != 'PRI']
        # Filter dataframe to only known columns
        df = df[[c for c in df.columns if c in non_pk_cols]]
        if df.empty:
            flash('No matching columns found in Excel file for this table.', 'error')
            cur.close()
            return redirect(url_for('database_management', table=table))
        placeholders = ','.join(['%s'] * len(df.columns))
        fields_sql = ','.join([f"`{c}`" for c in df.columns])
        sql = f"INSERT INTO `{table}` ({fields_sql}) VALUES ({placeholders})"
        for _, row in df.iterrows():
            cur.execute(sql, [None if pd.isna(v) else v for v in row.tolist()])
        mysql.connection.commit()
        cur.close()
        flash(f'Successfully imported {len(df)} rows into `{table}`.', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Import failed: {str(e)}', 'error')
    return redirect(url_for('database_management', table=table))

@app.route('/database-management/update-decision', methods=['POST'])
@login_required
def dbm_update_decision():
    try:
        payload = request.get_json(force=True) or {}
        table = (payload.get('table') or '').strip().lower()
        row_id = int(payload.get('row_id') or 0)
        decision = (payload.get('decision') or '').strip().upper()

        if table not in ('bid_incoming', 'go_bids'):
            return jsonify({'success': False, 'error': 'Invalid table'}), 400
        if row_id <= 0:
            return jsonify({'success': False, 'error': 'Invalid row id'}), 400
        # normalize decision
        mapping = {
            'GO': 'GO', 'NO GO': 'NO-GO', 'NO-GO': 'NO-GO', 'NOGO': 'NO-GO',
            'WON': 'WON', 'LOST': 'LOST', 'PENDING': 'PENDING'
        }
        decision = mapping.get(decision, 'PENDING')

        cur = mysql.connection.cursor(DictCursor)
        if table == 'bid_incoming':
            cur.execute("UPDATE bid_incoming SET decision=%s WHERE id=%s", (decision, row_id))
        else:
            cur.execute("UPDATE go_bids SET decision=%s WHERE g_id=%s", (decision, row_id))
        mysql.connection.commit()
        cur.close()

        # Keep go_bids synchronized with bid_incoming GO decisions
        try:
            sync_go_bids()
        except Exception as sync_err:
            # Do not fail the request if sync encounters an issue; just log it.
            print(f"Warning: sync_go_bids failed after decision update: {sync_err}")

        log_write('dbm_update_decision', f"table={table} id={row_id} decision={decision}")
        return jsonify({'success': True, 'decision': decision})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

def generate_complete_proposal(pdf_path, company, template, output_path, bid_name=None, sections_outline=None):
    """
    Generate a complete structured proposal document from RFP PDF using AI.
    Creates a comprehensive proposal with all required sections.
    """
    try:
        import os
        from datetime import datetime
        from rfp_analyzer_routes import extract_page_texts, ollama_json, trim_text_to_token_limit
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.enum.table import WD_TABLE_ALIGNMENT
        import re
        
        # Read PDF file
        with open(pdf_path, 'rb') as pdf_file:
            pages = extract_page_texts(pdf_file)
        
        if not pages or len(pages) == 0:
            raise Exception("Could not extract text from PDF")
        
        full_text = "\n\n".join(pages)
        
        # Extract key information from RFP
        rfp_text_limited = trim_text_to_token_limit(full_text, 8000)  # Limit for prompt
        
        # Extract project title
        project_title = bid_name or "Project"
        title_match = re.search(r'(?i)(?:project|title|proposal|rfp)[:\s]+([A-Z][^\n]{10,100})', full_text[:2000])
        if title_match:
            project_title = title_match.group(1).strip()
        
        # Extract deadline
        deadline = None
        deadline_patterns = [
            r'(?i)(?:deadline|due\s+date|submission\s+date)[:\s]+([A-Za-z]+\s+\d{1,2},?\s+\d{4})',
            r'(?i)(?:deadline|due\s+date)[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
            r'(?i)August\s+31,?\s+2026',
        ]
        for pattern in deadline_patterns:
            match = re.search(pattern, full_text[:5000])
            if match:
                deadline = match.group(1) if match.groups() else match.group(0)
                break
        
        # Build dynamic proposal prompt based on provided sections
        company_display = company or "Our Company"
        project_label = project_title or bid_name or "Project"

        if not sections_outline:
            sections_outline = [
                {'title': 'Executive Summary', 'guidance': 'Summarize the opportunity, objectives, and how the proposed solution delivers value.'},
                {'title': 'Company Profile', 'guidance': 'Describe the company background, capabilities, and relevant experience.'},
                {'title': 'Technical Approach & Scope', 'guidance': 'Detail the technical solution, methodology, schedule, and risk mitigation strategies.'},
                {'title': 'Project Management & Schedule', 'guidance': 'Outline management structure, milestones, resource plan, and communication strategy.'},
                {'title': 'Price Proposal Form (Attachment A)', 'guidance': 'Explain pricing assumptions, cost elements, and any exclusions.'},
                {'title': 'Compliance Matrix & Deliverables', 'guidance': 'Summarize compliance approach, required deliverables, and submission checklist.'},
                {'title': 'Project Team', 'guidance': '.'},
                {'title': 'Corporate Qualifications', 'guidance': ''}
            ]

        section_prompt_lines = []
        for sec in sections_outline:
            guidance = sec.get('guidance') or 'Provide a comprehensive, proposal-ready narrative aligned with this heading.'
            sub_titles = ', '.join(sub.get('title') for sub in sec.get('sub_categories') or [] if sub.get('title'))
            if sub_titles:
                guidance += f" Include coverage of: {sub_titles}."
            section_prompt_lines.append(f'- {sec.get("title", "Section")}: {guidance}')
        sections_prompt = "\n".join(section_prompt_lines)

        proposal_prompt = f"""You are a professional proposal generation assistant specialized in government and infrastructure RFPs. Analyze the uploaded RFP document in full detail and automatically create a draft proposal narrative that follows the section headers provided.

RFP Document Content:
{rfp_text_limited}

Company Name: {company_display}
Project Title: {project_label}
Deadline: {deadline or 'To be determined'}

Sections to include (in this exact order):
{sections_prompt}

Return JSON with the following structure:
{{
  "sections": [
    {{
      "title": "Exact Section Title",
      "content": [
        "Paragraph 1",
        "Paragraph 2",
        "Paragraph 3"
      ],
      "sub_sections": [
        {{
          "title": "Optional Subheading",
          "content": [
            "Paragraph 1",
            "Paragraph 2"
          ]
        }}
      ]
    }}
  ]
}}

Ensure each section contains multiple detailed paragraphs tailored to the RFP requirements, referencing specific compliance, schedule, and technical expectations whenever possible."""

        # Generate proposal content using AI
        print("Generating proposal content with AI...")
        proposal_result = ollama_json(proposal_prompt, 4000)

        if proposal_result.get('error'):
            raise Exception(f"AI generation error: {proposal_result['error']}")

        def normalize_title(value):
            return re.sub(r'\s+', ' ', str(value).strip().lower())

        sections_output = []
        if isinstance(proposal_result, dict):
            if isinstance(proposal_result.get('sections'), list):
                sections_output = proposal_result['sections']
            else:
                # Interpret as dictionary of section -> content
                sections_output = [
                    {'title': key, 'content': value}
                    for key, value in proposal_result.items()
                    if key not in ('raw', 'error')
                ]

        sections_map = {}
        for entry in sections_output:
            title = entry.get('title')
            if not title:
                continue
            sections_map[normalize_title(title)] = entry

        normalized_sections = []
        for sec in sections_outline:
            title = sec.get('title', 'Section').strip()
            entry = sections_map.get(normalize_title(title))

            content = []
            sub_sections = []

            if entry:
                entry_content = entry.get('content')
                if isinstance(entry_content, str):
                    content = [entry_content]
                elif isinstance(entry_content, list):
                    content = [item for item in entry_content if item]
                elif isinstance(entry_content, dict):
                    # Flatten dict content
                    content = [v for v in entry_content.values() if v]
                sub_sections = entry.get('sub_sections') or []

            if not content:
                fallback_content = sec.get('content') or []
                if isinstance(fallback_content, str):
                    fallback_content = [fallback_content]
                if not fallback_content and sec.get('guidance'):
                    fallback_content = [sec['guidance']]
                if not fallback_content:
                    fallback_content = ['Content will be added once capture information is finalized.']
                content = fallback_content

            normalized_subsections = []
            for sub in sub_sections:
                sub_title = sub.get('title')
                if not sub_title:
                    continue
                sub_content = sub.get('content')
                if isinstance(sub_content, str):
                    sub_content = [sub_content]
                elif isinstance(sub_content, list):
                    sub_content = [item for item in sub_content if item]
                else:
                    sub_content = []
                if not sub_content:
                    sub_content = [f'Additional guidance for {sub_title} will be incorporated.']
                normalized_subsections.append({
                    'title': sub_title,
                    'content': sub_content
                })

            normalized_sections.append({
                'title': title,
                'content': content,
                'sub_sections': normalized_subsections
            })

        # Create Word document (prefer H & F template to inherit header/footer)
        try:
            base_dir = app.root_path if hasattr(app, 'root_path') else os.getcwd()
            hf_template_path = os.path.join(base_dir, 'H & F.docx')
            if os.path.exists(hf_template_path):
                doc = Document(hf_template_path)
            else:
                doc = Document()
        except Exception:
            doc = Document()

        # Title Page
        doc.add_heading(project_label or 'Proposal Draft', 0)
        doc.add_paragraph(f'Prepared for: {project_label}')
        doc.add_paragraph(f'Prepared by: {company_display}')
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        doc.add_paragraph()
        doc.add_page_break()

        # Table of Contents
        doc.add_heading('Table of Contents', 1)
        for idx, section in enumerate(normalized_sections, 1):
            doc.add_paragraph(f'{idx}. {section["title"]}')
        doc.add_page_break()

        # Sections
        for idx, section in enumerate(normalized_sections):
            doc.add_heading(section['title'], level=1)
            for paragraph in section.get('content', []):
                for part in str(paragraph).split('\n\n'):
                    if part.strip():
                        doc.add_paragraph(part.strip())

            for sub_section in section.get('sub_sections', []):
                sub_title = sub_section.get('title')
                if sub_title:
                    doc.add_heading(sub_title, level=2)
                for paragraph in sub_section.get('content', []):
                    for part in str(paragraph).split('\n\n'):
                        if part.strip():
                            doc.add_paragraph(part.strip())

            if idx != len(normalized_sections) - 1:
                doc.add_page_break()

        # Save document
        try:
            # Link headers/footers to the first section so they are consistent
            for _s_idx, _section in enumerate(doc.sections):
                if _s_idx > 0:
                    try:
                        _section.header.is_linked_to_previous = True
                        _section.footer.is_linked_to_previous = True
                    except Exception:
                        pass
                # Disable special first/odd-even to make header/footer consistent on all pages
                try:
                    _section.different_first_page_header_footer = False
                except Exception:
                    pass
                try:
                    _section.odd_and_even_pages_header_footer = False
                except Exception:
                    pass
            # Ensure continuous page numbering across sections
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            for _s_idx, _section in enumerate(doc.sections):
                _sectPr = _section._sectPr
                _pgNumType = _sectPr.find(qn('w:pgNumType'))
                if _pgNumType is None:
                    _pgNumType = OxmlElement('w:pgNumType')
                    _sectPr.append(_pgNumType)
                if _s_idx == 0:
                    _pgNumType.set(qn('w:start'), "1")
                else:
                    try:
                        _pgNumType.attrib.pop(qn('w:start'), None)
                    except Exception:
                        pass
        except Exception:
            pass
        doc.save(output_path)
        print(f"Proposal saved to: {output_path}")
        return True
        
    except Exception as e:
        print(f"Error generating complete proposal: {str(e)}")
        import traceback
        traceback.print_exc()
        # Create a basic document as fallback
        try:
            from docx import Document
            from datetime import datetime
            doc = Document()
            doc.add_heading('Proposal Generation Error', 0)
            doc.add_paragraph(f'RFP Document: {os.path.basename(pdf_path)}')
            doc.add_paragraph(f'Company: {company}')
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph('\nNote: An error occurred during proposal generation.')
            doc.add_paragraph(f'Error: {str(e)}')
            doc.save(output_path)
            return True
        except:
            return False

def generate_executive_summary_with_toc(pdf_path, company, template, output_path):
    """
    Generate an executive summary document with table of contents from RFP PDF.
    Uses RFP analyzer functions to extract and summarize content.
    """
    try:
        import os
        from datetime import datetime
        from rfp_analyzer_routes import extract_page_texts, summarize_batch_with_llama, build_master_summary, ollama_json, trim_text_to_token_limit
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re
        
        # Read PDF file
        with open(pdf_path, 'rb') as pdf_file:
            pages = extract_page_texts(pdf_file)
        
        if not pages or len(pages) == 0:
            raise Exception("Could not extract text from PDF")
        
        full_text = "\n\n".join(pages)
        
        # Process in batches for summarization
        DEFAULT_BATCH_PAGES = 10
        batch_summaries = []
        total_pages = len(pages)
        
        for i in range(0, total_pages, DEFAULT_BATCH_PAGES):
            batch_pages = pages[i:min(i + DEFAULT_BATCH_PAGES, total_pages)]
            batch_text = "\n\n".join(batch_pages)
            batch_idx = (i // DEFAULT_BATCH_PAGES) + 1
            summary = summarize_batch_with_llama(batch_text, batch_idx)
            if summary and not summary.get('error'):
                batch_summaries.append(summary)
        
        # Build master summary
        if batch_summaries:
            master_summary = build_master_summary(batch_summaries)
        else:
            # Fallback: create basic summary from text
            master_summary = {
                "project_type": "General Project",
                "scope": full_text[:500] + "..." if len(full_text) > 500 else full_text,
                "summary": full_text[:1000] + "..." if len(full_text) > 1000 else full_text,
                "key_requirements": [],
                "technical_requirements": [],
                "bid_requirements": []
            }
        
        # Generate executive summary using LLM
        summary_text = trim_text_to_token_limit(master_summary.get('summary', full_text[:2000]), 4000)
        exec_prompt = f"""Create a comprehensive executive summary for this RFP document.
        
The summary should include:
1. Project Overview
2. Key Objectives
3. Scope of Work
4. Key Requirements
5. Timeline and Deadlines
6. Budget Information (if available)
7. Technical Specifications (highlights)
8. Submission Requirements

RFP Content Summary:
{summary_text}

Return a well-structured executive summary in clear, professional language suitable for senior management."""
        
        exec_summary_result = ollama_json(exec_prompt, 2000)
        executive_summary = exec_summary_result.get('summary') or exec_summary_result.get('raw') or master_summary.get('summary', 'Executive summary could not be generated.')
        
        # Extract table of contents from PDF text
        toc_items = []
        toc_patterns = [
            r'(?i)^\s*(?:table\s+of\s+contents?|contents?|toc)\s*$',
            r'(?i)^\s*\d+\.?\s+[A-Z][^\n]{10,100}\s+\.\.\.?\s*\d+',
            r'(?i)^\s*[A-Z][A-Z\s]{5,50}\s+\.\.\.?\s*\d+',
        ]
        
        # Try to find existing TOC in PDF
        lines = full_text.split('\n')
        in_toc = False
        for line in lines[:100]:  # Check first 100 lines
            line_clean = line.strip()
            if re.search(toc_patterns[0], line_clean):
                in_toc = True
                continue
            if in_toc and line_clean:
                # Look for TOC entries
                match = re.search(r'^(.+?)\s+\.\.\.?\s*(\d+)$', line_clean)
                if match:
                    toc_items.append({
                        'title': match.group(1).strip(),
                        'page': match.group(2).strip()
                    })
                elif len(toc_items) > 0 and len(line_clean) > 5:
                    # Continue previous entry
                    toc_items[-1]['title'] += ' ' + line_clean
                if len(toc_items) > 20:  # Limit TOC items
                    break
        
        # If no TOC found, generate one from sections
        if not toc_items:
            section_patterns = [
                r'(?i)^\s*\d+\.?\s+([A-Z][^\n]{10,80})',
                r'(?i)^\s*([A-Z][A-Z\s]{5,50})\s*$',
                r'(?i)^\s*(?:section|chapter|part)\s+\d+[:\s]+([^\n]{10,80})',
            ]
            seen_sections = set()
            for line in lines[:200]:
                for pattern in section_patterns:
                    match = re.search(pattern, line)
                    if match:
                        section_title = match.group(1).strip()
                        if len(section_title) > 10 and section_title not in seen_sections:
                            toc_items.append({'title': section_title, 'page': ''})
                            seen_sections.add(section_title)
                            if len(toc_items) >= 15:
                                break
                if len(toc_items) >= 15:
                    break
        
        # Create Word document (prefer H & F template to inherit header/footer)
        try:
            base_dir = app.root_path if hasattr(app, 'root_path') else os.getcwd()
            hf_template_path = os.path.join(base_dir, 'H & F.docx')
            if os.path.exists(hf_template_path):
                doc = Document(hf_template_path)
            else:
                doc = Document()
        except Exception:
            doc = Document()
        
        # Title page
        title_para = doc.add_heading('Executive Summary', 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'RFP Document: {os.path.basename(pdf_path)}')
        doc.add_paragraph(f'Company: {company}')
        doc.add_paragraph(f'Template: {template}')
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_page_break()
        
        # Table of Contents
        doc.add_heading('Table of Contents', 1)
        toc_para = doc.add_paragraph()
        
        if toc_items:
            for item in toc_items:
                p = doc.add_paragraph()
                run1 = p.add_run(item['title'])
                run1.font.size = Pt(11)
                if item['page']:
                    run2 = p.add_run(f'\t{item["page"]}')
                    run2.font.size = Pt(11)
                    # Add tab stop for alignment
                    p.paragraph_format.tab_stops.add_tab_stop(Inches(5.5))
        else:
            doc.add_paragraph('1. Executive Summary')
            doc.add_paragraph('2. Project Overview')
            doc.add_paragraph('3. Scope of Work')
            doc.add_paragraph('4. Key Requirements')
            doc.add_paragraph('5. Technical Specifications')
            doc.add_paragraph('6. Timeline and Deadlines')
            doc.add_paragraph('7. Budget Information')
            doc.add_paragraph('8. Submission Requirements')
        
        doc.add_page_break()
        
        # Executive Summary Section
        doc.add_heading('Executive Summary', 1)
        exec_para = doc.add_paragraph(executive_summary)
        exec_para_format = exec_para.paragraph_format
        exec_para_format.space_after = Pt(12)
        
        # Project Overview
        if master_summary.get('project_type') or master_summary.get('scope'):
            doc.add_heading('Project Overview', 1)
            if master_summary.get('project_type'):
                doc.add_paragraph(f"Project Type: {master_summary['project_type']}", style='List Bullet')
            if master_summary.get('scope'):
                scope_text = master_summary['scope'][:2000] if len(master_summary.get('scope', '')) > 2000 else master_summary.get('scope', '')
                doc.add_paragraph(scope_text)
        
        # Key Requirements
        if master_summary.get('key_requirements') and len(master_summary['key_requirements']) > 0:
            doc.add_heading('Key Requirements', 1)
            for req in master_summary['key_requirements'][:10]:  # Limit to 10
                if isinstance(req, str) and len(req.strip()) > 0:
                    doc.add_paragraph(req, style='List Bullet')
        
        # Technical Requirements
        if master_summary.get('technical_requirements') and len(master_summary['technical_requirements']) > 0:
            doc.add_heading('Technical Requirements', 1)
            for req in master_summary['technical_requirements'][:10]:  # Limit to 10
                if isinstance(req, str) and len(req.strip()) > 0:
                    doc.add_paragraph(req, style='List Bullet')
        
        # Bid Requirements
        if master_summary.get('bid_requirements') and len(master_summary['bid_requirements']) > 0:
            doc.add_heading('Bid Submission Requirements', 1)
            for req in master_summary['bid_requirements'][:10]:  # Limit to 10
                if isinstance(req, str) and len(req.strip()) > 0:
                    doc.add_paragraph(req, style='List Bullet')
        
        # Additional Information
        doc.add_heading('Document Information', 1)
        doc.add_paragraph(f'Total Pages: {total_pages}', style='List Bullet')
        doc.add_paragraph(f'Total Characters: {len(full_text):,}', style='List Bullet')
        if master_summary.get('due_date'):
            doc.add_paragraph(f'Due Date: {master_summary["due_date"]}', style='List Bullet')
        if master_summary.get('total_cost'):
            doc.add_paragraph(f'Estimated Cost: {master_summary["total_cost"]}', style='List Bullet')
        
        # Save document
        try:
            # Link headers/footers to the first section so they are consistent
            for _s_idx, _section in enumerate(doc.sections):
                if _s_idx > 0:
                    try:
                        _section.header.is_linked_to_previous = True
                        _section.footer.is_linked_to_previous = True
                    except Exception:
                        pass
                # Disable special first/odd-even to make header/footer consistent on all pages
                try:
                    _section.different_first_page_header_footer = False
                except Exception:
                    pass
                try:
                    _section.odd_and_even_pages_header_footer = False
                except Exception:
                    pass
            # Ensure continuous page numbering across sections
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            for _s_idx, _section in enumerate(doc.sections):
                _sectPr = _section._sectPr
                _pgNumType = _sectPr.find(qn('w:pgNumType'))
                if _pgNumType is None:
                    _pgNumType = OxmlElement('w:pgNumType')
                    _sectPr.append(_pgNumType)
                if _s_idx == 0:
                    _pgNumType.set(qn('w:start'), "1")
                else:
                    try:
                        _pgNumType.attrib.pop(qn('w:start'), None)
                    except Exception:
                        pass
        except Exception:
            pass
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error generating executive summary: {str(e)}")
        import traceback
        traceback.print_exc()
        # Create a basic document as fallback
        try:
            from docx import Document
            from datetime import datetime
            doc = Document()
            doc.add_heading('Executive Summary', 0)
            doc.add_paragraph(f'RFP Document: {os.path.basename(pdf_path)}')
            doc.add_paragraph(f'Company: {company}')
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph('\nNote: An error occurred during executive summary generation.')
            doc.add_paragraph(f'Error: {str(e)}')
            doc.save(output_path)
            return True
        except:
            return False

@app.route('/proposals-making', methods=['GET', 'POST'])
@login_required
def proposals_making():
    # Allow all logged-in users to access proposals-making (removed admin-only restriction)
    import os
    import re
    from datetime import datetime, date

    contact_email = getattr(current_user, 'email', '')

    def safe_str(value):
        if value is None:
            return ''
        return str(value).strip()

    def format_date_for_display(value):
        if not value:
            return ''
        if isinstance(value, (datetime, date)):
            return value.strftime('%B %d, %Y')
        try:
            parsed = datetime.strptime(str(value), '%Y-%m-%d')
            return parsed.strftime('%B %d, %Y')
        except Exception:
            return str(value)

    def split_text_blocks(text, max_blocks=3):
        if not text:
            return []
        blocks = [
            safe_str(piece).strip('•- ') for piece in re.split(r'(?:\r?\n){2,}', str(text))
            if safe_str(piece)
        ]
        if not blocks and safe_str(text):
            blocks = [safe_str(text)]
        return blocks[:max_blocks]

    def compact_list(items):
        result = []
        for item in items:
            if item is None:
                continue
            text = safe_str(item)
            if text:
                result.append(text)
        return result

    def get_bid_metadata(g_id):
        if not g_id:
            return None
        meta_cursor = mysql.connection.cursor(DictCursor)
        try:
            meta_cursor.execute("SELECT * FROM go_bids WHERE g_id=%s", (g_id,))
            return meta_cursor.fetchone()
        except Exception as meta_err:
            print(f"Error loading bid context for proposal view: {meta_err}")
            return None
        finally:
            meta_cursor.close()

    def build_sections_outline(bid_info, fallback_title, company_name, primary_contact):
        bid = bid_info or {}
        project_name = safe_str(bid.get('b_name')) or safe_str(fallback_title) or 'Current Opportunity'
        due_text = format_date_for_display(bid.get('due_date'))
        client_name = safe_str(
            bid.get('agency') or bid.get('customer') or bid.get('client') or bid.get('owner') or bid.get('comp_name')
        )
        scope_text = safe_str(bid.get('scope'))
        summary_text = safe_str(bid.get('summary'))
        type_text = safe_str(bid.get('type'))
        location_text = safe_str(bid.get('location') or bid.get('city') or bid.get('state'))
        stage_text = safe_str(bid.get('state'))
        company_display = safe_str(company_name) or safe_str(bid.get('company')) or 'Offer or'
        revenue_value = bid.get('revenue') if bid.get('revenue') not in (None, '') else bid.get('value')
        scoring_value = bid.get('scoring')

        summary_blocks = split_text_blocks(summary_text, 4)
        scope_blocks = split_text_blocks(scope_text, 5)

        revenue_text = ''
        if revenue_value not in (None, ''):
            try:
                revenue_text = f"${float(revenue_value):,.0f}"
            except Exception:
                revenue_text = safe_str(revenue_value)

        scoring_text = ''
        if scoring_value not in (None, ''):
            try:
                scoring_text = f"{float(scoring_value):.0f}"
            except Exception:
                scoring_text = safe_str(scoring_value)

        sections = []

        cover_content = compact_list([
            f"{company_display} is pleased to submit this proposal for {project_name}.",
            f"Issuing agency / customer: {client_name}." if client_name else '',
            f"Solicitation reference: {project_name}.",
            f"Proposal due date: {due_text}." if due_text else '',
            f"Primary point of contact: {primary_contact}." if primary_contact else '',
            f"Contract type: {type_text}." if type_text else '',
            f"Performance location: {location_text}." if location_text else '',
            f"Opportunity priority score: {scoring_text}." if scoring_text else ''
        ])
        sections.append({
            'id': 'section-cover',
            'title': 'Letter of Transmittal',
            'status': 'Saved' if cover_content else 'Draft',
            'guidance': 'Confirm the solicitation title, client information, offer or contact details, and submission identifiers before releasing the cover page.',
            'content': cover_content or ['Add cover page details once the capture team confirms the solicitation data.'],
            'sub_categories': [
                {'id': 'cover-company-info', 'title': 'Company Information'},
                {'id': 'cover-contact-details', 'title': 'Contact Details'},
                {'id': 'cover-solicitation-ref', 'title': 'Solicitation Reference'},
                {'id': 'cover-submission-date', 'title': 'Submission Date & Deadline'}
            ]
        })

        exec_content = summary_blocks or (
            [f"{company_display} will deliver a compliant, value-focused response for {project_name}."]
            if project_name else []
        )
        if scope_text and not summary_blocks:
            exec_content.append(scope_text)
        sections.append({
            'id': 'section-executive-summary',
            'title': 'Executive Summary',
            'status': 'In Progress' if summary_blocks else 'Draft',
            'guidance': "Summarize win themes, differentiators, and alignment to the customer's objectives captured during bid qualification.",
            'content': exec_content or ['Executive summary content will be generated as soon as capture notes are available.'],
            'sub_categories': [
                {'id': 'exec-overview', 'title': 'Project Overview'},
                {'id': 'exec-win-themes', 'title': 'Win Themes & Differentiators'},
                {'id': 'exec-value-proposition', 'title': 'Value Proposition'},
                {'id': 'exec-key-highlights', 'title': 'Key Highlights'}
            ]
        })

        compliance_content = compact_list([
            "Track each requirement in the compliance matrix to ensure every section of the RFP is addressed.",
            scope_blocks[0] if scope_blocks else '',
            summary_blocks[0] if summary_blocks else ''
        ])
        sections.append({
            'id': 'section-compliance',
            'title': 'Requirements Compliance',
            'status': 'In Progress',
            'guidance': 'Reference the compliance matrix and list every deliverable, attachment, and form demanded by the solicitation.',
            'content': compliance_content or [
                'Populate this section with compliance findings once the RFP analysis is complete.'
            ],
            'sub_categories': [
                {'id': 'comp-matrix', 'title': 'Compliance Matrix'},
                {'id': 'comp-deliverables', 'title': 'Deliverables List'},
                {'id': 'comp-attachments', 'title': 'Required Attachments'},
                {'id': 'comp-forms', 'title': 'Required Forms'}
            ]
        })

        sections.append({
            'id': 'section-scope-objectives',
            'title': 'Understanding of Scope & Objectives',
            'status': 'In Progress',
            'guidance': 'Demonstrate a clear understanding of the project scope, objectives, and expected outcomes.',
            'content': scope_blocks or [
                'Detail your understanding of the project scope and objectives based on the solicitation requirements.'
            ],
            'sub_categories': [
                {'id': 'scope-project-scope', 'title': 'Project Scope'},
                {'id': 'scope-objectives', 'title': 'Project Objectives'},
                {'id': 'scope-expected-outcomes', 'title': 'Expected Outcomes'},
                {'id': 'scope-assumptions', 'title': 'Key Assumptions'}
            ]
        })

        sections.append({
            'id': 'section-deviations',
            'title': 'Deviations, Assumptions, and Dependencies (if any)',
            'status': 'Draft',
            'guidance': 'Document any deviations from the solicitation requirements, key assumptions, and dependencies that may impact project delivery.',
            'content': [
                'List any deviations from the solicitation requirements, if applicable.',
                'Document key assumptions made during proposal development.',
                'Identify dependencies that may affect project execution.'
            ],
            'sub_categories': [
                {'id': 'dev-deviations', 'title': 'Deviations'},
                {'id': 'dev-assumptions', 'title': 'Assumptions'},
                {'id': 'dev-dependencies', 'title': 'Dependencies'},
                {'id': 'dev-risks', 'title': 'Risk Considerations'}
            ]
        })

        tech_content = scope_blocks or [
            'Detail the technical solution, materials, and implementation methodology once the RFP requirements are parsed.'
        ]
        sections.append({
            'id': 'section-technical',
            'title': 'Proposed Technical Solution',
            'status': 'In Progress' if scope_blocks else 'Draft',
            'guidance': 'Translate the solicitation Statement of Work into your delivery plan, highlighting compliance, innovations, and risk mitigations.',
            'content': tech_content,
            'sub_categories': [
                {'id': 'tech-approach', 'title': 'Technical Approach'},
                {'id': 'tech-methodology', 'title': 'Methodology'},
                {'id': 'tech-innovations', 'title': 'Innovations & Best Practices'},
                {'id': 'tech-compliance', 'title': 'Compliance & Standards'},
                {'id': 'tech-risk-mitigation', 'title': 'Risk Mitigation'}
            ]
        })

        sections.append({
            'id': 'section-implementation',
            'title': 'Implementation & Work Plan',
            'status': 'In Progress',
            'guidance': 'Provide a detailed implementation plan and work breakdown structure for project execution.',
            'content': [
                'Outline the step-by-step implementation approach.',
                'Detail the work breakdown structure and key activities.',
                'Describe the workflow and processes for project delivery.'
            ],
            'sub_categories': [
                {'id': 'impl-approach', 'title': 'Implementation Approach'},
                {'id': 'impl-work-breakdown', 'title': 'Work Breakdown Structure'},
                {'id': 'impl-phases', 'title': 'Implementation Phases'},
                {'id': 'impl-processes', 'title': 'Key Processes'}
            ]
        })

        sections.append({
            'id': 'section-deliverables',
            'title': 'Solution Deliverables',
            'status': 'In Progress',
            'guidance': 'List all deliverables that will be provided as part of the solution, including documentation, reports, and physical items.',
            'content': [
                'Comprehensive list of all solution deliverables.',
                'Documentation and reports to be provided.',
                'Physical deliverables, if applicable.'
            ],
            'sub_categories': [
                {'id': 'del-documentation', 'title': 'Documentation Deliverables'},
                {'id': 'del-reports', 'title': 'Reports & Analysis'},
                {'id': 'del-physical', 'title': 'Physical Deliverables'},
                {'id': 'del-schedule', 'title': 'Delivery Schedule'}
            ]
        })

        management_content = compact_list([
            f"Target submission date: {due_text}." if due_text else 'Submission date will be updated when the procurement schedule is confirmed.',
            f"Contract type guidance: {type_text}." if type_text else '',
            "Outline major milestones, design reviews, and government touchpoints that support a compliant response."
        ])
        sections.append({
            'id': 'section-management',
            'title': 'Project Management Plan',
            'status': 'In Progress',
            'guidance': 'Map the internal review cycle, color team dates, and delivery milestones to ensure on-time submission.',
            'content': management_content or [
                'Add the project management narrative and schedule once the capture calendar is finalized.'
            ],
            'sub_categories': [
                {'id': 'mgmt-organization', 'title': 'Project Organization'},
                {'id': 'mgmt-resources', 'title': 'Resource Management'},
                {'id': 'mgmt-quality', 'title': 'Quality Assurance'},
                {'id': 'mgmt-communication', 'title': 'Communication Plan'}
            ]
        })

        sections.append({
            'id': 'section-schedule',
            'title': 'Project Schedule & Milestones',
            'status': 'In Progress',
            'guidance': 'Provide a detailed project schedule with key milestones, deliverables, and timeline.',
            'content': [
                'Detailed project timeline with start and end dates.',
                'Key milestones and their target dates.',
                'Critical path and dependencies.'
            ],
            'sub_categories': [
                {'id': 'schedule-timeline', 'title': 'Project Timeline'},
                {'id': 'schedule-milestones', 'title': 'Key Milestones'},
                {'id': 'schedule-critical-path', 'title': 'Critical Path'},
                {'id': 'schedule-dependencies', 'title': 'Dependencies'}
            ]
        })

        sections.append({
            'id': 'section-team',
            'title': 'Project Team',
            'status': 'In Progress',
            'guidance': 'Introduce the project team members, their roles, qualifications, and responsibilities.',
            'content': [
                'List of key team members and their roles.',
                'Qualifications and experience of team members.',
                'Organizational structure and reporting lines.'
            ],
            'sub_categories': [
                {'id': 'team-members', 'title': 'Team Members'},
                {'id': 'team-roles', 'title': 'Roles & Responsibilities'},
                {'id': 'team-qualifications', 'title': 'Qualifications'},
                {'id': 'team-organization', 'title': 'Organization Structure'}
            ]
        })

        sections.append({
            'id': 'section-corporate',
            'title': 'Corporate Qualifications',
            'status': 'Saved',
            'guidance': 'Highlight corporate experience, past performance, certifications, and qualifications relevant to this opportunity.',
            'content': [
                'Company background and history.',
                'Relevant past performance and case studies.',
                'Certifications, accreditations, and qualifications.'
            ],
            'sub_categories': [
                {'id': 'corp-background', 'title': 'Company Background'},
                {'id': 'corp-past-performance', 'title': 'Past Performance'},
                {'id': 'corp-certifications', 'title': 'Certifications & Accreditations'},
                {'id': 'corp-capabilities', 'title': 'Key Capabilities'}
            ]
        })

        price_content = compact_list([
            f"Estimated contract value: {revenue_text}." if revenue_text else '',
            "Attachment A pricing will mirror the government-provided workbook with CLIN-level details.",
            "Document all pricing assumptions, escalation factors, and exclusions for contracting officer review."
        ])
        sections.append({
            'id': 'section-pricing',
            'title': 'Pricing Proposals',
            'status': 'Saved' if revenue_text else 'Draft',
            'guidance': 'Ensure Attachment A reflects current quantities, rates, and assumptions. Cross-check CLIN totals before submission.',
            'content': price_content or [
                'Populate Attachment A once final quantities and labor categories are confirmed.'
            ],
            'sub_categories': [
                {'id': 'price-clin-breakdown', 'title': 'CLIN Breakdown'},
                {'id': 'price-labor-rates', 'title': 'Labor Rates & Categories'},
                {'id': 'price-materials', 'title': 'Materials & Equipment'},
                {'id': 'price-assumptions', 'title': 'Pricing Assumptions'},
                {'id': 'price-total-summary', 'title': 'Total Cost Summary'}
            ]
        })

        sections.append({
            'id': 'section-contractual',
            'title': 'Contractual Terms',
            'status': 'Saved',
            'guidance': 'Outline key contractual terms, conditions, and agreements proposed for this engagement.',
            'content': [
                'Key contractual terms and conditions.',
                'Payment terms and conditions.',
                'Terms of service and deliverables.'
            ],
            'sub_categories': [
                {'id': 'contract-terms', 'title': 'Terms & Conditions'},
                {'id': 'contract-payment', 'title': 'Payment Terms'},
                {'id': 'contract-service', 'title': 'Service Terms'},
                {'id': 'contract-legal', 'title': 'Legal Considerations'}
            ]
        })

        sections.append({
            'id': 'section-project-schedule',
            'title': 'Project Schedule & Milestones',
            'status': 'Pending',
            'guidance': 'Confirm the bond amount, surety approval, and Treasury Circular 570 listing for the guarantor.',
            'content': [
                f"{company_display} will coordinate with the approved surety to furnish the required Project Schedule & Milestones prior to submission.",
                "Include the original bid bond or irrevocable letter of credit documentation in the final package."
            ],
           
        })

        sections.append({
            'id': 'requirements-attachments',
            'title': 'Requirements and Attachments',
            'status': 'Special Section',
            'guidance': 'List mandatory submission requirements, certifications, and attachment references extracted from the RFP.',
            'content': [
                'Compile all mandatory attachments, compliance documents, and certifications referenced throughout the solicitation.',
                'Ensure each attachment is labeled clearly and referenced in the compliance matrix.'
            ],
            'sub_categories': [
                {'id': 'req-mandatory', 'title': 'Mandatory Documents'},
                {'id': 'req-forms', 'title': 'Forms & Certifications'},
                {'id': 'req-checklist', 'title': 'Submission Checklist'}
            ]
        })

        sections.append({
            'id': 'section-amendments',
            'title': 'Addenda',
            'status': 'Pending',
            'guidance': 'Log each amendment number and release date. If no amendments were issued, include a statement confirming that status.',
            'content': [
                f"The offer or acknowledges receipt of all amendments released for {project_name} as of {datetime.now().strftime('%B %d, %Y')}." if project_name else
                f"The offer or acknowledges receipt of all solicitation amendments issued as of {datetime.now().strftime('%B %d, %Y')}.",
                "Update this section with individual amendment numbers and dates when the contracting officer publishes them."
            ],
            'sub_categories': [
                {'id': 'amend-list', 'title': 'Amendment List'},
                {'id': 'amend-acknowledgment', 'title': 'Acknowledgment Statement'},
                {'id': 'amend-impact', 'title': 'Impact Assessment'}
            ]
        })

        sections.append({
            'id': 'section-appendix',
            'title': 'Appendix',
            'status': 'Draft',
            'guidance': 'Include all supporting documents, references, and additional materials referenced in the proposal.',
            'content': [
                'Supporting documents and references.',
                'Additional technical specifications.',
                'Supplementary materials and exhibits.'
            ],
            'sub_categories': [
                {'id': 'appendix-documents', 'title': 'Supporting Documents'},
                {'id': 'appendix-references', 'title': 'References'},
                {'id': 'appendix-exhibits', 'title': 'Exhibits'},
                {'id': 'appendix-other', 'title': 'Other Materials'}
            ]
        })

        sections.append({
            'id': 'section-safety',
            'title': 'Safety',
            'status': 'Draft',
            'guidance': 'Describe safety management systems, OSHA compliance, hazard mitigation, and jobsite safety protocols. Include training, PPE, incident reporting, and continuous improvement practices.',
            'content': [
                'Overview of Safety Management System (SMS) and accountability.',
                'OSHA compliance approach and hazard identification and mitigation processes.',
                'Jobsite safety plan, tailgate meetings, and Job Hazard Analysis (JHA) procedures.',
                'Training matrix, certifications, and PPE standards.',
                'Incident reporting, root cause analysis, and corrective actions.'
            ],
            'sub_categories': [
                {'id': 'safety-plan', 'title': 'Safety Plan'},
                {'id': 'safety-compliance', 'title': 'Regulatory Compliance'},
                {'id': 'safety-training', 'title': 'Training & Certifications'},
                {'id': 'safety-incident', 'title': 'Incident Response & Reporting'}
            ]
        })

        sections.append({
            'id': 'section-interconnection',
            'title': 'Interconnection',
            'status': 'Draft',
            'guidance': 'Detail utility interconnection strategy including application, studies, protection coordination, metering, and commissioning. Align with applicable tariffs, IEEE 1547, and utility standards.',
            'content': [
                'Interconnection application workflow and milestone tracking.',
                'Load flow, short circuit, and protection coordination study assumptions.',
                'Protection scheme, relays, and anti-islanding compliance.',
                'Metering, telemetry/SCADA, and communications architecture.',
                'Commissioning, witness testing, and as-built documentation.'
            ],
            'sub_categories': [
                {'id': 'ic-applications', 'title': 'Applications & Permits'},
                {'id': 'ic-studies', 'title': 'Engineering Studies'},
                {'id': 'ic-protection', 'title': 'Protection & Controls'},
                {'id': 'ic-commissioning', 'title': 'Commissioning & Testing'}
            ]
        })

        return sections

    if request.method == 'POST':
        try:
            files = request.files.getlist('rfp_files')
            company = request.form.get('company', '').strip()
            template = request.form.get('template', 'standard').strip()
            bid_name = request.args.get('bid_name', '') or request.form.get('bid_name', '')
            bid_id_value = (
                request.args.get('bid_id', type=int)
                or request.form.get('bid_id', type=int)
                or request.args.get('g_id', type=int)
                or request.form.get('g_id', type=int)
            )
            bid_meta_for_generation = get_bid_metadata(bid_id_value)
            sections_outline_for_generation = build_sections_outline(
                bid_meta_for_generation,
                bid_name,
                company,
                contact_email
            )
            
            if not files or len(files) == 0:
                flash('Please select at least one PDF file.', 'error')
                return redirect(url_for('proposals_making'))
            
            if not company:
                flash('Please select a company.', 'error')
                return redirect(url_for('proposals_making'))
            
            import uuid
            from werkzeug.utils import secure_filename
            
            # Create proposals directory
            proposals_dir = 'uploads/proposals'
            os.makedirs(proposals_dir, exist_ok=True)
            
            processed_files = []
            errors = []
            
            for file in files:
                if file and file.filename:
                    if not file.filename.lower().endswith('.pdf'):
                        errors.append(f"{file.filename} is not a PDF file")
                        continue
                    
                    # Save uploaded file temporarily
                    filename = secure_filename(file.filename)
                    unique_id = str(uuid.uuid4())[:8]
                    temp_filename = f"{unique_id}_{filename}"
                    temp_path = os.path.join('uploads/rfp', temp_filename)
                    os.makedirs('uploads/rfp', exist_ok=True)
                    file.save(temp_path)
                    
                    try:
                        # Generate complete proposal with all sections
                        proposal_filename = f"proposal_{unique_id}_{filename.replace('.pdf', '.docx')}"
                        proposal_path = os.path.join(proposals_dir, proposal_filename)
                        
                        # Generate complete proposal using AI
                        success = generate_complete_proposal(
                            temp_path,
                            company,
                            template,
                            proposal_path,
                            bid_name=bid_name,
                            sections_outline=sections_outline_for_generation
                        )
                        
                        if success:
                            # Save proposal data to database
                            cur = mysql.connection.cursor(DictCursor)
                            try:
                                # Create proposals table if it doesn't exist
                                cur.execute("""
                                    CREATE TABLE IF NOT EXISTS proposals (
                                        id INT AUTO_INCREMENT PRIMARY KEY,
                                        unique_id VARCHAR(50) NOT NULL UNIQUE,
                                        rfp_name VARCHAR(500) NOT NULL,
                                        rfp_filename VARCHAR(500),
                                        proposal_filename VARCHAR(500) NOT NULL,
                                        proposal_path VARCHAR(1000) NOT NULL,
                                        company VARCHAR(100) NOT NULL,
                                        template VARCHAR(100),
                                        bid_id INT NULL,
                                        bid_name VARCHAR(500) NULL,
                                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                                        INDEX idx_unique_id (unique_id),
                                        INDEX idx_rfp_name (rfp_name),
                                        INDEX idx_company (company),
                                        INDEX idx_bid_id (bid_id)
                                    ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
                                """)
                                
                                # Get bid_id (aka g_id) and bid_name from query parameters if available
                                bid_id = (
                                    request.args.get('bid_id', type=int)
                                    or request.form.get('bid_id', type=int)
                                    or request.args.get('g_id', type=int)
                                    or request.form.get('g_id', type=int)
                                )
                                bid_name = request.args.get('bid_name', '') or request.form.get('bid_name', '')
                                
                                # Insert proposal record
                                cur.execute("""
                                    INSERT INTO proposals 
                                    (unique_id, rfp_name, rfp_filename, proposal_filename, proposal_path, company, template, bid_id, bid_name)
                                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """, (
                                    unique_id,
                                    filename,  # RFP name
                                    filename,  # RFP filename
                                    proposal_filename,
                                    proposal_path,
                                    company,
                                    template,
                                    bid_id if bid_id else None,
                                    bid_name if bid_name else None
                                ))
                                mysql.connection.commit()
                                
                                processed_files.append({
                                    'original': filename,
                                    'proposal': proposal_filename,
                                    'path': proposal_path,
                                    'unique_id': unique_id,
                                    'rfp_name': filename
                                })
                                log_write('proposal_generated', f"file={filename} company={company} template={template} unique_id={unique_id} bid_id={bid_id}")
                            except Exception as db_error:
                                mysql.connection.rollback()
                                print(f"Database error: {str(db_error)}")
                                # Still add to processed_files even if DB save fails
                                processed_files.append({
                                    'original': filename,
                                    'proposal': proposal_filename,
                                    'path': proposal_path,
                                    'unique_id': unique_id,
                                    'rfp_name': filename
                                })
                            finally:
                                cur.close()
                        else:
                            errors.append(f"Failed to generate executive summary for {filename}")
                    except Exception as e:
                        errors.append(f"Error processing {filename}: {str(e)}")
                        if os.path.exists(temp_path):
                            os.remove(temp_path)
            
            if processed_files:
                flash(f'Successfully generated {len(processed_files)} proposal(s).', 'success')
            if errors:
                flash(f'Errors: {"; ".join(errors)}', 'error')
            
            return redirect(url_for('proposals_making'))
        
        except Exception as e:
            flash(f'Error: {str(e)}', 'error')
            return redirect(url_for('proposals_making'))
    
    # GET request - show page
    import os
    from datetime import datetime
    
    # Get bid_id (a.k.a g_id) and bid_name from query parameters if available (from Generate Proposals button)
    bid_id = request.args.get('bid_id', type=int)
    if bid_id is None:
        bid_id = request.args.get('g_id', type=int)
    bid_name = request.args.get('bid_name', '')
    company = request.args.get('company', '')

    contact_email = getattr(current_user, 'email', '')

    def safe_str(value):
        if value is None:
            return ''
        return str(value).strip()

    def format_date_for_display(value):
        if not value:
            return ''
        if isinstance(value, (datetime, date)):
            return value.strftime('%B %d, %Y')
        try:
            parsed = datetime.strptime(str(value), '%Y-%m-%d')
            return parsed.strftime('%B %d, %Y')
        except Exception:
            return str(value)

    def split_text_blocks(text, max_blocks=3):
        if not text:
            return []
        blocks = [
            safe_str(piece).strip('•- ') for piece in re.split(r'(?:\r?\n){2,}', str(text))
            if safe_str(piece)
        ]
        if not blocks and safe_str(text):
            blocks = [safe_str(text)]
        return blocks[:max_blocks]

    def compact_list(items):
        result = []
        for item in items:
            if item is None:
                continue
            text = safe_str(item)
            if text:
                result.append(text)
        return result

    def build_sections_outline(bid_info, fallback_title, company_name, primary_contact):
        bid = bid_info or {}
        project_name = safe_str(bid.get('b_name')) or safe_str(fallback_title) or 'Current Opportunity'
        due_text = format_date_for_display(bid.get('due_date'))
        client_name = safe_str(
            bid.get('agency') or bid.get('customer') or bid.get('client') or bid.get('owner') or bid.get('comp_name')
        )
        scope_text = safe_str(bid.get('scope'))
        summary_text = safe_str(bid.get('summary'))
        type_text = safe_str(bid.get('type'))
        location_text = safe_str(bid.get('location') or bid.get('city') or bid.get('state'))
        stage_text = safe_str(bid.get('state'))
        company_display = safe_str(company_name) or safe_str(bid.get('company')) or 'Offer or'
        revenue_value = bid.get('revenue') if bid.get('revenue') not in (None, '') else bid.get('value')
        scoring_value = bid.get('scoring')

        summary_blocks = split_text_blocks(summary_text, 4)
        scope_blocks = split_text_blocks(scope_text, 5)

        revenue_text = ''
        if revenue_value not in (None, ''):
            try:
                revenue_text = f"${float(revenue_value):,.0f}"
            except Exception:
                revenue_text = safe_str(revenue_value)

        scoring_text = ''
        if scoring_value not in (None, ''):
            try:
                scoring_text = f"{float(scoring_value):.0f}"
            except Exception:
                scoring_text = safe_str(scoring_value)

        sections = []

        cover_content = compact_list([
            f"{company_display} is pleased to submit this proposal for {project_name}.",
            f"Issuing agency / customer: {client_name}." if client_name else '',
            f"Solicitation reference: {project_name}.",
            f"Proposal due date: {due_text}." if due_text else '',
            f"Primary point of contact: {primary_contact}." if primary_contact else '',
            f"Contract type: {type_text}." if type_text else '',
            f"Performance location: {location_text}." if location_text else '',
            f"Opportunity priority score: {scoring_text}." if scoring_text else ''
        ])
        sections.append({
            'id': 'section-cover',
            'title': 'Letter of Transmittal',
            'status': 'Saved' if cover_content else 'Draft',
            'guidance': 'Confirm the solicitation title, client information, offer or contact details, and submission identifiers before releasing the cover page.',
            'content': cover_content or ['Add cover page details once the capture team confirms the solicitation data.'],
            'sub_categories': [
                {'id': 'cover-company-info', 'title': 'Company Information'},
                {'id': 'cover-contact-details', 'title': 'Contact Details'},
                {'id': 'cover-solicitation-ref', 'title': 'Solicitation Reference'},
                {'id': 'cover-submission-date', 'title': 'Submission Date & Deadline'}
            ]
        })

        sections.append({
            'id': 'section-toc',
            'title': 'Table of Contents',
            'status': 'Saved',
            'guidance': 'Generate a comprehensive table of contents listing all proposal sections and their page numbers.',
            'content': [
                'The table of contents will be automatically generated once all sections are finalized.',
                'Ensure all major sections and subsections are included with accurate page references.'
            ],
            'sub_categories': [
                {'id': 'toc-main-sections', 'title': 'Main Sections'},
                {'id': 'toc-appendices', 'title': 'Appendices'},
                {'id': 'toc-figures-tables', 'title': 'Figures & Tables'}
            ]
        })

        exec_content = summary_blocks or (
            [f"{company_display} will deliver a compliant, value-focused response for {project_name}."]
            if project_name else []
        )
        if scope_text and not summary_blocks:
            exec_content.append(scope_text)
        sections.append({
            'id': 'section-executive-summary',
            'title': 'Executive Summary',
            'status': 'In Progress' if summary_blocks else 'Draft',
            'guidance': "Summarize win themes, differentiators, and alignment to the customer's objectives captured during bid qualification.",
            'content': exec_content or ['Executive summary content will be generated as soon as capture notes are available.'],
            'sub_categories': [
                {'id': 'exec-overview', 'title': 'Project Overview'},
                {'id': 'exec-win-themes', 'title': 'Win Themes & Differentiators'},
                {'id': 'exec-value-proposition', 'title': 'Value Proposition'},
                {'id': 'exec-key-highlights', 'title': 'Key Highlights'}
            ]
        })

        compliance_content = compact_list([
            "Track each requirement in the compliance matrix to ensure every section of the RFP is addressed.",
            scope_blocks[0] if scope_blocks else '',
            summary_blocks[0] if summary_blocks else ''
        ])
        sections.append({
            'id': 'section-compliance',
            'title': 'Requirements Compliance',
            'status': 'In Progress',
            'guidance': 'Reference the compliance matrix and list every deliverable, attachment, and form demanded by the solicitation.',
            'content': compliance_content or [
                'Populate this section with compliance findings once the RFP analysis is complete.'
            ],
            'sub_categories': [
                {'id': 'comp-matrix', 'title': 'Compliance Matrix'},
                {'id': 'comp-deliverables', 'title': 'Deliverables List'},
                {'id': 'comp-attachments', 'title': 'Required Attachments'},
                {'id': 'comp-forms', 'title': 'Required Forms'}
            ]
        })

        sections.append({
            'id': 'section-scope-objectives',
            'title': 'Understanding of Scope & Objectives',
            'status': 'In Progress',
            'guidance': 'Demonstrate a clear understanding of the project scope, objectives, and expected outcomes.',
            'content': scope_blocks or [
                'Detail your understanding of the project scope and objectives based on the solicitation requirements.'
            ],
            'sub_categories': [
                {'id': 'scope-project-scope', 'title': 'Project Scope'},
                {'id': 'scope-objectives', 'title': 'Project Objectives'},
                {'id': 'scope-expected-outcomes', 'title': 'Expected Outcomes'},
                {'id': 'scope-assumptions', 'title': 'Key Assumptions'}
            ]
        })

        sections.append({
            'id': 'section-deviations',
            'title': 'Deviations, Assumptions, and Dependencies (if any)',
            'status': 'Draft',
            'guidance': 'Document any deviations from the solicitation requirements, key assumptions, and dependencies that may impact project delivery.',
            'content': [
                'List any deviations from the solicitation requirements, if applicable.',
                'Document key assumptions made during proposal development.',
                'Identify dependencies that may affect project execution.'
            ],
            'sub_categories': [
                {'id': 'dev-deviations', 'title': 'Deviations'},
                {'id': 'dev-assumptions', 'title': 'Assumptions'},
                {'id': 'dev-dependencies', 'title': 'Dependencies'},
                {'id': 'dev-risks', 'title': 'Risk Considerations'}
            ]
        })

        tech_content = scope_blocks or [
            'Detail the technical solution, materials, and implementation methodology once the RFP requirements are parsed.'
        ]
        sections.append({
            'id': 'section-technical',
            'title': 'Proposed Technical Solution',
            'status': 'In Progress' if scope_blocks else 'Draft',
            'guidance': 'Translate the solicitation Statement of Work into your delivery plan, highlighting compliance, innovations, and risk mitigations.',
            'content': tech_content,
            'sub_categories': [
                {'id': 'tech-approach', 'title': 'Technical Approach'},
                {'id': 'tech-methodology', 'title': 'Methodology'},
                {'id': 'tech-innovations', 'title': 'Innovations & Best Practices'},
                {'id': 'tech-compliance', 'title': 'Compliance & Standards'},
                {'id': 'tech-risk-mitigation', 'title': 'Risk Mitigation'}
            ]
        })

        sections.append({
            'id': 'section-implementation',
            'title': 'Implementation & Work Plan',
            'status': 'In Progress',
            'guidance': 'Provide a detailed implementation plan and work breakdown structure for project execution.',
            'content': [
                'Outline the step-by-step implementation approach.',
                'Detail the work breakdown structure and key activities.',
                'Describe the workflow and processes for project delivery.'
            ],
            'sub_categories': [
                {'id': 'impl-approach', 'title': 'Implementation Approach'},
                {'id': 'impl-work-breakdown', 'title': 'Work Breakdown Structure'},
                {'id': 'impl-phases', 'title': 'Implementation Phases'},
                {'id': 'impl-processes', 'title': 'Key Processes'}
            ]
        })

        sections.append({
            'id': 'section-deliverables',
            'title': 'Solution Deliverables',
            'status': 'In Progress',
            'guidance': 'List all deliverables that will be provided as part of the solution, including documentation, reports, and physical items.',
            'content': [
                'Comprehensive list of all solution deliverables.',
                'Documentation and reports to be provided.',
                'Physical deliverables, if applicable.'
            ],
            'sub_categories': [
                {'id': 'del-documentation', 'title': 'Documentation Deliverables'},
                {'id': 'del-reports', 'title': 'Reports & Analysis'},
                {'id': 'del-physical', 'title': 'Physical Deliverables'},
                {'id': 'del-schedule', 'title': 'Delivery Schedule'}
            ]
        })

        management_content = compact_list([
            f"Target submission date: {due_text}." if due_text else 'Submission date will be updated when the procurement schedule is confirmed.',
          
            f"Contract type guidance: {type_text}." if type_text else '',
            "Outline major milestones, design reviews, and government touchpoints that support a compliant response."
        ])
        sections.append({
            'id': 'section-management',
            'title': 'Project Management Plan',
            'status': 'In Progress',
            'guidance': 'Map the internal review cycle, color team dates, and delivery milestones to ensure on-time submission.',
            'content': management_content or [
                'Add the project management narrative and schedule once the capture calendar is finalized.'
            ],
            'sub_categories': [
                {'id': 'mgmt-organization', 'title': 'Project Organization'},
                {'id': 'mgmt-schedule', 'title': 'Project Schedule & Milestones'},
                {'id': 'mgmt-resources', 'title': 'Resource Management'},
                {'id': 'mgmt-quality', 'title': 'Quality Assurance'},
                {'id': 'mgmt-communication', 'title': 'Communication Plan'}
            ]
        })

        price_content = compact_list([
            f"Estimated contract value: {revenue_text}." if revenue_text else '',
            "Attachment A pricing will mirror the government-provided workbook with CLIN-level details.",
            "Document all pricing assumptions, escalation factors, and exclusions for contracting officer review."
        ])
        sections.append({
            'id': 'section-pricing',
            'title': 'Pricing Proposals',
            'status': 'Saved' if revenue_text else 'Draft',
            'guidance': 'Ensure Attachment A reflects current quantities, rates, and assumptions. Cross-check CLIN totals before submission.',
            'content': price_content or [
                'Populate Attachment A once final quantities and labor categories are confirmed.'
            ],
            'sub_categories': [
                {'id': 'price-clin-breakdown', 'title': 'CLIN Breakdown'},
                {'id': 'price-labor-rates', 'title': 'Labor Rates & Categories'},
                {'id': 'price-materials', 'title': 'Materials & Equipment'},
                {'id': 'price-assumptions', 'title': 'Pricing Assumptions'},
                {'id': 'price-total-summary', 'title': 'Total Cost Summary'}
            ]
        })

        sections.append({
            'id': 'section-contractual',
            'title': 'Contractual Terms',
            'status': 'Saved',
            'guidance': 'Outline key contractual terms, conditions, and agreements proposed for this engagement.',
            'content': [
                'Key contractual terms and conditions.',
                'Payment terms and conditions.',
                'Terms of service and deliverables.'
            ],
            'sub_categories': [
                {'id': 'contract-terms', 'title': 'Terms & Conditions'},
                {'id': 'contract-payment', 'title': 'Payment Terms'},
                {'id': 'contract-service', 'title': 'Service Terms'},
                {'id': 'contract-legal', 'title': 'Legal Considerations'}
            ]
        })

        sections.append({
            'id': 'section-project-schedule',
            'title': 'Project Schedule & Milestones',
            'status': 'Pending',
            'guidance': 'Confirm the bond amount, surety approval, and Treasury Circular 570 listing for the guarantor.',
            'content': [
                f"{company_display} will coordinate with the approved surety to furnish the required Project Schedule & Milestones prior to submission.",
                "Include the original bid bond or irrevocable letter of credit documentation in the final package."
            ],
           
        })

        sections.append({
            'id': 'section-team',
            'title': 'Project Team',
            'status': 'In Progress',
            'guidance': 'Introduce the project team members, their roles, qualifications, and responsibilities.',
            'content': [
                'List of key team members and their roles.',
                'Qualifications and experience of team members.',
                'Organizational structure and reporting lines.'
            ],
            'sub_categories': [
                {'id': 'team-members', 'title': 'Team Members'},
                {'id': 'team-roles', 'title': 'Roles & Responsibilities'},
                {'id': 'team-qualifications', 'title': 'Qualifications'},
                {'id': 'team-organization', 'title': 'Organization Structure'}
            ]
        })

        sections.append({
            'id': 'section-corporate',
            'title': 'Corporate Qualifications',
            'status': 'Saved',
            'guidance': 'Highlight corporate experience, past performance, certifications, and qualifications relevant to this opportunity.',
            'content': [
                'Company background and history.',
                'Relevant past performance and case studies.',
                'Certifications, accreditations, and qualifications.'
            ],
            'sub_categories': [
                {'id': 'corp-background', 'title': 'Company Background'},
                {'id': 'corp-past-performance', 'title': 'Past Performance'},
                {'id': 'corp-certifications', 'title': 'Certifications & Accreditations'},
                {'id': 'corp-capabilities', 'title': 'Key Capabilities'}
            ]
        })

        sections.append({
            'id': 'section-reps-certs',
            'title': 'Representations and Certifications',
            'status': 'Saved',
            'guidance': 'Confirm FAR, DFARS, SAM.gov, and agency-specific representations are current and included.',
            'content': [
                'All FAR and DFARS representations will be validated against the active SAM registration prior to submission.',
                'Include any agency-specific attestations or forms requested in the solicitation instructions.'
            ],
            'sub_categories': [
                {'id': 'reps-far', 'title': 'FAR Representations'},
                {'id': 'reps-dfars', 'title': 'DFARS Representations'},
                {'id': 'reps-sam', 'title': 'SAM.gov Certifications'},
                {'id': 'reps-agency', 'title': 'Agency-Specific Certifications'}
            ]
        })

        sections.append({
            'id': 'section-amendments',
            'title': 'Addenda',
            'status': 'Pending',
            'guidance': 'Log each amendment number and release date. If no amendments were issued, include a statement confirming that status.',
            'content': [
                f"The offer or acknowledges receipt of all amendments released for {project_name} as of {datetime.now().strftime('%B %d, %Y')}." if project_name else
                f"The offer or acknowledges receipt of all solicitation amendments issued as of {datetime.now().strftime('%B %d, %Y')}.",
                "Update this section with individual amendment numbers and dates when the contracting officer publishes them."
            ],
            'sub_categories': [
                {'id': 'amend-list', 'title': 'Amendment List'},
                {'id': 'amend-acknowledgment', 'title': 'Acknowledgment Statement'},
                {'id': 'amend-impact', 'title': 'Impact Assessment'}
            ]
        })

        sections.append({
            'id': 'section-appendix',
            'title': 'Appendix',
            'status': 'Draft',
            'guidance': 'Include all supporting documents, references, and additional materials referenced in the proposal.',
            'content': [
                'Supporting documents and references.',
                'Additional technical specifications.',
                'Supplementary materials and exhibits.'
            ],
            'sub_categories': [
                {'id': 'appendix-documents', 'title': 'Supporting Documents'},
                {'id': 'appendix-references', 'title': 'References'},
                {'id': 'appendix-exhibits', 'title': 'Exhibits'},
                {'id': 'appendix-other', 'title': 'Other Materials'}
            ]
        })

        return sections

    def build_bid_context(bid_info, fallback_title, company_name):
        info = bid_info or {}
        project_name = safe_str(info.get('b_name')) or safe_str(fallback_title)
        due_text = format_date_for_display(info.get('due_date'))
        summary_text = safe_str(info.get('summary'))
        scope_text = safe_str(info.get('scope'))
        location_text = safe_str(info.get('location') or info.get('city') or info.get('state'))
        type_text = safe_str(info.get('type'))
        stage_text = safe_str(info.get('state'))
        revenue_value = info.get('revenue') if info.get('revenue') not in (None, '') else info.get('value')
        scoring_value = info.get('scoring')

        revenue_text = ''
        if revenue_value not in (None, ''):
            try:
                revenue_text = f"${float(revenue_value):,.0f}"
            except Exception:
                revenue_text = safe_str(revenue_value)

        scoring_text = ''
        if scoring_value not in (None, ''):
            try:
                scoring_text = f"{float(scoring_value):.0f}"
            except Exception:
                scoring_text = safe_str(scoring_value)

        summary_excerpt = ''
        source_text = summary_text or scope_text
        if source_text:
            summary_excerpt = source_text if len(source_text) <= 220 else source_text[:217].rstrip() + '…'

        info_chips = []
        if due_text:
            info_chips.append({'label': 'Due', 'value': due_text, 'icon': 'calendar-days'})
        if type_text:
            info_chips.append({'label': 'Type', 'value': type_text, 'icon': 'diagram-project'})
        if location_text:
            info_chips.append({'label': 'Location', 'value': location_text, 'icon': 'location-dot'})
        if revenue_text:
            info_chips.append({'label': 'Est. Value', 'value': revenue_text, 'icon': 'sack-dollar'})
        if stage_text:
            info_chips.append({'label': 'Stage', 'value': stage_text.title() if stage_text else stage_text, 'icon': 'gauge-high'})
        if scoring_text:
            info_chips.append({'label': 'Score', 'value': scoring_text, 'icon': 'star-half-stroke'})

        context = {
            'project_name': project_name or safe_str(fallback_title),
            'client': safe_str(
                info.get('agency') or info.get('customer') or info.get('client') or info.get('owner')
            ),
            'company': safe_str(company_name) or safe_str(info.get('company')),
            'due_date_display': due_text,
            'summary_excerpt': summary_excerpt,
            'scope': scope_text,
            'type': type_text,
            'location': location_text,
            'rfp_label': project_name or safe_str(fallback_title),
            'info_chips': info_chips,
            'bid_id': info.get('g_id') or info.get('id'),
            'g_id': info.get('g_id')
        }
        # Remove empty values for cleaner JSON payloads
        return {k: v for k, v in context.items() if v}

    bid_meta_raw = None
    if bid_id:
        meta_cursor = mysql.connection.cursor(DictCursor)
        try:
            meta_cursor.execute("SELECT * FROM go_bids WHERE g_id=%s", (bid_id,))
            bid_meta_raw = meta_cursor.fetchone()
        except Exception as meta_err:
            print(f"Error loading bid context for proposal view: {meta_err}")
        finally:
            meta_cursor.close()

    if not company and bid_meta_raw:
        company = safe_str(bid_meta_raw.get('company'))

    sections_outline = build_sections_outline(bid_meta_raw, bid_name, company or (bid_meta_raw or {}).get('company'), contact_email)
    bid_context = build_bid_context(bid_meta_raw, bid_name, company or (bid_meta_raw or {}).get('company'))

    # Load proposals from database
    proposals_list = []
    cur = mysql.connection.cursor(DictCursor)
    try:
        # Create table if it doesn't exist
        cur.execute("""
            CREATE TABLE IF NOT EXISTS proposals (
                id INT AUTO_INCREMENT PRIMARY KEY,
                unique_id VARCHAR(50) NOT NULL UNIQUE,
                rfp_name VARCHAR(500) NOT NULL,
                rfp_filename VARCHAR(500),
                proposal_filename VARCHAR(500) NOT NULL,
                proposal_path VARCHAR(1000) NOT NULL,
                company VARCHAR(100) NOT NULL,
                template VARCHAR(100),
                bid_id INT NULL,
                bid_name VARCHAR(500) NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                INDEX idx_unique_id (unique_id),
                INDEX idx_rfp_name (rfp_name),
                INDEX idx_company (company),
                INDEX idx_bid_id (bid_id)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """)
        
        # Query proposals from database
        if bid_id:
            cur.execute("""
                SELECT * FROM proposals 
                WHERE bid_id = %s 
                ORDER BY created_at DESC
            """, (bid_id,))
        else:
            cur.execute("""
                SELECT * FROM proposals 
                ORDER BY created_at DESC
                LIMIT 100
            """)
        
        db_proposals = cur.fetchall()
        
        # Convert database results to list format
        for prop in db_proposals:
            proposals_list.append({
                'id': prop['id'],
                'unique_id': prop['unique_id'],
                'name': prop['proposal_filename'],
                'rfp_name': prop['rfp_name'],
                'company': prop['company'],
                'template': prop.get('template', ''),
                'bid_id': prop.get('bid_id'),
                'bid_name': prop.get('bid_name', ''),
                'created_at': prop['created_at'].strftime('%Y-%m-%d %H:%M:%S') if prop['created_at'] else '',
                'download_url': url_for('download_proposal', filename=prop['proposal_filename'])
            })
    except Exception as e:
        print(f"Error loading proposals from database: {str(e)}")
        # Fallback to file system if database fails
        proposals_dir = 'uploads/proposals'
        if os.path.exists(proposals_dir):
            for filename in os.listdir(proposals_dir):
                if filename.endswith('.docx'):
                    filepath = os.path.join(proposals_dir, filename)
                    stat = os.stat(filepath)
                    proposals_list.append({
                        'name': filename,
                        'rfp_name': filename.replace('executive_summary_', '').replace('.docx', '.pdf'),
                        'created_at': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
                        'download_url': url_for('download_proposal', filename=filename)
                    })
    finally:
        cur.close()
    
    # Load RFP files from database for this bid
    rfp_files_list = []
    if bid_id:
        rfp_cursor = mysql.connection.cursor(DictCursor)
        try:
            _ensure_uploaded_rfp_table_exists(rfp_cursor)
            # Fetch all RFP files for this bid
            rfp_cursor.execute("""
                SELECT id, g_id, bid_id, filename, original_filename, saved_filename, file_path, file_size, uploaded_at
                FROM uploaded_rfp_files
                WHERE g_id = %s OR bid_id = %s
                ORDER BY uploaded_at DESC
            """, (bid_id, bid_id))
            rfp_files = rfp_cursor.fetchall()
            
            for rfp_file in rfp_files:
                file_path = rfp_file.get('file_path')
                if file_path and os.path.exists(file_path):
                    rfp_files_list.append({
                        'id': rfp_file.get('id'),
                        'filename': rfp_file.get('filename') or rfp_file.get('original_filename', ''),
                        'file_path': file_path,
                        'file_size': rfp_file.get('file_size', 0),
                        'uploaded_at': rfp_file.get('uploaded_at'),
                        'view_url': url_for('view_rfp_file', file_id=rfp_file.get('id')),
                        'parsed_url': url_for('api_rfp_file_parsed', g_id=bid_id)
                    })
        except Exception as e:
            print(f"Error loading RFP files from database: {str(e)}")
        finally:
            rfp_cursor.close()
    
    return render_template(
        'proposals_making.html',
        proposals=proposals_list,
        bid_id=bid_id,
        bid_name=bid_name,
        company=company,
        sections_outline=sections_outline,
        bid_context=bid_context,
        contact_email=contact_email,
        rfp_files=rfp_files_list
    )

@app.route('/api/proposals-agent', methods=['POST'])
@login_required
def proposals_agent():
    """Lightweight assistant endpoint for the proposals workspace."""
    payload = request.get_json(silent=True) or {}
    message = (payload.get('message') or '').strip()
    if not message:
        return jsonify({'error': 'missing_message', 'message': 'Please provide a question for the agent.'}), 400

    history = payload.get('history') or []
    if not isinstance(history, list):
        history = []
    history = history[-10:]

    context = payload.get('context') or {}
    try:
        context_summary = {
            'workspace': 'Proposals Making',
            'user': getattr(current_user, 'email', ''),
            'company': context.get('company') or '',
            'contact_email': context.get('contactEmail') or '',
        }
        bid_info = context.get('bid') or {}
        g_id = None
        if isinstance(bid_info, dict) and bid_info:
            g_id = bid_info.get('g_id') or bid_info.get('bid_id')
            context_summary['bid'] = {
                'project_name': bid_info.get('project_name') or bid_info.get('b_name') or '',
                'bid_id': bid_info.get('bid_id') or bid_info.get('g_id') or '',
                'due_date': bid_info.get('due_date') or '',
            }
        sections = context.get('sections') or []
        if isinstance(sections, list) and sections:
            context_summary['sections'] = [
                {
                    'title': sec.get('title', ''),
                    'status': sec.get('status', ''),
                    'guidance': (sec.get('guidance') or '')[:280],
                }
                for sec in sections[:5]
                if isinstance(sec, dict)
            ]
    except Exception:
        context_summary = {}

    # Retrieve RFP document content for RAG
    rfp_content = ""
    if g_id:
        cur = None
        try:
            cur = mysql.connection.cursor(DictCursor)
            _ensure_uploaded_rfp_table_exists(cur)
            # Get the latest RFP file for this bid
            rfp_file = _get_latest_rfp_file_for_bid(g_id)
            if rfp_file and rfp_file.get('file_path'):
                file_path = rfp_file['file_path']
                if os.path.exists(file_path):
                    # Extract text from first 50 pages (to avoid token limits)
                    pages, total_pages = _extract_pdf_pages(file_path, start_page=1, limit=50)
                    if pages:
                        # Combine all page texts
                        rfp_texts = [page.get('text', '') for page in pages if page.get('text')]
                        if rfp_texts:
                            # Limit total content to ~8000 characters to stay within token limits
                            combined_text = "\n\n--- PAGE BREAK ---\n\n".join(rfp_texts)
                            if len(combined_text) > 8000:
                                combined_text = combined_text[:8000] + "... [Content truncated]"
                            rfp_content = f"\n\n--- RFP DOCUMENT CONTENT (Pages 1-{min(50, total_pages)} of {total_pages}) ---\n{combined_text}\n--- END RFP CONTENT ---"
        except Exception as e:
            app.logger.warning(f"Error retrieving RFP content for RAG: {e}")
            rfp_content = ""
        finally:
            if cur:
                cur.close()

    system_prompt = (
        "You are the ESCO proposal workspace assistant. "
        "Provide concise, actionable support for crafting winning proposals. "
        "Use the workspace context when relevant, and be explicit when information is unavailable or uncertain."
    )
    if context_summary:
        try:
            system_prompt += "\n\nWorkspace context:\n" + json.dumps(context_summary, ensure_ascii=False, indent=2)
        except Exception:
            system_prompt += "\n\nWorkspace context:\n" + str(context_summary)
    
    # Add RFP document content for RAG
    if rfp_content:
        system_prompt += "\n\nIMPORTANT: The following RFP document content is available for reference. Use this content to answer questions about the RFP requirements, specifications, deadlines, and other details. Always base your responses on the actual RFP content when available:"
        system_prompt += rfp_content

    messages = [{'role': 'system', 'content': system_prompt}]
    for entry in history:
        if not isinstance(entry, dict):
            continue
        role = entry.get('role', 'user')
        if role not in {'user', 'assistant', 'system'}:
            role = 'user'
        content = entry.get('content')
        if not content:
            continue
        messages.append({'role': role, 'content': str(content)})
    messages.append({'role': 'user', 'content': message})

    api_key = (app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or '').strip()
    if not api_key:
        return jsonify({'error': 'agent_unavailable', 'message': 'OpenAI API key is not configured on the server.'}), 500

    base_url = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or 'https://api.groq.com/openai/v1').strip()
    model = (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'llama-3.1-70b-versatile').strip()
    endpoint = base_url.rstrip('/')
    if not endpoint.endswith('/chat/completions'):
        endpoint = f"{endpoint}/chat/completions"

    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json',
    }
    body = {
        'model': model,
        'messages': messages,
        'temperature': 0.2,
    }

    try:
        response = requests.post(endpoint, headers=headers, json=body, timeout=60)
    except requests.RequestException as err:
        app.logger.exception("OpenAI chat completion network failure: %s", err)
        return jsonify({'error': 'agent_unavailable', 'message': 'The proposal agent could not reach the language model service. Please try again shortly.'}), 502

    raw_text = response.text
    try:
        data = response.json()
    except ValueError:
        app.logger.error("OpenAI chat completion returned non-JSON response (status %s): %s", response.status_code, raw_text[:400])
        return jsonify({'error': 'agent_unavailable', 'message': 'Received an invalid response from the language model.'}), 502

    if response.status_code >= 400:
        err_message = data.get('error') if isinstance(data, dict) else None
        detail = ''
        if isinstance(err_message, dict):
            detail = err_message.get('message') or err_message.get('code') or ''
        elif isinstance(err_message, str):
            detail = err_message
        if not detail:
            detail = f"Upstream error {response.status_code}"
        app.logger.error("OpenAI chat completion error %s: %s", response.status_code, detail)
        return jsonify({'error': 'agent_unavailable', 'message': f'OpenAI returned an error: {detail}'}), 502

    choices = data.get('choices') or []
    if not choices:
        return jsonify({'error': 'agent_unavailable', 'message': 'No response returned by the language model.'}), 502

    reply = choices[0].get('message', {}).get('content', '')
    if not reply:
        return jsonify({'error': 'agent_unavailable', 'message': 'The language model returned an empty response.'}), 502

    return jsonify({
        'reply': reply.strip(),
        'model': data.get('model', model),
        'usage': data.get('usage', {}),
    })


def _fetch_rfp_content_for_bid(bid_id, page_limit=50, char_limit=12000):
    """Load RFP text from the uploaded files for a given bid or g_id."""
    if not bid_id:
        return ''

    rfp_content = ''
    try:
        cur = mysql.connection.cursor(DictCursor)
        cur.execute(
            """
            SELECT file_path
            FROM uploaded_rfp_files
            WHERE g_id = %s OR bid_id = %s
            ORDER BY uploaded_at DESC
            LIMIT 1
            """,
            (bid_id, bid_id),
        )
        rfp_file = cur.fetchone()
        cur.close()
    except Exception as exc:
        app.logger.warning(f"Error retrieving RFP content for bid {bid_id}: {exc}")
        return ''

    file_path = rfp_file.get('file_path') if rfp_file else None
    if not file_path or not os.path.exists(file_path):
        return ''

    try:
        if _FITZ_AVAILABLE:
            doc = _pymupdf.open(file_path)
            try:
                pages = []
                for page_num in range(min(page_limit, len(doc))):
                    page = doc[page_num]
                    pages.append(page.get_text())
                rfp_content = "\n\n--- PAGE BREAK ---\n\n".join(pages)
            finally:
                doc.close()
        else:
            app.logger.info("PyMuPDF not available; skipping PDF text extraction.")
    except Exception as exc:
        app.logger.warning(f"Error reading RFP PDF for bid {bid_id}: {exc}")
        rfp_content = ''

    if rfp_content and len(rfp_content) > char_limit:
        rfp_content = rfp_content[:char_limit] + "... [Content truncated]"

    return rfp_content


@app.route('/api/analyze-rfp-master', methods=['POST'])
@login_required
def analyze_rfp_master():
    """Run an end-to-end compliance and outline analysis for the active RFP."""
    payload = request.get_json(silent=True) or {}
    bid_id = payload.get('bid_id')

    if not bid_id:
        return jsonify({'error': 'missing_params', 'message': 'Bid ID is required.'}), 400

    rfp_content = _fetch_rfp_content_for_bid(bid_id, page_limit=120, char_limit=22000)
    if not rfp_content:
        return jsonify({'error': 'no_rfp', 'message': 'No RFP content found for this bid. Please upload an RFP file first.'}), 404

    truncated_rfp = rfp_content[:22000]

    analysis_instructions = """
You are an expert Proposal Analyst and Writer specialized in public procurement (RFP/RFQ/IFB) across construction, energy, IT, and professional services. Using the provided RFP corpus, produce a compliance-first analysis and proposal blueprint.

Requirements:
- Output valid HTML only. No Markdown, JSON, or plain text.
- Use <h2> headings in this exact order:
  1. Executive Brief
  2. Key Dates and Submission Rules
  3. Compliance Matrix
  4. Proposal Table of Contents
  5. Section-by-Section Guidance and Drafts
  6. Pricing Package Instructions
  7. SBE/DBE and Compliance Plan
  8. Bonds, Insurance, and Financials
  9. Risk Register and Clarifications
  10. Submission Checklist
- Provide concise, action-oriented guidance. Quote or paraphrase RFP sections with citations (e.g., Section 4, Instructions to Bidders).
- The Compliance Matrix must be a <table> with columns: Requirement, RFP Reference, Mandatory?, Proposal Location, Owner Role, Notes. Keep rows CSV-friendly (no extra commas unless inside quotes).
- In Section-by-Section Guidance, create a <section> per entry that includes:
    * <h3> heading for the section name.
    * <p> summarizing purpose and linkage.
    * <ul> of required content bullets.
    * <ul class="placeholders"> for data placeholders (e.g., [Vendor_Name]).
    * <pre class="boilerplate"> containing draft boilerplate text with placeholders.
    * <ul class="evidence"> listing evidence artifacts/forms to attach.
- Supply copy-paste templates wrapped in <pre class="template"> blocks for: Cover Letter, Executive Summary, Technical Approach, Safety Program, Small Business Plan, Pricing Cover, Bonds/Insurance Acknowledgment, References, Appendices List.
- Flag assumptions as “Assumption – verify”.
- If SBE/DBE goals exist, include a subcontracting plan template and tracking approach.
- Emphasize sealed bid compliance, bonding thresholds, signatures/notarization, pricing forms, safety stats, and portal rules when relevant to construction/MEP.
- Highlight phasing, technical standards, monitoring, and access coordination for solar/energy scopes when applicable.
- Never invent data; if missing, call it out explicitly.
- Finish with a concise checklist aligned to submission portal/file rules and include a Go/No-Go gate summary.
"""

    user_prompt = f"""
RFP / IFB Source Content:
{truncated_rfp}

Generate the HTML deliverable now, following every rule above.
"""

    try:
        from rfp_analyzer_routes import ollama_json
        llm_result = ollama_json(analysis_instructions + "\n\n" + user_prompt, 6000)
    except Exception as exc:
        app.logger.exception("Deep RFP analysis failed: %s", exc)
        return jsonify({'error': 'llm_failure', 'message': 'Unable to analyze the RFP at this time.'}), 500

    analysis_html = ''
    if isinstance(llm_result, dict):
        analysis_html = (
            llm_result.get('html')
            or llm_result.get('analysis')
            or llm_result.get('content')
            or llm_result.get('raw')
            or ''
        )
        if not analysis_html and llm_result:
            try:
                analysis_html = json.dumps(llm_result, indent=2)
            except Exception:
                analysis_html = str(llm_result)
    else:
        analysis_html = str(llm_result)

    if not analysis_html:
        analysis_html = "<p class='text-sm text-gray-500'>No analysis output was generated.</p>"

    return jsonify({'analysis_html': analysis_html})


@app.route('/api/analyze-subcategory', methods=['POST'])
@login_required
def analyze_subcategory():
    """Analyze a sub-category section using AI based on RFP content."""
    payload = request.get_json(silent=True) or {}
    section_id = payload.get('section_id', '').strip()
    subcategory_id = payload.get('subcategory_id', '').strip()
    subcategory_title = payload.get('subcategory_title', '').strip()
    bid_id = payload.get('bid_id')
    rfp_content = payload.get('rfp_content', '')
    
    if not section_id or not subcategory_id or not subcategory_title:
        return jsonify({'error': 'missing_params', 'message': 'Section ID, subcategory ID, and title are required.'}), 400
    
    # Get RFP content from database if bid_id is provided
    if bid_id and not rfp_content:
        rfp_content = _fetch_rfp_content_for_bid(bid_id, page_limit=50, char_limit=8000)
    
    # Build AI prompt for sub-category analysis
    system_prompt = (
        "You are an expert proposal writer helping to draft specific sections of a government proposal. "
        "Analyze the RFP content and generate comprehensive, compliant content for the requested sub-category. "
        "Your response should be professional, detailed, and aligned with the RFP requirements."
    )
    
    user_prompt = f"""Generate content for the sub-category: "{subcategory_title}" within the section "{section_id}".

Requirements:
1. Analyze the provided RFP content to understand requirements relevant to this sub-category
2. Generate professional, compliant content that addresses the sub-category topic
3. Ensure the content is specific, actionable, and aligned with government proposal standards
4. If RFP content is limited, provide a well-structured template that can be customized

RFP Content:
{rfp_content[:8000] if rfp_content else "No RFP content available. Generate a professional template for this sub-category."}

Please provide the content for "{subcategory_title}" in a clear, well-formatted manner."""
    
    api_key = (app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or '').strip()
    if not api_key:
        return jsonify({'error': 'api_unavailable', 'message': 'OpenAI API key is not configured on the server.'}), 500
    
    base_url = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or 'https://api.groq.com/openai/v1').strip()
    model = (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'llama-3.1-70b-versatile').strip()
    endpoint = base_url.rstrip('/')
    if not endpoint.endswith('/chat/completions'):
        endpoint = f"{endpoint}/chat/completions"
    
    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json',
    }
    body = {
        'model': model,
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt}
        ],
        'temperature': 0.3,
    }
    
    try:
        response = requests.post(endpoint, headers=headers, json=body, timeout=90)
    except requests.RequestException as err:
        app.logger.exception("OpenAI subcategory analysis network failure: %s", err)
        return jsonify({'error': 'api_unavailable', 'message': 'Could not reach the AI service. Please try again shortly.'}), 502
    
    try:
        data = response.json()
    except ValueError:
        app.logger.error("OpenAI subcategory analysis returned non-JSON response (status %s): %s", response.status_code, response.text[:400])
        return jsonify({'error': 'api_unavailable', 'message': 'Received an invalid response from the AI service.'}), 502
    
    if response.status_code >= 400:
        err_message = data.get('error') if isinstance(data, dict) else None
        detail = ''
        if isinstance(err_message, dict):
            detail = err_message.get('message') or err_message.get('code') or ''
        elif isinstance(err_message, str):
            detail = err_message
        if not detail:
            detail = f"Upstream error {response.status_code}"
        app.logger.error("OpenAI subcategory analysis error %s: %s", response.status_code, detail)
        return jsonify({'error': 'api_unavailable', 'message': f'AI service returned an error: {detail}'}), 502
    
    choices = data.get('choices') or []
    if not choices:
        return jsonify({'error': 'api_unavailable', 'message': 'No response returned by the AI service.'}), 502
    
    content = choices[0].get('message', {}).get('content', '')
    if not content:
        return jsonify({'error': 'api_unavailable', 'message': 'The AI service returned an empty response.'}), 502
    
    return jsonify({
        'content': content.strip(),
        'subcategory_id': subcategory_id,
        'subcategory_title': subcategory_title,
        'model': data.get('model', model),
        'usage': data.get('usage', {}),
    })


@app.route('/api/analyze-section-outline', methods=['POST'])
@login_required
def analyze_section_outline():
    """Extract key requirements and a summary for a proposal section directly from the RFP."""
    payload = request.get_json(silent=True) or {}
    section_id = (payload.get('section_id') or '').strip()
    section_title = (payload.get('section_title') or '').strip() or section_id or 'Proposal Section'
    bid_id = payload.get('bid_id')
    rfp_content = payload.get('rfp_content', '')

    if not section_id:
        return jsonify({'error': 'missing_params', 'message': 'Section ID is required.'}), 400

    if bid_id and not rfp_content:
        rfp_content = _fetch_rfp_content_for_bid(bid_id, page_limit=60, char_limit=12000)

    truncated_rfp = rfp_content[:12000] if rfp_content else ''

    system_prompt = (
        "You are an expert RFP and bid document analyst. "
        "Read RFP, IFB, and tender documents carefully and extract only the actionable requirements, eligibility criteria, "
        "and specifications needed for a compliant proposal submission. Present results in clear proposal-ready language."
    )

    user_prompt = f"""
Carefully read the following RFP content and produce an output that follows the exact template below. Use concise bullet points. Do not invent information that is not present.

Template to follow:

📘 **Project Overview**
- Project title, number, client/issuing agency, location
- Procurement officer/contact email
- Submission portal/platform
- Important bid dates and schedule (release, pre-bid, submission, opening, award)
- Estimated project value or budget (if mentioned)

📋 **Scope of Work / Technical Requirements**
- Summarize required deliverables, materials, installation, testing, safety, coordination, warranty, and quality standards.

🧾 **Eligibility & Qualification Requirements**
- Experience, licenses, certifications, references, safety, financial, insurance, staffing requirements, mandatory forms, bonds.

⚖️ **Bid Submission Requirements**
- Submission method, required documents, pricing forms, conflict forms, bonding instructions, clarification procedures.

📊 **Evaluation & Award Criteria**
- List all evaluation factors, scoring parameters, participation goals, best-value or low-bid rules.

💼 **Contract & Performance Requirements**
- Contract term, start/completion dates, liquidated damages, compliance obligations, minimum self-performed work.

🔒 **Bonds, Warranty & Insurance**
- Bid/performance/payment bonds with percentages, warranty coverage, insurance types and limits.

💡 **Added Value / Special Notes**
- Unique clauses, sustainability requirements, community programs, optional value-add proposals.

🗓️ **Key Dates Summary Table**
| Milestone | Date |
|------------|-------|
| IFB/RFP Release | |
| Pre-Bid Meeting | |
| Questions Deadline | |
| Bid Submission | |
| Bid Opening | |
| Board/Contract Award | |

✅ **Summary for Proposal Preparation**
- Summarize everything a bidder must ensure for compliance (eligibility, documents, bonds, pricing, schedule, goals).

Formatting rules:
- Use the headings exactly as shown.
- Use bullet points for lists; keep them short and factual.
- Skip any sub-bullet if the RFP does not mention it (do not fabricate content).
- End the response with the sentence: “These are the complete actionable requirements, eligibility criteria, and specifications extracted from this RFP for proposal preparation.”

RFP_CONTENT_START
{truncated_rfp if truncated_rfp else "No RFP content is available."}
RFP_CONTENT_END
"""

    api_key = (app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or '').strip()
    if not api_key:
        return jsonify({'error': 'api_unavailable', 'message': 'OpenAI API key is not configured on the server.'}), 500

    base_url = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or 'https://api.groq.com/openai/v1').strip()
    model = (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'llama-3.1-70b-versatile').strip()
    endpoint = base_url.rstrip('/')
    if not endpoint.endswith('/chat/completions'):
        endpoint = f"{endpoint}/chat/completions"

    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json',
    }
    body = {
        'model': model,
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt}
        ],
        'temperature': 0.2,
    }

    try:
        response = requests.post(endpoint, headers=headers, json=body, timeout=90)
    except requests.RequestException as err:
        app.logger.exception("OpenAI section outline network failure: %s", err)
        return jsonify({'error': 'api_unavailable', 'message': 'Could not reach the AI service. Please try again shortly.'}), 502

    try:
        data = response.json()
    except ValueError:
        app.logger.error("OpenAI section outline returned non-JSON response (status %s): %s", response.status_code, response.text[:400])
        return jsonify({'error': 'api_unavailable', 'message': 'Received an invalid response from the AI service.'}), 502

    if response.status_code >= 400:
        err_message = data.get('error') if isinstance(data, dict) else None
        detail = ''
        if isinstance(err_message, dict):
            detail = err_message.get('message') or err_message.get('code') or ''
        elif isinstance(err_message, str):
            detail = err_message
        if not detail:
            detail = f"Upstream error {response.status_code}"
        app.logger.error("OpenAI section outline error %s: %s", response.status_code, detail)
        return jsonify({'error': 'api_unavailable', 'message': f'AI service returned an error: {detail}'}), 502

    choices = data.get('choices') or []
    if not choices:
        return jsonify({'error': 'api_unavailable', 'message': 'No response returned by the AI service.'}), 502

    content = choices[0].get('message', {}).get('content', '')
    if not content:
        return jsonify({'error': 'api_unavailable', 'message': 'The AI service returned an empty response.'}), 502

    return jsonify({
        'output': content.strip(),
        'raw': content.strip(),
        'model': data.get('model', model),
        'usage': data.get('usage', {}),
    })

@app.route('/api/analyze-requirements-attachments', methods=['POST'])
@login_required
def analyze_requirements_attachments():
    """Extract requirements and required attachments from RFP document."""
    payload = request.get_json(silent=True) or {}
    bid_id = payload.get('bid_id')
    
    if not bid_id:
        return jsonify({'error': 'missing_params', 'message': 'Bid ID is required.'}), 400

    # Fetch RFP content
    rfp_content = _fetch_rfp_content_for_bid(bid_id, page_limit=80, char_limit=15000)
    
    if not rfp_content:
        return jsonify({'error': 'no_rfp', 'message': 'No RFP content found for this bid. Please upload an RFP file first.'}), 404

    truncated_rfp = rfp_content[:15000] if rfp_content else ''

    system_prompt = (
        "You are a professional RFP analysis assistant. "
        "Your only task is to analyze the provided RFP / IFB document and extract **only the mandatory attachments, forms, and documents** that bidders are required to submit with their bid or proposal. "
        "Do NOT include scope of work, eligibility, or background text. "
        "Do NOT infer or assume anything — extract only what is explicitly listed in the RFP as required attachments, forms, or bid documents."
    )

    user_prompt = f"""
Analyze the following RFP content and extract **only the mandatory attachments, forms, and documents** that bidders are required to submit.

### 🔹 OUTPUT FORMAT (strictly follow this structure)

📑 **Mandatory Attachments & Documents to Include with the Bid**

List each required document or form in clean bullet points.  
If attachment numbers or names are mentioned, list them exactly as written in the RFP.

Example output format:

- Attachment 1 – Contract Award Form  
- Attachment 2 – Acknowledgment Form  
- Attachment 3 – Bidder's Certifications  
- Attachment 4 – Conflict of Interest Questionnaire  
- Attachment 5 – Financial Interests and Potential Conflicts Form  
- Attachment 6 – References  
- Attachment 7 – Insurance Requirements  
- Attachment 8 – Small Business Development (SBD) Forms  
- Attachment 9 – Contractor Certification Sheet  
- Bid Bond (if required)  
- Payment Bond (if contract value > $25,000)  
- Performance Bond (if contract value > $100,000)  
- Company Profile / Cover Letter  
- Safety Record / EMR Report  
- Proof of Experience and References  
- Licenses / Certifications  
- Financial Stability Statement  
- Bid Pricing Sheet / Cost Proposal  
- Small Business Participation Documentation  
- Any other forms explicitly listed in the RFP

If no attachments are mentioned in the document, write:  
**"No specific attachments or mandatory bid documents listed in this RFP."**

### 🔹 RULES FOR EXTRACTION

- Include **only** attachment titles, form names, or document names mentioned in the RFP.  
- Remove all explanation text or surrounding sentences.  
- One line per attachment or required document.  
- Use plain, clean formatting — perfect for UI display next to an "Attach File" button.  
- Do not generate or assume missing data.

### 🔹 GOAL

Output must contain **only the explicit list of required attachments and documents** found in the RFP.  
No commentary. No assumptions. No descriptions.

End your response with:  
**"These are the mandatory attachments and documents explicitly required for bid submission in this RFP."**

RFP_CONTENT_START
{truncated_rfp if truncated_rfp else "No RFP content is available."}
RFP_CONTENT_END
"""

    api_key = (app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or '').strip()
    if not api_key:
        return jsonify({'error': 'api_unavailable', 'message': 'OpenAI API key is not configured on the server.'}), 500

    base_url = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or 'https://api.groq.com/openai/v1').strip()
    model = (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'llama-3.1-70b-versatile').strip()
    endpoint = base_url.rstrip('/')
    if not endpoint.endswith('/chat/completions'):
        endpoint = f"{endpoint}/chat/completions"

    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json',
    }
    body = {
        'model': model,
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt}
        ],
        'temperature': 0.2,
    }

    try:
        response = requests.post(endpoint, headers=headers, json=body, timeout=90)
    except requests.RequestException as err:
        app.logger.exception("OpenAI requirements analysis network failure: %s", err)
        return jsonify({'error': 'api_unavailable', 'message': 'Could not reach the AI service. Please try again shortly.'}), 502

    try:
        data = response.json()
    except ValueError:
        app.logger.error("OpenAI requirements analysis returned non-JSON response (status %s): %s", response.status_code, response.text[:400])
        return jsonify({'error': 'api_unavailable', 'message': 'Received an invalid response from the AI service.'}), 502

    if response.status_code >= 400:
        err_message = data.get('error') if isinstance(data, dict) else None
        detail = ''
        if isinstance(err_message, dict):
            detail = err_message.get('message') or err_message.get('code') or ''
        elif isinstance(err_message, str):
            detail = err_message
        if not detail:
            detail = f"Upstream error {response.status_code}"
        app.logger.error("OpenAI requirements analysis error %s: %s", response.status_code, detail)
        return jsonify({'error': 'api_unavailable', 'message': f'AI service returned an error: {detail}'}), 502

    choices = data.get('choices') or []
    if not choices:
        return jsonify({'error': 'api_unavailable', 'message': 'No response returned by the AI service.'}), 502

    content = choices[0].get('message', {}).get('content', '')
    if not content:
        return jsonify({'error': 'api_unavailable', 'message': 'The AI service returned an empty response.'}), 502

    return jsonify({
        'output': content.strip(),
        'requirements': content.strip(),
        'raw': content.strip(),
        'model': data.get('model', model),
        'usage': data.get('usage', {}),
    })

@app.route('/api/generate-section-content', methods=['POST'])
@login_required
def generate_section_content():
    """Generate content for a specific proposal section, especially Letter of Transmittal."""
    payload = request.get_json(silent=True) or {}
    section_id = (payload.get('section_id') or '').strip()
    section_title = (payload.get('section_title') or '').strip()
    bid_id = payload.get('bid_id')
    company = (payload.get('company') or '').strip()
    contact_email = (payload.get('contact_email') or '').strip()
    
    if not section_id:
        return jsonify({'error': 'missing_params', 'message': 'Section ID is required.'}), 400
    
    # Get RFP content from database if bid_id is provided
    rfp_content = ''
    if bid_id:
        # Use higher limits for sections that need thorough analysis
        rfp_content = _fetch_rfp_content_for_bid(bid_id, page_limit=100, char_limit=25000)
    
    # Get bid information for context
    bid_info = {}
    if bid_id:
        try:
            cur = mysql.connection.cursor(DictCursor)
            cur.execute("SELECT * FROM go_bids WHERE g_id=%s", (bid_id,))
            bid_info = cur.fetchone() or {}
            cur.close()
        except Exception as e:
            app.logger.warning(f"Error loading bid info: {e}")
    
    # Extract relevant bid information
    project_name = bid_info.get('b_name') or bid_info.get('project_name') or ''
    client_name = bid_info.get('agency') or bid_info.get('customer') or bid_info.get('client') or ''
    due_date = bid_info.get('due_date') or ''
    
    # Format due date
    if due_date:
        try:
            from datetime import datetime, date
            if isinstance(due_date, (datetime, date)):
                due_date_str = due_date.strftime('%B %d, %Y')
            else:
                parsed = datetime.strptime(str(due_date), '%Y-%m-%d')
                due_date_str = parsed.strftime('%B %d, %Y')
        except Exception:
            due_date_str = str(due_date)
    else:
        due_date_str = ''
    
    # ================= Two-pass Planner -> Writer (generic path for all sections) =================
    try:
        # Load company knowledge base if available
        company_texts = []
        if company:
            try:
                db_path = os.path.join('company_db', f"{company}.json")
                if os.path.exists(db_path):
                    with open(db_path, 'r', encoding='utf-8') as f:
                        db = json.load(f)
                        company_texts = db.get('texts', []) or []
            except Exception:
                company_texts = []
        
        # Retrieve focused company evidence for this section
        retrieval_query = section_title or "proposal section"
        company_evidence = scorer.retrieve_relevant_context(company_texts, retrieval_query) if company_texts else ""
        rfp_ctx = trim_text_to_token_limit(rfp_content or "", 6000)
        combined_ctx = (f"RFP Context:\n{rfp_ctx}\n\nCompany Evidence:\n{company_evidence}").strip()
        
        # Build planner prompt (JSON)
        planner_system = "You are a senior proposal planner. Return ONLY strict JSON, no markdown."
        planner_user = f"""
Plan a detailed outline for the section "{section_title}" of a proposal.
Use ONLY the provided context. If information is missing, set placeholders.
Return JSON with fields: title, objectives[], key_points[], evidence_needed[], tone, length_words, structure[{{heading, bullets[]}}].

Context:
{combined_ctx}
"""
        api_key = (app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or os.getenv('GROQ_API_KEY') or '').strip()
        if not api_key:
            return jsonify({'error': 'api_unavailable', 'message': 'AI API key not configured.'}), 500
        base_url = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or 'https://api.openai.com/v1').strip()
        model = (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'gpt-4o-mini').strip()
        endpoint = base_url.rstrip('/')
        if not endpoint.endswith('/chat/completions'):
            endpoint = f"{endpoint}/chat/completions"
        headers = {'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'}
        plan_body = {
            'model': model,
            'messages': [
                {'role': 'system', 'content': planner_system},
                {'role': 'user', 'content': planner_user},
            ],
            'temperature': 0.1,
        }
        plan_resp = requests.post(endpoint, headers=headers, json=plan_body, timeout=90)
        plan_data = plan_resp.json()
        if plan_resp.status_code >= 400:
            raise RuntimeError(plan_data.get('error', 'planner_error'))
        plan_text = (plan_data.get('choices') or [{}])[0].get('message', {}).get('content', '{}')
        try:
            plan = json.loads(plan_text)
        except Exception:
            # Try to extract JSON
            m = re.search(r'(\{.*\})', plan_text, flags=re.DOTALL)
            plan = json.loads(m.group(1)) if m else {}
        
        # Writer pass (uses the plan)
        writer_system = "You are an expert proposal writer. Produce clean, professional prose."
        writer_user = f"""
Write the full section "{section_title}" following this plan and using ONLY the context.
If evidence is insufficient, include concise 'Assumption – verify' notes.

Plan:
{json.dumps(plan, ensure_ascii=False)}

Context:
{combined_ctx}
"""
        write_body = {
            'model': model,
            'messages': [
                {'role': 'system', 'content': writer_system},
                {'role': 'user', 'content': writer_user},
            ],
            'temperature': 0.2,
        }
        write_resp = requests.post(endpoint, headers=headers, json=write_body, timeout=120)
        write_data = write_resp.json()
        if write_resp.status_code >= 400:
            raise RuntimeError(write_data.get('error', 'writer_error'))
        final_content = (write_data.get('choices') or [{}])[0].get('message', {}).get('content', '')
        if final_content:
            return jsonify({
                'output': final_content.strip(),
                'plan': plan,
                'model': write_data.get('model', model),
                'usage': write_data.get('usage', {}),
            })
    except Exception as _twopass_err:
        # Fall back to legacy per-section prompts below
        pass

    # Check if this is Letter of Transmittal section
    is_letter_of_transmittal = (
        section_id == 'section-cover' or 
        'letter of transmittal' in (section_title or '').lower() or
        'transmittal' in (section_title or '').lower()
    )
    
    # Check if this is Executive Summary section
    is_executive_summary = (
        section_id == 'section-executive-summary' or
        'executive summary' in (section_title or '').lower()
    )
    
    # Check if this is Requirements Compliance section
    is_requirements_compliance = (
        section_id == 'section-compliance' or
        'requirements compliance' in (section_title or '').lower() or
        'compliance' in (section_title or '').lower() and 'requirements' in (section_title or '').lower()
    )
    
    # Check if this is Understanding of Scope & Objectives section
    is_scope_objectives = (
        section_id == 'section-scope-objectives' or
        'understanding of scope' in (section_title or '').lower() or
        'scope & objectives' in (section_title or '').lower() or
        ('scope' in (section_title or '').lower() and 'objectives' in (section_title or '').lower())
    )
    
    # Check if this is Deviations, Assumptions, and Dependencies section
    is_deviations_assumptions = (
        section_id == 'section-deviations' or
        'deviations' in (section_title or '').lower() or
        ('deviations' in (section_title or '').lower() and 'assumptions' in (section_title or '').lower() and 'dependencies' in (section_title or '').lower())
    )
    
    # Check if this is Proposed Technical Solution section
    is_proposed_technical_solution = (
        section_id == 'section-technical' or
        'proposed technical solution' in (section_title or '').lower() or
        ('technical solution' in (section_title or '').lower() and 'proposed' in (section_title or '').lower())
    )
    
    # Check if this is Implementation & Work Plan section
    is_implementation_work_plan = (
        section_id == 'section-implementation' or
        'implementation & work plan' in (section_title or '').lower() or
        'implementation and work plan' in (section_title or '').lower() or
        ('implementation' in (section_title or '').lower() and 'work plan' in (section_title or '').lower())
    )
    
    # Check if this is Solution Deliverables section
    is_solution_deliverables = (
        section_id == 'section-deliverables' or
        'solution deliverables' in (section_title or '').lower() or
        ('solution' in (section_title or '').lower() and 'deliverables' in (section_title or '').lower())
    )
    
    # Check if this is Project Management Plan section
    is_project_management_plan = (
        section_id == 'section-management' or
        'project management plan' in (section_title or '').lower() or
        ('project management' in (section_title or '').lower() and 'plan' in (section_title or '').lower())
    )
    
    # Check if this is Pricing Proposals section
    is_pricing_proposals = (
        section_id == 'section-pricing' or
        'pricing proposals' in (section_title or '').lower() or
        ('pricing' in (section_title or '').lower() and 'proposals' in (section_title or '').lower())
    )
    
    # Check if this is Contractual Terms section
    is_contractual_terms = (
        section_id == 'section-contractual' or
        'contractual terms' in (section_title or '').lower() or
        ('contractual' in (section_title or '').lower() and 'terms' in (section_title or '').lower())
    )
    
    # Check if this is Corporate Qualifications section
    is_corporate_qualifications = (
        section_id == 'section-corporate-qualifications' or
        section_id == 'section-corporate' or
        'corporate qualifications' in (section_title or '').lower() or
        ('corporate' in (section_title or '').lower() and 'qualifications' in (section_title or '').lower())
    )
    
    # Check if this is Project Schedule & Milestones section
    is_project_schedule_milestones = (
        section_id == 'section-schedule-milestones' or
        section_id == 'section-schedule' or
        ('project schedule' in (section_title or '').lower() and 'milestones' in (section_title or '').lower()) or
        ('schedule' in (section_title or '').lower() and 'milestone' in (section_title or '').lower())
    )
    
    # Check if this is Project Team section
    is_project_team = (
        section_id == 'section-project-team' or
        section_id == 'section-team' or
        'project team' in (section_title or '').lower() or
        ('team' in (section_title or '').lower() and 'project' in (section_title or '').lower())
    )
    
    # Check if this is Addenda section
    is_addenda = (
        section_id == 'section-amendments' or
        'addenda' in (section_title or '').lower() or
        'amendments' in (section_title or '').lower()
    )
    
    # Check if this is Appendix section
    is_appendix = (
        section_id == 'section-appendix' or
        'appendix' in (section_title or '').lower()
    )
    
    if is_letter_of_transmittal:
        # Use the specific prompt for Letter of Transmittal
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to draft a formal, professional Letter of Transmittal that meets all standard business letter requirements. "
            "The letter must be written from our company to the client, be technical in nature, and include all necessary points."
        )
        
        user_prompt = f"""Draft a formal, one-page Letter of Transmittal on our company letterhead, to be signed by an authorized executive. The letter must be in a standard business-letter format with clear paragraphs. It must:

1. State our formal intent to bid for this specific RFP.

2. Acknowledge receipt and review of all addenda.

3. Provide a concise statement of compliance with all mandatory requirements.

4. Briefly highlight our single most compelling technical differentiator.

5. Identify our single point of contact for this proposal.

6. Conclude by affirming our enthusiasm for the partnership.

This must be from our company to the client. It should be technical and should include all necessary points.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (for context):
{rfp_content[:15000] if rfp_content else 'No RFP content available. Generate a professional Letter of Transmittal template.'}

Please generate a complete, professional Letter of Transmittal that addresses all six requirements listed above. The letter should be technical, comprehensive, and written from our company to the client. Format it as a standard business letter with proper paragraphs. Return the content as a single formatted text that can be split into paragraphs."""
    elif is_executive_summary:
        # Use the specific prompt for Executive Summary
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to draft a persuasive, one-page Executive Summary that is standalone and written for a high-level, non-technical decision-maker."
        )
        
        user_prompt = f"""With reference to the requirement mentioned in the RFP/RFI/RFQ/BID document, generate a persuasive, one-page Executive Summary. This summary must be standalone and written for a high-level, non-technical/technical decision-maker. It should be from our company to the client. Use the following for structure:

• In a single paragraph, summarize our understanding of the client's core problem, goals, and desired outcomes as stated in the bid.

• In two paragraphs, provide a high-level, benefit-focused overview of our technical and management solution. Avoid deep jargon.

• Highlight 3-5 specific, unique advantages our solution offers (e.g., proprietary technology, faster deployment, specific expertise).

• Mention the tangible, measurable results the client will receive.

• State the total, all-inclusive price in a clear, and concise manner.

• This should not have any subheadings and should be in paragraph only.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (for context):
{rfp_content[:15000] if rfp_content else 'No RFP content available. Generate a professional Executive Summary template.'}

Please generate a complete Executive Summary written entirely in paragraph format with NO subheadings. The content should follow the structure outlined above but present everything as continuous paragraphs without any headings, bullet points, or lists. All content should flow as continuous, well-structured paragraphs. The content should be persuasive, clear, and suitable for high-level decision-makers. Return the content as a single flowing narrative that can be split into paragraphs."""
    elif is_requirements_compliance:
        # Use the specific prompt for Requirements Compliance
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to draft a formal Statement of Compliance that thoroughly addresses all requirements and key activities from the RFP document."
        )
        
        user_prompt = f"""As per the requirement mentioned in the document, draft a formal 'Statement of Compliance.' State that we have read, understood, and will comply with all mandatory requirements, terms, and conditions set forth in the bid document. Include all key points in detail and in a technical manner. Do not skip any important points to be discussed.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify ALL requirements, key activities, deliverables, attachments, forms, technical specifications, and compliance obligations):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Statement of Compliance template.'}

Please generate the Statement of Compliance as exactly TWO paragraphs only, no headings, lists, bullets, or numbering:
- Paragraph 1: Provide the formal compliance assertion that we have read, understood, and will comply with all mandatory requirements, terms, and conditions of the bid.
- Paragraph 2: Note that any exceptions will be listed in the 'Deviations' section and summarize, in technical language, our comprehensive adherence to the RFP’s requirements, deliverables, forms, technical specifications, certifications, and compliance obligations.

Strictly limit the output to two paragraphs with complete sentences and professional tone."""
    elif is_scope_objectives:
        # Use the specific prompt for Understanding of Scope & Objectives
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to draft a comprehensive 'Understanding of Scope & Objectives' section that demonstrates deep subject-matter expertise by adding value and insight, not just repeating the RFP's text."
        )
        
        user_prompt = f"""As per the requirement mentioned in the document, draft a section titled 'Understanding of Scope & Objectives.' Using narrative paragraphs, analyse and re-state the client's core challenges, operational problems, and strategic goals. This section must demonstrate our deep subject-matter expertise by adding value and insight, not just repeating the RFP's text. Show that we understand the 'why' behind the 'what.' Include all key points in technical manner and in detail. Do not skip any important points to be discussed. Use pointers and table if necessary. Make it completely informative and elaborative.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to understand scope, objectives, challenges, phases, key activities, technical requirements, and all important points):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Understanding of Scope & Objectives template.'}

Please generate a comprehensive 'Understanding of Scope & Objectives' section that:
- Uses narrative paragraphs to demonstrate deep understanding, supplemented with pointers and tables where necessary
- Analyzes and re-states the client's core challenges, operational problems, and strategic goals in technical detail
- Adds value and insight beyond just repeating the RFP text - demonstrates subject-matter expertise
- Shows understanding of the 'why' behind the 'what' - explains the reasoning, context, strategic importance, and technical implications
- Includes ALL key points in technical manner and in detail - do not skip any important points to be discussed
- Uses pointers (bullet points) and tables where appropriate to organize information clearly and make it completely informative
- Makes it completely informative and elaborative - comprehensive coverage of all aspects with technical precision

Return the content as well-structured narrative paragraphs with appropriate use of pointers and tables to organize key activities, phases, technical requirements, and all important points. The content must be completely informative and elaborative, technically detailed, and comprehensive without omitting any important information."""
    elif is_deviations_assumptions:
        # Use the specific prompt for Deviations, Assumptions, and Dependencies
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a section that manages expectations by clearly documenting deviations, assumptions, and dependencies."
        )
        
        user_prompt = f"""As per mentioned in the document (if any), generate a section to manage expectations (if any). Use the following

• Use a numbered list to detail every item marked as 'Partial' or 'Exception'. Provide a clear justification for each.

• Use a numbered list to state all assumptions our technical solution and pricing are based on (e.g., 'We assume client will provide network access...').

• Use a numbered list to explicitly state all items, data, access, or resources we require from the client to ensure project success."

It should be completely technical and should include all key points. Do not skip any key points mentioned in RFP.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify ALL deviations, assumptions, dependencies, technical requirements, and key points):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Deviations, Assumptions, and Dependencies template.'}

Please generate a comprehensive, completely technical section that:
- Uses numbered lists for: Deviations and Exceptions; Technical & Operational Assumptions; Client-Furnished Dependencies
- Details every item marked as 'Partial' or 'Exception' in the compliance matrix (if applicable) with technical justifications
- States ALL assumptions our technical solution and pricing are based on
- Explicitly states ALL client-furnished items, data, access, resources, and dependencies required
- Includes ALL key points mentioned in the RFP - do not skip any important technical details, requirements, or specifications
- Is completely technical, specific, and actionable

Return the content with clearly separated parts and numbered lists as specified. The content must be completely technical, comprehensive, and include all key points from the RFP without omitting any important information."""
    elif is_proposed_technical_solution:
        # Use the specific prompt for Proposed Technical Solution
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive 'Proposed Technical Solution' section that addresses all technical requirements and work-streams from the RFP document."
        )
        
        user_prompt = f"""As per the requirement in the document, generate the core 'Proposed Technical Solution' section. Structure these using headings and subheadings that directly correspond to the technical work-streams defined in the bid document.

• Before generating, ensure the user has uploaded/attached the RFP/SoW/addenda via the Attach File button. Incorporate all details/information from the attached documents and any provided data inputs.

• For each subsection, use technical paragraphs, bulleted lists for specifications, and tables for technical data.

• This section must be highly detailed, addressing how our solution fulfils each technical requirement. Include specific product models, software versions, and methodologies where relevant.

• Include all key activities/technical requirements. Do not skip any important points to be discussed. Use pointers and table if necessary.

Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify ALL technical work-streams, requirements, specifications, deliverables, phases, key activities, and technical details):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Proposed Technical Solution template.'}

Please generate a comprehensive, highly detailed, and completely technical 'Proposed Technical Solution' section that:
- Uses headings and subheadings that directly correspond to technical work-streams from the RFP
- For each subsection, uses technical paragraphs, bulleted lists (pointers) for specifications, and tables for technical data
- Is highly detailed and addresses how our solution fulfils each technical requirement with technical precision
- Includes specific product models, software versions, hardware specifications, methodologies, standards, and protocols where relevant
- Includes ALL key activities/technical requirements, specifications, and important points - does not skip any important points
- Uses pointers (bullet points) and tables where required to organize and present technical information clearly
- Is completely technical and comprehensive
- Is aligned with the requirements in RFP - addresses every technical requirement, specification, and work-stream
- Makes it in very detail by including all necessary points
- Includes all necessary technical details: configurations, specifications, methodologies, protocols, standards, tools, technologies, and implementation approaches

Return the content as well-structured, highly detailed, and completely technical content with appropriate headings and subheadings, paragraphs, bulleted lists (pointers), and tables. The response must be aligned with all requirements in the RFP and include all necessary points without omitting any important information."""
    elif is_implementation_work_plan:
        # Use the specific prompt for Implementation & Work Plan
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive 'Implementation & Work Plan' section that details the methodology, project phases, and key activities for successful project delivery."
        )
        
        user_prompt = f"""Generate a comprehensive 'Implementation & Work Plan' section with two main subsections:

1. Methodology

As per the requirement in the document, draft methodology section based on the project scope. In narrative paragraphs, describe the specific project management and technical implementation methodology we will use. Include all key activities. Do not skip any important points to be discussed. Use pointers and table if necessary. Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP. Use headings, subheadings, pointers and tables where required.

2. Project Phases & Activities

As per the requirement in the document, "daft a 'Project Phases & Activities' subsection based on the project scope. Use subheadings for each distinct project phase. Use weeks for each phase in the form Week1-Week2 etc. These weeks should be practically feasible as per the requirement in the project document. Includer all key activities in each phase. Do not skip any important points to be discussed. Use pointers and tables where necessary. Make it in very detail by including all necessary points. Also, make sure this is completely technical and feasible. Include pointers and table where required. The response should be aligned with the requirement in the project.

Include all key activities in each phase. Do not skip any important points to be discussed. Use pointers and tables where necessary. Make it in very detail by including all necessary points. Also, make sure this is completely technical and feasible. Include pointers and table where required. The response should be aligned with the requirement in RFP.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify ALL project phases, activities, tasks, objectives, methodology requirements, technical implementation details, and all important points):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Implementation & Work Plan template.'}

Please generate a comprehensive 'Implementation & Work Plan' section that includes:

1. Methodology (use headings/subheadings as appropriate):
   - Narrative paragraphs describing the specific project management and technical implementation methodology in very detail
   - Justification for why this methodology is the best fit for this specific project and client
   - ALL key activities related to methodology included - no important points skipped
   - Use of pointers (bullet points) and tables where required to organize and present technical information clearly
   - Very detailed coverage including all necessary points
   - Completely technical content
   - Response aligned with the requirements in RFP

2. Project Phases & Activities (use subheadings):
   - Subheadings for each distinct project phase (e.g., 'Phase 1: Discovery & Design,' 'Phase 2: Build & Test,' 'Phase 3: Deployment & Go-Live')
   - For each phase, include a practical week range label (e.g., 'Week 1–Week 2', 'Week 3–Week 6') aligned with the bid timeline
   - Under each phase, bulleted lists detailing key activities, tasks, and objectives in very detail
   - ALL key activities in each phase included - no important points skipped
   - Use of pointers (bullet points) and tables where required to organize and present technical information clearly
   - Very detailed coverage including all necessary points
   - Completely technical content
   - Response aligned with the requirements in RFP

Return the content as well-structured, highly detailed, and completely technical content with appropriate headings and subheadings, narrative paragraphs, bulleted lists (pointers), and tables. Both the Methodology and Project Phases & Activities subsections must be in very detail, completely technical, include all key activities, use pointers and tables where required, and be aligned with all requirements in the RFP."""
    elif is_solution_deliverables:
        # Use the specific prompt for Solution Deliverables
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive 'Solution Deliverables' section that lists all tangible items the client will receive in a structured table format."
        )
        
        user_prompt = f"""As per the requirement in the project information document, generate a 'Solution Deliverables' section based on the scope of the work. Create a table that lists all tangible items the client will receive. The table columns must be:

1. Deliverable Title: must include what will be delivered.

2. Description: (A brief explanation).

3. Format: (e.g., 'PDF Document,' 'Online Dashboard').

4. Delivery Milestone: (When it will be delivered, e.g., 'End of Phase 1').

Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify ALL deliverables, their descriptions, formats, delivery milestones, technical specifications, and all important points):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Solution Deliverables template.'}

Please generate a comprehensive, highly detailed, and completely technical 'Solution Deliverables' section that:
- Creates a table with exactly 4 columns: Deliverable Title, Description, Format, Delivery Milestone
- Lists ALL tangible items the client will receive with technical precision
- Includes ALL deliverables from the RFP - does not skip any items or important points
- Provides very specific and detailed technical descriptions for each deliverable
- Matches delivery milestones to project phases from the RFP
- Uses appropriate technical formats for each deliverable
- Is well-structured, comprehensive, and includes all technical details
- Is aligned with the requirements in RFP - addresses every deliverable requirement, specification, and item
- Includes pointers (bullet points) and tables where required to organize and present technical information clearly
- Makes it in very detail by including all necessary points
- Is completely technical

Return the content as a well-structured, highly detailed, and completely technical table with all deliverables listed. The table format must be clear, comprehensive, include all required columns, and be aligned with all requirements in the RFP. Include pointers and tables where required to make it completely informative and technical."""
    elif is_project_management_plan:
        # Use the specific prompt for Project Management Plan
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive 'Project Management Plan' section that includes project governance, communication plan, risk management, and quality assurance."
        )
        
        user_prompt = f"""As per the requirement in the bid document, generate a comprehensive, highly detailed, and completely technical 'Project Management Plan' with four main subsections:

1. Project Governance

Prompt: As per the requirement in the project document, draft a 'Project Governance' subsection based on the scope of the project. In paragraphs, describe the management structure involved in the project. Include all key activities in each phase. Do not skip any important points to be discussed. Use pointers and table if necessary. Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

2. Communication Plan

Prompt: As per the requirement in the project document, generate a 'Communication Plan as per the scope of the work. Create a table that defines the communication protocols. Includer all key activities. Do not skip any important points to be discussed. Use pointers and table if necessary. Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

3. Risk Management Plan

Prompt: As per the requirement in the project document, generate a 'Risk Management Plan based on the scope of the work. Create a table. Do not skip any important points to be discussed. Use pointers and table if necessary. Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

4. Quality Assurance Plan

Prompt: A per the requirement in the bid document, draft a 'Quality Assurance Plan based on the scope of work. Includer all key activities. Do not skip any important points to be discussed. Use pointers and table if necessary. Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

Include all key activities. Do not skip any important points to be discussed. Use pointers and tables if necessary.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify all project management requirements, communication protocols, risks, quality standards, and key activities):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Project Management Plan template.'}

Please generate a comprehensive, highly detailed, and completely technical 'Project Management Plan' section. Use appropriate H3 subheadings for each subsection and ensure alignment with the RFP. Include all key activities and do not skip any important points. Use pointers and tables where required."""
    elif is_pricing_proposals:
        # Use the specific prompt for Pricing Proposals
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive, highly detailed, and completely technical 'Pricing Proposals' section aligned with RFP requirements."
        )
        
        user_prompt = f"""Generate a comprehensive, highly detailed, and completely technical 'Pricing Proposals' section with three subsections:

1. Pricing Summary

Generate a 'Pricing Summary.' Create a simple table that shows the main cost components and a 'Total All-Inclusive Price.' The structure should be clear and easy to read.

2. Detailed Cost Breakdown

Generate a 'Detailed Cost Breakdown.' You must use the exact pricing table format and line items requested in the bid document. If no specific format is provided, create a table that breaks down all costs by labor, materials, software, and other expenses. Columns should include: Cost Element, Unit (e.g., Hours, Each), Unit Cost, Quantity, and Total Line Item Cost.

3. Payment Schedule

Generate a 'Payment Schedule.' Propose a payment plan based on project milestones. Create a table with columns: Milestone (linked to the project schedule) and Payment Amount / Percentage.

Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and tables where required. The response should be aligned with the requirement in RFP.

IMPORTANT REQUIREMENTS:
1. Analyze the RFP thoroughly to identify all pricing requirements, cost breakdown formats, payment terms, and pricing assumptions
2. For the Pricing Summary subsection:
   - Create a simple, clear table showing main cost components
   - Include a 'Total All-Inclusive Price' row
   - Make the structure clear and easy to read
   - Include all major cost categories from the RFP
   - Make it in very detail by including all necessary points
   - Make sure this is completely technical
   - Include pointers and tables where required
   - The response should be aligned with the requirement in RFP
3. For the Detailed Cost Breakdown subsection:
   - Use the exact pricing table format and line items requested in the bid document if specified
   - If no specific format is provided, create a table with columns: Cost Element, Unit (e.g., Hours, Each), Unit Cost, Quantity, Total Line Item Cost
   - Break down all costs by labor, materials, software, and other expenses
   - Include all cost elements mentioned in the RFP - do not skip any important cost elements
   - Ensure accuracy and completeness
   - Make it in very detail by including all necessary points
   - Make sure this is completely technical
   - Include pointers and tables where required
   - The response should be aligned with the requirement in RFP
4. For the Payment Schedule subsection:
   - Create a table with columns: Milestone (linked to the project schedule) and Payment Amount / Percentage
   - Propose a payment plan based on project milestones
   - Link milestones to the project schedule from the RFP
   - Include all relevant payment milestones - do not skip any important milestones
   - Make it in very detail by including all necessary points
   - Make sure this is completely technical
   - Include pointers and tables where required
   - The response should be aligned with the requirement in RFP
5. Ensure comprehensive coverage of all pricing aspects from the RFP

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify all pricing requirements, cost breakdown formats, payment terms, milestones, and pricing assumptions):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Pricing Proposals template.'}

Please generate a comprehensive, highly detailed, and completely technical 'Pricing Proposals' section that includes:

1. Pricing Summary (H3 subheading):
   - Simple table showing main cost components
   - 'Total All-Inclusive Price' row
   - Clear and easy to read structure
   - All major cost categories included - do not skip any important cost categories
   - Make it in very detail by including all necessary points
   - Make sure this is completely technical
   - Include pointers and tables where required
   - The response should be aligned with the requirement in RFP

2. Detailed Cost Breakdown (H3 subheading):
   - Use exact pricing table format from bid document if specified
   - If no specific format, create table with columns: Cost Element, Unit (e.g., Hours, Each), Unit Cost, Quantity, Total Line Item Cost
   - Breakdown by labor, materials, software, and other expenses
   - All cost elements from RFP included - do not skip any important cost elements
   - Make it in very detail by including all necessary points
   - Make sure this is completely technical
   - Include pointers and tables where required
   - The response should be aligned with the requirement in RFP

3. Payment Schedule (H3 subheading):
   - Table with columns: Milestone (linked to project schedule) and Payment Amount / Percentage
   - Payment plan based on project milestones
   - Milestones linked to project schedule from RFP
   - All relevant payment milestones included - do not skip any important milestones
   - Make it in very detail by including all necessary points
   - Make sure this is completely technical
   - Include pointers and tables where required
   - The response should be aligned with the requirement in RFP

Structure the content with:
- H3 subheading: "Pricing Summary" followed by a simple table with main cost components and Total All-Inclusive Price
- H3 subheading: "Detailed Cost Breakdown" followed by a table with cost elements (using exact format from RFP if specified, or standard format with: Cost Element, Unit, Unit Cost, Quantity, Total Line Item Cost)
- H3 subheading: "Payment Schedule" followed by a table with columns: Milestone and Payment Amount / Percentage
- Clear, well-structured tables that are easy to read
- Comprehensive coverage of all pricing aspects from the RFP
- All content must be highly detailed, completely technical, and aligned with RFP requirements
- Appropriate use of pointers (bullet points) and tables to organize information clearly and comprehensively

Return the content as well-structured, highly detailed, and completely technical content with appropriate H3 subheadings, tables, and numbered lists. Ensure all key pricing elements are included, no important points are skipped, and the response is fully aligned with the requirements in the RFP."""
    elif is_contractual_terms:
        # Use the specific prompt for Contractual Terms
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive, highly detailed, and completely technical 'Acceptance of Terms and Conditions' section aligned with RFP requirements."
        )
        
        user_prompt = f"""As per the requirement in the bid document, generate an 'Acceptance of Terms and Conditions' section.

Prompt: "Analyze the contract and Terms & Conditions (T&Cs) from the bid document."

If we accept all terms: Draft a single paragraph stating: 'We have reviewed all terms and conditions outlined in the bid document and confirm our full acceptance and compliance without exception.'

If we have exceptions: Draft a paragraph stating we accept with exceptions, and then create a table with columns: Clause Number, Clause Title, Proposed Redline / Exception, and Business Justification.

Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify all contract terms and conditions, including indemnity, liability, payment terms, and any other contractual clauses):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Contractual Terms template.'}

Please generate a comprehensive, highly detailed, and completely technical 'Acceptance of Terms and Conditions' section. Use a clear heading and:
- If accepting all terms: include the single acceptance paragraph with the exact wording provided above.
- If there are exceptions: include the introductory paragraph and a table with 4 columns: Clause Number, Clause Title, Proposed Redline / Exception, Business Justification.
Ensure the language is professional, all key contractual items from the RFP are addressed, and use pointers and tables where required. The response must be fully aligned with the RFP."""
    elif is_corporate_qualifications:
        # Use the specific prompt for Corporate Qualifications
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive 'Corporate Qualifications' section that includes Company Overview, Relevant Past Performance, and Client References aligned with the RFP."
        )
        
        user_prompt = f"""PROMPT FOR Corporate Qualifications

o Company Overview

Prompt: "Draft a 'Company Overview.' In 2-3 concise paragraphs, describe our company's history, size, locations, and core competencies that are directly relevant to this project." Use technical statements.

o Relevant Past Performance

Prompt: As per the requirement in the bid document, generate a 'Relevant Past Performance' section. Under each client, create a table with rows for:

• Project Title:
• Period of Performance:
• Project Value:
• Challenge: (A brief description of the problem).
• Solution & Outcome: (A description of our work and the positive, quantifiable result)."

o Client References

Prompt: "Generate a 'Client References' section. Create a table providing 3-5 client references, corresponding to the projects listed in 'Past Performance.' Columns must be: Client Name, Contact Person, Title, Phone, and Email."

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content:
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Corporate Qualifications template.'}"""
    elif is_project_schedule_milestones:
        # Use the specific prompt for Project Schedule & Milestones
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive, highly detailed, and completely technical 'Project Schedule & Milestones' section that is complete, technical, and feasible practically, aligned with RFP requirements."
        )
        
        user_prompt = f"""As per the requirement in the bid document, generate a 'Project Schedule & Milestones' section. This should be complete technical and feasible practically. Use tables and pointers where necessary.

1. First, insert a placeholder for a visual Gantt chart (e.g., [Insert Gantt Chart Visual Here]).

2. Following the placeholder, create a table that lists the 'Critical Project Milestones.' Columns should be: Milestone, Description, and Target Completion Date (or 'Weeks from Kick-off').

Make it in very detail by including all necessary points. Also, make sure this is completely technical and feasible. Include pointers and table where required. The response should be aligned with the requirement in RFP.

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content (analyze thoroughly to identify all project phases, milestones, timelines, delivery dates, and schedule requirements):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Project Schedule & Milestones template.'}
"""
    elif is_project_team:
        # Use the specific prompt for Project Team
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive, highly detailed, and completely technical 'Project Team' section that includes organizational chart, roles & responsibilities, and key personnel information, aligned with RFP requirements."
        )
        
        user_prompt = f"""PROMPT FOR Project Team

o Project Team Organizational Chart

Prompt: "Insert a placeholder for a 'Project Team Organizational Chart.' This visual diagram must show the proposed team structure, reporting lines, and key interfaces with the client's team." 

o Roles & Responsibilities

Prompt: "Generate a 'Roles & Responsibilities' section. Create a table with columns: Project Role, Proposed Name, and Key Responsibilities." Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.

o Key Personnel Resumes

Prompt: insert a placeholder for key resumes or small bios

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

RFP Content:
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Project Team template.'}"""
    elif is_addenda:
        # Use the specific prompt for Addenda
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate an acknowledgment statement for the Addenda section that acknowledges all addenda received with their names and dates."
        )
        
        user_prompt = f"""PROMPT FOR  Contractual Terms:
 
 o Acceptance of Terms and Conditions
 
 Prompt: "Analyse the contract and Terms & Conditions (T&Cs) from the bid document.
 
 If we accept all terms: Draft a single paragraph stating that 'We have reviewed all terms and conditions outlined in the bid document and confirm our full acceptance and compliance without exception.' This should be technical in nature based on the scope of the work.
 
 If we have exceptions: Draft a paragraph stating we accept with exceptions, and then create a table with columns: Clause Number, Clause Title, Proposed Redline / Exception, and Business Justification." Make it in very detail by including all necessary points. Also, make sure this is completely technical. Include pointers and table where required. The response should be aligned with the requirement in RFP.
 
 Company Information:
 - Company Name: {company or 'Our Company'}
 - Primary Contact: {contact_email or 'Contact information to be provided'}
 
 RFP Information:
 - Project/Solicitation Title: {project_name or 'This RFP'}
 - Client/Issuing Agency: {client_name or 'The Client'}
 - Proposal Due Date: {due_date_str or 'As specified in the RFP'}
 
 RFP Content:
 {rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional Contractual Terms template.'}"""
    elif is_appendix:
        # Use the specific prompt for Appendix
        system_prompt = (
            "You are an expert proposal writer specializing in government and commercial RFP responses. "
            "Your task is to generate a comprehensive 'List of Appendices' section that identifies all supporting documents included at the end of the proposal."
        )
        
        user_prompt = f"""Generate a 'List of Appendices' section. Analyze the RFP thoroughly to identify all supporting documents, references, and additional materials that should be included in the appendices.

RFP Content (analyze thoroughly to identify all supporting documents, attachments, references, and additional materials referenced in the proposal):
{rfp_content[:25000] if rfp_content else 'No RFP content available. Generate a professional List of Appendices template.'}

Company Information:
- Company Name: {company or 'Our Company'}
- Primary Contact: {contact_email or 'Contact information to be provided'}

RFP Information:
- Project/Solicitation Title: {project_name or 'This RFP'}
- Client/Issuing Agency: {client_name or 'The Client'}
- Proposal Due Date: {due_date_str or 'As specified in the RFP'}

Prompt: "Generate a 'List of Appendices.' Use a lettered, nested list to identify all supporting documents included at the end of the proposal. This should include (but is not limited to):

- Appendix A: 
- Appendix B: 
- Appendix C: 
- Appendix D: 
- Appendix E:"

IMPORTANT REQUIREMENTS:
- Generate a 'List of Appendices' section
- Use a lettered, nested list format (Appendix A, Appendix B, Appendix C, etc.)
- Identify all supporting documents that should be included at the end of the proposal
- Include documents such as (but not limited to):
  - Technical specifications and drawings
  - Certifications and accreditations
  - Past performance documentation
  - Key personnel resumes
  - Organizational charts
  - Compliance matrices
  - Pricing details and breakdowns
  - References and testimonials
  - Any other supporting materials referenced in the proposal
- Analyze the RFP to identify which specific appendices are required or referenced
- Provide clear, descriptive titles for each appendix
- Ensure the list is comprehensive and includes all relevant supporting documents
- Use professional formatting with proper lettering (A, B, C, D, E, etc.)

Return the content as a well-structured 'List of Appendices' with a lettered, nested list format. Each appendix should be clearly identified with a letter (A, B, C, D, E, etc.) and a descriptive title indicating what document or material it contains."""
    else:
        # Generic section generation
        system_prompt = (
            "You are an expert proposal writer helping to draft specific sections of a government proposal. "
            "Analyze the RFP content and generate comprehensive, compliant content for the requested section."
        )
        
        user_prompt = f"""Generate content for the section: "{section_title or section_id}".

Requirements:
1. Analyze the provided RFP content to understand requirements relevant to this section
2. Generate professional, compliant content that addresses the section topic
3. Ensure the content is specific, actionable, and aligned with government proposal standards
4. If RFP content is limited, provide a well-structured template that can be customized

RFP Content:
{rfp_content[:15000] if rfp_content else "No RFP content available. Generate a professional template for this section."}

Company: {company or 'Our Company'}
Project: {project_name or 'This Project'}

Please provide the content for "{section_title or section_id}" in a clear, well-formatted manner."""
    
    # Enrich all prompts with company profile context from database
    try:
        cur = mysql.connection.cursor(DictCursor)
        cur.execute("SELECT * FROM company_details ORDER BY id DESC LIMIT 1")
        _details = cur.fetchone() or {}
        cur.execute("SELECT * FROM company_preferences ORDER BY id DESC LIMIT 1")
        _prefs = cur.fetchone() or {}
        try:
            cur.execute("SELECT * FROM company_capabilities ORDER BY uploaded_at DESC LIMIT 5")
        except Exception:
            cur.execute("SELECT * FROM company_capabilities ORDER BY id DESC LIMIT 5")
        _caps = cur.fetchall() or []
        cur.execute("SELECT * FROM company_performance ORDER BY year DESC, id DESC LIMIT 5")
        _perf = cur.fetchall() or []
        cur.close()
    except Exception:
        _details, _prefs, _caps, _perf = {}, {}, [], []
    
    company_context_lines = []
    if _details:
        company_context_lines.append(f"- Legal Name: {(_details.get('name') or company or 'Our Company')}")
        if _details.get('website'):
            company_context_lines.append(f"- Website: {_details.get('website')}")
        if _details.get('email'):
            company_context_lines.append(f"- Email: {_details.get('email')}")
        if _details.get('phone'):
            company_context_lines.append(f"- Phone: {_details.get('phone')}")
        if _details.get('about'):
            company_context_lines.append(f"- About: {_details.get('about')[:400]}")
    if _prefs and (_prefs.get('registered_states') or '').strip():
        company_context_lines.append(f"- Registered/Eligible States: {_prefs.get('registered_states')}")
    if _caps:
        company_context_lines.append("- Capabilities:")
        for c in _caps[:5]:
            desc = (c.get('description') or '').strip()
            if desc:
                company_context_lines.append(f"  • {desc[:160]}")
    if _perf:
        company_context_lines.append("- Past Performance Highlights:")
        for p in _perf[:5]:
            pname = (p.get('project_name') or '').strip() or 'Project'
            pyear = p.get('year')
            company_context_lines.append(f"  • {pname}{f' ({pyear})' if pyear else ''}")
    company_context_block = "\n".join(company_context_lines).strip()
    if company_context_block:
        user_prompt = f"{user_prompt}\n\nCompany Profile Context (from database):\n{company_context_block}"
    
    api_key = (app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or '').strip()
    if not api_key:
        return jsonify({'error': 'api_unavailable', 'message': 'OpenAI API key is not configured on the server.'}), 500
    
    base_url = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or 'https://api.groq.com/openai/v1').strip()
    model = (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'llama-3.1-70b-versatile').strip()
    endpoint = base_url.rstrip('/')
    if not endpoint.endswith('/chat/completions'):
        endpoint = f"{endpoint}/chat/completions"
    
    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json',
    }
    body = {
        'model': model,
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt}
        ],
        'temperature': 0.3,
    }
    
    try:
        response = requests.post(endpoint, headers=headers, json=body, timeout=90)
    except requests.RequestException as err:
        app.logger.exception("OpenAI section content generation network failure: %s", err)
        return jsonify({'error': 'api_unavailable', 'message': 'Could not reach the AI service. Please try again shortly.'}), 502
    
    try:
        data = response.json()
    except ValueError:
        app.logger.error("OpenAI section content generation returned non-JSON response (status %s): %s", response.status_code, response.text[:400])
        return jsonify({'error': 'api_unavailable', 'message': 'Received an invalid response from the AI service.'}), 502
    
    if response.status_code >= 400:
        err_message = data.get('error') if isinstance(data, dict) else None
        detail = ''
        if isinstance(err_message, dict):
            detail = err_message.get('message') or err_message.get('code') or ''
        elif isinstance(err_message, str):
            detail = err_message
        if not detail:
            detail = f"Upstream error {response.status_code}"
        app.logger.error("OpenAI section content generation error %s: %s", response.status_code, detail)
        return jsonify({'error': 'api_unavailable', 'message': f'AI service returned an error: {detail}'}), 502
    
    choices = data.get('choices') or []
    if not choices:
        return jsonify({'error': 'api_unavailable', 'message': 'No response returned by the AI service.'}), 502
    
    content = choices[0].get('message', {}).get('content', '')
    if not content:
        return jsonify({'error': 'api_unavailable', 'message': 'The AI service returned an empty response.'}), 502
    
    # Split content into paragraphs for frontend display
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', content.strip()) if p.strip()]
    if not paragraphs:
        paragraphs = [content.strip()]
    
    return jsonify({
        'content': paragraphs,
        'raw': content.strip(),
        'model': data.get('model', model),
        'usage': data.get('usage', {}),
    })


@app.route('/api/compile-proposal-sections', methods=['POST'])
@login_required
def compile_proposal_sections():
    """Compile all generated section content into a single downloadable Word document."""
    payload = request.get_json(silent=True) or {}
    sections = payload.get('sections') or []
    if not isinstance(sections, list) or not sections:
        return jsonify({'error': 'missing_sections', 'message': 'No sections were provided to compile.'}), 400

    project_title = (payload.get('project_title') or '').strip()
    company_name = (payload.get('company') or '').strip()
    html_fragments = []

    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return jsonify({'error': 'dependency_missing', 'message': 'python-docx library is required to compile proposals.'}), 500

    try:
        # Always prefer the H & F template for consistent formatting
        template_used = None
        document = None
        try:
            base_dir = app.root_path if hasattr(app, 'root_path') else os.getcwd()
            hf_template_path = os.path.join(base_dir, 'H & F.docx')
            if os.path.exists(hf_template_path):
                document = Document(hf_template_path)
                template_used = 'hf'
            else:
                # Fallback to blank document if template missing
                document = Document()
        except Exception:
            # Defensive fallback in case template loading fails
            document = Document()
        
        # Enforce a consistent Times New Roman font across the document
        PREFERRED_FONT_NAME = 'Times New Roman'
        try:
            normal_style = document.styles['Normal']
            normal_style.font.name = PREFERRED_FONT_NAME
        except Exception:
            pass
        # Also try to set common heading/title styles to Times New Roman
        for _style_name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6', 'Heading 7', 'Heading 8', 'Heading 9']:
            try:
                document.styles[_style_name].font.name = PREFERRED_FONT_NAME
            except Exception:
                continue
        
        # Helper: style existence
        def _has_style(doc, style_name):
            try:
                _ = doc.styles[style_name]
                return True
            except Exception:
                return False
        
        # Add title page with robust style fallback
        if project_title:
            try:
                if _has_style(document, 'Title'):
                    title_para = document.add_paragraph(project_title, style=document.styles['Title'])
                elif _has_style(document, 'METCO Title'):
                    title_para = document.add_paragraph(project_title, style=document.styles['METCO Title'])
                else:
                    # Fall back to first-level heading if Title style is unavailable
                    title_para = document.add_heading(project_title, level=1)
                try:
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
            except Exception:
                # Absolute fallback to avoid style-related failures
                p = document.add_paragraph(project_title)
                try:
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                except Exception:
                    pass
            html_fragments.append(f'<h1 style=\"text-align: center;\">{escape(project_title)}</h1>')
            document.add_paragraph()  # Add spacing
        
        # Collect section titles for Table of Contents (before processing content)
        section_titles = []
        for section in sections:
            if not isinstance(section, dict):
                continue
            title = (section.get('title') or 'Untitled Section').strip() or 'Untitled Section'
            content_blocks = section.get('content') if isinstance(section.get('content'), list) else []
            raw_html = section.get('raw_content') or ''
            if content_blocks or raw_html:
                section_titles.append(title)
        
        # Add Table of Contents page if we have sections
        if section_titles:
            document.add_page_break()
            toc_heading = document.add_heading('Table of Contents', level=1)
            try:
                for run in toc_heading.runs:
                    run.font.name = 'Times New Roman'
            except Exception:
                pass
            document.add_paragraph()  # Add spacing
            
            # Add each section to TOC
            for idx, toc_title in enumerate(section_titles, 1):
                toc_para = document.add_paragraph(f'{idx}. {toc_title}')
                toc_para.paragraph_format.space_after = Pt(6)
                try:
                    for run in toc_para.runs:
                        run.font.name = 'Times New Roman'
                except Exception:
                    pass
            
            document.add_page_break()
            html_fragments.append('<h2>Table of Contents</h2>')
            html_fragments.append('<ol>' + ''.join(f'<li>{escape(title)}</li>' for title in section_titles) + '</ol>')
        
        # Process each section with generated content
        first_section_added = False
        # Determine a table style available in the template, prefer METCO-specific if present
        table_style_name = None
        try:
            from docx.enum.style import WD_STYLE_TYPE as _WD_STYLE_TYPE
            for s in document.styles:
                try:
                    if getattr(s, "type", None) == _WD_STYLE_TYPE.TABLE:
                        if s.name in ("METCO Table", "Table Grid", "Light Grid Accent 1", "Grid Table 4 Accent 1"):
                            table_style_name = s.name
                            if table_style_name == "METCO Table":
                                break
                except Exception:
                    continue
            if not table_style_name:
                table_style_name = "Table Grid"
        except Exception:
            table_style_name = "Table Grid"

        # End of initial preparation block
    except Exception as e:
        app.logger.exception('Unexpected error preparing document: %s', e)
        return jsonify({'error': 'unexpected_error', 'message': f'Compilation failed: {str(e)}'}), 500

    # Continue with compilation
    def _render_markdown_to_docx_paragraph(doc, text: str):
        """
        Render a single text block into the Word document, interpreting a minimal Markdown subset:
        - Lines starting with '### ' become Heading 3.
        - Bold spans wrapped in **double asterisks** become bold runs.
        Any remaining stray '*' or '#' characters are removed.
        """
        if not text:
            return
        try:
            import re as __re_md
            heading_match = __re_md.match(r'^\s*#{3,}\s*(.+?)\s*$', text)
            if heading_match:
                heading_text = heading_match.group(1)
                # Add heading level 3
                para = doc.add_heading(heading_text, level=3)
                # Ensure bold emphasis for heading text
                try:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.name = 'Times New Roman'
                except Exception:
                    pass
                return
            # Regular paragraph with bold spans
            paragraph = doc.add_paragraph()
            parts = __re_md.split(r'(\*\*.+?\*\*)', text)
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**') and len(part) >= 4:
                    content = part[2:-2]
                    run = paragraph.add_run(content)
                    try:
                        run.bold = True
                        run.font.name = 'Times New Roman'
                    except Exception:
                        pass
                else:
                    # Remove leftover single '*' or '#' characters
                    cleaned = __re_md.sub(r'[#*]+', '', part)
                    run = paragraph.add_run(cleaned)
                    try:
                        run.font.name = 'Times New Roman'
                    except Exception:
                        pass
        except Exception:
            # Fallback to simple paragraph if anything fails
            doc.add_paragraph(text)

    def _plaintext_markdown_to_html(text: str) -> str:
        """
        Convert a minimal Markdown subset to clean HTML string for preview.
        - '### Title' -> <h3>Title</h3>
        - '**bold**'  -> <strong>bold</strong>
        - Remove stray '*' and '#' characters elsewhere.
        """
        if not text:
            return ''
        import re as __re_md
        # Heading
        heading_match = __re_md.match(r'^\s*#{3,}\s*(.+?)\s*$', text)
        if heading_match:
            return f'<h3>{escape(heading_match.group(1))}</h3>'
        # Bold spans
        def _bold_repl(m):
            return f'<strong>{escape(m.group(1))}</strong>'
        html = __re_md.sub(r'\*\*(.+?)\*\*', lambda m: _bold_repl(m), text)
        # Escape remaining content safely, but preserve the <strong> tags we just added
        # Strategy: temporarily replace strong tags, escape, then restore
        placeholder_open = '___STRONG_OPEN___'
        placeholder_close = '___STRONG_CLOSE___'
        html = html.replace('<strong>', placeholder_open).replace('</strong>', placeholder_close)
        html = escape(html)
        html = html.replace(placeholder_open, '<strong>').replace(placeholder_close, '</strong>')
        # Remove stray '*' and '#' characters
        html = __re_md.sub(r'[#*]+', '', html)
        # Wrap as paragraph by default
        return f'<p>{html}</p>'
    def _insert_html_with_tables(doc, html_text):
        """
        Render simple HTML content into the Word document.
        - Preserves tables (<table><tr><td>) as real Word tables with a template style.
        - Non-table HTML is converted to paragraphs using _strip_html_tags.
        """
        if not html_text:
            return
        # Helpers to style header rows with blue/green background and white bold text
        def _shade_cell(cell, fill_hex):
            try:
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fill_hex)
                tcPr.append(shd)
            except Exception:
                pass
        def _style_table_header_row(tbl):
            try:
                from docx.shared import RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH
                from docx.shared import Pt as _Pt
            except Exception:
                return
            if not getattr(tbl, "rows", None) or len(tbl.rows) == 0:
                return
            header_row = tbl.rows[0]
            header_colors = ['2E74B5', '70AD47']  # blue, green
            for idx, cell in enumerate(header_row.cells):
                _shade_cell(cell, header_colors[idx % 2])
                for paragraph in cell.paragraphs:
                    paragraph.alignment = _WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        paragraph.paragraph_format.space_before = _Pt(0)
                        paragraph.paragraph_format.space_after = _Pt(2)
                    except Exception:
                        pass
                    for run in paragraph.runs:
                        run.font.bold = True
                        try:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.name = 'Times New Roman'
                        except Exception:
                            pass
        def _style_table_body(tbl):
            # Ensure body cells have consistent spacing and vertical alignment
            try:
                from docx.shared import Pt as _Pt
                from docx.enum.table import WD_ALIGN_VERTICAL as _WD_ALIGN_VERTICAL
            except Exception:
                return
            if not getattr(tbl, "rows", None) or len(tbl.rows) <= 1:
                return
            for r_idx, row in enumerate(tbl.rows):
                if r_idx == 0:
                    continue  # header handled separately
                for cell in row.cells:
                    try:
                        cell.vertical_alignment = _WD_ALIGN_VERTICAL.TOP
                    except Exception:
                        pass
                    for paragraph in cell.paragraphs:
                        try:
                            paragraph.paragraph_format.space_before = _Pt(0)
                            paragraph.paragraph_format.space_after = _Pt(2)
                        except Exception:
                            pass
                        try:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                        except Exception:
                            pass
        # Render general delimited blocks (tabs, 2+ spaces, or key: value pairs)
        def _try_render_delimited_table_block(doc, block_text):
            """
            Detects:
            - Tab-delimited rows across multiple lines
            - Multi-space (2+) delimited rows across multiple lines
            - Repeated 'Key: Value' lines -> 2-column table
            Returns True if a table was rendered.
            """
            if not block_text or '\n' not in block_text:
                return False
            lines = [ln for ln in block_text.splitlines() if ln.strip()]
            if len(lines) < 2:
                return False
            import re as __re
            # 1) Key: Value pairs
            colon_pairs = [__re.split(r'\s*:\s+', ln, maxsplit=1) for ln in lines if ':' in ln]
            if len(colon_pairs) == len(lines) and all(len(p) == 2 for p in colon_pairs):
                # Build 2-col table, first row as header inferred if all keys look like titles
                doc.add_page_break()
                tbl = doc.add_table(rows=(1 + len(colon_pairs)), cols=2)
                try:
                    if table_style_name:
                        tbl.style = table_style_name
                except Exception:
                    pass
                try:
                    from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
                    tbl.alignment = _WD_TABLE_ALIGNMENT.CENTER
                except Exception:
                    pass
                # Header
                tbl.cell(0, 0).text = 'Key'
                tbl.cell(0, 1).text = 'Value'
                # Body
                for r_off, (k, v) in enumerate(colon_pairs, start=1):
                    tbl.cell(r_off, 0).text = _strip_html_tags(k.strip())
                    tbl.cell(r_off, 1).text = _strip_html_tags(v.strip())
                _style_table_header_row(tbl)
                _style_table_body(tbl)
                doc.add_paragraph()
                return True
            # 2) Tab-delimited rows
            if any('\t' in ln for ln in lines):
                rows = [ln.split('\t') for ln in lines]
                num_cols = max(len(r) for r in rows)
                if num_cols >= 2:
                    doc.add_page_break()
                    tbl = doc.add_table(rows=len(rows), cols=num_cols)
                    try:
                        if table_style_name:
                            tbl.style = table_style_name
                    except Exception:
                        pass
                    try:
                        from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
                        tbl.alignment = _WD_TABLE_ALIGNMENT.CENTER
                    except Exception:
                        pass
                    for r_idx, row_vals in enumerate(rows):
                        for c_idx in range(num_cols):
                            val = row_vals[c_idx] if c_idx < len(row_vals) else ''
                            tbl.cell(r_idx, c_idx).text = _strip_html_tags(val.strip())
                    _style_table_header_row(tbl)
                    _style_table_body(tbl)
                    doc.add_paragraph()
                    return True
            # 3) Multi-space-delimited rows
            space_split_rows = [__re.split(r'\s{2,}', ln.strip()) for ln in lines]
            num_cols = max(len(r) for r in space_split_rows)
            if num_cols >= 2 and all(len(r) >= 2 for r in space_split_rows):
                doc.add_page_break()
                tbl = doc.add_table(rows=len(space_split_rows), cols=num_cols)
                try:
                    if table_style_name:
                        tbl.style = table_style_name
                except Exception:
                    pass
                try:
                    from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
                    tbl.alignment = _WD_TABLE_ALIGNMENT.CENTER
                except Exception:
                    pass
                for r_idx, row_vals in enumerate(space_split_rows):
                    for c_idx in range(num_cols):
                        val = row_vals[c_idx] if c_idx < len(row_vals) else ''
                        tbl.cell(r_idx, c_idx).text = _strip_html_tags(val.strip())
                _style_table_header_row(tbl)
                _style_table_body(tbl)
                doc.add_paragraph()
                return True
            return False
        # Attempt to detect and render Markdown/pipe tables from plain text blocks
        def _try_render_markdown_table_block(doc, block_text):
            """
            Detects Markdown-style tables:
            | H1 | H2 |
            | --- | --- |
            | v1 | v2 |
            Or pipe-delimited rows without explicit alignment row.
            Returns True if a table was rendered, False otherwise.
            """
            lines = [ln for ln in (block_text or '').splitlines() if ln.strip()]
            if len(lines) < 2:
                return False
            import re as __re
            def split_cells(line):
                # Keep inner pipes by splitting and trimming; drop leading/trailing empty caused by outer pipes
                parts = [c.strip() for c in line.strip().split('|')]
                if parts and parts[0] == '':
                    parts = parts[1:]
                if parts and parts[-1] == '':
                    parts = parts[:-1]
                return parts
            def is_pipe_line(ln):
                return '|' in ln and len(split_cells(ln)) >= 2
            def is_align_row(ln, cols):
                # e.g., | --- | :---: | --- |
                cells = split_cells(ln)
                if len(cells) != cols:
                    return False
                for c in cells:
                    if not __re.fullmatch(r':?-{3,}:?', c.replace(' ', '')):
                        return False
                return True
            if not is_pipe_line(lines[0]) or not is_pipe_line(lines[1]):
                return False
            header_cells = split_cells(lines[0])
            cols = len(header_cells)
            has_align = is_align_row(lines[1], cols)
            remaining = lines[2:] if has_align else lines[1:]
            if not remaining:
                return False
            # Assemble rows allowing multi-line cell continuations (lines without pipes)
            assembled_rows = []
            current_row = None
            for raw_ln in remaining:
                ln = raw_ln.rstrip()
                if is_pipe_line(ln):
                    # flush any current row
                    if current_row is not None:
                        if len(current_row) < cols:
                            current_row += [''] * (cols - len(current_row))
                        assembled_rows.append(current_row[:cols])
                    current_row = split_cells(ln)
                else:
                    if current_row is None:
                        return False
                    if not current_row:
                        current_row.append(ln.strip())
                    else:
                        current_row[-1] = (current_row[-1] + '\n' + ln.strip()).strip()
            if current_row is not None:
                if len(current_row) < cols:
                    current_row += [''] * (cols - len(current_row))
                assembled_rows.append(current_row[:cols])
            if not assembled_rows:
                return False
            # Build table
            doc.add_page_break()
            tbl = doc.add_table(rows=(1 + len(assembled_rows)), cols=cols)
            try:
                if table_style_name:
                    tbl.style = table_style_name
            except Exception:
                pass
            try:
                from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
                tbl.alignment = _WD_TABLE_ALIGNMENT.CENTER
            except Exception:
                pass
            # Header
            for c_idx, text in enumerate(header_cells):
                tbl.cell(0, c_idx).text = _strip_html_tags(text)
            # Body
            for r_off, row_cells in enumerate(assembled_rows, start=1):
                # ensure length to cols
                if len(row_cells) < cols:
                    row_cells += [''] * (cols - len(row_cells))
                for c_idx, text in enumerate(row_cells[:cols]):
                    tbl.cell(r_off, c_idx).text = _strip_html_tags(text)
            _style_table_header_row(tbl)
            _style_table_body(tbl)
            doc.add_paragraph()
            return True
        import re as _re
        table_splitter = _re.compile(r'(<table[\s\S]*?</table>)', _re.IGNORECASE)
        table_rows_re = _re.compile(r'<tr[\s\S]*?</tr>', _re.IGNORECASE)
        cell_contents_re = _re.compile(r'<t[dh][^>]*?>([\s\S]*?)</t[dh]>', _re.IGNORECASE)

        segments = table_splitter.split(html_text)
        for segment in segments:
            if not segment:
                continue
            if segment.lstrip().lower().startswith('<table'):
                rows = table_rows_re.findall(segment) or []
                num_cols = 0
                for rhtml in rows:
                    num_cols = max(num_cols, len(cell_contents_re.findall(rhtml)))
                if num_cols <= 0:
                    # Fallback: treat as paragraph text
                    text_version = _strip_html_tags(segment)
                    for p in [p for p in text_version.split('\n\n') if p.strip()]:
                        para = doc.add_paragraph(p.strip())
                        try:
                            for run in para.runs:
                                run.font.name = 'Times New Roman'
                        except Exception:
                            pass
                    continue
                # Start each table on a new page
                doc.add_page_break()
                tbl = doc.add_table(rows=len(rows), cols=num_cols)
                try:
                    if table_style_name:
                        tbl.style = table_style_name
                except Exception:
                    pass
                try:
                    from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
                    tbl.alignment = _WD_TABLE_ALIGNMENT.CENTER
                except Exception:
                    pass
                for r_idx, rhtml in enumerate(rows):
                    cells = cell_contents_re.findall(rhtml)
                    for c_idx, cell_html in enumerate(cells):
                        if c_idx >= num_cols:
                            break
                        cell_text = _strip_html_tags(cell_html)
                        tbl.cell(r_idx, c_idx).text = cell_text
                # Style the header row with alternating blue/green colors
                _style_table_header_row(tbl)
                # Apply body cell formatting
                _style_table_body(tbl)
                # Add a small spacing after the table for readability
                doc.add_paragraph()
            else:
                text_version = _strip_html_tags(segment)
                # Try to identify markdown/pipe tables by blocks first
                blocks = [b for b in text_version.split('\n\n')]
                for block in blocks:
                    blk = block.strip()
                    if not blk:
                        continue
                    if _try_render_markdown_table_block(doc, blk) or _try_render_delimited_table_block(doc, blk):
                        continue
                    # Fallback: treat as plain paragraphs (split single newlines too)
                    paras = [p.strip() for p in blk.split('\n') if p.strip()]
                    for para_text in paras:
                        para = doc.add_paragraph(para_text)
                        try:
                            for run in para.runs:
                                run.font.name = 'Times New Roman'
                        except Exception:
                            pass

    for section in sections:
        if not isinstance(section, dict):
            continue
        title = (section.get('title') or 'Untitled Section').strip() or 'Untitled Section'
        content_blocks = section.get('content') if isinstance(section.get('content'), list) else []
        raw_html = section.get('raw_content') or ''
        
        # Skip sections with no content
        if not content_blocks and not raw_html:
            continue

        # Ensure every new section begins on a new page
        if first_section_added:
            document.add_page_break()
        first_section_added = True

        # Add section heading
        heading_para = document.add_heading(title, level=1)
        try:
            for run in heading_para.runs:
                run.font.name = 'Times New Roman'
        except Exception:
            pass
        html_fragments.append(f'<h2>{escape(title)}</h2>')

        # Process content - prioritize raw_html if available (for HTML tables, etc.)
        if raw_html:
            # Render HTML preserving tables; append raw_html to preview
            _insert_html_with_tables(document, raw_html)
            try:
                # Apply minimal markdown cleanup inside raw_html for preview only
                import re as __re_md
                rh = raw_html
                # Convert bold markers inside raw HTML text content crudely
                rh = __re_md.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', rh)
                # Convert '### ' headings at the start of lines to <h3>
                rh = __re_md.sub(r'(^|\n)\s*#{3,}\s*(.+?)\s*(?=\n|$)', lambda m: f"{m.group(1)}<h3>{escape(m.group(2))}</h3>", rh)
                # Remove remaining stray markers
                rh = __re_md.sub(r'[#*]+', '', rh)
                html_fragments.append(rh)
            except Exception:
                html_fragments.append(raw_html)
        else:
            # Process regular content blocks
            fragment_html = ''
            for paragraph in content_blocks:
                if paragraph and paragraph.strip():
                    para_text = paragraph.strip()
                    # If the paragraph looks like a table block (pipe, tabs, multi-space, or key:value across lines), render it as a real table
                    if '\n' in para_text:
                        try:
                            import re as __re2
                            looks_table = (
                                ('|' in para_text) or
                                ('\t' in para_text) or
                                (__re2.search(r'\s{2,}', para_text) is not None) or
                                (__re2.search(r'^[^:\n]+:\s*.+$', para_text, flags=__re2.M) is not None)
                            )
                        except Exception:
                            looks_table = ('|' in para_text) or ('\t' in para_text)
                    else:
                        looks_table = False
                    if looks_table:
                        _insert_html_with_tables(document, para_text)
                        # Keep preview minimal; embed as preformatted text block
                        fragment_html += f'<pre>{escape(para_text)}</pre>'
                    else:
                        # Render with Markdown interpretation
                        _render_markdown_to_docx_paragraph(document, para_text)
                        fragment_html += _plaintext_markdown_to_html(para_text)
            
            if fragment_html:
                html_fragments.append(fragment_html)
        # No extra paragraph needed here; page breaks handle separation

    if not html_fragments:
        return jsonify({'error': 'empty_sections', 'message': 'No compiled content available. Generate or enter section content first.'}), 400

    # Generate filename with project title if available
    os.makedirs('uploads/proposals', exist_ok=True)
    timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
    if project_title:
        # Sanitize project title for filename
        safe_title = ''.join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in project_title[:50])
        safe_title = safe_title.replace(' ', '_')
        filename = f"Proposal_{safe_title}_{timestamp}.docx"
    else:
        filename = f"Proposal_{timestamp}.docx"

    filepath = os.path.join('uploads/proposals', filename)

    try:
        # Ensure headers/footers are linked across sections and page numbering is continuous
        try:
            # Link headers/footers after all content is added
            for _s_idx, _section in enumerate(document.sections):
                if _s_idx > 0:
                    try:
                        _section.header.is_linked_to_previous = True
                        _section.footer.is_linked_to_previous = True
                    except Exception:
                        pass
            # Continuous page numbering: start at 1 on first section, continue thereafter
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            for _s_idx, _section in enumerate(document.sections):
                _sectPr = _section._sectPr
                _pgNumType = _sectPr.find(qn('w:pgNumType'))
                if _pgNumType is None:
                    _pgNumType = OxmlElement('w:pgNumType')
                    _sectPr.append(_pgNumType)
                if _s_idx == 0:
                    _pgNumType.set(qn('w:start'), "1")
                else:
                    try:
                        _pgNumType.attrib.pop(qn('w:start'), None)
                    except Exception:
                        pass
        except Exception:
            # If linking or numbering XML manipulation fails, proceed without blocking save
            pass
        document.save(filepath)
        app.logger.info(f'Successfully compiled proposal with {len(sections)} sections to {filename}')
    except Exception as error:
        app.logger.exception('Failed to save compiled proposal: %s', error)
        return jsonify({'error': 'save_failed', 'message': 'Could not save the compiled proposal. Please try again.'}), 500

    preview_html = ''.join(html_fragments)
    return jsonify({
        'message': f'Proposal compiled successfully with {len(sections)} section(s). Review the preview before downloading.',
        'filename': filename,
        'download_url': url_for('download_proposal', filename=filename),
        'html_preview': preview_html,
        'section_count': len(sections)
    })

@app.route('/api/validate-requirement-attachment', methods=['POST'])
@login_required
def validate_requirement_attachment():
    """Validate that an uploaded document matches the required form type."""
    if 'file' not in request.files:
        return jsonify({'error': 'missing_file', 'message': 'No file provided.'}), 400
    
    file = request.files['file']
    requirement_text = request.form.get('requirement_text', '').strip()
    
    if not file or not file.filename:
        return jsonify({'error': 'missing_file', 'message': 'No file provided.'}), 400
    
    if not requirement_text:
        return jsonify({'error': 'missing_requirement', 'message': 'Requirement text is required.'}), 400
    
    # Validate file type
    filename = file.filename.lower()
    if not (filename.endswith('.pdf') or filename.endswith('.jpg') or filename.endswith('.jpeg')):
        return jsonify({'error': 'invalid_format', 'message': 'Only PDF and JPG files are allowed.'}), 400
    
    try:
        # Save file temporarily
        import tempfile
        import os
        from werkzeug.utils import secure_filename
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp_file:
            file.save(tmp_file.name)
            tmp_path = tmp_file.name
        
        try:
            # Extract text from document
            extracted_text = ''
            
            if filename.endswith('.pdf'):
                # Extract text from PDF
                if _FITZ_AVAILABLE:
                    with _pymupdf.open(tmp_path) as doc:
                        for page_num in range(min(3, doc.page_count)):  # First 3 pages
                            page = doc.load_page(page_num)
                            extracted_text += page.get_text("text") + "\n\n"
                else:
                    # Fallback to PyPDF2
                    if _PdfReader is not None:
                        with open(tmp_path, "rb") as fh:
                            reader = _PdfReader(fh)
                            for page_num in range(min(3, len(reader.pages))):
                                extracted_text += reader.pages[page_num].extract_text() + "\n\n"
                    else:
                        return jsonify({'error': 'pdf_library_missing', 'message': 'PDF processing library not available.'}), 500
            
            elif filename.endswith('.jpg') or filename.endswith('.jpeg'):
                # For images, we'll use OpenAI Vision API or return limited validation
                # For now, we'll use AI to analyze the image directly
                extracted_text = '[IMAGE_FILE]'  # Placeholder for image files
            
            # Use AI to validate document content
            api_key = (app.config.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY') or '').strip()
            if not api_key:
                return jsonify({'error': 'api_unavailable', 'message': 'OpenAI API key is not configured.'}), 500
            
            base_url = (app.config.get('OPENAI_BASE_URL') or os.getenv('OPENAI_BASE_URL') or 'https://api.groq.com/openai/v1').strip()
            model = (app.config.get('OPENAI_MODEL') or os.getenv('OPENAI_MODEL') or 'llama-3.1-70b-versatile').strip()
            
            # For images, use vision model
            if filename.endswith('.jpg') or filename.endswith('.jpeg'):
                # Read image as base64
                import base64
                with open(tmp_path, 'rb') as img_file:
                    image_data = base64.b64encode(img_file.read()).decode('utf-8')
                
                endpoint = base_url.rstrip('/')
                if not endpoint.endswith('/chat/completions'):
                    endpoint = f"{endpoint}/chat/completions"
                
                headers = {
                    'Authorization': f'Bearer {api_key}',
                    'Content-Type': 'application/json',
                }
                
                system_prompt = (
                    "You are a document validation expert. Your task is to analyze uploaded documents "
                    "and determine if they match the required form type specified in the requirement text. "
                    "Return only 'YES' if the document matches the requirement, or 'NO' with a brief reason if it doesn't."
                )
                
                user_prompt = f"""
Analyze the uploaded image and determine if it matches the required document type.

REQUIREMENT: {requirement_text}

The requirement may specify forms like:
- W-9 Form (Request for Taxpayer Identification Number and Certification)
- W-8 Form (Certificate of Foreign Status)
- Insurance Certificate
- Bid Bond
- Performance Bond
- License/Certification
- Company Profile
- References
- Any other specific form or document

Examine the image carefully and check:
1. Does the document title/header match the required form name?
2. Are there identifying markers (form numbers, official headers) that match?
3. Is this the correct document type?

Respond with:
- "YES" if the document matches the requirement
- "NO: [brief reason]" if it doesn't match

Your response should be only "YES" or "NO: [reason]".
"""
                
                body = {
                    'model': 'gpt-4o',  # Use vision-capable model
                    'messages': [
                        {'role': 'system', 'content': system_prompt},
                        {
                            'role': 'user',
                            'content': [
                                {'type': 'text', 'text': user_prompt},
                                {
                                    'type': 'image_url',
                                    'image_url': {
                                        'url': f'data:image/jpeg;base64,{image_data}'
                                    }
                                }
                            ]
                        }
                    ],
                    'temperature': 0.1,
                    'max_tokens': 150
                }
            else:
                # For PDFs, use text-based validation
                endpoint = base_url.rstrip('/')
                if not endpoint.endswith('/chat/completions'):
                    endpoint = f"{endpoint}/chat/completions"
                
                headers = {
                    'Authorization': f'Bearer {api_key}',
                    'Content-Type': 'application/json',
                }
                
                system_prompt = (
                    "You are a document validation expert. Your task is to analyze extracted text from documents "
                    "and determine if they match the required form type specified in the requirement text. "
                    "Return only 'YES' if the document matches the requirement, or 'NO' with a brief reason if it doesn't."
                )
                
                user_prompt = f"""
Analyze the extracted text from the uploaded document and determine if it matches the required document type.

REQUIREMENT: {requirement_text}

The requirement may specify forms like:
- W-9 Form (Request for Taxpayer Identification Number and Certification)
- W-8 Form (Certificate of Foreign Status)
- Insurance Certificate
- Bid Bond
- Performance Bond
- License/Certification
- Company Profile
- References
- Any other specific form or document

EXTRACTED TEXT FROM DOCUMENT:
{extracted_text[:3000] if extracted_text else '[No text extracted]'}

Examine the text carefully and check:
1. Does the document title/header match the required form name?
2. Are there identifying markers (form numbers, official headers, form names) that match?
3. Is this the correct document type?

Respond with:
- "YES" if the document matches the requirement
- "NO: [brief reason]" if it doesn't match

Your response should be only "YES" or "NO: [reason]".
"""
                
                body = {
                    'model': model,
                    'messages': [
                        {'role': 'system', 'content': system_prompt},
                        {'role': 'user', 'content': user_prompt}
                    ],
                    'temperature': 0.1,
                    'max_tokens': 150
                }
            
            # Call OpenAI API
            import requests
            response = requests.post(endpoint, headers=headers, json=body, timeout=30)
            
            if response.status_code >= 400:
                data = response.json() if response.content else {}
                error_msg = data.get('error', {}).get('message', 'Validation failed') if isinstance(data, dict) else 'Validation failed'
                return jsonify({'error': 'validation_failed', 'message': error_msg}), 502
            
            data = response.json()
            choices = data.get('choices', [])
            if not choices:
                return jsonify({'error': 'validation_failed', 'message': 'No response from validation service.'}), 502
            
            validation_result = choices[0].get('message', {}).get('content', '').strip().upper()
            is_valid = validation_result.startswith('YES')
            
            return jsonify({
                'is_valid': is_valid,
                'validation_message': validation_result,
                'requirement': requirement_text
            })
        
        finally:
            # Clean up temporary file
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
    
    except Exception as e:
        app.logger.exception("Error validating attachment: %s", e)
        return jsonify({'error': 'validation_error', 'message': f'Validation failed: {str(e)}'}), 500

@app.route('/proposals-making/download/<filename>')
@login_required
def download_proposal(filename):
    # Allow all logged-in users to download proposals (removed admin-only restriction)
    from werkzeug.utils import secure_filename
    from flask import send_from_directory
    safe_filename = secure_filename(filename)
    proposals_dir = 'uploads/proposals'
    return send_from_directory(proposals_dir, safe_filename, as_attachment=True)

# --- API & Real-time Logic ---
@app.route('/api/update_stage/<int:bid_id>', methods=['POST'])
@login_required
def update_stage(bid_id):
    try:
        data = request.get_json()
        new_stage = (data.get('stage') or '').lower()
        
        # Validate stage
        allowed = {'analyzer', 'business', 'design', 'operations', 'engineer', 'handover'}
        if new_stage not in allowed:
            return jsonify({'error': 'invalid stage'}), 400
        
        cur = mysql.connection.cursor(DictCursor)
        
        # Get bid information from go_bids
        cur.execute("SELECT * FROM go_bids WHERE g_id=%s", (bid_id,))
        bid = cur.fetchone()
        
        if not bid:
            cur.close()
            return jsonify({'error': 'Bid not found or access denied'}), 404

        # Get old stage for logging
        old_stage = (bid.get('state') or 'analyzer').lower()
        
        # Update go_bids state
        cur.execute("UPDATE go_bids SET state=%s WHERE g_id=%s", (new_stage, bid_id))
        
        # Upsert assignment so the next team still sees it
        cur.execute("SELECT a_id FROM bid_assign WHERE g_id=%s", (bid_id,))
        row = cur.fetchone()
        if row:
            cur.execute("UPDATE bid_assign SET depart=%s, state=%s, status='pending' WHERE g_id=%s",
                        (new_stage, new_stage, bid_id))
        else:
            cur.execute("""
                INSERT INTO bid_assign (g_id, b_name, due_date, state, scope, type, company, depart,
                                       person_name, assignee_email, status, value, revenue)
                SELECT g_id, b_name, due_date, state, scope, type, company, %s, '', '', 'pending',
                       COALESCE(scoring, 0), COALESCE(revenue, 0)
                FROM go_bids WHERE g_id=%s
            """, (new_stage, bid_id))
        
        # Derive dynamic summary line
        from_txt = LABELS.get(old_stage, '')
        to_txt = LABELS.get(new_stage, '')
        summary_line = f"Updated by {from_txt} to {to_txt}"
        
        # Commit the transaction
        mysql.connection.commit()
        
        # Log the stage change
        log_write('stage_change', f"{bid.get('b_name')} | {old_stage} → {new_stage}")
        
        # Calculate dynamic progress and status texts for new stage
        pct = pct_for(new_stage)
        proj_status = 'completed' if new_stage == 'handover' else 'ongoing'
        
        # Recompute cards and analyzer stats after update
        # Summary cards across the three companies only
        cur2 = mysql.connection.cursor(DictCursor)
        cur2.execute("SELECT id FROM companies WHERE name IN ('Ikio','Metco','Sunsprint')")
        target_company_ids = [row['id'] for row in cur2.fetchall()]
        if not target_company_ids:
            target_company_ids = [-1]
        in_clause = ','.join(['%s'] * len(target_company_ids))

        # Compute totals from go_bids instead of bids
        cur2.execute("SELECT name FROM companies WHERE name IN ('Ikio','Metco','Sunsprint')")
        target_company_names = [row['name'] for row in cur2.fetchall()] or ['__none__']
        in_clause_names = ','.join(['%s'] * len(target_company_names))
        cur2.execute(f"SELECT COUNT(*) AS total_bids FROM go_bids WHERE company IN ({in_clause_names})", target_company_names)
        total_bids = cur2.fetchone()['total_bids']
        cur2.execute(f"SELECT COUNT(*) AS live_bids FROM go_bids WHERE COALESCE(state,'analyzer') IN ('business','design','operations','engineer') AND company IN ({in_clause_names})", target_company_names)
        live_bids = cur2.fetchone()['live_bids']
        cur2.execute(f"SELECT COUNT(*) AS bids_won FROM go_bids WHERE decision='WON' AND company IN ({in_clause_names})", target_company_names)
        bids_won = cur2.fetchone()['bids_won']
        # Total projects across target companies (fallback to all projects if none linked)
        cur2.execute(f"SELECT COUNT(*) AS projects_linked FROM projects WHERE company_id IN ({in_clause})", target_company_ids)
        projects_linked = cur2.fetchone()['projects_linked']
        if projects_linked > 0:
            cur2.execute(f"SELECT COUNT(*) AS projects_total FROM projects WHERE company_id IN ({in_clause})", target_company_ids)
            projects_total = cur2.fetchone()['projects_total']
        else:
            cur2.execute("SELECT COUNT(*) AS projects_total FROM projects")
            projects_total = cur2.fetchone()['projects_total']

        # Analyzer stats from bid_incoming table
        cur2.execute("SELECT COUNT(*) AS total_bids FROM bid_incoming")
        total_bids_analyzer = cur2.fetchone()['total_bids']
        
        cur2.execute("SELECT COUNT(*) AS bids_go FROM bid_incoming WHERE decision = 'GO'")
        bids_go_analyzer = cur2.fetchone()['bids_go']
        
        cur2.execute("SELECT COUNT(*) AS bids_no_go FROM bid_incoming WHERE decision = 'NO-GO'")
        bids_no_go_analyzer = cur2.fetchone()['bids_no_go']
        
        cur2.execute("SELECT COUNT(*) AS bids_submitted FROM bid_incoming WHERE state IN ('submitted', 'under_review')")
        bids_submitted_analyzer = cur2.fetchone()['bids_submitted']
        
        cur2.execute("SELECT COUNT(*) AS bids_won FROM bid_incoming WHERE decision = 'WON'")
        bids_won_analyzer = cur2.fetchone()['bids_won']
        
        cur2.execute("SELECT COUNT(*) AS bids_lost FROM bid_incoming WHERE decision = 'LOST'")
        bids_lost_analyzer = cur2.fetchone()['bids_lost']

        bid_stats = {
            'total_bids': total_bids_analyzer,
            'bids_go': bids_go_analyzer,
            'bids_no_go': bids_no_go_analyzer,
            'bids_submitted': bids_submitted_analyzer,
            'bids_won': bids_won_analyzer,
            'bids_lost': bids_lost_analyzer
        }

        socketio.emit('master_update', {
            'bid': {
                'id': bid_id,
                'name': bid.get('b_name'),
                'current_stage': new_stage,
                'user_email': getattr(current_user, 'email', '')
            },
            'summary': {
                'work_progress_pct': pct,
                'project_status': proj_status,
                'work_status': summary_line
            },
            'cards': {
                'total_bids': projects_total,
                'live_bids': 0,
                'bids_won': 0,
                'projects_completed': 0
            },
            'bid_stats': bid_stats
        })
        
        cur.close()
        cur2.close()
        return jsonify({'success': f'Bid {bid_id} updated to {new_stage}'})
    
    except Exception as e:
        mysql.connection.rollback()
        if 'cur' in locals():
            cur.close()
        if 'cur2' in locals():
            cur2.close()
        return jsonify({'error': f'Error updating stage: {str(e)}'}), 500
# --- Main execution ---
def _ensure_tables_exist():
    """Create tables if they don't exist"""
    try:
        cur = mysql.connection.cursor()
        
        # Create users table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id INT AUTO_INCREMENT PRIMARY KEY,
                email VARCHAR(100) UNIQUE NOT NULL,
                password VARCHAR(100) NOT NULL,
                is_admin BOOLEAN DEFAULT FALSE,
                role VARCHAR(50) DEFAULT 'member'
            )
        """)
        
        # Create bids table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS bids (
                id INT AUTO_INCREMENT PRIMARY KEY,
                name VARCHAR(100) NOT NULL,
                current_stage VARCHAR(50) DEFAULT 'analyzer',
                user_id INT,
                company_id INT,
                FOREIGN KEY (user_id) REFERENCES users(id)
            )
        """)
        
        # Create bid_incoming table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS bid_incoming (
                id INT AUTO_INCREMENT PRIMARY KEY,
                b_name VARCHAR(100),
               
                due_date DATE NOT NULL DEFAULT CURRENT_TIMESTAMP,
                state VARCHAR(100),
                scope TEXT,
                type VARCHAR(100),
                scoring INT,
                comp_name VARCHAR(100),
                decision VARCHAR(100),
                summary TEXT
            )
        """)
        
        # Create go_bids table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS go_bids (
                g_id INT AUTO_INCREMENT PRIMARY KEY,
                id INT,
                b_name VARCHAR(100),
               
                due_date DATE NOT NULL DEFAULT CURRENT_TIMESTAMP,
                state VARCHAR(100),
                scope TEXT,
                type VARCHAR(100),
                scoring INT,
                company TEXT,
                decision TEXT,
                summary TEXT,
                revenue DECIMAL(15,2) DEFAULT 0.00
            )
        """)
        
        # Ensure summary and scope columns exist and are TEXT type for detailed content
        # Check and alter go_bids table
        try:
            cur.execute("SHOW COLUMNS FROM go_bids LIKE 'summary'")
            summary_col = cur.fetchone()
            if not summary_col:
                cur.execute("ALTER TABLE go_bids ADD COLUMN summary TEXT")
            else:
                # Check if it's not TEXT type and alter it
                col_type = summary_col[1] if isinstance(summary_col, tuple) else summary_col.get('Type', '')
                if col_type and 'text' not in str(col_type).lower():
                    cur.execute("ALTER TABLE go_bids MODIFY COLUMN summary TEXT")
        except Exception:
            pass
        
        try:
            cur.execute("SHOW COLUMNS FROM go_bids LIKE 'scope'")
            scope_col = cur.fetchone()
            if not scope_col:
                cur.execute("ALTER TABLE go_bids ADD COLUMN scope TEXT")
            else:
                # Check if it's not TEXT type and alter it
                col_type = scope_col[1] if isinstance(scope_col, tuple) else scope_col.get('Type', '')
                if col_type and 'text' not in str(col_type).lower():
                    cur.execute("ALTER TABLE go_bids MODIFY COLUMN scope TEXT")
        except Exception:
            pass
        
        # Ensure summary and scope columns exist in bid_incoming table
        try:
            cur.execute("SHOW COLUMNS FROM bid_incoming LIKE 'summary'")
            summary_col = cur.fetchone()
            if not summary_col:
                cur.execute("ALTER TABLE bid_incoming ADD COLUMN summary TEXT")
            else:
                # Check if it's not TEXT type and alter it
                col_type = summary_col[1] if isinstance(summary_col, tuple) else summary_col.get('Type', '')
                if col_type and 'text' not in str(col_type).lower():
                    cur.execute("ALTER TABLE bid_incoming MODIFY COLUMN summary TEXT")
        except Exception:
            pass
        
        try:
            cur.execute("SHOW COLUMNS FROM bid_incoming LIKE 'scope'")
            scope_col = cur.fetchone()
            if not scope_col:
                cur.execute("ALTER TABLE bid_incoming ADD COLUMN scope TEXT")
            else:
                # Check if it's not TEXT type and alter it
                col_type = scope_col[1] if isinstance(scope_col, tuple) else scope_col.get('Type', '')
                if col_type and 'text' not in str(col_type).lower():
                    cur.execute("ALTER TABLE bid_incoming MODIFY COLUMN scope TEXT")
        except Exception:
            pass
        
        # Create bid_assign table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS bid_assign (
                a_id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT,
                b_name VARCHAR(100),
               
                due_date DATE NOT NULL DEFAULT CURRENT_TIMESTAMP,
                state VARCHAR(100),
                scope VARCHAR(100),
                type VARCHAR(100),
                company TEXT,
                depart TEXT,
                person_name TEXT,
                assignee_email VARCHAR(100),
                status TEXT,
                value INT,
                revenue DECIMAL(15,2) DEFAULT 0.00
            )
        """)
        
        # Create win_lost_results table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS win_lost_results (
                w_id INT AUTO_INCREMENT PRIMARY KEY,
                a_id INT,
                b_name TEXT,
              
                due_date INT,
                state TEXT,
                scope TEXT,
                value INT,
                company TEXT,
                department TEXT,
                person_name TEXT,
                status TEXT,
                result TEXT
            )
        """)
        
        # Create won_bids_result table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS won_bids_result (
                won_id INT AUTO_INCREMENT PRIMARY KEY,
                w_id INT,
                closure_status TEXT,
                work_progress_status TEXT
            )
        """)
        
        # Create assigned_bids table using the exact schema requested
        cur.execute("""
            CREATE TABLE IF NOT EXISTS assigned_bids (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT,
                b_name VARCHAR(100),
                company VARCHAR(100),
                revenue DECIMAL(15,2) DEFAULT 0.00,
                assigned_to VARCHAR(100),
                assigned_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Create work_progress_status table (extended schema used across the app)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS work_progress_status (
                p_id INT AUTO_INCREMENT PRIMARY KEY,
                won_id INT,
                company TEXT,
                b_name TEXT,
                dept_bde TEXT,
                dept_m_d TEXT,
                dept_op TEXT,
                dept_site TEXT,
                pr_completion_status TEXT
            )
        """)

    
        
        # Create logs table for tracking user actions
        cur.execute("""
            CREATE TABLE IF NOT EXISTS logs (
                id INT AUTO_INCREMENT PRIMARY KEY,
                action VARCHAR(255) NOT NULL,
                user_id INT,
                timestamp DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users(id)
            )
        """)
        
        # Create companies table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS companies (
                id INT AUTO_INCREMENT PRIMARY KEY,
                name VARCHAR(100) UNIQUE NOT NULL,
                description TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Create projects table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS projects (
                id INT AUTO_INCREMENT PRIMARY KEY,
                name VARCHAR(200) NOT NULL,
                company_id INT NOT NULL,
                start_date DATETIME,
                due_date DATETIME NOT NULL,
                revenue FLOAT DEFAULT 0.0,
                status VARCHAR(50) DEFAULT 'active',
                progress INT DEFAULT 0,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (company_id) REFERENCES companies(id)
            )
        """)
        
        # Create tasks table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS tasks (
                id INT AUTO_INCREMENT PRIMARY KEY,
                name VARCHAR(200) NOT NULL,
                project_id INT NOT NULL,
                assigned_user_id INT,
                due_date DATETIME NOT NULL,
                status VARCHAR(50) DEFAULT 'pending',
                priority VARCHAR(20) DEFAULT 'medium',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (project_id) REFERENCES projects(id),
                FOREIGN KEY (assigned_user_id) REFERENCES users(id)
            )
        """)
        
        # Create bid_timeline table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS bid_timeline (
                id INT AUTO_INCREMENT PRIMARY KEY,
                bid_id INT,
                event VARCHAR(200) NOT NULL,
                details TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Create employees table for team-specific employee management
        cur.execute("""
            CREATE TABLE IF NOT EXISTS employees (
                id INT AUTO_INCREMENT PRIMARY KEY,
                name VARCHAR(100) NOT NULL,
                email VARCHAR(100) UNIQUE NOT NULL,
                password VARCHAR(255) NOT NULL,
                department VARCHAR(50) NOT NULL,
                team_lead_id INT,
                is_active BOOLEAN DEFAULT TRUE,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (team_lead_id) REFERENCES users(id)
            )
        """)
        
        # Create bid_checklists table for task management per bid (with InnoDB orphan tablespace repair)
        create_sql = (
            """
            CREATE TABLE IF NOT EXISTS bid_checklists (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                task_name VARCHAR(200) NOT NULL,
                description TEXT,
                assigned_to INT,
                status VARCHAR(50) DEFAULT 'pending',
                progress_pct INT DEFAULT NULL,
                stage VARCHAR(50),
                priority VARCHAR(20) DEFAULT 'medium',
                due_date DATETIME,
                attachment_path VARCHAR(255),
                created_by INT,
                team_archive VARCHAR(50),
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                FOREIGN KEY (g_id) REFERENCES go_bids(g_id),
                FOREIGN KEY (assigned_to) REFERENCES employees(id),
                FOREIGN KEY (created_by) REFERENCES users(id)
            ) ENGINE=InnoDB
            """
        )
        try:
            cur.execute(create_sql)
        except Exception as e:
            try:
                err_no = e.args[0] if hasattr(e, 'args') and e.args else None
            except Exception:
                err_no = None
            # 1813: orphan tablespace exists on disk; attempt OS-level cleanup then recreate
            if err_no == 1813:
                try:
                    cur2 = mysql.connection.cursor()
                    cur2.execute("SHOW VARIABLES LIKE 'datadir'")
                    row = cur2.fetchone()
                    data_dir = None
                    if isinstance(row, (list, tuple)) and len(row) >= 2:
                        data_dir = row[1]
                    elif isinstance(row, dict):
                        data_dir = row.get('Value') or row.get('value')
                    cur2.close()
                    if data_dir:
                        db_name = app.config.get('MYSQL_DB', 'esco')
                        ibd_path = os.path.join(data_dir, db_name, 'bid_checklists.ibd')
                        cfg_path = os.path.join(data_dir, db_name, 'bid_checklists.cfg')
                        for p in (ibd_path, cfg_path):
                            try:
                                if os.path.exists(p):
                                    os.remove(p)
                            except Exception:
                                pass
                except Exception:
                    pass
                try:
                    cur.execute("DROP TABLE IF EXISTS bid_checklists")
                    mysql.connection.commit()
                except Exception:
                    pass
                # Retry create
                cur.execute(create_sql)
            else:
                raise

        # Ensure progress_pct column exists even on older databases
        cur.execute("SELECT COUNT(*) AS cnt FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME='bid_checklists' AND COLUMN_NAME='progress_pct'")
        row = cur.fetchone()
        if not row or int(row.get('cnt', 0)) == 0:
            try:
                cur.execute("ALTER TABLE bid_checklists ADD COLUMN progress_pct INT NULL AFTER status")
            except Exception:
                pass

        # Ensure stage column exists
        cur.execute("SELECT COUNT(*) AS cnt FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME='bid_checklists' AND COLUMN_NAME='stage'")
        row = cur.fetchone()
        if not row or int(row.get('cnt', 0)) == 0:
            try:
                cur.execute("ALTER TABLE bid_checklists ADD COLUMN stage VARCHAR(50) NULL AFTER progress_pct")
            except Exception:
                pass

        # Ensure attachment_path column exists
        cur.execute("SELECT COUNT(*) AS cnt FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME='bid_checklists' AND COLUMN_NAME='attachment_path'")
        row = cur.fetchone()
        if not row or int(row.get('cnt', 0)) == 0:
            try:
                cur.execute("ALTER TABLE bid_checklists ADD COLUMN attachment_path VARCHAR(255) NULL AFTER due_date")
            except Exception:
                pass
        # Create dynamic stage tables used across dashboards
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_stage_exclusions (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_bid_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_custom_stages (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_custom_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )

        # Create dynamic stage tables used across dashboards
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_stage_exclusions (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_bid_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS bid_custom_stages (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                stage VARCHAR(50) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_custom_stage (g_id, stage)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
            """
        )

        # Create project_transfers table for tracking project handoffs between teams
        cur.execute("""
            CREATE TABLE IF NOT EXISTS project_transfers (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                from_team VARCHAR(50) NOT NULL,
                to_team VARCHAR(50) NOT NULL,
                transferred_by INT,
                transfer_reason TEXT,
                status VARCHAR(50) DEFAULT 'pending',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (g_id) REFERENCES go_bids(g_id),
                FOREIGN KEY (transferred_by) REFERENCES users(id)
            )
        """)

        # Comments table for board modal
        cur.execute("""
            CREATE TABLE IF NOT EXISTS bid_comments (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                user_id INT,
                comment_text TEXT NOT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (g_id) REFERENCES go_bids(g_id),
                FOREIGN KEY (user_id) REFERENCES users(id)
            )
        """)

        # Comments table for board modal
        cur.execute("""
            CREATE TABLE IF NOT EXISTS bid_comments (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                user_id INT,
                comment_text TEXT NOT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (g_id) REFERENCES go_bids(g_id),
                FOREIGN KEY (user_id) REFERENCES users(id)
            )
        """)

        # --- Profiling tables ---
        cur.execute("""
            CREATE TABLE IF NOT EXISTS company_preferences (
                id INT AUTO_INCREMENT PRIMARY KEY,
                registered_states TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS company_capabilities (
                id INT AUTO_INCREMENT PRIMARY KEY,
                description TEXT,
                file_path VARCHAR(255),
                uploaded_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS company_performance (
                id INT AUTO_INCREMENT PRIMARY KEY,
                project_name VARCHAR(200) NOT NULL,
                year INT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        # Mapping table for multi-employee bid assignments
        cur.execute("""
            CREATE TABLE IF NOT EXISTS bid_assignment_members (
                id INT AUTO_INCREMENT PRIMARY KEY,
                g_id INT NOT NULL,
                employee_id INT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE KEY uniq_g_emp (g_id, employee_id)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
        """)

        # Company details (profile) table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS company_details (
                id INT AUTO_INCREMENT PRIMARY KEY,
                name VARCHAR(200),
                website VARCHAR(255),
                email VARCHAR(255),
                phone VARCHAR(50),
                about TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)

        mysql.connection.commit()
        cur.close()
        print("Database tables created/verified successfully")
    except Exception as e:
        print(f"Error creating tables: {e}")
        try:
            mysql.connection.rollback()
        except Exception:
            pass
        if 'cur' in locals():
            try:
                cur.close()
            except Exception:
                pass
if __name__ == '__main__':
    from datetime import datetime, timedelta
    
    with app.app_context():
        _ensure_tables_exist()
        
        cur = mysql.connection.cursor(DictCursor)
        
        # Always ensure companies exist
        cur.execute("SELECT COUNT(*) as count FROM companies")
        if cur.fetchone()['count'] == 0:
            cur.execute("""
                INSERT INTO companies (name, description) VALUES 
                ('Ikio', 'Renewable Energy Solutions'),
                ('Metco', 'Industrial Energy Management'),
                ('Sunsprint', 'Solar Power Systems')
            """)
            mysql.connection.commit()
            print("Companies created successfully")
        
        # Check if users exist
        cur.execute("SELECT COUNT(*) as count FROM users")
        if cur.fetchone()['count'] == 0:
            # Create admin user
            cur.execute("""
                INSERT INTO users (email, password, is_admin, role) VALUES 
                ('admin@example.com', 'admin', 1, 'admin')
            """)
            admin_user_id = cur.lastrowid
            
            # Create other users
            cur.execute("""
                INSERT INTO users (email, password, is_admin, role) VALUES 
                ('bd@example.com', 'user', 0, 'business dev'),
                ('designer@example.com', 'designer', 0, 'design'),
                ('ops@example.com', 'ops', 0, 'operations'),
                ('sitemgr@example.com', 'site', 0, 'site manager')
            """)
            # Fetch exact user IDs by email to avoid relying on lastrowid math
            cur.execute("SELECT id FROM users WHERE email=%s", ('bd@example.com',))
            user_bdm_id = cur.fetchone()['id']
            cur.execute("SELECT id FROM users WHERE email=%s", ('designer@example.com',))
            user_design_id = cur.fetchone()['id']
            cur.execute("SELECT id FROM users WHERE email=%s", ('ops@example.com',))
            user_ops_id = cur.fetchone()['id']
            cur.execute("SELECT id FROM users WHERE email=%s", ('sitemgr@example.com',))
            user_site_id = cur.fetchone()['id']
            
            # Get company ids for linking bids
            cur.execute("SELECT id FROM companies WHERE name='Ikio'")
            ikio_id = cur.fetchone()['id']
            cur.execute("SELECT id FROM companies WHERE name='Metco'")
            metco_id = cur.fetchone()['id']
            cur.execute("SELECT id FROM companies WHERE name='Sunsprint'")
            sunsprint_id = cur.fetchone()['id']

            # Create sample bids linked to companies
            cur.execute("""
            INSERT INTO bids (name, current_stage, user_id, company_id) VALUES 
            ('Project Alpha', 'business', %s, %s),
            ('Project Beta', 'design', %s, %s),
            ('Project Gamma', 'operations', %s, %s)
        """, (user_bdm_id, ikio_id, user_design_id, metco_id, user_ops_id, sunsprint_id))
            # Create sample bid_incoming data
            cur.execute("""
                INSERT INTO bid_incoming (b_name, due_date, state, scope, type, scoring, comp_name, decision, summary) VALUES 
                ('Solar Energy Project', %s, 'submitted', 'Installation of 500kW solar panels for commercial building', 'Renewable Energy', 85, 'Ikio', 'GO', 'High potential project with good ROI'),
                ('Wind Farm Development', %s, 'under_review', 'Development of 2MW wind farm in rural area', 'Wind Energy', 72, 'Metco', 'NO-GO', 'Land acquisition issues identified'),
                ('Energy Efficiency Audit', %s, 'pending', 'Comprehensive energy audit for manufacturing facility', 'Energy Management', 90, 'Sunsprint', 'WON', 'Excellent technical proposal with competitive pricing'),
                ('Battery Storage System', %s, 'submitted', 'Installation of 1MWh battery storage system', 'Energy Storage', 78, 'Ikio', 'LOST', 'Lost to competitor with lower bid'),
                ('Smart Grid Implementation', %s, 'completed', 'Implementation of smart grid technology for city', 'Smart Grid', 95, 'Metco', 'WON', 'Successfully completed project ahead of schedule')
            """, (
                datetime.now() + timedelta(days=30),
                datetime.now() + timedelta(days=45),
                datetime.now() + timedelta(days=15),
                datetime.now() + timedelta(days=60),
                datetime.now() + timedelta(days=90)
            ))
            
            # One sample bid ready for site manager handover
            cur.execute("UPDATE bids SET current_stage='site_manager' WHERE name='Project Beta'")
            
            mysql.connection.commit()
            
            # Company ids already loaded above
            
            # Create sample projects
            cur.execute("""
                INSERT INTO projects (name, company_id, start_date, due_date, revenue, status, progress) VALUES 
                ('Solar Farm Installation', %s, %s, %s, 50000, 'active', 45),
                ('Wind Energy Project', %s, %s, %s, 75000, 'active', 70),
                ('Energy Efficiency Audit', %s, %s, %s, 25000, 'active', 30),
                ('Industrial Solar Setup', %s, %s, %s, 100000, 'active', 15),
                ('Residential Solar Panel', %s, %s, %s, 30000, 'active', 80),
                ('Commercial Solar System', %s, %s, %s, 60000, 'active', 25)
            """, (
                ikio_id, datetime.now(), datetime.now() + timedelta(days=90),
                ikio_id, datetime.now() - timedelta(days=30), datetime.now() + timedelta(days=60),
                metco_id, datetime.now() - timedelta(days=15), datetime.now() + timedelta(days=45),
                metco_id, datetime.now(), datetime.now() + timedelta(days=120),
                sunsprint_id, datetime.now() - timedelta(days=10), datetime.now() + timedelta(days=30),
                sunsprint_id, datetime.now(), datetime.now() + timedelta(days=75)
            ))
            
            # Get project IDs for tasks
            cur.execute("SELECT id FROM projects ORDER BY id")
            project_ids = [row['id'] for row in cur.fetchall()]
            
            # Create sample tasks
            cur.execute("""
                INSERT INTO tasks (name, project_id, assigned_user_id, due_date, status, priority) VALUES 
                ('Site Survey', %s, %s, %s, 'in_progress', 'high'),
                ('Equipment Procurement', %s, %s, %s, 'pending', 'medium'),
                ('Installation Planning', %s, %s, %s, 'completed', 'high'),
                ('Energy Assessment', %s, %s, %s, 'in_progress', 'urgent'),
                ('Client Consultation', %s, %s, %s, 'pending', 'high'),
                ('System Testing', %s, %s, %s, 'in_progress', 'medium'),
                ('Documentation', %s, %s, %s, 'pending', 'low')
            """, (
                project_ids[0], user_bdm_id, datetime.now() + timedelta(days=5),
                project_ids[0], user_design_id, datetime.now() + timedelta(days=15),
                project_ids[1], admin_user_id, datetime.now() + timedelta(days=10),
                project_ids[2], user_ops_id, datetime.now() + timedelta(days=7),
                project_ids[3], user_site_id, datetime.now() + timedelta(days=3),
                project_ids[4], admin_user_id, datetime.now() + timedelta(days=2),
                project_ids[5], user_bdm_id, datetime.now() + timedelta(days=20)
            ))
            
            mysql.connection.commit()
            
        # Seed at least one log if none exist
        cur.execute("SELECT COUNT(*) as count FROM logs")
        if cur.fetchone()['count'] == 0:
            cur.execute("INSERT INTO logs (action) VALUES ('System initialized and sample data seeded.')")
            mysql.connection.commit()
        
        cur.close()
    
    socketio.run(app, debug=True, port=5001)